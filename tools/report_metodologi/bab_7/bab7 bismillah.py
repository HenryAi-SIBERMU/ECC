#!/usr/bin/env python3
"""
Generator Laporan Metodologi Bab 7: Kegagalan Tata Kelola - D3TLH Dalam Sistem Perizinan

Pilar 1 ditulis langsung dalam generator Python agar selaras dengan SOP dokumentasi Celios2.
Fokus awal: Sub-bab 7.1 Pembuktian Empiris: Status Ekologis vs Penerbitan Izin.
"""

import base64
import sys
from pathlib import Path

try:
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([
        sys.executable,
        "-m",
        "pip",
        "install",
        "pandas",
        "requests",
        "python-docx",
    ])
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, cfg):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        for key, val in cfg.items():
            el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def para_border_bottom(paragraph, color="2E7D32", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="2E7D32", sz="16"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="A5D6A7", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="1B5E20", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="2E7D32", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=G_MID)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)


def add_p(doc, parts, space_after=5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "EDF7EE")
    all_border_para(p)
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)

    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=G_DARK)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p)
    para_shd(p, "F1F8E9")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}

    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "2E7D32")
        set_cell_borders(cell, bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8.5, color=C_WHITE)

    for i, row_data in enumerate(rows):
        fill = "F5FBF5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            set_cell_borders(cell, bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def generate_all_bab7():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    data_dir = base_dir / "data" / "processed"
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_7"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 7 sub-bab 7.1...")
    df_izin = pd.read_csv(data_dir / "sulawesi_izin_baru_per_tahun.csv")
    df_gfw = pd.read_csv(data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv")
    df_panel_71 = pd.merge(df_gfw, df_izin, on=["Provinsi", "Tahun"], how="left").fillna({"Jumlah_Izin_Baru": 0, "Total_Luas_Konsesi_Baru_Ha": 0})

    indikator_71 = "Total_Deforestasi_Ha"
    label_indikator_71 = "Total Deforestasi Alam (Hektar)"
    tertekan_threshold_71 = df_panel_71[indikator_71].quantile(0.33)
    kritis_threshold_71 = df_panel_71[indikator_71].quantile(0.66)

    def classify_d3tlh_71(val):
        if val <= tertekan_threshold_71:
            return "Aman"
        elif val <= kritis_threshold_71:
            return "Tertekan"
        return "Kritis"

    df_panel_71["Status_D3TLH"] = df_panel_71[indikator_71].apply(classify_d3tlh_71)
    n_panel_71 = len(df_panel_71)

    agg_71 = df_panel_71.groupby("Status_D3TLH").agg(
        Total_IUP=("Jumlah_Izin_Baru", "sum"),
        Total_Luas_Ha=("Total_Luas_Konsesi_Baru_Ha", "sum"),
        Min_Def=(indikator_71, "min"),
        Max_Def=(indikator_71, "max"),
        N_Obs=(indikator_71, "count"),
    ).reset_index()
    order_map_71 = {"Aman": 1, "Tertekan": 2, "Kritis": 3}
    agg_71["Order"] = agg_71["Status_D3TLH"].map(order_map_71)
    agg_71 = agg_71.sort_values("Order")

    aturan_map_71 = {
        "Aman": ("Wajar diterbitkan izin", "Normal (Sesuai Aturan)"),
        "Tertekan": ("Izin mulai direm/dibatasi", "Anomali (Lampu Kuning)"),
        "Kritis": ("Moratorium / Evaluasi Ketat", "PERLU EVALUASI"),
    }
    status_rows_71 = []
    total_iup_kritis_71 = 0
    luas_kritis_71 = 0.0
    for _, row in agg_71.iterrows():
        if row["Status_D3TLH"] == "Kritis":
            total_iup_kritis_71 = int(row["Total_IUP"])
            luas_kritis_71 = float(row["Total_Luas_Ha"])
        aturan, kesimpulan = aturan_map_71[row["Status_D3TLH"]]
        status_rows_71.append([
            row["Status_D3TLH"],
            f"{row['Min_Def']:,.0f} - {row['Max_Def']:,.0f} Ha",
            f"{int(row['N_Obs'])}",
            aturan,
            f"{int(row['Total_IUP'])} Izin Baru Keluar",
            f"{row['Total_Luas_Ha']:,.0f}",
            kesimpulan,
        ])

    iup_per_status_71 = {row["Status_D3TLH"]: int(row["Total_IUP"]) for _, row in agg_71.iterrows()}
    n_per_status_71 = {row["Status_D3TLH"]: int(row["N_Obs"]) for _, row in agg_71.iterrows()}

    df_kritis_panel_71 = df_panel_71[(df_panel_71["Status_D3TLH"] == "Kritis") & (df_panel_71["Jumlah_Izin_Baru"] > 0)]
    kritis_pairs_71 = set(zip(df_kritis_panel_71["Provinsi"], df_kritis_panel_71["Tahun"]))

    df_raw_71 = pd.read_csv(data_dir / "sulawesi_izin_raw_details.csv")
    df_kritis_raw_71 = df_raw_71[df_raw_71.apply(lambda row: (row["Provinsi"], row["Tahun"]) in kritis_pairs_71, axis=1)].copy()
    n_izin_irisan_71 = len(df_kritis_raw_71)
    luas_irisan_71 = float(pd.to_numeric(df_kritis_raw_71["luas_ha"], errors="coerce").fillna(0).sum())

    deforestasi_map_71 = df_kritis_panel_71.set_index(["Provinsi", "Tahun"])[indikator_71].to_dict()
    df_top_irisan_71 = df_kritis_raw_71.copy()
    df_top_irisan_71["luas_ha"] = pd.to_numeric(df_top_irisan_71["luas_ha"], errors="coerce").fillna(0)
    df_top_irisan_71 = df_top_irisan_71.sort_values("luas_ha", ascending=False).head(10)
    irisan_rows_71 = []
    for _, row in df_top_irisan_71.iterrows():
        def_val_71 = deforestasi_map_71.get((row["Provinsi"], row["Tahun"]), 0)
        irisan_rows_71.append([
            row["nama_badan_usaha"],
            row["komoditas"],
            row["Provinsi"],
            str(int(row["Tahun"])),
            f"{def_val_71:,.0f}",
            f"{row['luas_ha']:,.2f}",
        ])

    mermaid_str_7_1 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Deforestasi GFW<br/><i>Provinsi, Tahun, Total Deforestasi (Ha)</i>"] --> C
        B["Data Izin Baru MODI ESDM<br/><i>Provinsi, Tahun, Jumlah IUP, Luas Konsesi</i>"] --> C
    end
    subgraph Compliance_Modeling["2. Rule-based Categorization"]
        C["Panel Join Provinsi-Tahun"] --> D["Binning Persentil Deforestasi<br/>Aman (P33); Tertekan (P33-P66); Kritis (P66+)"]
        D --> E["Agregasi IUP & luas konsesi<br/>per status daya dukung"]
    end
    subgraph Output_Analysis["3. Matriks Kepatuhan & Tabel Irisan"]
        E --> F["Matriks Seharusnya vs Kenyataan<br/>per status D3TLH"]
        E --> G["Tabel Irisan ESDM x GFW<br/>daftar IUP terbit di zona Kritis"]
    end
    F --> H["Pembacaan kepatuhan instrumen D3TLH"]
    G --> H"""
    mermaid_png_path_7_1 = str(tool_dir / "mermaid_flowchart_7_1.png")
    download_success_7_1 = download_mermaid_png(mermaid_str_7_1, mermaid_png_path_7_1)

    print("[1.5/4] Mengekstraksi dataset empiris Bab 7 sub-bab 7.2...")
    df_konflik_72 = pd.read_csv(data_dir / "sulawesi_konflik_hukum.csv")
    total_kasus_72 = len(df_konflik_72)

    df_tl_72 = df_konflik_72.dropna(subset=["Tahun"]).copy()
    df_tl_72["Tahun"] = df_tl_72["Tahun"].astype(int)
    df_tl_dekade_72 = df_tl_72[df_tl_72["Tahun"] >= 2014]
    n_kasus_dekade_72 = len(df_tl_dekade_72)
    tahun_min_72 = int(df_tl_72["Tahun"].min())
    tahun_max_72 = int(df_tl_72["Tahun"].max())

    sektor_counts_72 = df_konflik_72["Sektor"].value_counts().reset_index()
    sektor_counts_72.columns = ["Sektor", "Jumlah_Kasus"]
    sektor_rows_72 = []
    for _, row in sektor_counts_72.iterrows():
        sektor_rows_72.append([
            row["Sektor"],
            f"{int(row['Jumlah_Kasus'])}",
            f"{row['Jumlah_Kasus'] / total_kasus_72 * 100:.1f}%",
        ])
    top_sektor_72 = sektor_counts_72.iloc[0]["Sektor"]
    top_sektor_n_72 = int(sektor_counts_72.iloc[0]["Jumlah_Kasus"])
    top_sektor_pct_72 = top_sektor_n_72 / total_kasus_72 * 100
    sektor2_72 = sektor_counts_72.iloc[1]["Sektor"] if len(sektor_counts_72) > 1 else "-"
    sektor2_n_72 = int(sektor_counts_72.iloc[1]["Jumlah_Kasus"]) if len(sektor_counts_72) > 1 else 0

    prov_agg_72 = df_tl_72.groupby("Provinsi").agg(Jumlah=("ID_Konflik", "count"), Tahun_Awal=("Tahun", "min"), Tahun_Akhir=("Tahun", "max")).reset_index().sort_values("Jumlah", ascending=False)
    prov_rows_72 = []
    for _, row in prov_agg_72.iterrows():
        prov_rows_72.append([
            row["Provinsi"],
            f"{int(row['Jumlah'])}",
            f"{row['Jumlah'] / total_kasus_72 * 100:.1f}%",
            f"{int(row['Tahun_Awal'])}-{int(row['Tahun_Akhir'])}",
        ])
    top_prov_72 = prov_agg_72.iloc[0]["Provinsi"]
    top_prov_n_72 = int(prov_agg_72.iloc[0]["Jumlah"])

    df_sampel_72 = df_tl_72.sort_values("Tahun", ascending=False).head(10)
    sampel_rows_72 = []
    for _, row in df_sampel_72.iterrows():
        sampel_rows_72.append([
            str(int(row["Tahun"])),
            row["Provinsi"],
            row["Sektor"],
            row["Judul_Kasus"],
            row["Sumber"],
        ])

    mermaid_str_7_2 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Database Konflik Hukum LSM/KPA/TanahKita<br/><i>Provinsi, Sektor, Judul Kasus, Tahun, Sumber</i>"]
    end
    subgraph Thematic_Coding["2. Thematic Coding & Agregasi Insiden"]
        A --> B["Klasifikasi rekam jejak<br/>sengketa lahan, pelanggaran HAM, operasi ilegal"]
        B --> C["Agregasi kasus per Sektor"]
        B --> D["Agregasi kasus per Provinsi & Tahun"]
    end
    subgraph Visual_Output["3. Pemetaan Pembiaran (State Omission)"]
        C --> E["Tabel sebaran sektor konflik"]
        D --> F["Bubble Timeline 10 tahun terakhir<br/>Tahun × Provinsi × jumlah kasus"]
    end
    E --> G["Pembacaan impunitas dan pembiaran struktural"]
    F --> G"""
    mermaid_png_path_7_2 = str(tool_dir / "mermaid_flowchart_7_2.png")
    download_success_7_2 = download_mermaid_png(mermaid_str_7_2, mermaid_png_path_7_2)

    print("[1.8/4] Mengekstraksi dataset empiris Bab 7 sub-bab 7.3...")
    df_pltu_73 = pd.read_csv(data_dir / "sulawesi_pltu_captive.csv")
    total_unit_73 = len(df_pltu_73)
    total_mw_73 = float(df_pltu_73["Capacity (MW)"].sum())
    df_pltu_73["Provinsi"] = df_pltu_73["Subnational unit (province, state)"]
    df_pltu_73["Tahun"] = df_pltu_73["Start year"].fillna(2025).astype(int)

    prov_agg_73 = df_pltu_73.groupby("Provinsi").agg(Unit=("Capacity (MW)", "count"), MW=("Capacity (MW)", "sum")).reset_index().sort_values("MW", ascending=False)
    prov_rows_73 = []
    for _, row in prov_agg_73.iterrows():
        prov_rows_73.append([
            row["Provinsi"],
            f"{int(row['Unit'])}",
            f"{row['MW']:,.0f}",
            f"{row['MW'] / total_mw_73 * 100:.1f}%",
        ])
    top_prov_73 = prov_agg_73.iloc[0]["Provinsi"]
    top_prov_mw_73 = float(prov_agg_73.iloc[0]["MW"])
    top_prov_pct_73 = top_prov_mw_73 / total_mw_73 * 100

    status_counts_73 = df_pltu_73.groupby("Status").agg(Unit=("Capacity (MW)", "count"), MW=("Capacity (MW)", "sum")).reset_index().sort_values("MW", ascending=False)
    status_map_73 = {row["Status"]: (int(row["Unit"]), float(row["MW"])) for _, row in status_counts_73.iterrows()}
    op_unit_73, op_mw_73 = status_map_73.get("operating", (0, 0.0))

    df_tl_73 = df_pltu_73[df_pltu_73["Tahun"] <= 2024].groupby(["Provinsi", "Tahun"])["Capacity (MW)"].sum().reset_index()
    df_pivot_73 = df_tl_73.pivot(index="Tahun", columns="Provinsi", values="Capacity (MW)").fillna(0)
    all_years_73 = list(range(int(df_pivot_73.index.min()), 2025))
    df_pivot_73 = df_pivot_73.reindex(all_years_73, fill_value=0)
    df_cum_73 = df_pivot_73.cumsum()
    prov_cols_73 = list(df_cum_73.columns)
    timeline_rows_73 = []
    for tahun, row in df_cum_73.iterrows():
        timeline_rows_73.append([str(int(tahun))] + [f"{row[c]:,.0f}" for c in prov_cols_73] + [f"{row.sum():,.0f}"])

    df_top_unit_73 = df_pltu_73.sort_values("Capacity (MW)", ascending=False).head(10)
    unit_rows_73 = []
    for _, row in df_top_unit_73.iterrows():
        tahun_op_73 = f"{row['Start year']:.0f}" if pd.notnull(row["Start year"]) else "Belum Operasi"
        unit_rows_73.append([
            row["Plant name"],
            row["Unit name"],
            row["Provinsi"],
            f"{row['Capacity (MW)']:,.0f}",
            row["Status"],
            tahun_op_73,
        ])

    mermaid_str_7_3 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Global Coal Plant Tracker (GEM)<br/><i>Plant, Unit, Capacity (MW), Status, Start year, Provinsi</i>"]
    end
    subgraph Inventory_Aggregation["2. Isolasi Regional & Agregasi Kuantitatif"]
        A --> B["Filter unit PLTU captive<br/>kawasan industri Sulawesi"]
        B --> C["Agregasi unit & kapasitas MW<br/>per Provinsi dan Status"]
        B --> D["Timeline kumulatif kapasitas<br/>berdasarkan Start Year (s.d. 2024)"]
    end
    subgraph Visual_Output["3. Kuantifikasi Kontradiksi Karbon"]
        C --> E["Metrik total unit & MW pembangkitan kotor"]
        D --> F["Line Chart pertumbuhan kumulatif per provinsi"]
    end
    E --> G["Pembacaan paradoks hilirisasi hijau vs PLTU batubara"]
    F --> G"""
    mermaid_png_path_7_3 = str(tool_dir / "mermaid_flowchart_7_3.png")
    download_success_7_3 = download_mermaid_png(mermaid_str_7_3, mermaid_png_path_7_3)

    print("[2/4] Membangun DOCX Metodologi_Bab7_Kegagalan_Tata_Kelola.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)
    add_h1(doc, "BAB VII: METODOLOGI ANALISIS KEGAGALAN TATA KELOLA - D3TLH DALAM SISTEM PERIZINAN")
    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada ", False, False),
        ("Bab 7: Kegagalan Tata Kelola - D3TLH Dalam Sistem Perizinan", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.", False, False),
    ])

    add_h2(doc, "7.1 Pembuktian Empiris: Status Ekologis vs Penerbitan Izin")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Izin: data/processed/sulawesi_izin_baru_per_tahun.csv dan data/processed/sulawesi_izin_raw_details.csv (MODI ESDM); Data Deforestasi: data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv (GFW). Visualisasi dashboard menampilkan Matriks Kepatuhan D3TLH (Seharusnya vs Kenyataan per status Aman/Tertekan/Kritis) beserta Tabel Irisan daftar perusahaan penerima izin di zona kritis.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Daya Tampung dan Daya Dukung Lingkungan Hidup (D3TLH) dirancang sebagai instrumen pencegahan dan pengatur batas pengaman ekologis (ecological safeguard). Secara metodologis, penerbitan izin baru sepatutnya mempertimbangkan indikator daya dukung lingkungan guna mengantisipasi degradasi ekosistem. ", False, False),
        ("Penyandingan data deforestasi tahunan dari Global Forest Watch (GFW) dan data perizinan pertambangan dari Minerba One Data Indonesia (MODI) Kementerian ESDM menunjukkan bahwa penerbitan izin usaha pertambangan baru tetap tercatat pada kurun waktu ketika perubahan tutupan hutan meningkat, terlihat pada tren di wilayah Sulawesi Tengah dan Tenggara periode 2014-2023.", False, False),
    ])
    add_p(doc, [
        ("Kondisi ini menggarisbawahi pentingnya penguatan fungsi dokumen AMDAL, D3TLH, dan KLHS agar menjadi pertimbangan utama yang mengikat dalam pengambilan keputusan perizinan, demi menjaga keberlanjutan lingkungan dan kehidupan masyarakat sekitar.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Evaluasi Kepatuhan D3TLH (Rule-based Categorization)")
    add_p(doc, [
        ("Kerangka agregasi berbasis aturan untuk membedah ketidaksesuaian antara status kerusakan lingkungan dengan keputusan administratif perizinan diilustrasikan pada ", False, False),
        ("Bagan Alur 7.1", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Compliance Modeling deskriptif dengan binning persentil tiga kelas.", False, False),
    ])
    add_caption(doc, "Bagan Alur 7.1: Alur Logika Analisis Evaluasi Kepatuhan D3TLH (Rule-based Categorization)")
    if download_success_7_1:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_7_1, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Ambang Persentil, Klasifikasi Status, dan Kuantifikasi Pelanggaran")
    add_p(doc, [("Kuantifikasi status daya dukung dan pelanggaran ekologis dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Ambang Batas Persentil Distribusi Deforestasi", "Ambang_Tertekan = Persentil_33 ( D_p,t )   ;   Ambang_Kritis = Persentil_66 ( D_p,t )", [
        ("D_p,t", f"Nilai {label_indikator_71} pada provinsi p dan tahun t (basis status wilayah)."),
        ("Ambang_Tertekan", f"Persentil ke-33 distribusi panel: {tertekan_threshold_71:,.1f} Ha."),
        ("Ambang_Kritis", f"Persentil ke-66 distribusi panel: {kritis_threshold_71:,.1f} Ha."),
    ])
    add_formula(doc, "Persamaan Klasifikasi Status Daya Dukung (Fungsi Piecewise 3 Kelas)", "Status(x) = 'Aman' , jika x ≤ P33   |   'Tertekan' , jika P33 < x ≤ P66   |   'Kritis' , jika x > P66", [
        ("Status(x)", f"Klasifikasi tiap observasi provinsi-tahun (N={n_panel_71}) ke dalam tiga kelas status daya dukung lingkungan."),
    ])
    add_formula(doc, "Persamaan Kuantifikasi Izin di Zona Kritis", "Total_Izin_Zona_Kritis = Σ ( IUP_p,t )   ;   untuk seluruh observasi dengan Status = 'Kritis'", [
        ("Total_Izin_Zona_Kritis", f"Akumulasi Izin Usaha Pertambangan baru yang tetap terbit pada periode berstatus Kritis ({total_iup_kritis_71} izin)."),
        ("IUP_p,t", "Jumlah izin tambang baru pada provinsi p dan tahun t (variabel keputusan aktor)."),
    ])

    add_p(doc, [("Substitusi angka dari dataset aktual ke dalam rumus ambang persentil dan kuantifikasi pelanggaran adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Ambang Batas Persentil", f"Ambang_Tertekan = Persentil_33 ( D ) = {tertekan_threshold_71:,.1f} Ha   ;   Ambang_Kritis = Persentil_66 ( D ) = {kritis_threshold_71:,.1f} Ha")
    add_formula(doc, "Substitusi Klasifikasi Status Panel", f"N_Aman = {n_per_status_71.get('Aman', 0)} observasi   ;   N_Tertekan = {n_per_status_71.get('Tertekan', 0)} observasi   ;   N_Kritis = {n_per_status_71.get('Kritis', 0)} observasi   (Total N = {n_panel_71})")
    add_formula(doc, "Substitusi Kuantifikasi Izin per Status", f"Σ IUP_Aman = {iup_per_status_71.get('Aman', 0)} izin   ;   Σ IUP_Tertekan = {iup_per_status_71.get('Tertekan', 0)} izin   ;   Σ IUP_Kritis = {iup_per_status_71.get('Kritis', 0)} izin ({luas_kritis_71:,.0f} Ha)")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Kepatuhan D3TLH dan Tabel Irisan Zona Kritis")
    add_p(doc, [
        ("Matriks perbandingan aturan normatif dengan kenyataan penerbitan izin pada masing-masing status daya dukung disajikan pada ", False, False),
        ("Tabel 7.1", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 7.1: Matriks Kepatuhan D3TLH - Seharusnya vs Kenyataan per Status Daya Dukung")
    add_table_1col(doc, ["Status Daya Dukung", "Rentang Kerusakan Hutan", "N Observasi", "Seharusnya (Menurut Aturan)", "Kenyataan di Lapangan", "Total Luas Konsesi (Ha)", "Kesimpulan Tata Kelola"], status_rows_71, [2.2, 3.0, 1.6, 3.0, 2.6, 2.4, 2.6], ["C", "C", "C", "L", "L", "C", "L"])

    add_p(doc, [
        (f"Tabel Irisan (Intersection) menyatukan Data Satelit (GFW) dan Data Perizinan (ESDM Minerba One): sistem melacak nama-nama perusahaan yang SK IUP-nya ditandatangani persis pada Tahun dan Provinsi yang sedang berstatus Kritis akibat deforestasi. Sepuluh entri dengan luas konsesi terbesar (dari total {n_izin_irisan_71:,} izin di zona kritis) disajikan pada ", False, False),
        ("Tabel 7.2", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 7.2: Sampel Irisan Izin IUP Terbit di Zona Kritis - 10 Konsesi Terluas (dari {n_izin_irisan_71:,} izin)")
    add_table_1col(doc, ["Nama Perusahaan (IUP)", "Komoditas", "Provinsi", "Tahun Terbit", "Kehilangan Hutan Provinsi (Ha)", "Luas Konsesi (Ha)"], irisan_rows_71, [3.6, 2.6, 2.2, 1.8, 2.8, 2.4], ["L", "L", "L", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Konklusi Kepatuhan D3TLH Berdasarkan Data Historis")
    add_p(doc, [
        ("1. ", True, False), ("Fungsi Pembatas D3TLH Perlu Ditingkatkan: ", True, False),
        (f"Terdapat {total_iup_kritis_71} Izin Baru yang terbit pada periode berstatus deforestasi tinggi (Kritis), mencakup luasan konsesi {luas_kritis_71:,.0f} Ha. Apabila pada status Kritis izin masih diterbitkan, hal tersebut secara matematis mendiskualifikasi D3TLH sebagai instrumen perlindungan lingkungan.\n", False, False),
        ("2. ", True, False), ("Bukti Irisan 100% Data-Driven: ", True, False),
        (f"Tabel irisan ESDM x GFW mengidentifikasi {n_izin_irisan_71:,} izin IUP baru (total luas {luas_irisan_71:,.0f} Ha) yang tetap diterbitkan di tengah situasi kritis pada pasangan Provinsi-Tahun yang sama.\n", False, False),
        ("3. ", True, False), ("Implikasi Kebijakan: ", True, False),
        ("Diperlukan penguatan integrasi data lingkungan dalam keputusan perizinan agar dokumen AMDAL, D3TLH, dan KLHS menjadi pertimbangan utama yang mengikat.", False, False),
    ])

    add_h2(doc, "7.2 Tabrakan Hukum: Impunitas dan Pembiaran Operasi Ilegal")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Konflik Hukum: data/processed/sulawesi_konflik_hukum.csv (kompilasi Konsorsium Pembaruan Agraria / KPA, TanahKita, dan laporan LSM). Visualisasi dashboard menampilkan metrik total kasus, Bubble Timeline sebaran konflik agraria 10 tahun terakhir (Tahun × Provinsi), tabel sebaran sektor konflik, serta daftar rekam jejak kasus.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Konsep Daya Tampung dan Daya Dukung Lingkungan Hidup (D3TLH) mengukur kapasitas daya tahan ekosistem serta daya dukung sosial masyarakat di sekitar kawasan industri. Kompilasi laporan masyarakat sipil dan organisasi terkait mencatat adanya sengketa tanah dan dinamika sosial dalam ekspansi industri ekstraktif. ", False, False),
        ("Hal ini menunjukkan pentingnya kepatuhan perizinan dan penerapan sanksi administratif secara konsisten. Pengawasan terhadap batas wilayah perizinan (HGU/IUP) serta pelaksanaan konsultasi publik (FPIC) menjadi aspek penting dalam tata kelola pertanahan dan lingkungan.", False, False),
    ])
    add_p(doc, [
        ("Penguatan koordinasi antar-instansi serta penyelesaian sengketa tenurial secara adil menjadi langkah krusial untuk memastikan kepastian hukum dan perlindungan hak masyarakat di wilayah sekitar industri.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Pemetaan Impunitas Korporasi (Thematic Coding)")
    add_p(doc, [
        ("Kerangka agregasi pelaporan berbasis insiden (Incident-based Reporting Aggregation) untuk mengukur tingkat pembiaran penegakan hukum diilustrasikan pada ", False, False),
        ("Bagan Alur 7.2", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Thematic Coding dan analisis kasus deskriptif.", False, False),
    ])
    add_caption(doc, "Bagan Alur 7.2: Alur Logika Analisis Pemetaan Impunitas Korporasi (Thematic Coding)")
    if download_success_7_2:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_7_2, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Agregasi Insiden dan Volume Pembiaran Sektoral")
    add_p(doc, [("Kuantifikasi tingkat pembiaran penegakan hukum dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Total Kasus Impunitas", "Total_Kasus_Impunitas = Σ ( Kasus_i )   ;   untuk seluruh laporan insiden i dalam database", [
        ("Total_Kasus_Impunitas", "Jumlah agregat kasus konflik, pelanggaran, dan operasi ilegal yang terdokumentasi namun operasinya tetap dipertahankan."),
        ("Kasus_i", "Satu entri laporan insiden pada database konflik hukum (atribut: Provinsi, Sektor, Judul_Kasus, Tahun)."),
    ])
    add_formula(doc, "Persamaan Volume Pembiaran Sektoral", "Volume_Sektoral_s = Σ ( Kasus_i )   ;   untuk seluruh kasus i dengan Sektor = s", [
        ("Volume_Sektoral_s", "Jumlah kasus pembiaran yang diatribusikan pada sektor industri s (Pertambangan, Perkebunan, dst.)."),
    ])
    add_formula(doc, "Persamaan Proporsi Pembiaran Sektoral", "Proporsi_Sektoral_s (%) = ( Volume_Sektoral_s / Total_Kasus_Impunitas ) × 100", [
        ("Proporsi_Sektoral_s (%)", "Kontribusi sektor s terhadap total kasus impunitas terdokumentasi."),
    ])

    add_p(doc, [("Substitusi angka dari dataset aktual ke dalam rumus agregasi insiden adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Total Kasus Impunitas", f"Total_Kasus_Impunitas = {total_kasus_72} kasus   ;   {n_kasus_dekade_72} kasus di antaranya pada dekade 2014-{tahun_max_72}")
    add_formula(doc, "Substitusi Volume Sektoral Tertinggi", f"Volume_Sektoral_{top_sektor_72} = {top_sektor_n_72} kasus ({top_sektor_pct_72:.1f}%)   ;   Volume_Sektoral_{sektor2_72} = {sektor2_n_72} kasus")
    add_formula(doc, "Substitusi Sebaran Provinsi Tertinggi", f"Volume_Provinsi_{top_prov_72.replace(' ', '_')} = {top_prov_n_72} kasus   (dari total {total_kasus_72} kasus se-Sulawesi)")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Sebaran Sektor, Provinsi, dan Rekam Jejak Kasus")
    add_p(doc, [
        ("Sebaran kasus konflik dan pembiaran berdasarkan sektor penyebab disajikan pada ", False, False),
        ("Tabel 7.3", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 7.3: Sebaran Sektor Konflik dan Pembiaran Operasi Ilegal")
    add_table_1col(doc, ["Sektor (Penyebab)", "Jumlah Kasus", "Proporsi (%)"], sektor_rows_72, [6.0, 3.0, 3.0], ["L", "C", "C"])

    add_p(doc, [
        (f"Sebaran spasial kasus per provinsi beserta rentang tahun kejadiannya ({tahun_min_72}-{tahun_max_72}) disajikan pada ", False, False),
        ("Tabel 7.4", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 7.4: Sebaran Wilayah Konflik per Provinsi dan Rentang Tahun Kejadian")
    add_table_1col(doc, ["Provinsi", "Jumlah Kasus", "Proporsi (%)", "Rentang Tahun"], prov_rows_72, [4.0, 2.6, 2.6, 3.0], ["L", "C", "C", "C"])

    add_p(doc, [
        ("Sepuluh rekam jejak kasus terbaru dari daftar konflik agraria dan pelanggaran hak disajikan pada ", False, False),
        ("Tabel 7.5", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 7.5: Sampel Rekam Jejak Konflik Agraria & Pelanggaran Hak (10 Kasus Terbaru)")
    add_table_1col(doc, ["Tahun", "Provinsi", "Sektor", "Judul/Nama Kasus", "Sumber"], sampel_rows_72, [1.6, 2.6, 2.6, 5.6, 3.0], ["C", "L", "L", "L", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris: Impunitas dan Pembiaran Struktural")
    add_p(doc, [
        ("1. ", True, False), ("Bukti Impunitas Hukum: ", True, False),
        (f"Database kompilasi LSM/KPA mendokumentasikan {total_kasus_72} kasus konflik/pelanggaran yang dibiarkan di Sulawesi ({n_kasus_dekade_72} kasus pada dekade 2014-{tahun_max_72}), di mana korporasi yang terbukti bermasalah secara hukum tetap dipertahankan keberadaan operasinya.\n", False, False),
        ("2. ", True, False), ("Dominasi Sektor Ekstraktif: ", True, False),
        (f"Sektor {top_sektor_72} menjadi penyebab konflik terbanyak dengan {top_sektor_n_72} kasus ({top_sektor_pct_72:.1f}%), disusul {sektor2_72} ({sektor2_n_72} kasus) — konsisten dengan pola tekanan ekspansi industri ekstraktif pada bab-bab sebelumnya.\n", False, False),
        ("3. ", True, False), ("Konsentrasi Spasial: ", True, False),
        (f"Provinsi {top_prov_72} mencatat kasus terbanyak ({top_prov_n_72} kasus). Temuan ini menegaskan pentingnya penguatan koordinasi antar-instansi, penegakan sanksi administratif yang konsisten, serta penyelesaian sengketa tenurial secara adil.", False, False),
    ])

    add_h2(doc, "7.3 Inkonsistensi Iklim: Karpet Merah PLTU Captive")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data PLTU Captive: data/processed/sulawesi_pltu_captive.csv (Global Coal Plant Tracker / GEM, ekstraksi Januari 2026). Visualisasi dashboard menampilkan metrik total unit dan kapasitas pembangkitan kotor, Line Chart timeline pertumbuhan kapasitas kumulatif per provinsi hingga 2024, serta daftar lengkap unit PLTU batubara captive.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Komitmen transisi energi global dan pengembangan rantai pasok industri nikel memegang peranan strategis. Di saat yang sama, pemenuhan kebutuhan energi untuk fasilitas pengolahan nikel (smelter) di Sulawesi masih didominasi oleh Pembangkit Listrik Tenaga Uap (PLTU) Batubara Captive. ", False, False),
        (f"Data dari Global Coal Plant Tracker (GEM) mencatat keberadaan {total_unit_73} unit PLTU Captive yang beroperasi maupun direncanakan di kawasan industri Sulawesi dengan total kapasitas pembangkitan {total_mw_73:,.0f} MW. Pemanfaatan energi berbasis batu bara pada industri ini menghasilkan tantangan tersendiri bagi pengelolaan emisi gas rumah kaca dan kualitas udara ambien.", False, False),
    ])
    add_p(doc, [
        ("Kondisi ini menunjukkan perlunya strategi percepatan transisi energi bersih di sektor industri ekstraktif guna menyelaraskan target hilirisasi dengan komitmen penurunan emisi nasional.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Agregasi Beban Karbon PLTU Captive")
    add_p(doc, [
        ("Kerangka inventarisasi agregat kuantitatif (Quantitative Inventory Aggregation) dari database global PLTU batubara captive diilustrasikan pada ", False, False),
        ("Bagan Alur 7.3", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan penyaringan agregat dataset eksternal deskriptif.", False, False),
    ])
    add_caption(doc, "Bagan Alur 7.3: Alur Logika Analisis Agregasi Beban Karbon PLTU Captive (GEM)")
    if download_success_7_3:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_7_3, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Inventarisasi Unit, Beban Karbon, dan Timeline Kumulatif")
    add_p(doc, [("Kuantifikasi kontradiksi karbon dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Total Infrastruktur Kotor", "Total_Infrastruktur_Kotor = Σ ( Unit_u )   ;   untuk seluruh unit PLTU captive u di kawasan industri Sulawesi", [
        ("Total_Infrastruktur_Kotor", "Jumlah agregat unit pembangkit batubara captive (beroperasi, dibangun, maupun direncanakan)."),
        ("Unit_u", "Satu unit pembangkit pada inventaris Global Coal Plant Tracker."),
    ])
    add_formula(doc, "Persamaan Total Beban Karbon per Provinsi", "Total_Beban_Karbon_p = Σ ( Kapasitas_u )   ;   untuk seluruh unit u pada provinsi p", [
        ("Total_Beban_Karbon_p", "Agregat luaran listrik kotor (MW) yang menopang pabrik pemurnian nikel pada provinsi p."),
        ("Kapasitas_u", "Kapasitas terpasang (MW) unit pembangkit u."),
    ])
    add_formula(doc, "Persamaan Timeline Kumulatif Kapasitas", "Kumulatif_Kapasitas_p(T) = Σ Kapasitas_p(t)   ;   untuk t = tahun operasi awal s.d. T (T ≤ 2024)", [
        ("Kumulatif_Kapasitas_p(T)", "Akumulasi kapasitas PLTU captive provinsi p hingga tahun berjalan T berdasarkan Start Year unit."),
    ])

    add_p(doc, [("Substitusi angka dari dataset aktual ke dalam rumus inventarisasi adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Total Infrastruktur & Beban Karbon", f"Total_Infrastruktur_Kotor = {total_unit_73} unit   ;   Total_Beban_Karbon = {total_mw_73:,.0f} MW")
    add_formula(doc, "Substitusi Beban Karbon Provinsi Tertinggi", f"Total_Beban_Karbon_{top_prov_73.replace(' ', '_')} = {top_prov_mw_73:,.0f} MW ({top_prov_pct_73:.1f}% dari total)")
    add_formula(doc, "Substitusi Komposisi Status Unit", f"Operating = {op_unit_73} unit ({op_mw_73:,.0f} MW)   ;   Non-Operating (konstruksi/rencana/shelved/cancelled) = {total_unit_73 - op_unit_73} unit")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Beban Karbon per Provinsi, Timeline, dan Unit Terbesar")
    add_p(doc, [
        ("Agregat unit dan kapasitas PLTU captive per provinsi disajikan pada ", False, False),
        ("Tabel 7.6", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 7.6: Agregat Unit dan Kapasitas PLTU Captive per Provinsi (Total {total_mw_73:,.0f} MW)")
    add_table_1col(doc, ["Provinsi", "Jumlah Unit", "Kapasitas (MW)", "Proporsi (%)"], prov_rows_73, [4.2, 2.6, 3.0, 2.6], ["L", "C", "C", "C"])

    add_p(doc, [
        ("Timeline pertumbuhan kapasitas kumulatif per provinsi hingga 2024 disajikan pada ", False, False),
        ("Tabel 7.7", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 7.7: Timeline Pertumbuhan Kapasitas Kumulatif PLTU Captive per Provinsi (MW, s.d. 2024)")
    add_table_1col(doc, ["Tahun"] + prov_cols_73 + ["Total (MW)"], timeline_rows_73, [1.8] + [3.0] * len(prov_cols_73) + [2.6], ["C"] * (len(prov_cols_73) + 2))

    add_p(doc, [
        ("Sepuluh unit pembangkit dengan kapasitas terbesar disajikan pada ", False, False),
        ("Tabel 7.8", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 7.8: Sepuluh Unit PLTU Captive Berkapasitas Terbesar di Sulawesi")
    add_table_1col(doc, ["Nama Pembangkit", "Unit", "Provinsi", "Kapasitas (MW)", "Status", "Tahun Beroperasi"], unit_rows_73, [4.0, 2.0, 2.8, 2.4, 2.0, 2.4], ["L", "L", "L", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Paradoks Hilirisasi Hijau")
    add_p(doc, [
        ("1. ", True, False), ("Skala Infrastruktur Kotor: ", True, False),
        (f"GEM mencatat {total_unit_73} unit PLTU batubara captive di kawasan industri Sulawesi dengan total kapasitas {total_mw_73:,.0f} MW ({op_unit_73} unit beroperasi, {op_mw_73:,.0f} MW) — pembangkitan masif yang dibangun demi menopang pabrik pemurnian nikel yang dipromosikan sebagai proyek energi ramah lingkungan.\n", False, False),
        ("2. ", True, False), ("Konsentrasi Beban Karbon: ", True, False),
        (f"Provinsi {top_prov_73} menanggung beban karbon terbesar dengan {top_prov_mw_73:,.0f} MW ({top_prov_pct_73:.1f}% dari total), sejalan dengan konsentrasi kawasan industri hilirisasi.\n", False, False),
        ("3. ", True, False), ("Implikasi Transisi Energi: ", True, False),
        ("Kondisi ini menunjukkan perlunya strategi percepatan transisi energi bersih di sektor industri ekstraktif guna menyelaraskan target hilirisasi dengan komitmen penurunan emisi nasional.", False, False),
    ])

    docx_path = tool_dir / "Metodologi_Bab7_Kegagalan_Tata_Kelola.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 7...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 7 - Kegagalan Tata Kelola</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: Arial, sans-serif; max-width: 920px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #43A047; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; }}
.hdr-title {{ color: #81C784; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #2E7D32; padding-bottom: 8px; }}
h2 {{ color: #81C784; text-transform: uppercase; border-bottom: 1px solid #2E7D32; }}
h4 {{ color: #A5D6A7; }}
.note-box {{ background: #132213; border-left: 4px solid #2E7D32; padding: 10px 14px; margin: 12px 0; }}
.formula {{ background: #0D1B0E; border: 1px solid #2E7D32; color: #A5D6A7; padding: 8px 12px; font-family: monospace; }}
.table-caption {{ color: #A5D6A7; font-weight: 700; font-style: italic; margin-top: 14px; }}
.data-th {{ background: #1B5E20; color: white; padding: 6px; text-align: left; border: 1px solid #2E7D32; }}
.data-td {{ padding: 6px; border: 1px solid #243524; vertical-align: top; }}
.data-tr-even .data-td {{ background: #131B13; }}
.mermaid {{ background: #0D1610; border: 1px solid #2E7D32; padding: 12px; margin: 10px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB VII: Metodologi Analisis Kegagalan Tata Kelola - D3TLH Dalam Sistem Perizinan</div>
<p>Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada <strong>Bab 7: Kegagalan Tata Kelola - D3TLH Dalam Sistem Perizinan</strong>.</p>

<h2>7.1 Pembuktian Empiris: Status Ekologis vs Penerbitan Izin</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Izin: <code>data/processed/sulawesi_izin_baru_per_tahun.csv</code> dan <code>data/processed/sulawesi_izin_raw_details.csv</code> (MODI ESDM); Data Deforestasi: <code>data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv</code> (GFW). Visualisasi dashboard menampilkan Matriks Kepatuhan D3TLH (Seharusnya vs Kenyataan per status Aman/Tertekan/Kritis) beserta Tabel Irisan daftar perusahaan penerima izin di zona kritis.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>D3TLH dirancang sebagai instrumen pencegahan dan pengatur batas pengaman ekologis (ecological safeguard). Penyandingan data deforestasi tahunan GFW dan data perizinan MODI ESDM menunjukkan bahwa penerbitan izin usaha pertambangan baru tetap tercatat pada kurun waktu ketika perubahan tutupan hutan meningkat — terlihat pada tren Sulawesi Tengah dan Tenggara 2014-2023. Kondisi ini menggarisbawahi pentingnya penguatan fungsi AMDAL, D3TLH, dan KLHS sebagai pertimbangan yang mengikat dalam keputusan perizinan.</p>
<h4>B. Alur Logika Metodologis Evaluasi Kepatuhan D3TLH (Rule-based Categorization)</h4>
<p>Kerangka agregasi berbasis aturan diilustrasikan pada <strong>Bagan Alur 7.1</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Compliance Modeling deskriptif dengan binning persentil tiga kelas.</p>
<div class="table-caption">Bagan Alur 7.1: Alur Logika Analisis Evaluasi Kepatuhan D3TLH (Rule-based Categorization)</div>
<div class="mermaid">{mermaid_str_7_1}</div>
<h4>C. Formulasi Matematis: Ambang Persentil, Klasifikasi Status, dan Kuantifikasi Pelanggaran</h4>
<div class="formula">Ambang_Tertekan = Persentil_33 ( D_p,t )   ;   Ambang_Kritis = Persentil_66 ( D_p,t )</div>
<div class="formula">Status(x) = 'Aman' , jika x ≤ P33   |   'Tertekan' , jika P33 &lt; x ≤ P66   |   'Kritis' , jika x &gt; P66</div>
<div class="formula">Total_Izin_Zona_Kritis = Σ ( IUP_p,t )   ;   untuk seluruh observasi dengan Status = 'Kritis'</div>
<p>Substitusi angka dari dataset aktual ke dalam rumus ambang persentil dan kuantifikasi pelanggaran:</p>
<div class="formula">Ambang_Tertekan = Persentil_33 ( D ) = {tertekan_threshold_71:,.1f} Ha   ;   Ambang_Kritis = Persentil_66 ( D ) = {kritis_threshold_71:,.1f} Ha</div>
<div class="formula">N_Aman = {n_per_status_71.get('Aman', 0)} observasi   ;   N_Tertekan = {n_per_status_71.get('Tertekan', 0)} observasi   ;   N_Kritis = {n_per_status_71.get('Kritis', 0)} observasi   (Total N = {n_panel_71})</div>
<div class="formula">Σ IUP_Aman = {iup_per_status_71.get('Aman', 0)} izin   ;   Σ IUP_Tertekan = {iup_per_status_71.get('Tertekan', 0)} izin   ;   Σ IUP_Kritis = {iup_per_status_71.get('Kritis', 0)} izin ({luas_kritis_71:,.0f} Ha)</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 7.1: Matriks Kepatuhan D3TLH - Seharusnya vs Kenyataan per Status Daya Dukung</div>
{html_table(["Status Daya Dukung", "Rentang Kerusakan Hutan", "N Observasi", "Seharusnya (Menurut Aturan)", "Kenyataan di Lapangan", "Total Luas Konsesi (Ha)", "Kesimpulan Tata Kelola"], status_rows_71)}
<div class="table-caption">Tabel 7.2: Sampel Irisan Izin IUP Terbit di Zona Kritis - 10 Konsesi Terluas (dari {n_izin_irisan_71:,} izin)</div>
{html_table(["Nama Perusahaan (IUP)", "Komoditas", "Provinsi", "Tahun Terbit", "Kehilangan Hutan Provinsi (Ha)", "Luas Konsesi (Ha)"], irisan_rows_71)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Fungsi Pembatas D3TLH Perlu Ditingkatkan:</strong> terdapat <strong>{total_iup_kritis_71} Izin Baru</strong> yang terbit pada periode berstatus deforestasi tinggi (Kritis), mencakup luasan konsesi {luas_kritis_71:,.0f} Ha. <strong>2. Bukti Irisan 100% Data-Driven:</strong> tabel irisan ESDM x GFW mengidentifikasi <strong>{n_izin_irisan_71:,} izin IUP baru</strong> (total luas {luas_irisan_71:,.0f} Ha) yang tetap diterbitkan di tengah situasi kritis. <strong>3. Implikasi Kebijakan:</strong> diperlukan penguatan integrasi data lingkungan agar AMDAL, D3TLH, dan KLHS menjadi pertimbangan yang mengikat dalam keputusan perizinan.</p>

<h2>7.2 Tabrakan Hukum: Impunitas dan Pembiaran Operasi Ilegal</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Konflik Hukum: <code>data/processed/sulawesi_konflik_hukum.csv</code> (kompilasi Konsorsium Pembaruan Agraria / KPA, TanahKita, dan laporan LSM). Visualisasi dashboard menampilkan metrik total kasus, Bubble Timeline sebaran konflik agraria 10 tahun terakhir (Tahun × Provinsi), tabel sebaran sektor konflik, serta daftar rekam jejak kasus.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Konsep D3TLH mengukur kapasitas daya tahan ekosistem serta daya dukung sosial masyarakat di sekitar kawasan industri. Kompilasi laporan masyarakat sipil mencatat adanya sengketa tanah dan dinamika sosial dalam ekspansi industri ekstraktif — menunjukkan pentingnya kepatuhan perizinan, penerapan sanksi administratif yang konsisten, pengawasan batas wilayah perizinan (HGU/IUP), pelaksanaan konsultasi publik (FPIC), serta penyelesaian sengketa tenurial secara adil.</p>
<h4>B. Alur Logika Metodologis Pemetaan Impunitas Korporasi (Thematic Coding)</h4>
<p>Kerangka agregasi pelaporan berbasis insiden diilustrasikan pada <strong>Bagan Alur 7.2</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Thematic Coding dan analisis kasus deskriptif.</p>
<div class="table-caption">Bagan Alur 7.2: Alur Logika Analisis Pemetaan Impunitas Korporasi (Thematic Coding)</div>
<div class="mermaid">{mermaid_str_7_2}</div>
<h4>C. Formulasi Matematis: Agregasi Insiden dan Volume Pembiaran Sektoral</h4>
<div class="formula">Total_Kasus_Impunitas = Σ ( Kasus_i )   ;   untuk seluruh laporan insiden i dalam database</div>
<div class="formula">Volume_Sektoral_s = Σ ( Kasus_i )   ;   untuk seluruh kasus i dengan Sektor = s</div>
<div class="formula">Proporsi_Sektoral_s (%) = ( Volume_Sektoral_s / Total_Kasus_Impunitas ) × 100</div>
<p>Substitusi angka dari dataset aktual ke dalam rumus agregasi insiden:</p>
<div class="formula">Total_Kasus_Impunitas = {total_kasus_72} kasus   ;   {n_kasus_dekade_72} kasus pada dekade 2014-{tahun_max_72}</div>
<div class="formula">Volume_Sektoral_{top_sektor_72} = {top_sektor_n_72} kasus ({top_sektor_pct_72:.1f}%)   ;   Volume_Sektoral_{sektor2_72} = {sektor2_n_72} kasus</div>
<div class="formula">Volume_Provinsi_Tertinggi: {top_prov_72} = {top_prov_n_72} kasus (dari total {total_kasus_72} kasus se-Sulawesi)</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 7.3: Sebaran Sektor Konflik dan Pembiaran Operasi Ilegal</div>
{html_table(["Sektor (Penyebab)", "Jumlah Kasus", "Proporsi (%)"], sektor_rows_72)}
<div class="table-caption">Tabel 7.4: Sebaran Wilayah Konflik per Provinsi dan Rentang Tahun Kejadian</div>
{html_table(["Provinsi", "Jumlah Kasus", "Proporsi (%)", "Rentang Tahun"], prov_rows_72)}
<div class="table-caption">Tabel 7.5: Sampel Rekam Jejak Konflik Agraria & Pelanggaran Hak (10 Kasus Terbaru)</div>
{html_table(["Tahun", "Provinsi", "Sektor", "Judul/Nama Kasus", "Sumber"], sampel_rows_72)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Bukti Impunitas Hukum:</strong> database kompilasi LSM/KPA mendokumentasikan <strong>{total_kasus_72} kasus</strong> konflik/pelanggaran yang dibiarkan di Sulawesi ({n_kasus_dekade_72} kasus pada dekade 2014-{tahun_max_72}). <strong>2. Dominasi Sektor Ekstraktif:</strong> sektor <strong>{top_sektor_72}</strong> penyebab terbanyak dengan {top_sektor_n_72} kasus ({top_sektor_pct_72:.1f}%), disusul {sektor2_72} ({sektor2_n_72} kasus). <strong>3. Konsentrasi Spasial:</strong> {top_prov_72} mencatat kasus terbanyak ({top_prov_n_72} kasus) — menegaskan pentingnya penguatan koordinasi antar-instansi, sanksi administratif konsisten, dan penyelesaian sengketa tenurial yang adil.</p>

<h2>7.3 Inkonsistensi Iklim: Karpet Merah PLTU Captive</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data PLTU Captive: <code>data/processed/sulawesi_pltu_captive.csv</code> (Global Coal Plant Tracker / GEM, ekstraksi Januari 2026). Visualisasi dashboard menampilkan metrik total unit dan kapasitas pembangkitan kotor, Line Chart timeline pertumbuhan kapasitas kumulatif per provinsi hingga 2024, serta daftar lengkap unit PLTU batubara captive.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Pemenuhan kebutuhan energi untuk fasilitas pengolahan nikel (smelter) di Sulawesi masih didominasi oleh PLTU Batubara Captive. Data GEM mencatat <strong>{total_unit_73} unit PLTU Captive</strong> yang beroperasi maupun direncanakan di kawasan industri Sulawesi dengan total kapasitas <strong>{total_mw_73:,.0f} MW</strong> — tantangan tersendiri bagi pengelolaan emisi gas rumah kaca dan kualitas udara ambien, yang menunjukkan perlunya percepatan transisi energi bersih agar target hilirisasi selaras dengan komitmen penurunan emisi nasional.</p>
<h4>B. Alur Logika Metodologis Agregasi Beban Karbon PLTU Captive</h4>
<p>Kerangka inventarisasi agregat kuantitatif diilustrasikan pada <strong>Bagan Alur 7.3</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan penyaringan agregat dataset eksternal deskriptif.</p>
<div class="table-caption">Bagan Alur 7.3: Alur Logika Analisis Agregasi Beban Karbon PLTU Captive (GEM)</div>
<div class="mermaid">{mermaid_str_7_3}</div>
<h4>C. Formulasi Matematis: Inventarisasi Unit, Beban Karbon, dan Timeline Kumulatif</h4>
<div class="formula">Total_Infrastruktur_Kotor = Σ ( Unit_u )   ;   untuk seluruh unit PLTU captive u di kawasan industri Sulawesi</div>
<div class="formula">Total_Beban_Karbon_p = Σ ( Kapasitas_u )   ;   untuk seluruh unit u pada provinsi p</div>
<div class="formula">Kumulatif_Kapasitas_p(T) = Σ Kapasitas_p(t)   ;   untuk t = tahun operasi awal s.d. T (T ≤ 2024)</div>
<p>Substitusi angka dari dataset aktual ke dalam rumus inventarisasi:</p>
<div class="formula">Total_Infrastruktur_Kotor = {total_unit_73} unit   ;   Total_Beban_Karbon = {total_mw_73:,.0f} MW</div>
<div class="formula">Total_Beban_Karbon_Tertinggi: {top_prov_73} = {top_prov_mw_73:,.0f} MW ({top_prov_pct_73:.1f}% dari total)</div>
<div class="formula">Operating = {op_unit_73} unit ({op_mw_73:,.0f} MW)   ;   Non-Operating = {total_unit_73 - op_unit_73} unit</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 7.6: Agregat Unit dan Kapasitas PLTU Captive per Provinsi (Total {total_mw_73:,.0f} MW)</div>
{html_table(["Provinsi", "Jumlah Unit", "Kapasitas (MW)", "Proporsi (%)"], prov_rows_73)}
<div class="table-caption">Tabel 7.7: Timeline Pertumbuhan Kapasitas Kumulatif PLTU Captive per Provinsi (MW, s.d. 2024)</div>
{html_table(["Tahun"] + prov_cols_73 + ["Total (MW)"], timeline_rows_73)}
<div class="table-caption">Tabel 7.8: Sepuluh Unit PLTU Captive Berkapasitas Terbesar di Sulawesi</div>
{html_table(["Nama Pembangkit", "Unit", "Provinsi", "Kapasitas (MW)", "Status", "Tahun Beroperasi"], unit_rows_73)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Skala Infrastruktur Kotor:</strong> GEM mencatat <strong>{total_unit_73} unit</strong> PLTU batubara captive dengan total kapasitas <strong>{total_mw_73:,.0f} MW</strong> ({op_unit_73} unit beroperasi, {op_mw_73:,.0f} MW). <strong>2. Konsentrasi Beban Karbon:</strong> {top_prov_73} menanggung beban terbesar {top_prov_mw_73:,.0f} MW ({top_prov_pct_73:.1f}%), sejalan dengan konsentrasi kawasan industri hilirisasi. <strong>3. Implikasi Transisi Energi:</strong> diperlukan percepatan transisi energi bersih di sektor industri ekstraktif agar target hilirisasi selaras dengan komitmen penurunan emisi nasional.</p>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab7_Kegagalan_Tata_Kelola.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB VII: METODOLOGI ANALISIS KEGAGALAN TATA KELOLA - D3TLH DALAM SISTEM PERIZINAN",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada **Bab 7: Kegagalan Tata Kelola - D3TLH Dalam Sistem Perizinan**.",
        "",
        "## 7.1 Pembuktian Empiris: Status Ekologis vs Penerbitan Izin",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Izin: `data/processed/sulawesi_izin_baru_per_tahun.csv` dan `data/processed/sulawesi_izin_raw_details.csv` (MODI ESDM); Data Deforestasi: `data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv` (GFW). Visualisasi dashboard menampilkan Matriks Kepatuhan D3TLH (Seharusnya vs Kenyataan per status Aman/Tertekan/Kritis) beserta Tabel Irisan daftar perusahaan penerima izin di zona kritis.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "D3TLH dirancang sebagai instrumen pencegahan dan pengatur batas pengaman ekologis (*ecological safeguard*). Penyandingan data deforestasi tahunan GFW dan data perizinan MODI ESDM menunjukkan bahwa penerbitan izin usaha pertambangan baru tetap tercatat pada kurun waktu ketika perubahan tutupan hutan meningkat — terlihat pada tren Sulawesi Tengah dan Tenggara 2014-2023. Kondisi ini menggarisbawahi pentingnya penguatan fungsi AMDAL, D3TLH, dan KLHS sebagai pertimbangan yang mengikat dalam keputusan perizinan.",
        "",
        "#### B. Alur Logika Metodologis Evaluasi Kepatuhan D3TLH (Rule-based Categorization)",
        "Kerangka agregasi berbasis aturan diilustrasikan pada **Bagan Alur 7.1** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Compliance Modeling deskriptif dengan binning persentil tiga kelas.",
        "",
        "##### Bagan Alur 7.1: Alur Logika Analisis Evaluasi Kepatuhan D3TLH (Rule-based Categorization)",
        "```mermaid",
        mermaid_str_7_1,
        "```",
        "",
        "#### C. Formulasi Matematis: Ambang Persentil, Klasifikasi Status, dan Kuantifikasi Pelanggaran",
        "Kuantifikasi status daya dukung dan pelanggaran ekologis dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Ambang_Tertekan = Persentil_33 ( D_p,t )   ;   Ambang_Kritis = Persentil_66 ( D_p,t )",
        "Status(x) = 'Aman' , jika x ≤ P33   |   'Tertekan' , jika P33 < x ≤ P66   |   'Kritis' , jika x > P66",
        "Total_Izin_Zona_Kritis = Σ ( IUP_p,t )   ;   untuk seluruh observasi dengan Status = 'Kritis'",
        "```",
        "",
        "Substitusi angka dari dataset aktual ke dalam rumus ambang persentil dan kuantifikasi pelanggaran:",
        "",
        "```text",
        f"Ambang_Tertekan = Persentil_33 ( D ) = {tertekan_threshold_71:,.1f} Ha   ;   Ambang_Kritis = Persentil_66 ( D ) = {kritis_threshold_71:,.1f} Ha",
        f"N_Aman = {n_per_status_71.get('Aman', 0)} observasi   ;   N_Tertekan = {n_per_status_71.get('Tertekan', 0)} observasi   ;   N_Kritis = {n_per_status_71.get('Kritis', 0)} observasi   (Total N = {n_panel_71})",
        f"Σ IUP_Aman = {iup_per_status_71.get('Aman', 0)} izin   ;   Σ IUP_Tertekan = {iup_per_status_71.get('Tertekan', 0)} izin   ;   Σ IUP_Kritis = {iup_per_status_71.get('Kritis', 0)} izin ({luas_kritis_71:,.0f} Ha)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 7.1: Matriks Kepatuhan D3TLH - Seharusnya vs Kenyataan per Status Daya Dukung",
        markdown_table(["Status Daya Dukung", "Rentang Kerusakan Hutan", "N Observasi", "Seharusnya (Menurut Aturan)", "Kenyataan di Lapangan", "Total Luas Konsesi (Ha)", "Kesimpulan Tata Kelola"], status_rows_71),
        "",
        f"##### Tabel 7.2: Sampel Irisan Izin IUP Terbit di Zona Kritis - 10 Konsesi Terluas (dari {n_izin_irisan_71:,} izin)",
        markdown_table(["Nama Perusahaan (IUP)", "Komoditas", "Provinsi", "Tahun Terbit", "Kehilangan Hutan Provinsi (Ha)", "Luas Konsesi (Ha)"], irisan_rows_71),
        "",
        "#### E. Analisis Temuan Empiris: Konklusi Kepatuhan D3TLH Berdasarkan Data Historis",
        f"1. **Fungsi Pembatas D3TLH Perlu Ditingkatkan:** terdapat **{total_iup_kritis_71} Izin Baru** yang terbit pada periode berstatus deforestasi tinggi (Kritis), mencakup luasan konsesi {luas_kritis_71:,.0f} Ha. Apabila pada status Kritis izin masih diterbitkan, hal tersebut secara matematis mendiskualifikasi D3TLH sebagai instrumen perlindungan lingkungan.",
        f"2. **Bukti Irisan 100% Data-Driven:** tabel irisan ESDM x GFW mengidentifikasi **{n_izin_irisan_71:,} izin IUP baru** (total luas {luas_irisan_71:,.0f} Ha) yang tetap diterbitkan di tengah situasi kritis pada pasangan Provinsi-Tahun yang sama.",
        "3. **Implikasi Kebijakan:** diperlukan penguatan integrasi data lingkungan dalam keputusan perizinan agar dokumen AMDAL, D3TLH, dan KLHS menjadi pertimbangan utama yang mengikat.",
        "",
        "## 7.2 Tabrakan Hukum: Impunitas dan Pembiaran Operasi Ilegal",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Konflik Hukum: `data/processed/sulawesi_konflik_hukum.csv` (kompilasi Konsorsium Pembaruan Agraria / KPA, TanahKita, dan laporan LSM). Visualisasi dashboard menampilkan metrik total kasus, Bubble Timeline sebaran konflik agraria 10 tahun terakhir (Tahun × Provinsi), tabel sebaran sektor konflik, serta daftar rekam jejak kasus.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Konsep D3TLH mengukur kapasitas daya tahan ekosistem serta daya dukung sosial masyarakat di sekitar kawasan industri. Kompilasi laporan masyarakat sipil mencatat adanya sengketa tanah dan dinamika sosial dalam ekspansi industri ekstraktif — menunjukkan pentingnya kepatuhan perizinan, penerapan sanksi administratif yang konsisten, pengawasan batas wilayah perizinan (HGU/IUP), pelaksanaan konsultasi publik (FPIC), serta penyelesaian sengketa tenurial secara adil.",
        "",
        "#### B. Alur Logika Metodologis Pemetaan Impunitas Korporasi (Thematic Coding)",
        "Kerangka agregasi pelaporan berbasis insiden (*Incident-based Reporting Aggregation*) diilustrasikan pada **Bagan Alur 7.2** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Thematic Coding dan analisis kasus deskriptif.",
        "",
        "##### Bagan Alur 7.2: Alur Logika Analisis Pemetaan Impunitas Korporasi (Thematic Coding)",
        "```mermaid",
        mermaid_str_7_2,
        "```",
        "",
        "#### C. Formulasi Matematis: Agregasi Insiden dan Volume Pembiaran Sektoral",
        "Kuantifikasi tingkat pembiaran penegakan hukum dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Total_Kasus_Impunitas = Σ ( Kasus_i )   ;   untuk seluruh laporan insiden i dalam database",
        "Volume_Sektoral_s = Σ ( Kasus_i )   ;   untuk seluruh kasus i dengan Sektor = s",
        "Proporsi_Sektoral_s (%) = ( Volume_Sektoral_s / Total_Kasus_Impunitas ) × 100",
        "```",
        "",
        "Substitusi angka dari dataset aktual ke dalam rumus agregasi insiden:",
        "",
        "```text",
        f"Total_Kasus_Impunitas = {total_kasus_72} kasus   ;   {n_kasus_dekade_72} kasus pada dekade 2014-{tahun_max_72}",
        f"Volume_Sektoral_{top_sektor_72} = {top_sektor_n_72} kasus ({top_sektor_pct_72:.1f}%)   ;   Volume_Sektoral_{sektor2_72} = {sektor2_n_72} kasus",
        f"Volume_Provinsi_Tertinggi: {top_prov_72} = {top_prov_n_72} kasus (dari total {total_kasus_72} kasus se-Sulawesi)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 7.3: Sebaran Sektor Konflik dan Pembiaran Operasi Ilegal",
        markdown_table(["Sektor (Penyebab)", "Jumlah Kasus", "Proporsi (%)"], sektor_rows_72),
        "",
        "##### Tabel 7.4: Sebaran Wilayah Konflik per Provinsi dan Rentang Tahun Kejadian",
        markdown_table(["Provinsi", "Jumlah Kasus", "Proporsi (%)", "Rentang Tahun"], prov_rows_72),
        "",
        "##### Tabel 7.5: Sampel Rekam Jejak Konflik Agraria & Pelanggaran Hak (10 Kasus Terbaru)",
        markdown_table(["Tahun", "Provinsi", "Sektor", "Judul/Nama Kasus", "Sumber"], sampel_rows_72),
        "",
        "#### E. Analisis Temuan Empiris: Impunitas dan Pembiaran Struktural",
        f"1. **Bukti Impunitas Hukum:** database kompilasi LSM/KPA mendokumentasikan **{total_kasus_72} kasus** konflik/pelanggaran yang dibiarkan di Sulawesi ({n_kasus_dekade_72} kasus pada dekade 2014-{tahun_max_72}), di mana korporasi yang terbukti bermasalah secara hukum tetap dipertahankan keberadaan operasinya.",
        f"2. **Dominasi Sektor Ekstraktif:** sektor **{top_sektor_72}** menjadi penyebab konflik terbanyak dengan {top_sektor_n_72} kasus ({top_sektor_pct_72:.1f}%), disusul {sektor2_72} ({sektor2_n_72} kasus) — konsisten dengan pola tekanan ekspansi industri ekstraktif pada bab-bab sebelumnya.",
        f"3. **Konsentrasi Spasial:** provinsi **{top_prov_72}** mencatat kasus terbanyak ({top_prov_n_72} kasus) — menegaskan pentingnya penguatan koordinasi antar-instansi, penegakan sanksi administratif yang konsisten, serta penyelesaian sengketa tenurial secara adil.",
        "",
        "## 7.3 Inkonsistensi Iklim: Karpet Merah PLTU Captive",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data PLTU Captive: `data/processed/sulawesi_pltu_captive.csv` (Global Coal Plant Tracker / GEM, ekstraksi Januari 2026). Visualisasi dashboard menampilkan metrik total unit dan kapasitas pembangkitan kotor, Line Chart timeline pertumbuhan kapasitas kumulatif per provinsi hingga 2024, serta daftar lengkap unit PLTU batubara captive.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Pemenuhan kebutuhan energi untuk fasilitas pengolahan nikel (*smelter*) di Sulawesi masih didominasi oleh PLTU Batubara *Captive*. Data GEM mencatat **{total_unit_73} unit PLTU Captive** yang beroperasi maupun direncanakan di kawasan industri Sulawesi dengan total kapasitas **{total_mw_73:,.0f} MW** — tantangan tersendiri bagi pengelolaan emisi gas rumah kaca dan kualitas udara ambien, yang menunjukkan perlunya percepatan transisi energi bersih agar target hilirisasi selaras dengan komitmen penurunan emisi nasional.",
        "",
        "#### B. Alur Logika Metodologis Agregasi Beban Karbon PLTU Captive",
        "Kerangka inventarisasi agregat kuantitatif (*Quantitative Inventory Aggregation*) diilustrasikan pada **Bagan Alur 7.3** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan penyaringan agregat dataset eksternal deskriptif.",
        "",
        "##### Bagan Alur 7.3: Alur Logika Analisis Agregasi Beban Karbon PLTU Captive (GEM)",
        "```mermaid",
        mermaid_str_7_3,
        "```",
        "",
        "#### C. Formulasi Matematis: Inventarisasi Unit, Beban Karbon, dan Timeline Kumulatif",
        "Kuantifikasi kontradiksi karbon dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Total_Infrastruktur_Kotor = Σ ( Unit_u )   ;   untuk seluruh unit PLTU captive u di kawasan industri Sulawesi",
        "Total_Beban_Karbon_p = Σ ( Kapasitas_u )   ;   untuk seluruh unit u pada provinsi p",
        "Kumulatif_Kapasitas_p(T) = Σ Kapasitas_p(t)   ;   untuk t = tahun operasi awal s.d. T (T ≤ 2024)",
        "```",
        "",
        "Substitusi angka dari dataset aktual ke dalam rumus inventarisasi:",
        "",
        "```text",
        f"Total_Infrastruktur_Kotor = {total_unit_73} unit   ;   Total_Beban_Karbon = {total_mw_73:,.0f} MW",
        f"Total_Beban_Karbon_Tertinggi: {top_prov_73} = {top_prov_mw_73:,.0f} MW ({top_prov_pct_73:.1f}% dari total)",
        f"Operating = {op_unit_73} unit ({op_mw_73:,.0f} MW)   ;   Non-Operating = {total_unit_73 - op_unit_73} unit",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 7.6: Agregat Unit dan Kapasitas PLTU Captive per Provinsi (Total {total_mw_73:,.0f} MW)",
        markdown_table(["Provinsi", "Jumlah Unit", "Kapasitas (MW)", "Proporsi (%)"], prov_rows_73),
        "",
        "##### Tabel 7.7: Timeline Pertumbuhan Kapasitas Kumulatif PLTU Captive per Provinsi (MW, s.d. 2024)",
        markdown_table(["Tahun"] + prov_cols_73 + ["Total (MW)"], timeline_rows_73),
        "",
        "##### Tabel 7.8: Sepuluh Unit PLTU Captive Berkapasitas Terbesar di Sulawesi",
        markdown_table(["Nama Pembangkit", "Unit", "Provinsi", "Kapasitas (MW)", "Status", "Tahun Beroperasi"], unit_rows_73),
        "",
        "#### E. Analisis Temuan Empiris: Paradoks Hilirisasi Hijau",
        f"1. **Skala Infrastruktur Kotor:** GEM mencatat **{total_unit_73} unit** PLTU batubara captive di kawasan industri Sulawesi dengan total kapasitas **{total_mw_73:,.0f} MW** ({op_unit_73} unit beroperasi, {op_mw_73:,.0f} MW) — pembangkitan masif yang dibangun demi menopang pabrik pemurnian nikel yang dipromosikan sebagai proyek energi ramah lingkungan.",
        f"2. **Konsentrasi Beban Karbon:** provinsi **{top_prov_73}** menanggung beban karbon terbesar dengan {top_prov_mw_73:,.0f} MW ({top_prov_pct_73:.1f}% dari total), sejalan dengan konsentrasi kawasan industri hilirisasi.",
        "3. **Implikasi Transisi Energi:** diperlukan strategi percepatan transisi energi bersih di sektor industri ekstraktif guna menyelaraskan target hilirisasi dengan komitmen penurunan emisi nasional.",
        "",
    ]
    md_path = tool_dir / "Metodologi_Bab7_Kegagalan_Tata_Kelola.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")
    print("[4/4] Selesai membangun Bab 7.")


if __name__ == "__main__":
    generate_all_bab7()
