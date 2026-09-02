#!/usr/bin/env python3
"""
Generator Laporan Metodologi Bab 4: Ruang Hidup yang Terampas

Fokus awal: Sub-bab 4.1 Tren Eskalasi Konflik Agraria Seiring Ekspansi
Industri. Pilar 1 ditulis langsung dalam generator Python agar selaras
dengan SOP dokumentasi Celios2.
"""

import base64
import re
import sys
from pathlib import Path

try:
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([
        sys.executable,
        "-m",
        "pip",
        "install",
        "pandas",
        "requests",
        "python-docx",
    ])
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge, cfg in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        if cfg is None:
            el.set(qn("w:val"), "none")
        else:
            for key, val in cfg.items():
                el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_margin(cell, left=100, right=100, top=60, bottom=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margin.append(el)
    tc_pr.append(margin)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_border_bottom(paragraph, color="2E7D32", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="2E7D32", sz="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="444444", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="1B5E20", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="2E7D32", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=G_MID)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)


def add_p(doc, parts, space_after=5, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "EDF7EE")
    all_border_para(p, color="A5D6A7", sz="4")
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)
    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(2)
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=G_DARK)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p, color="2E7D32", sz="16")
    para_shd(p, "F1F8E9")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}
    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "2E7D32")
        cell_margin(cell, left=100, right=100, top=70, bottom=70)
        set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom={"val": "single", "sz": "8", "color": "1B5E20", "space": "0"}, right=bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8.5, color=C_WHITE)
    for i, row_data in enumerate(rows):
        fill = "F5FBF5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            cell_margin(cell, left=100, right=100, top=50, bottom=50)
            set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def map_sektor(status):
    status = str(status).lower()
    if "kebun" in status:
        return "Perkebunan"
    if "tambang" in status:
        return "Pertambangan"
    if "hutan" in status:
        return "Kehutanan"
    if any(x in status for x in ["infrastruktur", "bendungan", "transmigrasi", "energi", "fasilitas", "jalan", "industri"]):
        return "Infrastruktur & PSN"
    if any(x in status for x in ["pariwisata", "laut", "pesisir"]):
        return "Pariwisata & Pesisir"
    return "Lainnya"


def generate_all_bab4():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    data_dir = base_dir / "data" / "processed"
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_4"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 4 sub-bab 4.1...")
    df_raw = pd.read_csv(data_dir / "sulawesi_konflik_agraria_tanahkita.csv")
    keywords = r"\b(?:sulawesi|sulsel|sulteng|sultra|sulut|sulbar|gorontalo|morowali|konawe|kolaka|bombana|poso|donggala|makassar|manado|minahasa|sangihe|mamuju|majene|polewali|halmahera|maluku utara|weda|obi|soroako|luwu|bantaeng|buton|muna|wakatobi|banggai|buol|toli-toli|parigi|luwuk|kendari|baubau|palu|bitung|tomohon|kotamobagu|gowa|takalar|jeneponto|bulukumba|sinjai|bone|maros|pangkep|barru|pinrang|enrekang|toraja|palopo)\b"
    mask = (
        df_raw["judul"].str.contains(keywords, case=False, na=False, regex=True)
        | df_raw["deskripsi"].str.contains(keywords, case=False, na=False, regex=True)
        | df_raw["narasi"].str.contains(keywords, case=False, na=False, regex=True)
        | df_raw["lokasi"].str.contains(keywords, case=False, na=False, regex=True)
    )
    df_konflik = df_raw[mask].copy()
    df_konflik["tahun"] = pd.to_numeric(df_konflik["tahun"], errors="coerce")
    df_konflik = df_konflik.dropna(subset=["tahun"])
    df_konflik["tahun"] = df_konflik["tahun"].astype(int)
    df_konflik["Sektor_Grup"] = df_konflik["status"].apply(map_sektor)
    df_konflik["dampak_masyarakat_jiwa"] = pd.to_numeric(df_konflik["dampak_masyarakat_jiwa"], errors="coerce").fillna(0)

    total_konflik = len(df_konflik)
    konflik_kebun = len(df_konflik[df_konflik["status"].str.contains("Perkebunan", case=False, na=False)])
    konflik_tambang = len(df_konflik[df_konflik["status"].str.contains("Pertambangan", case=False, na=False)])
    konflik_hutan = len(df_konflik[df_konflik["status"].str.contains("Hutan", case=False, na=False)])
    konflik_infrastruktur = len(df_konflik[df_konflik["status"].str.contains("Infrastruktur|Bendungan|Transmigrasi|Energi|Fasilitas|Jalan", case=False, na=False)])
    konflik_pariwisata = len(df_konflik[df_konflik["status"].str.contains("Pariwisata|Konservasi Laut", case=False, na=False)])
    rasio_ekstraktif = ((konflik_tambang + konflik_kebun + konflik_hutan) / total_konflik) * 100 if total_konflik else 0
    total_jiwa = int(df_konflik["dampak_masyarakat_jiwa"].sum())
    status_belum_selesai = len(df_konflik[df_konflik["status_konflik"].str.contains("Belum Ditangani", na=False)])
    libat_masyarakat = df_konflik["keterlibatan_masyarakat"].notna().sum()

    df_ts_modern = df_konflik[df_konflik["tahun"] >= 1990].copy()
    total_ts = len(df_konflik)
    pasca_2005 = len(df_konflik[df_konflik["tahun"] >= 2005])
    pra_2005 = len(df_konflik[df_konflik["tahun"] < 2005])
    lonjakan = (pasca_2005 / pra_2005 * 100) if pra_2005 > 0 else 0

    df_agg = df_ts_modern.groupby(["tahun", "Sektor_Grup"]).size().reset_index(name="Jumlah")
    df_total_per_tahun = df_ts_modern.groupby("tahun").size().reset_index(name="Jumlah")
    peak_year = 0
    peak_value = 0
    if not df_total_per_tahun.empty:
        max_year_row = df_total_per_tahun.loc[df_total_per_tahun["Jumlah"].idxmax()]
        peak_year = int(max_year_row["tahun"])
        peak_value = int(max_year_row["Jumlah"])

    sector_rows = []
    sector_total = df_ts_modern.groupby("Sektor_Grup").size().reset_index(name="Jumlah").sort_values("Jumlah", ascending=False)
    for _, row in sector_total.iterrows():
        sector_rows.append([row["Sektor_Grup"], f"{int(row['Jumlah']):,}", f"{row['Jumlah'] / len(df_ts_modern) * 100:.1f}%"])

    annual_rows = []
    for _, row in df_total_per_tahun.sort_values("tahun").iterrows():
        annual_rows.append([str(int(row["tahun"])), f"{int(row['Jumlah']):,}"])

    peak_sector_rows = []
    if peak_year:
        peak_sector = df_agg[df_agg["tahun"] == peak_year].sort_values("Jumlah", ascending=False)
        for _, row in peak_sector.iterrows():
            peak_sector_rows.append([str(peak_year), row["Sektor_Grup"], f"{int(row['Jumlah']):,}"])

    mermaid_str_4_1 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Konflik Agraria KPA/Tanah Kita<br/><i>judul, deskripsi, lokasi, status, tahun</i>"]
    end
    subgraph Regional_Filter["2. Filter Regional & Klasifikasi Sektor"]
        A --> B["Filter keyword Sulawesi, Maluku Utara,<br/>dan sentra nikel terkait"]
        B --> C["Klasifikasi sektor pemicu<br/>Perkebunan, Kehutanan, Pertambangan, PSN, Pesisir"]
    end
    subgraph Trend_Analysis["3. Time-Series Trend Analysis"]
        C --> D["Agregasi konflik tahunan sejak 1990"]
        D --> E["Komparasi pra-2005 vs pasca-2005"]
        D --> F["Identifikasi puncak insidensi konflik"]
    end
    E --> G["Pembacaan eskalasi konflik agraria"]
    F --> G"""
    mermaid_png_path_4_1 = str(tool_dir / "mermaid_flowchart_4_1.png")
    download_success_4_1 = download_mermaid_png(mermaid_str_4_1, mermaid_png_path_4_1)

    print("[1.5/4] Mengekstraksi dataset empiris Bab 4 sub-bab 4.2...")
    df_dampak = df_konflik.copy()
    df_dampak["dampak_masyarakat_jiwa"] = pd.to_numeric(df_dampak["dampak_masyarakat_jiwa"], errors="coerce").fillna(0)
    df_dampak["luas_ha"] = pd.to_numeric(df_dampak["luas_ha"], errors="coerce").fillna(0)

    df_sektor_agg = df_dampak.groupby("Sektor_Grup").agg({
        "dampak_masyarakat_jiwa": "sum",
        "luas_ha": "sum",
    }).reset_index()
    df_sektor_agg = df_sektor_agg[df_sektor_agg["Sektor_Grup"] != "Lainnya"].copy()

    def sector_sum(col, sector):
        vals = df_sektor_agg.loc[df_sektor_agg["Sektor_Grup"] == sector, col]
        return float(vals.sum()) if not vals.empty else 0.0

    jiwa_kehutanan = sector_sum("dampak_masyarakat_jiwa", "Kehutanan")
    jiwa_tambang = sector_sum("dampak_masyarakat_jiwa", "Pertambangan")
    ha_kebun = sector_sum("luas_ha", "Perkebunan")
    ha_kehutanan = sector_sum("luas_ha", "Kehutanan")
    ha_tambang = sector_sum("luas_ha", "Pertambangan")
    total_jiwa_42 = df_sektor_agg["dampak_masyarakat_jiwa"].sum()
    total_ha_42 = df_sektor_agg["luas_ha"].sum()

    sektor_dampak_rows = []
    for _, row in df_sektor_agg.sort_values("dampak_masyarakat_jiwa", ascending=False).iterrows():
        sektor_dampak_rows.append([
            row["Sektor_Grup"],
            f"{row['dampak_masyarakat_jiwa']:,.0f}",
            f"{row['dampak_masyarakat_jiwa'] / total_jiwa_42 * 100:.1f}%" if total_jiwa_42 else "0.0%",
            f"{row['luas_ha']:,.0f}",
            f"{row['luas_ha'] / total_ha_42 * 100:.1f}%" if total_ha_42 else "0.0%",
        ])

    df_sektor_tahun = df_dampak[df_dampak["tahun"] >= 1990].groupby(["tahun", "Sektor_Grup"]).agg({
        "dampak_masyarakat_jiwa": "sum",
        "luas_ha": "sum",
    }).reset_index()
    df_sektor_tahun = df_sektor_tahun[df_sektor_tahun["Sektor_Grup"] != "Lainnya"].copy()

    top_jiwa = df_sektor_tahun.groupby("tahun")["dampak_masyarakat_jiwa"].sum().sort_values(ascending=False)
    top_jiwa = top_jiwa[top_jiwa > 0].head(2)
    top_ha = df_sektor_tahun.groupby("tahun")["luas_ha"].sum().sort_values(ascending=False)
    top_ha = top_ha[top_ha > 0].head(2)

    anomaly_jiwa_rows = []
    for i, (year, val) in enumerate(top_jiwa.items(), 1):
        cases = df_dampak[df_dampak["tahun"] == year].copy()
        cases["jiwa_num"] = pd.to_numeric(cases["dampak_masyarakat_jiwa"], errors="coerce").fillna(0)
        if cases.empty:
            continue
        top_case = cases.sort_values("jiwa_num", ascending=False).iloc[0]
        anomaly_jiwa_rows.append([
            f"Anomali Jiwa {i}",
            str(int(year)),
            top_case["Sektor_Grup"],
            f"{top_case['jiwa_num']:,.0f}",
            str(top_case["judul"]),
        ])

    anomaly_ha_rows = []
    for i, (year, val) in enumerate(top_ha.items(), 1):
        cases = df_dampak[df_dampak["tahun"] == year].copy()
        cases["ha_num"] = pd.to_numeric(cases["luas_ha"], errors="coerce").fillna(0)
        if cases.empty:
            continue
        top_case = cases.sort_values("ha_num", ascending=False).iloc[0]
        anomaly_ha_rows.append([
            f"Anomali Area {i}",
            str(int(year)),
            top_case["Sektor_Grup"],
            f"{top_case['ha_num']:,.0f}",
            str(top_case["judul"]),
        ])

    mermaid_str_4_2 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Konflik Agraria<br/><i>sektor, jiwa terdampak, luas ha, tahun</i>"]
    end
    subgraph Sectoral_Burden["2. Sectoral Burden Analysis"]
        A --> B["Klasifikasi sektor pemicu konflik"]
        B --> C["Agregasi korban terdampak (jiwa)<br/>per sektor dan tahun"]
        B --> D["Agregasi luas area konflik (Ha)<br/>per sektor dan tahun"]
    end
    subgraph Comparative_Output["3. Analisis Komparatif Dampak"]
        C --> E["Distribusi korban terdampak"]
        D --> F["Distribusi monopoli area konflik"]
        E --> G["Bedah anomali lonjakan jiwa"]
        F --> H["Bedah anomali lonjakan area"]
    end"""
    mermaid_png_path_4_2 = str(tool_dir / "mermaid_flowchart_4_2.png")
    download_success_4_2 = download_mermaid_png(mermaid_str_4_2, mermaid_png_path_4_2)

    print("[1.7/4] Mengekstraksi dataset empiris Bab 4 sub-bab 4.3...")
    df_dampak["jumlah_ditangkap"] = pd.to_numeric(df_dampak["jumlah_ditangkap"], errors="coerce").fillna(0)
    df_dampak["jumlah_luka"] = pd.to_numeric(df_dampak["jumlah_luka"], errors="coerce").fillna(0)
    df_dampak["jumlah_tewas"] = pd.to_numeric(df_dampak["jumlah_tewas"], errors="coerce").fillna(0)
    df_dampak["indikasi_kriminalisasi"] = df_dampak["indikasi_kriminalisasi"].fillna(False).astype(bool)

    total_kriminalisasi = df_dampak[df_dampak["indikasi_kriminalisasi"]].shape[0]
    total_ditangkap = int(df_dampak["jumlah_ditangkap"].sum())
    total_luka = int(df_dampak["jumlah_luka"].sum())
    total_tewas = int(df_dampak["jumlah_tewas"].sum())

    df_krim_tahun = df_dampak[(df_dampak["indikasi_kriminalisasi"]) & (df_dampak["tahun"] >= 2000)].groupby("tahun").size().reset_index(name="jumlah_kasus")
    df_krim_sektor = df_dampak[(df_dampak["indikasi_kriminalisasi"]) & (df_dampak["Sektor_Grup"] != "Lainnya")].groupby("Sektor_Grup").size().reset_index(name="jumlah_kasus").sort_values("jumlah_kasus", ascending=True)

    top_sektor = df_krim_sektor.iloc[-1]["Sektor_Grup"] if not df_krim_sektor.empty else "Industri"
    top_sektor_count = int(df_krim_sektor.iloc[-1]["jumlah_kasus"]) if not df_krim_sektor.empty else 0
    top_tahun = int(df_krim_tahun.loc[df_krim_tahun["jumlah_kasus"].idxmax()]["tahun"]) if not df_krim_tahun.empty else 0
    top_tahun_count = int(df_krim_tahun["jumlah_kasus"].max()) if not df_krim_tahun.empty else 0

    krim_tahun_rows = []
    for _, row in df_krim_tahun.sort_values("tahun").iterrows():
        krim_tahun_rows.append([str(int(row["tahun"])), f"{int(row['jumlah_kasus']):,}"])

    krim_sektor_rows = []
    for _, row in df_krim_sektor.sort_values("jumlah_kasus", ascending=False).iterrows():
        krim_sektor_rows.append([row["Sektor_Grup"], f"{int(row['jumlah_kasus']):,}"])

    df_kekerasan = df_dampak[(df_dampak["jumlah_ditangkap"] > 0) | (df_dampak["jumlah_tewas"] > 0)].sort_values(["jumlah_ditangkap", "jumlah_tewas"], ascending=[False, False])
    kekerasan_rows = []
    def clean_table_text(value):
        if pd.isna(value) or str(value).strip().lower() == "nan":
            return "-"
        return str(value).replace("|", ",").strip()

    for _, row in df_kekerasan.head(10).iterrows():
        perusahaan = clean_table_text(row["keterlibatan_perusahaan"] if pd.notna(row["keterlibatan_perusahaan"]) else "Tidak/Belum Teridentifikasi")
        deskripsi = clean_table_text(row["deskripsi"])
        if len(deskripsi) > 180:
            deskripsi = deskripsi[:177] + "..."
        kekerasan_rows.append([
            str(int(row["tahun"])),
            clean_table_text(row["Sektor_Grup"]),
            perusahaan,
            f"{int(row['jumlah_ditangkap']):,}",
            f"{int(row['jumlah_tewas']):,}",
            deskripsi,
        ])

    mermaid_str_4_3 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Konflik Agraria<br/><i>indikasi kriminalisasi, ditangkap, luka, tewas</i>"]
    end
    subgraph Violence_Tracking["2. Violence & Criminalization Tracking"]
        A --> B["Standarisasi kolom korban<br/>ditangkap, luka, tewas"]
        B --> C["Hitung total kasus kriminalisasi"]
        B --> D["Agregasi tren kriminalisasi pasca-2000"]
        B --> E["Agregasi sektor paling represif"]
    end
    subgraph Output["3. Pemetaan Ruang Sipil"]
        C --> F["Metrik represi agregat"]
        D --> G["Tren tahunan kriminalisasi"]
        E --> H["Sektor dominan represi"]
        F --> I["Pembacaan risiko HAM"]
        G --> I
        H --> I
    end"""
    mermaid_png_path_4_3 = str(tool_dir / "mermaid_flowchart_4_3.png")
    download_success_4_3 = download_mermaid_png(mermaid_str_4_3, mermaid_png_path_4_3)

    print("[1.9/4] Mengekstraksi dataset empiris Bab 4 sub-bab 4.4...")
    df_ba_44 = df_dampak[df_dampak["tahun"] >= 1990].copy()
    df_pra_44 = df_ba_44[df_ba_44["tahun"] < 2014]
    df_pasca_44 = df_ba_44[df_ba_44["tahun"] >= 2014]
    tahun_pra_44 = max(1, 2014 - int(df_pra_44["tahun"].min())) if not df_pra_44.empty else 24
    tahun_pasca_44 = max(1, int(df_pasca_44["tahun"].max()) - 2013) if not df_pasca_44.empty else 11
    avg_pra_44 = len(df_pra_44) / tahun_pra_44
    avg_pasca_44 = len(df_pasca_44) / tahun_pasca_44
    before_after_rows_44 = [
        ["Pra-Ekspansi (<2014)", f"{len(df_pra_44):,}", f"{tahun_pra_44:,}", f"{avg_pra_44:.1f}", f"{int(df_pra_44['jumlah_ditangkap'].sum()):,}", f"{int(df_pra_44['jumlah_tewas'].sum()):,}"],
        ["Pasca-Ekspansi (>=2014)", f"{len(df_pasca_44):,}", f"{tahun_pasca_44:,}", f"{avg_pasca_44:.1f}", f"{int(df_pasca_44['jumlah_ditangkap'].sum()):,}", f"{int(df_pasca_44['jumlah_tewas'].sum()):,}"],
    ]

    df_crosstab_44 = pd.read_csv(data_dir / "sulawesi_konflik_agraria_tanahkita.csv")
    df_crosstab_44["tahun"] = pd.to_numeric(df_crosstab_44["tahun"], errors="coerce")
    df_crosstab_44 = df_crosstab_44[df_crosstab_44["tahun"] >= 1990].copy()
    df_crosstab_44["Periode_Ekspansi"] = df_crosstab_44["tahun"].apply(lambda x: "Pasca-ekspansi (>= 2014)" if x >= 2014 else "Pra-ekspansi (< 2014)")
    df_crosstab_44["Sektor_Tambang"] = df_crosstab_44["status"].str.contains("Tambang|Pertambangan", case=False, na=False).apply(lambda x: "Sektor Pertambangan" if x else "Sektor Non-Tambang")
    df_crosstab_44["Keterlibatan_Pemerintah"] = df_crosstab_44["keterlibatan_pemerintah"].notna().apply(lambda x: "Terlibat Aparat/Negara" if x else "Tanpa Keterlibatan Negara")
    df_crosstab_44["Indikasi_Kriminalisasi"] = df_crosstab_44["indikasi_kriminalisasi"].fillna(False).astype(bool).apply(lambda x: "Ada Represi/Kriminalisasi" if x else "Baseline (Tanpa Kriminalisasi)")
    df_crosstab_44["Status_Penyelesaian"] = df_crosstab_44["status_konflik"].str.contains("Belum Ditangani", na=False).apply(lambda x: "Konflik Dibiarkan Terlantar" if x else "Konflik Selesai/Diproses")
    has_luka_44 = pd.to_numeric(df_crosstab_44["jumlah_luka"], errors="coerce").fillna(0) > 0
    has_tewas_44 = pd.to_numeric(df_crosstab_44["jumlah_tewas"], errors="coerce").fillna(0) > 0
    has_tangkap_44 = pd.to_numeric(df_crosstab_44["jumlah_ditangkap"], errors="coerce").fillna(0) > 0
    df_crosstab_44["Dampak_Kekerasan"] = (has_luka_44 | has_tewas_44 | has_tangkap_44).apply(lambda x: "Terjadi Kekerasan/Penangkapan" if x else "Tanpa Insiden Fisik")

    x_options_44 = {
        "Periode_Ekspansi": "Periode Ekspansi Industri",
        "Sektor_Tambang": "Tipe Sektor (Tambang vs Non-Tambang)",
        "Keterlibatan_Pemerintah": "Keterlibatan Aparat/Pemerintah",
    }
    y_options_44 = {
        "Indikasi_Kriminalisasi": "Tingkat Represi & Kriminalisasi",
        "Status_Penyelesaian": "Tingkat Penelantaran Kasus",
        "Dampak_Kekerasan": "Tingkat Insiden Fisik (Luka/Tewas/Ditangkap)",
    }
    x_order_44 = {
        "Periode_Ekspansi": ["Pra-ekspansi (< 2014)", "Pasca-ekspansi (>= 2014)"],
        "Sektor_Tambang": ["Sektor Non-Tambang", "Sektor Pertambangan"],
        "Keterlibatan_Pemerintah": ["Tanpa Keterlibatan Negara", "Terlibat Aparat/Negara"],
    }
    y_order_44 = {
        "Indikasi_Kriminalisasi": ["Baseline (Tanpa Kriminalisasi)", "Ada Represi/Kriminalisasi"],
        "Dampak_Kekerasan": ["Tanpa Insiden Fisik", "Terjadi Kekerasan/Penangkapan"],
        "Status_Penyelesaian": ["Konflik Selesai/Diproses", "Konflik Dibiarkan Terlantar"],
    }

    summary_rows_44 = []
    for k_x, v_x in x_options_44.items():
        for k_y, v_y in y_options_44.items():
            ct = pd.crosstab(df_crosstab_44[k_x], df_crosstab_44[k_y]).reindex(index=x_order_44[k_x], columns=y_order_44[k_y], fill_value=0)
            try:
                c2_val, pv_val, dof_val, _ = stats.chi2_contingency(ct)
            except Exception:
                c2_val, pv_val, dof_val = 0, 1.0, 0
            try:
                aa = ct.loc[x_order_44[k_x][0], y_order_44[k_y][0]]
                bb = ct.loc[x_order_44[k_x][0], y_order_44[k_y][1]]
                cc = ct.loc[x_order_44[k_x][1], y_order_44[k_y][0]]
                dd = ct.loc[x_order_44[k_x][1], y_order_44[k_y][1]]
                or_v = (aa * dd) / (bb * cc) if (bb * cc) > 0 else 0
            except Exception:
                or_v = 0
            p_disp = "p < 0.001" if pv_val < 0.001 else f"p = {pv_val:.3f}"
            summary_rows_44.append([v_x, v_y, f"{c2_val:.3f}", p_disp, f"{or_v:.2f}", "SIGNIFIKAN" if pv_val < 0.05 else "TIDAK SIGNIFIKAN"])

    sig_count_44 = sum(1 for row in summary_rows_44 if row[5] == "SIGNIFIKAN")
    total_scenarios_44 = len(summary_rows_44)
    valid_cases_44 = len(df_crosstab_44)

    konf_headers_44 = ["Komponen Uji", "Definisi Variabel (Sub-bab 4.4)"]
    konf_rows_44 = [
        ["Matriks Ekspansi (X)", "Periode Ekspansi Industri; Tipe Sektor (Tambang vs Non-Tambang); Keterlibatan Aparat/Pemerintah."],
        ["Matriks Eskalasi (Y)", "Tingkat Represi & Kriminalisasi; Tingkat Penelantaran Kasus; Tingkat Insiden Fisik."],
        ["Hipotesis Nol (H0)", "Variabel baris (Periode/Aktor) saling bebas secara absolut terhadap variabel kolom (Represi/Kematian)."],
        ["Decision Rule", "Chi-Square P-Value < 0.05, maka tolak H0 dan terdapat korelasi signifikan."],
        ["Unit Observasi", f"Catatan kejadian letupan konflik historis sejak 1990 (N={valid_cases_44})."],
    ]

    mermaid_str_4_4 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Konflik Agraria<br/><i>tahun, status, aktor, kriminalisasi, kekerasan</i>"]
    end
    subgraph Before_After["2. Before-After Analysis"]
        A --> B["Pisah periode<br/>Pra <2014 dan Pasca >=2014"]
        B --> C["Hitung rata-rata kasus/tahun<br/>ditangkap dan tewas"]
    end
    subgraph Crosstab["3. Crosstabulation"]
        A --> D["Bentuk variabel kategorikal X<br/>periode, sektor, pemerintah"]
        A --> E["Bentuk variabel kategorikal Y<br/>represi, penelantaran, kekerasan"]
        D --> F["Uji Chi-Square & Odds Ratio"]
        E --> F
    end
    C --> G["Pembacaan eskalasi konflik"]
    F --> G"""
    mermaid_png_path_4_4 = str(tool_dir / "mermaid_flowchart_4_4.png")
    download_success_4_4 = download_mermaid_png(mermaid_str_4_4, mermaid_png_path_4_4)

    print("[1.95/4] Mengekstraksi dataset empiris Bab 4 sub-bab 4.5...")
    df_nlp_45 = pd.read_csv(data_dir / "sulawesi_konflik_agraria_tanahkita.csv")
    n_kasus_45 = len(df_nlp_45)
    text_corpus_45 = " ".join((df_nlp_45["judul"].fillna("") + " " + df_nlp_45["deskripsi"].fillna("") + " " + df_nlp_45["narasi"].fillna("")).tolist())

    pattern_corp_45 = r'\b(?:PT|CV)\.?\s*[A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*){0,3}\b'
    pattern_civil_45 = r'\b(?:Preman|Ormas|Satgas|PAM Swakarsa|Pemuda Pancasila|GRIB|Laskar|Tandingan|Oknum|Security|Satpam|Pengamanan Swakarsa|Centeng|Beking)\b[^\.,;\!\?\(\)\[\]"\'\-]*'

    pts_45 = re.findall(pattern_corp_45, text_corpus_45)
    pts_45 = [" ".join(pt.split()) for pt in pts_45]
    pts_45 = [re.sub(r'\bPTPN(?:\s+(?:XIV|XII|VII|II|14|Unit\s*14))?\b', 'PT Perkebunan Nusantara (PTPN)', pt, flags=re.IGNORECASE) for pt in pts_45]
    df_aktor_perusahaan_45 = pd.Series(pts_45).value_counts().reset_index()
    df_aktor_perusahaan_45.columns = ["Aktor", "Frekuensi"]

    civils_raw_45 = re.findall(pattern_civil_45, text_corpus_45, flags=re.IGNORECASE)
    stopwords_45 = {"yang", "dan", "di", "dari", "dengan", "untuk", "pada", "ke", "dalam", "oleh", "serta", "sebagai", "adalah", "ini", "itu", "tersebut", "kepada", "saat", "ketika", "juga", "mengatasnamakan", "berjumlah", "melarang", "datang", "berupaya", "segera", "salah", "lainnya", "tak", "nya", "sedang", "akan", "karena", "sebab", "lalu", "kemudian", "mereka"}
    civils_clean_45 = []
    for phrase in civils_raw_45:
        clean_words = []
        for w in phrase.split():
            if w.lower() in stopwords_45:
                break
            clean_words.append(w.title())
        if clean_words:
            civils_clean_45.append(" ".join(clean_words))
    df_aktor_masyarakat_45 = pd.Series(civils_clean_45).value_counts().reset_index()
    df_aktor_masyarakat_45.columns = ["Aktor", "Frekuensi"]

    n_entitas_corp_45 = len(df_aktor_perusahaan_45)
    n_entitas_civ_45 = len(df_aktor_masyarakat_45)
    total_mentions_corp_45 = int(df_aktor_perusahaan_45["Frekuensi"].sum())
    total_mentions_civ_45 = int(df_aktor_masyarakat_45["Frekuensi"].sum())
    top1_corp_name_45 = df_aktor_perusahaan_45.iloc[0]["Aktor"] if not df_aktor_perusahaan_45.empty else "Korporasi"
    top1_corp_freq_45 = int(df_aktor_perusahaan_45.iloc[0]["Frekuensi"]) if not df_aktor_perusahaan_45.empty else 0
    top1_civ_name_45 = df_aktor_masyarakat_45.iloc[0]["Aktor"] if not df_aktor_masyarakat_45.empty else "Preman/Ormas"
    top1_civ_freq_45 = int(df_aktor_masyarakat_45.iloc[0]["Frekuensi"]) if not df_aktor_masyarakat_45.empty else 0

    corp_rows_45 = [[row["Aktor"], f"{int(row['Frekuensi']):,}"] for _, row in df_aktor_perusahaan_45.head(10).iterrows()]
    civil_rows_45 = [[row["Aktor"], f"{int(row['Frekuensi']):,}"] for _, row in df_aktor_masyarakat_45.head(10).iterrows()]

    mermaid_str_4_5 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Repositori Kasus TanahKita<br/><i>judul, deskripsi, narasi (free-text)</i>"] --> B["Penggabungan Korpus Teks<br/>seluruh kasus agraria nasional"]
    end
    subgraph NLP_Processing["2. Ekstraksi Entitas (RegEx NLP)"]
        B --> C["Pattern Matching Korporasi<br/>deteksi entitas PT/CV + normalisasi PTPN"]
        B --> D["Pattern Matching Aktor Proksi<br/>Preman, Ormas, Satgas, dst. + stopword cutoff"]
        C --> E["Token Counting<br/>frekuensi penyebutan per entitas"]
        D --> E
    end
    subgraph Visual_Output["3. Frequency Profiling"]
        E --> F["Dual Horizontal Bar Chart<br/>Top 10 korporasi vs Top 10 aktor proksi"]
        F --> G["Pembacaan orkestrasi konflik & pemetaan oligarki"]
    end"""
    mermaid_png_path_4_5 = str(tool_dir / "mermaid_flowchart_4_5.png")
    download_success_4_5 = download_mermaid_png(mermaid_str_4_5, mermaid_png_path_4_5)

    print("[2/4] Membangun DOCX Metodologi_Bab4_Ruang_Hidup.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)

    add_h1(doc, "BAB IV: METODOLOGI ANALISIS RUANG HIDUP YANG TERAMPAS")
    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada ", False, False),
        ("Bab 4: Ruang Hidup yang Terampas", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.", False, False),
    ])

    add_h2(doc, "4.1 Tren Eskalasi Konflik Agraria Seiring Ekspansi Industri")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Catatan Konflik Agraria: data/processed/sulawesi_konflik_agraria_tanahkita.csv. Visualisasi dashboard menggunakan Analisis Tren Time-Series untuk melacak eskalasi letupan konflik agraria historis berdasarkan tahun pencatatan dan sektor pemicu.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        (f"Ekspansi industri ekstraktif dan proyek strategis berimplikasi pada dinamika sosial dan penggunaan lahan masyarakat. Data empiris mencatat akumulasi {total_konflik:,} kasus konflik agraria, dengan estimasi {total_jiwa:,} jiwa terdampak dan {status_belum_selesai:,} kasus berstatus belum ditangani. ", False, False),
        (f"Tiga sektor utama (Perkebunan, Kehutanan, dan Pertambangan) menyumbang porsi {rasio_ekstraktif:.1f}% dari keseluruhan catatan konflik regional.", False, False),
    ])
    add_p(doc, [
        (f"Pada periode pra-2005, sistem pendataan mencatat {pra_2005:,} kasus konflik agraria. Pada periode pasca-2005 hingga saat ini, data mencatat {pasca_2005:,} kasus konflik lahan, atau setara peningkatan sebesar {lonjakan:,.1f}% dibandingkan periode sebelumnya.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Tren Time-Series")
    add_p(doc, [
        ("Kerangka analisis tren runtun waktu untuk melacak eskalasi kasus perampasan lahan secara historis diilustrasikan pada ", False, False),
        ("Bagan Alur 4.1", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan agregasi time-series dan komparasi periodik.", False, False),
    ])
    add_caption(doc, "Bagan Alur 4.1: Alur Logika Metodologis Time-Series Trend Analysis Konflik Agraria")
    if download_success_4_1:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_4_1, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 4.1 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Agregasi Konflik Tahunan dan Lonjakan Eskalasi")
    add_p(doc, [("Agregasi konflik berdasarkan tahun pencatatan dan sektor industri dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Total Konflik Tahunan per Sektor", "K_{t,s} = Σ c_i, untuk setiap kasus i pada tahun t dan sektor s", [
        ("K_{t,s}", "Total letupan konflik pada tahun t dan sektor pemicu s."),
        ("c_i", "Indikator kasus konflik ke-i; bernilai 1 jika kasus masuk tahun t dan sektor s."),
        ("t", "Tahun pencatatan konflik."),
        ("s", "Sektor pemicu konflik: Perkebunan, Kehutanan, Pertambangan, Infrastruktur & PSN, Pariwisata & Pesisir, atau Lainnya."),
    ])
    add_formula(doc, "Lonjakan Eskalasi Pasca-2005", "E (%) = ( K_Pasca / K_Pra ) × 100", [
        ("E (%)", "Rasio eskalasi konflik pasca-2005 terhadap periode pra-2005."),
        ("K_Pasca", f"Jumlah konflik periode pasca-2005 ({pasca_2005:,} kasus)."),
        ("K_Pra", f"Jumlah konflik periode pra-2005 ({pra_2005:,} kasus)."),
    ])
    add_formula(doc, "Substitusi Lonjakan Eskalasi", f"E = ({pasca_2005:,} / {pra_2005:,}) × 100 = {lonjakan:,.1f}%")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Tren Tahunan dan Sektor Pemicu Konflik")
    add_p(doc, [("Distribusi jumlah konflik agraria modern sejak 1990 disajikan pada Tabel 4.1 berikut:", False, False)])
    add_caption(doc, "Tabel 4.1: Tren Tahunan Konflik Agraria Regional (1990-2025)")
    add_table_1col(doc, ["Tahun", "Jumlah Konflik"], annual_rows, [3.0, 3.0], ["C", "C"])

    add_caption(doc, "Tabel 4.2: Distribusi Konflik Agraria menurut Sektor Pemicu")
    add_table_1col(doc, ["Sektor Pemicu", "Jumlah Konflik", "Proporsi"], sector_rows, [5.5, 3.0, 3.0], ["L", "C", "C"])

    add_caption(doc, f"Tabel 4.3: Komposisi Sektor pada Tahun Puncak Insidensi ({peak_year})")
    add_table_1col(doc, ["Tahun", "Sektor Pemicu", "Jumlah Konflik"], peak_sector_rows, [2.5, 5.5, 3.0], ["C", "L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Puncak Insidensi dan Eskalasi Konflik")
    add_p(doc, [
        (f"Grafik time-series pada dashboard memperlihatkan peningkatan insidensi konflik yang memuncak pada tahun {peak_year} dengan {peak_value:,} kasus konflik. Pembedahan data sektoral menunjukkan konsentrasi konflik pada sektor-sektor berbasis penguasaan ruang. ", False, False),
        ("Peningkatan insidensi konflik beririsan dengan dinamika perizinan kawasan, sehingga pengelolaan alokasi ruang dan perlindungan hak masyarakat di wilayah investasi menjadi faktor penting untuk meminimalkan dampak sosial.", False, False),
    ])

    add_h2(doc, "4.2 Sebaran Sektoral: Dampak Masyarakat dan Penggunaan Lahan")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Dampak Konflik: data/processed/sulawesi_konflik_agraria_tanahkita.csv. Visualisasi dashboard menggunakan Analisis Komparatif Dampak Sosial-Ekologis untuk membedah skala korban terdampak (jiwa) dan luas area konflik (hektar) antar sektor.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Visualisasi komparatif menggambarkan skala dampak sosial dan penggunaan lahan berdasarkan sektor industri. ", False, False),
        (f"Data menunjukkan bahwa Sektor Pertambangan mencatatkan jumlah warga terdampak terbesar, yaitu {jiwa_tambang:,.0f} jiwa, disusul sektor Kehutanan sebanyak {jiwa_kehutanan:,.0f} jiwa. ", False, False),
        ("Pola ini memperlihatkan bahwa konflik agraria tidak hanya terkait jumlah kejadian, tetapi juga skala sosial warga yang terdampak.", False, False),
    ])
    add_p(doc, [
        (f"Dari dimensi penggunaan lahan, Sektor Pertambangan juga mencatatkan luas sengketa terbesar yaitu {ha_tambang:,.0f} hektar, disusul Perkebunan seluas {ha_kebun:,.0f} ha dan Kehutanan seluas {ha_kehutanan:,.0f} ha.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Komparatif Dampak Sosial-Ekologis")
    add_p(doc, [
        ("Kerangka agregasi komparatif untuk membedah skala kehancuran sosial dan monopoli ruang antar sektor diilustrasikan pada ", False, False),
        ("Bagan Alur 4.2", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Sectoral Burden Analysis dan bedah anomali lonjakan.", False, False),
    ])
    add_caption(doc, "Bagan Alur 4.2: Alur Logika Metodologis Analisis Komparatif Dampak Sosial-Ekologis")
    if download_success_4_2:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_4_2, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 4.2 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Agregasi Jiwa Terdampak dan Monopoli Area")
    add_p(doc, [("Agregasi korban terdampak dan luas area konflik per sektor dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Total Jiwa Terdampak per Sektor", "J_s = Σ J_i, untuk setiap kasus i pada sektor s", [
        ("J_s", "Total warga terdampak pada sektor konflik s."),
        ("J_i", "Jumlah masyarakat terdampak pada kasus konflik ke-i."),
        ("s", "Sektor pemicu konflik."),
    ])
    add_formula(doc, "Total Monopoli Area Konflik per Sektor", "A_s = Σ A_i, untuk setiap kasus i pada sektor s", [
        ("A_s", "Total luas area konflik pada sektor s (hektar)."),
        ("A_i", "Luas area konflik pada kasus ke-i (hektar)."),
    ])
    add_formula(doc, "Proporsi Dampak Sektoral", "P_s (%) = ( Nilai_s / Nilai_Total ) × 100", [
        ("P_s (%)", "Persentase kontribusi sektor s terhadap total jiwa terdampak atau total luas area konflik."),
        ("Nilai_s", "Total jiwa atau total hektar pada sektor s."),
        ("Nilai_Total", "Total jiwa atau total hektar seluruh sektor yang dianalisis."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Dampak Sosial dan Monopoli Ruang")
    add_p(doc, [("Agregasi dampak masyarakat dan luas area konflik menurut sektor disajikan pada Tabel 4.4 berikut:", False, False)])
    add_caption(doc, "Tabel 4.4: Matriks Dampak Sosial-Ekologis Konflik Agraria menurut Sektor")
    add_table_1col(doc, ["Sektor Pemicu", "Jiwa Terdampak", "Proporsi Jiwa", "Luas Area (Ha)", "Proporsi Area"], sektor_dampak_rows, [3.8, 2.7, 2.2, 2.7, 2.2], ["L", "C", "C", "C", "C"])

    add_caption(doc, "Tabel 4.5: Bedah Anomali Lonjakan Korban Terdampak (Jiwa)")
    add_table_1col(doc, ["Anomali", "Tahun", "Sektor", "Korban Jiwa", "Kasus Utama"], anomaly_jiwa_rows, [2.3, 1.5, 2.5, 2.3, 6.0], ["L", "C", "L", "C", "L"])

    add_caption(doc, "Tabel 4.6: Bedah Anomali Monopoli Area Konflik (Hektar)")
    add_table_1col(doc, ["Anomali", "Tahun", "Sektor", "Luas Ha", "Kasus Utama"], anomaly_ha_rows, [2.3, 1.5, 2.5, 2.3, 6.0], ["L", "C", "L", "C", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris: Asimetri Dampak Sosial dan Penguasaan Ruang")
    add_p(doc, [
        (f"Matriks dampak sektoral menunjukkan bahwa sektor Pertambangan menjadi penyumbang utama korban terdampak sekaligus sektor dengan luas sengketa terbesar. Dengan demikian, dinamika konflik tidak hanya perlu dibaca dari jumlah kasus, tetapi juga dari skala korban dan luas ruang hidup yang diperebutkan. ", False, False),
        ("Bedah anomali tahunan memperlihatkan tahun-tahun tertentu sebagai titik lonjakan ekstrem yang mendorong grafik korban dan area konflik.", False, False),
    ])

    add_h2(doc, "4.3 Indikasi Represi dan Kriminalisasi dalam Konflik Agraria")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Represi dan Kriminalisasi: data/processed/sulawesi_konflik_agraria_tanahkita.csv. Visualisasi dashboard menggunakan Analisis Agregat Kasus Represi & Pelanggaran HAM untuk menghitung indikasi kriminalisasi, korban ditangkap, luka-luka, dan tewas.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        (f"Data kuantitatif di wilayah Sulawesi mencatat indikasi terjadinya represi dan tindakan kriminalisasi dalam sebagian sengketa agraria. Dari database yang didokumentasikan, terdapat {total_kriminalisasi:,} kasus indikasi kriminalisasi dan {total_ditangkap:,} warga/aktivis lingkungan yang tercatat pernah ditangkap dalam penanganan sengketa lahan. ", False, False),
        (f"Berdasarkan distribusi sektoral, Sektor {top_sektor} mencatatkan frekuensi indikasi represi tertinggi dengan {top_sektor_count:,} kasus. Tahun dengan jumlah catatan insiden represi tertinggi adalah {top_tahun} dengan {top_tahun_count:,} kasus.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Agregat Kasus Represi & Pelanggaran HAM")
    add_p(doc, [
        ("Kerangka pemodelan indikator kekerasan, kriminalisasi, dan penyempitan ruang sipil diilustrasikan pada ", False, False),
        ("Bagan Alur 4.3", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan agregasi kasus represi dan fatalitas.", False, False),
    ])
    add_caption(doc, "Bagan Alur 4.3: Alur Logika Metodologis Analisis Represi dan Kriminalisasi")
    if download_success_4_3:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_4_3, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 4.3 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Kriminalisasi dan Korban Represi")
    add_p(doc, [("Jumlah insiden kriminalisasi dan korban represi fisik dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Total Kasus Kriminalisasi", "K_krim = Σ I_i, untuk setiap kasus i dengan indikasi kriminalisasi", [
        ("K_krim", "Total kasus dengan indikasi kriminalisasi."),
        ("I_i", "Indikator kasus ke-i; bernilai 1 jika indikasi kriminalisasi = benar, dan 0 jika tidak."),
    ])
    add_formula(doc, "Total Korban Represi Fisik", "R = Σ ( D_i + L_i + T_i )", [
        ("R", "Total korban represi fisik yang terdokumentasi."),
        ("D_i", "Jumlah warga/aktivis ditangkap pada kasus i."),
        ("L_i", "Jumlah korban luka pada kasus i."),
        ("T_i", "Jumlah korban tewas pada kasus i."),
    ])
    add_formula(doc, "Substitusi Korban Represi", f"R = {total_ditangkap:,} + {total_luka:,} + {total_tewas:,} = {total_ditangkap + total_luka + total_tewas:,} orang")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Tren Represi dan Arsip Kekerasan")
    add_p(doc, [("Metrik agregat represi dan kriminalisasi disajikan pada Tabel 4.7 berikut:", False, False)])
    add_caption(doc, "Tabel 4.7: Metrik Agregat Represi dan Kriminalisasi")
    add_table_1col(doc, ["Indikator", "Nilai"], [["Kasus Indikasi Kriminalisasi", f"{total_kriminalisasi:,}"], ["Warga/Aktivis Ditangkap", f"{total_ditangkap:,}"], ["Korban Luka-luka", f"{total_luka:,}"], ["Korban Tewas", f"{total_tewas:,}"]], [6.0, 3.0], ["L", "C"])

    add_caption(doc, "Tabel 4.8: Tren Kasus Kriminalisasi dan Represi Pasca-2000")
    add_table_1col(doc, ["Tahun", "Jumlah Kasus"], krim_tahun_rows, [3.0, 3.0], ["C", "C"])

    add_caption(doc, "Tabel 4.9: Sektor Industri Paling Represif")
    add_table_1col(doc, ["Sektor Pemicu", "Jumlah Kasus Kriminalisasi"], krim_sektor_rows, [5.0, 4.0], ["L", "C"])

    add_caption(doc, "Tabel 4.10: Arsip Kasus Represi dan Kekerasan Fisik Tertinggi")
    add_table_1col(doc, ["Tahun", "Sektor", "Perusahaan Terlibat", "Ditangkap", "Tewas", "Narasi Singkat"], kekerasan_rows, [1.4, 2.0, 3.0, 1.6, 1.4, 6.0], ["C", "L", "L", "C", "C", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris: Penyempitan Ruang Sipil dan Risiko HAM")
    add_p(doc, [
        (f"Keberadaan kasus kriminalisasi di sekitar area konsesi, terutama pada sektor {top_sektor}, mengindikasikan pentingnya jaminan perlindungan ruang sipil dan penghormatan HAM dalam setiap proses pembangunan. ", False, False),
        ("Catatan ini menunjukkan perlunya pendekatan hukum yang adil, penyelesaian konflik secara ramah HAM, serta perlindungan bagi pejuang lingkungan dan komunitas lokal.", False, False),
    ])

    add_h2(doc, "4.4 Pembuktian Statistik: Ekspansi vs Eskalasi Konflik")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Base Data Cross-Section: data/processed/sulawesi_konflik_agraria_tanahkita.csv. Visualisasi dashboard menggunakan Before-After Analysis & Crosstabulation untuk menguji hubungan antara indikator ekspansi dan eskalasi konflik.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Hipotesis utama dalam evaluasi ini adalah bahwa industrialisasi dan ekspansi korporasi berbanding lurus dengan eskalasi konflik dan represi terhadap masyarakat. ", False, False),
        ("Analisis dibagi menjadi dua bagian: komparasi metrik Before-After dan uji signifikansi Crosstab Chi-Square. Unit observasinya adalah catatan kejadian letupan konflik historis.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Before-After Analysis & Crosstabulation")
    add_p(doc, [
        ("Kerangka komparasi pra/pasca ekspansi serta tabulasi silang indikator konflik diilustrasikan pada ", False, False),
        ("Bagan Alur 4.4", True, False),
        (" berikut. Konfigurasi variabel uji Chi-Square dirinci pada Tabel 4.4a di bawah gambar.", False, False),
    ])
    add_caption(doc, "Bagan Alur 4.4: Alur Logika Metodologis Before-After Analysis & Crosstabulation")
    if download_success_4_4:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_4_4, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 4.4 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_caption(doc, "Tabel 4.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 4.4)")
    add_table_1col(doc, konf_headers_44, konf_rows_44, [4.5, 11.0], ["L", "L"])

    add_h4(doc, "C. Formulasi Matematis: Before-After, Chi-Square, dan Odds Ratio")
    add_p(doc, [("Komparasi periode dan pengujian statistik dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Rata-rata Konflik per Tahun", "K̄_p = N_p / T_p", [
        ("K̄_p", "Rata-rata kasus konflik per tahun pada periode p."),
        ("N_p", "Total letupan konflik pada periode p."),
        ("T_p", "Jumlah tahun observasi pada periode p."),
    ])
    add_formula(doc, "Substitusi Before-After", f"K̄_Pra = {len(df_pra_44):,} / {tahun_pra_44} = {avg_pra_44:.1f}; K̄_Pasca = {len(df_pasca_44):,} / {tahun_pasca_44} = {avg_pasca_44:.1f}")
    add_formula(doc, "Uji Independensi Chi-Square Pearson", "χ² = Σ [ ( O_ij - E_ij )² / E_ij ]", [
        ("O_ij", "Frekuensi observasi pada sel baris i dan kolom j."),
        ("E_ij", "Frekuensi harapan jika variabel X dan Y saling independen."),
    ])
    add_formula(doc, "Odds Ratio", "OR = ( a × d ) / ( b × c )", [
        ("a,b,c,d", "Empat sel pada tabel kontinjensi 2x2."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Before-After dan Skenario Crosstab")
    add_caption(doc, "Tabel 4.11: Analisis Komparatif Before-After Pra vs Era Hilirisasi")
    add_table_1col(doc, ["Periode", "Total Konflik", "Jumlah Tahun", "Kasus/Tahun", "Ditangkap", "Tewas"], before_after_rows_44, [3.2, 2.3, 2.0, 2.2, 2.0, 1.8], ["L", "C", "C", "C", "C", "C"])

    add_caption(doc, "Tabel 4.12: Ringkasan Eksekutif Seluruh Skenario Crosstab Ekspansi vs Eskalasi Konflik")
    add_table_1col(doc, ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_44, [3.0, 3.3, 1.8, 1.8, 1.8, 2.3], ["L", "L", "C", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Validitas Statistik Eskalasi Konflik")
    if sig_count_44 > 0:
        finding_44 = f"Dari {total_scenarios_44} skenario pengujian, terdapat {sig_count_44} skenario yang terbukti SIGNIFIKAN. Tingginya Odds Ratio pada skenario yang signifikan menegaskan bahwa ekspansi operasi industri berasosiasi dengan peningkatan risiko sengketa lahan. Skenario yang tidak signifikan mengindikasikan bahwa dinamika sengketa lahan tersebar secara merata di berbagai sektor dan kurun waktu."
    else:
        finding_44 = f"Dari {total_scenarios_44} skenario pengujian, seluruhnya menunjukkan status TIDAK SIGNIFIKAN. Hal ini menunjukkan bahwa sengketa lahan dan tantangan penyelesaiannya terdistribusi secara konsisten di sepanjang waktu dan sektor."
    add_p(doc, [(finding_44, False, False)])

    add_h2(doc, "4.5 Peta Entitas Aktor: Korporasi dan Organisasi Masyarakat")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Teks Bebas (Free-Text): data/processed/sulawesi_konflik_agraria_tanahkita.csv. Visualisasi dashboard menampilkan dua Horizontal Bar Chart berdampingan (Top 10 Entitas Korporasi Paling Dominan dan Top Aktor Proksi & Vigilante Terdeteksi) hasil ekstraksi teks berbasis NLP Regex dari korpus narasi seluruh kasus agraria.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Analisis entitas aktor berbasis pemrosesan teks (string parsing) terhadap catatan kronologi dokumentasi TanahKita memetakan keterlibatan berbagai pihak dalam sengketa agraria. Hasil ekstraksi teks mengidentifikasi entitas korporasi, lembaga pemerintah, serta organisasi masyarakat sipil yang tercatat dalam dokumentasi kasus. ", False, False),
        (f"Korpus dibangun dari penggabungan kolom judul, deskripsi, dan narasi pada {n_kasus_45:,} kasus agraria (nasional) untuk memetakan orkestrasi struktural dan modus operandi aktor secara utuh, termasuk memvalidasi indikasi konsentrasi kekuasaan dan monopoli penguasaan ruang oleh segelintir konglomerasi besar melalui seberapa sering nama entitas muncul dalam sengketa tanah.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Frequency Profiling (Text Parsing NLP)")
    add_p(doc, [
        ("Kerangka ekstraksi entitas berbasis Regular Expressions (RegEx) dan penghitungan frekuensi penyebutan diilustrasikan pada ", False, False),
        ("Bagan Alur 4.5", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Frequency Profiling deskriptif atas kemunculan entitas dalam korpus teks.", False, False),
    ])
    add_caption(doc, "Bagan Alur 4.5: Alur Logika Analisis Frequency Profiling Entitas Aktor")
    if download_success_4_5:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_4_5, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 4.5 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Konstruksi Korpus dan Token Counting")
    add_p(doc, [("Kuantifikasi frekuensi penyebutan entitas dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Konstruksi Korpus Teks", "Korpus = Gabungan ( judul_k , deskripsi_k , narasi_k )   ;   untuk k = 1 s.d. N", [
        ("Korpus", "Teks gabungan seluruh dokumentasi kasus sebagai bahan ekstraksi entitas (Variabel Independen)."),
        ("judul_k, deskripsi_k, narasi_k", "Kolom teks bebas pada kasus ke-k dalam repositori TanahKita."),
        ("N", f"Jumlah kasus agraria dalam repositori ({n_kasus_45:,} kasus)."),
    ])
    add_formula(doc, "Persamaan Token Counting Frekuensi Entitas", "Frekuensi_a = Σ ( Match_i,a )   ;   untuk seluruh kemunculan pola entitas a dalam Korpus", [
        ("Frekuensi_a", "Jumlah absolut penyebutan (mentions) entitas a dalam korpus (Variabel Dependen)."),
        ("Match_i,a", "Kemunculan ke-i dari pola RegEx entitas a; pola korporasi mendeteksi awalan PT/CV diikuti nama kapital (maksimum 4 kata), dengan normalisasi varian PTPN menjadi satu entitas."),
        ("Pola Aktor Proksi", "Deteksi kata kunci Preman, Ormas, Satgas, PAM Swakarsa, Pemuda Pancasila, GRIB, Laskar, Oknum, Security, Satpam, Centeng, Beking beserta frasa lanjutannya, dipotong pada stopword pertama."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Frekuensi Entitas Korporasi dan Aktor Proksi")
    add_p(doc, [
        (f"Sepuluh entitas korporasi paling dominan (dari total {n_entitas_corp_45:,} entitas terdeteksi dengan {total_mentions_corp_45:,} penyebutan) disajikan pada ", False, False),
        ("Tabel 4.13", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 4.13: Top 10 Entitas Korporasi Paling Dominan dalam Dokumentasi Konflik")
    add_table_1col(doc, ["Entitas Korporasi", "Frekuensi Penyebutan"], corp_rows_45, [9.5, 6.0], ["L", "C"])

    add_p(doc, [
        (f"Sepuluh aktor proksi dan vigilante paling sering terdeteksi (dari total {n_entitas_civ_45:,} varian frasa dengan {total_mentions_civ_45:,} penyebutan) disajikan pada ", False, False),
        ("Tabel 4.14", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 4.14: Top Aktor Proksi & Vigilante Terdeteksi dalam Dokumentasi Konflik")
    add_table_1col(doc, ["Aktor Proksi / Vigilante", "Frekuensi Penyebutan"], civil_rows_45, [9.5, 6.0], ["L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Orkestrasi Konflik dan Pemetaan Oligarki")
    add_p(doc, [
        ("1. ", True, False), ("Dominasi Entitas Korporasi: ", True, False),
        (f"Ekstraksi teks mencatat frekuensi penyebutan entitas {top1_corp_name_45} tertinggi dengan {top1_corp_freq_45:,} catatan kasus terpisah, memvalidasi indikasi konsentrasi kekuasaan dan monopoli penguasaan ruang oleh segelintir konglomerasi besar.\n", False, False),
        ("2. ", True, False), ("Orkestrasi Horizontal Aktor Proksi: ", True, False),
        (f"Kemunculan kelompok sipil seperti {top1_civ_name_45} (terdeteksi hingga {top1_civ_freq_45:,} kali) menangkap besarnya skala orkestrasi horizontal. Korporasi seringkali menggunakan jasa pengamanan swakarsa, kelompok preman, hingga ormas vigilante sebagai 'bemper proksi' untuk mengintimidasi warga lokal dan memecah belah solidaritas akar rumput.", False, False),
    ])

    docx_path = tool_dir / "Metodologi_Bab4_Ruang_Hidup.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 4...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 4 - Ruang Hidup yang Terampas</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: Arial, sans-serif; max-width: 920px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #43A047; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; }}
.hdr-title {{ color: #81C784; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #2E7D32; padding-bottom: 8px; }}
h2 {{ color: #81C784; text-transform: uppercase; border-bottom: 1px solid #2E7D32; }}
h4 {{ color: #A5D6A7; }}
.note-box {{ background: #132213; border-left: 4px solid #2E7D32; padding: 10px 14px; margin: 12px 0; }}
.formula {{ background: #0D1B0E; border: 1px solid #2E7D32; color: #A5D6A7; padding: 8px 12px; font-family: monospace; }}
.data-th {{ background: #1B5E20; color: white; padding: 6px; text-align: left; border: 1px solid #2E7D32; }}
.data-td {{ padding: 6px; border: 1px solid #243524; vertical-align: top; }}
.data-tr-even .data-td {{ background: #131B13; }}
.mermaid {{ background: #0D1610; border: 1px solid #2E7D32; padding: 12px; margin: 10px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB IV: Metodologi Analisis Ruang Hidup yang Terampas</div>
<p>Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada <strong>Bab 4: Ruang Hidup yang Terampas</strong>.</p>
<h2>4.1 Tren Eskalasi Konflik Agraria Seiring Ekspansi Industri</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Catatan Konflik Agraria: <code>data/processed/sulawesi_konflik_agraria_tanahkita.csv</code>. Visualisasi dashboard menggunakan Analisis Tren Time-Series.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Ekspansi industri ekstraktif dan proyek strategis berimplikasi pada dinamika sosial dan penggunaan lahan masyarakat. Data empiris mencatat akumulasi <strong>{total_konflik:,} kasus konflik agraria</strong>, dengan estimasi <strong>{total_jiwa:,} jiwa terdampak</strong>. Tiga sektor utama (Perkebunan, Kehutanan, dan Pertambangan) menyumbang <strong>{rasio_ekstraktif:.1f}%</strong> dari keseluruhan catatan konflik regional.</p>
<h4>B. Alur Logika Metodologis Analisis Tren Time-Series</h4>
<div class="mermaid">{mermaid_str_4_1}</div>
<h4>C. Formulasi Matematis</h4>
<div class="formula">K_{{t,s}} = Σ c_i, untuk setiap kasus i pada tahun t dan sektor s</div>
<div class="formula">E (%) = ( K_Pasca / K_Pra ) × 100</div>
<div class="formula">E = ({pasca_2005:,} / {pra_2005:,}) × 100 = {lonjakan:,.1f}%</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 4.1: Tren Tahunan Konflik Agraria Regional (1990-2025)</div>
{html_table(["Tahun", "Jumlah Konflik"], annual_rows)}
<div class="table-caption">Tabel 4.2: Distribusi Konflik Agraria menurut Sektor Pemicu</div>
{html_table(["Sektor Pemicu", "Jumlah Konflik", "Proporsi"], sector_rows)}
<div class="table-caption">Tabel 4.3: Komposisi Sektor pada Tahun Puncak Insidensi ({peak_year})</div>
{html_table(["Tahun", "Sektor Pemicu", "Jumlah Konflik"], peak_sector_rows)}
<h4>E. Analisis Temuan Empiris</h4>
<p>Grafik time-series pada dashboard memperlihatkan peningkatan insidensi konflik yang memuncak pada tahun <strong>{peak_year}</strong> dengan <strong>{peak_value:,} kasus konflik</strong>. Peningkatan insidensi konflik beririsan dengan dinamika perizinan kawasan.</p>
<h2>4.2 Sebaran Sektoral: Dampak Masyarakat dan Penggunaan Lahan</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Dampak Konflik: <code>data/processed/sulawesi_konflik_agraria_tanahkita.csv</code>. Visualisasi dashboard menggunakan Analisis Komparatif Dampak Sosial-Ekologis.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Visualisasi komparatif menggambarkan skala dampak sosial dan penggunaan lahan berdasarkan sektor industri. Sektor Pertambangan mencatatkan jumlah warga terdampak terbesar, yaitu <strong>{jiwa_tambang:,.0f} jiwa</strong>, disusul sektor Kehutanan sebanyak <strong>{jiwa_kehutanan:,.0f} jiwa</strong>. Dari dimensi penggunaan lahan, Sektor Pertambangan juga mencatatkan luas sengketa terbesar yaitu <strong>{ha_tambang:,.0f} hektar</strong>.</p>
<h4>B. Alur Logika Metodologis Analisis Komparatif Dampak Sosial-Ekologis</h4>
<div class="mermaid">{mermaid_str_4_2}</div>
<h4>C. Formulasi Matematis</h4>
<div class="formula">J_s = Σ J_i, untuk setiap kasus i pada sektor s</div>
<div class="formula">A_s = Σ A_i, untuk setiap kasus i pada sektor s</div>
<div class="formula">P_s (%) = ( Nilai_s / Nilai_Total ) × 100</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 4.4: Matriks Dampak Sosial-Ekologis Konflik Agraria menurut Sektor</div>
{html_table(["Sektor Pemicu", "Jiwa Terdampak", "Proporsi Jiwa", "Luas Area (Ha)", "Proporsi Area"], sektor_dampak_rows)}
<div class="table-caption">Tabel 4.5: Bedah Anomali Lonjakan Korban Terdampak (Jiwa)</div>
{html_table(["Anomali", "Tahun", "Sektor", "Korban Jiwa", "Kasus Utama"], anomaly_jiwa_rows)}
<div class="table-caption">Tabel 4.6: Bedah Anomali Monopoli Area Konflik (Hektar)</div>
{html_table(["Anomali", "Tahun", "Sektor", "Luas Ha", "Kasus Utama"], anomaly_ha_rows)}
<h4>E. Analisis Temuan Empiris</h4>
<p>Matriks dampak sektoral menunjukkan bahwa sektor Pertambangan menjadi penyumbang utama korban terdampak sekaligus sektor dengan luas sengketa terbesar. Dinamika konflik perlu dibaca dari jumlah kasus, skala korban, dan luas ruang hidup yang diperebutkan.</p>
<h2>4.3 Indikasi Represi dan Kriminalisasi dalam Konflik Agraria</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Represi dan Kriminalisasi: <code>data/processed/sulawesi_konflik_agraria_tanahkita.csv</code>. Visualisasi dashboard menggunakan Analisis Agregat Kasus Represi & Pelanggaran HAM.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Dari database yang didokumentasikan, terdapat <strong>{total_kriminalisasi:,} kasus indikasi kriminalisasi</strong> dan <strong>{total_ditangkap:,} warga/aktivis lingkungan</strong> yang tercatat pernah ditangkap. Berdasarkan distribusi sektoral, <strong>Sektor {top_sektor}</strong> mencatatkan frekuensi indikasi represi tertinggi dengan <strong>{top_sektor_count:,} kasus</strong>. Tahun dengan jumlah catatan insiden represi tertinggi adalah <strong>{top_tahun}</strong> dengan <strong>{top_tahun_count:,} kasus</strong>.</p>
<h4>B. Alur Logika Metodologis Analisis Agregat Kasus Represi & Pelanggaran HAM</h4>
<div class="mermaid">{mermaid_str_4_3}</div>
<h4>C. Formulasi Matematis</h4>
<div class="formula">K_krim = Σ I_i, untuk setiap kasus i dengan indikasi kriminalisasi</div>
<div class="formula">R = Σ ( D_i + L_i + T_i )</div>
<div class="formula">R = {total_ditangkap:,} + {total_luka:,} + {total_tewas:,} = {total_ditangkap + total_luka + total_tewas:,} orang</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 4.7: Metrik Agregat Represi dan Kriminalisasi</div>
{html_table(["Indikator", "Nilai"], [["Kasus Indikasi Kriminalisasi", f"{total_kriminalisasi:,}"], ["Warga/Aktivis Ditangkap", f"{total_ditangkap:,}"], ["Korban Luka-luka", f"{total_luka:,}"], ["Korban Tewas", f"{total_tewas:,}"]])}
<div class="table-caption">Tabel 4.8: Tren Kasus Kriminalisasi dan Represi Pasca-2000</div>
{html_table(["Tahun", "Jumlah Kasus"], krim_tahun_rows)}
<div class="table-caption">Tabel 4.9: Sektor Industri Paling Represif</div>
{html_table(["Sektor Pemicu", "Jumlah Kasus Kriminalisasi"], krim_sektor_rows)}
<div class="table-caption">Tabel 4.10: Arsip Kasus Represi dan Kekerasan Fisik Tertinggi</div>
{html_table(["Tahun", "Sektor", "Perusahaan Terlibat", "Ditangkap", "Tewas", "Narasi Singkat"], kekerasan_rows)}
<h4>E. Analisis Temuan Empiris</h4>
<p>Keberadaan kasus kriminalisasi di sekitar area konsesi, terutama pada sektor <strong>{top_sektor}</strong>, mengindikasikan pentingnya jaminan perlindungan ruang sipil dan penghormatan HAM dalam setiap proses pembangunan.</p>
<h2>4.4 Pembuktian Statistik: Ekspansi vs Eskalasi Konflik</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Base Data Cross-Section: <code>data/processed/sulawesi_konflik_agraria_tanahkita.csv</code>. Visualisasi dashboard menggunakan Before-After Analysis & Crosstabulation.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Hipotesis utama dalam evaluasi ini adalah bahwa industrialisasi dan ekspansi korporasi berbanding lurus dengan eskalasi konflik dan represi terhadap masyarakat. Analisis dibagi menjadi komparasi metrik Before-After dan uji signifikansi Crosstab Chi-Square.</p>
<h4>B. Alur Logika Metodologis Before-After Analysis & Crosstabulation</h4>
<div class="mermaid">{mermaid_str_4_4}</div>
<div class="table-caption">Tabel 4.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 4.4)</div>
{html_table(konf_headers_44, konf_rows_44)}
<h4>C. Formulasi Matematis</h4>
<div class="formula">K̄_p = N_p / T_p</div>
<div class="formula">K̄_Pra = {len(df_pra_44):,} / {tahun_pra_44} = {avg_pra_44:.1f}; K̄_Pasca = {len(df_pasca_44):,} / {tahun_pasca_44} = {avg_pasca_44:.1f}</div>
<div class="formula">χ² = Σ [ ( O_ij - E_ij )² / E_ij ]</div>
<div class="formula">OR = ( a × d ) / ( b × c )</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 4.11: Analisis Komparatif Before-After Pra vs Era Hilirisasi</div>
{html_table(["Periode", "Total Konflik", "Jumlah Tahun", "Kasus/Tahun", "Ditangkap", "Tewas"], before_after_rows_44)}
<div class="table-caption">Tabel 4.12: Ringkasan Eksekutif Seluruh Skenario Crosstab Ekspansi vs Eskalasi Konflik</div>
{html_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_44)}
<h4>E. Analisis Temuan Empiris</h4>
<p>{finding_44}</p>

<h2>4.5 Peta Entitas Aktor: Korporasi dan Organisasi Masyarakat</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Teks Bebas (Free-Text): <code>data/processed/sulawesi_konflik_agraria_tanahkita.csv</code>. Visualisasi dashboard menampilkan dua Horizontal Bar Chart berdampingan (Top 10 Entitas Korporasi Paling Dominan dan Top Aktor Proksi & Vigilante Terdeteksi) hasil ekstraksi teks berbasis NLP Regex dari korpus narasi seluruh kasus agraria.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Analisis entitas aktor berbasis pemrosesan teks (string parsing) terhadap catatan kronologi dokumentasi TanahKita memetakan keterlibatan berbagai pihak dalam sengketa agraria. Korpus dibangun dari penggabungan kolom <code>judul</code>, <code>deskripsi</code>, dan <code>narasi</code> pada <strong>{n_kasus_45:,} kasus agraria</strong> (nasional) untuk memetakan orkestrasi struktural dan modus operandi aktor secara utuh, termasuk memvalidasi indikasi konsentrasi kekuasaan oleh segelintir konglomerasi besar.</p>
<h4>B. Alur Logika Metodologis Frequency Profiling (Text Parsing NLP)</h4>
<p>Kerangka ekstraksi entitas berbasis Regular Expressions (RegEx) dan penghitungan frekuensi penyebutan diilustrasikan pada <strong>Bagan Alur 4.5</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Frequency Profiling deskriptif atas kemunculan entitas dalam korpus teks.</p>
<div class="table-caption">Bagan Alur 4.5: Alur Logika Analisis Frequency Profiling Entitas Aktor</div>
<div class="mermaid">{mermaid_str_4_5}</div>
<h4>C. Formulasi Matematis: Konstruksi Korpus dan Token Counting</h4>
<div class="formula">Korpus = Gabungan ( judul_k , deskripsi_k , narasi_k )   ;   untuk k = 1 s.d. N</div>
<div class="formula">Frekuensi_a = Σ ( Match_i,a )   ;   untuk seluruh kemunculan pola entitas a dalam Korpus</div>
<p>Pola korporasi mendeteksi awalan PT/CV diikuti nama kapital (maksimum 4 kata) dengan normalisasi varian PTPN; pola aktor proksi mendeteksi kata kunci Preman, Ormas, Satgas, PAM Swakarsa, Pemuda Pancasila, GRIB, Laskar, Oknum, Security, Satpam, Centeng, Beking beserta frasa lanjutannya (dipotong pada stopword pertama).</p>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 4.13: Top 10 Entitas Korporasi Paling Dominan dalam Dokumentasi Konflik</div>
{html_table(["Entitas Korporasi", "Frekuensi Penyebutan"], corp_rows_45)}
<div class="table-caption">Tabel 4.14: Top Aktor Proksi & Vigilante Terdeteksi dalam Dokumentasi Konflik</div>
{html_table(["Aktor Proksi / Vigilante", "Frekuensi Penyebutan"], civil_rows_45)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Dominasi Entitas Korporasi:</strong> entitas <strong>{top1_corp_name_45}</strong> tercatat tertinggi dengan <strong>{top1_corp_freq_45:,} catatan kasus terpisah</strong> (dari {n_entitas_corp_45:,} entitas dan {total_mentions_corp_45:,} penyebutan), memvalidasi indikasi konsentrasi kekuasaan dan monopoli penguasaan ruang. <strong>2. Orkestrasi Horizontal Aktor Proksi:</strong> kelompok sipil seperti <strong>{top1_civ_name_45}</strong> terdeteksi hingga <strong>{top1_civ_freq_45:,} kali</strong> — korporasi seringkali menggunakan pengamanan swakarsa, kelompok preman, hingga ormas vigilante sebagai "bemper proksi" untuk mengintimidasi warga lokal dan memecah belah solidaritas akar rumput.</p>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab4_Ruang_Hidup.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB IV: METODOLOGI ANALISIS RUANG HIDUP YANG TERAMPAS",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada **Bab 4: Ruang Hidup yang Terampas**.",
        "",
        "## 4.1 Tren Eskalasi Konflik Agraria Seiring Ekspansi Industri",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Catatan Konflik Agraria: `data/processed/sulawesi_konflik_agraria_tanahkita.csv`. Visualisasi dashboard menggunakan *Analisis Tren Time-Series* untuk melacak eskalasi letupan konflik agraria historis berdasarkan tahun pencatatan dan sektor pemicu.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Ekspansi industri ekstraktif dan proyek strategis berimplikasi pada dinamika sosial dan penggunaan lahan masyarakat. Data empiris mencatat akumulasi **{total_konflik:,} kasus konflik agraria**, dengan estimasi **{total_jiwa:,} jiwa terdampak** dan **{status_belum_selesai:,} kasus** berstatus belum ditangani.",
        "",
        f"Pada periode pra-2005, sistem pendataan mencatat **{pra_2005:,} kasus** konflik agraria. Pada periode pasca-2005 hingga saat ini, data mencatat **{pasca_2005:,} kasus** konflik lahan, atau setara peningkatan sebesar **{lonjakan:,.1f}%** dibandingkan periode sebelumnya.",
        "",
        "#### B. Alur Logika Metodologis Analisis Tren Time-Series",
        "```mermaid",
        mermaid_str_4_1,
        "```",
        "",
        "#### C. Formulasi Matematis: Agregasi Konflik Tahunan dan Lonjakan Eskalasi",
        "```text",
        "K_{t,s} = Σ c_i, untuk setiap kasus i pada tahun t dan sektor s",
        "E (%) = ( K_Pasca / K_Pra ) × 100",
        f"E = ({pasca_2005:,} / {pra_2005:,}) × 100 = {lonjakan:,.1f}%",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 4.1: Tren Tahunan Konflik Agraria Regional (1990-2025)",
        markdown_table(["Tahun", "Jumlah Konflik"], annual_rows),
        "",
        "##### Tabel 4.2: Distribusi Konflik Agraria menurut Sektor Pemicu",
        markdown_table(["Sektor Pemicu", "Jumlah Konflik", "Proporsi"], sector_rows),
        "",
        f"##### Tabel 4.3: Komposisi Sektor pada Tahun Puncak Insidensi ({peak_year})",
        markdown_table(["Tahun", "Sektor Pemicu", "Jumlah Konflik"], peak_sector_rows),
        "",
        "#### E. Analisis Temuan Empiris: Puncak Insidensi dan Eskalasi Konflik",
        f"Grafik time-series pada dashboard memperlihatkan peningkatan insidensi konflik yang memuncak pada tahun **{peak_year}** dengan **{peak_value:,} kasus konflik**. Peningkatan insidensi konflik beririsan dengan dinamika perizinan kawasan, sehingga pengelolaan alokasi ruang dan perlindungan hak masyarakat di wilayah investasi menjadi faktor penting untuk meminimalkan dampak sosial.",
        "",
        "## 4.2 Sebaran Sektoral: Dampak Masyarakat dan Penggunaan Lahan",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Dampak Konflik: `data/processed/sulawesi_konflik_agraria_tanahkita.csv`. Visualisasi dashboard menggunakan *Analisis Komparatif Dampak Sosial-Ekologis* untuk membedah skala korban terdampak (jiwa) dan luas area konflik (hektar) antar sektor.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Visualisasi komparatif menggambarkan skala dampak sosial dan penggunaan lahan berdasarkan sektor industri. Data menunjukkan bahwa **Sektor Pertambangan** mencatatkan jumlah warga terdampak terbesar, yaitu **{jiwa_tambang:,.0f} jiwa**, disusul sektor Kehutanan sebanyak **{jiwa_kehutanan:,.0f} jiwa**.",
        "",
        f"Dari dimensi penggunaan lahan, **Sektor Pertambangan** juga mencatatkan luas sengketa terbesar yaitu **{ha_tambang:,.0f} hektar**, disusul Perkebunan seluas **{ha_kebun:,.0f} ha** dan Kehutanan seluas **{ha_kehutanan:,.0f} ha**.",
        "",
        "#### B. Alur Logika Metodologis Analisis Komparatif Dampak Sosial-Ekologis",
        "```mermaid",
        mermaid_str_4_2,
        "```",
        "",
        "#### C. Formulasi Matematis: Agregasi Jiwa Terdampak dan Monopoli Area",
        "```text",
        "J_s = Σ J_i, untuk setiap kasus i pada sektor s",
        "A_s = Σ A_i, untuk setiap kasus i pada sektor s",
        "P_s (%) = ( Nilai_s / Nilai_Total ) × 100",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 4.4: Matriks Dampak Sosial-Ekologis Konflik Agraria menurut Sektor",
        markdown_table(["Sektor Pemicu", "Jiwa Terdampak", "Proporsi Jiwa", "Luas Area (Ha)", "Proporsi Area"], sektor_dampak_rows),
        "",
        "##### Tabel 4.5: Bedah Anomali Lonjakan Korban Terdampak (Jiwa)",
        markdown_table(["Anomali", "Tahun", "Sektor", "Korban Jiwa", "Kasus Utama"], anomaly_jiwa_rows),
        "",
        "##### Tabel 4.6: Bedah Anomali Monopoli Area Konflik (Hektar)",
        markdown_table(["Anomali", "Tahun", "Sektor", "Luas Ha", "Kasus Utama"], anomaly_ha_rows),
        "",
        "#### E. Analisis Temuan Empiris: Asimetri Dampak Sosial dan Penguasaan Ruang",
        "Matriks dampak sektoral menunjukkan bahwa sektor Pertambangan menjadi penyumbang utama korban terdampak sekaligus sektor dengan luas sengketa terbesar. Dengan demikian, dinamika konflik tidak hanya perlu dibaca dari jumlah kasus, tetapi juga dari skala korban dan luas ruang hidup yang diperebutkan.",
        "",
        "## 4.3 Indikasi Represi dan Kriminalisasi dalam Konflik Agraria",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Represi dan Kriminalisasi: `data/processed/sulawesi_konflik_agraria_tanahkita.csv`. Visualisasi dashboard menggunakan *Analisis Agregat Kasus Represi & Pelanggaran HAM* untuk menghitung indikasi kriminalisasi, korban ditangkap, luka-luka, dan tewas.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Data kuantitatif di wilayah Sulawesi mencatat indikasi terjadinya represi dan tindakan kriminalisasi dalam sebagian sengketa agraria. Dari database yang didokumentasikan, terdapat **{total_kriminalisasi:,} kasus indikasi kriminalisasi** dan **{total_ditangkap:,} warga/aktivis lingkungan yang tercatat pernah ditangkap** dalam penanganan sengketa lahan.",
        "",
        f"Berdasarkan distribusi sektoral, **Sektor {top_sektor}** mencatatkan frekuensi indikasi represi tertinggi dengan **{top_sektor_count:,} kasus**. Tahun dengan jumlah catatan insiden represi tertinggi adalah **{top_tahun}** dengan **{top_tahun_count:,} kasus**.",
        "",
        "#### B. Alur Logika Metodologis Analisis Agregat Kasus Represi & Pelanggaran HAM",
        "```mermaid",
        mermaid_str_4_3,
        "```",
        "",
        "#### C. Formulasi Matematis: Kriminalisasi dan Korban Represi",
        "```text",
        "K_krim = Σ I_i, untuk setiap kasus i dengan indikasi kriminalisasi",
        "R = Σ ( D_i + L_i + T_i )",
        f"R = {total_ditangkap:,} + {total_luka:,} + {total_tewas:,} = {total_ditangkap + total_luka + total_tewas:,} orang",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 4.7: Metrik Agregat Represi dan Kriminalisasi",
        markdown_table(["Indikator", "Nilai"], [["Kasus Indikasi Kriminalisasi", f"{total_kriminalisasi:,}"], ["Warga/Aktivis Ditangkap", f"{total_ditangkap:,}"], ["Korban Luka-luka", f"{total_luka:,}"], ["Korban Tewas", f"{total_tewas:,}"]]),
        "",
        "##### Tabel 4.8: Tren Kasus Kriminalisasi dan Represi Pasca-2000",
        markdown_table(["Tahun", "Jumlah Kasus"], krim_tahun_rows),
        "",
        "##### Tabel 4.9: Sektor Industri Paling Represif",
        markdown_table(["Sektor Pemicu", "Jumlah Kasus Kriminalisasi"], krim_sektor_rows),
        "",
        "##### Tabel 4.10: Arsip Kasus Represi dan Kekerasan Fisik Tertinggi",
        markdown_table(["Tahun", "Sektor", "Perusahaan Terlibat", "Ditangkap", "Tewas", "Narasi Singkat"], kekerasan_rows),
        "",
        "#### E. Analisis Temuan Empiris: Penyempitan Ruang Sipil dan Risiko HAM",
        f"Keberadaan kasus kriminalisasi di sekitar area konsesi, terutama pada sektor **{top_sektor}**, mengindikasikan pentingnya jaminan perlindungan ruang sipil dan penghormatan HAM dalam setiap proses pembangunan. Catatan ini menunjukkan perlunya pendekatan hukum yang adil, penyelesaian konflik secara ramah HAM, serta perlindungan bagi pejuang lingkungan dan komunitas lokal.",
        "",
        "## 4.4 Pembuktian Statistik: Ekspansi vs Eskalasi Konflik",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Base Data Cross-Section: `data/processed/sulawesi_konflik_agraria_tanahkita.csv`. Visualisasi dashboard menggunakan *Before-After Analysis & Crosstabulation* untuk menguji hubungan antara indikator ekspansi dan eskalasi konflik.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Hipotesis utama dalam evaluasi ini adalah bahwa **industrialisasi dan ekspansi korporasi** berbanding lurus dengan **eskalasi konflik dan represi** terhadap masyarakat. Analisis dibagi menjadi dua bagian: komparasi metrik Before-After dan uji signifikansi Crosstab Chi-Square. Unit observasinya adalah catatan kejadian letupan konflik historis.",
        "",
        "#### B. Alur Logika Metodologis Before-After Analysis & Crosstabulation",
        "```mermaid",
        mermaid_str_4_4,
        "```",
        "",
        "##### Tabel 4.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 4.4)",
        markdown_table(konf_headers_44, konf_rows_44),
        "",
        "#### C. Formulasi Matematis: Before-After, Chi-Square, dan Odds Ratio",
        "```text",
        "K̄_p = N_p / T_p",
        f"K̄_Pra = {len(df_pra_44):,} / {tahun_pra_44} = {avg_pra_44:.1f}; K̄_Pasca = {len(df_pasca_44):,} / {tahun_pasca_44} = {avg_pasca_44:.1f}",
        "χ² = Σ [ ( O_ij - E_ij )² / E_ij ]",
        "OR = ( a × d ) / ( b × c )",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 4.11: Analisis Komparatif Before-After Pra vs Era Hilirisasi",
        markdown_table(["Periode", "Total Konflik", "Jumlah Tahun", "Kasus/Tahun", "Ditangkap", "Tewas"], before_after_rows_44),
        "",
        "##### Tabel 4.12: Ringkasan Eksekutif Seluruh Skenario Crosstab Ekspansi vs Eskalasi Konflik",
        markdown_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_44),
        "",
        "#### E. Analisis Temuan Empiris: Validitas Statistik Eskalasi Konflik",
        finding_44,
        "",
        "## 4.5 Peta Entitas Aktor: Korporasi dan Organisasi Masyarakat",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Teks Bebas (Free-Text): `data/processed/sulawesi_konflik_agraria_tanahkita.csv`. Visualisasi dashboard menampilkan dua Horizontal Bar Chart berdampingan (Top 10 Entitas Korporasi Paling Dominan dan Top Aktor Proksi & Vigilante Terdeteksi) hasil ekstraksi teks berbasis NLP Regex dari korpus narasi seluruh kasus agraria.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Analisis entitas aktor berbasis pemrosesan teks (*string parsing*) terhadap catatan kronologi dokumentasi TanahKita memetakan keterlibatan berbagai pihak dalam sengketa agraria. Korpus dibangun dari penggabungan kolom `judul`, `deskripsi`, dan `narasi` pada **{n_kasus_45:,} kasus agraria** (nasional) untuk memetakan orkestrasi struktural dan modus operandi aktor secara utuh, termasuk memvalidasi indikasi konsentrasi kekuasaan oleh segelintir konglomerasi besar.",
        "",
        "#### B. Alur Logika Metodologis Frequency Profiling (Text Parsing NLP)",
        "Kerangka ekstraksi entitas berbasis Regular Expressions (RegEx) dan penghitungan frekuensi penyebutan diilustrasikan pada **Bagan Alur 4.5** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Frequency Profiling deskriptif atas kemunculan entitas dalam korpus teks.",
        "",
        "##### Bagan Alur 4.5: Alur Logika Analisis Frequency Profiling Entitas Aktor",
        "```mermaid",
        mermaid_str_4_5,
        "```",
        "",
        "#### C. Formulasi Matematis: Konstruksi Korpus dan Token Counting",
        "Kuantifikasi frekuensi penyebutan entitas dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Korpus = Gabungan ( judul_k , deskripsi_k , narasi_k )   ;   untuk k = 1 s.d. N",
        "Frekuensi_a = Σ ( Match_i,a )   ;   untuk seluruh kemunculan pola entitas a dalam Korpus",
        "```",
        "",
        "Pola korporasi mendeteksi awalan PT/CV diikuti nama kapital (maksimum 4 kata) dengan normalisasi varian PTPN; pola aktor proksi mendeteksi kata kunci Preman, Ormas, Satgas, PAM Swakarsa, Pemuda Pancasila, GRIB, Laskar, Oknum, Security, Satpam, Centeng, Beking beserta frasa lanjutannya (dipotong pada stopword pertama).",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 4.13: Top 10 Entitas Korporasi Paling Dominan dalam Dokumentasi Konflik",
        markdown_table(["Entitas Korporasi", "Frekuensi Penyebutan"], corp_rows_45),
        "",
        "##### Tabel 4.14: Top Aktor Proksi & Vigilante Terdeteksi dalam Dokumentasi Konflik",
        markdown_table(["Aktor Proksi / Vigilante", "Frekuensi Penyebutan"], civil_rows_45),
        "",
        "#### E. Analisis Temuan Empiris: Orkestrasi Konflik dan Pemetaan Oligarki",
        f"1. **Dominasi Entitas Korporasi:** entitas **{top1_corp_name_45}** tercatat tertinggi dengan **{top1_corp_freq_45:,} catatan kasus terpisah** (dari {n_entitas_corp_45:,} entitas terdeteksi dan {total_mentions_corp_45:,} total penyebutan), memvalidasi indikasi konsentrasi kekuasaan dan monopoli penguasaan ruang oleh segelintir konglomerasi besar.",
        f"2. **Orkestrasi Horizontal Aktor Proksi:** kelompok sipil seperti **{top1_civ_name_45}** terdeteksi hingga **{top1_civ_freq_45:,} kali** — korporasi seringkali menggunakan pengamanan swakarsa, kelompok preman, hingga ormas vigilante sebagai 'bemper proksi' untuk mengintimidasi warga lokal dan memecah belah solidaritas akar rumput.",
        "",
    ]
    md_path = tool_dir / "Metodologi_Bab4_Ruang_Hidup.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")

    print("[4/4] Selesai membangun Bab 4 sub-bab 4.1.")


if __name__ == "__main__":
    generate_all_bab4()
