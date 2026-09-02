#!/usr/bin/env python3
"""
Generator Metodologi Versi Compact Bab 6 (Level Provinsi) — GAYA AKADEMIS TERPADU (CELIOS)
Mengadopsi arsitektur metodologi ringkas terstandarisasi konsisten dengan Bab 1 s.d. 5 dan Bab 6 Pulau:
- RUANG LINGKUP: Level 6 Provinsi Se-Pulau Sulawesi (Model Hybrid Z-Score Anomali & Entropy Weight Method / EWM)
- FORMAT: 1 KOLOM PENUH (Single Column Layout)
- PANJANG: 2–3 Halaman Maksimal (Elegan, proporsional, tanpa pemadatan berlebihan)
- PENOMORAN SEKSI UTAMA: Huruf kapital A, B, C, D, E, F
- SUB-BAB SEKSI D: Sub-bab 6.6 (6.6.1 s.d. 6.6.7 sesuai dokumen induk)
- OPERASIONALISASI INDIKATOR: 20 Indikator Empiris Kunci Terverifikasi (5 Kolom Baku tanpa kolom Periode)
- FORMULASI & TABEL SINTESIS: Formula matematis universal EWM/Z-Score dan Tabel 6.6a Komparasi 6 Provinsi
- KORESPONDENSI METODOLOGI: 3 kolom bersih (Sub-bab, Fokus Kajian Empiris Provinsi, Metode Analitis Utama)
- FLOWCHART: Mermaid JS horizontal (flowchart LR) dirender tajam ke DOCX (16.5 cm) dan blok kode di MD
- SINKRONISASI: Dual-save ke direktori versicompact/bab_6 dan bab_6.
"""

import os
import sys
import shutil
import base64
import requests
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Palet Warna Resmi CELIOS ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Forest Dark Green (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Celios Accent Green (#2E7D32)
G_LIGHT = RGBColor(0x38, 0x8E, 0x3C)  # Medium Green (#388E3C)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Charcoal Body Text (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Muted Grey Text (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Pure White

# ── Helper Fungsi XML Word ───────────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=80, right=80, top=40, bottom=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)

def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    pPr.append(s)

def para_border_bottom(p, color='1B5E20', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '3')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='1B5E20', sz='8'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Typography & Content ──────────────────────────────
def add_run(p, text, bold=False, italic=False, pt=8.5, color=C_BODY, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    else:
        r.font.name = 'Calibri'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    return r

def add_h2(doc, prefix, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_border_bottom(p, color='1B5E20', sz='8')
    add_run(p, f"{prefix}.  ", bold=True, pt=11, color=G_DARK)
    add_run(p, title.upper(), bold=True, pt=11, color=G_DARK)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, pt=9.5, color=G_MID)
    return p

def add_body(doc, parts, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        add_run(p, text, bold=bold, italic=italic, pt=8.5, color=C_BODY)
    return p

def add_formula(doc, text, ket=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    para_shd(p, 'EDF7EE')
    add_run(p, text, pt=8, color=G_MID, mono=True)
    if ket:
        add_run(p, f"\nKeterangan: {ket}", italic=True, pt=7.5, color=RGBColor(0x33, 0x55, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_styled(doc, headers, rows, col_widths_cm, alignments=None, font_pt=7):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    bd_subtle = {'val': 'single', 'sz': '4', 'color': 'D0D7DE', 'space': '0'}
    
    # Header Row
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=70, right=70, top=45, bottom=45)
        set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if alignments and j < len(alignments):
            align = alignments[j]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, pt=7.5, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F9FBF9' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_shd(c, fill)
            cell_margin(c, left=70, right=70, top=35, bottom=35)
            set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            
            if alignments and j < len(alignments):
                align = alignments[j]
                if align == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_run(p, str(val), pt=font_pt, color=C_BODY)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after  = Pt(3)
    return tbl

def download_mermaid_png(mermaid_str, filepath):
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        return True
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False


def generate_bab6_provinsi_compact():
    print("[1/3] Membangun dokumen compact Bab 6 Provinsi (Format 1-Kolom, 2-3 Halaman)...")
    
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab6    = out_dir_compact.parent.parent / "bab_6"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab6.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER DOKUMEN ──────────────────────────────────────────
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after  = Pt(1)
    add_run(p_t, "RINGKASAN EKSEKUTIF METODOLOGIS", bold=True, pt=8.5, color=G_MID)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after  = Pt(2)
    para_border_bottom(p_h, color='1B5E20', sz='12')
    add_run(p_h, "BAB 6: AUDIT FORENSIK METODOLOGI D3TLH (LEVEL PROVINSI)", bold=True, pt=15, color=G_DARK)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(1)
    p_meta.paragraph_format.space_after  = Pt(5)
    add_run(p_meta, "Sub-bab 6.6: Model Skoring Spasial-Statistik Z-Score & Entropy Weight Method (EWM) 6 Provinsi · ", italic=True, pt=8, color=C_GREY)
    add_run(p_meta, "Center of Economic and Law Studies (CELIOS)", bold=True, italic=True, pt=8, color=G_DARK)

    # ── A. DESAIN PENELITIAN & TUJUAN ───────────────────────────
    add_h2(doc, "A", "Desain Penelitian & Tujuan")
    add_body(doc, [
        ("Kajian tingkat provinsi ini menggunakan ", False, False),
        ("desain evaluasi multi-kriteria regional, standardisasi deviasi spasial (Z-Score), dan pembobotan objektif berbasis dispersi informasi Shannon (Entropy Weight Method / EWM)", True, False),
        (" sesuai kaidah ", False, False),
        ("Nature Scientific Reports", True, True),
        (" (Sun et al., 2024). Pendekatan ini dirancang untuk mengatasi kelemahan teknik perataan wilayah (*dilution effect*) pada dokumen D3TLH konvensional pemerintah yang kerap mengaburkan episentrum kerusakan lingkungan lokal. Tiga tujuan utama metodologis Bab 6 (Level Provinsi) mencakup:", False, False)
    ])
    add_body(doc, [
        ("1. ", True, False), ("Standardisasi Anomali Spasial Lintas Wilayah (Z-Score): ", True, False),
        ("Mengukur deviasi empiris 20 indikator multisektor pada masing-masing dari 6 provinsi terhadap nilai rata-rata regional se-Pulau Sulawesi, termasuk perlakuan inversi tanda matematis untuk indikator kualitas air (IKA).\n", False, False),
        ("2. ", True, False), ("Pembobotan Objektif Shannon Entropy (EWM): ", True, False),
        ("Menetapkan bobot signifikansi masing-masing indikator secara murni berbasis dispersi informasi data aktual tanpa intervensi bobot subjektif, sehingga indikator dengan ketimpangan tertinggi (seperti B3, tailing, dan konflik agraria) memperoleh bobot analitis terbesar.\n", False, False),
        ("3. ", True, False), ("Tipologi & Peringkat Kerentanan Ekologis Komparatif: ", True, False),
        ("Mengagregasikan skor pilar ke dalam Indeks Komposit Likert (0–5) dan Weighted Sum Model (0–10) guna memetakan polarisasi status daya dukung 6 provinsi antara episentrum industri nikel vs zona agromaritim berdaya lentur.", False, False)
    ])

    # ── B. SUMBER DATA & CAKUPAN WILAYAH ─────────────────────────
    add_h2(doc, "B", "Sumber Data & Cakupan Wilayah")
    add_body(doc, [
        ("Analisis komparatif tingkat provinsi mengolah matriks data panel regional yang mencakup seluruh 6 provinsi di Pulau Sulawesi (Sulawesi Tengah, Sulawesi Tenggara, Sulawesi Selatan, Sulawesi Barat, Gorontalo, dan Sulawesi Utara) bersumber dari:", False, False)
    ])
    add_body(doc, [
        ("• ", True, False), ("Ditjen Minerba ESDM & Global Energy Monitor (GEM 2023): ", True, False),
        ("Kapasitas operasional PLTU captive per provinsi dan sebaran 574 izin tambang nikel aktif.\n", False, False),
        ("• ", True, False), ("Satelit Copernicus Sentinel-5P (NASA/ESA TROPOMI): ", True, False),
        ("Ekstraksi data troposferik densitas kolom gas nitrogen dioksida (NO2 rasio mol/m²) per yurisdiksi provinsi.\n", False, False),
        ("• ", True, False), ("Kementerian Kesehatan RI & Profil Kesehatan Daerah: ", True, False),
        ("Data morbiditas klinis ISPA dan Diare (Incidence Rate Ratio / IRR) serta audit kepatuhan faskes ASPAK SPA.\n", False, False),
        ("• ", True, False), ("Kementerian Lingkungan Hidup dan Kehutanan (KLHK): ", True, False),
        ("Indeks Kualitas Air (IKA), neraca timbulan limbah B3, dan batas daya tampung residu tailing/slag per provinsi.\n", False, False),
        ("• ", True, False), ("Global Forest Watch (GFW / Hansen UMD) & DIBI BNPB: ", True, False),
        ("Luasan deforestasi primer, emisi karbon FOLU, tutupan hutan lindung terambah, dan kejadian bencana hidrometeorologi.\n", False, False),
        ("• ", True, False), ("Konsorsium Pembaruan Agraria (CATAHU KPA) & Satya Bumi: ", True, False),
        ("Sebaran korban jiwa konflik agraria, manipulasi persetujuan FPIC, dan insiden kriminalisasi pembela HAM.", False, False)
    ])

    # ── C. OPERASIONALISASI VARIABEL & INDIKATOR RISET ──────────
    add_h2(doc, "C", "Operasionalisasi Variabel & Indikator Riset")
    add_body(doc, [
        ("Merujuk pada Tabel Verifikasi Threshold model evaluasi D3TLH, seluruh parameter bio-fisik cerobong, neraca perairan, kerusakan tutupan lahan, kerentanan hak sosial, hingga instrumen veto perizinan dioperasionalkan secara terstruktur ke dalam ", False, False),
        ("20 indikator riset empiris terverifikasi", True, False),
        (" sebagaimana dirangkum pada matriks operasional berikut:", False, False)
    ])

    add_caption(doc, "Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 6 (Level Provinsi)")
    table_indikator_data = [
        ["1a", "Kapasitas PLTU Captive (Udara 1a)", "Beban Pembakaran Batubara Industri Off-Grid", "Megawatt (MW)", "Global Energy Monitor (GEM 2023)"],
        ["1b", "Polusi NO2 Satelit TROPOMI (Udara 1b)", "Densitas Kolom Troposferik Gas NO2 Atmosfer", "µmol / m²", "Copernicus Sentinel-5P (NASA/ESA)"],
        ["2", "Rasio Morbiditas ISPA (Udara 2)", "Anomali Morbiditas Saluran Pernapasan (IRR)", "Rasio Peluang (IRR)", "Kemenkes RI & WHO EHC 6"],
        ["3", "Proporsi Timbulan Limbah B3 (Udara 3)", "Beban Residu B3 terhadap Agregat Nasional", "Persen (%)", "Laporan Kinerja Ditjen PSLB3 KLHK"],
        ["4", "Defisit Emisi Karbon CO2 (Udara 4)", "Pelepasan Karbon Deforestasi vs Target NDC", "Juta Ton CO2e", "GFW & SK MenLHK 168/2022 (FOLU)"],
        ["5", "Kualitas Air IKA & Cr6+ (Air 1)", "Status Mutu Air Sungai & Paparan Logam Berat", "Poin & mg/L", "Ditjen PPKL KLHK & Uji Lab AEER"],
        ["6", "Rasio Morbiditas Diare (Air 2)", "Anomali Morbiditas Saluran Pencernaan (IRR)", "Rasio Peluang (IRR)", "Kemenkes RI & Profil Kesehatan 2023"],
        ["7", "Konflik Ruang Air Pesisir (Air 3)", "Letupan Sengketa Ruang Tangkap Nelayan", "Kasus", "Konsorsium Pembaruan Agraria (KPA)"],
        ["8", "Beban Residu Tailing & Slag (Air 4)", "Akumulasi Timbulan Tailing Dam & Slag", "Juta Ton / Tahun", "PT HPI-IMIP & AEER 2020"],
        ["9", "Bencana Hidrometeorologi (Lahan 1)", "Frekuensi Kejadian Banjir & Tanah Longsor", "Kejadian", "Data Informasi Bencana Indonesia BNPB"],
        ["10", "Deforestasi Hutan Primer (Lahan 2)", "Kehilangan Tutupan vs Kuota FOLU Net Sink", "Hektar (Ha)", "GFW Hansen & Renops FOLU 2030"],
        ["11", "Perambahan Hutan Lindung (Lahan 3)", "Pelanggaran Kawasan Lindung (Nol Toleransi)", "Hektar (Ha)", "GFW Overlay & UU No. 41/1999"],
        ["12", "Aktor Tambang & Sawit (Lahan 4)", "Deforestasi Akibat Komoditas Ekstraktif", "Hektar (Ha)", "GFW Dominant Drivers of Loss"],
        ["13", "Kepadatan Konsesi Tambang (Lahan 5)", "Rasio Konsesi IUP terhadap Luas Daratan", "Persen (%)", "Ditjen Minerba ESDM & BPS"],
        ["14", "Pelanggaran Asas FPIC (Sosial 1)", "Manipulasi Persetujuan Bebas Awal Warga", "Kasus", "Koalisi Sipil (JATAM, WALHI, AMAN)"],
        ["15", "Masyarakat Terdampak (Sosial 2)", "Korban Penggusuran & Perampasan Ruang", "Jiwa", "CATAHU KPA 2023"],
        ["16", "Kriminalisasi Pembela HAM (Sosial 3)", "Serangan & Penuntutan Hukum Warga/Aktivis", "Insiden", "Laporan Satya Bumi & KPA"],
        ["17", "Defisit Sarana Faskes SPA (Sosial 4)", "Kesenjangan Pemenuhan Standar Puskesmas", "Persen Kesenjangan (%)", "ASPAK Kemenkes & Permenkes 6/2024"],
        ["18", "Obral Perizinan IUP Baru (Veto 1)", "Penerbitan IUP Baru di Zona Kritis Ekologis", "Unit Izin", "Data Registry Ditjen Minerba ESDM"],
        ["19", "Impunitas Tambang Ilegal (Veto 2)", "Pembiaran Korporasi Pelanggar Tata Ruang", "Korporasi", "Catatan Akhir Tahun (CATAHU) KPA"],
        ["20", "Ekspansi PLTU Captive (Veto 3)", "Pemberian Izin PLTU Batubara Off-Grid", "Megawatt (MW)", "Global Energy Monitor (GEM 2023)"]
    ]

    add_table_styled(
        doc,
        headers=["No", "Indikator Riset", "Fokus Pengukuran", "Satuan", "Sumber Data Primer Resmi"],
        rows=table_indikator_data,
        col_widths_cm=[0.8, 4.5, 4.5, 2.2, 5.0],
        alignments=['C', 'L', 'L', 'C', 'L'],
        font_pt=7
    )

    # ── D. KERANGKA ANALISIS & FORMULASI MATEMATIS ──────────────
    add_h2(doc, "D", "Kerangka Analisis & Formulasi Matematis")

    add_h3(doc, "Sub-bab 6.6: Algoritma Skoring Tingkat Provinsi (Model Hybrid Z-Score & EWM)")
    add_body(doc, [
        ("Model evaluasi regional mengombinasikan standardisasi Z-Score untuk mendeteksi outlier deviasi spasial dengan pembobotan objektif Entropy Weight Method (EWM) berbasis dispersi informasi Shannon, dilanjutkan pemetaan skala Likert diskret (0–5):", False, False)
    ])
    add_formula(doc, "1. Z-Score Regional: Z_ij = (x_ij - mean_j) / std_j   |   Khusus IKA: Z_ika = - (ika_i - mean_ika) / std_ika\n"
                     "2. Min-Max Normalisasi: r_ij = (x_ij - min_j) / (max_j - min_j)   ;   P_ij = r_ij / Σ r_ij\n"
                     "3. Entropi Shannon: E_j = - (1 / ln(n)) × Σ [ P_ij × ln(P_ij + ε) ]   ;   D_j = 1 - E_j   ;   W_j = D_j / Σ D_j\n"
                     "4. Pemetaan Likert Diskret: Z >= +1.0σ -> 5.0 ; 0.5 <= Z < 1.0 -> 4.0 ; 0.0 <= Z < 0.5 -> 3.0 ; -0.5 <= Z < 0.0 -> 2.0 ; -1.0 <= Z < -0.5 -> 1.0 ; Z < -1.0 -> 0.0\n"
                     "5. Skor Pilar Terbobot EWM: Skor_Pilar_i = Σ [ L_ij × W_j ] / Σ W_j   |   Skor Komposit = [ Σ Skor_Pilar (1..5) ] / 5.0",
                ket="n = 6 Provinsi; ε = 1e-12; W_j = Bobot objektif Shannon EWM (B3 = 8,29%, Tailing = 8,22%, Korban Agraria = 7,81%, PLTU = 7,73%); L_ij = Skor Likert diskret (0-5); Nilai Z >= +1.0σ merefleksikan anomali krisis ekstrem (Red Alert).")

    add_body(doc, [
        ("Penerapan model ini pada 6 provinsi se-Pulau Sulawesi menghasilkan sintesis komparatif tingkat kerentanan regional sebagaimana dirangkum pada tabel berikut:", False, False)
    ])

    add_caption(doc, "Tabel 6.6a: Matriks Sintesis Komparatif Skor D3TLH 6 Provinsi Se-Pulau Sulawesi")
    tabel_6_6a_rows = [
        ["1", "Sulawesi Tengah", "4.9", "3.3", "4.7", "2.5", "4.4", "4.0 / 5", "7.92", "Melampaui Batas", "Episentrum PLTU Captive (7.325 MW), B3 (25,3 Jt Ton), Deforestasi Masif"],
        ["2", "Sulawesi Tenggara", "3.1", "2.8", "3.6", "4.5", "3.0", "3.4 / 5", "6.78", "Mendekati Batas", "Krisis Agraria (39.821 Jiwa), Kepadatan IUP Ekstrem (11,72%), Sengketa FPIC"],
        ["3", "Sulawesi Selatan", "2.1", "2.9", "2.7", "2.4", "3.1", "2.6 / 5", "5.29", "Mendekati Batas", "Rekor Bencana Alam (669 Kejadian), Konflik Pesisir Nelayan, Kriminalisasi HAM"],
        ["4", "Sulawesi Utara", "0.9", "1.5", "1.8", "2.4", "1.3", "1.6 / 5", "3.13", "Tidak Melampaui Batas", "Outlier Kesenjangan Faskes SPA Kepulauan (25,16%), Isu Tambang Sangihe"],
        ["5", "Sulawesi Barat", "1.2", "1.9", "0.8", "1.0", "1.0", "1.2 / 5", "2.36", "Tidak Melampaui Batas", "Bioregion Agromaritim Terjaga, Tekanan Mutu Air Sungai Akibat PKS Sawit"],
        ["6", "Gorontalo", "1.4", "1.4", "0.7", "1.0", "1.3", "1.2 / 5", "2.31", "Tidak Melampaui Batas", "Atmosfer Satelit NO2 Terbersih, Deforestasi & Emisi Karbon Terendah"]
    ]
    add_table_styled(
        doc,
        headers=["Rank", "Provinsi", "Udara", "Air", "Lahan", "Sosial", "Veto", "Likert", "WSM", "Status Ekologis", "Faktor Determinan Utama"],
        rows=tabel_6_6a_rows,
        col_widths_cm=[0.8, 2.5, 1.0, 1.0, 1.0, 1.0, 1.0, 1.3, 1.1, 2.3, 4.0],
        alignments=['C', 'L', 'C', 'C', 'C', 'C', 'C', 'C', 'C', 'C', 'L'],
        font_pt=6.5
    )

    # ── E. KORESPONDENSI METODOLOGI TERHADAP SUB-BAB LAPORAN ────
    add_h2(doc, "E", "Korespondensi Metodologi terhadap Sub-bab Laporan Bab 6")
    add_body(doc, [
        ("Setiap sub-bab analitis tingkat provinsi pada Bab 6 ditopang oleh metode empiris terstandarisasi sebagaimana dirangkum pada matriks korespondensi berikut:", False, False)
    ])

    table_korespondensi = [
        ["Sub-bab 6.6.1", "Evaluasi D3TLH Sulawesi Tengah", "Z-Score Anomaly Mapping, EWM Weighting, Red Alert Asimilasi Udara & Tekanan Lahan"],
        ["Sub-bab 6.6.2", "Evaluasi D3TLH Sulawesi Tenggara", "Spatial Mining Density Audit, Agrarian Dispossession Scaling, Coastal FPIC Evaluation"],
        ["Sub-bab 6.6.3", "Evaluasi D3TLH Sulawesi Selatan", "Hydrometeorological Outlier Normalization, Heavy Metal Cr6+ Detection, SLAPP Tracking"],
        ["Sub-bab 6.6.4", "Evaluasi D3TLH Sulawesi Barat", "Baseline Control Group Analysis, PKS Water Quality Deficit Audit, Non-Smelter Modeling"],
        ["Sub-bab 6.6.5", "Evaluasi D3TLH Provinsi Gorontalo", "Clean Atmosphere Baseline Tracking, Inversion Air Dispersion, Low-Stress Resilience"],
        ["Sub-bab 6.6.6", "Evaluasi D3TLH Sulawesi Utara", "Small Island Vulnerability Assessment, Health Infrastructure (SPA) Gap Outlier"],
        ["Sub-bab 6.6.7", "Sintesis Komparatif 6 Provinsi", "Multi-Criteria Regional Ranking, Spatial Typology Classification, Moratorium Mandate"]
    ]

    add_table_styled(
        doc,
        headers=["Sub-bab", "Fokus Kajian Empiris Provinsi", "Metode Analitis Utama"],
        rows=table_korespondensi,
        col_widths_cm=[2.5, 5.5, 9.0],
        alignments=['C', 'L', 'L'],
        font_pt=7.5
    )

    # ── F. BAGAN ALUR KERANGKA KERJA RISET BAB 6 PROVINSI ────────
    add_h2(doc, "F", "Bagan Alur Kerangka Kerja Riset (Research Workflow)")
    add_body(doc, [
        ("Alur komputasi analitis tingkat provinsi dijalankan secara terintegrasi melalui empat tahapan metodologis sebagaimana divisualisasikan pada bagan berikut:", False, False)
    ])

    mermaid_str_f = """flowchart LR
    subgraph F1["Fase I: Matriks 20 Indikator"]
        A1["Data 6 Provinsi<br/><i>ESDM, GEM, Satelit NO2</i>"]
        A2["Data Air & Limbah<br/><i>IKA, Tailing, Cr6+</i>"]
        A3["Data Lahan & Bencana<br/><i>GFW, BNPB, Lindung</i>"]
        A4["Data Sosial & Veto<br/><i>KPA, FPIC, Izin Baru</i>"]
    end
    subgraph F2["Fase II: Z-Score & EWM"]
        B1["Standardisasi Deviasi<br/><i>Z = (x - mean) / std</i>"]
        B2["Inversi Mutu IKA<br/><i>Z_ika = - (ika - mean)/std</i>"]
        B3["Entropi Shannon<br/><i>Ej = -(1/ln n) Σ P ln P</i>"]
        B4["Bobot Objektif Wj<br/><i>B3, Tailing, Korban, PLTU</i>"]
    end
    subgraph F3["Fase III: Transformasi Likert"]
        C1["Mapping Z >= +1.0σ<br/><i>Skor 5.0 (Red Alert)</i>"]
        C2["Mapping 0.5 <= Z < 1.0<br/><i>Skor 4.0 (Melampaui)</i>"]
        C3["Mapping -0.5 <= Z < 0.5<br/><i>Skor 2.0 - 3.0 (Waspada)</i>"]
        C4["Mapping Z < -0.5σ<br/><i>Skor 0.0 - 1.0 (Aman)</i>"]
    end
    subgraph F4["Fase IV: Sintesis 6 Provinsi"]
        D1["Sulteng (4.0/5: Red Alert)<br/><i>Krisis Asimilasi Udara & B3</i>"]
        D2["Sultra (3.4/5: Waspada)<br/><i>Krisis Agraria & Konsesi IUP</i>"]
        D3["Sulsel (2.6/5: Rentan)<br/><i>Bencana & Represi Hukum</i>"]
        D4["Sulut, Sulbar, Gorontalo<br/><i>1.2 - 1.6/5 (Zona Resiliensi)</i>"]
    end
    F1 --> F2 --> F3 --> F4"""

    png_workflow_path = str(out_dir_compact / "mermaid_workflow_bab6_provinsi.png")
    is_downloaded = download_mermaid_png(mermaid_str_f, png_workflow_path)

    add_caption(doc, "Bagan Alur 6.2: Alur Logika Model Hybrid Z-Score & EWM Tingkat Provinsi (Research Workflow)")
    if is_downloaded and os.path.exists(png_workflow_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(3)
        p_img.paragraph_format.space_after  = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(png_workflow_path, width=Cm(16.5))
        try:
            shutil.copyfile(png_workflow_path, str(out_dir_bab6 / "mermaid_workflow_bab6_provinsi.png"))
        except Exception:
            pass

    # Box Output Kesimpulan
    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after  = Pt(4)
    all_border_para(p_box, color='1B5E20', sz='8')
    para_shd(p_box, 'F1F8E9')
    add_run(p_box, "KERANGKA KELUARAN METODOLOGIS BAB 6 (LEVEL PROVINSI):\n", bold=True, pt=8.5, color=G_DARK)
    add_run(p_box, "1. Eliminasi Bias Perataan Spasial: Penerapan model Hybrid Z-Score dan pembobotan entropi Shannon (EWM) berhasil mendeteksi anomali krisis ekstrem yang selama ini tertutupi oleh teknik agregasi makro dokumen D3TLH pemerintah.\n"
                   "2. Polarisasi Tipologi Ekologis: Mengonfirmasi adanya jurang pemisah tajam antara provinsi episentrum hilirisasi nikel (Sulteng: Skor 4,0/5 Red Alert dan Sultra: Skor 3,4/5 Krisis Agraria & Kepadatan Konsesi) dibandingkan provinsi agromaritim (Sulbar, Gorontalo, Sulut: Skor 1,2–1,6/5 Terjaga).\n"
                   "3. Dasar Intervensi Moratorium Terarah: Menyediakan justifikasi kuantitatif objektif bagi pembuat kebijakan untuk segera memberlakukan moratorium total penerbitan IUP dan penghentian pembangunan PLTU captive baru di provinsi-provinsi berstatus Red Alert.",
            pt=8, color=C_BODY)

    # ── SIMPAN DOKUMEN DOCX (DUAL SAVE) ─────────────────────────
    docx_compact = out_dir_compact / "Metodologi_Bab6_Provinsi_Skoring_Compact.docx"
    docx_bab6    = out_dir_bab6 / "Metodologi_Bab6_Provinsi_Skoring_Compact.docx"
    
    doc.save(str(docx_compact))
    shutil.copyfile(docx_compact, docx_bab6)
    print(f"  [OK] Tersimpan DOCX: {docx_compact}")
    print(f"  [OK] Salinan DOCX : {docx_bab6}")

    # ── GENERATE MARKDOWN PADANAN ───────────────────────────────
    print("[2/3] Membangun dokumen Markdown padanan...")
    MD_CONTENT = """# METODOLOGI PENELITIAN: BAB 6 — AUDIT FORENSIK D3TLH (LEVEL PROVINSI)
*Sub-bab 6.6: Model Skoring Spasial-Statistik Z-Score & Entropy Weight Method (EWM) 6 Provinsi · CELIOS*

---

## A. Desain Penelitian & Tujuan
Kajian tingkat provinsi ini menggunakan **desain evaluasi multi-kriteria regional, standardisasi deviasi spasial (Z-Score), dan pembobotan objektif berbasis dispersi informasi Shannon (Entropy Weight Method / EWM)** sesuai kaidah *Nature Scientific Reports* (Sun et al., 2024). Pendekatan ini dirancang untuk mengatasi kelemahan teknik perataan wilayah (*dilution effect*) pada dokumen D3TLH konvensional pemerintah yang kerap mengaburkan episentrum kerusakan lingkungan lokal. Tiga tujuan utama metodologis Bab 6 (Level Provinsi) mencakup:

1. **Standardisasi Anomali Spasial Lintas Wilayah (Z-Score):** Mengukur deviasi empiris 20 indikator multisektor pada masing-masing dari 6 provinsi terhadap nilai rata-rata regional se-Pulau Sulawesi, termasuk perlakuan inversi tanda matematis untuk indikator kualitas air (IKA).
2. **Pembobotan Objektif Shannon Entropy (EWM):** Menetapkan bobot signifikansi masing-masing indikator secara murni berbasis dispersi informasi data aktual tanpa intervensi bobot subjektif, sehingga indikator dengan ketimpangan tertinggi (seperti B3, tailing, dan konflik agraria) memperoleh bobot analitis terbesar.
3. **Tipologi & Peringkat Kerentanan Ekologis Komparatif:** Mengagregasikan skor pilar ke dalam Indeks Komposit Likert (0–5) dan Weighted Sum Model (0–10) guna memetakan polarisasi status daya dukung 6 provinsi antara episentrum industri nikel vs zona agromaritim berdaya lentur.

---

## B. Sumber Data & Cakupan Wilayah
Analisis komparatif tingkat provinsi mengolah matriks data panel regional yang mencakup seluruh 6 provinsi di Pulau Sulawesi (Sulawesi Tengah, Sulawesi Tenggara, Sulawesi Selatan, Sulawesi Barat, Gorontalo, dan Sulawesi Utara) bersumber dari:

- **Ditjen Minerba ESDM & Global Energy Monitor (GEM 2023):** Kapasitas operasional PLTU captive per provinsi dan sebaran 574 izin tambang nikel aktif.
- **Satelit Copernicus Sentinel-5P (NASA/ESA TROPOMI):** Ekstraksi data troposferik densitas kolom gas nitrogen dioksida (NO2 rasio mol/m²) per yurisdiksi provinsi.
- **Kementerian Kesehatan RI & Profil Kesehatan Daerah:** Data morbiditas klinis ISPA dan Diare (Incidence Rate Ratio / IRR) serta audit kepatuhan faskes ASPAK SPA.
- **Kementerian Lingkungan Hidup dan Kehutanan (KLHK):** Indeks Kualitas Air (IKA), neraca timbulan limbah B3, dan batas daya tampung residu tailing/slag per provinsi.
- **Global Forest Watch (GFW / Hansen UMD) & DIBI BNPB:** Luasan deforestasi primer, emisi karbon FOLU, tutupan hutan lindung terambah, dan kejadian bencana hidrometeorologi.
- **Konsorsium Pembaruan Agraria (CATAHU KPA) & Satya Bumi:** Sebaran korban jiwa konflik agraria, manipulasi persetujuan FPIC, dan insiden kriminalisasi pembela HAM.

---

## C. Operasionalisasi Variabel & Indikator Riset
Merujuk pada Tabel Verifikasi Threshold model evaluasi D3TLH, seluruh parameter bio-fisik cerobong, neraca perairan, kerusakan tutupan lahan, kerentanan hak sosial, hingga instrumen veto perizinan dioperasionalkan secara terstruktur ke dalam **20 indikator riset empiris terverifikasi** sebagaimana dirangkum pada matriks operasional berikut:

##### Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 6 (Level Provinsi)
| No | Indikator Riset | Fokus Pengukuran | Satuan | Sumber Data Primer Resmi |
| :-: | :--- | :--- | :-: | :--- |
| 1a | Kapasitas PLTU Captive (Udara 1a) | Beban Pembakaran Batubara Industri Off-Grid | Megawatt (MW) | Global Energy Monitor (GEM 2023) |
| 1b | Polusi NO2 Satelit TROPOMI (Udara 1b) | Densitas Kolom Troposferik Gas NO2 Atmosfer | µmol / m² | Copernicus Sentinel-5P (NASA/ESA) |
| 2 | Rasio Morbiditas ISPA (Udara 2) | Anomali Morbiditas Saluran Pernapasan (IRR) | Rasio Peluang (IRR) | Kemenkes RI & WHO EHC 6 |
| 3 | Proporsi Timbulan Limbah B3 (Udara 3) | Beban Residu B3 terhadap Agregat Nasional | Persen (%) | Laporan Kinerja Ditjen PSLB3 KLHK |
| 4 | Defisit Emisi Karbon CO2 (Udara 4) | Pelepasan Karbon Deforestasi vs Target NDC | Juta Ton CO2e | GFW & SK MenLHK 168/2022 (FOLU) |
| 5 | Kualitas Air IKA & Cr6+ (Air 1) | Status Mutu Air Sungai & Paparan Logam Berat | Poin & mg/L | Ditjen PPKL KLHK & Uji Lab AEER |
| 6 | Rasio Morbiditas Diare (Air 2) | Anomali Morbiditas Saluran Pencernaan (IRR) | Rasio Peluang (IRR) | Kemenkes RI & Profil Kesehatan 2023 |
| 7 | Konflik Ruang Air Pesisir (Air 3) | Letupan Sengketa Ruang Tangkap Nelayan | Kasus | Konsorsium Pembaruan Agraria (KPA) |
| 8 | Beban Residu Tailing & Slag (Air 4) | Akumulasi Timbulan Tailing Dam & Slag | Juta Ton / Tahun | PT HPI-IMIP & AEER 2020 |
| 9 | Bencana Hidrometeorologi (Lahan 1) | Frekuensi Kejadian Banjir & Tanah Longsor | Kejadian | Data Informasi Bencana Indonesia BNPB |
| 10 | Deforestasi Hutan Primer (Lahan 2) | Kehilangan Tutupan vs Kuota FOLU Net Sink | Hektar (Ha) | GFW Hansen & Renops FOLU 2030 |
| 11 | Perambahan Hutan Lindung (Lahan 3) | Pelanggaran Kawasan Lindung (Nol Toleransi) | Hektar (Ha) | GFW Overlay & UU No. 41/1999 |
| 12 | Aktor Tambang & Sawit (Lahan 4) | Deforestasi Akibat Komoditas Ekstraktif | Hektar (Ha) | GFW Dominant Drivers of Loss |
| 13 | Kepadatan Konsesi Tambang (Lahan 5) | Rasio Konsesi IUP terhadap Luas Daratan | Persen (%) | Ditjen Minerba ESDM & BPS |
| 14 | Pelanggaran Asas FPIC (Sosial 1) | Manipulasi Persetujuan Bebas Awal Warga | Kasus | Koalisi Sipil (JATAM, WALHI, AMAN) |
| 15 | Masyarakat Terdampak (Sosial 2) | Korban Penggusuran & Perampasan Ruang | Jiwa | CATAHU KPA 2023 |
| 16 | Kriminalisasi Pembela HAM (Sosial 3) | Serangan & Penuntutan Hukum Warga/Aktivis | Insiden | Laporan Satya Bumi & KPA |
| 17 | Defisit Sarana Faskes SPA (Sosial 4) | Kesenjangan Pemenuhan Standar Puskesmas | Persen Kesenjangan (%) | ASPAK Kemenkes & Permenkes 6/2024 |
| 18 | Obral Perizinan IUP Baru (Veto 1) | Penerbitan IUP Baru di Zona Kritis Ekologis | Unit Izin | Data Registry Ditjen Minerba ESDM |
| 19 | Impunitas Tambang Ilegal (Veto 2) | Pembiaran Korporasi Pelanggar Tata Ruang | Korporasi | Catatan Akhir Tahun (CATAHU) KPA |
| 20 | Ekspansi PLTU Captive (Veto 3) | Pemberian Izin PLTU Batubara Off-Grid | Megawatt (MW) | Global Energy Monitor (GEM 2023) |

---

## D. Kerangka Analisis & Formulasi Matematis

### Sub-bab 6.6: Algoritma Skoring Tingkat Provinsi (Model Hybrid Z-Score & EWM)
Model evaluasi regional mengombinasikan standardisasi Z-Score untuk mendeteksi outlier deviasi spasial dengan pembobotan objektif Entropy Weight Method (EWM) berbasis dispersi informasi Shannon, dilanjutkan pemetaan skala Likert diskret (0–5):

> `1. Z-Score Regional: Z_ij = (x_ij - mean_j) / std_j   |   Khusus IKA: Z_ika = - (ika_i - mean_ika) / std_ika`  
> `2. Min-Max Normalisasi: r_ij = (x_ij - min_j) / (max_j - min_j)   ;   P_ij = r_ij / Σ r_ij`  
> `3. Entropi Shannon: E_j = - (1 / ln(n)) × Σ [ P_ij × ln(P_ij + ε) ]   ;   D_j = 1 - E_j   ;   W_j = D_j / Σ D_j`  
> `4. Pemetaan Likert Diskret: Z >= +1.0σ -> 5.0 ; 0.5 <= Z < 1.0 -> 4.0 ; 0.0 <= Z < 0.5 -> 3.0 ; -0.5 <= Z < 0.0 -> 2.0 ; -1.0 <= Z < -0.5 -> 1.0 ; Z < -1.0 -> 0.0`  
> `5. Skor Pilar Terbobot EWM: Skor_Pilar_i = Σ [ L_ij × W_j ] / Σ W_j   |   Skor Komposit = [ Σ Skor_Pilar (1..5) ] / 5.0`  
> *Keterangan: n = 6 Provinsi; ε = 1e-12; W_j = Bobot objektif Shannon EWM (B3 = 8,29%, Tailing = 8,22%, Korban Agraria = 7,81%, PLTU = 7,73%); L_ij = Skor Likert diskret (0-5); Nilai Z >= +1.0σ merefleksikan anomali krisis ekstrem (Red Alert).*

##### Tabel 6.6a: Matriks Sintesis Komparatif Skor D3TLH 6 Provinsi Se-Pulau Sulawesi
| Rank | Provinsi | Udara | Air | Lahan | Sosial | Veto | Likert | WSM | Status Ekologis | Faktor Determinan Utama |
| :---: | :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :--- |
| 1 | Sulawesi Tengah | 4.9 | 3.3 | 4.7 | 2.5 | 4.4 | 4.0 / 5 | 7.92 | Melampaui Batas | Episentrum PLTU Captive (7.325 MW), B3 (25,3 Jt Ton), Deforestasi Masif |
| 2 | Sulawesi Tenggara | 3.1 | 2.8 | 3.6 | 4.5 | 3.0 | 3.4 / 5 | 6.78 | Mendekati Batas | Krisis Agraria (39.821 Jiwa), Kepadatan IUP Ekstrem (11,72%), Sengketa FPIC |
| 3 | Sulawesi Selatan | 2.1 | 2.9 | 2.7 | 2.4 | 3.1 | 2.6 / 5 | 5.29 | Mendekati Batas | Rekor Bencana Alam (669 Kejadian), Konflik Pesisir Nelayan, Kriminalisasi HAM |
| 4 | Sulawesi Utara | 0.9 | 1.5 | 1.8 | 2.4 | 1.3 | 1.6 / 5 | 3.13 | Tidak Melampaui Batas | Outlier Kesenjangan Faskes SPA Kepulauan (25,16%), Isu Tambang Sangihe |
| 5 | Sulawesi Barat | 1.2 | 1.9 | 0.8 | 1.0 | 1.0 | 1.2 / 5 | 2.36 | Tidak Melampaui Batas | Bioregion Agromaritim Terjaga, Tekanan Mutu Air Sungai Akibat PKS Sawit |
| 6 | Gorontalo | 1.4 | 1.4 | 0.7 | 1.0 | 1.3 | 1.2 / 5 | 2.31 | Tidak Melampaui Batas | Atmosfer Satelit NO2 Terbersih, Deforestasi & Emisi Karbon Terendah |

---

## E. Korespondensi Metodologi terhadap Sub-bab Laporan Bab 6
Setiap sub-bab analitis tingkat provinsi pada Bab 6 ditopang oleh metode empiris terstandarisasi sebagaimana dirangkum pada matriks korespondensi berikut:

##### Matriks Korespondensi Metodologis Bab 6 (Level Provinsi)
| Sub-bab | Fokus Kajian Empiris Provinsi | Metode Analitis Utama |
| :-: | :--- | :--- |
| Sub-bab 6.6.1 | Evaluasi D3TLH Sulawesi Tengah | Z-Score Anomaly Mapping, EWM Weighting, Red Alert Asimilasi Udara & Tekanan Lahan |
| Sub-bab 6.6.2 | Evaluasi D3TLH Sulawesi Tenggara | Spatial Mining Density Audit, Agrarian Dispossession Scaling, Coastal FPIC Evaluation |
| Sub-bab 6.6.3 | Evaluasi D3TLH Sulawesi Selatan | Hydrometeorological Outlier Normalization, Heavy Metal Cr6+ Detection, SLAPP Tracking |
| Sub-bab 6.6.4 | Evaluasi D3TLH Sulawesi Barat | Baseline Control Group Analysis, PKS Water Quality Deficit Audit, Non-Smelter Modeling |
| Sub-bab 6.6.5 | Evaluasi D3TLH Provinsi Gorontalo | Clean Atmosphere Baseline Tracking, Inversion Air Dispersion, Low-Stress Resilience |
| Sub-bab 6.6.6 | Evaluasi D3TLH Sulawesi Utara | Small Island Vulnerability Assessment, Health Infrastructure (SPA) Gap Outlier |
| Sub-bab 6.6.7 | Sintesis Komparatif 6 Provinsi | Multi-Criteria Regional Ranking, Spatial Typology Classification, Moratorium Mandate |

---

## F. Bagan Alur Kerangka Kerja Riset (Research Workflow)
Alur komputasi analitis tingkat provinsi dijalankan secara terintegrasi melalui empat tahapan metodologis sebagaimana divisualisasikan pada bagan berikut:

```mermaid
flowchart LR
    subgraph F1["Fase I: Matriks 20 Indikator"]
        A1["Data 6 Provinsi<br/><i>ESDM, GEM, Satelit NO2</i>"]
        A2["Data Air & Limbah<br/><i>IKA, Tailing, Cr6+</i>"]
        A3["Data Lahan & Bencana<br/><i>GFW, BNPB, Lindung</i>"]
        A4["Data Sosial & Veto<br/><i>KPA, FPIC, Izin Baru</i>"]
    end
    subgraph F2["Fase II: Z-Score & EWM"]
        B1["Standardisasi Deviasi<br/><i>Z = (x - mean) / std</i>"]
        B2["Inversi Mutu IKA<br/><i>Z_ika = - (ika - mean)/std</i>"]
        B3["Entropi Shannon<br/><i>Ej = -(1/ln n) Σ P ln P</i>"]
        B4["Bobot Objektif Wj<br/><i>B3, Tailing, Korban, PLTU</i>"]
    end
    subgraph F3["Fase III: Transformasi Likert"]
        C1["Mapping Z >= +1.0σ<br/><i>Skor 5.0 (Red Alert)</i>"]
        C2["Mapping 0.5 <= Z < 1.0<br/><i>Skor 4.0 (Melampaui)</i>"]
        C3["Mapping -0.5 <= Z < 0.5<br/><i>Skor 2.0 - 3.0 (Waspada)</i>"]
        C4["Mapping Z < -0.5σ<br/><i>Skor 0.0 - 1.0 (Aman)</i>"]
    end
    subgraph F4["Fase IV: Sintesis 6 Provinsi"]
        D1["Sulteng (4.0/5: Red Alert)<br/><i>Krisis Asimilasi Udara & B3</i>"]
        D2["Sultra (3.4/5: Waspada)<br/><i>Krisis Agraria & Konsesi IUP</i>"]
        D3["Sulsel (2.6/5: Rentan)<br/><i>Bencana & Represi Hukum</i>"]
        D4["Sulut, Sulbar, Gorontalo<br/><i>1.2 - 1.6/5 (Zona Resiliensi)</i>"]
    end
    F1 --> F2 --> F3 --> F4
```

> **KERANGKA KELUARAN METODOLOGIS BAB 6 (LEVEL PROVINSI):**  
> 1. **Eliminasi Bias Perataan Spasial:** Penerapan model Hybrid Z-Score dan pembobotan entropi Shannon (EWM) berhasil mendeteksi anomali krisis ekstrem yang selama ini tertutupi oleh teknik agregasi makro dokumen D3TLH pemerintah.  
> 2. **Polarisasi Tipologi Ekologis:** Mengonfirmasi adanya jurang pemisah tajam antara provinsi episentrum hilirisasi nikel (Sulteng: Skor 4,0/5 Red Alert dan Sultra: Skor 3,4/5 Krisis Agraria & Kepadatan Konsesi) dibandingkan provinsi agromaritim (Sulbar, Gorontalo, Sulut: Skor 1,2–1,6/5 Terjaga).  
> 3. **Dasar Intervensi Moratorium Terarah:** Menyediakan justifikasi kuantitatif objektif bagi pembuat kebijakan untuk segera memberlakukan moratorium total penerbitan IUP dan penghentian pembangunan PLTU captive baru di provinsi-provinsi berstatus Red Alert.
"""

    md_compact = out_dir_compact / "Metodologi_Bab6_Provinsi_Skoring_Compact.md"
    md_bab6    = out_dir_bab6 / "Metodologi_Bab6_Provinsi_Skoring_Compact.md"
    with open(md_compact, 'w', encoding='utf-8') as f:
        f.write(MD_CONTENT)
    shutil.copyfile(md_compact, md_bab6)
    print(f"  [OK] Tersimpan MD  : {md_compact}")
    print(f"  [OK] Salinan MD   : {md_bab6}")

    print("[3/3] Selesai menghasilkan dokumen metodologi Bab 6 Provinsi versi compact (1-Kolom, 2-3 Halaman).\n")


if __name__ == "__main__":
    generate_bab6_provinsi_compact()
