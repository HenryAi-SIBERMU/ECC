#!/usr/bin/env python3
"""
Generator Metodologi Versi Compact Bab 6 (Level Pulau / Bioregion) — GAYA AKADEMIS TERPADU (CELIOS)
Mengadopsi arsitektur metodologi ringkas terstandarisasi konsisten dengan Bab 1 s.d. 5:
- RUANG LINGKUP: Level Bioregion Makro Pulau Sulawesi (Audit Forensik D3TLH Nasional vs Fakta Empiris)
- FORMAT: 1 KOLOM PENUH (Single Column Layout)
- PANJANG: 2–3 Halaman Maksimal (Elegan, proporsional, tanpa pemadatan berlebihan)
- PENOMORAN SEKSI UTAMA: Huruf kapital A, B, C, D, E, F
- SUB-BAB SEKSI D: 6.1, 6.2, 6.3, 6.4, 6.5 (Judul persis dokumen induk)
- OPERASIONALISASI INDIKATOR: 10 Indikator Empiris Kunci Bioregion Pulau (5 Kolom Baku tanpa kolom Periode)
- FORMULASI & TABEL REKAPITULASI: Formula matematis lengkap dan Tabel 6.5a Sintesis Skoring 5 Dimensi
- KORESPONDENSI METODOLOGI: 3 kolom bersih (Sub-bab, Fokus Kajian Empiris, Metode Analitis Utama)
- FLOWCHART: Mermaid JS horizontal (flowchart LR) dirender tajam ke DOCX (16.5 cm) dan blok kode di MD
- SINKRONISASI: Dual-save ke direktori versicompact/bab_6 dan bab_6.
"""

import os
import sys
import shutil
import base64
import requests
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Palet Warna Resmi CELIOS ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Forest Dark Green (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Celios Accent Green (#2E7D32)
G_LIGHT = RGBColor(0x38, 0x8E, 0x3C)  # Medium Green (#388E3C)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Charcoal Body Text (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Muted Grey Text (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Pure White

# ── Helper Fungsi XML Word ───────────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=80, right=80, top=40, bottom=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)

def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    pPr.append(s)

def para_border_bottom(p, color='1B5E20', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '3')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def para_border_left(p, color='2E7D32', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:left')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='1B5E20', sz='8'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Typography & Content ──────────────────────────────
def add_run(p, text, bold=False, italic=False, pt=8.5, color=C_BODY, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    else:
        r.font.name = 'Calibri'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    return r

def add_h2(doc, prefix, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_border_bottom(p, color='1B5E20', sz='8')
    add_run(p, f"{prefix}.  ", bold=True, pt=11, color=G_DARK)
    add_run(p, title.upper(), bold=True, pt=11, color=G_DARK)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, pt=9.5, color=G_MID)
    return p

def add_body(doc, parts, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        add_run(p, text, bold=bold, italic=italic, pt=8.5, color=C_BODY)
    return p

def add_formula(doc, text, ket=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    para_shd(p, 'EDF7EE')
    add_run(p, text, pt=8, color=G_MID, mono=True)
    if ket:
        add_run(p, f"\nKeterangan: {ket}", italic=True, pt=7.5, color=RGBColor(0x33, 0x55, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_styled(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    bd_subtle = {'val': 'single', 'sz': '4', 'color': 'D0D7DE', 'space': '0'}
    
    # Header Row
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=80, right=80, top=50, bottom=50)
        set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if alignments and j < len(alignments):
            align = alignments[j]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, pt=8, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F9FBF9' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_shd(c, fill)
            cell_margin(c, left=80, right=80, top=40, bottom=40)
            set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            
            if alignments and j < len(alignments):
                align = alignments[j]
                if align == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_run(p, str(val), pt=7.5, color=C_BODY)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after  = Pt(3)
    return tbl

def download_mermaid_png(mermaid_str, filepath):
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        return True
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False


def generate_bab6_pulau_compact():
    print("[1/3] Membangun dokumen compact Bab 6 Pulau (Format 1-Kolom, 2-3 Halaman)...")
    
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab6    = out_dir_compact.parent.parent / "bab_6"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab6.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER DOKUMEN ──────────────────────────────────────────
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after  = Pt(1)
    add_run(p_t, "RINGKASAN EKSEKUTIF METODOLOGIS", bold=True, pt=8.5, color=G_MID)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after  = Pt(2)
    para_border_bottom(p_h, color='1B5E20', sz='12')
    add_run(p_h, "BAB 6: AUDIT FORENSIK METODOLOGI D3TLH (LEVEL BIOREGION PULAU)", bold=True, pt=15, color=G_DARK)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(1)
    p_meta.paragraph_format.space_after  = Pt(5)
    add_run(p_meta, "Studi Daya Dukung & Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi · ", italic=True, pt=8, color=C_GREY)
    add_run(p_meta, "Center of Economic and Law Studies (CELIOS)", bold=True, italic=True, pt=8, color=G_DARK)

    # ── A. DESAIN PENELITIAN & TUJUAN ───────────────────────────
    add_h2(doc, "A", "Desain Penelitian & Tujuan")
    add_body(doc, [
        ("Penelitian ini menggunakan ", False, False),
        ("desain audit forensik metodologis, pembuktian terbalik berbasis data empiris, dan agregasi Simple Additive Weighting (SAW) kuantitatif terpadu", True, False),
        (" untuk menguji secara kritis keabsahan dokumen Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) resmi di tingkat bioregion makro Pulau Sulawesi sepanjang kurun waktu pengamatan (", False, False),
        ("2014–2024", True, False),
        ("). Tiga tujuan utama metodologis Bab 6 (Level Pulau) meliputi:", False, False)
    ])
    add_body(doc, [
        ("1. ", True, False), ("Dekonstruksi Bias Jasa Ekosistem Pasif: ", True, False),
        ("Membongkar kelemahan pemodelan D3TLH resmi yang semata-mata mengandalkan tutupan vegetasi hutan statis tanpa memperhitungkan beban emisi PLTU captive, konsentrasi gas NO2 satelit, timbulan B3, dan morbiditas ISPA.\n", False, False),
        ("2. ", True, False), ("Audit Lima Dimensi Daya Lentur Ekologis Pulau: ", True, False),
        ("Mengoperasionalkan skoring kuantitatif terhadap 5 pilar daya lentur bioregion: Udara, Air, Lahan, Sosial, dan Veto Kebijakan menggunakan ambang batas baku mutu nasional (PP 22/2021, PermenLHK 27/2021, UU 41/1999) dan standar global (WHO, GEM, IFC PS7).\n", False, False),
        ("3. ", True, False), ("Formulasi Indeks Komposit & Pembuktian Status Kolaps: ", True, False),
        ("Mengagregasikan skor multi-metrik kontinu (0–10) dan skala Likert diskrit (1–5) guna menetapkan status ambang batas ekologis Pulau Sulawesi sebagai landasan mandat moratorium izin eksploitasi.", False, False)
    ])

    # ── B. SUMBER DATA & CAKUPAN WILAYAH ─────────────────────────
    add_h2(doc, "B", "Sumber Data & Cakupan Wilayah")
    add_body(doc, [
        ("Kajian audit bioregion pulau mengintegrasikan enam klaster data resmi kementerian teknis, observasi satelit resolusi tinggi, dan basis data independen terverifikasi:", False, False)
    ])
    add_body(doc, [
        ("• ", True, False), ("Global Energy Monitor (GEM 2023) & Ditjen Minerba ESDM: ", True, False),
        ("Inventarisasi 10,26 GW (10.255 MW) kapasitas PLTU captive batubara dan registri 574 IUP nikel baru se-Sulawesi.\n", False, False),
        ("• ", True, False), ("Satelit Copernicus Sentinel-5P (NASA/ESA TROPOMI): ", True, False),
        ("Pengukuran densitas konsentrasi troposferik nitrogen dioksida (NO2 rasio µmol/m²) di atas kawasan industri nikel.\n", False, False),
        ("• ", True, False), ("Kementerian Kesehatan RI & Dinas Kesehatan Provinsi: ", True, False),
        ("Data epidemiologis insidensi ISPA dan Diare (Incidence Rate Ratio / IRR) serta evaluasi kelayakan sarana-prasarana faskes (ASPAK SPA).\n", False, False),
        ("• ", True, False), ("Kementerian Lingkungan Hidup dan Kehutanan (KLHK): ", True, False),
        ("Indeks Kualitas Air (IKA Ditjen PPKL), neraca timbulan limbah B3 (Ditjen PSLB3), dan dokumen batas daya dukung AMDAL.\n", False, False),
        ("• ", True, False), ("Global Forest Watch (GFW / Hansen UMD) & BNPB: ", True, False),
        ("Time-series kehilangan 1,38 juta Ha tutupan hutan, emisi 804 juta ton CO2e, perambahan hutan lindung, dan 1.609 insiden bencana hidrometeorologi.\n", False, False),
        ("• ", True, False), ("Konsorsium Pembaruan Agraria (KPA) & Koalisi Sipil (JATAM, WALHI): ", True, False),
        ("Dokumentasi 8 kasus manipulasi persetujuan awal (FPIC), 54.310 jiwa korban konflik agraria (505.192 Ha), dan catatan represi aparat.", False, False)
    ])

    # ── C. OPERASIONALISASI VARIABEL & INDIKATOR RISET ──────────
    add_h2(doc, "C", "Operasionalisasi Variabel & Indikator Riset")
    add_body(doc, [
        ("Seluruh parameter bio-fisik cerobong, neraca kualitas perairan, kerusakan tutupan lahan, kerentanan hak sosial, hingga instrumen pembatasan izin dioperasionalkan secara terstruktur ke dalam ", False, False),
        ("indikator riset empiris", True, False),
        (" sebagaimana dirangkum pada matriks operasional berikut:", False, False)
    ])

    add_caption(doc, "Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 6 (Level Pulau)")
    table_indikator_data = [
        ["1", "Kapasitas Emisi PLTU Captive", "Beban Pembakaran Batubara Industri Off-Grid", "Megawatt (MW)", "Global Energy Monitor (GEM)"],
        ["2", "Densitas Polutan NO2 Satelit", "Konsentrasi Troposferik Gas NO2 Atmosfer", "µmol / m²", "Satelit Sentinel-5P TROPOMI"],
        ["3", "Rasio Morbiditas ISPA (IRR)", "Anomali Morbiditas Saluran Pernapasan Warga", "Rasio Peluang (IRR)", "Kemenkes RI & Profil Dinkes"],
        ["4", "Status Mutu Air Sungai (IKA)", "Kondisi Baku Mutu Perairan Regional Pulau", "Poin Indeks (0–100)", "Ditjen PPKL KLHK (IKLH)"],
        ["5", "Beban Timbulan Tailing & Slag B3", "Akumulasi Limbah Pirometalurgi & HPAL", "Juta Ton / Tahun", "Amdal KLHK, AEER & WALHI"],
        ["6", "Laju Deforestasi & Bencana Lahan", "Kehilangan Hutan Primer & Kejadian Banjir", "Ha & Kejadian", "GFW Hansen & DIBI BNPB"],
        ["7", "Pelanggaran Hutan Lindung (Zero Tol)", "Perambahan Kawasan Lindung dan Konservasi", "Hektar (Ha)", "GFW Overlay Kawasan Lindung"],
        ["8", "Skala Pelanggaran Persetujuan FPIC", "Manipulasi Konsultasi Masyarakat Adat/Lokal", "Kasus", "KPA, JATAM & WALHI"],
        ["9", "Defisit Sarana-Prasarana Faskes", "Kesenjangan Standar SPA Puskesmas Tapak", "Persen Kesenjangan (%)", "ASPAK Kemenkes RI"],
        ["10", "Obral Perizinan Baru & Impunitas", "Penerbitan IUP Baru & Pembiaran Korporasi", "Unit Izin & Korporasi", "ESDM MODI & CATAHU KPA"]
    ]

    add_table_styled(
        doc,
        headers=["No", "Indikator Riset", "Fokus Pengukuran", "Satuan", "Sumber Data Primer Resmi"],
        rows=table_indikator_data,
        col_widths_cm=[0.8, 4.5, 4.5, 2.2, 5.0],
        alignments=['C', 'L', 'L', 'C', 'L']
    )

    # ── D. KERANGKA ANALISIS & FORMULASI MATEMATIS ──────────────
    add_h2(doc, "D", "Kerangka Analisis & Formulasi Matematis")

    # 6.1
    add_h3(doc, "6.1 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Udara")
    add_body(doc, [
        ("Daya tampung udara dinilai melalui 4 sub-metrik pembuktian terbalik berbasis Simple Additive Weighting (SAW): kombinasi kapasitas PLTU dan densitas NO2 satelit, anomali morbiditas ISPA klinis, rasio timbulan limbah B3, dan defisit emisi karbon FOLU Net Sink 2030:", False, False)
    ])
    add_formula(doc, "Skor_PLTU = min(5, [Cap / 5000] × 5)   ;   Skor_NO2 = min(5, [(NO2 - 4e-6) / 2e-6] × 5)   |   Skor_Udara1 = min(10, Skor_PLTU + Skor_NO2)\n"
                     "Skor_Udara2 (ISPA) = min(10, [IRR - 1] × 10)   ;   Skor_Udara3 (B3) = min(10, [B3 / 5] × 10)   ;   Skor_Udara4 (CO2) = min(10, [CO2 / 150] × 10)\n"
                     "Skor_Akumulasi_Udara = Σ [ Skor_Udara_i ] / 4.0   |   Skor Likert = Skor_Akumulasi_Udara / 2.0",
                ket="Cap = Kapasitas PLTU (9.825 MW); NO2 = Densitas satelit (5,56e-6 mol/m²); IRR = Rasio insidensi ISPA (3,50x); B3 = Pangsa limbah B3 (7,93%); CO2 = Pelepasan emisi (804,05 Juta Ton); Skor Udara = 9,73 / 10 (Likert 4,9 / 5,0: Kapasitas Asimilasi Habis).")

    # 6.2
    add_h3(doc, "6.2 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Air")
    add_body(doc, [
        ("Daya tampung air diukur berdasarkan defisit baku mutu IKA perairan terhadap ambang batas PermenLHK 27/2021, risiko morbiditas diare klinis, konflik perampasan ruang pesisir nelayan, dan beban timbulan residu tailing HPAL/slag nikel terhadap daya tampung AMDAL:", False, False)
    ])
    add_formula(doc, "Skor_Air1 (IKA) = min(10, max(0, [80 - IKA] / 30) × 10)   ;   Skor_Air2 (Diare) = min(10, [IRR - 1] × 10)\n"
                     "Skor_Air3 (Konflik) = min(10, [Kasus / 15] × 10)   ;   Skor_Air4 (Tailing) = min(10, [Tailing / 25] × 10)\n"
                     "Skor_Akumulasi_Air = Σ [ Skor_Air_i ] / 4.0   |   Skor Likert = Skor_Akumulasi_Air / 2.0",
                ket="IKA = Rata-rata mutu air (59,69 / Kategori Sedang); IRR = Rasio insidensi diare (1,52x); Konflik = Sengketa pesisir nelayan (15 kasus); Tailing = Residu tailing/slag (32,0 Jt Ton/Thn); Skor Air = 8,19 / 10 (Likert 4,2 / 5,0: Penetralan Limbah Melampaui Batas).")

    # 6.3
    add_h3(doc, "6.3 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Lahan")
    add_body(doc, [
        ("Daya dukung lahan dievaluasi menggunakan normalisasi Z-Score frekuensi bencana alam BNPB, deforestasi primer terhadap kuota iklim FOLU 2030, aturan nol-toleransi perambahan hutan lindung (UU 41/1999), dominasi korporasi tambang/sawit, dan rasio konsesi IUP nikel daratan:", False, False)
    ])
    add_formula(doc, "Skor_Lahan1 = min(10, [Bencana / 877] × 10)   ;   Skor_Lahan2 = min(10, [Loss / 638000] × 10)   ;   Skor_Lahan3 = 10 if Loss_Lindung > 0 else 0\n"
                     "Skor_Lahan4 = min(10, [Driver / 500000] × 10)   ;   Skor_Lahan5 = min(10, [Rasio_IUP / 0.10] × 10)\n"
                     "Skor_Akumulasi_Lahan = Σ [ Skor_Lahan_i ] / 5.0   |   Skor Likert = Skor_Akumulasi_Lahan / 2.0",
                ket="Bencana = Banjir/longsor (1.609 kejadian); Loss = Deforestasi (1.386.055 Ha); Loss_Lindung = Hutan lindung hilang (41.785 Ha); Driver = Monopoli tambang/sawit (1.001.654 Ha); Rasio_IUP = 6,3% daratan (1,18 Jt Ha); Skor Lahan = 9,25 / 10 (Likert 4,6 / 5,0: Darurat Lahan).")

    # 6.4
    add_h3(doc, "6.4 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Sosial")
    add_body(doc, [
        ("Daya dukung sosial mengukur batas toleransi kedaulatan warga melalui manipulasi asas FPIC (standar IFC PS7), skala perampasan ruang hidup demografis, insidensi represi kriminalisasi pejuang hak lingkungan, dan kesenjangan pemenuhan sarana-prasarana kesehatan (SPA):", False, False)
    ])
    add_formula(doc, "Skor_Sosial1 = min(10, [FPIC / 3] × 10)   ;   Skor_Sosial2 = min(10, [Jiwa / 40000] × 10)\n"
                     "Skor_Sosial3 = min(10, [Kasus / 10] × 10)   ;   Skor_Sosial4 = min(10, [Gap_SPA / 45] × 10)\n"
                     "Skor_Akumulasi_Sosial = Σ [ Skor_Sosial_i ] / 4.0   |   Skor Likert = Skor_Akumulasi_Sosial / 2.0",
                ket="FPIC = Pelanggaran persetujuan awal (8 kasus); Jiwa = Korban sengketa agraria (54.310 jiwa); Kasus = Represi/kriminalisasi (21 insiden); Gap_SPA = Defisit kelayakan Puskesmas (5,65% di bawah target 80%); Skor Sosial = 7,81 / 10 (Likert 3,9 / 5,0: Perlu Pengawasan).")

    # 6.5
    add_h3(doc, "6.5 Algoritma Skoring Bioregion Pulau: Matriks Veto Kebijakan & Sintesis Komposit")
    add_body(doc, [
        ("Matriks Veto Kebijakan menguji efektivitas fungsi kontrol hukum (Pasal 12 UU 32/2009) terhadap penerbitan izin baru di zona kritis, pembiaran korporasi pelanggar hukum, dan ekspansi PLTU captive, yang kemudian disintesiskan ke dalam Indeks Komposit Bioregion Pulau Sulawesi:", False, False)
    ])
    add_formula(doc, "Skor_Veto1 = min(10, [IUP / 100] × 10)   ;   Skor_Veto2 = min(10, [Korporat / 10] × 10)   ;   Skor_Veto3 = min(10, [MW / 5000] × 10)\n"
                     "Skor_Akumulasi_Veto = Σ [ Skor_Veto_i ] / 3.0   |   Skor_Komposit_Pulau = [ Σ Skor_Dimensi (1..5) ] / 5.0",
                ket="IUP = Izin baru pasca-2014 (574 izin); Korporat = Entitas pelanggar hukum (21 korporasi); MW = PLTU captive (10.255 MW); Skor Veto = 10,00 / 10 (Likert 5,0 / 5,0); Skor Komposit Pulau = 9,00 / 10 (Likert 4,5 / 5,0: Kolaps Daya Dukung Sistemik).")

    add_caption(doc, "Tabel 6.5a: Rekapitulasi Sintesis Skoring 5 Dimensi Bioregion Pulau Sulawesi")
    tabel_6_5a_rows = [
        ["Dimensi 1: Udara", "Kapasitas Asimilasi Habis", "9.825 MW PLTU, NO2 5,56e-6, ISPA IRR 3,5x, B3 7,93%", "9.73 / 10.0", "4.9 / 5.0", "Darurat Udara"],
        ["Dimensi 2: Air", "Penetralan Limbah Terlampaui", "IKA 59,69 (Sedang), Diare IRR 1,5x, Tailing 32 Jt Ton", "8.19 / 10.0", "4.2 / 5.0", "Darurat Air"],
        ["Dimensi 3: Lahan", "Evaluasi Pengelolaan Lanskap", "1.609 Bencana, Deforestasi 1,38 Jt Ha, Lindung 41 Ribu Ha", "9.25 / 10.0", "4.6 / 5.0", "Darurat Lahan"],
        ["Dimensi 4: Sosial", "Pelibatan Masyarakat Lokal", "8 Kasus FPIC, 54.310 Jiwa Tergusur, 21 Represi HAM", "7.81 / 10.0", "3.9 / 5.0", "Perlu Pengawasan"],
        ["Dimensi 5: Veto", "Penguatan Pengawasan Kebijakan", "574 IUP Baru, 21 Korporasi Ilegal, 10,26 GW PLTU", "10.00 / 10.0", "5.0 / 5.0", "Perlu Reformasi"],
        ["TOTAL BIOREGION", "STATUS D3TLH MAKRO SULAWESI", "Agregasi 5 Pilar Daya Dukung & Daya Tampung Ekologis", "9.00 / 10.0", "4.5 / 5.0", "KOLAPS SISTEMIK"]
    ]
    add_table_styled(
        doc,
        headers=["Dimensi Evaluasi", "Status Ekologis Dashboard", "Kondisi Aktual Empiris Terukur", "Skor WSM", "Skor Likert", "Vonis D3TLH"],
        rows=tabel_6_5a_rows,
        col_widths_cm=[2.8, 3.8, 4.8, 1.8, 1.8, 2.0],
        alignments=['L', 'L', 'L', 'C', 'C', 'C']
    )

    # ── E. KORESPONDENSI METODOLOGI TERHADAP SUB-BAB LAPORAN ────
    add_h2(doc, "E", "Korespondensi Metodologi terhadap Sub-bab Laporan Bab 6")
    add_body(doc, [
        ("Setiap sub-bab analitis pada Bab 6 (Level Pulau) ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:", False, False)
    ])

    table_korespondensi = [
        ["Sub-bab 6.1", "Daya Tampung Udara Bioregion", "Multi-Pollutant Aggregation, Threshold Normalization, Simple Additive Weighting (SAW)"],
        ["Sub-bab 6.2", "Daya Tampung Air & Beban Tailing", "Water Quality Deficit Modeling, Epidemiological IRR Mapping, Assimilation Capacity Audit"],
        ["Sub-bab 6.3", "Daya Dukung Lahan & Lanskap", "Geospatial Z-Score Disaster Normalization, FOLU Net Sink Deficit, Zero-Tolerance Rule"],
        ["Sub-bab 6.4", "Daya Dukung Sosial & Hak Tenurial", "Human Rights Violation Tracking, Population Displacement Scaling, Health Facility Gap Audit"],
        ["Sub-bab 6.5", "Veto Kebijakan & Komposit Pulau", "Policy Failure Index, Regulatory Impunity Audit, 5-Dimension Bioregional Synthesis"]
    ]

    add_table_styled(
        doc,
        headers=["Sub-bab", "Fokus Kajian Empiris", "Metode Analitis Utama"],
        rows=table_korespondensi,
        col_widths_cm=[2.5, 5.5, 9.0],
        alignments=['C', 'L', 'L']
    )

    # ── F. BAGAN ALUR KERANGKA KERJA RISET BAB 6 PULAU ──────────
    add_h2(doc, "F", "Bagan Alur Kerangka Kerja Riset (Research Workflow)")
    add_body(doc, [
        ("Kerangka operasional metodologi Bab 6 (Level Pulau) berjalan secara terpadu melalui empat fase berurutan sebagaimana divisualisasikan pada bagan alur kerja riset berikut:", False, False)
    ])

    mermaid_str_f = """flowchart LR
    subgraph F1["Fase I: Input Multi-Pilar"]
        A1["PLTU & Satelit NO2<br/><i>GEM 9.825 MW & TROPOMI</i>"]
        A2["IKA & Tailing HPAL<br/><i>Mutu Air 59.69 & 32 Jt Ton</i>"]
        A3["Deforestasi & Bencana<br/><i>1.38 Jt Ha & 1.609 Banjir</i>"]
        A4["FPIC & Veto Kebijakan<br/><i>8 FPIC & 574 Izin Baru</i>"]
    end
    subgraph F2["Fase II: Thresholding Regulasi"]
        B1["Baku Mutu Udara PP 22/2021<br/><i>Target ISPA IRR WHO</i>"]
        B2["Baku Mutu Air PermenLHK 27<br/><i>AMDAL Tailing Dam</i>"]
        B3["Aturan UU 41/1999 & FOLU<br/><i>Target Iklim NDC 2030</i>"]
        B4["Standar HAM IFC PS7<br/><i>Veto Kuota Izin ESDM</i>"]
    end
    subgraph F3["Fase III: Kalkulasi 5 Matriks"]
        C1["Skor Udara: 9.73 / 10.0<br/><i>(Likert: 4.9 / 5.0)</i>"]
        C2["Skor Air: 8.19 / 10.0<br/><i>(Likert: 4.2 / 5.0)</i>"]
        C3["Skor Lahan: 9.25 / 10.0<br/><i>(Likert: 4.6 / 5.0)</i>"]
        C4["Skor Sosial: 7.81 / 10.0<br/><i>(Likert: 3.9 / 5.0)</i>"]
        C5["Skor Veto: 10.00 / 10.0<br/><i>(Likert: 5.0 / 5.0)</i>"]
    end
    subgraph F4["Fase IV: Agregasi Bioregion"]
        D1["Simple Additive Weighting<br/><i>Bobot Equal 20% per Dimensi</i>"]
        D2["Skor Komposit: 4.5 / 5.0<br/><i>(Skor WSM 9.00 / 10.0)</i>"]
        D3["Vonis: KOLAPS SISTEMIK<br/><i>Mandat Moratorium Total</i>"]
    end
    F1 --> F2 --> F3 --> F4"""

    png_workflow_path = str(out_dir_compact / "mermaid_workflow_bab6_pulau.png")
    is_downloaded = download_mermaid_png(mermaid_str_f, png_workflow_path)

    add_caption(doc, "Bagan Alur 6.1: Alur Logika Kerangka Kerja Riset Bab 6 Pulau (Research Workflow)")
    if is_downloaded and os.path.exists(png_workflow_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(3)
        p_img.paragraph_format.space_after  = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(png_workflow_path, width=Cm(16.5))
        try:
            shutil.copyfile(png_workflow_path, str(out_dir_bab6 / "mermaid_workflow_bab6_pulau.png"))
        except Exception:
            pass

    # Box Output Kesimpulan
    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after  = Pt(4)
    all_border_para(p_box, color='1B5E20', sz='8')
    para_shd(p_box, 'F1F8E9')
    add_run(p_box, "KERANGKA KELUARAN METODOLOGIS BAB 6 (LEVEL BIOREGION PULAU):\n", bold=True, pt=8.5, color=G_DARK)
    add_run(p_box, "1. Dekonstruksi Ilmiah D3TLH Pasif: Membuktikan bahwa pemodelan resmi D3TLH pemerintah menyembunyikan krisis lingkungan akut akibat pengabaian variabel cerobong PLTU (9.825 MW), densitas polutan satelit NO2, residu tailing/slag (32 Jt Ton), dan lonjakan insidensi penyakit saluran napas (ISPA IRR 3,5x).\n"
                   "2. Bukti Pelanggaran Batas Daya Lentur Fisik & Sosial: Mengonfirmasi status darurat pada pilar Udara (Skor 4,9/5), Air (Skor 4,2/5), Lahan (Skor 4,6/5), serta kerentanan hak tenurial sosial (Skor 3,9/5) yang mengorbankan 54.310 jiwa warga terdampak.\n"
                   "3. Vonis Kolaps Sistemik & Mandat Moratorium: Menghasilkan Skor Komposit Bioregion Pulau sebesar 4,5 / 5,0 (Status: Melampaui Batas / Kolaps Ekologis Sistemik), membuktikan bahwa seluruh ruang hidup Pulau Sulawesi telah kehilangan kapasitas pemulihan alami dan menuntut pemberlakuan moratorium izin eksploitasi baru secara mutlak.",
            pt=8, color=C_BODY)

    # ── SIMPAN DOKUMEN DOCX (DUAL SAVE) ─────────────────────────
    docx_compact = out_dir_compact / "Metodologi_Bab6_Pulau_Audit_D3TLH_Compact.docx"
    docx_bab6    = out_dir_bab6 / "Metodologi_Bab6_Pulau_Audit_D3TLH_Compact.docx"
    
    doc.save(str(docx_compact))
    shutil.copyfile(docx_compact, docx_bab6)
    print(f"  [OK] Tersimpan DOCX: {docx_compact}")
    print(f"  [OK] Salinan DOCX : {docx_bab6}")

    # ── GENERATE MARKDOWN PADANAN ───────────────────────────────
    print("[2/3] Membangun dokumen Markdown padanan...")
    MD_CONTENT = """# METODOLOGI PENELITIAN: BAB 6 — AUDIT FORENSIK D3TLH (LEVEL BIOREGION PULAU)
*CELIOS (Center of Economic and Law Studies) · Audit Spasial-Statistik D3TLH Sulawesi (2014–2024) · Ringkasan Eksekutif Metodologis*

---

## A. Desain Penelitian & Tujuan
Penelitian ini menggunakan **desain audit forensik metodologis, pembuktian terbalik berbasis data empiris, dan agregasi Simple Additive Weighting (SAW) kuantitatif terpadu** untuk menguji secara kritis keabsahan dokumen Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) resmi di tingkat bioregion makro Pulau Sulawesi sepanjang kurun waktu pengamatan (**2014–2024**). Tiga tujuan utama metodologis Bab 6 (Level Pulau) meliputi:

1. **Dekonstruksi Bias Jasa Ekosistem Pasif:** Membongkar kelemahan pemodelan D3TLH resmi yang semata-mata mengandalkan tutupan vegetasi hutan statis tanpa memperhitungkan beban emisi PLTU captive, konsentrasi gas NO2 satelit, timbulan B3, dan morbiditas ISPA.
2. **Audit Lima Dimensi Daya Lentur Ekologis Pulau:** Mengoperasionalkan skoring kuantitatif terhadap 5 pilar daya lentur bioregion: Udara, Air, Lahan, Sosial, dan Veto Kebijakan menggunakan ambang batas baku mutu nasional (PP 22/2021, PermenLHK 27/2021, UU 41/1999) dan standar global (WHO, GEM, IFC PS7).
3. **Formulasi Indeks Komposit & Pembuktian Status Kolaps:** Mengagregasikan skor multi-metrik kontinu (0–10) dan skala Likert diskrit (1–5) guna menetapkan status ambang batas ekologis Pulau Sulawesi sebagai landasan mandat moratorium izin eksploitasi.

---

## B. Sumber Data & Cakupan Wilayah
Kajian audit bioregion pulau mengintegrasikan enam klaster data resmi kementerian teknis, observasi satelit resolusi tinggi, dan basis data independen terverifikasi:

- **Global Energy Monitor (GEM 2023) & Ditjen Minerba ESDM:** Inventarisasi 10,26 GW (10.255 MW) kapasitas PLTU captive batubara dan registri 574 IUP nikel baru se-Sulawesi.
- **Satelit Copernicus Sentinel-5P (NASA/ESA TROPOMI):** Pengukuran densitas konsentrasi troposferik nitrogen dioksida (NO2 rasio µmol/m²) di atas kawasan industri nikel.
- **Kementerian Kesehatan RI & Dinas Kesehatan Provinsi:** Data epidemiologis insidensi ISPA dan Diare (Incidence Rate Ratio / IRR) serta evaluasi kelayakan sarana-prasarana faskes (ASPAK SPA).
- **Kementerian Lingkungan Hidup dan Kehutanan (KLHK):** Indeks Kualitas Air (IKA Ditjen PPKL), neraca timbulan limbah B3 (Ditjen PSLB3), dan dokumen batas daya dukung AMDAL.
- **Global Forest Watch (GFW / Hansen UMD) & BNPB:** Time-series kehilangan 1,38 juta Ha tutupan hutan, emisi 804 juta ton CO2e, perambahan hutan lindung, dan 1.609 insiden bencana hidrometeorologi.
- **Konsorsium Pembaruan Agraria (KPA) & Koalisi Sipil (JATAM, WALHI):** Dokumentasi 8 kasus manipulasi persetujuan awal (FPIC), 54.310 jiwa korban konflik agraria (505.192 Ha), dan catatan represi aparat.

---

## C. Operasionalisasi Variabel & Indikator Riset
Seluruh parameter bio-fisik cerobong, neraca kualitas perairan, kerusakan tutupan lahan, kerentanan hak sosial, hingga instrumen pembatasan izin dioperasionalkan secara terstruktur ke dalam **indikator riset empiris** sebagaimana dirangkum pada matriks operasional berikut:

##### Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 6 (Level Pulau)
| No | Indikator Riset | Fokus Pengukuran | Satuan | Sumber Data Primer Resmi |
| :-: | :--- | :--- | :-: | :--- |
| 1 | Kapasitas Emisi PLTU Captive | Beban Pembakaran Batubara Industri Off-Grid | Megawatt (MW) | Global Energy Monitor (GEM) |
| 2 | Densitas Polutan NO2 Satelit | Konsentrasi Troposferik Gas NO2 Atmosfer | µmol / m² | Satelit Sentinel-5P TROPOMI |
| 3 | Rasio Morbiditas ISPA (IRR) | Anomali Morbiditas Saluran Pernapasan Warga | Rasio Peluang (IRR) | Kemenkes RI & Profil Dinkes |
| 4 | Status Mutu Air Sungai (IKA) | Kondisi Baku Mutu Perairan Regional Pulau | Poin Indeks (0–100) | Ditjen PPKL KLHK (IKLH) |
| 5 | Beban Timbulan Tailing & Slag B3 | Akumulasi Limbah Pirometalurgi & HPAL | Juta Ton / Tahun | Amdal KLHK, AEER & WALHI |
| 6 | Laju Deforestasi & Bencana Lahan | Kehilangan Hutan Primer & Kejadian Banjir | Ha & Kejadian | GFW Hansen & DIBI BNPB |
| 7 | Pelanggaran Hutan Lindung (Zero Tol) | Perambahan Kawasan Lindung dan Konservasi | Hektar (Ha) | GFW Overlay Kawasan Lindung |
| 8 | Skala Pelanggaran Persetujuan FPIC | Manipulasi Konsultasi Masyarakat Adat/Lokal | Kasus | KPA, JATAM & WALHI |
| 9 | Defisit Sarana-Prasarana Faskes | Kesenjangan Standar SPA Puskesmas Tapak | Persen Kesenjangan (%) | ASPAK Kemenkes RI |
| 10 | Obral Perizinan Baru & Impunitas | Penerbitan IUP Baru & Pembiaran Korporasi | Unit Izin & Korporasi | ESDM MODI & CATAHU KPA |

---

## D. Kerangka Analisis & Formulasi Matematis

### 6.1 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Udara
Daya tampung udara dinilai melalui 4 sub-metrik pembuktian terbalik berbasis Simple Additive Weighting (SAW): kombinasi kapasitas PLTU dan densitas NO2 satelit, anomali morbiditas ISPA klinis, rasio timbulan limbah B3, dan defisit emisi karbon FOLU Net Sink 2030:

> `Skor_PLTU = min(5, [Cap / 5000] × 5)   ;   Skor_NO2 = min(5, [(NO2 - 4e-6) / 2e-6] × 5)   |   Skor_Udara1 = min(10, Skor_PLTU + Skor_NO2)`  
> `Skor_Udara2 (ISPA) = min(10, [IRR - 1] × 10)   ;   Skor_Udara3 (B3) = min(10, [B3 / 5] × 10)   ;   Skor_Udara4 (CO2) = min(10, [CO2 / 150] × 10)`  
> `Skor_Akumulasi_Udara = Σ [ Skor_Udara_i ] / 4.0   |   Skor Likert = Skor_Akumulasi_Udara / 2.0`  
> *Keterangan: Cap = Kapasitas PLTU (9.825 MW); NO2 = Densitas satelit (5,56e-6 mol/m²); IRR = Rasio insidensi ISPA (3,50x); B3 = Pangsa limbah B3 (7,93%); CO2 = Pelepasan emisi (804,05 Juta Ton); Skor Udara = 9,73 / 10 (Likert 4,9 / 5,0: Kapasitas Asimilasi Habis).*

### 6.2 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Air
Daya tampung air diukur berdasarkan defisit baku mutu IKA perairan terhadap ambang batas PermenLHK 27/2021, risiko morbiditas diare klinis, konflik perampasan ruang pesisir nelayan, dan beban timbulan residu tailing HPAL/slag nikel terhadap daya tampung AMDAL:

> `Skor_Air1 (IKA) = min(10, max(0, [80 - IKA] / 30) × 10)   ;   Skor_Air2 (Diare) = min(10, [IRR - 1] × 10)`  
> `Skor_Air3 (Konflik) = min(10, [Kasus / 15] × 10)   ;   Skor_Air4 (Tailing) = min(10, [Tailing / 25] × 10)`  
> `Skor_Akumulasi_Air = Σ [ Skor_Air_i ] / 4.0   |   Skor Likert = Skor_Akumulasi_Air / 2.0`  
> *Keterangan: IKA = Rata-rata mutu air (59,69 / Kategori Sedang); IRR = Rasio insidensi diare (1,52x); Konflik = Sengketa pesisir nelayan (15 kasus); Tailing = Residu tailing/slag (32,0 Jt Ton/Thn); Skor Air = 8,19 / 10 (Likert 4,2 / 5,0: Penetralan Limbah Melampaui Batas).*

### 6.3 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Lahan
Daya dukung lahan dievaluasi menggunakan normalisasi Z-Score frekuensi bencana alam BNPB, deforestasi primer terhadap kuota iklim FOLU 2030, aturan nol-toleransi perambahan hutan lindung (UU 41/1999), dominasi korporasi tambang/sawit, dan rasio konsesi IUP nikel daratan:

> `Skor_Lahan1 = min(10, [Bencana / 877] × 10)   ;   Skor_Lahan2 = min(10, [Loss / 638000] × 10)   ;   Skor_Lahan3 = 10 if Loss_Lindung > 0 else 0`  
> `Skor_Lahan4 = min(10, [Driver / 500000] × 10)   ;   Skor_Lahan5 = min(10, [Rasio_IUP / 0.10] × 10)`  
> `Skor_Akumulasi_Lahan = Σ [ Skor_Lahan_i ] / 5.0   |   Skor Likert = Skor_Akumulasi_Lahan / 2.0`  
> *Keterangan: Bencana = Banjir/longsor (1.609 kejadian); Loss = Deforestasi (1.386.055 Ha); Loss_Lindung = Hutan lindung hilang (41.785 Ha); Driver = Monopoli tambang/sawit (1.001.654 Ha); Rasio_IUP = 6,3% daratan (1,18 Jt Ha); Skor Lahan = 9,25 / 10 (Likert 4,6 / 5,0: Darurat Lahan).*

### 6.4 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Sosial
Daya dukung sosial mengukur batas toleransi kedaulatan warga melalui manipulasi asas FPIC (standar IFC PS7), skala perampasan ruang hidup demografis, insidensi represi kriminalisasi pejuang hak lingkungan, dan kesenjangan pemenuhan sarana-prasarana kesehatan (SPA):

> `Skor_Sosial1 = min(10, [FPIC / 3] × 10)   ;   Skor_Sosial2 = min(10, [Jiwa / 40000] × 10)`  
> `Skor_Sosial3 = min(10, [Kasus / 10] × 10)   ;   Skor_Sosial4 = min(10, [Gap_SPA / 45] × 10)`  
> `Skor_Akumulasi_Sosial = Σ [ Skor_Sosial_i ] / 4.0   |   Skor Likert = Skor_Akumulasi_Sosial / 2.0`  
> *Keterangan: FPIC = Pelanggaran persetujuan awal (8 kasus); Jiwa = Korban sengketa agraria (54.310 jiwa); Kasus = Represi/kriminalisasi (21 insiden); Gap_SPA = Defisit kelayakan Puskesmas (5,65% di bawah target 80%); Skor Sosial = 7,81 / 10 (Likert 3,9 / 5,0: Perlu Pengawasan).*

### 6.5 Algoritma Skoring Bioregion Pulau: Matriks Veto Kebijakan & Sintesis Komposit
Matriks Veto Kebijakan menguji efektivitas fungsi kontrol hukum (Pasal 12 UU 32/2009) terhadap penerbitan izin baru di zona kritis, pembiaran korporasi pelanggar hukum, dan ekspansi PLTU captive, yang kemudian disintesiskan ke dalam Indeks Komposit Bioregion Pulau Sulawesi:

> `Skor_Veto1 = min(10, [IUP / 100] × 10)   ;   Skor_Veto2 = min(10, [Korporat / 10] × 10)   ;   Skor_Veto3 = min(10, [MW / 5000] × 10)`  
> `Skor_Akumulasi_Veto = Σ [ Skor_Veto_i ] / 3.0   |   Skor_Komposit_Pulau = [ Σ Skor_Dimensi (1..5) ] / 5.0`  
> *Keterangan: IUP = Izin baru pasca-2014 (574 izin); Korporat = Entitas pelanggar hukum (21 korporasi); MW = PLTU captive (10.255 MW); Skor Veto = 10,00 / 10 (Likert 5,0 / 5,0); Skor Komposit Pulau = 9,00 / 10 (Likert 4,5 / 5,0: Kolaps Daya Dukung Sistemik).*

##### Tabel 6.5a: Rekapitulasi Sintesis Skoring 5 Dimensi Bioregion Pulau Sulawesi
| Dimensi Evaluasi | Status Ekologis Dashboard | Kondisi Aktual Empiris Terukur | Skor WSM | Skor Likert | Vonis D3TLH |
| :--- | :--- | :--- | :---: | :---: | :---: |
| Dimensi 1: Udara | Kapasitas Asimilasi Habis | 9.825 MW PLTU, NO2 5,56e-6, ISPA IRR 3,5x, B3 7,93% | 9.73 / 10.0 | 4.9 / 5.0 | Darurat Udara |
| Dimensi 2: Air | Penetralan Limbah Terlampaui | IKA 59,69 (Sedang), Diare IRR 1,5x, Tailing 32 Jt Ton | 8.19 / 10.0 | 4.2 / 5.0 | Darurat Air |
| Dimensi 3: Lahan | Evaluasi Pengelolaan Lanskap | 1.609 Bencana, Deforestasi 1,38 Jt Ha, Lindung 41 Ribu Ha | 9.25 / 10.0 | 4.6 / 5.0 | Darurat Lahan |
| Dimensi 4: Sosial | Pelibatan Masyarakat Lokal | 8 Kasus FPIC, 54.310 Jiwa Tergusur, 21 Represi HAM | 7.81 / 10.0 | 3.9 / 5.0 | Perlu Pengawasan |
| Dimensi 5: Veto | Penguatan Pengawasan Kebijakan | 574 IUP Baru, 21 Korporasi Ilegal, 10,26 GW PLTU | 10.00 / 10.0 | 5.0 / 5.0 | Perlu Reformasi |
| **TOTAL BIOREGION** | **STATUS D3TLH MAKRO SULAWESI** | **Agregasi 5 Pilar Daya Dukung & Daya Tampung Ekologis** | **9.00 / 10.0** | **4.5 / 5.0** | **KOLAPS SISTEMIK** |

---

## E. Korespondensi Metodologi terhadap Sub-bab Laporan Bab 6
Setiap sub-bab analitis pada Bab 6 (Level Pulau) ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:

##### Matriks Korespondensi Metodologis Bab 6 (Level Pulau)
| Sub-bab | Fokus Kajian Empiris | Metode Analitis Utama |
| :-: | :--- | :--- |
| Sub-bab 6.1 | Daya Tampung Udara Bioregion | Multi-Pollutant Aggregation, Threshold Normalization, Simple Additive Weighting (SAW) |
| Sub-bab 6.2 | Daya Tampung Air & Beban Tailing | Water Quality Deficit Modeling, Epidemiological IRR Mapping, Assimilation Capacity Audit |
| Sub-bab 6.3 | Daya Dukung Lahan & Lanskap | Geospatial Z-Score Disaster Normalization, FOLU Net Sink Deficit, Zero-Tolerance Rule |
| Sub-bab 6.4 | Daya Dukung Sosial & Hak Tenurial | Human Rights Violation Tracking, Population Displacement Scaling, Health Facility Gap Audit |
| Sub-bab 6.5 | Veto Kebijakan & Komposit Pulau | Policy Failure Index, Regulatory Impunity Audit, 5-Dimension Bioregional Synthesis |

---

## F. Bagan Alur Kerangka Kerja Riset (Research Workflow)
Kerangka operasional metodologi Bab 6 (Level Pulau) berjalan secara terpadu melalui empat fase berurutan sebagaimana divisualisasikan pada bagan alur kerja riset berikut:

```mermaid
flowchart LR
    subgraph F1["Fase I: Input Multi-Pilar"]
        A1["PLTU & Satelit NO2<br/><i>GEM 9.825 MW & TROPOMI</i>"]
        A2["IKA & Tailing HPAL<br/><i>Mutu Air 59.69 & 32 Jt Ton</i>"]
        A3["Deforestasi & Bencana<br/><i>1.38 Jt Ha & 1.609 Banjir</i>"]
        A4["FPIC & Veto Kebijakan<br/><i>8 FPIC & 574 Izin Baru</i>"]
    end
    subgraph F2["Fase II: Thresholding Regulasi"]
        B1["Baku Mutu Udara PP 22/2021<br/><i>Target ISPA IRR WHO</i>"]
        B2["Baku Mutu Air PermenLHK 27<br/><i>AMDAL Tailing Dam</i>"]
        B3["Aturan UU 41/1999 & FOLU<br/><i>Target Iklim NDC 2030</i>"]
        B4["Standar HAM IFC PS7<br/><i>Veto Kuota Izin ESDM</i>"]
    end
    subgraph F3["Fase III: Kalkulasi 5 Matriks"]
        C1["Skor Udara: 9.73 / 10.0<br/><i>(Likert: 4.9 / 5.0)</i>"]
        C2["Skor Air: 8.19 / 10.0<br/><i>(Likert: 4.2 / 5.0)</i>"]
        C3["Skor Lahan: 9.25 / 10.0<br/><i>(Likert: 4.6 / 5.0)</i>"]
        C4["Skor Sosial: 7.81 / 10.0<br/><i>(Likert: 3.9 / 5.0)</i>"]
        C5["Skor Veto: 10.00 / 10.0<br/><i>(Likert: 5.0 / 5.0)</i>"]
    end
    subgraph F4["Fase IV: Agregasi Bioregion"]
        D1["Simple Additive Weighting<br/><i>Bobot Equal 20% per Dimensi</i>"]
        D2["Skor Komposit: 4.5 / 5.0<br/><i>(Skor WSM 9.00 / 10.0)</i>"]
        D3["Vonis: KOLAPS SISTEMIK<br/><i>Mandat Moratorium Total</i>"]
    end
    F1 --> F2 --> F3 --> F4
```

> **KERANGKA KELUARAN METODOLOGIS BAB 6 (LEVEL BIOREGION PULAU):**  
> 1. **Dekonstruksi Ilmiah D3TLH Pasif:** Membuktikan bahwa pemodelan resmi D3TLH pemerintah menyembunyikan krisis lingkungan akut akibat pengabaian variabel cerobong PLTU (9.825 MW), densitas polutan satelit NO2, residu tailing/slag (32 Jt Ton), dan lonjakan insidensi penyakit saluran napas (ISPA IRR 3,5x).  
> 2. **Bukti Pelanggaran Batas Daya Lentur Fisik & Sosial:** Mengonfirmasi status darurat pada pilar Udara (Skor 4,9/5), Air (Skor 4,2/5), Lahan (Skor 4,6/5), serta kerentanan hak tenurial sosial (Skor 3,9/5) yang mengorbankan 54.310 jiwa warga terdampak.  
> 3. **Vonis Kolaps Sistemik & Mandat Moratorium:** Menghasilkan Skor Komposit Bioregion Pulau sebesar 4,5 / 5,0 (Status: Melampaui Batas / Kolaps Ekologis Sistemik), membuktikan bahwa seluruh ruang hidup Pulau Sulawesi telah kehilangan kapasitas pemulihan alami dan menuntut pemberlakuan moratorium izin eksploitasi baru secara mutlak.
"""

    md_compact = out_dir_compact / "Metodologi_Bab6_Pulau_Audit_D3TLH_Compact.md"
    md_bab6    = out_dir_bab6 / "Metodologi_Bab6_Pulau_Audit_D3TLH_Compact.md"
    with open(md_compact, 'w', encoding='utf-8') as f:
        f.write(MD_CONTENT)
    shutil.copyfile(md_compact, md_bab6)
    print(f"  [OK] Tersimpan MD  : {md_compact}")
    print(f"  [OK] Salinan MD   : {md_bab6}")

    print("[3/3] Selesai menghasilkan dokumen metodologi Bab 6 Pulau versi compact (1-Kolom, 2-3 Halaman).\n")


if __name__ == "__main__":
    generate_bab6_pulau_compact()
