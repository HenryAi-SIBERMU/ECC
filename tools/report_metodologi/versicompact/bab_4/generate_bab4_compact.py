#!/usr/bin/env python3
"""
Generator Metodologi Versi Compact Bab 4 — GAYA AKADEMIS TERPADU (CELIOS)
Mengadopsi arsitektur metodologi ringkas terstandarisasi konsisten dengan Bab 1, Bab 2, dan Bab 3:
- FORMAT: 1 KOLOM PENUH (Single Column Layout)
- PANJANG: 2–3 Halaman Maksimal (Elegan, proporsional, tanpa pemadatan berlebihan)
- PENOMORAN SEKSI UTAMA: Huruf kapital A, B, C, D, E, F
- SUB-BAB SEKSI D: 4.1, 4.2, 4.3, 4.4, 4.5 (Judul persis dokumen induk)
- OPERASIONALISASI INDIKATOR: 10 Indikator Empiris Lengkap (Matriks Indikator & Sumber Data Resmi)
- FORMULASI & TABEL CROSSTAB: Format standar Tabel 1.5b dengan keterangan definisi variabel lengkap
- KORESPONDENSI METODOLOGI: 3 kolom bersih (Sub-bab, Fokus Kajian Empiris, Metode Analitis Utama)
- FLOWCHART: Mermaid JS horizontal (flowchart LR) dirender tajam ke DOCX (16.5 cm) dan blok kode di MD
- SINKRONISASI: Dual-save ke direktori versicompact/bab_4 dan bab_4.
"""

import os
import sys
import shutil
import base64
import requests
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Palet Warna Resmi CELIOS ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Forest Dark Green (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Celios Accent Green (#2E7D32)
G_LIGHT = RGBColor(0x38, 0x8E, 0x3C)  # Medium Green (#388E3C)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Charcoal Body Text (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Muted Grey Text (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Pure White

# ── Helper Fungsi XML Word ───────────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=80, right=80, top=40, bottom=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)

def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    pPr.append(s)

def para_border_bottom(p, color='1B5E20', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '3')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def para_border_left(p, color='2E7D32', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:left')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='1B5E20', sz='8'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Typography & Content ──────────────────────────────
def add_run(p, text, bold=False, italic=False, pt=8.5, color=C_BODY, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    else:
        r.font.name = 'Calibri'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    return r

def add_h2(doc, prefix, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_border_bottom(p, color='1B5E20', sz='8')
    add_run(p, f"{prefix}.  ", bold=True, pt=11, color=G_DARK)
    add_run(p, title.upper(), bold=True, pt=11, color=G_DARK)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, pt=9.5, color=G_MID)
    return p

def add_body(doc, parts, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        add_run(p, text, bold=bold, italic=italic, pt=8.5, color=C_BODY)
    return p

def add_formula(doc, text, ket=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    para_shd(p, 'EDF7EE')
    add_run(p, text, pt=8, color=G_MID, mono=True)
    if ket:
        add_run(p, f"\nKeterangan: {ket}", italic=True, pt=7.5, color=RGBColor(0x33, 0x55, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_styled(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    bd_subtle = {'val': 'single', 'sz': '4', 'color': 'D0D7DE', 'space': '0'}
    
    # Header Row
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=80, right=80, top=50, bottom=50)
        set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if alignments and j < len(alignments):
            align = alignments[j]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, pt=8, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F9FBF9' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_shd(c, fill)
            cell_margin(c, left=80, right=80, top=40, bottom=40)
            set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            
            if alignments and j < len(alignments):
                align = alignments[j]
                if align == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_run(p, str(val), pt=7.5, color=C_BODY)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after  = Pt(3)
    return tbl

def download_mermaid_png(mermaid_str, filepath):
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        return True
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False


def generate_bab4_compact():
    print("[1/3] Membangun dokumen compact Bab 4 (Format 1-Kolom, 2-3 Halaman)...")
    
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab4    = out_dir_compact.parent.parent / "bab_4"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab4.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER DOKUMEN ──────────────────────────────────────────
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after  = Pt(1)
    add_run(p_t, "RINGKASAN EKSEKUTIF METODOLOGIS", bold=True, pt=8.5, color=G_MID)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after  = Pt(2)
    para_border_bottom(p_h, color='1B5E20', sz='12')
    add_run(p_h, "BAB 4: METODOLOGI ANALISIS RUANG HIDUP YANG TERAMPAS", bold=True, pt=15, color=G_DARK)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(1)
    p_meta.paragraph_format.space_after  = Pt(5)
    add_run(p_meta, "Studi Daya Dukung & Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi · ", italic=True, pt=8, color=C_GREY)
    add_run(p_meta, "Center of Economic and Law Studies (CELIOS)", bold=True, italic=True, pt=8, color=G_DARK)

    # ── A. DESAIN PENELITIAN & TUJUAN ───────────────────────────
    add_h2(doc, "A", "Desain Penelitian & Tujuan")
    add_body(doc, [
        ("Penelitian ini menggunakan ", False, False),
        ("desain sosiologi hukum kritis, audit sengketa agraria, dan analisis inferensial kuantitatif terintegrasi", True, False),
        (" untuk mengidentifikasi perampasan ruang hidup masyarakat adat dan komunitas lokal akibat ekspansi masif korporasi tambang nikel dan kawasan industri penunjang di Pulau Sulawesi sepanjang kurun waktu pengamatan (", False, False),
        ("1990–2024", True, False),
        ("). Tiga tujuan utama metodologis Bab 4 meliputi:", False, False)
    ])
    add_body(doc, [
        ("1. ", True, False), ("Kuantifikasi Asimetri Penguasaan Ruang & Korban Terdampak: ", True, False),
        ("Mengevaluasi distribusi sektoral sengketa agraria guna membuktikan dominasi monopoli lahan dan skala korban masyarakat terdampak pada sektor pertambangan nikel dibandingkan sektor lainnya.\n", False, False),
        ("2. ", True, False), ("Pembuktian Inferensial Kausalitas Ekspansi vs Eskalasi Represi: ", True, False),
        ("Menguji signifikansi hubungan antara periode hilirisasi dan keterlibatan korporasi ekstraktif terhadap peningkatan risiko kekerasan, penangkapan, serta kriminalisasi warga melalui matriks Chi-Square dan Odds Ratio (OR).\n", False, False),
        ("3. ", True, False), ("Pemetaan Entitas Oligarki & Aktor Proksi (NLP Text Parsing): ", True, False),
        ("Mengekstraksi jaringan entitas korporasi dominan dan mendeteksi keterlibatan aktor vigilante/pengamanan swakarsa dalam kronologi sengketa agraria melalui penambangan teks reguler (NLP Regex).", False, False)
    ])

    # ── B. SUMBER DATA & CAKUPAN WILAYAH ─────────────────────────
    add_h2(doc, "B", "Sumber Data & Cakupan Wilayah")
    add_body(doc, [
        ("Penelitian mencakup catatan letupan konflik agraria terdata di seluruh wilayah administratif ", False, False),
        ("Pulau Sulawesi dan pulau-pulau kecil penyangga sentra nikel", True, False),
        (" (seperti Pulau Wawonii dan pesisir Morowali). Basis data dibangun dari repositori dokumentasi konflik agraria nasional dan advokasi masyarakat sipil:", False, False)
    ])
    add_body(doc, [
        ("• ", True, False), ("Konsorsium Pembaruan Agraria (KPA) / Basis Data TanahKita: ", True, False),
        ("Dokumentasi 95 kasus konflik agraria regional Sulawesi (dan korpus 568 narasi nasional) mencakup luas sengketa (Ha), jiwa terdampak, sektor industri, serta status penanganan hukum.\n", False, False),
        ("• ", True, False), ("Koalisi Masyarakat Sipil (WALHI, JATAM, AMAN): ", True, False),
        ("Kronologi advokasi hak tenurial masyarakat adat, kasus kekerasan fisik, dan pemantauan perampasan wilayah kelola rakyat.\n", False, False),
        ("• ", True, False), ("Komisi Nasional Hak Asasi Manusia (Komnas HAM): ", True, False),
        ("Registri pengaduan pelanggaran hak sipil, kriminalisasi pejuang lingkungan, warga ditangkap, luka-luka, dan korban tewas.\n", False, False),
        ("• ", True, False), ("Kementerian ATR/BPN & ESDM (MODI): ", True, False),
        ("Status perizinan hak guna usaha (HGU), izin usaha pertambangan (IUP), dan konsesi kawasan industri.", False, False)
    ])

    # ── C. OPERASIONALISASI VARIABEL & INDIKATOR RISET ──────────
    add_h2(doc, "C", "Operasionalisasi Variabel & Indikator Riset")
    add_body(doc, [
        ("Seluruh dinamika perampasan ruang, eskalasi konflik, dan keterlibatan aktor dioperasionalkan ke dalam ", False, False),
        ("10 indikator empiris terpadu", True, False),
        (" sebagaimana dirangkum pada matriks operasional berikut:", False, False)
    ])

    add_caption(doc, "Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 4")
    table_indikator_data = [
        ["1", "Insidensi Konflik Agraria", "Frekuensi Kejadian Letupan Sengketa Lahan", "Kasus", "1990–2024", "KPA / TanahKita"],
        ["2", "Sektor Pemicu Sengketa", "Klasifikasi Sektoral (Tambang, Sawit, Hutan)", "Kategori", "1990–2024", "KPA / TanahKita"],
        ["3", "Skala Korban Terdampak", "Masyarakat Adat & Komunitas Lokal Terdampak", "Jiwa", "1990–2024", "KPA & Koalisi Sipil"],
        ["4", "Luas Monopoli Area Sengketa", "Luas Ruang Hidup & Lahan Diperebutkan", "Hektar (Ha)", "1990–2024", "TanahKita & ATR/BPN"],
        ["5", "Kasus Indikasi Kriminalisasi", "Penuntutan Hukum Terhadap Warga/Aktivis", "Kasus", "2000–2024", "KPA & Komnas HAM"],
        ["6", "Korban Represi & Kekerasan", "Warga Ditangkap, Mengalami Luka, & Tewas", "Orang", "2000–2024", "KPA & Komnas HAM"],
        ["7", "Laju Eskalasi Hilirisasi", "Rasio Before-After Kasus Pra vs Pasca 2014", "Kasus / Tahun", "1990 vs 2024", "Data Panel Tahunan"],
        ["8", "Tingkat Penelantaran Kasus", "Sengketa Lahan Berstatus Belum Ditangani", "Persen (%)", "1990–2024", "KPA / TanahKita"],
        ["9", "Frekuensi Entitas Korporasi", "Keterlibatan Konglomerasi dalam Konflik", "Token Count", "Korpus NLP", "NLP Regex TanahKita"],
        ["10", "Frekuensi Aktor Proksi/Vigilante", "Keterlibatan Pengamanan Swakarsa/Preman", "Token Count", "Korpus NLP", "NLP Regex TanahKita"]
    ]

    add_table_styled(
        doc,
        headers=["No", "Indikator Riset", "Fokus Pengukuran", "Satuan", "Periode", "Sumber Data Primer Resmi"],
        rows=table_indikator_data,
        col_widths_cm=[0.8, 3.8, 3.8, 1.8, 1.8, 5.0],
        alignments=['C', 'L', 'L', 'C', 'C', 'L']
    )

    # ── D. KERANGKA ANALISIS & FORMULASI MATEMATIS ──────────────
    add_h2(doc, "D", "Kerangka Analisis & Formulasi Matematis")

    # 4.1
    add_h3(doc, "4.1 Tren Eskalasi Konflik Agraria Seiring Ekspansi Industri")
    add_body(doc, [
        ("Kuantifikasi eskalasi sengketa ruang hidup dihitung berdasarkan deret waktu tahunan dan rasio laju peningkatan kasus antara era pra-ekspansi dan era hilirisasi:", False, False)
    ])
    add_formula(doc, "Agregasi Konflik (K_t,s) = Σ c_i   ;   Laju Eskalasi (E %) = [ K_Pasca / K_Pra ] × 100",
                ket="c_i = Kasus konflik agraria i pada tahun t dan sektor s; K_Pasca = Total kasus pasca ekspansi industri; K_Pra = Total kasus pra ekspansi; E = Laju persentase lonjakan eskalasi konflik.")

    # 4.2
    add_h3(doc, "4.2 Sebaran Sektoral: Dampak Masyarakat dan Penggunaan Lahan")
    add_body(doc, [
        ("Dekomposisi beban dampak sosiologis dan monopoli penguasaan ruang dihitung per sektor industri pemicu konflik guna mengukur asimetri dampak:", False, False)
    ])
    add_formula(doc, "Total Jiwa Terdampak (J_s) = Σ J_i   ;   Total Luas Area (A_s) = Σ A_i   ;   Porsi Sektoral (P_s %) = [ Nilai_s / Nilai_Total ] × 100",
                ket="J_i = Warga terdampak pada kasus i; A_i = Luas lahan sengketa kasus i (Ha); J_s & A_s = Total korban jiwa dan luas sengketa sektor s; P_s = Pangsa persentase sektor terhadap total regional.")

    # 4.3
    add_h3(doc, "4.3 Indikasi Represi dan Kriminalisasi dalam Konflik Agraria")
    add_body(doc, [
        ("Kuantifikasi penyempitan ruang sipil menghitung total kasus kriminalisasi serta menjumlahkan seluruh korban represi fisik yang terdokumentasi:", False, False)
    ])
    add_formula(doc, "Total Kriminalisasi = Σ I_i   ;   Total Korban Represi (R) = Σ [ D_i + L_i + T_i ]",
                ket="I_i = Indikator biner kriminalisasi pada kasus i (1 jika ada); D_i = Korban ditangkap; L_i = Korban luka-luka; T_i = Korban tewas; R = Akumulasi total korban pelanggaran HAM.")

    # 4.4
    add_h3(doc, "4.4 Pembuktian Statistik: Ekspansi vs Eskalasi Konflik")
    add_body(doc, [
        ("Pengujian komparatif Before-After dan uji independensi Chi-Square (χ²) tabulasi silang diterapkan pada basis data kejadian konflik (N=523) untuk membuktikan korelasi ekspansi industri terhadap eskalasi represi:", False, False)
    ])
    add_formula(doc, "Rata-rata Konflik (K̄_p) = N_p / T_p   ;   χ² = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   OR = ( a × d ) / ( b × c )",
                ket="N_p = Total konflik periode p; T_p = Jumlah tahun periode p; K̄_p = Kasus per tahun; χ² = Statistik Chi-Square (O_ij = observasi, E_ij = harapan); OR = Odds Ratio kelipatan peluang represi pada sektor tambang.")

    add_caption(doc, "Tabel 4.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 4.4)")
    tabel_4_4a_rows = [
        ["Variabel Independen (X)", "Periode Ekspansi Industri (Pasca vs Pra 2014); Tipe Sektor (Tambang vs Non-Tambang); Keterlibatan Aparat/Pemerintah."],
        ["Variabel Dependen (Y)", "Tingkat Represi & Kriminalisasi; Tingkat Penelantaran Kasus; Tingkat Insiden Fisik (Ditangkap/Luka/Tewas)."],
        ["Hipotesis Nol (H0)", "Faktor ekspansi industri dan tipe sektor saling bebas secara absolut terhadap tingkat represi dan kriminalisasi."],
        ["Hipotesis Alternatif (H1)", "Ekspansi industri pertambangan berasosiasi signifikan dengan peningkatan risiko represi dan kriminalisasi pejuang hak tenurial."],
        ["Decision Rule (Alpha 5%)", "Chi-Square P-Value < 0.05 (Tolak H0) dan rasio peluang Odds Ratio (OR) > 1.0."],
        ["Threshold Kategori", "Klasifikasi biner data cross-section (N=523 kejadian letupan konflik historis): Periode (Pasca-2014 vs Pra-2014), Sektor (Tambang vs Lainnya), Represi (Ada vs Tidak Ada)."],
        ["Orientasi Odds Ratio", "OR = ( a × d ) / ( b × c ) dengan a = Sektor Tambang/Era Pasca & Ada Represi; mengukur kelipatan risiko kekerasan pada aktivitas industri ekstraktif."]
    ]
    add_table_styled(
        doc,
        headers=["Komponen Uji", "Definisi Variabel (Sub-bab 4.4)"],
        rows=tabel_4_4a_rows,
        col_widths_cm=[4.5, 12.5],
        alignments=['L', 'L']
    )

    # 4.5
    add_h3(doc, "4.5 Peta Entitas Aktor: Korporasi dan Organisasi Masyarakat")
    add_body(doc, [
        ("Penambangan teks berbasis Regular Expressions (RegEx NLP) membedah korpus narasi kronologi sengketa lahan (N=568 dokumen kasus) untuk memetakan frekuensi keterlibatan entitas korporasi dan aktor proksi swakarsa:", False, False)
    ])
    add_formula(doc, "Korpus = Gabungan ( Judul_k , Deskripsi_k , Narasi_k )   ;   Frekuensi Entitas_a = Σ [ Match_i,a ]",
                ket="Judul_k, Deskripsi_k, Narasi_k = Teks narasi bebas kasus k; Korpus = Kumpulan seluruh narasi kasus agraria; Match_i,a = Kecocokan pola teks entitas a pada dokumen i; Frekuensi = Total kemunculan entitas dalam korpus.")

    # ── E. KORESPONDENSI METODOLOGI TERHADAP SUB-BAB LAPORAN ────
    add_h2(doc, "E", "Korespondensi Metodologi terhadap Sub-bab Laporan Bab 4")
    add_body(doc, [
        ("Setiap sub-bab analitis pada Bab 4 ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:", False, False)
    ])

    table_korespondensi = [
        ["Sub-bab 4.1", "Eskalasi Konflik Agraria Historis", "Time-Series Trend Analysis, Laju Pertumbuhan Kasus Pra vs Pasca Hilirisasi"],
        ["Sub-bab 4.2", "Asimetri Dampak Sosial & Penguasaan Ruang", "Sectoral Burden Analysis, Agregasi Korban Jiwa & Monopoli Hektar Lahan"],
        ["Sub-bab 4.3", "Ruang Sipil, Represi & Kriminalisasi", "Violence Tracking Analysis, Agregasi Kriminalisasi & Korban Pelanggaran HAM"],
        ["Sub-bab 4.4", "Pembuktian Statistik Relasi Kausalitas", "Before-After Cross-Section, Uji Chi-Square (χ²), Odds Ratio Risiko (OR)"],
        ["Sub-bab 4.5", "Orkestrasi Oligarki & Aktor Proksi", "Text Parsing NLP RegEx, Frequency Profiling Korporasi & Kelompok Vigilante"]
    ]

    add_table_styled(
        doc,
        headers=["Sub-bab", "Fokus Kajian Empiris", "Metode Analitis Utama"],
        rows=table_korespondensi,
        col_widths_cm=[2.5, 5.5, 9.0],
        alignments=['C', 'L', 'L']
    )

    # ── F. BAGAN ALUR KERANGKA KERJA RISET BAB 4 ────────────────
    add_h2(doc, "F", "Bagan Alur Kerangka Kerja Riset (Research Workflow)")
    add_body(doc, [
        ("Kerangka operasional metodologi Bab 4 berjalan secara terpadu melalui empat fase berurutan sebagaimana divisualisasikan pada bagan alur kerja riset berikut:", False, False)
    ])

    mermaid_str_f = """flowchart LR
    subgraph F1["Fase I: Akuisisi Data"]
        A1["Kurasi Basis Data Konflik<br/><i>TanahKita KPA, Komnas HAM, WALHI</i>"]
        A2["Korpus Narasi Sengketa<br/><i>568 Kronologi Kasus & Arsip Advokasi</i>"]
    end
    subgraph F2["Fase II: Reklasifikasi & NLP"]
        B1["Reklasifikasi Sektoral & Waktu<br/><i>Tambang, Sawit vs Pra & Pasca 2014</i>"]
        B2["Ekstraksi Pola RegEx NLP<br/><i>Entitas Korporasi & Aktor Proksi</i>"]
    end
    subgraph F3["Fase III: Uji Statistik"]
        C1["Tabel Kontinjensi 2x2<br/><i>Sektor & Periode vs Represi</i>"]
        C2["Uji Chi-Square & Odds Ratio<br/><i>Signifikansi & Kelipatan Risiko HAM</i>"]
    end
    subgraph F4["Fase IV: Profiling & Sintesis"]
        D1["Frequency Profiling Aktor<br/><i>Top 10 Oligarki & Vigilante</i>"]
        D2["Bukti Kausalitas D3TLH<br/><i>Monopoli Ruang & Kriminalisasi</i>"]
    end
    F1 --> F2 --> F3 --> F4"""

    png_workflow_path = str(out_dir_compact / "mermaid_workflow_bab4.png")
    is_downloaded = download_mermaid_png(mermaid_str_f, png_workflow_path)

    add_caption(doc, "Bagan Alur 4.1: Alur Logika Kerangka Kerja Riset Bab 4 (Research Workflow)")
    if is_downloaded and os.path.exists(png_workflow_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(3)
        p_img.paragraph_format.space_after  = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(png_workflow_path, width=Cm(16.5))
        try:
            shutil.copyfile(png_workflow_path, str(out_dir_bab4 / "mermaid_workflow_bab4.png"))
        except Exception:
            pass

    # Box Output Kesimpulan
    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after  = Pt(4)
    all_border_para(p_box, color='1B5E20', sz='8')
    para_shd(p_box, 'F1F8E9')
    add_run(p_box, "KERANGKA KELUARAN METODOLOGIS BAB 4:\n", bold=True, pt=8.5, color=G_DARK)
    add_run(p_box, "1. Konfigurasi Asimetri Penguasaan Ruang: Membuktikan secara empiris bahwa sektor pertambangan nikel memonopoli 73,1% luasan sengketa lahan (441.286 Ha) dan menumbalkan 60,3% korban terdampak (54.658 jiwa).\n"
                   "2. Konfigurasi Inferensial Eskalasi Represi: Membuktikan korelasi kausalitas signifikan secara statistik antara ekspansi industri tambang dan keterlibatan aparat terhadap lonjakan risiko kriminalisasi warga (Odds Ratio hingga 4,8 kali lipat).\n"
                   "3. Konfigurasi Profiling Aktor & Oligarki: Mengungkap modus operandi pengamanan swakarsa dan orkestrasi aktor proksi vigilante di balik perampasan ruang hidup masyarakat lingkar tambang.",
            pt=8, color=C_BODY)

    # ── SIMPAN DOKUMEN DOCX (DUAL SAVE) ─────────────────────────
    docx_compact = out_dir_compact / "Metodologi_Bab4_Ruang_Hidup_Compact.docx"
    docx_bab4    = out_dir_bab4 / "Metodologi_Bab4_Ruang_Hidup_Compact.docx"
    
    doc.save(str(docx_compact))
    shutil.copyfile(docx_compact, docx_bab4)
    print(f"  [OK] Tersimpan DOCX: {docx_compact}")
    print(f"  [OK] Salinan DOCX : {docx_bab4}")

    # ── GENERATE MARKDOWN PADANAN ───────────────────────────────
    print("[2/3] Membangun dokumen Markdown padanan...")
    MD_CONTENT = """# METODOLOGI PENELITIAN: BAB 4 — ANALISIS RUANG HIDUP YANG TERAMPAS
*CELIOS (Center of Economic and Law Studies) · Audit Spasial-Statistik D3TLH Sulawesi (2014–2024) · Ringkasan Eksekutif Metodologis*

---

## A. Desain Penelitian & Tujuan
Penelitian ini menggunakan **desain sosiologi hukum kritis, audit sengketa agraria, dan analisis inferensial kuantitatif terintegrasi** untuk mengidentifikasi perampasan ruang hidup masyarakat adat dan komunitas lokal akibat ekspansi masif korporasi tambang nikel dan kawasan industri penunjang di Pulau Sulawesi sepanjang kurun waktu pengamatan (**1990–2024**). Tiga tujuan utama metodologis Bab 4 meliputi:

1. **Kuantifikasi Asimetri Penguasaan Ruang & Korban Terdampak:** Mengevaluasi distribusi sektoral sengketa agraria guna membuktikan dominasi monopoli lahan dan skala korban masyarakat terdampak pada sektor pertambangan nikel dibandingkan sektor lainnya.
2. **Pembuktian Inferensial Kausalitas Ekspansi vs Eskalasi Represi:** Menguji signifikansi hubungan antara periode hilirisasi dan keterlibatan korporasi ekstraktif terhadap peningkatan risiko kekerasan, penangkapan, serta kriminalisasi warga melalui matriks Chi-Square dan Odds Ratio (OR).
3. **Pemetaan Entitas Oligarki & Aktor Proksi (NLP Text Parsing):** Mengekstraksi jaringan entitas korporasi dominan dan mendeteksi keterlibatan aktor vigilante/pengamanan swakarsa dalam kronologi sengketa agraria melalui penambangan teks reguler (NLP Regex).

---

## B. Sumber Data & Cakupan Wilayah
Penelitian mencakup catatan letupan konflik agraria terdata di seluruh wilayah administratif **Pulau Sulawesi dan pulau-pulau kecil penyangga sentra nikel** (seperti Pulau Wawonii dan pesisir Morowali). Basis data dibangun dari repositori dokumentasi konflik agraria nasional dan advokasi masyarakat sipil:

- **Konsorsium Pembaruan Agraria (KPA) / Basis Data TanahKita:** Dokumentasi 95 kasus konflik agraria regional Sulawesi (dan korpus 568 narasi nasional) mencakup luas sengketa (Ha), jiwa terdampak, sektor industri, serta status penanganan hukum.
- **Koalisi Masyarakat Sipil (WALHI, JATAM, AMAN):** Kronologi advokasi hak tenurial masyarakat adat, kasus kekerasan fisik, dan pemantauan perampasan wilayah kelola rakyat.
- **Komisi Nasional Hak Asasi Manusia (Komnas HAM):** Registri pengaduan pelanggaran hak sipil, kriminalisasi pejuang lingkungan, warga ditangkap, luka-luka, dan korban tewas.
- **Kementerian ATR/BPN & ESDM (MODI):** Status perizinan hak guna usaha (HGU), izin usaha pertambangan (IUP), dan konsesi kawasan industri.

---

## C. Operasionalisasi Variabel & Indikator Riset
Seluruh dinamika perampasan ruang, eskalasi konflik, dan keterlibatan aktor dioperasionalkan ke dalam **10 indikator empiris terpadu** sebagaimana dirangkum pada matriks operasional berikut:

##### Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 4
| No | Indikator Riset | Fokus Pengukuran | Satuan | Periode | Sumber Data Primer Resmi |
| :-: | :--- | :--- | :-: | :-: | :--- |
| 1 | Insidensi Konflik Agraria | Frekuensi Kejadian Letupan Sengketa Lahan | Kasus | 1990–2024 | KPA / TanahKita |
| 2 | Sektor Pemicu Sengketa | Klasifikasi Sektoral (Tambang, Sawit, Hutan) | Kategori | 1990–2024 | KPA / TanahKita |
| 3 | Skala Korban Terdampak | Masyarakat Adat & Komunitas Lokal Terdampak | Jiwa | 1990–2024 | KPA & Koalisi Sipil |
| 4 | Luas Monopoli Area Sengketa | Luas Ruang Hidup & Lahan Diperebutkan | Hektar (Ha) | 1990–2024 | TanahKita & ATR/BPN |
| 5 | Kasus Indikasi Kriminalisasi | Penuntutan Hukum Terhadap Warga/Aktivis | Kasus | 2000–2024 | KPA & Komnas HAM |
| 6 | Korban Represi & Kekerasan | Warga Ditangkap, Mengalami Luka, & Tewas | Orang | 2000–2024 | KPA & Komnas HAM |
| 7 | Laju Eskalasi Hilirisasi | Rasio Before-After Kasus Pra vs Pasca 2014 | Kasus / Tahun | 1990 vs 2024 | Data Panel Tahunan |
| 8 | Tingkat Penelantaran Kasus | Sengketa Lahan Berstatus Belum Ditangani | Persen (%) | 1990–2024 | KPA / TanahKita |
| 9 | Frekuensi Entitas Korporasi | Keterlibatan Konglomerasi dalam Konflik | Token Count | Korpus NLP | NLP Regex TanahKita |
| 10 | Frekuensi Aktor Proksi/Vigilante | Keterlibatan Pengamanan Swakarsa/Preman | Token Count | Korpus NLP | NLP Regex TanahKita |

---

## D. Kerangka Analisis & Formulasi Matematis

### 4.1 Tren Eskalasi Konflik Agraria Seiring Ekspansi Industri
Kuantifikasi eskalasi sengketa ruang hidup dihitung berdasarkan deret waktu tahunan dan rasio laju peningkatan kasus antara era pra-ekspansi dan era hilirisasi:

> `Agregasi Konflik (K_t,s) = Σ c_i   ;   Laju Eskalasi (E %) = [ K_Pasca / K_Pra ] × 100`  
> *Keterangan: c_i = Kasus konflik agraria i pada tahun t dan sektor s; K_Pasca = Total kasus pasca ekspansi industri; K_Pra = Total kasus pra ekspansi; E = Laju persentase lonjakan eskalasi konflik.*

### 4.2 Sebaran Sektoral: Dampak Masyarakat dan Penggunaan Lahan
Dekomposisi beban dampak sosiologis dan monopoli penguasaan ruang dihitung per sektor industri pemicu konflik guna mengukur asimetri dampak:

> `Total Jiwa Terdampak (J_s) = Σ J_i   ;   Total Luas Area (A_s) = Σ A_i   ;   Porsi Sektoral (P_s %) = [ Nilai_s / Nilai_Total ] × 100`  
> *Keterangan: J_i = Warga terdampak pada kasus i; A_i = Luas lahan sengketa kasus i (Ha); J_s & A_s = Total korban jiwa dan luas sengketa sektor s; P_s = Pangsa persentase sektor terhadap total regional.*

### 4.3 Indikasi Represi dan Kriminalisasi dalam Konflik Agraria
Kuantifikasi penyempitan ruang sipil menghitung total kasus kriminalisasi serta menjumlahkan seluruh korban represi fisik yang terdokumentasi:

> `Total Kriminalisasi = Σ I_i   ;   Total Korban Represi (R) = Σ [ D_i + L_i + T_i ]`  
> *Keterangan: I_i = Indikator biner kriminalisasi pada kasus i (1 jika ada); D_i = Korban ditangkap; L_i = Korban luka-luka; T_i = Korban tewas; R = Akumulasi total korban pelanggaran HAM.*

### 4.4 Pembuktian Statistik: Ekspansi vs Eskalasi Konflik
Pengujian komparatif Before-After dan uji independensi Chi-Square (χ²) tabulasi silang diterapkan pada basis data kejadian konflik (N=523) untuk membuktikan korelasi ekspansi industri terhadap eskalasi represi:

> `Rata-rata Konflik (K̄_p) = N_p / T_p   ;   χ² = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   OR = ( a × d ) / ( b × c )`  
> *Keterangan: N_p = Total konflik periode p; T_p = Jumlah tahun periode p; K̄_p = Kasus per tahun; χ² = Statistik Chi-Square (O_ij = observasi, E_ij = harapan); OR = Odds Ratio kelipatan peluang represi pada sektor tambang.*

##### Tabel 4.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 4.4)
| Komponen Uji | Definisi Variabel (Sub-bab 4.4) |
| :--- | :--- |
| **Variabel Independen (X)** | Periode Ekspansi Industri (Pasca vs Pra 2014); Tipe Sektor (Tambang vs Non-Tambang); Keterlibatan Aparat/Pemerintah. |
| **Variabel Dependen (Y)** | Tingkat Represi & Kriminalisasi; Tingkat Penelantaran Kasus; Tingkat Insiden Fisik (Ditangkap/Luka/Tewas). |
| **Hipotesis Nol (H0)** | Faktor ekspansi industri dan tipe sektor saling bebas secara absolut terhadap tingkat represi dan kriminalisasi. |
| **Hipotesis Alternatif (H1)** | Ekspansi industri pertambangan berasosiasi signifikan dengan peningkatan risiko represi dan kriminalisasi pejuang hak tenurial. |
| **Decision Rule (Alpha 5%)** | Chi-Square P-Value < 0.05 (Tolak H0) dan rasio peluang Odds Ratio (OR) > 1.0. |
| **Threshold Kategori** | Klasifikasi biner data cross-section (N=523 kejadian letupan konflik historis): Periode (Pasca-2014 vs Pra-2014), Sektor (Tambang vs Lainnya), Represi (Ada vs Tidak Ada). |
| **Orientasi Odds Ratio** | OR = ( a × d ) / ( b × c ) dengan a = Sektor Tambang/Era Pasca & Ada Represi; mengukur kelipatan risiko kekerasan pada aktivitas industri ekstraktif. |

### 4.5 Peta Entitas Aktor: Korporasi dan Organisasi Masyarakat
Penambangan teks berbasis Regular Expressions (RegEx NLP) membedah korpus narasi kronologi sengketa lahan (N=568 dokumen kasus) untuk memetakan frekuensi keterlibatan entitas korporasi dan aktor proksi swakarsa:

> `Korpus = Gabungan ( Judul_k , Deskripsi_k , Narasi_k )   ;   Frekuensi Entitas_a = Σ [ Match_i,a ]`  
> *Keterangan: Judul_k, Deskripsi_k, Narasi_k = Teks narasi bebas kasus k; Korpus = Kumpulan seluruh narasi kasus agraria; Match_i,a = Kecocokan pola teks entitas a pada dokumen i; Frekuensi = Total kemunculan entitas dalam korpus.*

---

## E. Korespondensi Metodologi terhadap Sub-bab Laporan Bab 4
Setiap sub-bab analitis pada Bab 4 ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:

##### Matriks Korespondensi Sub-bab terhadap Metode Analitis
| Sub-bab | Fokus Kajian Empiris | Metode Analitis Utama |
| :---: | :--- | :--- |
| **Sub-bab 4.1** | Eskalasi Konflik Agraria Historis | Time-Series Trend Analysis, Laju Pertumbuhan Kasus Pra vs Pasca Hilirisasi |
| **Sub-bab 4.2** | Asimetri Dampak Sosial & Penguasaan Ruang | Sectoral Burden Analysis, Agregasi Korban Jiwa & Monopoli Hektar Lahan |
| **Sub-bab 4.3** | Ruang Sipil, Represi & Kriminalisasi | Violence Tracking Analysis, Agregasi Kriminalisasi & Korban Pelanggaran HAM |
| **Sub-bab 4.4** | Pembuktian Statistik Relasi Kausalitas | Before-After Cross-Section, Uji Chi-Square (χ²), Odds Ratio Risiko (OR) |
| **Sub-bab 4.5** | Orkestrasi Oligarki & Aktor Proksi | Text Parsing NLP RegEx, Frequency Profiling Korporasi & Kelompok Vigilante |

---

## F. Bagan Alur Kerangka Kerja Riset (Research Workflow)

```mermaid
flowchart LR
    subgraph F1["Fase I: Akuisisi Data"]
        A1["Kurasi Basis Data Konflik<br/><i>TanahKita KPA, Komnas HAM, WALHI</i>"]
        A2["Korpus Narasi Sengketa<br/><i>568 Kronologi Kasus & Arsip Advokasi</i>"]
    end
    subgraph F2["Fase II: Reklasifikasi & NLP"]
        B1["Reklasifikasi Sektoral & Waktu<br/><i>Tambang, Sawit vs Pra & Pasca 2014</i>"]
        B2["Ekstraksi Pola RegEx NLP<br/><i>Entitas Korporasi & Aktor Proksi</i>"]
    end
    subgraph F3["Fase III: Uji Statistik"]
        C1["Tabel Kontinjensi 2x2<br/><i>Sektor & Periode vs Represi</i>"]
        C2["Uji Chi-Square & Odds Ratio<br/><i>Signifikansi & Kelipatan Risiko HAM</i>"]
    end
    subgraph F4["Fase IV: Profiling & Sintesis"]
        D1["Frequency Profiling Aktor<br/><i>Top 10 Oligarki & Vigilante</i>"]
        D2["Bukti Kausalitas D3TLH<br/><i>Monopoli Ruang & Kriminalisasi</i>"]
    end
    F1 --> F2 --> F3 --> F4
```

> **KERANGKA KELUARAN METODOLOGIS BAB 4:**  
> 1. **Konfigurasi Asimetri Penguasaan Ruang:** Membuktikan secara empiris bahwa sektor pertambangan nikel memonopoli 73,1% luasan sengketa lahan (441.286 Ha) dan menumbalkan 60,3% korban terdampak (54.658 jiwa).  
> 2. **Konfigurasi Inferensial Eskalasi Represi:** Membuktikan korelasi kausalitas signifikan secara statistik antara ekspansi industri tambang dan keterlibatan aparat terhadap lonjakan risiko kriminalisasi warga (Odds Ratio hingga 4,8 kali lipat).  
> 3. **Konfigurasi Profiling Aktor & Oligarki:** Mengungkap modus operandi pengamanan swakarsa dan orkestrasi aktor proksi vigilante di balik perampasan ruang hidup masyarakat lingkar tambang.
"""

    md_compact = out_dir_compact / "Metodologi_Bab4_Ruang_Hidup_Compact.md"
    md_bab4    = out_dir_bab4 / "Metodologi_Bab4_Ruang_Hidup_Compact.md"
    for pth in [md_compact, md_bab4]:
        with open(pth, "w", encoding="utf-8") as f:
            f.write(MD_CONTENT)
    print(f"  [OK] Tersimpan MD  : {md_compact}")
    print(f"  [OK] Salinan MD   : {md_bab4}")

    print("[3/3] Selesai menghasilkan dokumen metodologi Bab 4 versi compact (1-Kolom, 2-3 Halaman).")


if __name__ == "__main__":
    generate_bab4_compact()
