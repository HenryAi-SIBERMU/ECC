#!/usr/bin/env python3
"""
Generator Metodologi Versi Compact Bab 7 — GAYA AKADEMIS TERPADU (CELIOS)
Mengadopsi arsitektur metodologi ringkas terstandarisasi konsisten dengan Bab 1 s.d. 6:
- RUANG LINGKUP: Bab 7 — Kegagalan Tata Kelola: D3TLH dalam Sistem Perizinan
- FORMAT: 1 KOLOM PENUH (Single Column Layout)
- PANJANG: 2–3 Halaman Maksimal (Elegan, proporsional, tanpa pemadatan berlebihan)
- PENOMORAN SEKSI UTAMA: Huruf kapital A, B, C, D, E, F
- SUB-BAB SEKSI D: Sub-bab 7.1, 7.2, 7.3 sesuai dokumen induk
- OPERASIONALISASI INDIKATOR: 9 Indikator Riset Empiris Kunci Terverifikasi (5 Kolom Baku tanpa kolom Periode)
- NOTASI MATEMATIKA: Bahasa intuitif dan ramah pembaca awam dengan penjelasan penalaran logis
- KORESPONDENSI METODOLOGI: 3 kolom bersih (Sub-bab, Fokus Kajian Empiris, Metode Analitis Utama)
- FLOWCHART: Mermaid JS horizontal (flowchart LR) dirender tajam ke DOCX (16.5 cm) dan blok kode di MD
- SINKRONISASI: Dual-save ke direktori versicompact/bab_7 dan bab_7.
"""

import os
import sys
import shutil
import base64
import requests
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Palet Warna Resmi CELIOS ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Forest Dark Green (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Celios Accent Green (#2E7D32)
G_LIGHT = RGBColor(0x38, 0x8E, 0x3C)  # Medium Green (#388E3C)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Charcoal Body Text (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Muted Grey Text (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Pure White

# ── Helper Fungsi XML Word ───────────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=80, right=80, top=40, bottom=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)

def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    pPr.append(s)

def para_border_bottom(p, color='1B5E20', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '3')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='1B5E20', sz='8'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Typography & Content ──────────────────────────────
def add_run(p, text, bold=False, italic=False, pt=8.5, color=C_BODY, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    else:
        r.font.name = 'Calibri'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    return r

def add_h2(doc, prefix, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_border_bottom(p, color='1B5E20', sz='8')
    add_run(p, f"{prefix}.  ", bold=True, pt=11, color=G_DARK)
    add_run(p, title.upper(), bold=True, pt=11, color=G_DARK)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, pt=9.5, color=G_MID)
    return p

def add_body(doc, parts, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        add_run(p, text, bold=bold, italic=italic, pt=8.5, color=C_BODY)
    return p

def add_formula(doc, text, ket=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    para_shd(p, 'EDF7EE')
    add_run(p, text, pt=8, color=G_MID, mono=True)
    if ket:
        add_run(p, f"\nKeterangan: {ket}", italic=True, pt=7.5, color=RGBColor(0x33, 0x55, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_styled(doc, headers, rows, col_widths_cm, alignments=None, font_pt=7):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    bd_subtle = {'val': 'single', 'sz': '4', 'color': 'D0D7DE', 'space': '0'}
    
    # Header Row
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=70, right=70, top=45, bottom=45)
        set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if alignments and j < len(alignments):
            align = alignments[j]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, pt=7.5, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F9FBF9' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_shd(c, fill)
            cell_margin(c, left=70, right=70, top=35, bottom=35)
            set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            
            if alignments and j < len(alignments):
                align = alignments[j]
                if align == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_run(p, str(val), pt=font_pt, color=C_BODY)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after  = Pt(3)
    return tbl

def download_mermaid_png(mermaid_str, filepath):
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        return True
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False


def generate_bab7_compact():
    print("[1/3] Membangun dokumen compact Bab 7 (Format 1-Kolom, 2-3 Halaman)...")
    
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab7    = out_dir_compact.parent.parent / "bab_7"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab7.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER DOKUMEN ──────────────────────────────────────────
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after  = Pt(1)
    add_run(p_t, "RINGKASAN EKSEKUTIF METODOLOGIS", bold=True, pt=8.5, color=G_MID)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after  = Pt(2)
    para_border_bottom(p_h, color='1B5E20', sz='12')
    add_run(p_h, "BAB 7: METODOLOGI ANALISIS KEGAGALAN TATA KELOLA", bold=True, pt=15, color=G_DARK)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(1)
    p_meta.paragraph_format.space_after  = Pt(5)
    add_run(p_meta, "D3TLH dalam Sistem Perizinan: Evaluasi Kepatuhan, Impunitas Hukum, dan Inkonsistensi Iklim · ", italic=True, pt=8, color=C_GREY)
    add_run(p_meta, "Center of Economic and Law Studies (CELIOS)", bold=True, italic=True, pt=8, color=G_DARK)

    # ── A. DESAIN PENELITIAN & TUJUAN ───────────────────────────
    add_h2(doc, "A", "Desain Penelitian & Tujuan")
    add_body(doc, [
        ("Penelitian Bab 7 menerapkan ", False, False),
        ("desain audit forensik kepatuhan kebijakan (Compliance & Institutional Failure Audit)", True, False),
        (" untuk menguji efektivitas instrumen Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) dalam tata kelola perizinan industri ekstraktif di Pulau Sulawesi. Melalui integrasi spasial data perizinan tambang, tutupan hutan, sengketa tenurial, dan infrastruktur energi fosil, kajian ini membuktikan secara empiris tiga kegagalan struktural tata kelola:", False, False)
    ])
    add_body(doc, [
        ("1. ", True, False), ("Evaluasi Kepatuhan D3TLH vs Penerbitan Izin (Rule-based Categorization): ", True, False),
        ("Menganalisis sinkronisasi antara batas pengaman daya dukung hutan (GFW) dan izin baru pertambangan (MODI ESDM), guna membuktikan apakah instrumen D3TLH berfungsi sebagai pembatas atau diabaikan dalam zona krisis ekologis.\n", False, False),
        ("2. ", True, False), ("Audit Impunitas Hukum & Pembiaran Korporasi (Thematic Coding): ", True, False),
        ("Menginventarisasi rekam jejak pelanggaran hukum, perampasan ruang hidup, dan sengketa agraria yang dibiarkan tanpa penegakan sanksi administratif maupun pidana lingkungan (*state omission*).\n", False, False),
        ("3. ", True, False), ("Kuantifikasi Kontradiksi Karbon PLTU Captive (Asset-level Inventory): ", True, False),
        ("Mendokumentasikan paradoks hilirisasi hijau melalui inventarisasi aset 67 unit pembangkit listrik tenaga uap (PLTU) batubara *off-grid* captive yang beroperasi di dalam kawasan industri nikel.", False, False)
    ])

    # ── B. SUMBER DATA & CAKUPAN WILAYAH ─────────────────────────
    add_h2(doc, "B", "Sumber Data & Cakupan Wilayah")
    add_body(doc, [
        ("Audit tata kelola perizinan ini menggabungkan 4 basis data resmi lintas kementerian dan lembaga masyarakat sipil yang mencakup seluruh yurisdiksi 6 provinsi se-Pulau Sulawesi kurun 2014–2024:", False, False)
    ])
    add_body(doc, [
        ("• ", True, False), ("MODI Ditjen Minerba ESDM RI: ", True, False),
        ("Data geospasial izin baru pertambangan (IUP), tahun penerbitan, luas konsesi (hektar), dan komoditas tambang aktif.\n", False, False),
        ("• ", True, False), ("Global Forest Watch (GFW / Hansen UMD): ", True, False),
        ("Time-series deforestasi tahunan tingkat provinsi guna menetapkan ambang persentil daya dukung hutan alam.\n", False, False),
        ("• ", True, False), ("Konsorsium Pembaruan Agraria (CATAHU KPA), TanahKita, & Koalisi Sipil: ", True, False),
        ("Dokumentasi kasus konflik tenurial, pelanggaran izin di kawasan lindung, kriminalisasi warga, dan impunitas korporasi.\n", False, False),
        ("• ", True, False), ("Global Energy Monitor (GEM Coal Plant Tracker, Jan 2026): ", True, False),
        ("Inventarisasi aset 67 unit PLTU captive batubara di kawasan industri Morowali, Konawe, Bantaeng, dan sekitarnya.", False, False)
    ])

    # ── C. OPERASIONALISASI VARIABEL & INDIKATOR RISET ──────────
    add_h2(doc, "C", "Operasionalisasi Variabel & Indikator Riset")
    add_body(doc, [
        ("Seluruh parameter tata kelola, pelanggaran batas daya dukung, sengketa tenurial, hingga aset energi fosil dioperasionalkan ke dalam ", False, False),
        ("9 indikator riset empiris kunci", True, False),
        (" sebagaimana dirangkum pada matriks operasional berikut:", False, False)
    ])

    add_caption(doc, "Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 7 (Tata Kelola Perizinan)")
    table_indikator_data = [
        ["1", "Kepatuhan D3TLH Perizinan (7.1)", "Rasio Penerbitan IUP Baru pada Zona Deforestasi Kritis", "Unit Izin", "MODI ESDM & GFW Hansen"],
        ["2", "Luas Konsesi Zona Kritis (7.1)", "Total Luasan Konsesi Tambang Terbit di Status Kritis", "Hektar (Ha)", "Ditjen Minerba ESDM & GFW"],
        ["3", "Ambang Persentil Daya Dukung (7.1)", "Klasifikasi Status Hutan: Aman (P33), Tertekan, Kritis (P66)", "Hektar / Tahun", "GFW Master Time-Series 2014-2023"],
        ["4", "Kasus Impunitas Penegakan Hukum (7.2)", "Frekuensi Pembiaran Pelanggaran Korporasi Ekstraktif", "Kasus", "CATAHU KPA, TanahKita, YLBHI"],
        ["5", "Dominasi Konflik Sektoral (7.2)", "Proporsi Sengketa Agraria Akibat Ekspansi Tambang", "Persen (%)", "Database Kasus Tenurial KPA"],
        ["6", "Konsentrasi Spasial Sengketa (7.2)", "Sebaran Wilayah Konflik Tertinggi Lintas Provinsi", "Kasus & %", "Laporan Koalisi Sipil & KPA"],
        ["7", "Agregat Aset PLTU Captive (7.3)", "Total Unit & Kapasitas Pembangkitan Kotor Batubara Off-Grid", "Unit & Megawatt (MW)", "Global Energy Monitor (GEM 2026)"],
        ["8", "Kapasitas PLTU Captive Aktif (7.3)", "Kapasitas Terpasang Unit Batubara Beroperasi (Operating)", "Megawatt (MW)", "GEM Global Coal Plant Tracker"],
        ["9", "Dominasi Beban Karbon Spasial (7.3)", "Konsentrasi Kapasitas Pembangkitan Fosil per Provinsi", "Persen (%)", "GEM Asset Inventory Tracker"]
    ]

    add_table_styled(
        doc,
        headers=["No", "Indikator Riset", "Fokus Pengukuran", "Satuan", "Sumber Data Primer Resmi"],
        rows=table_indikator_data,
        col_widths_cm=[0.8, 4.5, 4.5, 2.2, 5.0],
        alignments=['C', 'L', 'L', 'C', 'L'],
        font_pt=7
    )

    # ── D. KERANGKA ANALISIS & FORMULASI MATEMATIS ──────────────
    add_h2(doc, "D", "Kerangka Analisis & Formulasi Matematis")

    # Sub-bab 7.1
    add_h3(doc, "Sub-bab 7.1: Evaluasi Kepatuhan D3TLH vs Penerbitan Izin Tambang Baru")
    add_body(doc, [
        ("Penilaian kepatuhan sistem perizinan terhadap batas pengaman lingkungan hidup dilakukan dengan mengelompokkan data deforestasi tahunan ke dalam 3 kelas persentil daya dukung, kemudian menyandingkannya dengan penerbitan IUP baru:", False, False)
    ])
    add_formula(doc,
        "1. Penentuan Ambang Batas Daya Dukung (Binning Persentil):\n"
        "   • Ambang Tertekan = Persentil-33 (P33) dari Deforestasi Tahunan = 12.898 Ha/tahun\n"
        "   • Ambang Kritis   = Persentil-66 (P66) dari Deforestasi Tahunan = 26.453 Ha/tahun\n\n"
        "2. Aturan Klasifikasi Status Wilayah:\n"
        "   • Status Aman     : Kerusakan Hutan <= 12.898 Ha/tahun (Izin Wajar Diterbitkan)\n"
        "   • Status Tertekan : 12.898 Ha < Kerusakan <= 26.453 Ha/tahun (Penerbitan Izin Mulai Dibatasi)\n"
        "   • Status Kritis   : Kerusakan Hutan > 26.453 Ha/tahun (Wajib Moratorium Total)\n\n"
        "3. Kuantifikasi Pelanggaran Tata Kelola:\n"
        "   Total Izin Zona Kritis = Jumlah IUP Baru yang Tetap Diterbitkan saat Status Hutan Kritis",
        ket="Fakta Empiris Lapangan: Pada saat status hutan dinyatakan KRITIS (> 26.453 Ha deforestasi), pemerintah justru menerbitkan 277 IUP BARU dengan total luas konsesi mencapai 440.998 Hektar (mayoritas di Sulteng dan Sulsel)."
    )

    add_caption(doc, "Tabel 7.1: Matriks Kepatuhan D3TLH — Seharusnya vs Kenyataan Penerbitan Izin")
    tabel_7_1_rows = [
        ["Aman", "<= 12.898 Ha", "20", "Wajar diterbitkan izin", "26 Izin Baru Keluar", "87.070 Ha", "Normal (Sesuai Aturan)"],
        ["Tertekan", "12.898 - 26.453 Ha", "19", "Izin mulai dibatasi ketat", "77 Izin Baru Keluar", "107.377 Ha", "Anomali (Lampu Kuning)"],
        ["Kritis", "> 26.453 Ha", "21", "Moratorium Total / Larangan Izin", "277 Izin Baru Keluar", "440.998 Ha", "PELANGGARAN STRUKTURAL"]
    ]
    add_table_styled(
        doc,
        headers=["Status Daya Dukung", "Deforestasi Tahunan", "N Observasi", "Aturan Seharusnya", "Kenyataan di Lapangan", "Luas Konsesi (Ha)", "Kesimpulan Tata Kelola"],
        rows=tabel_7_1_rows,
        col_widths_cm=[2.2, 2.5, 1.5, 3.2, 2.8, 2.2, 2.6],
        alignments=['C', 'C', 'C', 'L', 'C', 'R', 'C'],
        font_pt=6.5
    )

    # Sub-bab 7.2
    add_h3(doc, "Sub-bab 7.2: Pemetaan Impunitas Hukum & Pembiaran Kasus Tenurial")
    add_body(doc, [
        ("Tingkat pembiaran negara (*state omission*) dihitung melalui klasifikasi tematik terhadap seluruh insiden konflik agraria, pelanggaran kawasan lindung, dan intimidasi warga yang terdokumentasi tanpa adanya sanksi hukum tegas:", False, False)
    ])
    add_formula(doc,
        "1. Penghitungan Proporsi Konflik Sektoral:\n"
        "   Proporsi Sektor (%) = (Jumlah Kasus pada Sektor Tertentu / Total Kasus Terdata) × 100%\n\n"
        "2. Identifikasi Wilayah Episentrum Pembiaran:\n"
        "   Proporsi Wilayah (%) = (Jumlah Kasus di Suatu Provinsi / Total Kasus Terdata) × 100%",
        ket="Fakta Empiris: Dari 32 kasus impunitas yang terdata di Sulawesi, Sektor Pertambangan menjadi penyumbang terbesar dengan 11 kasus (34,4%), disusul Perkebunan Sawit 6 kasus (18,8%). Sulawesi Tenggara mencatat kasus terbanyak (8 kasus / 25,0%), disusul Sulawesi Selatan (7 kasus / 21,9%)."
    )

    add_caption(doc, "Tabel 7.2: Sebaran Sektor Konflik dan Pembiaran Operasi Ilegal di Sulawesi")
    tabel_7_2_rows = [
        ["Pertambangan (Nikel & Batuan)", "11 Kasus", "34,4%", "Sultra, Sulteng, Sulsel"],
        ["Perkebunan Kelapa Sawit", "6 Kasus", "18,8%", "Sulbar, Sulteng, Gorontalo"],
        ["Perambahan Hutan Lindung", "5 Kasus", "15,6%", "Sulteng, Sultra, Sulut"],
        ["Hutan Produksi & Konservasi", "5 Kasus", "15,6%", "Gorontalo, Sulsel"],
        ["Infrastruktur & Kawasan Industri", "5 Kasus", "15,6%", "Sulut, Sulteng, Sulsel"],
        ["TOTAL KESELURUHAN", "32 Kasus", "100,0%", "Pulau Sulawesi"]
    ]
    add_table_styled(
        doc,
        headers=["Sektor Penyebab Konflik", "Jumlah Kasus", "Porsi (%)", "Wilayah Terdampak Utama"],
        rows=tabel_7_2_rows,
        col_widths_cm=[5.5, 2.5, 2.0, 7.0],
        alignments=['L', 'C', 'C', 'L'],
        font_pt=7
    )

    # Sub-bab 7.3
    add_h3(doc, "Sub-bab 7.3: Inkonsistensi Iklim — Karpet Merah PLTU Batubara Captive")
    add_body(doc, [
        ("Kontradiksi hilirisasi hijau dihitung dengan menginventarisasi seluruh unit PLTU batubara *off-grid* yang beroperasi khusus melayani pabrik pemurnian nikel, serta melacak timeline penambahan kapasitas kumulatifnya:", False, False)
    ])
    add_formula(doc,
        "1. Akumulasi Kapasitas Pembangkit Fosil:\n"
        "   Total Kapasitas Provinsi = Penjumlahan Seluruh Unit PLTU Captive di Kawasan Industri\n\n"
        "2. Pertumbuhan Kapasitas Kumulatif Tahunan:\n"
        "   Kapasitas Kumulatif (Tahun T) = Total Kapasitas Seluruh Unit yang Mulai Beroperasi s.d. Tahun T",
        ket="Fakta Empiris: Terdata 67 unit PLTU captive batubara dengan total kapasitas 12.245 MW. Sebanyak 55 unit (9.825 MW) telah aktif beroperasi. Sulawesi Tengah menanggung beban terbesar yakni 44 unit (9.365 MW atau 76,5% dari total kapasitas pulau)."
    )

    add_caption(doc, "Tabel 7.3: Agregat Unit dan Kapasitas PLTU Captive Kawasan Industri per Provinsi")
    tabel_7_3_rows = [
        ["Sulawesi Tengah", "44 Unit", "9.365 MW", "7.325 MW", "76,5%", "Episentrum Kawasan Industri IMIP Morowali"],
        ["Sulawesi Tenggara", "13 Unit", "2.280 MW", "1.900 MW", "18,6%", "Sentra Smelter VDNI/OSS Konawe"],
        ["Sulawesi Selatan", "10 Unit", "600 MW", "600 MW", "4,9%", "Kawasan Industri Huadi Bantaeng"],
        ["TOTAL SULAWESI", "67 Unit", "12.245 MW", "9.825 MW", "100,0%", "Paradoks Hilirisasi Bersih"]
    ]
    add_table_styled(
        doc,
        headers=["Provinsi", "Total Unit", "Kapasitas Total", "Sudah Beroperasi", "Porsi Beban", "Keterangan Kawasan"],
        rows=tabel_7_3_rows,
        col_widths_cm=[3.2, 2.0, 2.5, 2.5, 2.0, 4.8],
        alignments=['L', 'C', 'C', 'C', 'C', 'L'],
        font_pt=7
    )

    # ── E. KORESPONDENSI METODOLOGI TERHADAP SUB-BAB LAPORAN ────
    add_h2(doc, "E", "Korespondensi Metodologi terhadap Sub-bab Laporan Bab 7")
    add_body(doc, [
        ("Setiap sub-bab analitis pada Bab 7 dibangun menggunakan metodologi evaluasi empiris yang terstandarisasi sebagaimana dirangkum pada tabel berikut:", False, False)
    ])

    table_korespondensi = [
        ["Sub-bab 7.1", "Status Ekologis vs Penerbitan Izin Tambang", "Spatial Overlay Panel Join, Persentil Deforestasi Binning, Compliance Audit Modeling"],
        ["Sub-bab 7.2", "Tabrakan Hukum & Impunitas Operasi Ilegal", "Incident-based Aggregation, Thematic Coding Kasus Tenurial, Sektoral Disparity Tracking"],
        ["Sub-bab 7.3", "Inkonsistensi Iklim: Karpet Merah PLTU Captive", "Quantitative Asset Inventory, Timeline Tracking Kapasitas Kumulatif, Decoupling Contrast Analysis"]
    ]

    add_table_styled(
        doc,
        headers=["Sub-bab", "Fokus Kajian Empiris", "Metode Analitis Utama"],
        rows=table_korespondensi,
        col_widths_cm=[2.5, 6.0, 8.5],
        alignments=['C', 'L', 'L'],
        font_pt=7.5
    )

    # ── F. BAGAN ALUR KERANGKA KERJA RISET BAB 7 ────────────────
    add_h2(doc, "F", "Bagan Alur Kerangka Kerja Riset (Research Workflow)")
    add_body(doc, [
        ("Kerangka investigasi forensik tata kelola perizinan dijalankan secara terpadu melalui empat tahapan analisis sebagaimana divisualisasikan pada diagram alur berikut:", False, False)
    ])

    mermaid_str_f = """flowchart LR
    subgraph F1["Fase I: Input Multi-Domain"]
        A1["Izin MODI ESDM<br/><i>IUP Baru & Luas Konsesi</i>"]
        A2["Hutan GFW Hansen<br/><i>Deforestasi Master 10 Thn</i>"]
        A3["Konflik KPA & TanahKita<br/><i>32 Kasus Tenurial/HAM</i>"]
        A4["Aset GEM Tracker<br/><i>67 Unit PLTU Captive</i>"]
    end
    subgraph F2["Fase II: Pemrosesan Analitis"]
        B1["Binning Persentil<br/><i>P33 Aman vs P66 Kritis</i>"]
        B2["Panel Join Spasial<br/><i>Provinsi x Tahun (GFW-ESDM)</i>"]
        B3["Thematic Coding<br/><i>Sektor & Impunitas Korporasi</i>"]
        B4["Tracking Timeline MW<br/><i>Pertumbuhan 2013-2024</i>"]
    end
    subgraph F3["Fase III: Uji Kepatuhan & Paradoks"]
        C1["Matriks Seharusnya vs Realita<br/><i>277 IUP di Zona Kritis</i>"]
        C2["Volume Pembiaran Sektoral<br/><i>Tambang 34% & Sultra 25%</i>"]
        C3["Konsentrasi Beban Karbon<br/><i>Sulteng 76.5% (9.365 MW)</i>"]
    end
    subgraph F4["Fase IV: Kesimpulan Tata Kelola"]
        D1["Disregard D3TLH<br/><i>Instrumen Lingkungan Diabaikan</i>"]
        D2["State Omission<br/><i>Impunitas Pelanggaran Korporasi</i>"]
        D3["Paradoks Iklim<br/><i>Hilirisasi Bertenaga Batubara</i>"]
    end
    F1 --> F2 --> F3 --> F4"""

    png_workflow_path = str(out_dir_compact / "mermaid_workflow_bab7_compact.png")
    is_downloaded = download_mermaid_png(mermaid_str_f, png_workflow_path)

    add_caption(doc, "Bagan Alur 7.1: Alur Logika Evaluasi Kegagalan Tata Kelola Perizinan (Research Workflow)")
    if is_downloaded and os.path.exists(png_workflow_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(3)
        p_img.paragraph_format.space_after  = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(png_workflow_path, width=Cm(16.5))
        try:
            shutil.copyfile(png_workflow_path, str(out_dir_bab7 / "mermaid_workflow_bab7_compact.png"))
        except Exception:
            pass

    # Box Output Kesimpulan
    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after  = Pt(4)
    all_border_para(p_box, color='1B5E20', sz='8')
    para_shd(p_box, 'F1F8E9')
    add_run(p_box, "KESIMPULAN METODOLOGIS BAB 7 (KEGAGALAN TATA KELOLA PERIZINAN):\n", bold=True, pt=8.5, color=G_DARK)
    add_run(p_box, "1. Disregard Terhadap D3TLH: Data membuktikan sebanyak 277 izin tambang baru (luas 440.998 Ha) tetap diterbitkan pada kurun waktu ketika deforestasi provinsi berada pada status Kritis, mengonfirmasi bahwa D3TLH dan AMDAL tidak difungsikan sebagai batas pengaman perizinan.\n"
                   "2. Pembiaran Hukum (State Omission): Sebanyak 32 kasus sengketa tenurial dan pelanggaran izin terdata tanpa penegakan sanksi tegas, di mana sektor pertambangan menjadi aktor penyumbang konflik terbesar (34,4%).\n"
                   "3. Paradoks Transisi Energi: Ketergantungan terhadap 67 unit PLTU captive batubara (12.245 MW total kapasitas) menegaskan bahwa rantai pasok hilirisasi nikel beroperasi di atas paradoks emisi karbon fosil yang bertolak belakang dengan komitmen iklim nasional.",
            pt=8, color=C_BODY)

    # ── SIMPAN DOKUMEN DOCX (DUAL SAVE) ─────────────────────────
    docx_compact = out_dir_compact / "Metodologi_Bab7_Kegagalan_Tata_Kelola_Compact.docx"
    docx_bab7    = out_dir_bab7 / "Metodologi_Bab7_Kegagalan_Tata_Kelola_Compact.docx"
    
    doc.save(str(docx_compact))
    shutil.copyfile(docx_compact, docx_bab7)
    print(f"  [OK] Tersimpan DOCX: {docx_compact}")
    print(f"  [OK] Salinan DOCX : {docx_bab7}")

    # ── GENERATE MARKDOWN PADANAN ───────────────────────────────
    print("[2/3] Membangun dokumen Markdown padanan...")
    MD_CONTENT = """# BAB VII: METODOLOGI ANALISIS KEGAGALAN TATA KELOLA — D3TLH DALAM SISTEM PERIZINAN
*Ringkasan Eksekutif Metodologis · Center of Economic and Law Studies (CELIOS)*

---

## A. Desain Penelitian & Tujuan
Penelitian Bab 7 menerapkan **desain audit forensik kepatuhan kebijakan (Compliance & Institutional Failure Audit)** untuk menguji efektivitas instrumen Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) dalam tata kelola perizinan industri ekstraktif di Pulau Sulawesi. Melalui integrasi spasial data perizinan tambang, tutupan hutan, sengketa tenurial, dan infrastruktur energi fosil, kajian ini membuktikan secara empiris tiga kegagalan struktural tata kelola:

1. **Evaluasi Kepatuhan D3TLH vs Penerbitan Izin (Rule-based Categorization):** Menganalisis sinkronisasi antara batas pengaman daya dukung hutan (GFW) dan izin baru pertambangan (MODI ESDM), guna membuktikan apakah instrumen D3TLH berfungsi sebagai pembatas atau diabaikan dalam zona krisis ekologis.
2. **Audit Impunitas Hukum & Pembiaran Korporasi (Thematic Coding):** Menginventarisasi rekam jejak pelanggaran hukum, perampasan ruang hidup, dan sengketa agraria yang dibiarkan tanpa penegakan sanksi administratif maupun pidana lingkungan (*state omission*).
3. **Kuantifikasi Kontradiksi Karbon PLTU Captive (Asset-level Inventory):** Mendokumentasikan paradoks hilirisasi hijau melalui inventarisasi aset 67 unit pembangkit listrik tenaga uap (PLTU) batubara *off-grid* captive yang beroperasi di dalam kawasan industri nikel.

---

## B. Sumber Data & Cakupan Wilayah
Audit tata kelola perizinan ini menggabungkan 4 basis data resmi lintas kementerian dan lembaga masyarakat sipil yang mencakup seluruh yurisdiksi 6 provinsi se-Pulau Sulawesi kurun 2014–2024:

- **MODI Ditjen Minerba ESDM RI:** Data geospasial izin baru pertambangan (IUP), tahun penerbitan, luas konsesi (hektar), dan komoditas tambang aktif.
- **Global Forest Watch (GFW / Hansen UMD):** Time-series deforestasi tahunan tingkat provinsi guna menetapkan ambang persentil daya dukung hutan alam.
- **Konsorsium Pembaruan Agraria (CATAHU KPA), TanahKita, & Koalisi Sipil:** Dokumentasi kasus konflik tenurial, pelanggaran izin di kawasan lindung, kriminalisasi warga, dan impunitas korporasi.
- **Global Energy Monitor (GEM Coal Plant Tracker, Jan 2026):** Inventarisasi aset 67 unit PLTU captive batubara di kawasan industri Morowali, Konawe, Bantaeng, dan sekitarnya.

---

## C. Operasionalisasi Variabel & Indikator Riset
Seluruh parameter tata kelola, pelanggaran batas daya dukung, sengketa tenurial, hingga aset energi fosil dioperasionalkan ke dalam **9 indikator riset empiris kunci** sebagaimana dirangkum pada matriks operasional berikut:

##### Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 7 (Tata Kelola Perizinan)
| No | Indikator Riset | Fokus Pengukuran | Satuan | Sumber Data Primer Resmi |
| :-: | :--- | :--- | :-: | :--- |
| 1 | Kepatuhan D3TLH Perizinan (7.1) | Rasio Penerbitan IUP Baru pada Zona Deforestasi Kritis | Unit Izin | MODI ESDM & GFW Hansen |
| 2 | Luas Konsesi Zona Kritis (7.1) | Total Luasan Konsesi Tambang Terbit di Status Kritis | Hektar (Ha) | Ditjen Minerba ESDM & GFW |
| 3 | Ambang Persentil Daya Dukung (7.1) | Klasifikasi Status Hutan: Aman (P33), Tertekan, Kritis (P66) | Hektar / Tahun | GFW Master Time-Series 2014-2023 |
| 4 | Kasus Impunitas Penegakan Hukum (7.2) | Frekuensi Pembiaran Pelanggaran Korporasi Ekstraktif | Kasus | CATAHU KPA, TanahKita, YLBHI |
| 5 | Dominasi Konflik Sektoral (7.2) | Proporsi Sengketa Agraria Akibat Ekspansi Tambang | Persen (%) | Database Kasus Tenurial KPA |
| 6 | Konsentrasi Spasial Sengketa (7.2) | Sebaran Wilayah Konflik Tertinggi Lintas Provinsi | Kasus & % | Laporan Koalisi Sipil & KPA |
| 7 | Agregat Aset PLTU Captive (7.3) | Total Unit & Kapasitas Pembangkitan Kotor Batubara Off-Grid | Unit & Megawatt (MW) | Global Energy Monitor (GEM 2026) |
| 8 | Kapasitas PLTU Captive Aktif (7.3) | Kapasitas Terpasang Unit Batubara Beroperasi (Operating) | Megawatt (MW) | GEM Global Coal Plant Tracker |
| 9 | Dominasi Beban Karbon Spasial (7.3) | Konsentrasi Kapasitas Pembangkitan Fosil per Provinsi | Persen (%) | GEM Asset Inventory Tracker |

---

## D. Kerangka Analisis & Formulasi Matematis

### Sub-bab 7.1: Evaluasi Kepatuhan D3TLH vs Penerbitan Izin Tambang Baru
Penilaian kepatuhan sistem perizinan terhadap batas pengaman lingkungan hidup dilakukan dengan mengelompokkan data deforestasi tahunan ke dalam 3 kelas persentil daya dukung, kemudian menyandingkannya dengan penerbitan IUP baru:

> **1. Penentuan Ambang Batas Daya Dukung (Binning Persentil):**  
> • Ambang Tertekan = Persentil-33 (P33) dari Deforestasi Tahunan = 12.898 Ha/tahun  
> • Ambang Kritis = Persentil-66 (P66) dari Deforestasi Tahunan = 26.453 Ha/tahun  
>  
> **2. Aturan Klasifikasi Status Wilayah:**  
> • **Status Aman:** Kerusakan Hutan ≤ 12.898 Ha/tahun (Izin Wajar Diterbitkan)  
> • **Status Tertekan:** 12.898 Ha < Kerusakan ≤ 26.453 Ha/tahun (Penerbitan Izin Mulai Dibatasi)  
> • **Status Kritis:** Kerusakan Hutan > 26.453 Ha/tahun (Wajib Moratorium Total)  
>  
> **3. Kuantifikasi Pelanggaran Tata Kelola:**  
> `Total Izin Zona Kritis = Jumlah IUP Baru yang Tetap Diterbitkan saat Status Hutan Kritis`  
> *Fakta Empiris Lapangan: Pada saat status hutan dinyatakan KRITIS (> 26.453 Ha deforestasi), pemerintah justru menerbitkan 277 IUP BARU dengan total luas konsesi mencapai 440.998 Hektar (mayoritas di Sulteng dan Sulsel).*

##### Tabel 7.1: Matriks Kepatuhan D3TLH — Seharusnya vs Kenyataan Penerbitan Izin
| Status Daya Dukung | Deforestasi Tahunan | N Observasi | Aturan Seharusnya | Kenyataan di Lapangan | Luas Konsesi (Ha) | Kesimpulan Tata Kelola |
| :---: | :---: | :---: | :--- | :---: | :---: | :---: |
| Aman | <= 12.898 Ha | 20 | Wajar diterbitkan izin | 26 Izin Baru Keluar | 87.070 Ha | Normal (Sesuai Aturan) |
| Tertekan | 12.898 - 26.453 Ha | 19 | Izin mulai dibatasi ketat | 77 Izin Baru Keluar | 107.377 Ha | Anomali (Lampu Kuning) |
| Kritis | > 26.453 Ha | 21 | Moratorium Total / Larangan Izin | 277 Izin Baru Keluar | 440.998 Ha | PELANGGARAN STRUKTURAL |

---

### Sub-bab 7.2: Pemetaan Impunitas Hukum & Pembiaran Kasus Tenurial
Tingkat pembiaran negara (*state omission*) dihitung melalui klasifikasi tematik terhadap seluruh insiden konflik agraria, pelanggaran kawasan lindung, dan intimidasi warga yang terdokumentasi tanpa adanya sanksi hukum tegas:

> **1. Penghitungan Proporsi Konflik Sektoral:**  
> `Proporsi Sektor (%) = (Jumlah Kasus pada Sektor Tertentu / Total Kasus Terdata) × 100%`  
>  
> **2. Identifikasi Wilayah Episentrum Pembiaran:**  
> `Proporsi Wilayah (%) = (Jumlah Kasus di Suatu Provinsi / Total Kasus Terdata) × 100%`  
>  
> *Fakta Empiris: Dari 32 kasus impunitas yang terdata di Sulawesi, Sektor Pertambangan menjadi penyumbang terbesar dengan 11 kasus (34,4%), disusul Perkebunan Sawit 6 kasus (18,8%). Sulawesi Tenggara mencatat kasus terbanyak (8 kasus / 25,0%), disusul Sulawesi Selatan (7 kasus / 21,9%).*

##### Tabel 7.2: Sebaran Sektor Konflik dan Pembiaran Operasi Ilegal di Sulawesi
| Sektor Penyebab Konflik | Jumlah Kasus | Porsi (%) | Wilayah Terdampak Utama |
| :--- | :---: | :---: | :--- |
| Pertambangan (Nikel & Batuan) | 11 Kasus | 34,4% | Sultra, Sulteng, Sulsel |
| Perkebunan Kelapa Sawit | 6 Kasus | 18,8% | Sulbar, Sulteng, Gorontalo |
| Perambahan Hutan Lindung | 5 Kasus | 15,6% | Sulteng, Sultra, Sulut |
| Hutan Produksi & Konservasi | 5 Kasus | 15,6% | Gorontalo, Sulsel |
| Infrastruktur & Kawasan Industri | 5 Kasus | 15,6% | Sulut, Sulteng, Sulsel |
| **TOTAL KESELURUHAN** | **32 Kasus** | **100,0%** | **Pulau Sulawesi** |

---

### Sub-bab 7.3: Inkonsistensi Iklim — Karpet Merah PLTU Batubara Captive
Kontradiksi hilirisasi hijau dihitung dengan menginventarisasi seluruh unit PLTU batubara *off-grid* yang beroperasi khusus melayani pabrik pemurnian nikel, serta melacak timeline penambahan kapasitas kumulatifnya:

> **1. Akumulasi Kapasitas Pembangkit Fosil:**  
> `Total Kapasitas Provinsi = Penjumlahan Seluruh Unit PLTU Captive di Kawasan Industri`  
>  
> **2. Pertumbuhan Kapasitas Kumulatif Tahunan:**  
> `Kapasitas Kumulatif (Tahun T) = Total Kapasitas Seluruh Unit yang Mulai Beroperasi s.d. Tahun T`  
>  
> *Fakta Empiris: Terdata 67 unit PLTU captive batubara dengan total kapasitas 12.245 MW. Sebanyak 55 unit (9.825 MW) telah aktif beroperasi. Sulawesi Tengah menanggung beban terbesar yakni 44 unit (9.365 MW atau 76,5% dari total kapasitas pulau).*

##### Tabel 7.3: Agregat Unit dan Kapasitas PLTU Captive Kawasan Industri per Provinsi
| Provinsi | Total Unit | Kapasitas Total | Sudah Beroperasi | Porsi Beban | Keterangan Kawasan |
| :--- | :---: | :---: | :---: | :---: | :--- |
| Sulawesi Tengah | 44 Unit | 9.365 MW | 7.325 MW | 76,5% | Episentrum Kawasan Industri IMIP Morowali |
| Sulawesi Tenggara | 13 Unit | 2.280 MW | 1.900 MW | 18,6% | Sentra Smelter VDNI/OSS Konawe |
| Sulawesi Selatan | 10 Unit | 600 MW | 600 MW | 4,9% | Kawasan Industri Huadi Bantaeng |
| **TOTAL SULAWESI** | **67 Unit** | **12.245 MW** | **9.825 MW** | **100,0%** | **Paradoks Hilirisasi Bersih** |

---

## E. Korespondensi Metodologi terhadap Sub-bab Laporan Bab 7
Setiap sub-bab analitis pada Bab 7 dibangun menggunakan metodologi evaluasi empiris yang terstandarisasi sebagaimana dirangkum pada tabel berikut:

##### Matriks Korespondensi Metodologis Bab 7
| Sub-bab | Fokus Kajian Empiris | Metode Analitis Utama |
| :-: | :--- | :--- |
| Sub-bab 7.1 | Status Ekologis vs Penerbitan Izin Tambang | Spatial Overlay Panel Join, Persentil Deforestasi Binning, Compliance Audit Modeling |
| Sub-bab 7.2 | Tabrakan Hukum & Impunitas Operasi Ilegal | Incident-based Aggregation, Thematic Coding Kasus Tenurial, Sektoral Disparity Tracking |
| Sub-bab 7.3 | Inkonsistensi Iklim: Karpet Merah PLTU Captive | Quantitative Asset Inventory, Timeline Tracking Kapasitas Kumulatif, Decoupling Contrast Analysis |

---

## F. Bagan Alur Kerangka Kerja Riset (Research Workflow)
Kerangka investigasi forensik tata kelola perizinan dijalankan secara terpadu melalui empat tahapan analisis sebagaimana divisualisasikan pada diagram alur berikut:

```mermaid
flowchart LR
    subgraph F1["Fase I: Input Multi-Domain"]
        A1["Izin MODI ESDM<br/><i>IUP Baru & Luas Konsesi</i>"]
        A2["Hutan GFW Hansen<br/><i>Deforestasi Master 10 Thn</i>"]
        A3["Konflik KPA & TanahKita<br/><i>32 Kasus Tenurial/HAM</i>"]
        A4["Aset GEM Tracker<br/><i>67 Unit PLTU Captive</i>"]
    end
    subgraph F2["Fase II: Pemrosesan Analitis"]
        B1["Binning Persentil<br/><i>P33 Aman vs P66 Kritis</i>"]
        B2["Panel Join Spasial<br/><i>Provinsi x Tahun (GFW-ESDM)</i>"]
        B3["Thematic Coding<br/><i>Sektor & Impunitas Korporasi</i>"]
        B4["Tracking Timeline MW<br/><i>Pertumbuhan 2013-2024</i>"]
    end
    subgraph F3["Fase III: Uji Kepatuhan & Paradoks"]
        C1["Matriks Seharusnya vs Realita<br/><i>277 IUP di Zona Kritis</i>"]
        C2["Volume Pembiaran Sektoral<br/><i>Tambang 34% & Sultra 25%</i>"]
        C3["Konsentrasi Beban Karbon<br/><i>Sulteng 76.5% (9.365 MW)</i>"]
    end
    subgraph F4["Fase IV: Kesimpulan Tata Kelola"]
        D1["Disregard D3TLH<br/><i>Instrumen Lingkungan Diabaikan</i>"]
        D2["State Omission<br/><i>Impunitas Pelanggaran Korporasi</i>"]
        D3["Paradoks Iklim<br/><i>Hilirisasi Bertenaga Batubara</i>"]
    end
    F1 --> F2 --> F3 --> F4
```

> **KESIMPULAN METODOLOGIS BAB 7 (KEGAGALAN TATA KELOLA PERIZINAN):**  
> 1. **Disregard Terhadap D3TLH:** Data membuktikan sebanyak 277 izin tambang baru (luas 440.998 Ha) tetap diterbitkan pada kurun waktu ketika deforestasi provinsi berada pada status Kritis, mengonfirmasi bahwa D3TLH dan AMDAL tidak difungsikan sebagai batas pengaman perizinan.  
> 2. **Pembiaran Hukum (State Omission):** Sebanyak 32 kasus sengketa tenurial dan pelanggaran izin terdata tanpa penegakan sanksi tegas, di mana sektor pertambangan menjadi aktor penyumbang konflik terbesar (34,4%).  
> 3. **Paradoks Transisi Energi:** Ketergantungan terhadap 67 unit PLTU captive batubara (12.245 MW total kapasitas) menegaskan bahwa rantai pasok hilirisasi nikel beroperasi di atas paradoks emisi karbon fosil yang bertolak belakang dengan komitmen iklim nasional.
"""

    md_compact = out_dir_compact / "Metodologi_Bab7_Kegagalan_Tata_Kelola_Compact.md"
    md_bab7    = out_dir_bab7 / "Metodologi_Bab7_Kegagalan_Tata_Kelola_Compact.md"
    with open(md_compact, 'w', encoding='utf-8') as f:
        f.write(MD_CONTENT)
    shutil.copyfile(md_compact, md_bab7)
    print(f"  [OK] Tersimpan MD  : {md_compact}")
    print(f"  [OK] Salinan MD   : {md_bab7}")

    print("[3/3] Selesai menghasilkan dokumen metodologi Bab 7 versi compact (1-Kolom, 2-3 Halaman).\n")


if __name__ == "__main__":
    generate_bab7_compact()
