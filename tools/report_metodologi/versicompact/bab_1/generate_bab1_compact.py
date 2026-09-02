#!/usr/bin/env python3
"""
Generator Laporan Metodologi Statistik (Versi Compact) Bab 1:
Ekspansi Industri Ekstraktif dan Infrastruktur Penunjang di Pulau Sulawesi

Format: Standar Versi Compact CELIOS
- Header dan Judul SAMA PERSIS dengan dokumen non-compact / root
- Penomoran sub-bab persis: 1.1, 1.2, 1.3, dst.
- Tanpa icon / emoji sama sekali
- Tanpa improvisasi narasi (semua narasi, rumus, dan tabel murni bersumber dari Metodologi_Bab1_Ekspansi_Industri.md)

Output:
1. tools/report_metodologi/versicompact/bab_1/Metodologi_Bab1_Ekspansi_Industri_Compact.docx
2. tools/report_metodologi/versicompact/bab_1/Metodologi_Bab1_Ekspansi_Industri_Compact.md
3. tools/report_metodologi/bab_1/Metodologi_Bab1_Ekspansi_Industri_Compact.docx
4. tools/report_metodologi/bab_1/Metodologi_Bab1_Ekspansi_Industri_Compact.md
"""

import os
import sys
import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    print("[INFO] Memasang modul python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Warna Tema CELIOS D3TLH ─────────────────────────────────
G_DARK   = RGBColor(0x1B, 0x5E, 0x20)  # Hijau Hutan Gelap (#1B5E20)
G_MID    = RGBColor(0x2E, 0x7D, 0x32)  # Hijau Utama CELIOS (#2E7D32)
G_ACC    = RGBColor(0x43, 0xA0, 0x47)  # Hijau Aksen (#43A047)
C_BODY   = RGBColor(0x22, 0x22, 0x22)  # Abu Gelap Teks (#222222)
C_GREY   = RGBColor(0x55, 0x55, 0x55)  # Abu Sekunder (#555555)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)  # Putih
C_RED    = RGBColor(0xB7, 0x1C, 0x1C)  # Merah Kritis (#B71C1C)

# ── Helper XML Word Formatting ──────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=100, right=100, top=60, bottom=60):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def para_shd(p, fill):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    pPr.append(s)

def cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    tcPr.append(s)

def para_border_bottom(p, color='2E7D32', sz='6'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '2')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def para_border_left(p, color='2E7D32', sz='18'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:left')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '6')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='444444', sz='4'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Konten & Tipografi ──────────────────────────────
def run(p, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = 'Courier New'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Courier New')
    return r

def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    para_border_bottom(p, color='1B5E20', sz='12')
    run(p, title.upper(), bold=True, pt=12, color=G_DARK)

def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    para_border_bottom(p, color='2E7D32', sz='6')
    run(p, title, bold=True, pt=11, color=G_DARK)

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after  = Pt(3)
    run(p, title, bold=True, pt=10, color=G_MID)

def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_DARK)

def add_p(doc, parts, space_after=4, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p

def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Pt(10)
    para_shd(p, 'EDF7EE')
    all_border_para(p, color='A5D6A7', sz='4')
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)

    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(1)
        p_desc.paragraph_format.space_after  = Pt(4)
        p_desc.paragraph_format.left_indent  = Pt(12)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=RGBColor(0x1B, 0x5E, 0x20))
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"{item[1]}{trailing}", italic=False, pt=8, color=RGBColor(0x44, 0x44, 0x44))

def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(10)
    para_border_left(p, color='2E7D32', sz='16')
    para_shd(p, 'F1F8E9')
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.autofit = False

    bd_cfg = {'val': 'single', 'sz': '4', 'color': 'D0D0D0', 'space': '0'}

    # Header
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=80, right=80, top=60, bottom=60)
        set_cell_borders(c, top=bd_cfg, left=bd_cfg, bottom={'val': 'single', 'sz': '8', 'color': '1B5E20', 'space': '0'}, right=bd_cfg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (alignments and alignments[j] == 'C') else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run(p, h, bold=True, pt=8.5, color=C_WHITE)

    # Rows
    for i, row in enumerate(rows):
        bg = 'F9FBF9' if i % 2 == 1 else 'FFFFFF'
        for j, (val, w) in enumerate(zip(row, col_widths_cm)):
            c = tbl.rows[1 + i].cells[j]
            c.width = Cm(w)
            cell_shd(c, bg)
            cell_margin(c, left=80, right=80, top=40, bottom=40)
            set_cell_borders(c, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = c.paragraphs[0]
            align = alignments[j] if alignments else 'L'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            is_bold = (j == 0) or ('Total' in str(val)) or ('Ekstraktif' in str(val))
            run(p, str(val), bold=is_bold, pt=8.5, color=C_BODY)

# ── Main Generator ──────────────────────────────────────────
def build_compact_report():
    print("[1/4] Menginisialisasi dokumen Word python-docx...")
    doc = Document()

    # Margin Halaman (2.0 cm sekeliling)
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    # ── HEADER BANNER ───────────────────────────────────────────
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(2)
    run(p_hdr, "CELIOS — CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)

    add_h1(doc, "BAB I: METODOLOGI ANALISIS EKSPANSI INDUSTRI EKSTRAKTIF DAN INFRASTRUKTUR PENUNJANG DI PULAU SULAWESI")

    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, landasan regulasi, formulasi matematis, prosedur analisis statistik, serta metodologi pembuktian berbasis data terbuka yang dioperasionalkan pada ", False, False),
        ("Bab 1: Ekspansi Industri Ekstraktif", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi periode 2014–2024.", False, False),
    ], space_after=8)

    # ═══════════════════════════════════════════════════════════
    # 1.1 KONTEKS MAKRO: BREAKDOWN PDRB PER KOMODITAS
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.1 Konteks Makro: Breakdown PDRB per Komoditas")

    # 1.1.1
    add_h3(doc, "1.1.1 Konteks Makro: Dominasi Ekstraktif vs Ekonomi Akar Rumput")
    add_p(doc, [
        ("Bagian ini menganalisis struktur Produk Domestik Regional Bruto (PDRB) pada enam provinsi di Pulau Sulawesi sepanjang periode 2016–2024 menggunakan visualisasi grafik area bertumpuk (*Stacked Area Chart*). Analisis ini ditujukan untuk menguji secara empiris apakah percepatan pertumbuhan ekonomi daerah benar-benar bersumber dari sektor produktif masyarakat lokal atau didominasi oleh industri ekstraktif padat modal yang mengalihkan pemanfaatan ruang dan sumber daya alam.", False, False)
    ])
    add_note_box(
        doc,
        "Sumber Data",
        "Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi Stacked Area Chart memetakan dinamika Produk Domestik Regional Bruto (PDRB) berdasarkan klasifikasi rantai pasok hukum (Legal Supply-Chain) untuk membandingkan trajektori Sektor Ekstraktif, Ekonomi Akar Rumput, dan Sektor Jasa & Lainnya."
    )

    add_caption(doc, "Tabel 1.1: Reklasifikasi Sektoral PDRB KBLI 2020 Berdasarkan Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)")
    reklas_headers = ["Kategori BPS", "Sektor Lapangan Usaha", "Klasifikasi Analisis", "Dasar Regulasi & Mandat Hukum", "Intisari Ketentuan Hukum"]
    reklas_rows = [
        ["Kategori B", "Pertambangan dan Penggalian", "Ekstraktif", "Perpres No. 26 Tahun 2010", "Ketentuan Pasal 1 Ayat (2) mengenai pengambilan komoditas tambang dari dalam bumi."],
        ["Kategori C", "Industri Pengolahan (Smelter Logam)", "Ekstraktif", "UU No. 3 Tahun 2020 & PP No. 96 Tahun 2021", "Pasal 102–103 mewajibkan pengolahan dan pemurnian di dalam negeri sebagai kesatuan pertambangan."],
        ["Kategori D", "Pengadaan Listrik & Gas (PLTU Captive)", "Ekstraktif", "Perpres No. 112 Tahun 2022 & RUPTL PLN", "Pasal 3 Ayat (4) huruf b mengecualikan PLTU baru hanya bagi yang terintegrasi melayani smelter."],
        ["Kategori A", "Pertanian, Kehutanan, Perikanan", "Ekonomi Akar Rumput", "KBLI 2020 BPS", "Sektor pemanfaatan sumber daya hayati terbarukan dan penyerap tenaga kerja lokal."],
        ["Kategori E–U", "13 Sektor Jasa & Konstruksi", "Sektor Jasa & Lainnya", "Klasifikasi Standar BPS", "Sektor sekunder dan tersier penunjang perekonomian daerah."]
    ]
    add_table(doc, reklas_headers, reklas_rows, [2.2, 3.8, 2.5, 4.2, 4.3], ['C', 'L', 'C', 'L', 'L'])

    add_formula(
        doc,
        "Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)",
        "Sektor_Ekstraktif = PDRB(Kat.B: Pertambangan) + PDRB(Kat.C: Ind. Pengolahan) + PDRB(Kat.D: Listrik)\n"
        "Sektor_Akar_Rumput = PDRB(Kat.A: Pertanian, Kehutanan, dan Perikanan)\n"
        "Sektor_Jasa = Jumlah PDRB (Kategori E sampai dengan Kategori U)\n"
        "Total_PDRB = Sektor_Ekstraktif + Sektor_Akar_Rumput + Sektor_Jasa\n"
        "Pangsa_Ekstraktif (%) = ( Sektor_Ekstraktif / Total_PDRB ) * 100\n"
        "Laju_Pertumbuhan_Tahunan (%) = [ ( Nilai_Tahun_t - Nilai_Tahun_{t-1} ) / Nilai_Tahun_{t-1} ] * 100",
        [
            ("Sektor_Ekstraktif", "Total nilai tambah bruto dari klaster industri ekstraktif yang saling terintegrasi (Triliun Rupiah)."),
            ("Sektor_Akar_Rumput", "Nilai PDRB pemanfaatan sumber daya hayati terbarukan (Triliun Rupiah)."),
            ("Sektor_Jasa", "Nilai tambah 13 sektor penunjang sekunder dan tersier (Triliun Rupiah)."),
            ("Total_PDRB", "Total nilai Produk Domestik Regional Bruto wilayah atas dasar harga berlaku (Triliun Rupiah)."),
            ("Pangsa_Ekstraktif (%)", "Persentase pangsa dominasi sektor ekstraktif terhadap total ekonomi (%)."),
            ("Laju_Pertumbuhan_Tahunan (%)", "Tingkat percepatan/perlambatan ekspansi tahunan sektor ekonomi (%).")
        ]
    )

    add_caption(doc, "Tabel 1.2: Definisi Operasional Komponen Makroekonomi dan Sumber Data PDRB Sektoral")
    def_headers = ["Komponen Analisis", "Cakupan Lapangan Usaha", "Definisi Operasional", "Satuan Nilai", "Sumber Data Primer"]
    def_rows = [
        ["Sektor Ekstraktif", "Kategori B, Kategori C, Kategori D", "Akumulasi nilai tambah pertambangan nikel, smelter logam dasar, dan PLTU captive.", "Triliun Rupiah", "BPS Provinsi (SIMDASI)"],
        ["Ekonomi Akar Rumput", "Kategori A", "Nilai tambah pertanian, perkebunan, kehutanan, dan perikanan.", "Triliun Rupiah", "BPS Provinsi"],
        ["Sektor Jasa & Lainnya", "Kategori E hingga U", "Nilai tambah gabungan perdagangan, konstruksi, transportasi, keuangan, dan jasa.", "Triliun Rupiah", "BPS Provinsi"],
        ["Total PDRB Wilayah", "Seluruh 17 Kategori", "Total nilai PDRB wilayah atas dasar harga berlaku pada tahun berjalan.", "Triliun Rupiah", "BPS Provinsi"],
        ["Pangsa Ekstraktif (%)", "Rasio Kontribusi", "Persentase kontribusi sektor ekstraktif terhadap total perekonomian.", "Persen (%)", "Hasil Olahan CELIOS"]
    ]
    add_table(doc, def_headers, def_rows, [3.2, 3.8, 4.5, 2.2, 3.3], ['L', 'L', 'L', 'C', 'L'])

    add_h4(doc, "Analisis Temuan Empiris: Ketimpangan Struktural Sulawesi Tengah")
    add_p(doc, [
        ("Penerapan formulasi di atas menunjukkan bahwa di ", False, False),
        ("Sulawesi Tengah (sebagai pusat hilirisasi)", True, False),
        (", ekspansi industri ekstraktif menguasai ", False, False),
        ("55.8% dari total PDRB provinsi", True, False),
        (" pada tahun 2024, memperlihatkan dominasi yang sangat kuat dibanding provinsi lainnya. Sektor ekstraktif melonjak dari Rp28,45 triliun pada 2016 menjadi Rp210,51 triliun pada 2024.", False, False)
    ])

    # 1.1.2
    add_h3(doc, "1.1.2 Pemusatan Sektor Ekstraktif di Kabupaten se-Sulawesi Tengah")
    add_p(doc, [
        ("Jika dianalisis secara spasial pada tingkat kabupaten di Sulawesi Tengah, terlihat konsentrasi kegiatan industri ekstraktif. Kabupaten ", False, False),
        ("Morowali", True, False),
        (" dan ", False, False),
        ("Morowali Utara", True, False),
        (" mendominasi struktur PDRB provinsi melalui pengembangan kawasan industri hilirisasi dan PLTU Captive. Analisis ini membandingkan komposisi ketiga sektor advokatif di seluruh 13 kabupaten/kota se-Sulawesi Tengah pada tahun terbaru (2024).", False, False)
    ])
    add_note_box(
        doc,
        "Sumber Data",
        "Badan Pusat Statistik (BPS) Kabupaten se-Sulawesi Tengah (diolah CELIOS). Visualisasi Stacked Bar Chart memetakan struktur Produk Domestik Regional Bruto (PDRB) tahun 2024 pada seluruh 13 kabupaten/kota untuk mengidentifikasi tingkat konsentrasi sektoral dan polarisasi spasial antara sentra industri pengolahan nikel dengan daerah non-sentra."
    )

    add_formula(
        doc,
        "Agregasi Sektoral Kabupaten",
        "Sektor_Ekstraktif_Kabupaten = PDRB_Kab(Kat.B: Pertambangan) + PDRB_Kab(Kat.C: Ind. Pengolahan) + PDRB_Kab(Kat.D: Listrik)\n"
        "Total_PDRB_Kabupaten = Sektor_Ekstraktif_Kabupaten + Sektor_Akar_Rumput_Kabupaten + Sektor_Jasa_Kabupaten\n"
        "Porsi_Sektor_Kabupaten (%) = ( Nilai_Sektor_Kabupaten / Total_PDRB_Kabupaten ) * 100",
        [
            ("Sektor_Ekstraktif_Kabupaten", "Total nilai tambah sektor ekstraktif di tingkat kabupaten target (Triliun Rupiah)."),
            ("Total_PDRB_Kabupaten", "Total output perekonomian bruto kabupaten target atas dasar harga berlaku (Triliun Rupiah)."),
            ("Porsi_Sektor_Kabupaten (%)", "Persentase kontribusi sektor target terhadap total PDRB kabupaten.")
        ]
    )

    add_caption(doc, "Tabel 1.3: Distribusi Nilai Tambah Bruto dan Komposisi Sektoral PDRB 13 Kabupaten/Kota di Sulawesi Tengah (Tahun 2024)")
    kab_headers = ["Kabupaten / Kota", "Akar Rumput (T Rp)", "Ekstraktif (T Rp)", "Jasa (T Rp)", "Total PDRB (T Rp)", "Porsi Akar (%)", "Porsi Eks (%)", "Porsi Jasa (%)", "Basis Utama Ekonomi"]
    kab_rows = [
        ["Morowali", "2.70", "157.17", "187.85", "347.72", "0.8%", "45.2%", "54.0%", "Hilirisasi Nikel (Smelter & PLTU)"],
        ["Banggai", "8.85", "20.63", "51.99", "81.47", "10.9%", "25.3%", "63.8%", "Migas, Tambang & Perdagangan"],
        ["Palu", "1.24", "4.56", "60.03", "65.84", "1.9%", "6.9%", "91.2%", "Jasa, Perdagangan & Pemerintahan"],
        ["Morowali Utara", "5.17", "19.22", "36.08", "60.47", "8.5%", "31.8%", "59.7%", "Hilirisasi Nikel (Smelter GNI)"],
        ["Parigi Moutong", "9.97", "1.93", "35.05", "46.95", "21.2%", "4.1%", "74.7%", "Pertanian Pangan & Hortikultura"],
        ["Donggala", "5.96", "3.53", "23.57", "33.05", "18.0%", "10.7%", "71.3%", "Pertanian, Perkebunan & Galian C"],
        ["Poso", "4.96", "0.40", "20.12", "25.48", "19.5%", "1.6%", "79.0%", "Pertanian & Perkebunan Kakao"],
        ["Sigi", "5.17", "0.83", "19.13", "25.12", "20.6%", "3.3%", "76.1%", "Pertanian Pangan & Hortikultura"],
        ["Toli-Toli", "4.35", "0.44", "17.36", "22.15", "19.7%", "2.0%", "78.4%", "Perkebunan Cengkeh & Perikanan"],
        ["Buol", "3.77", "1.15", "10.67", "15.58", "24.2%", "7.4%", "68.5%", "Kelapa Sawit & Tanaman Pangan"],
        ["Tojo Una-Una", "2.88", "0.61", "11.21", "14.71", "19.6%", "4.2%", "76.2%", "Pertanian & Pariwisata Bahari"],
        ["Banggai Kepulauan", "2.53", "0.18", "8.04", "10.75", "23.5%", "1.7%", "74.8%", "Perikanan Tangkap & Kelautan"],
        ["Banggai Laut", "1.80", "0.12", "4.52", "6.45", "27.9%", "1.9%", "70.2%", "Perikanan & Budidaya Laut"]
    ]
    add_table(doc, kab_headers, kab_rows, [2.5, 1.6, 1.6, 1.6, 1.8, 1.5, 1.5, 1.5, 3.4], ['L', 'R', 'R', 'R', 'R', 'C', 'C', 'C', 'L'])

    add_h4(doc, "Analisis Temuan Empiris: Polarisasi Ekstrem Morowali vs Daerah Non-Smelter")
    add_p(doc, [
        ("Data empiris pada Tabel 1.3 mengungkap bukti polarisasi ekonomi wilayah yang sangat ekstrem di Sulawesi Tengah:\n", False, False),
        ("1. Dominasi Sektor Ekstraktif Morowali: ", True, False),
        ("Kabupaten Morowali mencatatkan nilai sektor ekstraktif sebesar Rp 157.17 Triliun atau menguasai porsi 45.2% dari total kue ekonomi kabupatennya (Rp 347.72 Triliun). Nilai sektor ekstraktif Morowali saja melampaui gabungan total PDRB dari delapan kabupaten lainnya di Sulawesi Tengah.\n", False, False),
        ("2. Pemusatan pada Dua Sentra Hilirisasi: ", True, False),
        ("Kabupaten Morowali dan Morowali Utara merupakan dua daerah dengan nilai Sektor Ekstraktif tertinggi di Sulawesi Tengah, membuktikan bahwa percepatan output industri pertambangan dan hilirisasi terkunci pada kawasan industri smelter.\n", False, False),
        ("3. Ketertinggalan Daerah Non-Sentra: ", True, False),
        ("Sebaliknya, delapan kabupaten lainnya (seperti Banggai Laut, Banggai Kepulauan, Tojo Una-Una, Buol, Toli-Toli, Sigi, Poso, dan Donggala) memiliki porsi Sektor Ekstraktif yang sangat rendah (<11%) dan tetap bergantung pada sektor pertanian rakyat (Akar Rumput) berproduktivitas rendah dengan keterbatasan akses terhadap nilai tambah modal.", False, False)
    ])

    # 1.1.3
    add_h3(doc, "1.1.3 Perbandingan Distribusi 17 Sektor Komoditas per Provinsi (Small Multiples, Tahun Terbaru)")
    add_p(doc, [
        ("Visualisasi komparatif ", False, False),
        ("Small Multiples Horizontal Bar Chart", True, False),
        (" membedah struktur 17 sektor lapangan usaha KBLI 2020 secara terpisah pada enam provinsi di Pulau Sulawesi pada tahun terbaru (2024). Setiap panel provinsi menampilkan sektor yang diurutkan dari penyumbang terbesar hingga terkecil dengan skala sumbu nilai yang disetarakan secara seragam untuk memastikan validitas komparasi lintas wilayah.", False, False)
    ])

    add_caption(doc, "Tabel 1.4: 5 Sektor Lapangan Usaha Penyumbang Utama PDRB di 6 Provinsi Sulawesi (Tahun 2024)")
    s17_headers = ["Provinsi", "Sektor Utama 1", "Porsi (%)", "Sektor Utama 2", "Porsi (%)", "Tipologi Wilayah"]
    s17_rows = [
        ["Sulawesi Tengah", "Industri Pengolahan (Smelter)", "44,1%", "Pertambangan & Penggalian", "11,8%", "Didominasi Industri Pengolahan Smelter & Pertambangan (Ekstraktif)"],
        ["Sulawesi Tenggara", "Pertanian & Perikanan", "22,4%", "Pertambangan Logam", "20,9%", "Didominasi Pertanian & Pertambangan Logam (Campuran)"],
        ["Sulawesi Selatan", "Pertanian & Perikanan", "21,8%", "Perdagangan Besar & Eceran", "14,6%", "Didominasi Pertanian, Perdagangan & Konstruksi (Agraris & Jasa)"],
        ["Sulawesi Utara", "Pertanian & Perikanan", "20,5%", "Perdagangan Besar & Eceran", "13,2%", "Didominasi Pertanian, Perdagangan & Transportasi (Jasa & Maritim)"],
        ["Sulawesi Barat", "Pertanian & Perkebunan", "38,2%", "Industri Pengolahan Sawit", "11,4%", "Didominasi Pertanian Tanaman Pangan & Perkebunan (Agraris)"],
        ["Gorontalo", "Pertanian & Tanaman Pangan", "36,4%", "Perdagangan Besar & Eceran", "14,1%", "Didominasi Pertanian, Perdagangan & Konstruksi (Agraris)"]
    ]
    add_table(doc, s17_headers, s17_rows, [2.8, 3.8, 1.6, 3.6, 1.6, 3.6], ['L', 'L', 'C', 'L', 'C', 'L'])

    # ═══════════════════════════════════════════════════════════
    # 1.2 KONSENTRASI KAWASAN INDUSTRI & PLTU CAPTIVE
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.2 Konsentrasi Kawasan Industri & PLTU Captive")
    add_p(doc, [
        ("Intensifikasi industri pengolahan mineral di Pulau Sulawesi berpusat pada pembangunan mega-smelter yang ditopang secara mutlak oleh pembangkit listrik tenaga uap khusus (*PLTU Captive*) batu bara non-jaringan (*off-grid*). Bagian ini mengombinasikan ", False, False),
        ("Analisis Spasial Deskriptif", True, False),
        (" untuk mengidentifikasi tingkat pemusatan fasilitas dan kapasitas energi fosil di enam provinsi, dengan ", False, False),
        ("Uji Tabulasi Silang Panel (Inferential Spatiotemporal Crosstabulation)", True, False),
        (" berstandar SPSS guna membuktikan secara ilmiah keterkaitan antara ekspansi PLTU Captive terhadap kehilangan tutupan hutan di Pulau Sulawesi.", False, False)
    ])
    add_note_box(
        doc,
        "Sumber Data Resmi & Deskripsi Metodologis",
        "Kementerian Energi dan Sumber Daya Mineral (ESDM / Minerbaone), Global Energy Monitor (GEM Coal Plant Tracker), dan Global Forest Watch (GFW / University of Maryland) (diolah CELIOS). Visualisasi Bar Chart Konsentrasi Industri dan Pemetaan Spasial menyajikan distribusi 778 unit fasilitas smelter serta 9,825 MW kapasitas terpasang aktif PLTU captive di 6 provinsi se-Pulau Sulawesi. Analisis dipadukan dengan Uji Tabulasi Silang Data Panel Spasiotemporal (Chi-Square Test & Risk Odds Ratio, N=60) untuk menguji keterkaitan ekspansi energi fosil industri terhadap eskalasi deforestasi komoditas."
    )

    add_formula(
        doc,
        "Kalkulasi Konsentrasi Spasial & Uji Chi-Square",
        "Kapasitas_PLTU_Kumulatif_t (MW) = Jumlah Kapasitas Aktif Baru (MW) dari Tahun 2014 hingga Tahun t\n"
        "Porsi_Smelter_Provinsi (%) = ( Jumlah_Smelter_Provinsi / Total_Smelter_Sulawesi ) * 100\n"
        "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]\n"
        "Odds_Ratio (OR) = ( a * d ) / ( b * c )",
        [
            ("Chi_Square (χ²)", "Nilai statistik uji kecocokan Pearson untuk membuktikan hubungan antara ekspansi PLTU Captive dengan lonjakan deforestasi pada panel spasiotemporal (N=60)."),
            ("Frekuensi_Harapan (E)", "Jumlah kasus teoretis jika kedua variabel saling independen: E = (Total Baris * Total Kolom) / N."),
            ("Odds_Ratio (OR)", "Ukuran kelipatan risiko peluang terjadinya deforestasi komoditas tinggi pada kelompok dengan PLTU Captive aktif (>0 MW) dibanding kelompok tanpa PLTU Captive (≤0 MW).")
        ]
    )

    add_caption(doc, "Tabel 1.5: Matriks Tabulasi Silang 2×2, Uji Chi-Square (χ²), dan Estimasi Odds Ratio Panel PLTU Captive vs Deforestasi Komoditas (2014–2023)")
    pltu_headers = ["Kategori Kapasitas PLTU (X)", "Deforestasi Rendah (<10,961 Ha)", "Deforestasi Tinggi (≥10,961 Ha)", "Total Kasus", "Parameter Statistik Uji", "Nilai / df", "Signifikansi / Kesimpulan"]
    pltu_rows = [
        ["Rendah (≤0 MW)", "27 [Exp: 18.5]", "10 [Exp: 18.5]", "37 (100%)", "Pearson Chi-Square (χ²)", "18.049 (df=1)", "p < 0.0001 (Signifikan)"],
        ["Tinggi (>0 MW)", "3 [Exp: 11.5]", "20 [Exp: 11.5]", "23 (100%)", "Likelihood Ratio", "19.420 (df=1)", "p < 0.0001 (Signifikan)"],
        ["Total Observasi Panel", "30 [Exp: 30.0]", "30 [Exp: 30.0]", "60 (100%)", "Linear-by-Linear Association", "20.036 (df=1)", "p < 0.0001 (Signifikan)"],
        ["Ukuran Risiko (Risk Estimate)", "Cross-Product: (27×20)/(10×3)", "Rasio Peluang Risiko", "OR = 18.00", "Odds Ratio (OR)", "18.00x", "Risiko Lonjakan 18x Lipat"]
    ]
    add_table(doc, pltu_headers, pltu_rows, [3.2, 2.2, 2.2, 1.8, 3.0, 2.0, 2.6], ['L', 'C', 'C', 'C', 'L', 'C', 'L'])

    add_h4(doc, "Interpretasi Spasial Industri: Eksternalitas dan Efek Meluber (Spillover)")
    add_p(doc, [
        ("Hasil pengujian empiris pada Tabel 1.5 membuktikan secara meyakinkan keterkaitan langsung antara ekspansi PLTU Captive dan kerusakan tutupan hutan di Pulau Sulawesi:\n", False, False),
        ("1. Pemusatan Ekstrem di 3 Sentra Ekstraktif Utama: ", True, False),
        ("100% kapasitas PLTU Captive dan mayoritas smelter berpusat di wilayah ini, memicu akumulasi deforestasi komoditas hingga ratusan ribu hektar, berbanding terbalik dengan 'Area Non-Smelter'.\n", False, False),
        ("2. Signifikansi Statistik yang Sangat Kuat (p < 0.0001): ", True, False),
        ("Hipotesis Nol (H0) ditolak mutlak. Bukti empiris mengonfirmasi bahwa penambahan kapasitas PLTU Captive berkorelasi langsung dengan lonjakan kehilangan tutupan hutan.\n", False, False),
        ("3. Kelipatan Risiko Bencana Ekologis (Odds Ratio = 18.00x): ", True, False),
        ("Wilayah dengan PLTU Captive memiliki risiko deforestasi komoditas 18 KALI LIPAT lebih besar. Hal ini didorong konversi masif untuk infrastruktur pendukung (coal yard, jalur transmisi, dan jalan logistik).\n", False, False),
        ("4. Efek Meluber Lintas Batas (Spillover Effect) & Emisi Karbon Terkunci: ", True, False),
        ("Eksternalitas destruktif proyek merambat luas mendegradasi DAS dan laut, mengorbankan ruang hidup lokal, serta mengunci emisi dari ketergantungan puluhan juta ton batu bara per tahun.", False, False)
    ])

    # ═══════════════════════════════════════════════════════════
    # 1.3 TREN PERTUMBUHAN IZIN TAMBANG BARU & UJI STATISTIK
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.3 Tren Pertumbuhan Izin Tambang Baru & Uji Signifikansi Statistik")
    add_p(doc, [
        ("Pola perizinan pertambangan di Pulau Sulawesi selama satu dekade terakhir menunjukkan peningkatan alokasi ruang yang signifikan. Berdasarkan data agregat ", False, False),
        ("Minerbaone", True, False),
        (", tercatat ", False, False),
        ("574 Izin Usaha Pertambangan (IUP) baru", True, False),
        (" sepanjang 2014-2024, dengan total luas konsesi mencapai ", False, False),
        ("819,452 Hektar", True, False),
        (". Terjadi lonjakan sebesar 246% pada periode 2022-2024.", False, False)
    ])

    add_caption(doc, "Tabel 1.5b: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.3)")
    var_headers = ["Komponen Uji", "Definisi Variabel (Sub-bab 1.3)"]
    var_rows = [
        ["Variabel Independen (X)", "Frekuensi Penerbitan Izin Tambang Baru (IUP) / Luas Konsesi Baru (Ha)"],
        ["Variabel Dependen (Y)", "Deforestasi Komoditas (Ha) / Total Deforestasi Alam (Ha)"],
        ["Hipotesis Nol (H0)", "Tingkat penerbitan izin/luas konsesi tidak berhubungan dengan laju deforestasi."],
        ["Hipotesis Alternatif (H1)", "Ada hubungan positif antara tingginya penerbitan izin dengan tingginya laju deforestasi."],
        ["Decision Rule (Alpha 5%)", "Jika P-Value < 0.05, maka Tolak H0 (terbukti signifikan bahwa ekspansi perizinan mendorong deforestasi)."],
        ["Threshold Kategori", "Nilai Median Data Panel (N=60): X ≥ 2.0 izin; Y ≥ 10,961.8 Ha."],
        ["Orientasi Odds Ratio", "OR = ( a × d ) / ( b × c ) dengan a = Izin Tinggi & Deforestasi Tinggi; mengukur risiko deforestasi tinggi pada kelompok penerbitan izin tinggi."]
    ]
    add_table(doc, var_headers, var_rows, [4.5, 12.5], ['L', 'L'])

    add_formula(
        doc,
        "Analisis Tren & Uji Chi-Square",
        "Pertumbuhan_Izin (%) = [ ( Jumlah_Izin_t - Jumlah_Izin_t-1 ) / Jumlah_Izin_t-1 ] × 100\n"
        "Kategori(x) = 'Tinggi' , jika x ≥ Median(Panel)   |   'Rendah' , jika x < Median(Panel)",
        [
            ("Pertumbuhan_Izin (%)", "Persentase perubahan laju penerbitan izin tambang baru antar-tahun (%)."),
            ("Kategori(x)", "Data panel spasial-temporal diubah menjadi dua tingkatan untuk uji tabulasi silang (Tinggi vs Rendah).")
        ]
    )

    add_caption(doc, "Tabel 1.6: Ringkasan Hasil Uji Independensi Chi-Square (χ²) dan Odds Ratio (OR) Data Panel Bab 1")
    crosstab_headers = ["Variabel Faktor Tekanan", "Variabel Dampak Lingkungan", "Nilai Chi-Square (χ²)", "Nilai Signifikansi (p)", "Odds Ratio (OR)", "Derajat Bebas (df)", "Kesimpulan Ilmiah"]
    crosstab_rows = [
        ["Jumlah Izin Tambang Baru (IUP)", "Total Deforestasi Alam (Ha)", "17.239", "< 0.0001", "13.75", "1", "SIGNIFIKAN"],
        ["Jumlah Izin Tambang Baru (IUP)", "Deforestasi Komoditas (Ha)", "21.818", "< 0.0001", "21.36", "1", "SIGNIFIKAN"],
        ["Luas Konsesi Tambang Baru (Ha)", "Total Deforestasi Alam (Ha)", "19.267", "< 0.0001", "16.00", "1", "SIGNIFIKAN"],
        ["Luas Konsesi Tambang Baru (Ha)", "Deforestasi Komoditas (Ha)", "19.267", "< 0.0001", "16.00", "1", "SIGNIFIKAN"]
    ]
    add_table(doc, crosstab_headers, crosstab_rows, [4.5, 3.8, 2.0, 2.0, 1.8, 1.2, 1.7], ['L', 'L', 'C', 'C', 'C', 'C', 'C'])

    add_h4(doc, "Analisis Temuan Empiris: Pembedahan Realitas Ekologis")
    add_p(doc, [
        ("Data panel membedah realitas di lapangan: lonjakan izin di wilayah pusat ekstraksi sejalan dengan tingginya nilai Chi-Square. Nilai Odds Ratio menegaskan bahwa wilayah dengan tren izin tambang yang tinggi memiliki peluang lebih besar untuk mengalami tekanan deforestasi tinggi pada tahun-tahun berjalan dan berikutnya. Secara spesifik, terjadi ", False, False),
        ("lonjakan absolut sebesar 246%", True, False),
        (" dalam penerbitan izin tambang baru pada rentang 2022 hingga 2024.", False, False)
    ])

    # ═══════════════════════════════════════════════════════════
    # 1.4 ANALISIS REALISASI INVESTASI PMDN & TUTUPAN HUTAN
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.4 Analisis Realisasi Investasi PMDN dan Dampak Terhadap Tutupan Hutan")
    add_p(doc, [
        ("Akumulasi Penanaman Modal Dalam Negeri sebesar ", False, False),
        ("Rp 218 Triliun", True, False),
        (" (Kementerian Investasi / BKPM) yang masuk dari tahun 2016-2024 berbanding lurus dengan ", False, False),
        ("1,001,654 Hektar", True, False),
        (" kehilangan tutupan hutan komoditas (Global Forest Watch). Terlihat adanya fenomena ", False, False),
        ("Efek Jeda Waktu (Time-Lagging Effect)", True, False),
        (", di mana peningkatan realisasi modal pada tahap awal perizinan dan konstruksi diikuti oleh lonjakan pembukaan lahan hutan fisik pada 1 hingga 2 tahun berikutnya.", False, False)
    ])

    add_caption(doc, "Tabel 1.7c: Matriks Pembedahan Ekologis Aktor & Emisi Karbon (Periode 2001-2025)")
    aktor_headers = ["Kategori Aktor / Metrik Ekologis", "Nilai Agregat", "Persentase dari Total Kehilangan"]
    aktor_rows = [
        ["Ekspansi Komoditas (Tambang & Sawit)", "1,890,659 Hektar", "48.4%"],
        ["Kehutanan (Logging)", "247,011 Hektar", "6.3%"],
        ["Pertanian Berpindah", "115,404 Hektar", "2.9%"],
        ["Total Kehilangan Hutan Primer", "3,904,079 Hektar", "100.0%"],
        ["Estimasi Emisi Karbon Komoditas", "1,282,195,705 Mg CO2", "-"]
    ]
    add_table(doc, aktor_headers, aktor_rows, [7.0, 4.5, 5.5], ['L', 'C', 'C'])

    add_caption(doc, "Tabel 1.8: Ringkasan Eksekutif Seluruh Skenario Crosstab Realisasi Investasi PMDN Bab 1")
    pmdn_crosstab_headers = ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"]
    pmdn_crosstab_rows = [
        ["Realisasi Investasi PMDN (Juta Rp)", "Total Deforestasi Alam (Hektar)", "2.083", "p = 0.1489", "2.8", "TIDAK SIGNIFIKAN"],
        ["Realisasi Investasi PMDN (Juta Rp)", "Deforestasi Komoditas Tambang/Sawit (Hektar)", "2.083", "p = 0.1489", "2.8", "TIDAK SIGNIFIKAN"]
    ]
    add_table(doc, pmdn_crosstab_headers, pmdn_crosstab_rows, [5.0, 5.0, 2.0, 2.0, 1.5, 1.5], ['L', 'L', 'C', 'C', 'C', 'C'])

    add_h4(doc, "Analisis Temuan Empiris: Efek Jeda Waktu (Time-Lagging)")
    add_p(doc, [
        ("Hasil pengujian seluruh skenario tabulasi silang PMDN mengungkap fenomena yang kompleks dalam alur investasi ekstraktif:\n", False, False),
        ("1. Ketidaksignifikanan Simultan & Variasi P-Value: ", True, False),
        ("Tingkat signifikansi yang bervariasi menyingkap tabir jeda waktu (lagging effect) dalam eksekusi investasi di lapangan.\n", False, False),
        ("2. Jeda Waktu Eksekusi Investasi (Lagging Effect): ", True, False),
        ("Suntikan modal masif di tahun tertentu tidak secara instan berwujud pembabatan lahan di tahun yang sama. Modal tersebut tertahan pada fase birokrasi, pembebasan lahan, dan pengadaan infrastruktur, sebelum daya rusaknya mengonversi lanskap hutan pada tahun-tahun berikutnya.\n", False, False),
        ("3. Konsentrasi Modal Ekstrem di 3 Provinsi: ", True, False),
        ("Data spasial membuktikan 89% dari total modal PMDN ekstraktif se-Sulawesi hanya tersedot ke tiga provinsi sentra (Sulteng, Sultra, Sulsel), mengakibatkan polarisasi pertumbuhan dan mengunci ketimpangan spasial.", False, False)
    ])

    # ═══════════════════════════════════════════════════════════
    # 1.5 PELABUHAN EKSPOR & LOGISTIK NIKEL
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.5 Pelabuhan Ekspor & Peta Jalur Distribusi Logistik Nikel Sulawesi")
    add_p(doc, [
        ("Eksploitasi nikel di Sulawesi tidak berhenti di tapak darat, melainkan terhubung langsung ke pasar global melalui infrastruktur pelabuhan. Bagian ini memetakan simpul logistik maritim yang mendistribusikan produk ekstraktif (NPI, Matte, MHP) dari pesisir Sulawesi menuju negara tujuan utama seperti Tiongkok dan Jepang. Verifikasi dilakukan melalui protokol triangulasi informasi publik: Laporan KNKT, Regulasi PSN (Perpres No. 109 Tahun 2020), Laporan Keberlanjutan Korporasi, dan Riset Independen.", False, False)
    ])

    add_caption(doc, "Tabel 1.7: Inventarisasi Enam Simpul Pelabuhan dan Terminal Khusus Ekspor Nikel di Pulau Sulawesi")
    port_headers = ["Simpul Kawasan Industri", "Wilayah Administrasi", "Fasilitas Pelabuhan / Terminal", "Status Regulasi", "Kapasitas Kapal", "Tujuan Utama Ekspor"]
    port_rows = [
        ["IMIP Morowali", "Morowali, Sulawesi Tengah", "Pelabuhan Samudera & Dermaga Curah", "PSN (Perpres 109/2020)", "Hingga 52.378 DWT", "Pasar Global (Tiongkok)"],
        ["GNI Morowali Utara", "Morowali Utara, Sulteng", "Terminal Khusus Pesisir Tomori", "Izin Industri Mandiri", "Hingga 30.000 DWT", "Pasar Global (Tiongkok)"],
        ["VDNI Konawe", "Konawe, Sulawesi Tenggara", "Dermaga Khusus Curah & Kargo", "PSN (Perpres 109/2020)", "Hingga 50.000 DWT", "Pasar Global (Tiongkok)"],
        ["OSS Konawe", "Konawe, Sulawesi Tenggara", "Dermaga Terintegrasi Konawe", "PSN (Perpres 109/2020)", "Hingga 50.000 DWT", "Pasar Global (Tiongkok)"],
        ["Pomalaa (ANTAM)", "Kolaka, Sulawesi Tenggara", "Dermaga Pomalaa & Konveyor", "Kawasan BUMN Industri", "Hingga 12.000 DWT", "Jepang & Korsel"],
        ["Sorowako (Vale)", "Luwu Timur, Sulawesi Selatan", "Pelabuhan Balantang Malili", "Kontrak Karya Tambang", "Hingga 15.000 DWT", "Jepang & Skandinavia"]
    ]
    add_table(doc, port_headers, port_rows, [2.8, 3.2, 3.5, 2.5, 2.2, 2.8], ['L', 'L', 'L', 'C', 'C', 'L'])

    # ═══════════════════════════════════════════════════════════
    # 1.6 PETA JALUR DISTRIBUSI LOGISTIK NIKEL SULAWESI
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.6 Peta Jalur Distribusi Logistik Nikel Sulawesi")
    add_p(doc, [
        ("Metode analisis pada tahapan ini difokuskan pada ", False, False),
        ("Pemetaan Kausalitas (Spasial)", True, False),
        (" untuk membedah asimetri penguasaan ruang antara wilayah hulu (origin: sumber ekstraksi di Sulawesi) dan hilir (destination: pusat industrialisasi luar negeri). Garis diplot menggunakan rute pelayaran untuk merepresentasikan jarak tempuh aktual kapal logistik di permukaan bumi.", False, False)
    ])

    add_formula(
        doc,
        "Kurva Parametrik Alur Pelayaran Maritim",
        "Kurva(t) = (1 - t)^2 * Titik_Asal + 2 * (1 - t) * t * Titik_Kontrol + t^2 * Titik_Tujuan",
        [
            ("Kurva(t)", "Vektor posisi koordinat geografis lintasan kapal pada parameter waktu t (rentang kontinu [0, 1])."),
            ("Titik_Asal", "Titik koordinat geografis pelabuhan muat khusus di pesisir Sulawesi."),
            ("Titik_Kontrol", "Titik koordinat jangkar pemandu kurva lengkung di perairan internasional."),
            ("Titik_Tujuan", "Titik koordinat geografis pelabuhan bongkar di negara tujuan ekspor.")
        ]
    )

    add_h4(doc, "Interpretasi Spasial Industri (Anatomi Rantai Pasok)")
    add_p(doc, [
        ("Peta rute logistik maritim mengilustrasikan alur distribusi produk olahan nikel dari kawasan industri di Sulawesi:\n", False, False),
        ("1. Orientasi Ekspor: ", True, False),
        ("Kawasan industri utama yang berstatus Proyek Strategis Nasional (PSN) mengalirkan produk olahan ke sentra-sentra industri manufaktur di pasar internasional.\n", False, False),
        ("2. Integrasi Rantai Pasok: ", True, False),
        ("Mayoritas rute pengapalan terhubung langsung dengan pelabuhan ekspor tujuan, yang mengindikasikan posisi kawasan pemurnian di Sulawesi sebagai pemasok bahan baku setengah jadi.\n", False, False),
        ("3. Dinamika Rute Maritim: ", True, False),
        ("Peta rute mencerminkan diversifikasi pasar ekspor (Asia Timur) dan jaringan logistik kawasan.", False, False)
    ])

    # ═══════════════════════════════════════════════════════════
    # 1.7 MATRIKS INDIKATOR DAN SUMBER DATA RESMI BAB 1
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.7 Matriks Indikator dan Sumber Data Resmi Bab 1")
    add_p(doc, [
        ("Seluruh variabel kuantitatif, kategori analisis, satuan ukur, periode tahun observasi, dan institusi penyedia data primer resmi yang digunakan dalam Bab 1 dikompilasikan pada ", False, False),
        ("Tabel 1.8", True, False),
        (" berikut:", False, False)
    ])

    add_caption(doc, "Tabel 1.8: Matriks Indikator dan Sumber Data Primer Resmi Bab 1")
    master_headers = ["No", "Nama Indikator", "Kategori Analisis", "Satuan Ukur", "Cakupan Tahun", "Institusi & Sumber Data Resmi", "Data File Asli"]
    master_rows = [
        ["1", "Izin Usaha Pertambangan (IUP) Baru", "Faktor Tekanan Ekstraktif", "Unit Izin", "2014-2024", "Data Registry ESDM MODI (Minerbaone)", "sulawesi_izin_baru_per_tahun.csv"],
        ["2", "Luas Wilayah Konsesi Tambang Baru", "Faktor Tekanan Ekstraktif", "Hektar (Ha)", "2014-2024", "Data Registry ESDM MODI (Minerbaone)", "sulawesi_kawasan_nikel_luas.csv"],
        ["3", "Kapasitas Terpasang PLTU Captive", "Infrastruktur Energi Khusus", "Megawatt (MW)", "2014-2024", "NGO (Global Energy Monitor / GEM)", "sulawesi_pltu_captive.csv"],
        ["4", "Fasilitas Smelter Nikel", "Fasilitas Industri Hilir", "Unit Fasilitas", "2014-2024", "Database Smelter CGS & ESDM MODI", "sulawesi_esdm_nikel.csv"],
        ["5", "Realisasi Investasi PMDN & Nikel", "Arus Modal Domestik", "Triliun Rupiah", "2016-2024", "API BPS & BKPM", "sulawesi_investasi_pmdn_2016_2024.csv"],
        ["6", "PDRB Provinsi (Ekstraktif vs Akar Rumput)", "Struktur Ekonomi Makro", "Triliun Rupiah", "2016-2024", "API BPS (Subject 52)", "sulawesi_pdrb_sektoral_2016_2024.csv"],
        ["7", "PDRB Kabupaten Sentra Tambang", "Struktur Ekonomi Daerah", "Triliun Rupiah", "2016-2024", "API BPS (Subject 52)", "sulawesi_pdrb_sektoral_kabupaten_2016_2024.csv"],
        ["8", "Pendapatan Asli Daerah (PAD) & Breakdown Pajak", "Kapasitas Fiskal Daerah", "Triliun Rupiah", "2016-2024", "API BPS", "sulawesi_pad_breakdown_2016_2024.csv"],
        ["9", "Luas Total Deforestasi Alam & Komoditas", "Dampak Ekologis", "Hektar (Ha)", "2014-2023", "Global Forest Watch (GFW API)", "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv"],
        ["10", "Simpul Pelabuhan & Terminal Logistik", "Infrastruktur Rantai Pasok", "Titik Koordinat & DWT", "2014-2024", "Laporan Publik (KNKT, Perpres PSN, Korporasi)", "sulawesi_logistik_simpul_nikel.csv"]
    ]
    add_table(doc, master_headers, master_rows, [0.8, 3.2, 2.2, 1.8, 1.8, 3.8, 3.4], ['C', 'L', 'L', 'C', 'C', 'L', 'L'])

    # ═══════════════════════════════════════════════════════════
    # 1.8 BAGAN ALUR KERANGKA KERJA RISET BAB 1
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.8 Bagan Alur Kerangka Kerja Riset Bab 1")
    add_p(doc, [
        ("Keseluruhan struktur metodologis riset Bab 1 dioperasionalkan melalui empat fase kerja berurutan sebagaimana disajikan pada ", False, False),
        ("Tabel 1.9", True, False),
        (" berikut:", False, False)
    ])

    add_caption(doc, "Tabel 1.9: Matriks Tahapan dan Alur Kerangka Kerja Riset Bab 1")
    flow_headers = ["Tahapan Riset", "Fokus Metodologis", "Bahan & Sumber Data", "Keluaran / Hasil Analisis"]
    flow_rows = [
        ["Fase I: Pengumpulan Data", "Kurasi data resmi lintas kementerian dan lembaga", "Publikasi BPS, Minerbaone, BKPM, GEM, dan GFW", "Basis Data Tabular Panel Provinsi (2014–2024)"],
        ["Fase II: Reklasifikasi Hukum", "Penyusunan kerangka rantai pasok hukum terintegrasi", "UU No. 3/2020, PP No. 96/2021, Perpres No. 112/2022", "3 Klaster Makro (Ekstraktif, Akar Rumput, Jasa)"],
        ["Fase III: Pengujian Statistik", "Uji signifikansi hubungan dan rasio peluang", "Tabel Kontinjensi, Uji Chi-Square, Odds Ratio", "Bukti Kausalitas Signifikan Tekanan vs Deforestasi"],
        ["Fase IV: Pemetaan Rantai Pasok", "Triangulasi data logistik dan pemodelan maritim", "Laporan KNKT, Perpres PSN, Kurva Parametrik Bézier", "Peta Alur Rantai Pasok Ekspor & Konsentrasi Spasial 78%"]
    ]
    add_table(doc, flow_headers, flow_rows, [3.0, 4.2, 4.8, 5.0], ['L', 'L', 'L', 'L'])

    # Simpan File DOCX
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab1    = out_dir_compact.parent.parent / "bab_1"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab1.mkdir(parents=True, exist_ok=True)

    docx_path_compact = out_dir_compact / "Metodologi_Bab1_Ekspansi_Industri_Compact.docx"
    docx_path_bab1    = out_dir_bab1 / "Metodologi_Bab1_Ekspansi_Industri_Compact.docx"

    doc.save(str(docx_path_compact))
    shutil.copyfile(docx_path_compact, docx_path_bab1)
    print(f"[OK] Berhasil menyimpan DOCX di: {docx_path_compact}")
    print(f"[OK] Salinan tersimpan di: {docx_path_bab1}")

# ── Generator Naskah Markdown Compact ───────────────────────
def generate_compact_markdown():
    print("[2/4] Menyusun naskah Markdown Metodologi Versi Compact...")
    md_content = """# BAB I: METODOLOGI ANALISIS EKSPANSI INDUSTRI EKSTRAKTIF DAN INFRASTRUKTUR PENUNJANG DI PULAU SULAWESI

Dokumen laporan metodologi ini menyajikan kerangka ilmiah, landasan regulasi, formulasi matematis, prosedur analisis statistik, serta metodologi pembuktian berbasis data terbuka yang dioperasionalkan pada **Bab 1: Ekspansi Industri Ekstraktif** dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi periode 2014–2024.

---

## 1.1 Konteks Makro: Breakdown PDRB per Komoditas

### 1.1.1 Konteks Makro: Dominasi Ekstraktif vs Ekonomi Akar Rumput
Bagian ini menganalisis struktur Produk Domestik Regional Bruto (PDRB) pada enam provinsi di Pulau Sulawesi sepanjang periode 2016–2024 menggunakan visualisasi grafik area bertumpuk (*Stacked Area Chart*). Analisis ini ditujukan untuk menguji secara empiris apakah percepatan pertumbuhan ekonomi daerah benar-benar bersumber dari sektor produktif masyarakat lokal atau didominasi oleh industri ekstraktif padat modal yang mengalihkan pemanfaatan ruang dan sumber daya alam.

> **Sumber Data:** Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi *Stacked Area Chart* memetakan dinamika Produk Domestik Regional Bruto (PDRB) berdasarkan klasifikasi rantai pasok hukum (*Legal Supply-Chain*) untuk membandingkan trajektori Sektor Ekstraktif, Ekonomi Akar Rumput, dan Sektor Jasa & Lainnya.

#### A. Kerangka Dekomposisi Sektoral & Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)
Sistem KBLI 2020 BPS membagi 17 sektor PDRB. Melalui pendekatan Legal Supply-Chain, 17 sektor direklasifikasi menjadi 3 Klaster Makro (Ekstraktif, Akar Rumput, Jasa). Rincian pembagian sektor, dasar regulasi, serta intisari ketentuan hukum disajikan secara lengkap pada **Tabel 1.1** berikut:

##### Tabel 1.1: Reklasifikasi Sektoral PDRB KBLI 2020 Berdasarkan Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)
| Kategori BPS | Sektor Lapangan Usaha | Klasifikasi Analisis | Dasar Regulasi & Mandat Hukum | Intisari Ketentuan Hukum |
| :--- | :--- | :---: | :--- | :--- |
| **Kategori B** | Pertambangan dan Penggalian | Ekstraktif | Perpres No. 26 Tahun 2010 | Ketentuan Pasal 1 Ayat (2) mengenai pengambilan komoditas tambang dari dalam bumi. |
| **Kategori C** | Industri Pengolahan (Smelter Logam) | Ekstraktif | UU No. 3 Tahun 2020 & PP No. 96 Tahun 2021 | Pasal 102–103 mewajibkan pengolahan dan pemurnian di dalam negeri sebagai kesatuan pertambangan. |
| **Kategori D** | Pengadaan Listrik & Gas (PLTU Captive) | Ekstraktif | Perpres No. 112 Tahun 2022 & RUPTL PLN | Pasal 3 Ayat (4) huruf b mengecualikan PLTU baru hanya bagi yang terintegrasi melayani smelter. |
| **Kategori A** | Pertanian, Kehutanan, Perikanan | Ekonomi Akar Rumput | KBLI 2020 BPS | Sektor pemanfaatan sumber daya hayati terbarukan dan penyerap tenaga kerja lokal. |
| **Kategori E–U** | 13 Sektor Jasa & Konstruksi | Sektor Jasa & Lainnya | Klasifikasi Standar BPS | Sektor sekunder dan tersier penunjang perekonomian daerah. |

#### B. Formulasi Matematis: Persamaan Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)
```text
Sektor_Ekstraktif = PDRB(Kat.B: Pertambangan) + PDRB(Kat.C: Ind. Pengolahan) + PDRB(Kat.D: Listrik)
Sektor_Akar_Rumput = PDRB(Kat.A: Pertanian, Kehutanan, dan Perikanan)
Sektor_Jasa = Jumlah PDRB (Kategori E sampai dengan Kategori U)
Total_PDRB = Sektor_Ekstraktif + Sektor_Akar_Rumput + Sektor_Jasa
Pangsa_Ekstraktif (%) = ( Sektor_Ekstraktif / Total_PDRB ) * 100
Laju_Pertumbuhan_Tahunan (%) = [ ( Nilai_Tahun_t - Nilai_Tahun_{t-1} ) / Nilai_Tahun_{t-1} ] * 100
```
- `Sektor_Ekstraktif`: Total nilai tambah bruto dari klaster industri ekstraktif yang saling terintegrasi (Triliun Rupiah).
- `PDRB(Kat.B: Pertambangan)`: Nilai tambah kegiatan eksplorasi dan ekstraksi bijih mineral (BPS Kategori B).
- `PDRB(Kat.C: Ind. Pengolahan)`: Nilai tambah pemurnian logam dasar di smelter nikel (BPS Kategori C / Golongan 24).
- `PDRB(Kat.D: Listrik)`: Nilai tambah penyediaan listrik batubara khusus smelter / PLTU captive (BPS Kategori D).
- `Total_PDRB`: Total nilai Produk Domestik Regional Bruto wilayah atas dasar harga berlaku (Triliun Rupiah).
- `Pangsa_Ekstraktif (%)`: Persentase pangsa dominasi sektor ekstraktif terhadap total ekonomi (%).

##### Tabel 1.2: Definisi Operasional Komponen Makroekonomi dan Sumber Data PDRB Sektoral
| Komponen Analisis | Cakupan Lapangan Usaha | Definisi Operasional | Satuan Nilai | Sumber Data Primer |
| :--- | :--- | :--- | :---: | :--- |
| **Sektor Ekstraktif** | Kategori B, Kategori C, Kategori D | Akumulasi nilai tambah pertambangan nikel, smelter logam dasar, dan PLTU captive. | Triliun Rupiah | BPS Provinsi (SIMDASI) |
| **Ekonomi Akar Rumput** | Kategori A | Nilai tambah pertanian, perkebunan, kehutanan, dan perikanan. | Triliun Rupiah | BPS Provinsi |
| **Sektor Jasa & Lainnya** | Kategori E hingga U | Nilai tambah gabungan perdagangan, konstruksi, transportasi, keuangan, dan jasa. | Triliun Rupiah | BPS Provinsi |
| **Total PDRB Wilayah** | Seluruh 17 Kategori | Total nilai PDRB wilayah atas dasar harga berlaku pada tahun berjalan. | Triliun Rupiah | BPS Provinsi |
| **Pangsa Ekstraktif (%)** | Rasio Kontribusi | Persentase kontribusi sektor ekstraktif terhadap total perekonomian. | Persen (%) | Hasil Olahan CELIOS |

#### C. Analisis Temuan Empiris: Ketimpangan Struktural Sulawesi Tengah
Penerapan formulasi di atas menunjukkan bahwa di **Sulawesi Tengah (sebagai pusat hilirisasi)**, ekspansi industri ekstraktif menguasai **55.8% dari total PDRB provinsi** pada tahun 2024, memperlihatkan dominasi yang sangat kuat dibanding provinsi lainnya. Sektor ekstraktif melonjak dari Rp28,45 triliun pada 2016 menjadi Rp210,51 triliun pada 2024.

---

### 1.1.2 Pemusatan Sektor Ekstraktif di Kabupaten se-Sulawesi Tengah
Jika dianalisis secara spasial pada tingkat kabupaten di Sulawesi Tengah, terlihat konsentrasi kegiatan industri ekstraktif. Kabupaten **Morowali** dan **Morowali Utara** mendominasi struktur PDRB provinsi melalui pengembangan kawasan industri hilirisasi dan PLTU Captive. Analisis ini membandingkan komposisi ketiga sektor advokatif di seluruh 13 kabupaten/kota se-Sulawesi Tengah pada tahun terbaru (2024).

> **Sumber Data:** Badan Pusat Statistik (BPS) Kabupaten se-Sulawesi Tengah (diolah CELIOS). Visualisasi *Stacked Bar Chart* memetakan struktur Produk Domestik Regional Bruto (PDRB) tahun 2024 pada seluruh 13 kabupaten/kota untuk mengidentifikasi tingkat konsentrasi sektoral dan polarisasi spasial antara sentra industri pengolahan nikel dengan daerah non-sentra.

#### A. Formulasi Matematis: Persamaan Agregasi Sektoral Kabupaten
```text
Sektor_Ekstraktif_Kabupaten = PDRB_Kab(Kat.B: Pertambangan) + PDRB_Kab(Kat.C: Ind. Pengolahan) + PDRB_Kab(Kat.D: Listrik)
Total_PDRB_Kabupaten = Sektor_Ekstraktif_Kabupaten + Sektor_Akar_Rumput_Kabupaten + Sektor_Jasa_Kabupaten
Porsi_Sektor_Kabupaten (%) = ( Nilai_Sektor_Kabupaten / Total_PDRB_Kabupaten ) * 100
```

##### Tabel 1.3: Distribusi Nilai Tambah Bruto dan Komposisi Sektoral PDRB 13 Kabupaten/Kota di Sulawesi Tengah (Tahun 2024)
| Kabupaten / Kota | Akar Rumput (T Rp) | Ekstraktif (T Rp) | Jasa (T Rp) | Total PDRB (T Rp) | Porsi Akar Rumput (%) | Porsi Ekstraktif (%) | Porsi Jasa (%) | Basis Utama Ekonomi |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :--- |
| **Morowali** | 2.70 | 157.17 | 187.85 | **347.72** | 0.8% | 45.2% | 54.0% | Hilirisasi Nikel (Smelter & PLTU) |
| **Banggai** | 8.85 | 20.63 | 51.99 | **81.47** | 10.9% | 25.3% | 63.8% | Migas, Tambang & Perdagangan |
| **Palu** | 1.24 | 4.56 | 60.03 | **65.84** | 1.9% | 6.9% | 91.2% | Jasa, Perdagangan & Pemerintahan |
| **Morowali Utara** | 5.17 | 19.22 | 36.08 | **60.47** | 8.5% | 31.8% | 59.7% | Hilirisasi Nikel (Smelter GNI) |
| **Parigi Moutong** | 9.97 | 1.93 | 35.05 | **46.95** | 21.2% | 4.1% | 74.7% | Pertanian Pangan & Hortikultura |
| **Donggala** | 5.96 | 3.53 | 23.57 | **33.05** | 18.0% | 10.7% | 71.3% | Pertanian, Perkebunan & Galian C |
| **Poso** | 4.96 | 0.40 | 20.12 | **25.48** | 19.5% | 1.6% | 79.0% | Pertanian & Perkebunan Kakao |
| **Sigi** | 5.17 | 0.83 | 19.13 | **25.12** | 20.6% | 3.3% | 76.1% | Pertanian Pangan & Hortikultura |
| **Toli-Toli** | 4.35 | 0.44 | 17.36 | **22.15** | 19.7% | 2.0% | 78.4% | Perkebunan Cengkeh & Perikanan |
| **Buol** | 3.77 | 1.15 | 10.67 | **15.58** | 24.2% | 7.4% | 68.5% | Kelapa Sawit & Tanaman Pangan |
| **Tojo Una-Una** | 2.88 | 0.61 | 11.21 | **14.71** | 19.6% | 4.2% | 76.2% | Pertanian & Pariwisata Bahari |
| **Banggai Kepulauan** | 2.53 | 0.18 | 8.04 | **10.75** | 23.5% | 1.7% | 74.8% | Perikanan Tangkap & Kelautan |
| **Banggai Laut** | 1.80 | 0.12 | 4.52 | **6.45** | 27.9% | 1.9% | 70.2% | Perikanan & Budidaya Laut |

#### B. Analisis Temuan Empiris: Polarisasi Ekstrem Morowali vs Daerah Non-Smelter
Data empiris pada Tabel 1.3 mengungkap bukti polarisasi ekonomi wilayah yang sangat ekstrem di Sulawesi Tengah:
1. **Dominasi Sektor Ekstraktif Morowali:** Kabupaten Morowali mencatatkan nilai sektor ekstraktif sebesar Rp 157.17 Triliun atau menguasai porsi 45.2% dari total kue ekonomi kabupatennya (Rp 347.72 Triliun). Nilai sektor ekstraktif Morowali saja melampaui gabungan total PDRB dari delapan kabupaten lainnya di Sulawesi Tengah.
2. **Pemusatan pada Dua Sentra Hilirisasi:** Kabupaten Morowali dan Morowali Utara merupakan dua daerah dengan nilai Sektor Ekstraktif tertinggi di Sulawesi Tengah, membuktikan bahwa percepatan output industri pertambangan dan hilirisasi terkunci pada kawasan industri smelter.
3. **Ketertinggalan Daerah Non-Sentra:** Sebaliknya, delapan kabupaten lainnya (seperti Banggai Laut, Banggai Kepulauan, Tojo Una-Una, Buol, Toli-Toli, Sigi, Poso, dan Donggala) memiliki porsi Sektor Ekstraktif yang sangat rendah (<11%) dan tetap bergantung pada sektor pertanian rakyat (Akar Rumput) berproduktivitas rendah dengan keterbatasan akses terhadap nilai tambah modal.

---

### 1.1.3 Perbandingan Distribusi 17 Sektor Komoditas per Provinsi (Small Multiples, Tahun Terbaru)
Visualisasi komparatif **Small Multiples Horizontal Bar Chart** membedah struktur 17 sektor lapangan usaha KBLI 2020 secara terpisah pada enam provinsi di Pulau Sulawesi pada tahun terbaru (2024). Setiap panel provinsi menampilkan sektor yang diurutkan dari penyumbang terbesar hingga terkecil dengan skala sumbu nilai yang disetarakan secara seragam untuk memastikan validitas komparasi lintas wilayah.

##### Tabel 1.4: 5 Sektor Lapangan Usaha Penyumbang Utama PDRB di 6 Provinsi Sulawesi (Tahun 2024)
| Provinsi | Sektor Utama 1 | Porsi (%) | Sektor Utama 2 | Porsi (%) | Tipologi Wilayah |
| :--- | :--- | :---: | :--- | :---: | :--- |
| **Sulawesi Tengah** | Industri Pengolahan (Smelter) | 44,1% | Pertambangan & Penggalian | 11,8% | Didominasi Industri Pengolahan Smelter & Pertambangan (Ekstraktif) |
| **Sulawesi Tenggara** | Pertanian & Perikanan | 22,4% | Pertambangan Logam | 20,9% | Didominasi Pertanian & Pertambangan Logam (Campuran) |
| **Sulawesi Selatan** | Pertanian & Perikanan | 21,8% | Perdagangan Besar & Eceran | 14,6% | Didominasi Pertanian, Perdagangan & Konstruksi (Agraris & Jasa) |
| **Sulawesi Utara** | Pertanian & Perikanan | 20,5% | Perdagangan Besar & Eceran | 13,2% | Didominasi Pertanian, Perdagangan & Transportasi (Jasa & Maritim) |
| **Sulawesi Barat** | Pertanian & Perkebunan | 38,2% | Industri Pengolahan Sawit | 11,4% | Didominasi Pertanian Tanaman Pangan & Perkebunan (Agraris) |
| **Gorontalo** | Pertanian & Tanaman Pangan | 36,4% | Perdagangan Besar & Eceran | 14,1% | Didominasi Pertanian, Perdagangan & Konstruksi (Agraris) |

---

## 1.2 Konsentrasi Kawasan Industri & PLTU Captive

Intensifikasi industri pengolahan mineral di Pulau Sulawesi berpusat pada pembangunan mega-smelter yang ditopang secara mutlak oleh pembangkit listrik tenaga uap khusus (*PLTU Captive*) batu bara non-jaringan (*off-grid*). Bagian ini mengombinasikan **Analisis Spasial Deskriptif** untuk mengidentifikasi tingkat pemusatan fasilitas dan kapasitas energi fosil di enam provinsi, dengan **Uji Tabulasi Silang Panel (Inferential Spatiotemporal Crosstabulation)** berstandar SPSS guna membuktikan secara ilmiah keterkaitan antara ekspansi PLTU Captive terhadap kehilangan tutupan hutan di Pulau Sulawesi.

> **Sumber Data Resmi & Deskripsi Metodologis:** Kementerian Energi dan Sumber Daya Mineral (ESDM / Minerbaone), Global Energy Monitor (GEM Coal Plant Tracker), dan Global Forest Watch (GFW / University of Maryland) (diolah CELIOS). Visualisasi *Bar Chart* Konsentrasi Industri dan Pemetaan Spasial menyajikan distribusi 778 unit fasilitas smelter serta 9,825 MW kapasitas terpasang aktif PLTU captive di 6 provinsi se-Pulau Sulawesi. Analisis dipadukan dengan Uji Tabulasi Silang Data Panel Spasiotemporal (Chi-Square Test & Risk Odds Ratio, N=60) untuk menguji keterkaitan ekspansi energi fosil industri terhadap eskalasi deforestasi komoditas.

#### A. Formulasi Matematis: Kalkulasi Konsentrasi Spasial & Uji Chi-Square
```text
Kapasitas_PLTU_Kumulatif_t (MW) = Jumlah Kapasitas Aktif Baru (MW) dari Tahun 2014 hingga Tahun t
Porsi_Smelter_Provinsi (%) = ( Jumlah_Smelter_Provinsi / Total_Smelter_Sulawesi ) * 100
Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]
Odds_Ratio (OR) = ( a * d ) / ( b * c )
```
- `Chi_Square (χ²)`: Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara ekspansi PLTU Captive dengan lonjakan deforestasi pada panel spasiotemporal (N=60).
- `Frekuensi_Harapan (E)`: Jumlah kasus teoretis jika kedua variabel saling independen: E = (Total Baris * Total Kolom) / N.
- `Odds_Ratio (OR)`: Ukuran kelipatan risiko peluang terjadinya deforestasi komoditas tinggi pada kelompok dengan PLTU Captive aktif (>0 MW) dibanding kelompok tanpa PLTU Captive (≤0 MW).

##### Tabel 1.5: Matriks Tabulasi Silang 2×2, Uji Chi-Square (χ²), dan Estimasi Odds Ratio Panel PLTU Captive vs Deforestasi Komoditas (2014–2023)
| Kategori Kapasitas PLTU (X) | Deforestasi Rendah (<10,961 Ha) | Deforestasi Tinggi (≥10,961 Ha) | Total Kasus | Parameter Statistik Uji | Nilai / df | Signifikansi / Kesimpulan |
| :--- | :---: | :---: | :---: | :--- | :--- | :--- |
| **Rendah (≤0 MW)** | 27 [Exp: 18.5] | 10 [Exp: 18.5] | 37 (100%) | **Pearson Chi-Square (χ²)** | **18.049** (df=1) | p < 0.0001 (Signifikan) |
| **Tinggi (>0 MW)** | 3 [Exp: 11.5] | 20 [Exp: 11.5] | 23 (100%) | **Likelihood Ratio** | **19.420** (df=1) | p < 0.0001 (Signifikan) |
| **Total Observasi Panel** | **30** [Exp: 30.0] | **30** [Exp: 30.0] | **60** (100%) | **Linear-by-Linear Association** | **20.036** (df=1) | p < 0.0001 (Signifikan) |
| **Ukuran Risiko (Risk Estimate)** | Cross-Product: (27×20)/(10×3) | Rasio Peluang Risiko | OR = 18.00 | **Odds Ratio (OR)** | **18.00x** | **Risiko Lonjakan 18x Lipat** |

#### B. Interpretasi Spasial Industri: Eksternalitas dan Efek Meluber (Spillover)
Hasil pengujian empiris pada Tabel 1.5 membuktikan secara meyakinkan keterkaitan langsung antara ekspansi PLTU Captive dan kerusakan tutupan hutan di Pulau Sulawesi:
1. **Pemusatan Ekstrem di 3 Sentra Ekstraktif Utama:** 100% kapasitas PLTU Captive dan mayoritas smelter berpusat di wilayah ini, memicu akumulasi deforestasi komoditas hingga ratusan ribu hektar, berbanding terbalik dengan "Area Non-Smelter".
2. **Signifikansi Statistik yang Sangat Kuat (p < 0.0001):** Hipotesis Nol (H0) ditolak mutlak. Bukti empiris mengonfirmasi bahwa penambahan kapasitas PLTU Captive berkorelasi langsung dengan lonjakan kehilangan tutupan hutan.
3. **Kelipatan Risiko Bencana Ekologis (Odds Ratio = 18.00x):** Wilayah dengan PLTU Captive memiliki risiko deforestasi komoditas 18 KALI LIPAT lebih besar. Hal ini didorong konversi masif untuk infrastruktur pendukung (coal yard, jalur transmisi, dan jalan logistik).
4. **Efek Meluber Lintas Batas (Spillover Effect) & Emisi Karbon Terkunci:** Eksternalitas destruktif proyek merambat luas mendegradasi DAS dan laut, mengorbankan ruang hidup lokal, serta mengunci emisi dari ketergantungan puluhan juta ton batu bara per tahun.

---

## 1.3 Tren Pertumbuhan Izin Tambang Baru & Uji Signifikansi Statistik

Pola perizinan pertambangan di Pulau Sulawesi selama satu dekade terakhir menunjukkan peningkatan alokasi ruang yang signifikan. Berdasarkan data agregat **Minerbaone**, tercatat 574 Izin Usaha Pertambangan (IUP) baru sepanjang 2014-2024, dengan total luas konsesi mencapai 819,452 Hektar. Terjadi lonjakan sebesar **246% pada periode 2022-2024**.

##### Tabel 1.5b: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.3)
| Komponen Uji | Definisi Variabel (Sub-bab 1.3) |
| :--- | :--- |
| Variabel Independen (X) | Frekuensi Penerbitan Izin Tambang Baru (IUP) / Luas Konsesi Baru (Ha) |
| Variabel Dependen (Y) | Deforestasi Komoditas (Ha) / Total Deforestasi Alam (Ha) |
| Hipotesis Nol (H0) | Tingkat penerbitan izin/luas konsesi tidak berhubungan dengan laju deforestasi. |
| Hipotesis Alternatif (H1) | Ada hubungan positif antara tingginya penerbitan izin dengan tingginya laju deforestasi. |
| Decision Rule (Alpha 5%) | Jika P-Value < 0.05, maka Tolak H0 (terbukti signifikan bahwa ekspansi perizinan mendorong deforestasi). |
| Threshold Kategori | Nilai Median Data Panel (N=60): X ≥ 2.0 izin; Y ≥ 10,961.8 Ha. |
| Orientasi Odds Ratio | OR = ( a × d ) / ( b × c ) dengan a = Izin Tinggi & Deforestasi Tinggi; mengukur risiko deforestasi tinggi pada kelompok penerbitan izin tinggi. |

#### A. Formulasi Matematis: Analisis Tren & Uji Chi-Square
```text
Pertumbuhan_Izin (%) = [ ( Jumlah_Izin_t - Jumlah_Izin_t-1 ) / Jumlah_Izin_t-1 ] × 100
Kategori(x) = 'Tinggi' , jika x ≥ Median(Panel)   |   'Rendah' , jika x < Median(Panel)
```

##### Tabel 1.6: Ringkasan Hasil Uji Independensi Chi-Square (χ²) dan Odds Ratio (OR) Data Panel Bab 1
| Variabel Faktor Tekanan | Variabel Dampak Lingkungan | Nilai Chi-Square (χ²) | Nilai Signifikansi (p) | Odds Ratio (OR) | Derajat Bebas (df) | Kesimpulan Ilmiah |
| :--- | :--- | :---: | :---: | :---: | :---: | :--- |
| Jumlah Izin Tambang Baru (IUP) | Total Deforestasi Alam (Ha) | 17.239 | < 0.0001 | 13.75 | 1 | SIGNIFIKAN |
| Jumlah Izin Tambang Baru (IUP) | Deforestasi Komoditas (Ha) | 21.818 | < 0.0001 | 21.36 | 1 | SIGNIFIKAN |
| Luas Konsesi Tambang Baru (Ha) | Total Deforestasi Alam (Ha) | 19.267 | < 0.0001 | 16.00 | 1 | SIGNIFIKAN |
| Luas Konsesi Tambang Baru (Ha) | Deforestasi Komoditas (Ha) | 19.267 | < 0.0001 | 16.00 | 1 | SIGNIFIKAN |

#### B. Analisis Temuan Empiris: Pembedahan Realitas Ekologis
Data panel membedah realitas di lapangan: lonjakan izin di wilayah pusat ekstraksi sejalan dengan tingginya nilai Chi-Square. Nilai Odds Ratio menegaskan bahwa wilayah dengan tren izin tambang yang tinggi memiliki peluang lebih besar untuk mengalami tekanan deforestasi tinggi pada tahun-tahun berjalan dan berikutnya. Secara spesifik, terjadi **lonjakan absolut sebesar 246%** dalam penerbitan izin tambang baru pada rentang 2022 hingga 2024.

---

## 1.4 Analisis Realisasi Investasi PMDN dan Dampak Terhadap Tutupan Hutan

Akumulasi Penanaman Modal Dalam Negeri sebesar **Rp 218 Triliun** (Kementerian Investasi / BKPM) yang masuk dari tahun 2016-2024 berbanding lurus dengan **1,001,654 Hektar** kehilangan tutupan hutan komoditas (Global Forest Watch). Terlihat adanya fenomena **Efek Jeda Waktu (Time-Lagging Effect)**, di mana peningkatan realisasi modal pada tahap awal perizinan dan konstruksi diikuti oleh lonjakan pembukaan lahan hutan fisik pada 1 hingga 2 tahun berikutnya.

##### Tabel 1.7c: Matriks Pembedahan Ekologis Aktor & Emisi Karbon (Periode 2001-2025)
| Kategori Aktor / Metrik Ekologis | Nilai Agregat | Persentase dari Total Kehilangan |
| :--- | :---: | :---: |
| **Ekspansi Komoditas (Tambang & Sawit)** | 1,890,659 Hektar | 48.4% |
| **Kehutanan (Logging)** | 247,011 Hektar | 6.3% |
| **Pertanian Berpindah** | 115,404 Hektar | 2.9% |
| **Total Kehilangan Hutan Primer** | **3,904,079 Hektar** | **100.0%** |
| **Estimasi Emisi Karbon Komoditas** | 1,282,195,705 Mg CO2 | - |

##### Tabel 1.8: Ringkasan Eksekutif Seluruh Skenario Crosstab Realisasi Investasi PMDN Bab 1
| Variabel Independen (X) | Variabel Dependen (Y) | Chi-Square (χ²) | P-Value | Odds Ratio | Kesimpulan |
| :--- | :--- | :---: | :---: | :---: | :--- |
| Realisasi Investasi PMDN (Juta Rp) | Total Deforestasi Alam (Hektar) | 2.083 | p = 0.1489 | 2.8 | TIDAK SIGNIFIKAN |
| Realisasi Investasi PMDN (Juta Rp) | Deforestasi Komoditas Tambang/Sawit (Hektar) | 2.083 | p = 0.1489 | 2.8 | TIDAK SIGNIFIKAN |

#### A. Analisis Temuan Empiris: Efek Jeda Waktu (Time-Lagging)
Hasil pengujian seluruh skenario tabulasi silang PMDN mengungkap fenomena yang kompleks dalam alur investasi ekstraktif:
1. **Ketidaksignifikanan Simultan & Variasi P-Value:** Tingkat signifikansi yang bervariasi menyingkap tabir jeda waktu (lagging effect) dalam eksekusi investasi di lapangan.
2. **Jeda Waktu Eksekusi Investasi (Lagging Effect):** Suntikan modal masif di tahun tertentu tidak secara instan berwujud pembabatan lahan di tahun yang sama. Modal tersebut tertahan pada fase birokrasi, pembebasan lahan, dan pengadaan infrastruktur, sebelum daya rusaknya mengonversi lanskap hutan pada tahun-tahun berikutnya.
3. **Konsentrasi Modal Ekstrem di 3 Provinsi:** Data spasial membuktikan 89% dari total modal PMDN ekstraktif se-Sulawesi hanya tersedot ke tiga provinsi sentra (Sulteng, Sultra, Sulsel), mengakibatkan polarisasi pertumbuhan dan mengunci ketimpangan spasial.

---

## 1.5 Pelabuhan Ekspor & Peta Jalur Distribusi Logistik Nikel Sulawesi

Eksploitasi nikel di Sulawesi tidak berhenti di tapak darat, melainkan terhubung langsung ke pasar global melalui infrastruktur pelabuhan. Bagian ini memetakan simpul logistik maritim yang mendistribusikan produk ekstraktif (NPI, Matte, MHP) dari pesisir Sulawesi menuju negara tujuan utama seperti Tiongkok dan Jepang. Verifikasi dilakukan melalui protokol triangulasi informasi publik: Laporan KNKT, Regulasi PSN (Perpres No. 109 Tahun 2020), Laporan Keberlanjutan Korporasi, dan Riset Independen.

##### Tabel 1.7: Inventarisasi Enam Simpul Pelabuhan dan Terminal Khusus Ekspor Nikel di Pulau Sulawesi
| Simpul Kawasan Industri | Wilayah Administrasi | Fasilitas Pelabuhan / Terminal | Status Regulasi | Kapasitas Kapal | Tujuan Utama Ekspor |
| :--- | :--- | :--- | :--- | :---: | :--- |
| **IMIP Morowali** | Morowali, Sulawesi Tengah | Pelabuhan Samudera & Dermaga Curah | PSN (Perpres 109/2020) | Hingga 52.378 DWT | Pasar Global (Tiongkok) |
| **GNI Morowali Utara** | Morowali Utara, Sulteng | Terminal Khusus Pesisir Tomori | Izin Industri Mandiri | Hingga 30.000 DWT | Pasar Global (Tiongkok) |
| **VDNI Konawe** | Konawe, Sulawesi Tenggara | Dermaga Khusus Curah & Kargo | PSN (Perpres 109/2020) | Hingga 50.000 DWT | Pasar Global (Tiongkok) |
| **OSS Konawe** | Konawe, Sulawesi Tenggara | Dermaga Terintegrasi Konawe | PSN (Perpres 109/2020) | Hingga 50.000 DWT | Pasar Global (Tiongkok) |
| **Pomalaa (ANTAM)** | Kolaka, Sulawesi Tenggara | Dermaga Pomalaa & Konveyor | Kawasan BUMN Industri | Hingga 12.000 DWT | Jepang & Korsel |
| **Sorowako (Vale)** | Luwu Timur, Sulawesi Selatan | Pelabuhan Balantang Malili | Kontrak Karya Tambang | Hingga 15.000 DWT | Jepang & Skandinavia |

---

## 1.6 Peta Jalur Distribusi Logistik Nikel Sulawesi

Metode analisis pada tahapan ini difokuskan pada **Pemetaan Kausalitas (Spasial)** untuk membedah asimetri penguasaan ruang antara wilayah hulu (origin: sumber ekstraksi di Sulawesi) dan hilir (destination: pusat industrialisasi luar negeri). Garis diplot menggunakan rute pelayaran untuk merepresentasikan jarak tempuh aktual kapal logistik di permukaan bumi.

#### A. Formulasi Matematis: Kurva Parametrik Alur Pelayaran
```text
Kurva(t) = (1 - t)^2 * Titik_Asal + 2 * (1 - t) * t * Titik_Kontrol + t^2 * Titik_Tujuan
```
- `Kurva(t)`: Vektor posisi koordinat geografis lintasan kapal pada parameter waktu t (rentang kontinu [0, 1]).
- `Titik_Asal`: Titik koordinat geografis pelabuhan muat khusus di pesisir Sulawesi.
- `Titik_Kontrol`: Titik koordinat jangkar pemandu kurva lengkung di perairan internasional.
- `Titik_Tujuan`: Titik koordinat geografis pelabuhan bongkar di negara tujuan ekspor.

#### B. Interpretasi Spasial Industri (Anatomi Rantai Pasok)
Peta rute logistik maritim mengilustrasikan alur distribusi produk olahan nikel dari kawasan industri di Sulawesi:
1. **Orientasi Ekspor:** Kawasan industri utama yang berstatus Proyek Strategis Nasional (PSN) mengalirkan produk olahan ke sentra-sentra industri manufaktur di pasar internasional.
2. **Integrasi Rantai Pasok:** Mayoritas rute pengapalan terhubung langsung dengan pelabuhan ekspor tujuan, yang mengindikasikan posisi kawasan pemurnian di Sulawesi sebagai pemasok bahan baku setengah jadi.
3. **Dinamika Rute Maritim:** Peta rute mencerminkan diversifikasi pasar ekspor (Asia Timur) dan jaringan logistik kawasan.

---

## 1.7 Matriks Indikator dan Sumber Data Resmi Bab 1

Seluruh variabel kuantitatif, kategori analisis, satuan ukur, periode tahun observasi, dan institusi penyedia data primer resmi yang digunakan dalam Bab 1 dikompilasikan pada **Tabel 1.8** berikut:

##### Tabel 1.8: Matriks Indikator dan Sumber Data Primer Resmi Bab 1
| No | Nama Indikator | Kategori Analisis | Satuan Ukur | Cakupan Tahun | Institusi & Sumber Data Resmi | Data File Asli |
| :---: | :--- | :--- | :---: | :---: | :--- | :--- |
| 1 | Izin Usaha Pertambangan (IUP) Baru | Faktor Tekanan Ekstraktif | Unit Izin | 2014-2024 | Data Registry ESDM MODI (Minerbaone) | `sulawesi_izin_baru_per_tahun.csv` |
| 2 | Luas Wilayah Konsesi Tambang Baru | Faktor Tekanan Ekstraktif | Hektar (Ha) | 2014-2024 | Data Registry ESDM MODI (Minerbaone) | `sulawesi_kawasan_nikel_luas.csv` |
| 3 | Kapasitas Terpasang PLTU Captive | Infrastruktur Energi Khusus | Megawatt (MW) | 2014-2024 | NGO (Global Energy Monitor / GEM) | `sulawesi_pltu_captive.csv` |
| 4 | Fasilitas Smelter Nikel | Fasilitas Industri Hilir | Unit Fasilitas | 2014-2024 | Database Smelter CGS & ESDM MODI | `sulawesi_esdm_nikel.csv` |
| 5 | Realisasi Investasi PMDN & Nikel | Arus Modal Domestik | Triliun Rupiah | 2016-2024 | API BPS & BKPM | `sulawesi_investasi_pmdn_2016_2024.csv` |
| 6 | PDRB Provinsi (Ekstraktif vs Akar Rumput) | Struktur Ekonomi Makro | Triliun Rupiah | 2016-2024 | API BPS (Subject 52) | `sulawesi_pdrb_sektoral_2016_2024.csv` |
| 7 | PDRB Kabupaten Sentra Tambang | Struktur Ekonomi Daerah | Triliun Rupiah | 2016-2024 | API BPS (Subject 52) | `sulawesi_pdrb_sektoral_kabupaten_2016_2024.csv` |
| 8 | Pendapatan Asli Daerah (PAD) & Breakdown Pajak | Kapasitas Fiskal Daerah | Triliun Rupiah | 2016-2024 | API BPS | `sulawesi_pad_breakdown_2016_2024.csv` |
| 9 | Luas Total Deforestasi Alam & Komoditas | Dampak Ekologis | Hektar (Ha) | 2014-2023 | Global Forest Watch (GFW API) | `sulawesi_gfw_master_1_dekade_2014_2023_v3.csv` |
| 10 | Simpul Pelabuhan & Terminal Logistik | Infrastruktur Rantai Pasok | Titik Koordinat & DWT | 2014-2024 | Laporan Publik (KNKT, Perpres PSN, Korporasi) | `sulawesi_logistik_simpul_nikel.csv` |

---

## 1.8 Bagan Alur Kerangka Kerja Riset Bab 1

Keseluruhan struktur metodologis riset Bab 1 dioperasionalkan melalui empat fase kerja berurutan sebagaimana disajikan pada **Tabel 1.9** berikut:

##### Tabel 1.9: Matriks Tahapan dan Alur Kerangka Kerja Riset Bab 1
| Tahapan Riset | Fokus Metodologis | Bahan & Sumber Data | Keluaran / Hasil Analisis |
| :--- | :--- | :--- | :--- |
| **Fase I: Pengumpulan Data** | Kurasi data resmi lintas kementerian dan lembaga | Publikasi BPS, Minerbaone, BKPM, GEM, dan GFW | Basis Data Tabular Panel Provinsi (2014–2024) |
| **Fase II: Reklasifikasi Hukum** | Penyusunan kerangka rantai pasok hukum terintegrasi | UU No. 3/2020, PP No. 96/2021, Perpres No. 112/2022 | 3 Klaster Makro (Ekstraktif, Akar Rumput, Jasa) |
| **Fase III: Pengujian Statistik** | Uji signifikansi hubungan dan rasio peluang | Tabel Kontinjensi, Uji Chi-Square, Odds Ratio | Bukti Kausalitas Signifikan Tekanan vs Deforestasi |
| **Fase IV: Pemetaan Rantai Pasok** | Triangulasi data logistik dan pemodelan maritim | Laporan KNKT, Perpres PSN, Kurva Parametrik Bézier | Peta Alur Rantai Pasok Ekspor & Konsentrasi Spasial 78% |
"""

    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab1    = out_dir_compact.parent.parent / "bab_1"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab1.mkdir(parents=True, exist_ok=True)
    
    md_path_compact = out_dir_compact / "Metodologi_Bab1_Ekspansi_Industri_Compact.md"
    md_path_bab1    = out_dir_bab1 / "Metodologi_Bab1_Ekspansi_Industri_Compact.md"

    with open(md_path_compact, "w", encoding="utf-8") as f:
        f.write(md_content)
    with open(md_path_bab1, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(f"[OK] Berhasil menyimpan Markdown di: {md_path_compact}")
    print(f"[OK] Salinan tersimpan di: {md_path_bab1}")

if __name__ == "__main__":
    print("=" * 70)
    print("GENERATOR METODOLOGI STATISTIK VERSI COMPACT - BAB 1")
    print("=" * 70)
    build_compact_report()
    generate_compact_markdown()
    print("=" * 70)
    print("SELESAI! Seluruh dokumen Versi Compact Bab 1 berhasil digenerate.")
    print("=" * 70)
