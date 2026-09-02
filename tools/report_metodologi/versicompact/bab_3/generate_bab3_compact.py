#!/usr/bin/env python3
"""
Generator Metodologi Versi Compact Bab 3 — GAYA AKADEMIS TERPADU (CELIOS)
Mengadopsi arsitektur metodologi ringkas terstandarisasi konsisten dengan Bab 1 & Bab 2:
- FORMAT: 1 KOLOM PENUH (Single Column Layout)
- PANJANG: 2–3 Halaman Maksimal (Elegan, proporsional, tanpa pemadatan berlebihan)
- PENOMORAN SEKSI UTAMA: Huruf kapital A, B, C, D, E, F
- SUB-BAB SEKSI D: 3.1, 3.2, 3.3, 3.4, 3.5, 3.6, 3.7 (Judul persis dokumen induk)
- OPERASIONALISASI INDIKATOR: 10 Indikator Empiris Lengkap (Matriks Indikator & Sumber Data Resmi)
- FORMULASI & TABEL CROSSTAB: Format standar Tabel 1.5b (Komponen Uji per baris, murni konfigurasi teknis tanpa hasil empiris)
- KORESPONDENSI METODOLOGI: 3 kolom bersih (Sub-bab, Fokus Kajian Empiris, Metode Analitis Utama)
- FLOWCHART: Mermaid JS horizontal (flowchart LR) dirender tajam ke DOCX (16.5 cm) dan blok kode di MD
- SINKRONISASI: Dual-save ke direktori versicompact/bab_3 dan bab_3.
"""

import os
import sys
import shutil
import base64
import requests
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Palet Warna Resmi CELIOS ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Forest Dark Green (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Celios Accent Green (#2E7D32)
G_LIGHT = RGBColor(0x38, 0x8E, 0x3C)  # Medium Green (#388E3C)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Charcoal Body Text (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Muted Grey Text (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Pure White

# ── Helper Fungsi XML Word ───────────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=80, right=80, top=40, bottom=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)

def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    pPr.append(s)

def para_border_bottom(p, color='1B5E20', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '3')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def para_border_left(p, color='2E7D32', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:left')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='1B5E20', sz='8'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Typography & Content ──────────────────────────────
def add_run(p, text, bold=False, italic=False, pt=8.5, color=C_BODY, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    else:
        r.font.name = 'Calibri'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    return r

def add_h2(doc, prefix, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_border_bottom(p, color='1B5E20', sz='8')
    add_run(p, f"{prefix}.  ", bold=True, pt=11, color=G_DARK)
    add_run(p, title.upper(), bold=True, pt=11, color=G_DARK)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, pt=9.5, color=G_MID)
    return p

def add_body(doc, parts, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        add_run(p, text, bold=bold, italic=italic, pt=8.5, color=C_BODY)
    return p

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    para_shd(p, 'EDF7EE')
    add_run(p, text, pt=8, color=G_MID, mono=True)

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_styled(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    bd_subtle = {'val': 'single', 'sz': '4', 'color': 'D0D7DE', 'space': '0'}
    
    # Header Row
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=80, right=80, top=50, bottom=50)
        set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if alignments and j < len(alignments):
            align = alignments[j]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, pt=8, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F9FBF9' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_shd(c, fill)
            cell_margin(c, left=80, right=80, top=40, bottom=40)
            set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            
            if alignments and j < len(alignments):
                align = alignments[j]
                if align == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_run(p, str(val), pt=7.5, color=C_BODY)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after  = Pt(3)
    return tbl

def download_mermaid_png(mermaid_str, filepath):
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        return True
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False


def generate_bab3_compact():
    print("[1/3] Membangun dokumen compact Bab 3 (Format 1-Kolom, 2-3 Halaman)...")
    
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab3    = out_dir_compact.parent.parent / "bab_3"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab3.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER DOKUMEN ──────────────────────────────────────────
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after  = Pt(1)
    add_run(p_t, "RINGKASAN EKSEKUTIF METODOLOGIS", bold=True, pt=8.5, color=G_MID)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after  = Pt(2)
    para_border_bottom(p_h, color='1B5E20', sz='12')
    add_run(p_h, "BAB 3: METODOLOGI ANALISIS BEBAN KESEHATAN MASYARAKAT TERDAMPAK", bold=True, pt=15, color=G_DARK)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(1)
    p_meta.paragraph_format.space_after  = Pt(5)
    add_run(p_meta, "Studi Daya Dukung & Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi · ", italic=True, pt=8, color=C_GREY)
    add_run(p_meta, "Center of Economic and Law Studies (CELIOS)", bold=True, italic=True, pt=8, color=G_DARK)

    # ── A. DESAIN PENELITIAN & TUJUAN ───────────────────────────
    add_h2(doc, "A", "Desain Penelitian & Tujuan")
    add_body(doc, [
        ("Penelitian ini menggunakan ", False, False),
        ("desain epidemiologi lingkungan dan audit spasial-statistik kuantitatif terintegrasi", True, False),
        (" untuk mengukur beban morbiditas kesehatan masyarakat, defisit fasilitas layanan kesehatan, serta paparan eksternalitas limbah beracun di enam provinsi Pulau Sulawesi sepanjang satu dekade (", False, False),
        ("2014–2024", True, False),
        ("). Tiga tujuan utama metodologis Bab 3 meliputi:", False, False)
    ])
    add_body(doc, [
        ("1. ", True, False), ("Membuktikan Disparitas Fasilitas & Morbiditas Kesehatan: ", True, False),
        ("Mengevaluasi kesenjangan rasio ketersediaan faskes (Puskesmas vs Rumah Sakit) dan membandingkan rata-rata beban penyakit pernapasan (ISPA) serta pencernaan (Diare) antara provinsi Sentra Industri nikel vs Non-Sentra.\n", False, False),
        ("2. ", True, False), ("Analisis Inferensial Panel & Dinamika Zoonosis Tapak: ", True, False),
        ("Menguji signifikansi korelasi antara penurunan indeks kualitas lingkungan (IKU & IKA) terhadap lonjakan kasus penyakit melalui uji Chi-Square dan Odds Ratio, serta mengisolasi anomali vektor zoonosis di kabupaten lingkar tambang.\n", False, False),
        ("3. ", True, False), ("Validasi Toksisitas Dua Lensa & Neraca Limbah B3: ", True, False),
        ("Memadukan analisis makro provinsi dengan pembuktian klinis mikroskopis logam berat karsinogenik Kromium Heksavalen (Cr6+) di muara tambang, serta mengagregasi timbulan 32,8 juta ton limbah B3 slag dan tailing HPAL.", False, False)
    ])

    # ── B. SUMBER DATA & CAKUPAN WILAYAH ─────────────────────────
    add_h2(doc, "B", "Sumber Data & Cakupan Wilayah")
    add_body(doc, [
        ("Penelitian mencakup analisis lintas provinsi pada ", False, False),
        ("6 provinsi Pulau Sulawesi", True, False),
        (" (Sulawesi Tengah, Sulawesi Tenggara, Sulawesi Selatan, Sulawesi Barat, Gorontalo, Sulawesi Utara) serta ", False, False),
        ("deep-dive case study tingkat kabupaten/distrik lingkar industri", True, False),
        (" (Morowali, Morowali Utara, Banggai, Konawe, Bantaeng). Data dihimpun dari sumber data primer resmi kementerian, dinas kesehatan daerah, registri BPS, dan audit laboratorium independen:", False, False)
    ])
    add_body(doc, [
        ("• ", True, False), ("Badan Pusat Statistik (BPS) & Kementerian Kesehatan RI: ", True, False),
        ("Registri unit fasilitas kesehatan (Puskesmas dan Rumah Sakit) serta sensus populasi denominator per kapita.\n", False, False),
        ("• ", True, False), ("Dinas Kesehatan Provinsi Se-Sulawesi (Profil Kesehatan 2014–2024): ", True, False),
        ("Data time-series insidensi penyakit ISPA/Pneumonia, Diare terlayani, Malaria, DBD, Filariasis, dan Rabies.\n", False, False),
        ("• ", True, False), ("Kementerian Lingkungan Hidup dan Kehutanan (Ditjen PPKL): ", True, False),
        ("Indeks Kualitas Udara (IKU) dan Indeks Kualitas Air (IKA) time-series panel provinsi-tahun (2015–2024).\n", False, False),
        ("• ", True, False), ("Audit Fisik Laboratorium Independen (AEER & WALHI): ", True, False),
        ("Uji konsentrasi Kromium Heksavalen (Cr6+ dalam satuan mg/L) pada 12 titik sampling sungai dan pesisir lingkar smelter.\n", False, False),
        ("• ", True, False), ("Registri Audit Limbah B3 (KLHK, AEER, WALHI, JATAM): ", True, False),
        ("Neraca timbulan terak slag nikel, tailing HPAL (asam sulfat), air limbah tambang, dan limbah EAF per fasilitas mayor industri.", False, False)
    ])

    # ── C. OPERASIONALISASI VARIABEL & INDIKATOR RISET ──────────
    add_h2(doc, "C", "Operasionalisasi Variabel & Indikator Riset")
    add_body(doc, [
        ("Seluruh variabel kesehatan masyarakat, kualitas sanitasi, toksisitas klinis, dan limbah industri dioperasionalkan ke dalam ", False, False),
        ("10 indikator empiris terpadu", True, False),
        (" sebagaimana dirangkum pada matriks operasional berikut:", False, False)
    ])

    add_caption(doc, "Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 3")
    table_indikator_data = [
        ["1", "Ketersediaan Fasilitas Kesehatan", "Rasio Rumah Sakit & Puskesmas per Zona", "Unit Faskes", "2024", "BPS & Kemenkes RI"],
        ["2", "Beban Penyakit ISPA / Pneumonia", "Morbiditas Saluran Pernapasan Akut", "Kasus Absolut", "2014–2024", "Dinas Kesehatan Provinsi"],
        ["3", "Beban Kasus Diare Terlayani", "Morbiditas Saluran Pencernaan & Sanitasi", "Kasus Absolut", "2014–2024", "Dinas Kesehatan Provinsi"],
        ["4", "Tingkat Insidensi per Kapita", "Normalisasi Beban Penyakit terhadap Populasi", "Kasus / 10.000 Jiwa", "2014–2024", "Dinkes & Populasi BPS"],
        ["5", "Indeks Kualitas Udara (IKU)", "Kondisi Baku Mutu Udara Ambien Agregat", "Poin Skor (0–100)", "2015–2024", "Ditjen PPKL KLHK"],
        ["6", "Indeks Kualitas Air (IKA)", "Kondisi Baku Mutu Air Sungai & DAS Agregat", "Poin Skor (0–100)", "2016–2024", "Ditjen PPKL KLHK"],
        ["7", "Prevalensi Vektor Zoonosis", "Insidensi DBD, Malaria, & Filariasis Tapak", "Kasus / Distrik", "2015–2024", "Dinkes Sulteng (Tapak)"],
        ["8", "Kadar Kromium Heksavalen (Cr6+)", "Toksisitas Logam Berat Karsinogenik Tapak", "mg / Liter", "2022–2024", "Uji Lab AEER & WALHI"],
        ["9", "Timbulan Limbah B3 Industri", "Volume Residu Slag & Tailing HPAL", "Juta Ton / Tahun", "2024–2025", "KLHK, AEER, WALHI, JATAM"],
        ["10", "Dinamika Spasial Before-After", "Pergeseran Spasial Morbiditas Ekologis", "Rasio Pertumbuhan (%)", "2015 vs 2024", "GeoJSON & Profil Dinkes"]
    ]

    add_table_styled(
        doc,
        headers=["No", "Indikator Riset", "Fokus Pengukuran", "Satuan", "Periode", "Sumber Data Primer Resmi"],
        rows=table_indikator_data,
        col_widths_cm=[0.8, 3.7, 3.8, 1.8, 1.9, 5.0],
        alignments=['C', 'L', 'L', 'C', 'C', 'L']
    )

    # ── D. KERANGKA ANALISIS & FORMULASI MATEMATIS ──────────────
    add_h2(doc, "D", "Kerangka Analisis & Formulasi Matematis")

    # 3.1
    add_h3(doc, "3.1 Kesenjangan Fasilitas Kesehatan di Kawasan Ekstraktif")
    add_body(doc, [
        ("Kesenjangan fasilitas pelayanan kesehatan dianalisis melalui segmentasi cross-sectional per jenis fasilitas (Puskesmas vs Rumah Sakit) antara zona sentra industri ekstraktif dan zona non-sentra:", False, False)
    ])
    add_formula(doc, "Rata-rata Faskes (F̄_z,j) = [ Σ F_p,j ] / n_z   ;   Rasio Disparitas (D_j) = F̄_Sentra,j / F̄_Non-Sentra,j")

    # 3.2
    add_h3(doc, "3.2 Ketimpangan Beban Penyakit: Sentra Industri vs Non-Sentra")
    add_body(doc, [
        ("Komparasi beban morbiditas penyakit pernapasan dan pencernaan dihitung guna mengukur disparitas kelipatan risiko kesehatan pada provinsi lingkar hilirisasi:", False, False)
    ])
    add_formula(doc, "Beban Rata-rata (B̄_z) = [ Σ B_p,t ] / N_z   ;   Kelipatan Disparitas (Q) = B̄_Sentra / B̄_Non-Sentra")

    # 3.3
    add_h3(doc, "3.3 Lintasan Waktu Ekologis & Dinamika Penyakit di Kawasan Industri Ekstraktif")
    add_body(doc, [
        ("Normalisasi beban penyakit per 10.000 penduduk dan protokol pengujian independensi Chi-Square (χ²) tabulasi silang diterapkan pada matriks panel provinsi-tahun berbasis ambang batas median spesifik provinsi:", False, False)
    ])
    add_formula(doc, "Insiden_10K = ( Kasus_p,t / Populasi_p ) × 10.000   ;   χ² = Σ [ ( O_ij - E_ij )² / E_ij ]")

    add_caption(doc, "Tabel 3.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 3.3)")
    tabel_3_3a_rows = [
        ["Variabel Independen (X)", "IKU Wilayah Sentra Tambang / IKU Wilayah Non-Sentra (indeks tekanan kualitas lingkungan)."],
        ["Variabel Dependen (Y)", "Total Kasus ISPA/Pneumonia (insidensi penyakit pernapasan dan lingkungan)."],
        ["Hipotesis Nol (H0)", "Penurunan kualitas lingkungan (IKU/IKA) tidak berkorelasi dengan peningkatan insidensi penyakit pernapasan dan pencernaan."],
        ["Hipotesis Alternatif (H1)", "Penurunan kualitas udara ambien (IKU) berbanding lurus dengan peningkatan insidensi penyakit pernapasan dan lingkungan (ISPA dan Diare)."],
        ["Decision Rule (Alpha 5%)", "Chi-Square P-Value < 0.05 (Tolak H0) dan kalkulasi Odds Ratio."],
        ["Threshold Kategori", "Median per-provinsi data panel Provinsi-Tahun (N=18 observasi valid skenario Sentra); binning 'Tinggi'/'Rendah' per provinsi untuk menghilangkan bias besaran absolut antar wilayah."],
        ["Orientasi Odds Ratio", "Untuk variabel X berjenis indeks kualitas (IKU/IKA), risiko dihitung saat indeks Rendah: OR = ( b × c ) / ( a × d )."]
    ]
    add_table_styled(
        doc,
        headers=["Komponen Uji", "Definisi Variabel (Sub-bab 3.3)"],
        rows=tabel_3_3a_rows,
        col_widths_cm=[4.5, 12.5],
        alignments=['L', 'L']
    )

    # 3.4
    add_h3(doc, "3.4 Anomali Zoonosis: Dampak Kritis Ekspansi Industri di Level Tapak (Studi Kasus Sulteng)")
    add_body(doc, [
        ("Isolasi data tapak tingkat distrik/kabupaten sentra tambang aktif (Morowali, Morowali Utara, Banggai) mengukur lonjakan vektor zoonosis akibat genangan lubang tambang dan sanitasi industri dibandingkan kabupaten kontrol:", False, False)
    ])
    add_formula(doc, "Akumulasi Zoonosis (Z_w,t,d) = Σ C_r,t,d   ;   Rasio Zoonosis Tapak (R_d) = Z̄_Tambang / Z̄_Kontrol")

    # 3.5
    add_h3(doc, "3.5 Pemetaan Geospasial: Distribusi Spasial Beban Penyakit")
    add_body(doc, [
        ("Pemodelan Before-After Analysis memproyeksikan pergeseran spasial intensitas morbiditas antara tahun acuan awal ekstraksi (2015) dan kondisi terkini (2024):", False, False)
    ])
    add_formula(doc, "Radius Bubble Diare (r_p,t) = √D_p,t / K   ;   Laju Pertumbuhan (G_p %) = [ ( X_2024 - X_2015 ) / X_2015 ] × 100")

    # 3.6
    add_h3(doc, "3.6 Krisis Air Bersih: Tinjauan Makro Provinsi dan Bukti Uji Klinis Lingkar Tambang")
    add_body(doc, [
        ("Pendekatan komplementer dua lensa memadukan benchmark rasio toksisitas laboratorium Kromium Heksavalen (Cr6+) pada muara sungai tambang dengan uji Chi-Square data panel makro IKA vs Diare:", False, False)
    ])
    add_formula(doc, "Rasio Pelanggaran Toksisitas = Konsentrasi Cr6+ / Baku Mutu Biota Laut (0.005 mg/L)")

    add_caption(doc, "Tabel 3.6a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 3.6)")
    tabel_3_6a_rows = [
        ["Variabel Independen (X)", "IKA Wilayah Sentra Tambang / IKA Wilayah Non-Sentra (Indeks Kualitas Air BPS/KLHK)."],
        ["Variabel Dependen (Y)", "Total Kasus Diare (kasus infeksi saluran pencernaan yang dilayani, Kemenkes)."],
        ["Hipotesis Nol (H0)", "Rendahnya Indeks Kualitas Air (IKA) tidak berhubungan dengan tingginya kasus Diare."],
        ["Hipotesis Alternatif (H1)", "Provinsi dengan IKA rendah berasosiasi signifikan dengan peningkatan kasus Diare."],
        ["Decision Rule (Alpha 5%)", "Jika P-Value < 0.05, maka Tolak H0 (terbukti signifikan bahwa pencemaran air meningkatkan kasus Diare)."],
        ["Threshold Kategori", "Median per-provinsi data panel Provinsi-Tahun (N=16 observasi valid skenario Sentra dari 6 provinsi × 8 tahun); binning 'Tinggi'/'Rendah' per provinsi."],
        ["Orientasi Odds Ratio", "Karena IKA indikator positif (semakin tinggi semakin baik), risiko dihitung saat IKA Rendah: OR = ( b × c ) / ( a × d )."]
    ]
    add_table_styled(
        doc,
        headers=["Komponen Uji", "Definisi Variabel (Sub-bab 3.6)"],
        rows=tabel_3_6a_rows,
        col_widths_cm=[4.5, 12.5],
        alignments=['L', 'L']
    )

    # 3.7
    add_h3(doc, "3.7 Beban Limbah Beracun (B3): Eksternalitas Kesehatan yang Diabaikan")
    add_body(doc, [
        ("Kuantifikasi neraca timbulan limbah bahan berbahaya dan beracun (B3) mengagregasi volume pelepasan residu padat dan cair per provinsi serta menghitung proporsi jenis limbah spesifik:", False, False)
    ])
    add_formula(doc, "Total Timbulan B3_p = Σ [ Timbulan Fasilitas Mayor_i ]   ;   Proporsi Jenis B3 (%) = [ Total B3_j / Total B3 ] × 100")

    # ── E. KORESPONDENSI METODOLOGI TERHADAP SUB-BAB LAPORAN ────
    add_h2(doc, "E", "Korespondensi Metodologi terhadap Sub-bab Laporan Bab 3")
    add_body(doc, [
        ("Setiap sub-bab analitis pada Bab 3 ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:", False, False)
    ])

    table_korespondensi = [
        ["Sub-bab 3.1", "Kesenjangan Fasilitas Kesehatan", "Grouped Horizontal Bar Chart, Rasio Disparitas Faskes per Zona Industri"],
        ["Sub-bab 3.2", "Ketimpangan Beban Morbiditas", "Comparative Spatial Analysis Sentra vs Non-Sentra, Kelipatan Disparitas Morbiditas"],
        ["Sub-bab 3.3", "Lintasan Waktu Insidensi & IKU", "Time-Series Line Chart per 10.000 Jiwa, Uji Chi-Square (χ²), Odds Ratio (OR)"],
        ["Sub-bab 3.4", "Anomali Vektor Zoonosis Tapak", "Deep-Dive Case Study Distrik Tambang Sulteng, Analisis Komparasi Wilayah Kontrol"],
        ["Sub-bab 3.5", "Pemetaan Spasial Beban Penyakit", "Choropleth Poligon ISPA & Radius Bubble Diare, Before-After Analysis (2015 vs 2024)"],
        ["Sub-bab 3.6", "Krisis Air Bersih & Toksisitas Cr6+", "Pendekatan Dua Lensa (Mikro Lab vs Makro Panel), Regresi OLS, Uji Chi-Square (χ²)"],
        ["Sub-bab 3.7", "Eksternalitas Limbah Beracun B3", "Agregasi Timbulan Regional, Profiling Fasilitas Mayor, Analisis Komposisi Residu"]
    ]

    add_table_styled(
        doc,
        headers=["Sub-bab", "Fokus Kajian Empiris", "Metode Analitis Utama"],
        rows=table_korespondensi,
        col_widths_cm=[2.5, 5.5, 9.0],
        alignments=['C', 'L', 'L']
    )

    # ── F. BAGAN ALUR KERANGKA KERJA RISET BAB 3 ────────────────
    add_h2(doc, "F", "Bagan Alur Kerangka Kerja Riset (Research Workflow)")
    add_body(doc, [
        ("Kerangka operasional metodologi Bab 3 berjalan secara terpadu melalui empat fase berurutan sebagaimana divisualisasikan pada bagan alur kerja riset berikut:", False, False)
    ])

    mermaid_str_f = """flowchart LR
    subgraph F1["Fase I: Akuisisi Data"]
        A1["Kurasi Data Resmi Terbuka<br/><i>BPS, Kemenkes, KLHK, Lab AEER & WALHI</i>"]
        A2["Panel Provinsi-Tahun<br/><i>6 Provinsi Se-Sulawesi (2014–2024)</i>"]
    end
    subgraph F2["Fase II: Segmentasi & Normalisasi"]
        B1["Segmentasi Spasial Ekstraktif<br/><i>Sentra Tambang vs Non-Sentra Kontrol</i>"]
        B2["Standarisasi per Kapita<br/><i>Rasio Insiden per 10.000 Penduduk</i>"]
    end
    subgraph F3["Fase III: Uji Inferensial & Klinis"]
        C1["Tabel Kontinjensi 2x2<br/><i>Ambang Median IKU/IKA vs Morbiditas</i>"]
        C2["Uji Chi-Square & Toksisitas<br/><i>Benchmark Cr6+ & Odds Ratio Risiko</i>"]
    end
    subgraph F4["Fase IV: Pemodelan & Sintesis"]
        D1["Pemetaan Spasial Dual-Lensa<br/><i>Choropleth ISPA & Bubble Diare</i>"]
        D2["Bukti Kausalitas D3TLH<br/><i>Defisit Faskes & Ancaman Limbah B3</i>"]
    end
    F1 --> F2 --> F3 --> F4"""

    png_workflow_path = str(out_dir_compact / "mermaid_workflow_bab3.png")
    is_downloaded = download_mermaid_png(mermaid_str_f, png_workflow_path)

    add_caption(doc, "Bagan Alur 3.1: Alur Logika Kerangka Kerja Riset Bab 3 (Research Workflow)")
    if is_downloaded and os.path.exists(png_workflow_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(3)
        p_img.paragraph_format.space_after  = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(png_workflow_path, width=Cm(16.5))
        try:
            shutil.copyfile(png_workflow_path, str(out_dir_bab3 / "mermaid_workflow_bab3.png"))
        except Exception:
            pass

    # Box Output Kesimpulan
    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after  = Pt(4)
    all_border_para(p_box, color='1B5E20', sz='8')
    para_shd(p_box, 'F1F8E9')
    add_run(p_box, "KERANGKA KELUARAN METODOLOGIS BAB 3:\n", bold=True, pt=8.5, color=G_DARK)
    add_run(p_box, "1. Konfigurasi Disparitas Infrastruktur & Morbiditas: Membuktikan paradoks kawasan sentra hilirisasi yang memikul beban penyakit pernapasan 2x lipat lebih tinggi di tengah defisit fasilitas rumah sakit sekunder.\n"
                   "2. Konfigurasi Dua Lensa Krisis Sanitasi & Toksisitas: Memadukan pengujian inferensial data panel makro dengan bukti mikroskopis pelanggaran baku mutu logam berat karsinogenik Kromium Heksavalen (Cr6+) hingga 200 kali lipat di muara tambang.\n"
                   "3. Konfigurasi Neraca Eksternalitas Limbah Beracun (B3): Mengkuantifikasi timbulan 32,8 juta ton per tahun limbah slag dan tailing HPAL sebagai ancaman risiko kesehatan masyarakat jangka panjang bagi dokumen D3TLH.",
            pt=8, color=C_BODY)

    # ── SIMPAN DOKUMEN DOCX (DUAL SAVE) ─────────────────────────
    docx_compact = out_dir_compact / "Metodologi_Bab3_Beban_Kesehatan_Compact.docx"
    docx_bab3    = out_dir_bab3 / "Metodologi_Bab3_Beban_Kesehatan_Compact.docx"
    
    doc.save(str(docx_compact))
    shutil.copyfile(docx_compact, docx_bab3)
    print(f"  [OK] Tersimpan DOCX: {docx_compact}")
    print(f"  [OK] Salinan DOCX : {docx_bab3}")

    # ── GENERATE MARKDOWN PADANAN ───────────────────────────────
    print("[2/3] Membangun dokumen Markdown padanan...")
    MD_CONTENT = """# METODOLOGI PENELITIAN: BAB 3 — ANALISIS BEBAN KESEHATAN MASYARAKAT TERDAMPAK
*CELIOS (Center of Economic and Law Studies) · Audit Spasial-Statistik D3TLH Sulawesi (2014–2024) · Ringkasan Eksekutif Metodologis*

---

## A. Desain Penelitian & Tujuan
Penelitian ini menggunakan **desain epidemiologi lingkungan dan audit spasial-statistik kuantitatif terintegrasi** untuk mengukur beban morbiditas kesehatan masyarakat, defisit fasilitas layanan kesehatan, serta paparan eksternalitas limbah beracun di enam provinsi Pulau Sulawesi sepanjang satu dekade (**2014–2024**). Tiga tujuan utama metodologis Bab 3 meliputi:

1. **Membuktikan Disparitas Fasilitas & Morbiditas Kesehatan:** Mengevaluasi kesenjangan rasio ketersediaan faskes (Puskesmas vs Rumah Sakit) dan membandingkan rata-rata beban penyakit pernapasan (ISPA) serta pencernaan (Diare) antara provinsi Sentra Industri nikel vs Non-Sentra.
2. **Analisis Inferensial Panel & Dinamika Zoonosis Tapak:** Menguji signifikansi korelasi antara penurunan indeks kualitas lingkungan (IKU & IKA) terhadap lonjakan kasus penyakit melalui uji Chi-Square dan Odds Ratio, serta mengisolasi anomali vektor zoonosis di kabupaten lingkar tambang.
3. **Validasi Toksisitas Dua Lensa & Neraca Limbah B3:** Memadukan analisis makro provinsi dengan pembuktian klinis mikroskopis logam berat karsinogenik Kromium Heksavalen (Cr6+) di muara tambang, serta mengagregasi timbulan 32,8 juta ton limbah B3 slag dan tailing HPAL.

---

## B. Sumber Data & Cakupan Wilayah
Penelitian mencakup analisis lintas provinsi pada **6 provinsi Pulau Sulawesi** (Sulawesi Tengah, Sulawesi Tenggara, Sulawesi Selatan, Sulawesi Barat, Gorontalo, Sulawesi Utara) serta **deep-dive case study tingkat kabupaten/distrik lingkar industri** (Morowali, Morowali Utara, Banggai, Konawe, Bantaeng). Data dihimpun dari sumber data primer resmi kementerian, dinas kesehatan daerah, registri BPS, dan audit laboratorium independen:

- **Badan Pusat Statistik (BPS) & Kementerian Kesehatan RI:** Registri unit fasilitas kesehatan (Puskesmas dan Rumah Sakit) serta sensus populasi denominator per kapita.
- **Dinas Kesehatan Provinsi Se-Sulawesi (Profil Kesehatan 2014–2024):** Data time-series insidensi penyakit ISPA/Pneumonia, Diare terlayani, Malaria, DBD, Filariasis, dan Rabies.
- **Kementerian Lingkungan Hidup dan Kehutanan (Ditjen PPKL):** Indeks Kualitas Udara (IKU) dan Indeks Kualitas Air (IKA) time-series panel provinsi-tahun (2015–2024).
- **Audit Fisik Laboratorium Independen (AEER & WALHI):** Uji konsentrasi Kromium Heksavalen (Cr6+ dalam satuan mg/L) pada 12 titik sampling sungai dan pesisir lingkar smelter.
- **Registri Audit Limbah B3 (KLHK, AEER, WALHI, JATAM):** Neraca timbulan terak slag nikel, tailing HPAL (asam sulfat), air limbah tambang, dan limbah EAF per fasilitas mayor industri.

---

## C. Operasionalisasi Variabel & Indikator Riset
Seluruh variabel kesehatan masyarakat, kualitas sanitasi, toksisitas klinis, dan limbah industri dioperasionalkan ke dalam **10 indikator empiris terpadu** sebagaimana dirangkum pada matriks operasional berikut:

##### Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 3
| No | Indikator Riset | Fokus Pengukuran | Satuan | Periode | Sumber Data Primer Resmi |
| :-: | :--- | :--- | :-: | :-: | :--- |
| 1 | Ketersediaan Fasilitas Kesehatan | Rasio Rumah Sakit & Puskesmas per Zona | Unit Faskes | 2024 | BPS & Kemenkes RI |
| 2 | Beban Penyakit ISPA / Pneumonia | Morbiditas Saluran Pernapasan Akut | Kasus Absolut | 2014–2024 | Dinas Kesehatan Provinsi |
| 3 | Beban Kasus Diare Terlayani | Morbiditas Saluran Pencernaan & Sanitasi | Kasus Absolut | 2014–2024 | Dinas Kesehatan Provinsi |
| 4 | Tingkat Insidensi per Kapita | Normalisasi Beban Penyakit terhadap Populasi | Kasus / 10.000 Jiwa | 2014–2024 | Dinkes & Populasi BPS |
| 5 | Indeks Kualitas Udara (IKU) | Kondisi Baku Mutu Udara Ambien Agregat | Poin Skor (0–100) | 2015–2024 | Ditjen PPKL KLHK |
| 6 | Indeks Kualitas Air (IKA) | Kondisi Baku Mutu Air Sungai & DAS Agregat | Poin Skor (0–100) | 2016–2024 | Ditjen PPKL KLHK |
| 7 | Prevalensi Vektor Zoonosis | Insidensi DBD, Malaria, & Filariasis Tapak | Kasus / Distrik | 2015–2024 | Dinkes Sulteng (Tapak) |
| 8 | Kadar Kromium Heksavalen (Cr6+) | Toksisitas Logam Berat Karsinogenik Tapak | mg / Liter | 2022–2024 | Uji Lab AEER & WALHI |
| 9 | Timbulan Limbah B3 Industri | Volume Residu Slag & Tailing HPAL | Juta Ton / Tahun | 2024–2025 | KLHK, AEER, WALHI, JATAM |
| 10 | Dinamika Spasial Before-After | Pergeseran Spasial Morbiditas Ekologis | Rasio Pertumbuhan (%) | 2015 vs 2024 | GeoJSON & Profil Dinkes |

---

## D. Kerangka Analisis & Formulasi Matematis

### 3.1 Kesenjangan Fasilitas Kesehatan di Kawasan Ekstraktif
Kesenjangan fasilitas pelayanan kesehatan dianalisis melalui segmentasi cross-sectional per jenis fasilitas (Puskesmas vs Rumah Sakit) antara zona sentra industri ekstraktif dan zona non-sentra:

> `Rata-rata Faskes (F̄_z,j) = [ Σ F_p,j ] / n_z   ;   Rasio Disparitas (D_j) = F̄_Sentra,j / F̄_Non-Sentra,j`

### 3.2 Ketimpangan Beban Penyakit: Sentra Industri vs Non-Sentra
Komparasi beban morbiditas penyakit pernapasan dan pencernaan dihitung guna mengukur disparitas kelipatan risiko kesehatan pada provinsi lingkar hilirisasi:

> `Beban Rata-rata (B̄_z) = [ Σ B_p,t ] / N_z   ;   Kelipatan Disparitas (Q) = B̄_Sentra / B̄_Non-Sentra`

### 3.3 Lintasan Waktu Ekologis & Dinamika Penyakit di Kawasan Industri Ekstraktif
Normalisasi beban penyakit per 10.000 penduduk dan protokol pengujian independensi Chi-Square (χ²) tabulasi silang diterapkan pada matriks panel provinsi-tahun berbasis ambang batas median spesifik provinsi:

> `Insiden_10K = ( Kasus_p,t / Populasi_p ) × 10.000   ;   χ² = Σ [ ( O_ij - E_ij )² / E_ij ]`

##### Tabel 3.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 3.3)
| Komponen Uji | Definisi Variabel (Sub-bab 3.3) |
| :--- | :--- |
| **Variabel Independen (X)** | IKU Wilayah Sentra Tambang / IKU Wilayah Non-Sentra (indeks tekanan kualitas lingkungan). |
| **Variabel Dependen (Y)** | Total Kasus ISPA/Pneumonia (insidensi penyakit pernapasan dan lingkungan). |
| **Hipotesis Nol (H0)** | Penurunan kualitas lingkungan (IKU/IKA) tidak berkorelasi dengan peningkatan insidensi penyakit pernapasan dan pencernaan. |
| **Hipotesis Alternatif (H1)** | Penurunan kualitas udara ambien (IKU) berbanding lurus dengan peningkatan insidensi penyakit pernapasan dan lingkungan (ISPA dan Diare). |
| **Decision Rule (Alpha 5%)** | Chi-Square P-Value < 0.05 (Tolak H0) dan kalkulasi Odds Ratio. |
| **Threshold Kategori** | Median per-provinsi data panel Provinsi-Tahun (N=18 observasi valid skenario Sentra); binning 'Tinggi'/'Rendah' per provinsi untuk menghilangkan bias besaran absolut antar wilayah. |
| **Orientasi Odds Ratio** | Untuk variabel X berjenis indeks kualitas (IKU/IKA), risiko dihitung saat indeks Rendah: OR = ( b × c ) / ( a × d ). |

### 3.4 Anomali Zoonosis: Dampak Kritis Ekspansi Industri di Level Tapak (Studi Kasus Sulteng)
Isolasi data tapak tingkat distrik/kabupaten sentra tambang aktif (Morowali, Morowali Utara, Banggai) mengukur lonjakan vektor zoonosis akibat genangan lubang tambang dan sanitasi industri dibandingkan kabupaten kontrol:

> `Akumulasi Zoonosis (Z_w,t,d) = Σ C_r,t,d   ;   Rasio Zoonosis Tapak (R_d) = Z̄_Tambang / Z̄_Kontrol`

### 3.5 Pemetaan Geospasial: Distribusi Spasial Beban Penyakit
Pemodelan Before-After Analysis memproyeksikan pergeseran spasial intensitas morbiditas antara tahun acuan awal ekstraksi (2015) dan kondisi terkini (2024):

> `Radius Bubble Diare (r_p,t) = √D_p,t / K   ;   Laju Pertumbuhan (G_p %) = [ ( X_2024 - X_2015 ) / X_2015 ] × 100`

### 3.6 Krisis Air Bersih: Tinjauan Makro Provinsi dan Bukti Uji Klinis Lingkar Tambang
Pendekatan komplementer dua lensa memadukan benchmark rasio toksisitas laboratorium Kromium Heksavalen (Cr6+) pada muara sungai tambang dengan uji Chi-Square data panel makro IKA vs Diare:

> `Rasio Pelanggaran Toksisitas = Konsentrasi Cr6+ / Baku Mutu Biota Laut (0.005 mg/L)`

##### Tabel 3.6a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 3.6)
| Komponen Uji | Definisi Variabel (Sub-bab 3.6) |
| :--- | :--- |
| **Variabel Independen (X)** | IKA Wilayah Sentra Tambang / IKA Wilayah Non-Sentra (Indeks Kualitas Air BPS/KLHK). |
| **Variabel Dependen (Y)** | Total Kasus Diare (kasus infeksi saluran pencernaan yang dilayani, Kemenkes). |
| **Hipotesis Nol (H0)** | Rendahnya Indeks Kualitas Air (IKA) tidak berhubungan dengan tingginya kasus Diare. |
| **Hipotesis Alternatif (H1)** | Provinsi dengan IKA rendah berasosiasi signifikan dengan peningkatan kasus Diare. |
| **Decision Rule (Alpha 5%)** | Jika P-Value < 0.05, maka Tolak H0 (terbukti signifikan bahwa pencemaran air meningkatkan kasus Diare). |
| **Threshold Kategori** | Median per-provinsi data panel Provinsi-Tahun (N=16 observasi valid skenario Sentra dari 6 provinsi × 8 tahun); binning 'Tinggi'/'Rendah' per provinsi. |
| **Orientasi Odds Ratio** | Karena IKA indikator positif (semakin tinggi semakin baik), risiko dihitung saat IKA Rendah: OR = ( b × c ) / ( a × d ). |

### 3.7 Beban Limbah Beracun (B3): Eksternalitas Kesehatan yang Diabaikan
Kuantifikasi neraca timbulan limbah bahan berbahaya dan beracun (B3) mengagregasi volume pelepasan residu padat dan cair per provinsi serta menghitung proporsi jenis limbah spesifik:

> `Total Timbulan B3_p = Σ [ Timbulan Fasilitas Mayor_i ]   ;   Proporsi Jenis B3 (%) = [ Total B3_j / Total B3 ] × 100`

---

## E. Korespondensi Metodologi terhadap Sub-bab Laporan Bab 3
Setiap sub-bab analitis pada Bab 3 ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:

##### Matriks Korespondensi Sub-bab terhadap Metode Analitis
| Sub-bab | Fokus Kajian Empiris | Metode Analitis Utama |
| :---: | :--- | :--- |
| **Sub-bab 3.1** | Kesenjangan Fasilitas Kesehatan | Grouped Horizontal Bar Chart, Rasio Disparitas Faskes per Zona Industri |
| **Sub-bab 3.2** | Ketimpangan Beban Morbiditas | Comparative Spatial Analysis Sentra vs Non-Sentra, Kelipatan Disparitas Morbiditas |
| **Sub-bab 3.3** | Lintasan Waktu Insidensi & IKU | Time-Series Line Chart per 10.000 Jiwa, Uji Chi-Square (χ²), Odds Ratio (OR) |
| **Sub-bab 3.4** | Anomali Vektor Zoonosis Tapak | Deep-Dive Case Study Distrik Tambang Sulteng, Analisis Komparasi Wilayah Kontrol |
| **Sub-bab 3.5** | Pemetaan Spasial Beban Penyakit | Choropleth Poligon ISPA & Radius Bubble Diare, Before-After Analysis (2015 vs 2024) |
| **Sub-bab 3.6** | Krisis Air Bersih & Toksisitas Cr6+ | Pendekatan Dua Lensa (Mikro Lab vs Makro Panel), Regresi OLS, Uji Chi-Square (χ²) |
| **Sub-bab 3.7** | Eksternalitas Limbah Beracun B3 | Agregasi Timbulan Regional, Profiling Fasilitas Mayor, Analisis Komposisi Residu |

---

## F. Bagan Alur Kerangka Kerja Riset (Research Workflow)

```mermaid
flowchart LR
    subgraph F1["Fase I: Akuisisi Data"]
        A1["Kurasi Data Resmi Terbuka<br/><i>BPS, Kemenkes, KLHK, Lab AEER & WALHI</i>"]
        A2["Panel Provinsi-Tahun<br/><i>6 Provinsi Se-Sulawesi (2014–2024)</i>"]
    end
    subgraph F2["Fase II: Segmentasi & Normalisasi"]
        B1["Segmentasi Spasial Ekstraktif<br/><i>Sentra Tambang vs Non-Sentra Kontrol</i>"]
        B2["Standarisasi per Kapita<br/><i>Rasio Insiden per 10.000 Penduduk</i>"]
    end
    subgraph F3["Fase III: Uji Inferensial & Klinis"]
        C1["Tabel Kontinjensi 2x2<br/><i>Ambang Median IKU/IKA vs Morbiditas</i>"]
        C2["Uji Chi-Square & Toksisitas<br/><i>Benchmark Cr6+ & Odds Ratio Risiko</i>"]
    end
    subgraph F4["Fase IV: Pemodelan & Sintesis"]
        D1["Pemetaan Spasial Dual-Lensa<br/><i>Choropleth ISPA & Bubble Diare</i>"]
        D2["Bukti Kausalitas D3TLH<br/><i>Defisit Faskes & Ancaman Limbah B3</i>"]
    end
    F1 --> F2 --> F3 --> F4
```

> **KERANGKA KELUARAN METODOLOGIS BAB 3:**  
> 1. **Konfigurasi Disparitas Infrastruktur & Morbiditas:** Membuktikan paradoks kawasan sentra hilirisasi yang memikul beban penyakit pernapasan 2x lipat lebih tinggi di tengah defisit fasilitas rumah sakit sekunder.  
> 2. **Konfigurasi Dua Lensa Krisis Sanitasi & Toksisitas:** Memadukan pengujian inferensial data panel makro dengan bukti mikroskopis pelanggaran baku mutu logam berat karsinogenik Kromium Heksavalen (Cr6+) hingga 200 kali lipat di muara tambang.  
> 3. **Konfigurasi Neraca Eksternalitas Limbah Beracun (B3):** Mengkuantifikasi timbulan 32,8 juta ton per tahun limbah slag dan tailing HPAL sebagai ancaman risiko kesehatan masyarakat jangka panjang bagi dokumen D3TLH.
"""

    md_compact = out_dir_compact / "Metodologi_Bab3_Beban_Kesehatan_Compact.md"
    md_bab3    = out_dir_bab3 / "Metodologi_Bab3_Beban_Kesehatan_Compact.md"
    for pth in [md_compact, md_bab3]:
        with open(pth, "w", encoding="utf-8") as f:
            f.write(MD_CONTENT)
    print(f"  [OK] Tersimpan MD  : {md_compact}")
    print(f"  [OK] Salinan MD   : {md_bab3}")

    print("[3/3] Selesai menghasilkan dokumen metodologi Bab 3 versi compact (1-Kolom, 2-3 Halaman).")


if __name__ == "__main__":
    generate_bab3_compact()
