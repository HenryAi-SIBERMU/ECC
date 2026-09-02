#!/usr/bin/env python3
"""
Generator Metodologi Versi Compact Bab 5 — GAYA AKADEMIS TERPADU (CELIOS)
Mengadopsi arsitektur metodologi ringkas terstandarisasi konsisten dengan Bab 1, 2, 3, dan 4:
- FORMAT: 1 KOLOM PENUH (Single Column Layout)
- PANJANG: 2–3 Halaman Maksimal (Elegan, proporsional, tanpa pemadatan berlebihan)
- PENOMORAN SEKSI UTAMA: Huruf kapital A, B, C, D, E, F
- SUB-BAB SEKSI D: 5.1, 5.2, 5.3, 5.4 (Judul persis dokumen induk)
- OPERASIONALISASI INDIKATOR: 9 Indikator Empiris Lengkap (Matriks Indikator & Sumber Data Resmi)
- FORMULASI & TABEL CROSSTAB: Format standar Tabel 1.5b/4.4a/5.4a dengan keterangan definisi variabel lengkap
- KORESPONDENSI METODOLOGI: 3 kolom bersih (Sub-bab, Fokus Kajian Empiris, Metode Analitis Utama)
- FLOWCHART: Mermaid JS horizontal (flowchart LR) dirender tajam ke DOCX (16.5 cm) dan blok kode di MD
- SINKRONISASI: Dual-save ke direktori versicompact/bab_5 dan bab_5.
"""

import os
import sys
import shutil
import base64
import requests
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Palet Warna Resmi CELIOS ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Forest Dark Green (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Celios Accent Green (#2E7D32)
G_LIGHT = RGBColor(0x38, 0x8E, 0x3C)  # Medium Green (#388E3C)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Charcoal Body Text (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Muted Grey Text (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Pure White

# ── Helper Fungsi XML Word ───────────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=80, right=80, top=40, bottom=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)

def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    pPr.append(s)

def para_border_bottom(p, color='1B5E20', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '3')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def para_border_left(p, color='2E7D32', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:left')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='1B5E20', sz='8'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Typography & Content ──────────────────────────────
def add_run(p, text, bold=False, italic=False, pt=8.5, color=C_BODY, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    else:
        r.font.name = 'Calibri'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    return r

def add_h2(doc, prefix, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_border_bottom(p, color='1B5E20', sz='8')
    add_run(p, f"{prefix}.  ", bold=True, pt=11, color=G_DARK)
    add_run(p, title.upper(), bold=True, pt=11, color=G_DARK)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, pt=9.5, color=G_MID)
    return p

def add_body(doc, parts, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        add_run(p, text, bold=bold, italic=italic, pt=8.5, color=C_BODY)
    return p

def add_formula(doc, text, ket=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    para_shd(p, 'EDF7EE')
    add_run(p, text, pt=8, color=G_MID, mono=True)
    if ket:
        add_run(p, f"\nKeterangan: {ket}", italic=True, pt=7.5, color=RGBColor(0x33, 0x55, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_styled(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    bd_subtle = {'val': 'single', 'sz': '4', 'color': 'D0D7DE', 'space': '0'}
    
    # Header Row
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=80, right=80, top=50, bottom=50)
        set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if alignments and j < len(alignments):
            align = alignments[j]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, pt=8, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F9FBF9' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_shd(c, fill)
            cell_margin(c, left=80, right=80, top=40, bottom=40)
            set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            
            if alignments and j < len(alignments):
                align = alignments[j]
                if align == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_run(p, str(val), pt=7.5, color=C_BODY)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after  = Pt(3)
    return tbl

def download_mermaid_png(mermaid_str, filepath):
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        return True
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False


def generate_bab5_compact():
    print("[1/3] Membangun dokumen compact Bab 5 (Format 1-Kolom, 2-3 Halaman)...")
    
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab5    = out_dir_compact.parent.parent / "bab_5"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab5.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER DOKUMEN ──────────────────────────────────────────
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after  = Pt(1)
    add_run(p_t, "RINGKASAN EKSEKUTIF METODOLOGIS", bold=True, pt=8.5, color=G_MID)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after  = Pt(2)
    para_border_bottom(p_h, color='1B5E20', sz='12')
    add_run(p_h, "BAB 5: METODOLOGI ANALISIS POLA PENERBITAN IZIN DI ZONA KRITIS EKOLOGIS", bold=True, pt=15, color=G_DARK)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(1)
    p_meta.paragraph_format.space_after  = Pt(5)
    add_run(p_meta, "Studi Daya Dukung & Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi · ", italic=True, pt=8, color=C_GREY)
    add_run(p_meta, "Center of Economic and Law Studies (CELIOS)", bold=True, italic=True, pt=8, color=G_DARK)

    # ── A. DESAIN PENELITIAN & TUJUAN ───────────────────────────
    add_h2(doc, "A", "Desain Penelitian & Tujuan")
    add_body(doc, [
        ("Penelitian ini menggunakan ", False, False),
        ("desain audit perizinan geospasial-temporal, telaah kepatuhan tata ruang, dan pengujian inferensial kuantitatif terpadu", True, False),
        (" untuk membedah relasi kausal antara eskalasi penerbitan Izin Usaha Pertambangan (IUP) baru dengan kehancuran ekosistem kritis di enam provinsi Pulau Sulawesi sepanjang satu dekade (", False, False),
        ("2014–2024", True, False),
        ("). Tiga tujuan utama metodologis Bab 5 meliputi:", False, False)
    ])
    add_body(doc, [
        ("1. ", True, False), ("Sinkronisasi Waktu & Akselerasi Izin (Timeline Mapping): ", True, False),
        ("Menguji sinkronisasi temporal antara tren deforestasi tahunan dan pelepasan konsesi tambang baru, serta mengukur rasio percepatan izin pra vs pasca penetapan kebijakan hilirisasi.\n", False, False),
        ("2. ", True, False), ("Audit Tabrakan Tata Ruang Spasial (Livelihood Overlay): ", True, False),
        ("Mengisolasi dan menghitung secara kumulatif luas tutupan lahan yang hilang pada kawasan livelihood produktif warga (pertanian, peternakan, perkebunan) akibat penetrasi konsesi pertambangan.\n", False, False),
        ("3. ", True, False), ("Evaluasi Tata Kelola FPIC & Pembuktian Kausalitas Inferensial: ", True, False),
        ("Mendokumentasikan pelanggaran prosedur persetujuan awal (FPIC) serta membuktikan signifikansi hubungan kausal antara penerbitan izin dan deforestasi kritis melalui uji Pearson Chi-Square dan rasio peluang (Odds Ratio).", False, False)
    ])

    # ── B. SUMBER DATA & CAKUPAN WILAYAH ─────────────────────────
    add_h2(doc, "B", "Sumber Data & Cakupan Wilayah")
    add_body(doc, [
        ("Kajian ini mengintegrasikan lima klaster basis data resmi kementerian teknis, platform satelit global, dan registri advokasi masyarakat sipil yang telah divalidasi silang:", False, False)
    ])
    add_body(doc, [
        ("• ", True, False), ("Kementerian ESDM (MODI & MinerbaOne): ", True, False),
        ("Registri 574 IUP baru mencakup luas konsesi 819.452,5 Ha terdistribusi menurut provinsi dan tahun penerbitan (2014–2024).\n", False, False),
        ("• ", True, False), ("Global Forest Watch (GFW / Hansen UMD) & KLHK: ", True, False),
        ("Data time-series deforestasi total (1,38 juta Ha) dan kehilangan tutupan pohon akibat pendorong komoditas tambang/sawit (2014–2023).\n", False, False),
        ("• ", True, False), ("Batas Geospasial Kawasan Livelihood & Penyangga Pangan: ", True, False),
        ("Poligon spasial peruntukan ruang kelola warga mencakup Zona Pertanian-Peternakan dan Perkebunan Rakyat.\n", False, False),
        ("• ", True, False), ("Konsorsium Pembaruan Agraria (CATAHU KPA): ", True, False),
        ("Audit rekam jejak korporasi tambang bermasalah, izin ilegal, dan kasus tumpang tindih kawasan hutan di Sulawesi.\n", False, False),
        ("• ", True, False), ("Koalisi Masyarakat Sipil (JATAM, WALHI, AMAN): ", True, False),
        ("Dokumentasi 12 kasus konflik pertambangan spesifik di tapak industri dan catatan pelanggaran asas FPIC terhadap masyarakat adat/lokal.", False, False)
    ])

    # ── C. OPERASIONALISASI VARIABEL & INDIKATOR RISET ──────────
    add_h2(doc, "C", "Operasionalisasi Variabel & Indikator Riset")
    add_body(doc, [
        ("Seluruh dinamika perizinan tambang, laju deforestasi tutupan hutan, perambahan zona penyangga livelihood, pelanggaran konsultasi FPIC, hingga pengujian korelasi statistik dioperasionalkan secara terstruktur ke dalam ", False, False),
        ("indikator riset empiris", True, False),
        (" sebagaimana dirangkum pada matriks operasional berikut:", False, False)
    ])

    add_caption(doc, "Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 5")
    table_indikator_data = [
        ["1", "Akumulasi Penerbitan IUP Baru", "Frekuensi Izin Usaha Pertambangan Baru Terbit", "Unit Izin", "Data Registry ESDM MODI"],
        ["2", "Luas Alokasi Konsesi Tambang", "Bentang Konsesi Pertambangan Baru", "Hektar (Ha)", "Data Registry ESDM MODI"],
        ["3", "Laju Deforestasi Hutan Alam", "Kehilangan Tutupan Pohon Alami Tahunan", "Hektar (Ha)", "Global Forest Watch (Hansen UMD)"],
        ["4", "Deforestasi Driver Komoditas", "Kehilangan Tutupan Akibat Tambang & Sawit", "Hektar (Ha)", "GFW Commodity Drivers"],
        ["5", "Perambahan Kawasan Livelihood", "Kerusakan Zona Pertanian & Peternakan Warga", "Hektar (Ha)", "GFW Overlay Livelihood Zone"],
        ["6", "Perambahan Perkebunan Rakyat", "Kerusakan Zona Perkebunan Warga Produktif", "Hektar (Ha)", "GFW Overlay Livelihood Zone"],
        ["7", "Insidensi Pelanggaran Asas FPIC", "Konflik Tambang Tanpa Persetujuan Awal Warga", "Kasus", "Koalisi Sipil (JATAM & WALHI)"],
        ["8", "Anomali Legalitas & Tata Kelola Izin", "Pelanggaran Prosedur & Rekam Jejak Korporasi", "Kasus Korporasi", "CATAHU KPA & TanahKita"],
        ["9", "Rasio Peluang Risiko Ekologis (OR)", "Magnitudo Kelipatan Risiko Deforestasi Kritis", "Rasio Peluang (Odds)", "Panel Data Join ESDM-GFW"]
    ]

    add_table_styled(
        doc,
        headers=["No", "Indikator Riset", "Fokus Pengukuran", "Satuan", "Sumber Data Primer Resmi"],
        rows=table_indikator_data,
        col_widths_cm=[0.8, 4.5, 4.5, 2.2, 5.0],
        alignments=['C', 'L', 'L', 'C', 'L']
    )

    # ── D. KERANGKA ANALISIS & FORMULASI MATEMATIS ──────────────
    add_h2(doc, "D", "Kerangka Analisis & Formulasi Matematis")

    # 5.1
    add_h3(doc, "5.1 Fakta Penyebab: Sinkronisasi Waktu (Timeline Mapping)")
    add_body(doc, [
        ("Sinkronisasi waktu memetakan relasi temporal antara lonjakan izin pertambangan dengan eskalasi kehilangan tutupan hutan tahunan, serta menghitung rasio laju akselerasi izin era pra vs pasca-2020:", False, False)
    ])
    add_formula(doc, "Agregasi Tahunan: D_t = Σ D_{p,t}   ;   I_t = Σ I_{p,t}   ;   L_t = Σ L_{p,t}   |   Rasio Akselerasi (R) = I_Pasca / I_Pra",
                ket="D_{p,t} = Deforestasi provinsi p tahun t (Ha); I_{p,t} = IUP terbit tahun t; L_{p,t} = Luas konsesi (Ha); I_Pasca = Total izin pasca-2020 (468 IUP); I_Pra = Total izin pra-2020 (106 IUP); R = Rasio lonjakan akselerasi izin (4,4 kali lipat).")

    # 5.2
    add_h3(doc, "5.2 Fakta Spasial: Tabrakan Tata Ruang di Kawasan Konservasi & Livelihood")
    add_body(doc, [
        ("Penapisan spasial (spatial overlay intersection) mengisolasi poligon tree cover loss yang beririsan dengan kawasan livelihood produktif warga dan menghitung laju kerusakan kumulatif antar-kategori:", False, False)
    ])
    add_formula(doc, "Kehancuran Tahunan: H_c(t) = Σ Loss_i   ;   Akumulasi: K_c(T) = Σ H_c(t)   ;   Total Kumulatif(T) = K_Tani(T) + K_Kebun(T)",
                ket="Loss_i = Luas tutupan hilang pada poligon livelihood i (Ha); c = Kategori livelihood (1 = Pertanian/Peternakan, 2 = Perkebunan); K_c(T) = Akumulasi tutupan hilang s.d. tahun T; Total Kumulatif = Total kerusakan ruang pangan (41.785,1 Ha).")

    # 5.3
    add_h3(doc, "5.3 Realitas Lapangan: Izin Bermasalah, FPIC Diabaikan, Masyarakat Dikorbankan")
    add_body(doc, [
        ("Integrasi data lintas registri (cross-dataset audit) mengukur proporsi kasus konflik pertambangan yang secara eksplisit mencatatkan indikasi pengabaian hak persetujuan awal masyarakat (FPIC):", False, False)
    ])
    add_formula(doc, "Total Konflik = Σ K_i   ;   Pelanggaran FPIC = Σ K_{i,FPIC=True}   ;   Rasio Pengabaian (%) = [ Pelanggaran FPIC / Total Konflik ] × 100",
                ket="K_i = Kasus sengketa pertambangan di Sulawesi (N = 12); Pelanggaran FPIC = Kasus sengketa dengan indikasi pelanggaran FPIC (N = 8); Rasio Pengabaian = Tingkat pengabaian persetujuan awal masyarakat adat/lokal (66,7%).")

    # 5.4
    add_h3(doc, "5.4 Pembuktian Empiris: Uji Statistik Korelasi Penerbitan Izin & Deforestasi")
    add_body(doc, [
        ("Pengujian statistik inferensial non-parametrik Pearson Chi-Square independensi (df = 1, α = 5%) diterapkan pada matriks kontinjensi 2×2 berbasis ambang median data panel provinsi-tahun (N = 60 observasi: 6 provinsi × 10 tahun). Rasio peluang Odds Ratio (OR) mengukur magnitudo kelipatan risiko deforestasi kritis pada wilayah/tahun dengan penerbitan izin tinggi:", False, False)
    ])
    add_formula(doc, "Kategori(X) = Tinggi jika X ≥ Median ; Rendah jika X < Median   |   χ² = Σ [ (O_ij - E_ij)² / E_ij ]   |   Odds Ratio (OR) = (a × d) / (b × c)",
                ket="X = Nilai observasi panel provinsi-tahun; Median = Ambang batas klasifikasi biner distribusi panel; O_ij & E_ij = Frekuensi teramati dan ekspektasi pada sel ij; a, b, c, d = Sel kontinjensi 2×2; OR = Rasio kelipatan peluang risiko deforestasi.")

    add_caption(doc, "Tabel 5.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 5.4)")
    tabel_5_4a_rows = [
        ["Variabel Independen (X)", "Jumlah Izin Baru (IUP) / Total Luas Konsesi Baru (Ha) per provinsi-tahun."],
        ["Variabel Dependen (Y)", "Total Deforestasi Alam (Ha) / Deforestasi Komoditas Tambang & Sawit (Ha)."],
        ["Hipotesis Nol (H0)", "Tidak terdapat hubungan signifikan antara tingginya penerbitan IUP baru dan tingginya laju deforestasi."],
        ["Hipotesis Alternatif (H1)", "Tingginya penerbitan IUP baru berasosiasi signifikan dengan peningkatan risiko laju deforestasi kritis."],
        ["Decision Rule (Alpha 5%)", "Tolak H0 jika Pearson Chi-Square P-Value < 0.05 dan Odds Ratio (OR) > 1.0."],
        ["Threshold Kategori (Median Panel)", "Median Jumlah IUP = 2,0 izin/tahun; Median Luas Konsesi = 2.011,5 Ha; Median Deforestasi Total = 15.917,7 Ha; Median Deforestasi Komoditas = 10.961,8 Ha (N = 60 observasi)."],
        ["Orientasi Odds Ratio (OR)", "OR = (a × d) / (b × c) dengan a = Kuadran Izin Tinggi & Deforestasi Tinggi; membuktikan kelipatan risiko kehancuran hutan (OR terhitung berkisar 9,04 s.d. 16,00 kali lipat)."]
    ]
    add_table_styled(
        doc,
        headers=["Komponen Uji", "Definisi Variabel (Sub-bab 5.4)"],
        rows=tabel_5_4a_rows,
        col_widths_cm=[4.5, 12.5],
        alignments=['L', 'L']
    )

    # ── E. KORESPONDENSI METODOLOGI TERHADAP SUB-BAB LAPORAN ────
    add_h2(doc, "E", "Korespondensi Metodologi terhadap Sub-bab Laporan Bab 5")
    add_body(doc, [
        ("Setiap sub-bab analitis pada Bab 5 ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:", False, False)
    ])

    table_korespondensi = [
        ["Sub-bab 5.1", "Akselerasi & Sinkronisasi Waktu Izin", "Timeline Alignment, Multi-Axis Combo Analysis, Rasio Akselerasi Pra vs Pasca 2020"],
        ["Sub-bab 5.2", "Tabrakan Tata Ruang Kawasan Livelihood", "Geospatial Intersection Overlay, Akumulasi Kerusakan Livelihood Zone (Pertanian & Perkebunan)"],
        ["Sub-bab 5.3", "Anomali Tata Kelola & Pengabaian FPIC", "Cross-Dataset Integration, Case Tracking Pelanggaran FPIC & Rekam Jejak Korporasi CATAHU"],
        ["Sub-bab 5.4", "Pembuktian Korelasi Kausalitas Spasial", "Panel Data Crosstabulation (N=60), Median Binning, Pearson Chi-Square, Odds Ratio Analysis"]
    ]

    add_table_styled(
        doc,
        headers=["Sub-bab", "Fokus Kajian Empiris", "Metode Analitis Utama"],
        rows=table_korespondensi,
        col_widths_cm=[2.5, 5.5, 9.0],
        alignments=['C', 'L', 'L']
    )

    # ── F. BAGAN ALUR KERANGKA KERJA RISET BAB 5 ────────────────
    add_h2(doc, "F", "Bagan Alur Kerangka Kerja Riset (Research Workflow)")
    add_body(doc, [
        ("Kerangka operasional metodologi Bab 5 berjalan secara terpadu melalui empat fase berurutan sebagaimana divisualisasikan pada bagan alur kerja riset berikut:", False, False)
    ])

    mermaid_str_f = """flowchart LR
    subgraph F1["Fase I: Kurasi Data"]
        A1["Registry IUP ESDM MODI<br/><i>574 Izin & 819.452 Ha</i>"]
        A2["GFW Tree Cover Loss<br/><i>Deforestasi Alam & Komoditas</i>"]
        A3["Livelihood & CATAHU KPA<br/><i>Kawasan Pangan & Audit FPIC</i>"]
    end
    subgraph F2["Fase II: Harmonisasi & Overlay"]
        B1["Timeline Alignment 2014-2024<br/><i>Sinkronisasi Izin vs Krisis</i>"]
        B2["Spatial Overlay Intersection<br/><i>Isolasi Zona Livelihood</i>"]
        B3["Panel Join Provinsi-Tahun<br/><i>6 Provinsi × 10 Tahun (N=60)</i>"]
    end
    subgraph F3["Fase III: Komputasi & Inferensi"]
        C1["Rasio Akselerasi Izin (4.4x)<br/><i>Pra vs Pasca Hilirisasi 2020</i>"]
        C2["Kumulatif Kerusakan Ruang<br/><i>Pertanian & Perkebunan Warga</i>"]
        C3["Uji Pearson Chi-Square<br/><i>Signifikansi & Odds Ratio (OR)</i>"]
    end
    subgraph F4["Fase IV: Sintesis Temuan"]
        D1["Sinkronisasi Krisis-Izin<br/><i>86.8% Izin di Atas Median</i>"]
        D2["Tabrakan Tata Ruang Livelihood<br/><i>41.8 Ribu Ha Terdegradasi</i>"]
        D3["Kausalitas Deforestasi Kritis<br/><i>4/4 Skenario Signifikan p<0.001</i>"]
    end
    F1 --> F2 --> F3 --> F4"""

    png_workflow_path = str(out_dir_compact / "mermaid_workflow_bab5.png")
    is_downloaded = download_mermaid_png(mermaid_str_f, png_workflow_path)

    add_caption(doc, "Bagan Alur 5.1: Alur Logika Kerangka Kerja Riset Bab 5 (Research Workflow)")
    if is_downloaded and os.path.exists(png_workflow_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(3)
        p_img.paragraph_format.space_after  = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(png_workflow_path, width=Cm(16.5))
        try:
            shutil.copyfile(png_workflow_path, str(out_dir_bab5 / "mermaid_workflow_bab5.png"))
        except Exception:
            pass

    # Box Output Kesimpulan
    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after  = Pt(4)
    all_border_para(p_box, color='1B5E20', sz='8')
    para_shd(p_box, 'F1F8E9')
    add_run(p_box, "KERANGKA KELUARAN METODOLOGIS BAB 5:\n", bold=True, pt=8.5, color=G_DARK)
    add_run(p_box, "1. Konfigurasi Sinkronisasi Krisis & Akselerasi Izin: Menunjukkan bahwa 86,8% izin tambang terbit pada tahun-tahun deforestasi provinsi di atas median historis, dengan lonjakan akselerasi izin era pasca-2020 mencapai 4,4 kali lipat (468 izin vs 106 izin pra-2020).\n"
                   "2. Konfigurasi Tabrakan Tata Ruang Livelihood: Mengkuantifikasi kerusakan permanen seluas lebih dari 41,8 ribu hektar kawasan penyangga livelihood pangan masyarakat (57,7% Pertanian-Peternakan dan 42,3% Perkebunan Rakyat) akibat penetrasi izin konsesi ekstraktif.\n"
                   "3. Konfigurasi Pembuktian Kausalitas Inferensial: Membuktikan secara matematis melalui pengujian Chi-Square bahwa seluruh 4 skenario perizinan vs deforestasi terbukti signifikan (p < 0,001) dengan magnitudo risiko kerusakan ekologis (Odds Ratio) hingga 16,0 kali lipat.",
            pt=8, color=C_BODY)

    # ── SIMPAN DOKUMEN DOCX (DUAL SAVE) ─────────────────────────
    docx_compact = out_dir_compact / "Metodologi_Bab5_Pola_Perizinan_Compact.docx"
    docx_bab5    = out_dir_bab5 / "Metodologi_Bab5_Pola_Perizinan_Compact.docx"
    
    doc.save(str(docx_compact))
    shutil.copyfile(docx_compact, docx_bab5)
    print(f"  [OK] Tersimpan DOCX: {docx_compact}")
    print(f"  [OK] Salinan DOCX : {docx_bab5}")

    # ── GENERATE MARKDOWN PADANAN ───────────────────────────────
    print("[2/3] Membangun dokumen Markdown padanan...")
    MD_CONTENT = """# METODOLOGI PENELITIAN: BAB 5 — POLA PENERBITAN IZIN DI ZONA KRITIS EKOLOGIS
*CELIOS (Center of Economic and Law Studies) · Audit Spasial-Statistik D3TLH Sulawesi (2014–2024) · Ringkasan Eksekutif Metodologis*

---

## A. Desain Penelitian & Tujuan
Penelitian ini menggunakan **desain audit perizinan geospasial-temporal, telaah kepatuhan tata ruang, dan pengujian inferensial kuantitatif terpadu** untuk membedah relasi kausal antara eskalasi penerbitan Izin Usaha Pertambangan (IUP) baru dengan kehancuran ekosistem kritis di enam provinsi Pulau Sulawesi sepanjang satu dekade (**2014–2024**). Tiga tujuan utama metodologis Bab 5 meliputi:

1. **Sinkronisasi Waktu & Akselerasi Izin (Timeline Mapping):** Menguji sinkronisasi temporal antara tren deforestasi tahunan dan pelepasan konsesi tambang baru, serta mengukur rasio percepatan izin pra vs pasca penetapan kebijakan hilirisasi.
2. **Audit Tabrakan Tata Ruang Spasial (Livelihood Overlay):** Mengisolasi dan menghitung secara kumulatif luas tutupan lahan yang hilang pada kawasan livelihood produktif warga (pertanian, peternakan, perkebunan) akibat penetrasi konsesi pertambangan.
3. **Evaluasi Tata Kelola FPIC & Pembuktian Kausalitas Inferensial:** Mendokumentasikan pelanggaran prosedur persetujuan awal (FPIC) serta membuktikan signifikansi hubungan kausal antara penerbitan izin dan deforestasi kritis melalui uji Pearson Chi-Square dan rasio peluang (Odds Ratio).

---

## B. Sumber Data & Cakupan Wilayah
Kajian ini mengintegrasikan lima klaster basis data resmi kementerian teknis, platform satelit global, dan registri advokasi masyarakat sipil yang telah divalidasi silang:

- **Kementerian ESDM (MODI & MinerbaOne):** Registri 574 IUP baru mencakup luas konsesi 819.452,5 Ha terdistribusi menurut provinsi dan tahun penerbitan (2014–2024).
- **Global Forest Watch (GFW / Hansen UMD) & KLHK:** Data time-series deforestasi total (1,38 juta Ha) dan kehilangan tutupan pohon akibat pendorong komoditas tambang/sawit (2014–2023).
- **Batas Geospasial Kawasan Livelihood & Penyangga Pangan:** Poligon spasial peruntukan ruang kelola warga mencakup Zona Pertanian-Peternakan dan Perkebunan Rakyat.
- **Konsorsium Pembaruan Agraria (CATAHU KPA):** Audit rekam jejak korporasi tambang bermasalah, izin ilegal, dan kasus tumpang tindih kawasan hutan di Sulawesi.
- **Koalisi Masyarakat Sipil (JATAM, WALHI, AMAN):** Dokumentasi 12 kasus konflik pertambangan spesifik di tapak industri dan catatan pelanggaran asas FPIC terhadap masyarakat adat/lokal.

---

## C. Operasionalisasi Variabel & Indikator Riset
Seluruh dinamika perizinan tambang, laju deforestasi tutupan hutan, perambahan zona penyangga livelihood, pelanggaran konsultasi FPIC, hingga pengujian korelasi statistik dioperasionalkan secara terstruktur ke dalam **indikator riset empiris** sebagaimana dirangkum pada matriks operasional berikut:

##### Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 5
| No | Indikator Riset | Fokus Pengukuran | Satuan | Sumber Data Primer Resmi |
| :-: | :--- | :--- | :-: | :--- |
| 1 | Akumulasi Penerbitan IUP Baru | Frekuensi Izin Usaha Pertambangan Baru Terbit | Unit Izin | Data Registry ESDM MODI |
| 2 | Luas Alokasi Konsesi Tambang | Bentang Konsesi Pertambangan Baru | Hektar (Ha) | Data Registry ESDM MODI |
| 3 | Laju Deforestasi Hutan Alam | Kehilangan Tutupan Pohon Alami Tahunan | Hektar (Ha) | Global Forest Watch (Hansen UMD) |
| 4 | Deforestasi Driver Komoditas | Kehilangan Tutupan Akibat Tambang & Sawit | Hektar (Ha) | GFW Commodity Drivers |
| 5 | Perambahan Kawasan Livelihood | Kerusakan Zona Pertanian & Peternakan Warga | Hektar (Ha) | GFW Overlay Livelihood Zone |
| 6 | Perambahan Perkebunan Rakyat | Kerusakan Zona Perkebunan Warga Produktif | Hektar (Ha) | GFW Overlay Livelihood Zone |
| 7 | Insidensi Pelanggaran Asas FPIC | Konflik Tambang Tanpa Persetujuan Awal Warga | Kasus | Koalisi Sipil (JATAM & WALHI) |
| 8 | Anomali Legalitas & Tata Kelola Izin | Pelanggaran Prosedur & Rekam Jejak Korporasi | Kasus Korporasi | CATAHU KPA & TanahKita |
| 9 | Rasio Peluang Risiko Ekologis (OR) | Magnitudo Kelipatan Risiko Deforestasi Kritis | Rasio Peluang (Odds) | Panel Data Join ESDM-GFW |

---

## D. Kerangka Analisis & Formulasi Matematis

### 5.1 Fakta Penyebab: Sinkronisasi Waktu (Timeline Mapping)
Sinkronisasi waktu memetakan relasi temporal antara lonjakan izin pertambangan dengan eskalasi kehilangan tutupan hutan tahunan, serta menghitung rasio laju akselerasi izin era pra vs pasca-2020:

> `Agregasi Tahunan: D_t = Σ D_{p,t}   ;   I_t = Σ I_{p,t}   ;   L_t = Σ L_{p,t}   |   Rasio Akselerasi (R) = I_Pasca / I_Pra`  
> *Keterangan: D_{p,t} = Deforestasi provinsi p tahun t (Ha); I_{p,t} = IUP terbit tahun t; L_{p,t} = Luas konsesi (Ha); I_Pasca = Total izin pasca-2020 (468 IUP); I_Pra = Total izin pra-2020 (106 IUP); R = Rasio lonjakan akselerasi izin (4,4 kali lipat).*

### 5.2 Fakta Spasial: Tabrakan Tata Ruang di Kawasan Konservasi & Livelihood
Penapisan spasial (spatial overlay intersection) mengisolasi poligon tree cover loss yang beririsan dengan kawasan livelihood produktif warga dan menghitung laju kerusakan kumulatif antar-kategori:

> `Kehancuran Tahunan: H_c(t) = Σ Loss_i   ;   Akumulasi: K_c(T) = Σ H_c(t)   ;   Total Kumulatif(T) = K_Tani(T) + K_Kebun(T)`  
> *Keterangan: Loss_i = Luas tutupan hilang pada poligon livelihood i (Ha); c = Kategori livelihood (1 = Pertanian/Peternakan, 2 = Perkebunan); K_c(T) = Akumulasi tutupan hilang s.d. tahun T; Total Kumulatif = Total kerusakan ruang pangan (41.785,1 Ha).*

### 5.3 Realitas Lapangan: Izin Bermasalah, FPIC Diabaikan, Masyarakat Dikorbankan
Integrasi data lintas registri (cross-dataset audit) mengukur proporsi kasus konflik pertambangan yang secara eksplisit mencatatkan indikasi pengabaian hak persetujuan awal masyarakat (FPIC):

> `Total Konflik = Σ K_i   ;   Pelanggaran FPIC = Σ K_{i,FPIC=True}   ;   Rasio Pengabaian (%) = [ Pelanggaran FPIC / Total Konflik ] × 100`  
> *Keterangan: K_i = Kasus sengketa pertambangan di Sulawesi (N = 12); Pelanggaran FPIC = Kasus sengketa dengan indikasi pelanggaran FPIC (N = 8); Rasio Pengabaian = Tingkat pengabaian persetujuan awal masyarakat adat/lokal (66,7%).*

### 5.4 Pembuktian Empiris: Uji Statistik Korelasi Penerbitan Izin & Deforestasi
Pengujian statistik inferensial non-parametrik Pearson Chi-Square independensi (df = 1, α = 5%) diterapkan pada matriks kontinjensi 2×2 berbasis ambang median data panel provinsi-tahun (N = 60 observasi: 6 provinsi × 10 tahun). Rasio peluang Odds Ratio (OR) mengukur magnitudo kelipatan risiko deforestasi kritis pada wilayah/tahun dengan penerbitan izin tinggi:

> `Kategori(X) = Tinggi jika X ≥ Median ; Rendah jika X < Median   |   χ² = Σ [ (O_ij - E_ij)² / E_ij ]   |   Odds Ratio (OR) = (a × d) / (b × c)`  
> *Keterangan: X = Nilai observasi panel provinsi-tahun; Median = Ambang batas klasifikasi biner distribusi panel; O_ij & E_ij = Frekuensi teramati dan ekspektasi pada sel ij; a, b, c, d = Sel kontinjensi 2×2; OR = Rasio kelipatan peluang risiko deforestasi.*

##### Tabel 5.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 5.4)
| Komponen Uji | Definisi Variabel (Sub-bab 5.4) |
| :--- | :--- |
| Variabel Independen (X) | Jumlah Izin Baru (IUP) / Total Luas Konsesi Baru (Ha) per provinsi-tahun. |
| Variabel Dependen (Y) | Total Deforestasi Alam (Ha) / Deforestasi Komoditas Tambang & Sawit (Ha). |
| Hipotesis Nol (H0) | Tidak terdapat hubungan signifikan antara tingginya penerbitan IUP baru dan tingginya laju deforestasi. |
| Hipotesis Alternatif (H1) | Tingginya penerbitan IUP baru berasosiasi signifikan dengan peningkatan risiko laju deforestasi kritis. |
| Decision Rule (Alpha 5%) | Tolak H0 jika Pearson Chi-Square P-Value < 0.05 dan Odds Ratio (OR) > 1.0. |
| Threshold Kategori (Median Panel) | Median Jumlah IUP = 2,0 izin/tahun; Median Luas Konsesi = 2.011,5 Ha; Median Deforestasi Total = 15.917,7 Ha; Median Deforestasi Komoditas = 10.961,8 Ha (N = 60 observasi). |
| Orientasi Odds Ratio (OR) | OR = (a × d) / (b × c) dengan a = Kuadran Izin Tinggi & Deforestasi Tinggi; membuktikan kelipatan risiko kehancuran hutan (OR terhitung berkisar 9,04 s.d. 16,00 kali lipat). |

---

## E. Korespondensi Metodologi terhadap Sub-bab Laporan Bab 5
Setiap sub-bab analitis pada Bab 5 ditopang oleh metode kuantitatif yang terukur dan menghasilkan sintesis bukti empiris terstandarisasi sebagaimana dirangkum pada matriks berikut:

##### Matriks Korespondensi Metodologis Bab 5
| Sub-bab | Fokus Kajian Empiris | Metode Analitis Utama |
| :-: | :--- | :--- |
| Sub-bab 5.1 | Akselerasi & Sinkronisasi Waktu Izin | Timeline Alignment, Multi-Axis Combo Analysis, Rasio Akselerasi Pra vs Pasca 2020 |
| Sub-bab 5.2 | Tabrakan Tata Ruang Kawasan Livelihood | Geospatial Intersection Overlay, Akumulasi Kerusakan Livelihood Zone (Pertanian & Perkebunan) |
| Sub-bab 5.3 | Anomali Tata Kelola & Pengabaian FPIC | Cross-Dataset Integration, Case Tracking Pelanggaran FPIC & Rekam Jejak Korporasi CATAHU |
| Sub-bab 5.4 | Pembuktian Korelasi Kausalitas Spasial | Panel Data Crosstabulation (N=60), Median Binning, Pearson Chi-Square, Odds Ratio Analysis |

---

## F. Bagan Alur Kerangka Kerja Riset (Research Workflow)
Kerangka operasional metodologi Bab 5 berjalan secara terpadu melalui empat fase berurutan sebagaimana divisualisasikan pada bagan alur kerja riset berikut:

```mermaid
flowchart LR
    subgraph F1["Fase I: Kurasi Data"]
        A1["Registry IUP ESDM MODI<br/><i>574 Izin & 819.452 Ha</i>"]
        A2["GFW Tree Cover Loss<br/><i>Deforestasi Alam & Komoditas</i>"]
        A3["Livelihood & CATAHU KPA<br/><i>Kawasan Pangan & Audit FPIC</i>"]
    end
    subgraph F2["Fase II: Harmonisasi & Overlay"]
        B1["Timeline Alignment 2014-2024<br/><i>Sinkronisasi Izin vs Krisis</i>"]
        B2["Spatial Overlay Intersection<br/><i>Isolasi Zona Livelihood</i>"]
        B3["Panel Join Provinsi-Tahun<br/><i>6 Provinsi × 10 Tahun (N=60)</i>"]
    end
    subgraph F3["Fase III: Komputasi & Inferensi"]
        C1["Rasio Akselerasi Izin (4.4x)<br/><i>Pra vs Pasca Hilirisasi 2020</i>"]
        C2["Kumulatif Kerusakan Ruang<br/><i>Pertanian & Perkebunan Warga</i>"]
        C3["Uji Pearson Chi-Square<br/><i>Signifikansi & Odds Ratio (OR)</i>"]
    end
    subgraph F4["Fase IV: Sintesis Temuan"]
        D1["Sinkronisasi Krisis-Izin<br/><i>86.8% Izin di Atas Median</i>"]
        D2["Tabrakan Tata Ruang Livelihood<br/><i>41.8 Ribu Ha Terdegradasi</i>"]
        D3["Kausalitas Deforestasi Kritis<br/><i>4/4 Skenario Signifikan p<0.001</i>"]
    end
    F1 --> F2 --> F3 --> F4
```

> **KERANGKA KELUARAN METODOLOGIS BAB 5:**  
> 1. **Konfigurasi Sinkronisasi Krisis & Akselerasi Izin:** Menunjukkan bahwa 86,8% izin tambang terbit pada tahun-tahun deforestasi provinsi di atas median historis, dengan lonjakan akselerasi izin era pasca-2020 mencapai 4,4 kali lipat (468 izin vs 106 izin pra-2020).  
> 2. **Konfigurasi Tabrakan Tata Ruang Livelihood:** Mengkuantifikasi kerusakan permanen seluas lebih dari 41,8 ribu hektar kawasan penyangga livelihood pangan masyarakat (57,7% Pertanian-Peternakan dan 42,3% Perkebunan Rakyat) akibat penetrasi izin konsesi ekstraktif.  
> 3. **Konfigurasi Pembuktian Kausalitas Inferensial:** Membuktikan secara matematis melalui pengujian Chi-Square bahwa seluruh 4 skenario perizinan vs deforestasi terbukti signifikan (p < 0,001) dengan magnitudo risiko kerusakan ekologis (Odds Ratio) hingga 16,0 kali lipat.
"""

    md_compact = out_dir_compact / "Metodologi_Bab5_Pola_Perizinan_Compact.md"
    md_bab5    = out_dir_bab5 / "Metodologi_Bab5_Pola_Perizinan_Compact.md"
    with open(md_compact, 'w', encoding='utf-8') as f:
        f.write(MD_CONTENT)
    shutil.copyfile(md_compact, md_bab5)
    print(f"  [OK] Tersimpan MD  : {md_compact}")
    print(f"  [OK] Salinan MD   : {md_bab5}")

    print("[3/3] Selesai menghasilkan dokumen metodologi Bab 5 versi compact (1-Kolom, 2-3 Halaman).\n")


if __name__ == "__main__":
    generate_bab5_compact()
