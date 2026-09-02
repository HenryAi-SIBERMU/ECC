#!/usr/bin/env python3
"""
Generator Metodologi Versi Compact Bab 9 — GAYA AKADEMIS TERPADU (CELIOS)
Mengadopsi arsitektur metodologi ringkas terstandarisasi konsisten dengan Bab 1 s.d. 8:
- RUANG LINGKUP: Bab 9 — Demografi Sosial: Ketika Hilirisasi Mengubah Struktur Masyarakat
- FORMAT: 1 KOLOM PENUH (Single Column Layout)
- PANJANG: 2–3 Halaman Maksimal (Elegan, proporsional, tanpa pemadatan berlebihan)
- PENOMORAN SEKSI UTAMA: Huruf kapital A, B, C, D, E, F
- SUB-BAB SEKSI D: Sub-bab 9.1, 9.2, 9.3 sesuai dokumen induk
- OPERASIONALISASI INDIKATOR: 9 Indikator Riset Empiris Kunci Terverifikasi (5 Kolom Baku tanpa kolom Periode)
- NOTASI MATEMATIKA: Bahasa intuitif dan ramah pembaca awam dengan penjelasan penalaran logis
- KORESPONDENSI METODOLOGI: 3 kolom bersih (Sub-bab, Fokus Kajian Empiris, Metode Analitis Utama)
- FLOWCHART: Mermaid JS horizontal (flowchart LR) dirender tajam ke DOCX (16.5 cm) dan blok kode di MD
- SINKRONISASI: Dual-save ke direktori versicompact/bab_9 dan bab_9.
"""

import os
import sys
import shutil
import base64
import requests
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Palet Warna Resmi CELIOS ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Forest Dark Green (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Celios Accent Green (#2E7D32)
G_LIGHT = RGBColor(0x38, 0x8E, 0x3C)  # Medium Green (#388E3C)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Charcoal Body Text (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Muted Grey Text (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Pure White

# ── Helper Fungsi XML Word ───────────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=80, right=80, top=40, bottom=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)

def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    pPr.append(s)

def para_border_bottom(p, color='1B5E20', sz='12'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '3')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='1B5E20', sz='8'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Typography & Content ──────────────────────────────
def add_run(p, text, bold=False, italic=False, pt=8.5, color=C_BODY, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    else:
        r.font.name = 'Calibri'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    return r

def add_h2(doc, prefix, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_border_bottom(p, color='1B5E20', sz='8')
    add_run(p, f"{prefix}.  ", bold=True, pt=11, color=G_DARK)
    add_run(p, title.upper(), bold=True, pt=11, color=G_DARK)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, bold=True, pt=9.5, color=G_MID)
    return p

def add_body(doc, parts, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        add_run(p, text, bold=bold, italic=italic, pt=8.5, color=C_BODY)
    return p

def add_formula(doc, text, ket=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    para_shd(p, 'EDF7EE')
    add_run(p, text, pt=8, color=G_MID, mono=True)
    if ket:
        add_run(p, f"\nKeterangan: {ket}", italic=True, pt=7.5, color=RGBColor(0x33, 0x55, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_styled(doc, headers, rows, col_widths_cm, alignments=None, font_pt=7):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    bd_subtle = {'val': 'single', 'sz': '4', 'color': 'D0D7DE', 'space': '0'}
    
    # Header Row
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=70, right=70, top=45, bottom=45)
        set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if alignments and j < len(alignments):
            align = alignments[j]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, pt=7.5, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F9FBF9' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_shd(c, fill)
            cell_margin(c, left=70, right=70, top=35, bottom=35)
            set_cell_borders(c, top=bd_subtle, left=bd_subtle, bottom=bd_subtle, right=bd_subtle)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            
            if alignments and j < len(alignments):
                align = alignments[j]
                if align == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            add_run(p, str(val), pt=font_pt, color=C_BODY)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after  = Pt(3)
    return tbl

def download_mermaid_png(mermaid_str, filepath):
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        return True
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False


def generate_bab9_compact():
    print("[1/3] Membangun dokumen compact Bab 9 (Format 1-Kolom, 2-3 Halaman)...")
    
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab9    = out_dir_compact.parent.parent / "bab_9"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab9.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER DOKUMEN ──────────────────────────────────────────
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after  = Pt(1)
    add_run(p_t, "RINGKASAN EKSEKUTIF METODOLOGIS", bold=True, pt=8.5, color=G_MID)

    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after  = Pt(2)
    para_border_bottom(p_h, color='1B5E20', sz='12')
    add_run(p_h, "BAB 9: DEMOGRAFI SOSIAL — DISRUPSI STRUKTUR MASYARAKAT", bold=True, pt=15, color=G_DARK)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(1)
    p_meta.paragraph_format.space_after  = Pt(5)
    add_run(p_meta, "Dampak Hilirisasi Nikel: Tekanan Demografi, Pemadatan Ruang, dan Pergeseran Sektor PDRB · ", italic=True, pt=8, color=C_GREY)
    add_run(p_meta, "Center of Economic and Law Studies (CELIOS)", bold=True, italic=True, pt=8, color=G_DARK)

    # ── A. DESAIN PENELITIAN & TUJUAN ───────────────────────────
    add_h2(doc, "A", "Desain Penelitian & Tujuan")
    add_body(doc, [
        ("Penelitian Bab 9 menerapkan ", False, False),
        ("desain analisis demografi spasial dan transformasi struktural ekonomi (Spatial Demography & Structural Transformation Analysis)", True, False),
        (" guna menguji disrupsi sosial yang terjadi akibat penetrasi industri hilirisasi nikel skala masif di Pulau Sulawesi. Melalui pembacaan data deret waktu populasi, pemodelan sebaran kuantil, dan rasio pergeseran sektoral PDRB, kajian ini membuktikan tiga fenomena perubahan sosial-spasial:", False, False)
    ])
    add_body(doc, [
        ("1. ", True, False), ("Tekanan Demografi & Proxy Migrasi (Hazen Quantile Boxplot Analysis): ", True, False),
        ("Menganalisis anomali lonjakan penduduk tahunan (YoY) pada 7 kabupaten prioritas smelter dibandingkan kabupaten non-smelter, guna membuktikan fenomena tarikan migrasi tenaga kerja dan siklus fluktuasi tajam (*boom and bust*).\n", False, False),
        ("2. ", True, False), ("Intensifikasi Ruang & Beban Layanan Publik (Comparative Density Analysis): ", True, False),
        ("Mengukur laju pemadatan penduduk per kilometer persegi pada kawasan industri ekstraktif yang semula berpenduduk jarang, sebagai indikator stres daya dukung sarana air bersih, sanitasi, dan perumahan lokal.\n", False, False),
        ("3. ", True, False), ("Pergeseran Gravitasi Ekonomi Sektoral (PDRB Sector Shift Index): ", True, False),
        ("Mengkuantifikasi transformasi struktur perekonomian daerah dari basis agraris (Sektor A: Pertanian, Kehutanan, Perikanan) menuju dominasi blok ekstraktif-industrial (Sektor B: Pertambangan dan C: Industri Pengolahan).", False, False)
    ])

    # ── B. SUMBER DATA & CAKUPAN WILAYAH ─────────────────────────
    add_h2(doc, "B", "Sumber Data & Cakupan Wilayah")
    add_body(doc, [
        ("Analisis demografi sosial ini mengolah basis data panel resmi Badan Pusat Statistik (BPS) kurun waktu 2014–2024 yang mencakup seluruh wilayah kabupaten/kota dan 6 provinsi se-Pulau Sulawesi:", False, False)
    ])
    add_body(doc, [
        ("• ", True, False), ("BPS SIMDASI (Sistem Informasi Rujukan Statistik Terintegrasi): ", True, False),
        ("Data deret waktu populasi penduduk kabupaten/kota, luas daratan yurisdiksi, dan laju pertumbuhan penduduk YoY.\n", False, False),
        ("• ", True, False), ("Klasifikasi 7 Kabupaten Prioritas Smelter (Fase 4): ", True, False),
        ("Klaster kabupaten sentra industri pengolahan nikel: Banggai, Kolaka, Konawe, Konawe Utara, Luwu Timur, Morowali, dan Morowali Utara (total populasi 2024 mencapai 1,59 juta jiwa).\n", False, False),
        ("• ", True, False), ("BPS PDRB Sektoral Seri 2010 (Tahun 2014–2024): ", True, False),
        ("Struktur Produk Domestik Regional Bruto menurut lapangan usaha: Sektor A (Pertanian), Sektor B (Pertambangan), dan Sektor C (Industri Pengolahan).\n", False, False),
        ("• ", True, False), ("Statistik Perikanan Tangkap BPS: ", True, False),
        ("Dekomposisi estimasi kontribusi sub-sektor perikanan tangkap laut (~22% dari Sektor A) pada provinsi-provinsi pesisir Sulawesi.", False, False)
    ])

    # ── C. OPERASIONALISASI VARIABEL & INDIKATOR RISET ──────────
    add_h2(doc, "C", "Operasionalisasi Variabel & Indikator Riset")
    add_body(doc, [
        ("Seluruh parameter demografi, kepadatan spasial, hingga pergeseran struktur produksi dioperasionalkan ke dalam ", False, False),
        ("9 indikator riset empiris terverifikasi", True, False),
        (" sebagaimana dirangkum pada matriks operasional berikut:", False, False)
    ])

    add_caption(doc, "Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 9 (Demografi Sosial)")
    table_indikator_data = [
        ["1", "Laju Pertumbuhan Penduduk YoY (9.1)", "Pertumbuhan Tahunan Populasi Kabupaten (Proxy Migrasi)", "Persen (% / Tahun)", "BPS SIMDASI Demografi"],
        ["2", "Anatomi Sebaran Boxplot Hazen (9.1)", "Median, Q1, Q3, IQR, dan Rentang Pagar Kewajaran", "Nilai Persentil (%)", "BPS SIMDASI (Kuantil Hazen)"],
        ["3", "Anomali Fluktuasi Boom-and-Bust (9.1)", "Lonjakan Ekstrem Tertinggi vs Kejatuhan Terendah", "Persen Ekstrem (%)", "Dataset Deret Waktu BPS"],
        ["4", "Kepadatan Penduduk Ekstraktif (9.2)", "Jumlah Penduduk per Luas Daratan Kabupaten Smelter", "Jiwa / km²", "BPS SIMDASI & Luas Wilayah"],
        ["5", "Rasio Kepadatan Wilayah (9.2)", "Rasio Densitas Kabupaten Smelter vs Non-Smelter", "Rasio Kelipatan (x)", "Analisis Komparasi Kepadatan"],
        ["6", "Laju Intensifikasi Pemadatan (9.2)", "Peningkatan Densitas Penduduk Selama Periode 2016-2024", "Jiwa / km² & Kelipatan", "Tracking Time-Series BPS"],
        ["7", "Kontribusi PDRB Basis Agraris (9.3)", "Pangsa PDRB Sektor Pertanian, Kehutanan, Perikanan (A)", "Persen PDRB (%)", "BPS PDRB Sektoral Seri 2010"],
        ["8", "Kontribusi Tambang & Industri (9.3)", "Pangsa PDRB Blok Ekstraktif-Industrial (Sektor B + C)", "Persen PDRB (%)", "BPS PDRB Sektoral Seri 2010"],
        ["9", "Indeks Pergeseran Ekonomi / Shift (9.3)", "Rasio Pangsa (B + C) terhadap Sektor Agraris (A)", "Rasio Indeks", "BPS PDRB Sektoral (Ambang 1,0)"]
    ]

    add_table_styled(
        doc,
        headers=["No", "Indikator Riset", "Fokus Pengukuran", "Satuan", "Sumber Data Primer Resmi"],
        rows=table_indikator_data,
        col_widths_cm=[0.8, 4.5, 4.5, 2.2, 5.0],
        alignments=['C', 'L', 'L', 'C', 'L'],
        font_pt=7
    )

    # ── D. KERANGKA ANALISIS & FORMULASI MATEMATIS ──────────────
    add_h2(doc, "D", "Kerangka Analisis & Formulasi Matematis")

    # Sub-bab 9.1
    add_h3(doc, "Sub-bab 9.1: Tekanan Demografi di Kabupaten Industri Ekstraktif")
    add_body(doc, [
        ("Penilaian tekanan demografi dilakukan dengan membaca perubahan jumlah penduduk sebagai sinyal tarikan tenaga kerja (proxy migrasi) dan membandingkan distribusinya melalui algoritma boxplot kuantil Hazen:", False, False)
    ])
    add_formula(doc,
        "1. Laju Pertumbuhan Penduduk Tahunan (YoY %):\n"
        "   Pertumbuhan YoY (%) = [ (Jumlah Penduduk Tahun Ini - Penduduk Tahun Lalu) / Penduduk Tahun Lalu ] × 100%\n\n"
        "2. Rentang Pagar Sebaran Normal (Fences Boxplot Hazen):\n"
        "   • Rentang Interkuartil (IQR) = Kuartil Atas (Q3) - Kuartil Bawah (Q1)\n"
        "   • Pagar Atas (Upper Fence)  = Q3 + (1,5 × IQR)   |   Pagar Bawah (Lower Fence) = Q1 - (1,5 × IQR)",
        ket="Fakta Empiris: Rata-rata pertumbuhan kabupaten smelter mencapai 3,36% (median 2,00%) vs non-smelter 2,03% (median 1,15%). Wilayah smelter membuktikan fenomena Boom and Bust tajam: lonjakan tertinggi mencapai +20,34% di awal fase proyek, disusul kejatuhan terendah hingga -7,76% saat fase konstruksi mereda."
    )

    add_caption(doc, "Tabel 9.1: Rincian Anatomi Boxplot Laju Pertumbuhan Penduduk YoY (%)")
    tabel_9_1_rows = [
        ["Kabupaten Industri Ekstraktif (7 Kab)", "20,34%", "4,22%", "2,78%", "2,00%", "1,50%", "-0,10%", "-7,76%", "3,36%"],
        ["Kabupaten Non-Ekstraktif (Lainnya)", "14,80%", "3,61%", "1,90%", "1,15%", "0,69%", "-0,89%", "-6,73%", "2,03%"]
    ]
    add_table_styled(
        doc,
        headers=["Kategori Wilayah", "Maksimum", "Pagar Atas", "Q3", "Median", "Q1", "Pagar Bawah", "Minimum", "Mean"],
        rows=tabel_9_1_rows,
        col_widths_cm=[4.5, 1.4, 1.5, 1.3, 1.3, 1.3, 1.5, 1.4, 1.4],
        alignments=['L', 'C', 'C', 'C', 'C', 'C', 'C', 'C', 'C'],
        font_pt=6.5
    )

    # Sub-bab 9.2
    add_h3(doc, "Sub-bab 9.2: Intensifikasi Ruang — Kepadatan Industri Ekstraktif vs Non-Ekstraktif")
    add_body(doc, [
        ("Pengukuran intensifikasi ruang menilai laju pemadatan penduduk pada kabupaten industri ekstraktif yang memiliki wilayah daratan luas dan semula berpenduduk jarang:", False, False)
    ])
    add_formula(doc,
        "1. Kepadatan Penduduk Rata-rata Kategori:\n"
        "   Kepadatan Kategori = Total Jumlah Penduduk Seluruh Kabupaten / Total Luas Daratan Seluruh Kabupaten\n\n"
        "2. Rasio Intensifikasi Ruang:\n"
        "   Rasio Kepadatan = Kepadatan Kabupaten Smelter / Kepadatan Kabupaten Non-Smelter",
        ket="Fakta Empiris: Pada 2024, kepadatan kabupaten ekstraktif tercatat 42,7 jiwa/km² vs non-ekstraktif 438,3 jiwa/km² (rasio 0,10x). Kendati rasionya tampak kecil karena luas wilayahnya besar, laju pemadatan di kawasan ekstraktif melipat 14,2 kali (dari 3,0 ke 42,7 jiwa/km²) sepanjang 2016-2024, membuktikan adanya pemadatan ruang drastis tanpa kesiapan fasilitas dasar."
    )

    add_caption(doc, "Tabel 9.2: Tren Kepadatan Penduduk dan Rasio Intensifikasi Ruang (2016–2024)")
    tabel_9_2_rows = [
        ["2016", "3,0 jiwa/km²", "212,8 jiwa/km²", "0,01x", "Awal Ekspansi Smelter Terbuka"],
        ["2018", "39,6 jiwa/km²", "434,0 jiwa/km²", "0,09x", "Arus Masuk Tenaga Kerja Masif"],
        ["2020", "34,2 jiwa/km²", "470,2 jiwa/km²", "0,07x", "Restriksi Mobilitas Pandemi"],
        ["2022", "50,0 jiwa/km²", "347,4 jiwa/km²", "0,14x", "Puncak Operasi Smelter Baru"],
        ["2024", "42,7 jiwa/km²", "438,3 jiwa/km²", "0,10x", "Pemadatan Ruang Meningkat 14,2x Lipat"]
    ]
    add_table_styled(
        doc,
        headers=["Tahun", "Kabupaten Smelter", "Kabupaten Non-Smelter", "Rasio (x)", "Konteks Dinamika Lapangan"],
        rows=tabel_9_2_rows,
        col_widths_cm=[1.8, 3.2, 3.5, 1.8, 6.7],
        alignments=['C', 'C', 'C', 'C', 'L'],
        font_pt=7
    )

    # Sub-bab 9.3
    add_h3(doc, "Sub-bab 9.3: Pergeseran Ekonomi Agraris ke Tambang dan Industri Pengolahan")
    add_body(doc, [
        ("Pergeseran gravitasi ekonomi daerah diukur melalui rasio kontribusi blok ekstraktif-industrial terhadap basis agraris tradisional:", False, False)
    ])
    add_formula(doc,
        "1. Rumus Shift Index Sektoral PDRB:\n"
        "   Shift Index = [ PDRB Sektor B (Pertambangan) + PDRB Sektor C (Industri) ] / PDRB Sektor A (Pertanian)\n\n"
        "2. Garis Ambang Batas Dominasi (Threshold = 1,0):\n"
        "   • Shift Index > 1,0 : Kontribusi Tambang & Industri telah MELAMPAUI basis pangan agraris\n"
        "   • Shift Index <= 1,0 : Perekonomian daerah masih bertumpu pada basis agromaritim tradisional",
        ket="Fakta Empiris: Sulawesi Tengah menjadi episentrum pergeseran dengan lonjakan Shift Index dari 0,449 (2014) menjadi 3,533 (2024), atau melipat 7,9 kali lipat! Pangsa pertanian Sulteng terpangkas separuh (dari 34,39% ke 15,80%), sementara tambang+industri meroket menguasai 55,82% PDRB."
    )

    add_caption(doc, "Tabel 9.3: Ringkasan Pergeseran Struktur Ekonomi (Shift Index) 6 Provinsi Se-Sulawesi")
    tabel_9_3_rows = [
        ["Sulawesi Tengah", "0,449", "3,533", "7,9x", "MELAMPAUI AMBANG (Dominasi Tambang 55,8%)"],
        ["Sulawesi Tenggara", "1,009", "1,300", "1,3x", "MELAMPAUI AMBANG (Dominasi Tambang & Smelter)"],
        ["Sulawesi Selatan", "0,918", "0,804", "0,9x", "Di Bawah Ambang (Basis Agraris & Jasa Kuat)"],
        ["Sulawesi Utara", "0,661", "0,754", "1,1x", "Di Bawah Ambang (Ekonomi Agromaritim & Perikanan)"],
        ["Sulawesi Barat", "0,298", "0,274", "0,9x", "Di Bawah Ambang (Sentra Perkebunan Rakyat)"],
        ["Gorontalo", "0,145", "0,147", "1,0x", "Di Bawah Ambang (Basis Pertanian Jagung & Pangan)"]
    ]
    add_table_styled(
        doc,
        headers=["Provinsi", "Shift Index 2014", "Shift Index 2024", "Multiplier", "Status Ambang Dominasi (B+C > A)"],
        rows=tabel_9_3_rows,
        col_widths_cm=[3.5, 2.5, 2.5, 2.0, 6.5],
        alignments=['L', 'C', 'C', 'C', 'L'],
        font_pt=7
    )

    # ── E. KORESPONDENSI METODOLOGI TERHADAP SUB-BAB LAPORAN ────
    add_h2(doc, "E", "Korespondensi Metodologi terhadap Sub-bab Laporan Bab 9")
    add_body(doc, [
        ("Setiap sub-bab analitis pada Bab 9 ditopang oleh metode empiris terstandarisasi sebagaimana dirangkum pada matriks korespondensi berikut:", False, False)
    ])

    table_korespondensi = [
        ["Sub-bab 9.1", "Tekanan Demografi di Kabupaten Smelter", "Population Time-Series Proxy, Hazen Quantile Boxplot, Boom-and-Bust Disparity Audit"],
        ["Sub-bab 9.2", "Intensifikasi Ruang & Beban Layanan Publik", "Comparative Density Modeling, Spatial Intensification Tracking, Public Service Stress Audit"],
        ["Sub-bab 9.3", "Pergeseran Ekonomi Agraris ke Ekstraktif", "PDRB Sectoral Share Ratio, Shift Index Threshold Analysis, Agrarian Displacement Tracking"]
    ]

    add_table_styled(
        doc,
        headers=["Sub-bab", "Fokus Kajian Empiris", "Metode Analitis Utama"],
        rows=table_korespondensi,
        col_widths_cm=[2.5, 6.0, 8.5],
        alignments=['C', 'L', 'L'],
        font_pt=7.5
    )

    # ── F. BAGAN ALUR KERANGKA KERJA RISET BAB 9 ────────────────
    add_h2(doc, "F", "Bagan Alur Kerangka Kerja Riset (Research Workflow)")
    add_body(doc, [
        ("Kerangka penyelidikan demografi sosial dijalankan secara berjenjang melalui empat tahapan metodologis berikut:", False, False)
    ])

    mermaid_str_f = """flowchart LR
    subgraph F1["Fase I: Input Data BPS"]
        A1["Populasi Penduduk<br/><i>BPS SIMDASI 2014-2024</i>"]
        A2["Klaster 7 Kab Smelter<br/><i>1.59 Juta Jiwa (Fase 4)</i>"]
        A3["Luas Daratan Wilayah<br/><i>Luas km2 Kabupaten</i>"]
        A4["PDRB Sektoral Seri 2010<br/><i>Sektor A vs B+C</i>"]
    end
    subgraph F2["Fase II: Segmentasi & Komputasi"]
        B1["Laju Pertumbuhan YoY<br/><i>Proxy Tarikan Migrasi</i>"]
        B2["Pemodelan Kepadatan<br/><i>Jiwa per km2 Kategori</i>"]
        B3["PDRB Shift Ratio<br/><i>(B+C) dibagi Sektor A</i>"]
    end
    subgraph F3["Fase III: Evaluasi Transformasi"]
        C1["Boxplot Hazen<br/><i>Boom & Bust Ekstraktif</i>"]
        C2["Intensifikasi Ruang<br/><i>Pemadatan 14.2x Lipat</i>"]
        C3["Ambang Batas Index 1.0<br/><i>Sulteng 3.53 & Sultra 1.30</i>"]
    end
    subgraph F4["Fase IV: Sintesis Disrupsi Sosial"]
        D1["Tekanan Demografi<br/><i>Tarikan Pekerja Migran</i>"]
        D2["Stres Layanan Publik<br/><i>Sanitasi, Air & Faskes</i>"]
        D3["Alih Gravitasi Ekonomi<br/><i>Tergusurnya Basis Agraris</i>"]
    end
    F1 --> F2 --> F3 --> F4"""

    png_workflow_path = str(out_dir_compact / "mermaid_workflow_bab9_compact.png")
    is_downloaded = download_mermaid_png(mermaid_str_f, png_workflow_path)

    add_caption(doc, "Bagan Alur 9.1: Alur Logika Analisis Demografi Sosial dan Transformasi Struktural (Research Workflow)")
    if is_downloaded and os.path.exists(png_workflow_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(3)
        p_img.paragraph_format.space_after  = Pt(4)
        r_img = p_img.add_run()
        r_img.add_picture(png_workflow_path, width=Cm(16.5))
        try:
            shutil.copyfile(png_workflow_path, str(out_dir_bab9 / "mermaid_workflow_bab9_compact.png"))
        except Exception:
            pass

    # Box Output Kesimpulan
    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after  = Pt(4)
    all_border_para(p_box, color='1B5E20', sz='8')
    para_shd(p_box, 'F1F8E9')
    add_run(p_box, "KESIMPULAN METODOLOGIS BAB 9 (DISRUPSI DEMOGRAFI SOSIAL & PDRB):\n", bold=True, pt=8.5, color=G_DARK)
    add_run(p_box, "1. Anomali Boom and Bust Demografis: Kawasan industri smelter mencatat rata-rata pertumbuhan penduduk 3,36% dengan rentang variabilitas ekstrem (+20,34% ke -7,76%), membuktikan adanya tarikan migrasi massal di awal proyek yang rentan terhadap guncangan PHK dan eksodus pasca-konstruksi.\n"
                   "2. Akselerasi Intensifikasi Ruang: Kepadatan penduduk di kabupaten smelter melesat 14,2 kali lipat sepanjang 2016–2024, memicu tekanan berat terhadap daya dukung infrastruktur perumahan, air bersih, sanitasi, dan fasilitas kesehatan perdesaan.\n"
                   "3. Pergeseran Gravitasi Ekonomi: Hilirisasi memicu pergeseran struktural tajam di Sulawesi Tengah (Shift Index naik 7,9x ke 3,533) dan Sulawesi Tenggara (1,300), di mana dominasi industri ekstraktif menggeser ruang produksi agraris dan memperbesar ketergantungan daerah pada rantai pasok modal besar.",
            pt=8, color=C_BODY)

    # ── SIMPAN DOKUMEN DOCX (DUAL SAVE) ─────────────────────────
    docx_compact = out_dir_compact / "Metodologi_Bab9_Demografi_Sosial_Compact.docx"
    docx_bab9    = out_dir_bab9 / "Metodologi_Bab9_Demografi_Sosial_Compact.docx"
    
    doc.save(str(docx_compact))
    shutil.copyfile(docx_compact, docx_bab9)
    print(f"  [OK] Tersimpan DOCX: {docx_compact}")
    print(f"  [OK] Salinan DOCX : {docx_bab9}")

    # ── GENERATE MARKDOWN PADANAN ───────────────────────────────
    print("[2/3] Membangun dokumen Markdown padanan...")
    MD_CONTENT = """# BAB IX: METODOLOGI ANALISIS DEMOGRAFI SOSIAL — KETIKA HILIRISASI MENGUBAH STRUKTUR MASYARAKAT
*Ringkasan Eksekutif Metodologis · Center of Economic and Law Studies (CELIOS)*

---

## A. Desain Penelitian & Tujuan
Penelitian Bab 9 menerapkan **desain analisis demografi spasial dan transformasi struktural ekonomi (Spatial Demography & Structural Transformation Analysis)** guna menguji disrupsi sosial yang terjadi akibat penetrasi industri hilirisasi nikel skala masif di Pulau Sulawesi. Melalui pembacaan data deret waktu populasi, pemodelan sebaran kuantil, dan rasio pergeseran sektoral PDRB, kajian ini membuktikan tiga fenomena perubahan sosial-spasial:

1. **Tekanan Demografi & Proxy Migrasi (Hazen Quantile Boxplot Analysis):** Menganalisis anomali lonjakan penduduk tahunan (YoY) pada 7 kabupaten prioritas smelter dibandingkan kabupaten non-smelter, guna membuktikan fenomena tarikan migrasi tenaga kerja dan siklus fluktuasi tajam (*boom and bust*).
2. **Intensifikasi Ruang & Beban Layanan Publik (Comparative Density Analysis):** Mengukur laju pemadatan penduduk per kilometer persegi pada kawasan industri ekstraktif yang semula berpenduduk jarang, sebagai indikator stres daya dukung sarana air bersih, sanitasi, dan perumahan lokal.
3. **Pergeseran Gravitasi Ekonomi Sektoral (PDRB Sector Shift Index):** Mengkuantifikasi transformasi struktur perekonomian daerah dari basis agraris (Sektor A: Pertanian, Kehutanan, Perikanan) menuju dominasi blok ekstraktif-industrial (Sektor B: Pertambangan dan C: Industri Pengolahan).

---

## B. Sumber Data & Cakupan Wilayah
Analisis demografi sosial ini mengolah basis data panel resmi Badan Pusat Statistik (BPS) kurun waktu 2014–2024 yang mencakup seluruh wilayah kabupaten/kota dan 6 provinsi se-Pulau Sulawesi:

- **BPS SIMDASI (Sistem Informasi Rujukan Statistik Terintegrasi):** Data deret waktu populasi penduduk kabupaten/kota, luas daratan yurisdiksi, dan laju pertumbuhan penduduk YoY.
- **Klasifikasi 7 Kabupaten Prioritas Smelter (Fase 4):** Klaster kabupaten sentra industri pengolahan nikel: Banggai, Kolaka, Konawe, Konawe Utara, Luwu Timur, Morowali, dan Morowali Utara (total populasi 2024 mencapai 1,59 juta jiwa).
- **BPS PDRB Sektoral Seri 2010 (Tahun 2014–2024):** Struktur Produk Domestik Regional Bruto menurut lapangan usaha: Sektor A (Pertanian), Sektor B (Pertambangan), dan Sektor C (Industri Pengolahan).
- **Statistik Perikanan Tangkap BPS:** Dekomposisi estimasi kontribusi sub-sektor perikanan tangkap laut (~22% dari Sektor A) pada provinsi-provinsi pesisir Sulawesi.

---

## C. Operasionalisasi Variabel & Indikator Riset
Seluruh parameter demografi, kepadatan spasial, hingga pergeseran struktur produksi dioperasionalkan ke dalam **9 indikator riset empiris terverifikasi** sebagaimana dirangkum pada matriks operasional berikut:

##### Matriks Operasionalisasi Variabel dan Sumber Data Resmi Bab 9 (Demografi Sosial)
| No | Indikator Riset | Fokus Pengukuran | Satuan | Sumber Data Primer Resmi |
| :-: | :--- | :--- | :-: | :--- |
| 1 | Laju Pertumbuhan Penduduk YoY (9.1) | Pertumbuhan Tahunan Populasi Kabupaten (Proxy Migrasi) | Persen (% / Tahun) | BPS SIMDASI Demografi |
| 2 | Anatomi Sebaran Boxplot Hazen (9.1) | Median, Q1, Q3, IQR, dan Rentang Pagar Kewajaran | Nilai Persentil (%) | BPS SIMDASI (Kuantil Hazen) |
| 3 | Anomali Fluktuasi Boom-and-Bust (9.1) | Lonjakan Ekstrem Tertinggi vs Kejatuhan Terendah | Persen Ekstrem (%) | Dataset Deret Waktu BPS |
| 4 | Kepadatan Penduduk Ekstraktif (9.2) | Jumlah Penduduk per Luas Daratan Kabupaten Smelter | Jiwa / km² | BPS SIMDASI & Luas Wilayah |
| 5 | Rasio Kepadatan Wilayah (9.2) | Rasio Densitas Kabupaten Smelter vs Non-Smelter | Rasio Kelipatan (x) | Analisis Komparasi Kepadatan |
| 6 | Laju Intensifikasi Pemadatan (9.2) | Peningkatan Densitas Penduduk Selama Periode 2016-2024 | Jiwa / km² & Kelipatan | Tracking Time-Series BPS |
| 7 | Kontribusi PDRB Basis Agraris (9.3) | Pangsa PDRB Sektor Pertanian, Kehutanan, Perikanan (A) | Persen PDRB (%) | BPS PDRB Sektoral Seri 2010 |
| 8 | Kontribusi Tambang & Industri (9.3) | Pangsa PDRB Blok Ekstraktif-Industrial (Sektor B + C) | Persen PDRB (%) | BPS PDRB Sektoral Seri 2010 |
| 9 | Indeks Pergeseran Ekonomi / Shift (9.3) | Rasio Pangsa (B + C) terhadap Sektor Agraris (A) | Rasio Indeks | BPS PDRB Sektoral (Ambang 1,0) |

---

## D. Kerangka Analisis & Formulasi Matematis

### Sub-bab 9.1: Tekanan Demografi di Kabupaten Industri Ekstraktif
Penilaian tekanan demografi dilakukan dengan membaca perubahan jumlah penduduk sebagai sinyal tarikan tenaga kerja (proxy migrasi) dan membandingkan distribusinya melalui algoritma boxplot kuantil Hazen:

> **1. Laju Pertumbuhan Penduduk Tahunan (YoY %):**  
> `Pertumbuhan YoY (%) = [ (Jumlah Penduduk Tahun Ini - Penduduk Tahun Lalu) / Penduduk Tahun Lalu ] × 100%`  
>  
> **2. Rentang Pagar Sebaran Normal (Fences Boxplot Hazen):**  
> • Rentang Interkuartil (IQR) = Kuartil Atas (Q3) - Kuartil Bawah (Q1)  
> • Pagar Atas (Upper Fence) = Q3 + (1,5 × IQR)  
> • Pagar Bawah (Lower Fence) = Q1 - (1,5 × IQR)  
>  
> *Fakta Empiris: Rata-rata pertumbuhan kabupaten smelter mencapai 3,36% (median 2,00%) vs non-smelter 2,03% (median 1,15%). Wilayah smelter membuktikan fenomena Boom and Bust tajam: lonjakan tertinggi mencapai +20,34% di awal fase proyek, disusul kejatuhan terendah hingga -7,76% saat fase konstruksi mereda.*

##### Tabel 9.1: Rincian Anatomi Boxplot Laju Pertumbuhan Penduduk YoY (%)
| Kategori Wilayah | Maksimum | Pagar Atas | Q3 | Median | Q1 | Pagar Bawah | Minimum | Mean |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: |
| Kabupaten Industri Ekstraktif (7 Kab) | 20,34% | 4,22% | 2,78% | 2,00% | 1,50% | -0,10% | -7,76% | 3,36% |
| Kabupaten Non-Ekstraktif (Lainnya) | 14,80% | 3,61% | 1,90% | 1,15% | 0,69% | -0,89% | -6,73% | 2,03% |

---

### Sub-bab 9.2: Intensifikasi Ruang — Kepadatan Industri Ekstraktif vs Non-Ekstraktif
Pengukuran intensifikasi ruang menilai laju pemadatan penduduk pada kabupaten industri ekstraktif yang memiliki wilayah daratan luas dan semula berpenduduk jarang:

> **1. Kepadatan Penduduk Rata-rata Kategori:**  
> `Kepadatan Kategori = Total Jumlah Penduduk Seluruh Kabupaten / Total Luas Daratan Seluruh Kabupaten`  
>  
> **2. Rasio Intensifikasi Ruang:**  
> `Rasio Kepadatan = Kepadatan Kabupaten Smelter / Kepadatan Kabupaten Non-Smelter`  
>  
> *Fakta Empiris: Pada 2024, kepadatan kabupaten ekstraktif tercatat 42,7 jiwa/km² vs non-ekstraktif 438,3 jiwa/km² (rasio 0,10x). Kendati rasionya tampak kecil karena luas wilayahnya besar, laju pemadatan di kawasan ekstraktif melipat 14,2 kali (dari 3,0 ke 42,7 jiwa/km²) sepanjang 2016-2024, membuktikan adanya pemadatan ruang drastis tanpa kesiapan fasilitas dasar.*

##### Tabel 9.2: Tren Kepadatan Penduduk dan Rasio Intensifikasi Ruang (2016–2024)
| Tahun | Kabupaten Smelter | Kabupaten Non-Smelter | Rasio (x) | Konteks Dinamika Lapangan |
| :---: | :---: | :---: | :---: | :--- |
| 2016 | 3,0 jiwa/km² | 212,8 jiwa/km² | 0,01x | Awal Ekspansi Smelter Terbuka |
| 2018 | 39,6 jiwa/km² | 434,0 jiwa/km² | 0,09x | Arus Masuk Tenaga Kerja Masif |
| 2020 | 34,2 jiwa/km² | 470,2 jiwa/km² | 0,07x | Restriksi Mobilitas Pandemi |
| 2022 | 50,0 jiwa/km² | 347,4 jiwa/km² | 0,14x | Puncak Operasi Smelter Baru |
| 2024 | 42,7 jiwa/km² | 438,3 jiwa/km² | 0,10x | Pemadatan Ruang Meningkat 14,2x Lipat |

---

### Sub-bab 9.3: Pergeseran Ekonomi Agraris ke Tambang dan Industri Pengolahan
Pergeseran gravitasi ekonomi daerah diukur melalui rasio kontribusi blok ekstraktif-industrial terhadap basis agraris tradisional:

> **1. Rumus Shift Index Sektoral PDRB:**  
> `Shift Index = [ PDRB Sektor B (Pertambangan) + PDRB Sektor C (Industri) ] / PDRB Sektor A (Pertanian)`  
>  
> **2. Garis Ambang Batas Dominasi (Threshold = 1,0):**  
> • **Shift Index > 1,0** : Kontribusi Tambang & Industri telah MELAMPAUI basis pangan agraris  
> • **Shift Index ≤ 1,0** : Perekonomian daerah masih bertumpu pada basis agromaritim tradisional  
>  
> *Fakta Empiris: Sulawesi Tengah menjadi episentrum pergeseran dengan lonjakan Shift Index dari 0,449 (2014) menjadi 3,533 (2024), atau melipat 7,9 kali lipat! Pangsa pertanian Sulteng terpangkas separuh (dari 34,39% ke 15,80%), sementara tambang+industri meroket menguasai 55,82% PDRB.*

##### Tabel 9.3: Ringkasan Pergeseran Struktur Ekonomi (Shift Index) 6 Provinsi Se-Sulawesi
| Provinsi | Shift Index 2014 | Shift Index 2024 | Multiplier | Status Ambang Dominasi (B+C > A) |
| :--- | :---: | :---: | :---: | :--- |
| **Sulawesi Tengah** | **0,449** | **3,533** | **7,9x** | **MELAMPAUI AMBANG (Dominasi Tambang 55,8%)** |
| **Sulawesi Tenggara** | **1,009** | **1,300** | **1,3x** | **MELAMPAUI AMBANG (Dominasi Tambang & Smelter)** |
| Sulawesi Selatan | 0,918 | 0,804 | 0,9x | Di Bawah Ambang (Basis Agraris & Jasa Kuat) |
| Sulawesi Utara | 0,661 | 0,754 | 1,1x | Di Bawah Ambang (Ekonomi Agromaritim & Perikanan) |
| Sulawesi Barat | 0,298 | 0,274 | 0,9x | Di Bawah Ambang (Sentra Perkebunan Rakyat) |
| Gorontalo | 0,145 | 0,147 | 1,0x | Di Bawah Ambang (Basis Pertanian Jagung & Pangan) |

---

## E. Korespondensi Metodologi terhadap Sub-bab Laporan Bab 9
Setiap sub-bab analitis pada Bab 9 ditopang oleh metode empiris terstandarisasi sebagaimana dirangkum pada matriks korespondensi berikut:

##### Matriks Korespondensi Metodologis Bab 9
| Sub-bab | Fokus Kajian Empiris | Metode Analitis Utama |
| :-: | :--- | :--- |
| Sub-bab 9.1 | Tekanan Demografi di Kabupaten Smelter | Population Time-Series Proxy, Hazen Quantile Boxplot, Boom-and-Bust Disparity Audit |
| Sub-bab 9.2 | Intensifikasi Ruang & Beban Layanan Publik | Comparative Density Modeling, Spatial Intensification Tracking, Public Service Stress Audit |
| Sub-bab 9.3 | Pergeseran Ekonomi Agraris ke Ekstraktif | PDRB Sectoral Share Ratio, Shift Index Threshold Analysis, Agrarian Displacement Tracking |

---

## F. Bagan Alur Kerangka Kerja Riset (Research Workflow)
Kerangka penyelidikan demografi sosial dijalankan secara berjenjang melalui empat tahapan metodologis berikut:

```mermaid
flowchart LR
    subgraph F1["Fase I: Input Data BPS"]
        A1["Populasi Penduduk<br/><i>BPS SIMDASI 2014-2024</i>"]
        A2["Klaster 7 Kab Smelter<br/><i>1.59 Juta Jiwa (Fase 4)</i>"]
        A3["Luas Daratan Wilayah<br/><i>Luas km2 Kabupaten</i>"]
        A4["PDRB Sektoral Seri 2010<br/><i>Sektor A vs B+C</i>"]
    end
    subgraph F2["Fase II: Segmentasi & Komputasi"]
        B1["Laju Pertumbuhan YoY<br/><i>Proxy Tarikan Migrasi</i>"]
        B2["Pemodelan Kepadatan<br/><i>Jiwa per km2 Kategori</i>"]
        B3["PDRB Shift Ratio<br/><i>(B+C) dibagi Sektor A</i>"]
    end
    subgraph F3["Fase III: Evaluasi Transformasi"]
        C1["Boxplot Hazen<br/><i>Boom & Bust Ekstraktif</i>"]
        C2["Intensifikasi Ruang<br/><i>Pemadatan 14.2x Lipat</i>"]
        C3["Ambang Batas Index 1.0<br/><i>Sulteng 3.53 & Sultra 1.30</i>"]
    end
    subgraph F4["Fase IV: Sintesis Disrupsi Sosial"]
        D1["Tekanan Demografi<br/><i>Tarikan Pekerja Migran</i>"]
        D2["Stres Layanan Publik<br/><i>Sanitasi, Air & Faskes</i>"]
        D3["Alih Gravitasi Ekonomi<br/><i>Tergusurnya Basis Agraris</i>"]
    end
    F1 --> F2 --> F3 --> F4
```

> **KESIMPULAN METODOLOGIS BAB 9 (DISRUPSI DEMOGRAFI SOSIAL & PDRB):**  
> 1. **Anomali Boom and Bust Demografis:** Kawasan industri smelter mencatat rata-rata pertumbuhan penduduk 3,36% dengan rentang variabilitas ekstrem (+20,34% ke -7,76%), membuktikan adanya tarikan migrasi massal di awal proyek yang rentan terhadap guncangan PHK dan eksodus pasca-konstruksi.  
> 2. **Akselerasi Intensifikasi Ruang:** Kepadatan penduduk di kabupaten smelter melesat 14,2 kali lipat sepanjang 2016–2024, memicu tekanan berat terhadap daya dukung infrastruktur perumahan, air bersih, sanitasi, dan fasilitas kesehatan perdesaan.  
> 3. **Pergeseran Gravitasi Ekonomi:** Hilirisasi memicu pergeseran struktural tajam di Sulawesi Tengah (Shift Index naik 7,9x ke 3,533) dan Sulawesi Tenggara (1,300), di mana dominasi industri ekstraktif menggeser ruang produksi agraris dan memperbesar ketergantungan daerah pada rantai pasok modal besar.
"""

    md_compact = out_dir_compact / "Metodologi_Bab9_Demografi_Sosial_Compact.md"
    md_bab9    = out_dir_bab9 / "Metodologi_Bab9_Demografi_Sosial_Compact.md"
    with open(md_compact, 'w', encoding='utf-8') as f:
        f.write(MD_CONTENT)
    shutil.copyfile(md_compact, md_bab9)
    print(f"  [OK] Tersimpan MD  : {md_compact}")
    print(f"  [OK] Salinan MD   : {md_bab9}")

    print("[3/3] Selesai menghasilkan dokumen metodologi Bab 9 versi compact (1-Kolom, 2-3 Halaman).\n")


if __name__ == "__main__":
    generate_bab9_compact()
