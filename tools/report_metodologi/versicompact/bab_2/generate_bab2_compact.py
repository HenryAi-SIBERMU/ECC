#!/usr/bin/env python3
"""
Generator Laporan Metodologi Statistik (Versi Compact) Bab 2:
Metodologi Analisis Kualitas Lingkungan di Kawasan Smelter

Standar Baku CELIOS Versi Compact:
- Highlight pada Formulasi Matematis dan Persamaan Substitusi substantif (angka nyata)
- Notasi formulasi WAJIB "mudah dibaca publik" ala referensi Celios-Metodologi-Statistik-2:
  nama variabel deskriptif bahasa sehari-hari, operator ÷ dan ×, TANPA simbol Σ / subscript {p,t}
- Untuk Crosstab / Chi-Square / Odds Ratio: TIDAK MENGGUNAKAN persamaan substitusi aritmatika;
  WAJIB menyertakan Tabel Konfigurasi Variabel Uji (X, Y, H1, decision rule, threshold median, orientasi OR)
- Sumber data ditulis sebagai INSTITUSI RESMI (ESDM, KLHK, BPS, GEM, GFW, NASA TROPOMI, GBIF, IUCN),
  DILARANG menuliskan nama file dataset .csv (dokumen dibaca publik)
- Tanpa label artifisial 'Analisis Temuan Empiris:' (narasi mengalir natural)
- Header dan Judul SAMA PERSIS dengan dokumen non-compact root
- Penomoran sub-bab langsung: 2.1, 2.2, 2.3, 2.4, 2.5
- Target Panjang: maksimal 3-4 lembar di Microsoft Word (tidak perlu pemadatan ekstrem)
- Tanpa icon / emoji
- Narasi, rumus, dan tabel murni dari Metodologi_Bab2_Kualitas_Lingkungan.md
"""

import os
import sys
import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Warna Tema CELIOS D3TLH ─────────────────────────────────
G_DARK   = RGBColor(0x1B, 0x5E, 0x20)  # Hijau Hutan Gelap (#1B5E20)
G_MID    = RGBColor(0x2E, 0x7D, 0x32)  # Hijau Utama CELIOS (#2E7D32)
C_BODY   = RGBColor(0x22, 0x22, 0x22)  # Abu Gelap Teks (#222222)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)  # Putih

# ── Helper XML Word Formatting (Ultra-Compact) ──────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=40, right=40, top=18, bottom=18):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    tcPr.append(s)

def para_shd(p, fill):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    pPr.append(s)

def para_border_left(p, color='2E7D32', sz='14'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:left')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def para_border_bottom(p, color='2E7D32', sz='6'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '1')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Tipografi Padat ──────────────────────────────────
def run(p, text, bold=False, italic=False, pt=8.5, color=None, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    return r

def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1.5)
    para_border_bottom(p, color='1B5E20', sz='8')
    run(p, title.upper(), bold=True, pt=10.0, color=G_DARK)

def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3.5)
    p.paragraph_format.space_after  = Pt(1)
    para_border_bottom(p, color='2E7D32', sz='4')
    run(p, title, bold=True, pt=9.0, color=G_DARK)

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2.5)
    p.paragraph_format.space_after  = Pt(1)
    run(p, title, bold=True, pt=8.5, color=G_MID)

def add_p(doc, parts, space_after=2, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=8.5)
    return p

def add_note_inline(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    para_border_left(p, color='2E7D32', sz='10')
    para_shd(p, 'F1F8E9')
    run(p, f"{title}: ", bold=True, pt=7.5, color=G_DARK)
    run(p, text, italic=True, pt=7.5, color=RGBColor(0x33, 0x33, 0x33))

def add_math_block(doc, title, formula_str, sub_str=None, ket_str=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    para_shd(p, 'EDF7EE')
    para_border_left(p, color='2E7D32', sz='12')
    run(p, f"Formulasi Matematis ({title}):\n", bold=True, pt=7.5, color=G_DARK)
    run(p, f"{formula_str}\n", pt=7.2, color=RGBColor(0x10, 0x40, 0x10), mono=True)
    if sub_str:
        run(p, "Persamaan Substitusi:\n", bold=True, pt=7.5, color=RGBColor(0x1B, 0x5E, 0x20))
        run(p, f"{sub_str}", pt=7.2, color=RGBColor(0x22, 0x22, 0x22), mono=True)
    if ket_str:
        run(p, f"\nKet: {ket_str}", italic=True, pt=7.0, color=RGBColor(0x55, 0x55, 0x55))

def add_caption_compact(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    run(p, caption_text, bold=True, italic=True, pt=7.5, color=G_MID)

def add_table_compact(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.autofit = False

    bd_cfg = {'val': 'single', 'sz': '4', 'color': 'D8D8D8', 'space': '0'}

    # Header
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=40, right=40, top=18, bottom=18)
        set_cell_borders(c, top=bd_cfg, left=bd_cfg, bottom={'val': 'single', 'sz': '8', 'color': '1B5E20', 'space': '0'}, right=bd_cfg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (alignments and alignments[j] == 'C') else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run(p, h, bold=True, pt=7.5, color=C_WHITE)

    # Rows
    for i, row in enumerate(rows):
        bg = 'F9FBF9' if i % 2 == 1 else 'FFFFFF'
        for j, (val, w) in enumerate(zip(row, col_widths_cm)):
            c = tbl.rows[1 + i].cells[j]
            c.width = Cm(w)
            cell_shd(c, bg)
            cell_margin(c, left=40, right=40, top=16, bottom=16)
            set_cell_borders(c, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = c.paragraphs[0]
            align = alignments[j] if alignments else 'L'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'C' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'R' else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            is_bold = (j == 0) or ('Total' in str(val)) or ('SIGNIFIKAN' in str(val))
            run(p, str(val), bold=is_bold, pt=7.0, color=C_BODY)

# ── Main Generator ──────────────────────────────────────────
def build_compact_report():
    print("[1/4] Menginisialisasi dokumen Word Bab 2 dengan fokus Formulasi & Persamaan Substitusi...")
    doc = Document()

    # Margin Halaman Padat (1.2 cm)
    for section in doc.sections:
        section.top_margin    = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin   = Cm(1.2)
        section.right_margin  = Cm(1.2)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(8.5)

    # ── HEADER BANNER ───────────────────────────────────────────
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(1)
    run(p_hdr, "CELIOS — CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=7.5, color=G_MID)

    add_h1(doc, "BAB II: METODOLOGI ANALISIS KUALITAS LINGKUNGAN DI KAWASAN SMELTER")

    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, formulasi matematis, prosedur pengolahan data, dan pengujian statistik yang dioperasionalkan pada ", False, False),
        ("Bab 2: Kualitas Lingkungan di Kawasan Smelter", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi periode 2014–2024.", False, False),
    ], space_after=2)

    # ═══════════════════════════════════════════════════════════
    # 2.1 DAMPAK LIMBAH TAILING: SMELTER VS IKA
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "2.1. Dampak Limbah Tailing: Konsentrasi Smelter vs Indeks Kualitas Air (IKA)")
    add_p(doc, [
        ("Pengoperasian ", False, False),
        ("778 fasilitas mega-smelter", True, False),
        (" yang didukung oleh kapasitas ", False, False),
        ("9.825 MW PLTU Captive", True, False),
        (" meningkatkan intensitas emisi dan beban lingkungan di Pulau Sulawesi. Konversi tutupan hutan mencapai ", False, False),
        ("1.001.654 Hektar", True, False),
        (", estimasi timbulan limbah B3/tailing sebesar ", False, False),
        ("20,9 Juta Ton per tahun", True, False),
        (", dan rata-rata Indeks Kualitas Air (IKA) tahun 2024 sebesar ", False, False),
        ("59,7", True, False),
        (".", False, False)
    ])
    add_note_inline(doc, "Sumber Data", "Kementerian ESDM & Center for Global Sustainability (fasilitas smelter); KLHK & BPS (Indeks Kualitas Air 2016–2024); laporan lembaga masyarakat sipil (proksi timbulan limbah B3/tailing) dan inventarisasi laporan pencemaran sungai-pesisir (diolah CELIOS).")

    add_caption_compact(doc, "Tabel 2.1: Rincian Empiris Konsentrasi Smelter, IKA, Limbah B3, dan Sungai Tercemar per Provinsi (2024)")
    t1_headers = ["Provinsi", "Smelter (Unit)", "Skor IKA", "Limbah B3 (Ton/Thn)", "Sungai Tercemar", "Daftar Sungai / Pesisir Terdampak"]
    t1_rows = [
        ["Sulawesi Tengah", "344", "62.1", "12,000,000", "4", "Sungai Bahodopi, Laroenai, Morowali, Pesisir Fatufia"],
        ["Sulawesi Tenggara", "262", "65.3", "7,700,000", "3", "Sungai Lasolo, Sungai Lalindu, Sungai Konaweha"],
        ["Sulawesi Selatan", "111", "58.5", "1,000,000", "1", "Pesisir dan Sungai Bantaeng"],
        ["Sulawesi Barat", "39", "55.9", "0", "0", "Tidak teridentifikasi pembuangan tailing smelter"],
        ["Sulawesi Utara", "15", "58.2", "0", "0", "Tidak teridentifikasi pembuangan tailing smelter"],
        ["Gorontalo", "7", "58.1", "0", "0", "Tidak teridentifikasi pembuangan tailing smelter"]
    ]
    add_table_compact(doc, t1_headers, t1_rows, [3.2, 2.2, 1.8, 3.2, 2.2, 6.0], ['L', 'C', 'C', 'R', 'C', 'L'])

    # Sesuai arahan user: Untuk Crosstab TIDAK USAH persamaan substitusi aritmatika
    add_math_block(
        doc,
        "Konsentrasi Smelter & Beban Limbah B3 Tailing",
        "Jumlah Smelter Provinsi = Penjumlahan seluruh fasilitas smelter di provinsi tersebut  |  Rata-rata IKA Provinsi = Jumlah skor IKA seluruh titik pantau ÷ Banyaknya titik pantau\n"
        "Chi-Square (χ²) = Jumlah dari [ ( Frekuensi Observasi - Frekuensi Harapan )² ÷ Frekuensi Harapan ]  |  Odds Ratio (OR) = ( a × d ) ÷ ( b × c )",
        "Gabungan Smelter Sulteng + Sultra = 344 unit + 262 unit = 606 unit (77,89% dari total se-Sulawesi)\n"
        "Estimasi Limbah B3 Tailing = 12.000.000 Ton/Thn (Sulteng) + 7.700.000 Ton/Thn (Sultra) = 19.700.000 Ton/Thn (94,26% dari total limbah B3)",
        "Hasil uji Chi-Square dan Odds Ratio data panel disajikan pada Tabel 2.5, dengan konfigurasi variabel uji pada Tabel 2.4. Kegagalan signifikansi statistik membuktikan Aggregate Dilution Bias (pencemaran fatal sungai industri Bahodopi & Lasolo terencerkan oleh stasiun pemantau sungai non-industri)."
    )

    # ═══════════════════════════════════════════════════════════
    # 2.2 KEPUNGAN ASAP: KAPASITAS PLTU VS IKU
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "2.2. Kepungan Asap: Kapasitas PLTU vs Indeks Kualitas Udara (IKU)")
    add_p(doc, [
        ("Sebanyak ", False, False),
        ("9.825 MW PLTU Captive batu bara off-grid", True, False),
        (" beroperasi di kawasan hilirisasi nikel menyumbang polusi udara ambien dan gas buang NO₂.", False, False)
    ])
    add_note_inline(doc, "Sumber Data", "Kementerian ESDM & Global Energy Monitor (kapasitas PLTU Captive & Grid PLN); KLHK (Indeks Kualitas Udara 2015–2024); pantauan emisi Satelit NASA TROPOMI (konsentrasi NO₂) (diolah CELIOS).")

    add_caption_compact(doc, "Tabel 2.2: Rincian Empiris Kapasitas PLTU (Captive & Grid), IKU, dan Konsentrasi NO₂ NASA (2024)")
    t2_headers = ["Provinsi", "Kapasitas PLTU Captive (MW)", "PLTU Grid PLN (MW)", "Total Daya (MW)", "Skor IKU", "NASA TROPOMI NO₂ (mol/m²)"]
    t2_rows = [
        ["Sulawesi Tengah", "9,365", "0", "9,365", "92.9", "6.50e-06"],
        ["Sulawesi Tenggara", "2,280", "100", "2,380", "93.0", "6.62e-06"],
        ["Sulawesi Selatan", "600", "920", "1,520", "91.5", "6.40e-06"],
        ["Sulawesi Utara", "0", "220", "220", "93.4", "4.09e-06"],
        ["Gorontalo", "0", "100", "100", "93.5", "3.76e-06"],
        ["Sulawesi Barat", "0", "0", "0", "92.5", "6.00e-06"]
    ]
    add_table_compact(doc, t2_headers, t2_rows, [3.8, 3.2, 2.6, 2.6, 2.2, 4.2], ['L', 'R', 'R', 'R', 'C', 'C'])

    # Sesuai arahan user: Untuk Crosstab TIDAK USAH persamaan substitusi aritmatika
    add_math_block(
        doc,
        "Kapasitas Energi PLTU & Parameter Kualitas Udara",
        "Kapasitas PLTU Provinsi = Penjumlahan kapasitas seluruh unit PLTU di provinsi tersebut  |  Rata-rata IKU = Jumlah skor IKU ÷ Banyaknya observasi provinsi-tahun\n"
        "Chi-Square (χ²) = Jumlah dari [ ( Frekuensi Observasi - Frekuensi Harapan )² ÷ Frekuensi Harapan ]  |  Odds Ratio (OR) = ( a × d ) ÷ ( b × c )",
        "Total PLTU Captive 3 Sentra = 9.365 MW (Sulteng) + 2.280 MW (Sultra) + 600 MW (Sulsel) = 12.245 MW terpasang",
        "Hasil pengujian statistik tabulasi silang dirinci pada Tabel 2.5, dengan konfigurasi variabel uji pada Tabel 2.4. Ketidaksignifikanan membuktikan Efek Pengenceran Udara Ambien karena sensor IKU tersebar merata di hutan berudara bersih."
    )

    # ═══════════════════════════════════════════════════════════
    # 2.3 EKSEKUSI RUANG: EKSPANSI INDUSTRI VS DEFORESTASI
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "2.3. Eksekusi Ruang: Ekspansi Kawasan Industri vs Tekanan Ekologis (Deforestasi)")
    add_p(doc, [
        ("Alokasi konsesi IUP dan Kawasan Industri mencakup ", False, False),
        ("1.185.174 Hektar", True, False),
        (" di Sulawesi. Sepanjang dekade 2014–2023, data Global Forest Watch (GFW) merekam akumulasi kehilangan tutupan pohon sebesar ", False, False),
        ("1.386.055 Hektar", True, False),
        (" (terbesar di Sulawesi Tengah dan Tenggara).", False, False)
    ])
    add_note_inline(doc, "Sumber Data", "Kementerian ESDM — MODI/Minerbaone (izin konsesi IUP & kawasan industri); Global Forest Watch — University of Maryland (kehilangan tutupan pohon 2014–2023) (diolah CELIOS).")

    add_caption_compact(doc, "Tabel 2.3: Rincian Empiris Luas Konsesi IUP-Kawasan Industri dan Deforestasi Kumulatif (2014–2023)")
    t3_headers = ["Provinsi", "Luas IUP & Kawasan (Ha)", "Konsesi Baru Kumulatif (Ha)", "Deforestasi Kumulatif 1 Dekade (Ha)"]
    t3_rows = [
        ["Sulawesi Tengah", "453,216", "387,124", "481,908"],
        ["Sulawesi Tenggara", "446,025", "212,717", "337,434"],
        ["Sulawesi Selatan", "181,469", "123,065", "261,147"],
        ["Sulawesi Utara", "94,829", "89,170", "74,240"],
        ["Gorontalo", "5,212", "5,212", "98,063"],
        ["Sulawesi Barat", "4,424", "2,163", "133,263"]
    ]
    add_table_compact(doc, t3_headers, t3_rows, [4.2, 4.6, 4.6, 5.2], ['L', 'R', 'R', 'R'])

    # KONFIGURASI VARIABEL UJI CROSSTAB BAB 2 (SUB-BAB 2.1, 2.2, 2.3)
    add_caption_compact(doc, "Tabel 2.4: Konfigurasi Variabel Uji Tabulasi Silang (Crosstab) Bab 2 — Skenario Sub-bab 2.1, 2.2, dan 2.3")
    cfg_headers = ["Komponen Uji", "2.1 Smelter vs IKA", "2.2 PLTU vs IKU", "2.3 Ekspansi Industri vs Deforestasi"]
    cfg_rows = [
        ["Variabel Independen (X)", "Jumlah Smelter (fasilitas beroperasi & konstruksi)", "Kapasitas PLTU (MW)", "Luas Ekspansi Industri / IUP & Kawasan (Ha)"],
        ["Variabel Dependen (Y)", "Indeks Kualitas Air (skor baku mutu air provinsi)", "Indeks Kualitas Udara (skor baku mutu udara ambien)", "Kehilangan Tutupan Pohon / Total Deforestasi Alam (Ha)"],
        ["Hipotesis Alternatif (H1)", "Hubungan negatif: semakin padat smelter, semakin kritis mutu air", "Hubungan negatif: semakin besar kapasitas PLTU, semakin kritis mutu udara", "Korelasi positif: ekspansi izin lahan mendorong laju deforestasi"],
        ["Decision Rule (Alpha 5%)", "Tolak H0 jika P-Value < 0.05", "Tolak H0 jika P-Value < 0.05", "Tolak H0 jika P-Value < 0.05"],
        ["Threshold Kategori (Median Panel)", "X >= 75.0 fasilitas; Y >= 55.9 poin (N=54)", "X >= 220.0 MW; Y >= 91.0 poin (N=54)", "X >= 138,148.8 Ha; Y >= 15,917.7 Ha (N=60)"],
        ["Orientasi Odds Ratio", "a = Smelter Tinggi & IKA Kritis (risiko IKA kritis)", "a = PLTU Tinggi & IKU Kritis (risiko IKU kritis)", "a = IUP Tinggi & Deforestasi Parah (risiko deforestasi parah)"]
    ]
    add_table_compact(doc, cfg_headers, cfg_rows, [3.2, 4.9, 5.0, 5.5], ['L', 'L', 'L', 'L'])

    # SINTESIS TABEL INFERENSIAL BAB 2
    add_caption_compact(doc, "Tabel 2.5: Ringkasan Hasil Uji Independensi Chi-Square (χ²) dan Odds Ratio (OR) Data Panel Bab 2 (N=54 s.d. 60)")
    inf_headers = ["Faktor Tekanan Lingkungan (X)", "Indikator Dampak (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan Ilmiah"]
    inf_rows = [
        ["Kepadatan Smelter (Fasilitas)", "Indeks Kualitas Air (IKA)", "2.667", "0.102", "0.35x", "TIDAK SIGNIFIKAN (Pengenceran Agregat)"],
        ["Kapasitas PLTU (MW)", "Indeks Kualitas Udara (IKU)", "0.000", "1.000", "1.18x", "TIDAK SIGNIFIKAN (Pengenceran Ambien)"],
        ["Luas Ekspansi Industri (Ha)", "Kehilangan Tutupan Pohon (Ha)", "35.267", "p < 0.001", "81.0x", "SIGNIFIKAN (Risiko Deforestasi 81x Lipat)"]
    ]
    add_table_compact(doc, inf_headers, inf_rows, [4.5, 4.2, 2.0, 1.8, 1.8, 4.3], ['L', 'L', 'C', 'C', 'C', 'L'])

    # Sesuai arahan user: Untuk Crosstab TIDAK USAH persamaan substitusi aritmatika
    add_math_block(
        doc,
        "Eksekusi Ruang & Konsentrasi Deforestasi Sentra",
        "Luas IUP & Kawasan Provinsi = Penjumlahan luas seluruh izin konsesi di provinsi tersebut  |  Deforestasi Kumulatif = Penjumlahan deforestasi tahunan 2014 s.d. 2023\n"
        "Chi-Square (χ²) = Jumlah dari [ ( Frekuensi Observasi - Frekuensi Harapan )² ÷ Frekuensi Harapan ]  |  Odds Ratio (OR) = ( a × d ) ÷ ( b × c )",
        "Gabungan Konsesi Sulteng + Sultra = 453.216 Ha + 446.025 Ha = 899.241 Hektar (75,87% dari konsesi se-Sulawesi)\n"
        "Gabungan Deforestasi Sulteng + Sultra = 481.908 Ha + 337.434 Ha = 819.342 Hektar (59,11% dari deforestasi se-Sulawesi)",
        "Uji Chi-Square membuktikan secara sangat signifikan (χ² = 35.267, p < 0.001, OR = 81.0x) bahwa wilayah konsesi industri nikel menghadapi risiko deforestasi 81 kali lipat lebih tinggi (lihat Tabel 2.5; konfigurasi variabel uji pada Tabel 2.4)."
    )

    # ═══════════════════════════════════════════════════════════
    # 2.4 DRIVER DEFORESTASI & ATRIBUSI EMISI CO2
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "2.4. Driver Deforestasi: Analisis Faktor Pendorong Perubahan Tutupan Hutan")
    add_p(doc, [
        ("Atribusi kausalitas membedah kontribusi faktor pendorong terhadap ", False, False),
        ("1,22+ juta hektar deforestasi di Sulawesi", True, False),
        (" (GFW 2014–2023) antara industri komoditas ekstraktif skala besar vs pertanian berpindah masyarakat.", False, False)
    ])
    add_note_inline(doc, "Sumber Data", "Global Forest Watch — University of Maryland (klasifikasi faktor pendorong deforestasi & estimasi pelepasan emisi CO₂ kumulatif 2014–2023) (diolah CELIOS).")

    add_caption_compact(doc, "Tabel 2.6: Matriks Atribusi Deforestasi dan Pelepasan Emisi CO₂ per Faktor Pendorong (Kumulatif 2014–2023)")
    driver_headers = ["Faktor Pendorong Deforestasi", "Total Deforestasi (Ha)", "Proporsi (%)", "Estimasi Emisi Karbon CO₂ (Mg)", "Proporsi Emisi (%)"]
    driver_rows = [
        ["Pertambangan dan Sawit (Ekstraktif)", "1,001,654", "82.2%", "664,472,885", "82.6%"],
        ["Kehutanan Komersial (Logging)", "134,637", "11.1%", "87,138,022", "10.8%"],
        ["Pertanian Berpindah (Masyarakat)", "55,905", "4.6%", "38,215,565", "4.8%"],
        ["Tidak Teridentifikasi / Lainnya", "25,738", "2.1%", "14,225,278", "1.8%"],
        ["Total Agregat Sulawesi", "1,217,934", "100.0%", "804,051,750", "100.0%"]
    ]
    add_table_compact(doc, driver_headers, driver_rows, [5.2, 3.0, 2.2, 5.0, 3.2], ['L', 'R', 'C', 'R', 'C'])

    add_math_block(
        doc,
        "Atribusi Deforestasi Komoditas & Pelepasan Karbon",
        "Proporsi Faktor Pendorong (%) = ( Deforestasi Faktor Tersebut ÷ Total Deforestasi Kumulatif ) × 100%\n"
        "Rasio Kerusakan = Deforestasi Tambang & Sawit ÷ Deforestasi Pertanian Rakyat",
        "Proporsi Tambang & Sawit = ( 1.001.654 Ha ÷ 1.217.934 Ha ) × 100% = 82,24% (Emisi: 664.472.885 Mg CO2 / 82,64%)\n"
        "Proporsi Pertanian Rakyat = ( 55.905 Ha ÷ 1.217.934 Ha ) × 100% = 4,59% (Emisi: 38.215.565 Mg CO2 / 4,75%)\n"
        "Rasio Kerusakan = 1.001.654 Ha ÷ 55.905 Ha = 17,92 Kali Lipat Lebih Masif",
        "Fakta empiris membantah tudingan deforestasi akibat perladangan rakyat: industri tambang & sawit merusak hutan 18 kali lipat lebih masif."
    )

    # ═══════════════════════════════════════════════════════════
    # 2.5 KEHANCURAN BIODIVERSITAS: HABITAT SATWA ENDEMIK
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "2.5. Kehancuran Biodiversitas: Dampak Terhadap Habitat Satwa Endemik")
    add_p(doc, [
        ("Ekspansi pertambangan nikel dan kawasan industri mengancam keanekaragaman hayati kawasan biogeografi Wallacea. Data spasial ", False, False),
        ("GBIF", True, False),
        (" memetakan ", False, False),
        ("269 titik koordinat perjumpaan (occurrence)", True, False),
        (" dari 7 spesies endemik kunci. Validasi ", False, False),
        ("IUCN Red List", True, False),
        (" menunjukkan seluruh spesies mengalami tren penurunan populasi (*Decreasing*) dan ", False, False),
        ("4 dari 7 spesies terkonfirmasi menghadapi Mining Threat", True, False),
        (".", False, False)
    ])
    add_note_inline(doc, "Sumber Data", "GBIF — Global Biodiversity Information Facility (titik koordinat perjumpaan satwa); IUCN Red List (status konservasi, tren populasi, dan penanda Mining Threat) (diolah CELIOS).")

    add_caption_compact(doc, "Tabel 2.7: Matriks Spesies Endemik Wallacea, Status IUCN, Penanda Mining Threat, dan Titik GBIF")
    bio_headers = ["Nama Ilmiah (Scientific Name)", "Nama Umum (Common Name)", "Status IUCN", "Tren Populasi", "Mining Threat", "Titik GBIF"]
    bio_rows = [
        ["Macaca nigra", "Celebes Crested Macaque", "Critically Endangered", "Decreasing", "Yes", "87"],
        ["Macrocephalon maleo", "Maleo", "Critically Endangered", "Decreasing", "No", "95"],
        ["Bubalus depressicornis", "Lowland Anoa", "Endangered", "Decreasing", "Yes", "18"],
        ["Bubalus quarlesi", "Mountain Anoa", "Endangered", "Decreasing", "Yes", "10"],
        ["Babyrousa celebensis", "Sulawesi Babirusa", "Vulnerable", "Decreasing", "Yes", "33"],
        ["Babyrousa babyrussa", "Hairy Babirusa", "Vulnerable", "Decreasing", "No", "14"],
        ["Tarsius tarsier", "Spectral Tarsier", "Vulnerable", "Decreasing", "No", "12"]
    ]
    add_table_compact(doc, bio_headers, bio_rows, [4.2, 4.2, 3.6, 2.4, 2.0, 2.2], ['L', 'L', 'C', 'C', 'C', 'C'])

    add_math_block(
        doc,
        "Keterancaman Spesies Wallacea & Penanda Mining Threat",
        "Proporsi Status Kritis (%) = ( Jumlah Spesies CR & EN ÷ Total Spesies ) × 100%  |  Proporsi Mining Threat (%) = ( Jumlah Spesies Terancam Tambang ÷ Total Spesies ) × 100%",
        "Titik Occurrence GBIF = 269 Titik Terverifikasi (Maleo 95, Macaca 87, Babirusa 47, Anoa 28, Tarsius 12)\n"
        "Status Konservasi Kritis = ( ( 2 Spesies CR + 2 Spesies EN ) ÷ 7 Spesies ) × 100% = 57,14% Sangat Terancam Punah\n"
        "Penanda Mining Threat = ( 4 Spesies ÷ 7 Spesies ) × 100% = 57,14% Beririsan Langsung dengan Konsesi IUP Nikel",
        "Spesies Anoa dan Babirusa terbukti secara empiris menghadapi ancaman kepunahan langsung akibat fragmentasi habitat konsesi pertambangan."
    )

    # Simpan File DOCX
    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab2    = out_dir_compact.parent.parent / "bab_2"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab2.mkdir(parents=True, exist_ok=True)

    docx_path_compact = out_dir_compact / "Metodologi_Bab2_Kualitas_Lingkungan_Compact.docx"
    docx_path_bab2    = out_dir_bab2 / "Metodologi_Bab2_Kualitas_Lingkungan_Compact.docx"

    doc.save(str(docx_path_compact))
    shutil.copyfile(docx_path_compact, docx_path_bab2)
    print(f"[OK] Berhasil menyimpan DOCX di: {docx_path_compact}")
    print(f"[OK] Salinan tersimpan di: {docx_path_bab2}")

# ── Generator Naskah Markdown Compact ───────────────────────
def generate_compact_markdown():
    print("[2/4] Menyusun naskah Markdown Bab 2 Versi Compact...")
    md_content = """# BAB II: METODOLOGI ANALISIS KUALITAS LINGKUNGAN DI KAWASAN SMELTER

Dokumen laporan metodologi ini menyajikan kerangka ilmiah, formulasi matematis, prosedur pengolahan data, dan pengujian statistik yang dioperasionalkan pada **Bab 2: Kualitas Lingkungan di Kawasan Smelter** dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi periode 2014–2024.

---

## 2.1. Dampak Limbah Tailing: Konsentrasi Smelter vs Indeks Kualitas Air (IKA)

Pengoperasian **778 fasilitas mega-smelter** yang didukung oleh kapasitas **9.825 MW PLTU Captive** meningkatkan intensitas emisi dan beban lingkungan di Pulau Sulawesi. Konversi tutupan hutan mencapai **1.001.654 Hektar**, estimasi timbulan limbah B3/tailing sebesar **20,9 Juta Ton per tahun**, dan rata-rata Indeks Kualitas Air (IKA) tahun 2024 sebesar **59,7**.

> **Sumber Data:** Kementerian ESDM & Center for Global Sustainability (fasilitas smelter); KLHK & BPS (Indeks Kualitas Air 2016–2024); laporan lembaga masyarakat sipil (proksi timbulan limbah B3/tailing) dan inventarisasi laporan pencemaran sungai-pesisir (diolah CELIOS).

##### Tabel 2.1: Rincian Empiris Konsentrasi Smelter, IKA, Limbah B3, dan Sungai Tercemar per Provinsi (2024)
| Provinsi | Smelter (Unit) | Skor IKA | Limbah B3 (Ton/Thn) | Sungai Tercemar | Daftar Sungai / Pesisir Terdampak |
| :--- | :---: | :---: | :---: | :---: | :--- |
| **Sulawesi Tengah** | 344 | 62.1 | 12,000,000 | 4 | Sungai Bahodopi, Laroenai, Morowali, Pesisir Fatufia |
| **Sulawesi Tenggara** | 262 | 65.3 | 7,700,000 | 3 | Sungai Lasolo, Sungai Lalindu, Sungai Konaweha |
| **Sulawesi Selatan** | 111 | 58.5 | 1,000,000 | 1 | Pesisir dan Sungai Bantaeng |
| **Sulawesi Barat** | 39 | 55.9 | 0 | 0 | Tidak teridentifikasi pembuangan tailing smelter |
| **Sulawesi Utara** | 15 | 58.2 | 0 | 0 | Tidak teridentifikasi pembuangan tailing smelter |
| **Gorontalo** | 7 | 58.1 | 0 | 0 | Tidak teridentifikasi pembuangan tailing smelter |

**Formulasi Matematis (Konsentrasi Smelter & Beban Limbah B3 Tailing):**
```text
Jumlah Smelter Provinsi = Penjumlahan seluruh fasilitas smelter di provinsi tersebut  |  Rata-rata IKA Provinsi = Jumlah skor IKA seluruh titik pantau ÷ Banyaknya titik pantau
Chi-Square (χ²) = Jumlah dari [ ( Frekuensi Observasi - Frekuensi Harapan )² ÷ Frekuensi Harapan ]  |  Odds Ratio (OR) = ( a × d ) ÷ ( b × c )
```
**Persamaan Substitusi:**
```text
Gabungan Smelter Sulteng + Sultra = 344 unit + 262 unit = 606 unit (77,89% dari total se-Sulawesi)
Estimasi Limbah B3 Tailing = 12.000.000 Ton/Thn (Sulteng) + 7.700.000 Ton/Thn (Sultra) = 19.700.000 Ton/Thn (94,26% dari total limbah B3)
```
*Hasil uji Chi-Square dan Odds Ratio data panel disajikan pada Tabel 2.5, dengan konfigurasi variabel uji pada Tabel 2.4. Kegagalan signifikansi statistik membuktikan Aggregate Dilution Bias (pencemaran fatal sungai industri Bahodopi & Lasolo terencerkan oleh stasiun pemantau sungai non-industri).*

---

## 2.2. Kepungan Asap: Kapasitas PLTU vs Indeks Kualitas Udara (IKU)

Sebanyak **9.825 MW PLTU Captive batu bara off-grid** beroperasi di kawasan hilirisasi nikel menyumbang polusi udara ambien dan gas buang NO₂.

> **Sumber Data:** Kementerian ESDM & Global Energy Monitor (kapasitas PLTU Captive & Grid PLN); KLHK (Indeks Kualitas Udara 2015–2024); pantauan emisi Satelit NASA TROPOMI (konsentrasi NO₂) (diolah CELIOS).

##### Tabel 2.2: Rincian Empiris Kapasitas PLTU (Captive & Grid), IKU, dan Konsentrasi NO₂ NASA (2024)
| Provinsi | Kapasitas PLTU Captive (MW) | PLTU Grid PLN (MW) | Total Daya (MW) | Skor IKU | NASA TROPOMI NO₂ (mol/m²) |
| :--- | :---: | :---: | :---: | :---: | :---: |
| **Sulawesi Tengah** | 9,365 | 0 | 9,365 | 92.9 | 6.50e-06 |
| **Sulawesi Tenggara** | 2,280 | 100 | 2,380 | 93.0 | 6.62e-06 |
| **Sulawesi Selatan** | 600 | 920 | 1,520 | 91.5 | 6.40e-06 |
| **Sulawesi Utara** | 0 | 220 | 220 | 93.4 | 4.09e-06 |
| **Gorontalo** | 0 | 100 | 100 | 93.5 | 3.76e-06 |
| **Sulawesi Barat** | 0 | 0 | 0 | 92.5 | 6.00e-06 |

**Formulasi Matematis (Kapasitas Energi PLTU & Parameter Kualitas Udara):**
```text
Kapasitas PLTU Provinsi = Penjumlahan kapasitas seluruh unit PLTU di provinsi tersebut  |  Rata-rata IKU = Jumlah skor IKU ÷ Banyaknya observasi provinsi-tahun
Chi-Square (χ²) = Jumlah dari [ ( Frekuensi Observasi - Frekuensi Harapan )² ÷ Frekuensi Harapan ]  |  Odds Ratio (OR) = ( a × d ) ÷ ( b × c )
```
**Persamaan Substitusi:**
```text
Total PLTU Captive 3 Sentra = 9.365 MW (Sulteng) + 2.280 MW (Sultra) + 600 MW (Sulsel) = 12.245 MW terpasang
```
*Hasil pengujian statistik tabulasi silang dirinci pada Tabel 2.5, dengan konfigurasi variabel uji pada Tabel 2.4. Ketidaksignifikanan membuktikan Efek Pengenceran Udara Ambien karena sensor IKU tersebar merata di hutan berudara bersih.*

---

## 2.3. Eksekusi Ruang: Ekspansi Kawasan Industri vs Tekanan Ekologis (Deforestasi)

Alokasi konsesi IUP dan Kawasan Industri mencakup **1.185.174 Hektar** di Sulawesi. Sepanjang dekade 2014–2023, data Global Forest Watch (GFW) merekam akumulasi kehilangan tutupan pohon sebesar **1.386.055 Hektar** (terbesar di Sulawesi Tengah dan Tenggara).

> **Sumber Data:** Kementerian ESDM — MODI/Minerbaone (izin konsesi IUP & kawasan industri); Global Forest Watch — University of Maryland (kehilangan tutupan pohon 2014–2023) (diolah CELIOS).

##### Tabel 2.3: Rincian Empiris Luas Konsesi IUP-Kawasan Industri dan Deforestasi Kumulatif (2014–2023)
| Provinsi | Luas IUP & Kawasan (Ha) | Konsesi Baru Kumulatif (Ha) | Deforestasi Kumulatif 1 Dekade (Ha) |
| :--- | :---: | :---: | :---: |
| **Sulawesi Tengah** | 453,216 | 387,124 | 481,908 |
| **Sulawesi Tenggara** | 446,025 | 212,717 | 337,434 |
| **Sulawesi Selatan** | 181,469 | 123,065 | 261,147 |
| **Sulawesi Utara** | 94,829 | 89,170 | 74,240 |
| **Gorontalo** | 5,212 | 5,212 | 98,063 |
| **Sulawesi Barat** | 4,424 | 2,163 | 133,263 |

##### Tabel 2.4: Konfigurasi Variabel Uji Tabulasi Silang (Crosstab) Bab 2 — Skenario Sub-bab 2.1, 2.2, dan 2.3
| Komponen Uji | 2.1 Smelter vs IKA | 2.2 PLTU vs IKU | 2.3 Ekspansi Industri vs Deforestasi |
| :--- | :--- | :--- | :--- |
| **Variabel Independen (X)** | Jumlah Smelter (fasilitas beroperasi & konstruksi) | Kapasitas PLTU (MW) | Luas Ekspansi Industri / IUP & Kawasan (Ha) |
| **Variabel Dependen (Y)** | Indeks Kualitas Air (skor baku mutu air provinsi) | Indeks Kualitas Udara (skor baku mutu udara ambien) | Kehilangan Tutupan Pohon / Total Deforestasi Alam (Ha) |
| **Hipotesis Alternatif (H1)** | Hubungan negatif: semakin padat smelter, semakin kritis mutu air | Hubungan negatif: semakin besar kapasitas PLTU, semakin kritis mutu udara | Korelasi positif: ekspansi izin lahan mendorong laju deforestasi |
| **Decision Rule (Alpha 5%)** | Tolak H0 jika P-Value < 0.05 | Tolak H0 jika P-Value < 0.05 | Tolak H0 jika P-Value < 0.05 |
| **Threshold Kategori (Median Panel)** | X ≥ 75.0 fasilitas; Y ≥ 55.9 poin (N=54) | X ≥ 220.0 MW; Y ≥ 91.0 poin (N=54) | X ≥ 138,148.8 Ha; Y ≥ 15,917.7 Ha (N=60) |
| **Orientasi Odds Ratio** | a = Smelter Tinggi & IKA Kritis (risiko IKA kritis) | a = PLTU Tinggi & IKU Kritis (risiko IKU kritis) | a = IUP Tinggi & Deforestasi Parah (risiko deforestasi parah) |

##### Tabel 2.5: Ringkasan Hasil Uji Independensi Chi-Square (χ²) dan Odds Ratio (OR) Data Panel Bab 2 (N=54 s.d. 60)
| Faktor Tekanan Lingkungan (X) | Indikator Dampak (Y) | Chi-Square (χ²) | P-Value | Odds Ratio | Kesimpulan Ilmiah |
| :--- | :--- | :---: | :---: | :---: | :--- |
| **Kepadatan Smelter (Fasilitas)** | Indeks Kualitas Air (IKA) | 2.667 | 0.102 | 0.35x | TIDAK SIGNIFIKAN (Pengenceran Agregat) |
| **Kapasitas PLTU (MW)** | Indeks Kualitas Udara (IKU) | 0.000 | 1.000 | 1.18x | TIDAK SIGNIFIKAN (Pengenceran Ambien) |
| **Luas Ekspansi Industri (Ha)** | Kehilangan Tutupan Pohon (Ha) | 35.267 | p < 0.001 | 81.0x | SIGNIFIKAN (Risiko Deforestasi 81x Lipat) |

**Formulasi Matematis (Eksekusi Ruang & Konsentrasi Deforestasi Sentra):**
```text
Luas IUP & Kawasan Provinsi = Penjumlahan luas seluruh izin konsesi di provinsi tersebut  |  Deforestasi Kumulatif = Penjumlahan deforestasi tahunan 2014 s.d. 2023
Chi-Square (χ²) = Jumlah dari [ ( Frekuensi Observasi - Frekuensi Harapan )² ÷ Frekuensi Harapan ]  |  Odds Ratio (OR) = ( a × d ) ÷ ( b × c )
```
**Persamaan Substitusi:**
```text
Gabungan Konsesi Sulteng + Sultra = 453.216 Ha + 446.025 Ha = 899.241 Hektar (75,87% dari konsesi se-Sulawesi)
Gabungan Deforestasi Sulteng + Sultra = 481.908 Ha + 337.434 Ha = 819.342 Hektar (59,11% dari deforestasi se-Sulawesi)
```
*Uji Chi-Square membuktikan secara sangat signifikan (χ² = 35.267, p < 0.001, OR = 81.0x) bahwa wilayah konsesi industri nikel menghadapi risiko deforestasi 81 kali lipat lebih tinggi (lihat Tabel 2.5; konfigurasi variabel uji pada Tabel 2.4).*

---

## 2.4. Driver Deforestasi: Analisis Faktor Pendorong Perubahan Tutupan Hutan

Atribusi kausalitas membedah kontribusi faktor pendorong terhadap **1,22+ juta hektar deforestasi di Sulawesi** (GFW 2014–2023) antara industri komoditas ekstraktif skala besar vs pertanian berpindah masyarakat.

> **Sumber Data:** Global Forest Watch — University of Maryland (klasifikasi faktor pendorong deforestasi & estimasi pelepasan emisi CO₂ kumulatif 2014–2023) (diolah CELIOS).

##### Tabel 2.6: Matriks Atribusi Deforestasi dan Pelepasan Emisi CO₂ per Faktor Pendorong (Kumulatif 2014–2023)
| Faktor Pendorong Deforestasi | Total Deforestasi (Ha) | Proporsi (%) | Estimasi Emisi Karbon CO₂ (Mg) | Proporsi Emisi (%) |
| :--- | :---: | :---: | :---: | :---: |
| **Pertambangan dan Sawit (Ekstraktif)** | 1,001,654 | 82.2% | 664,472,885 | 82.6% |
| **Kehutanan Komersial (Logging)** | 134,637 | 11.1% | 87,138,022 | 10.8% |
| **Pertanian Berpindah (Masyarakat)** | 55,905 | 4.6% | 38,215,565 | 4.8% |
| **Tidak Teridentifikasi / Lainnya** | 25,738 | 2.1% | 14,225,278 | 1.8% |
| **Total Agregat Sulawesi** | **1,217,934** | **100.0%** | **804,051,750** | **100.0%** |

**Formulasi Matematis (Atribusi Deforestasi Komoditas & Pelepasan Karbon):**
```text
Proporsi Faktor Pendorong (%) = ( Deforestasi Faktor Tersebut ÷ Total Deforestasi Kumulatif ) × 100%
Rasio Kerusakan = Deforestasi Tambang & Sawit ÷ Deforestasi Pertanian Rakyat
```
**Persamaan Substitusi:**
```text
Proporsi Tambang & Sawit = ( 1.001.654 Ha ÷ 1.217.934 Ha ) × 100% = 82,24% (Emisi: 664.472.885 Mg CO2 / 82,64%)
Proporsi Pertanian Rakyat = ( 55.905 Ha ÷ 1.217.934 Ha ) × 100% = 4,59% (Emisi: 38.215.565 Mg CO2 / 4,75%)
Rasio Kerusakan = 1.001.654 Ha ÷ 55.905 Ha = 17,92 Kali Lipat Lebih Masif
```
*Fakta empiris membantah tudingan deforestasi akibat perladangan rakyat: industri tambang & sawit merusak hutan 18 kali lipat lebih masif.*

---

## 2.5. Kehancuran Biodiversitas: Dampak Terhadap Habitat Satwa Endemik

Ekspansi pertambangan nikel dan kawasan industri mengancam keanekaragaman hayati kawasan biogeografi Wallacea. Data spasial **GBIF** memetakan **269 titik koordinat perjumpaan (*occurrence*)** dari 7 spesies endemik kunci. Validasi **IUCN Red List** menunjukkan seluruh spesies mengalami tren penurunan populasi (*Decreasing*) dan **4 dari 7 spesies terkonfirmasi menghadapi Mining Threat**.

> **Sumber Data:** GBIF — Global Biodiversity Information Facility (titik koordinat perjumpaan satwa); IUCN Red List (status konservasi, tren populasi, dan penanda Mining Threat) (diolah CELIOS).

##### Tabel 2.7: Matriks Spesies Endemik Wallacea, Status IUCN, Penanda Mining Threat, dan Titik GBIF
| Nama Ilmiah (Scientific Name) | Nama Umum (Common Name) | Status IUCN | Tren Populasi | Mining Threat | Titik GBIF |
| :--- | :--- | :---: | :---: | :---: | :---: |
| **Macaca nigra** | Celebes Crested Macaque | Critically Endangered | Decreasing | Yes | 87 |
| **Macrocephalon maleo** | Maleo | Critically Endangered | Decreasing | No | 95 |
| **Bubalus depressicornis** | Lowland Anoa | Endangered | Decreasing | Yes | 18 |
| **Bubalus quarlesi** | Mountain Anoa | Endangered | Decreasing | Yes | 10 |
| **Babyrousa celebensis** | Sulawesi Babirusa | Vulnerable | Decreasing | Yes | 33 |
| **Babyrousa babyrussa** | Hairy Babirusa | Vulnerable | Decreasing | No | 14 |
| **Tarsius tarsier** | Spectral Tarsier | Vulnerable | Decreasing | No | 12 |

**Formulasi Matematis (Keterancaman Spesies Wallacea & Penanda Mining Threat):**
```text
Proporsi Status Kritis (%) = ( Jumlah Spesies CR & EN ÷ Total Spesies ) × 100%  |  Proporsi Mining Threat (%) = ( Jumlah Spesies Terancam Tambang ÷ Total Spesies ) × 100%
```
**Persamaan Substitusi:**
```text
Titik Occurrence GBIF = 269 Titik Terverifikasi (Maleo 95, Macaca 87, Babirusa 47, Anoa 28, Tarsius 12)
Status Konservasi Kritis = ( ( 2 Spesies CR + 2 Spesies EN ) ÷ 7 Spesies ) × 100% = 57,14% Sangat Terancam Punah
Penanda Mining Threat = ( 4 Spesies ÷ 7 Spesies ) × 100% = 57,14% Beririsan Langsung dengan Konsesi IUP Nikel
```
*Spesies Anoa dan Babirusa terbukti secara empiris menghadapi ancaman kepunahan langsung akibat fragmentasi habitat konsesi pertambangan.*
"""

    out_dir_compact = Path(__file__).resolve().parent
    out_dir_bab2    = out_dir_compact.parent.parent / "bab_2"
    out_dir_compact.mkdir(parents=True, exist_ok=True)
    out_dir_bab2.mkdir(parents=True, exist_ok=True)
    
    md_path_compact = out_dir_compact / "Metodologi_Bab2_Kualitas_Lingkungan_Compact.md"
    md_path_bab2    = out_dir_bab2 / "Metodologi_Bab2_Kualitas_Lingkungan_Compact.md"

    with open(md_path_compact, "w", encoding="utf-8") as f:
        f.write(md_content)
    with open(md_path_bab2, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(f"[OK] Berhasil menyimpan Markdown di: {md_path_compact}")
    print(f"[OK] Salinan tersimpan di: {md_path_bab2}")

if __name__ == "__main__":
    print("=" * 70)
    print("GENERATOR METODOLOGI STATISTIK VERSI COMPACT - BAB 2")
    print("=" * 70)
    build_compact_report()
    generate_compact_markdown()
    print("=" * 70)
    print("SELESAI! Dokumen Versi Compact Bab 2 berhasil diperbarui.")
    print("=" * 70)
