# ==============================================================================
# BAB 6 SUB-BAB 6.6: GENERATOR METODOLOGI SKORING D3TLH TINGKAT PROVINSI
# DOKUMEN SATUAN KOMPILASI TINGKAT PROVINSI (TERPUSAT & MODULAR)
# CAKUPAN SEKSI:
#   6.6   : KERANGKA METODOLOGI & FORMULASI MATEMATIS UNIVERSAL SE-SULAWESI
#   6.6.1 : EVALUASI EMPIRIS PROVINSI SULAWESI TENGAH (SULTENG)
#   6.6.2 : EVALUASI EMPIRIS PROVINSI SULAWESI TENGGARA (SULTRA)
# METODE: HYBRID Z-SCORE ANOMALI STANDAR DEVIASI + ENTROPY WEIGHT METHOD (EWM)
# SINKRONISASI: 100% PERSIS TAB 3 STREAMLIT PAGE 6 (pages/6_Audit_D3TLH.py)
# ==============================================================================

import sys
import os
import urllib.request
import base64
from pathlib import Path

# Setup paths
current_file = Path(__file__).resolve()
project_root = current_file.parents[3]
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

import tools.algo_skoring_provinsi_ZscoreEWM.kalkulasi_provinsi_sulawesi as algo_prov_mod

# -----------------------------------------------------------------------------
# PALET WARNA RESMI CELIOS (SOP KORPORAT & LAPORAN AKADEMIK)
# -----------------------------------------------------------------------------
C_NAVY = RGBColor(0x14, 0x36, 0x42)
C_DARK = RGBColor(0x20, 0x24, 0x26)
C_RED = RGBColor(0xA8, 0x20, 0x1A)
C_TEAL = RGBColor(0x0F, 0x4C, 0x5C)
C_AMBER = RGBColor(0xE3, 0x64, 0x14)
C_MUTED = RGBColor(0x55, 0x55, 0x55)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

C_HEX_PRIMARY = "143642"
C_HEX_LIGHT_BG = "F4F5F6"
C_HEX_RED = "A8201A"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_table_borders(table, border_color="D3D3D3"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{C_HEX_PRIMARY}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = C_NAVY
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = C_RED
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = C_TEAL
    return p

def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = C_NAVY
    return p

def add_p(doc, runs, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for text, bold, italic in runs:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(9.0)
        r.font.color.rgb = C_DARK
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = C_NAVY
    return p

def add_note_box(doc, title, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    c.width = Cm(17.0)
    set_cell_shading(c, C_HEX_LIGHT_BG)
    set_cell_margins(c, top=120, bottom=120, left=160, right=160)
    tcPr = c._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{C_HEX_PRIMARY}"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(f"{title}\n")
    r_t.bold = True
    r_t.font.size = Pt(8.5)
    r_t.font.color.rgb = C_NAVY
    r_b = p.add_run(body)
    r_b.font.size = Pt(8.0)
    r_b.font.color.rgb = C_DARK
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(4)

def add_formula_box(doc, title, formula_text, variable_tuples):
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(5)
    p_t.paragraph_format.space_after = Pt(2)
    p_t.paragraph_format.keep_with_next = True
    r_t = p_t.add_run(f"• {title}")
    r_t.bold = True
    r_t.font.size = Pt(9.0)
    r_t.font.color.rgb = C_TEAL

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    c.width = Cm(17.0)
    set_cell_shading(c, "F8F9FA")
    set_cell_margins(c, top=80, bottom=80, left=120, right=120)
    tcPr = c._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="16" w:space="0" w:color="{C_HEX_PRIMARY}"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(formula_text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.bold = True
    r.font.color.rgb = C_NAVY

    if variable_tuples:
        p_v = doc.add_paragraph()
        p_v.paragraph_format.space_before = Pt(2)
        p_v.paragraph_format.space_after = Pt(5)
        p_v.paragraph_format.left_indent = Cm(0.5)
        p_v.paragraph_format.line_spacing = 1.1
        for var_name, var_desc in variable_tuples:
            r1 = p_v.add_run(f"  - {var_name}: ")
            r1.bold = True
            r1.font.size = Pt(8.0)
            r1.font.color.rgb = C_DARK
            r2 = p_v.add_run(f"{var_desc}\n")
            r2.font.size = Pt(8.0)
            r2.font.color.rgb = C_MUTED

def add_table_styled(doc, headers, rows, col_widths, col_alignments):
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    hdr_cells = tbl.rows[0].cells
    for i, h_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Cm(col_widths[i])
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        set_cell_shading(cell, C_HEX_PRIMARY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8.0)
        r.font.color.rgb = C_WHITE

    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = C_HEX_LIGHT_BG if r_idx % 2 == 1 else "FFFFFF"
        is_total = (str(row_data[0]).strip().upper() == "TOTAL" or "SKOR KOMPOSIT" in str(row_data[0]).upper())
        if is_total:
            bg = "E6ECF0"

        for c_idx, val in enumerate(row_data):
            cell = row_cells[c_idx]
            cell.width = Cm(col_widths[c_idx])
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            align = col_alignments[c_idx]
            if align == "C":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "R":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            r = p.add_run(str(val))
            r.font.size = Pt(7.5 if len(headers) > 6 else 8.0)
            r.font.color.rgb = C_DARK
            if is_total:
                r.bold = True

    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(6)

def download_mermaid_png(mermaid_str, output_path):
    try:
        b64 = base64.b64encode(mermaid_str.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{b64}"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=12) as resp, open(output_path, 'wb') as f:
            f.write(resp.read())
        return True
    except Exception as e:
        print(f"[WARN] Gagal download Mermaid: {e}")
        return False

def get_likert_label(skor_likert):
    val = round(float(skor_likert))
    if val >= 4:
        return "Melampaui Batas"
    elif val == 3:
        return "Mendekati Batas"
    else:
        return "Tidak Melampaui Batas"

def html_table(headers, rows):
    out = ['<div class="table-container"><table>']
    out.append('<thead><tr>')
    for h in headers:
        out.append(f'<th class="data-th">{h}</th>')
    out.append('</tr></thead><tbody>')
    for r in rows:
        is_total = (str(r[0]).strip().upper() == "TOTAL" or "SKOR KOMPOSIT" in str(r[0]).upper())
        row_cls = ' class="total-row"' if is_total else ''
        out.append(f'<tr{row_cls}>')
        for idx, val in enumerate(r):
            val_str = str(val)
            align_cls = 'text-center'
            if idx in [1, 4] and len(headers) == 5:
                align_cls = 'text-left'
            elif idx == 1:
                align_cls = 'text-left'
            elif idx in [2, 3]:
                align_cls = 'text-right'

            if "Melampaui Batas" in val_str:
                val_str = f'<span class="badge-danger">{val_str}</span>'
            elif "Mendekati Batas" in val_str:
                val_str = f'<span class="badge-warning">{val_str}</span>'
            elif "Tidak Melampaui Batas" in val_str:
                val_str = f'<span class="badge-success">{val_str}</span>'
            out.append(f'<td class="{align_cls}">{val_str}</td>')
        out.append('</tr>')
    out.append('</tbody></table></div>')
    return "\n".join(out)

def markdown_table(headers, rows):
    lines = []
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join([":---" if i in [1, len(headers)-1] else ":---:" for i in range(len(headers))]) + " |")
    for r in rows:
        lines.append("| " + " | ".join([str(x) for x in r]) + " |")
    return "\n".join(lines)


# =============================================================================
# METADATA INDIKATOR EMPIRIS (20 INDIKATOR KUNCI MULTIKRITERIA)
# =============================================================================
INDICATOR_META = {
    'pltu_mw': ('Pilar Udara', 'Kapasitas PLTU Captive Beroperasi', 'MW'),
    'no2': ('Pilar Udara', 'Konsentrasi Gas NO2 Troposferik Satelit', 'mol/m²'),
    'ispa_irr': ('Pilar Udara', 'Morbiditas ISPA (Incidence Rate Ratio)', 'x lipat'),
    'b3_ton': ('Pilar Udara', 'Proporsi Timbulan Limbah B3 Industri', 'Jt Ton'),
    'co2_mton': ('Pilar Udara', 'Pelepasan Emisi Karbon Deforestasi GFW', 'Jt Ton CO2e'),
    'ika': ('Pilar Air', 'Indeks Kualitas Air (IKA) Terkini', 'Poin'),
    'diare_irr': ('Pilar Air', 'Morbiditas Diare (Incidence Rate Ratio)', 'x lipat'),
    'konflik_pesisir': ('Pilar Air', 'Konflik Ruang Laut Nelayan vs Tambang', 'Kasus'),
    'tailing_ton': ('Pilar Air', 'Akumulasi Beban Tailing, Slag & DSTP', 'Jt Ton'),
    'bencana': ('Pilar Lahan', 'Bencana Hidrometeorologi (Banjir & Longsor)', 'Kejadian'),
    'deforestasi_ha': ('Pilar Lahan', 'Deforestasi Hutan Alam Primer GFW', 'Ha'),
    'lindung_ha': ('Pilar Lahan', 'Perambahan Tambang di Kawasan Hutan Lindung', 'Ha'),
    'driver_ha': ('Pilar Lahan', 'Aktor Deforestasi Komoditas Tambang & Sawit', 'Ha'),
    'kepadatan_iup': ('Pilar Lahan', 'Kepadatan Konsesi IUP Nikel vs Daratan', '% Daratan'),
    'fpic': ('Pilar Sosial', 'Manipulasi Persetujuan Konsultasi Warga (FPIC)', 'Kasus'),
    'jiwa_terdampak': ('Pilar Sosial', 'Korban Perampasan Ruang Hidup & Krisis Agraria', 'Jiwa'),
    'kriminalisasi': ('Pilar Sosial', 'Insiden Kriminalisasi Warga & Pembela HAM', 'Insiden'),
    'gap_spa': ('Pilar Sosial', 'Defisit Kelayakan Standar Faskes SPA', '% Gap'),
    'izin_baru': ('Pilar Veto', 'Penerbitan Obral Konsesi WIUP Baru Pasca-2014', 'Izin'),
    'ilegal': ('Pilar Veto', 'Korporat Tambang Pelanggar Hukum Beroperasi Ilegal', 'Korporasi')
}

def extract_prov_eval_rows(prov_data):
    """
    Mengekstrak baris tabel evaluasi empiris spesifik provinsi:
    Columns: [Pilar, Indikator Empiris, Fakta Mentah A, Z-Score, Bobot EWM, Likert, Status Ekologis]
    """
    raw_absolut = prov_data['raw_absolut']
    raw_zscores = prov_data['raw_zscores']
    likert_dict = prov_data['math_details']['likert']
    ewm_weights = prov_data['math_details']['ewm_weights']

    table_rows = []
    for col, (pilar, nama_ind, satuan) in INDICATOR_META.items():
        val_a = raw_absolut.get(col, 0.0)
        z_val = raw_zscores.get(col, 0.0)
        w_val = ewm_weights.get(col, 0.0)
        l_val = likert_dict.get(col, 0.0)
        label_ekologis = get_likert_label(l_val)

        if col == 'no2':
            str_a = f"{val_a:.2e} {satuan}"
        elif col in ['pltu_mw', 'deforestasi_ha', 'lindung_ha', 'driver_ha', 'jiwa_terdampak']:
            str_a = f"{val_a:,.0f} {satuan}"
        elif col in ['b3_ton', 'co2_mton', 'tailing_ton']:
            str_a = f"{val_a:,.2f} {satuan}"
        elif col == 'kepadatan_iup':
            str_a = f"{val_a*100:.2f}%"
        elif col in ['ispa_irr', 'diare_irr']:
            str_a = f"{val_a:.2f}x lipat"
        else:
            str_a = f"{val_a:,.1f} {satuan}" if isinstance(val_a, float) else f"{val_a} {satuan}"

        table_rows.append([
            pilar,
            nama_ind,
            str_a,
            f"{z_val:+.2f}σ",
            f"{w_val:.4f}",
            f"{l_val:.1f} / 5",
            label_ekologis
        ])
    return table_rows


# =============================================================================
# EKSEKUTOR UTAMA PEMBANGUN DOKUMEN SATUAN TINGKAT PROVINSI
# =============================================================================
def generate_bab6_skoring_provinsi():
    print("[1/4] Mengekstraksi dataset empiris seluruh provinsi (Model Z-Score EWM)...")
    tool_dir = Path(__file__).parent
    
    # 1. Jalankan Engine Algoritma Resmi ZscoreEWM
    all_prov_results = algo_prov_mod.kalkulasi_skor_provinsi_sulawesi()
    sulteng = all_prov_results['Sulawesi Tengah']
    sultra = all_prov_results['Sulawesi Tenggara']

    # 2. Hitung Matriks Parameter Regional Se-Sulawesi (Universal Table 6.12)
    df_prov = algo_prov_mod.pd.DataFrame.from_dict(algo_prov_mod.RAW_DATA, orient='index')
    means = df_prov.mean(axis=0)
    stds = df_prov.std(axis=0)
    stds[stds == 0] = 1.0

    n, m = df_prov.shape
    min_vals = df_prov.min(axis=0)
    max_vals = df_prov.max(axis=0)
    range_vals = max_vals - min_vals
    range_vals[range_vals == 0] = 1.0
    r_matrix = (df_prov - min_vals) / range_vals
    sum_r = r_matrix.sum(axis=0)
    sum_r[sum_r == 0] = 1.0
    p_matrix = r_matrix / sum_r
    k_entropy = 1.0 / algo_prov_mod.np.log(n)
    p_log_p = p_matrix * algo_prov_mod.np.log(p_matrix.replace(0, 1e-12))
    e_j = -k_entropy * p_log_p.sum(axis=0)
    d_j = 1.0 - e_j
    w_j = d_j / d_j.sum()

    table_regional_meta = []
    for col, (pilar, nama_ind, satuan) in INDICATOR_META.items():
        m_val = means.get(col, 0.0)
        s_val = stds.get(col, 1.0)
        ej_val = e_j.get(col, 0.0)
        dj_val = d_j.get(col, 0.0)
        wj_val = w_j.get(col, 0.0)

        if col == 'no2':
            str_m = f"{m_val:.2e}"
            str_s = f"{s_val:.2e}"
        elif col in ['pltu_mw', 'deforestasi_ha', 'lindung_ha', 'driver_ha', 'jiwa_terdampak']:
            str_m = f"{m_val:,.0f} {satuan}"
            str_s = f"{s_val:,.0f} {satuan}"
        elif col in ['b3_ton', 'co2_mton', 'tailing_ton']:
            str_m = f"{m_val:,.2f} {satuan}"
            str_s = f"{s_val:,.2f} {satuan}"
        elif col == 'kepadatan_iup':
            str_m = f"{m_val*100:.2f}%"
            str_s = f"{s_val*100:.2f}%"
        elif col in ['ispa_irr', 'diare_irr']:
            str_m = f"{m_val:.2f}x"
            str_s = f"{s_val:.2f}x"
        else:
            str_m = f"{m_val:,.1f} {satuan}" if isinstance(m_val, float) else f"{m_val} {satuan}"
            str_s = f"{s_val:,.1f} {satuan}" if isinstance(s_val, float) else f"{s_val} {satuan}"

        table_regional_meta.append([
            pilar,
            nama_ind,
            str_m,
            str_s,
            f"{ej_val:.4f}",
            f"{dj_val:.4f}",
            f"{wj_val:.4f}"
        ])

    # 3. Ekstrak Data Empiris Sulteng (6.6.1)
    table_eval_sulteng = extract_prov_eval_rows(sulteng)
    rekap_sulteng = [
        ["Pilar 1: Udara", "PLTU (7.325 MW), NO2 (6.5e-6), ISPA (3.5x), B3 (25.3 Jt Ton), CO2 (291 Jt Ton)", f"{sulteng['udara']:.1f} / 5", get_likert_label(sulteng['udara']), "Episentrum PLTU Captive Terbesar & Konsentrasi Limbah B3"],
        ["Pilar 2: Air", "IKA (62.07), Diare (1.52x), Tailing (24.5 Jt Ton), Toksisitas Cr6+", f"{sulteng['air']:.1f} / 5", get_likert_label(sulteng['air']), "Beban Tailing Ekstrem & Morbiditas Penyakit Pencernaan"],
        ["Pilar 3: Lahan", "Bencana (458), Deforestasi (481k Ha), Lindung (19.8k Ha), Driver (383k Ha)", f"{sulteng['lahan']:.1f} / 5", get_likert_label(sulteng['lahan']), "Deforestasi Primer Masif & Perambahan Kawasan Lindung"],
        ["Pilar 4: Sosial", "FPIC (1 Kasus), Korban (12.231 Jiwa), Kriminalisasi (6 Insiden), Defisit SPA", f"{sulteng['sosial']:.1f} / 5", get_likert_label(sulteng['sosial']), "Kriminalisasi Warga Pembela HAM & Hilangnya Ruang Hidup"],
        ["Pilar 5: Veto", "Obral Izin (260 IUP Baru), Korporat Ilegal (3 Perusahaan), PLTU Ekspansi", f"{sulteng['veto']:.1f} / 5", get_likert_label(sulteng['veto']), "Kegagalan Pengendalian Izin & Impunitas Pelanggaran"],
        ["SKOR KOMPOSIT SULTENG", "Agregasi 5 Pilar EWM Weighted Average (Z-Score Standardization)", f"{sulteng['total_likert']:.1f} / 5", sulteng['likert_label'], "STATUS RED ALERT: DARURAT DAYA DUKUNG LINGKUNGAN"]
    ]

    # 4. Ekstrak Data Empiris Sultra (6.6.2)
    table_eval_sultra = extract_prov_eval_rows(sultra)
    rekap_sultra = [
        ["Pilar 1: Udara", "PLTU (1.900 MW), NO2 (6.6e-6), ISPA (0.91x), B3 (6.5 Jt Ton), CO2 (189 Jt Ton)", f"{sultra['udara']:.1f} / 5", get_likert_label(sultra['udara']), "Emisi Smelter Morosi/VDNI & Peningkatan NO2 Satelit"],
        ["Pilar 2: Air", "IKA (65.32), Diare (1.11x), Tailing (6.5 Jt Ton), Konflik Pesisir (5 Kasus)", f"{sultra['air']:.1f} / 5", get_likert_label(sultra['air']), "Sedimentasi Tailing Pesisir & Konflik Ruang Tangkap Nelayan"],
        ["Pilar 3: Lahan", "Bencana (158), Deforestasi (337k Ha), Lindung (8.2k Ha), Kepadatan IUP (11.72%)", f"{sultra['lahan']:.1f} / 5", get_likert_label(sultra['lahan']), "Outlier Kepadatan Konsesi Tambang Nikel Terpadat Se-Sulawesi"],
        ["Pilar 4: Sosial", "FPIC (5 Kasus), Korban (39.821 Jiwa), Kriminalisasi (4 Insiden), Gap SPA (17.9%)", f"{sultra['sosial']:.1f} / 5", get_likert_label(sultra['sosial']), "Krisis Kemanusiaan & Perampasan Ruang Hidup Terparah Se-Sulawesi"],
        ["Pilar 5: Veto", "Obral Izin (160 IUP Baru), Korporat Ilegal (1 Perusahaan), Ekspansi Smelter", f"{sultra['veto']:.1f} / 5", get_likert_label(sultra['veto']), "Obral Izin Tambang Nikel Baru Pasca-2014 di Pesisir & Pulau"],
        ["SKOR KOMPOSIT SULTRA", "Agregasi 5 Pilar EWM Weighted Average (Z-Score Standardization)", f"{sultra['total_likert']:.1f} / 5", sultra['likert_label'], "STATUS AMBANG BATAS: KRISIS RUANG HIDUP & KEPADATAN IUP"]
    ]

    # 5. Ekstrak Data Empiris Sulsel (6.6.3)
    sulsel = all_prov_results['Sulawesi Selatan']
    table_eval_sulsel = extract_prov_eval_rows(sulsel)
    rekap_sulsel = [
        ["Pilar 1: Udara", "PLTU (600 MW), NO2 (6.4e-6), ISPA (0.39x), B3 (1.0 Jt Ton), CO2 (139 Jt Ton)", f"{sulsel['udara']:.1f} / 5", get_likert_label(sulsel['udara']), "Konsentrasi NO2 Satelit & Emisi PLTU Jeneponto/Barru"],
        ["Pilar 2: Air", "IKA (58.50), Diare (0.91x), Tailing (1.0 Jt Ton), Cr6+ (1.0), Konflik Laut (7 Kasus)", f"{sulsel['air']:.1f} / 5", get_likert_label(sulsel['air']), "Toksisitas Logam Berat Cr6+ & Konflik Ruang Tangkap Nelayan Terbanyak"],
        ["Pilar 3: Lahan", "Bencana (669 Kejadian), Deforestasi (261k Ha), Lindung (5.3k Ha), IUP (3.88%)", f"{sulsel['lahan']:.1f} / 5", get_likert_label(sulsel['lahan']), "Kerentanan Hidrometeorologi Terparah Se-Sulawesi & Banjir Bandang DAS"],
        ["Pilar 4: Sosial", "FPIC (0 Kasus), Korban (2.257 Jiwa), Kriminalisasi (9 Insiden), Gap SPA (12.3%)", f"{sulsel['sosial']:.1f} / 5", get_likert_label(sulsel['sosial']), "Angka Kriminalisasi Petani & Pembela Lingkungan Tertinggi Se-Sulawesi"],
        ["Pilar 5: Veto", "Obral Izin (105 IUP Baru), Korporat Ilegal (10 Perusahaan), Pengawasan Lemah", f"{sulsel['veto']:.1f} / 5", get_likert_label(sulsel['veto']), "Aktivitas Tambang Ilegal Terbanyak Se-Sulawesi di Kawasan DAS & Hutan"],
        ["SKOR KOMPOSIT SULSEL", "Agregasi 5 Pilar EWM Weighted Average (Z-Score Standardization)", f"{sulsel['total_likert']:.1f} / 5", sulsel['likert_label'], "STATUS AMBANG BATAS: OUTLIER BENCANA, KRIMINALISASI & CR6+"]
    ]

    # 6. Ekstrak Data Empiris Sulbar (6.6.4)
    sulbar = all_prov_results['Sulawesi Barat']
    table_eval_sulbar = extract_prov_eval_rows(sulbar)
    rekap_sulbar = [
        ["Pilar 1: Udara", "PLTU (0 MW), NO2 (6.0e-6), ISPA (0.77x), B3 (0 Jt Ton), CO2 (82.5 Jt Ton)", f"{sulbar['udara']:.1f} / 5", get_likert_label(sulbar['udara']), "Bebas Polusi PLTU Captive & Nihil Timbulan Limbah B3 Smelter"],
        ["Pilar 2: Air", "IKA (55.93), Diare (1.27x), Tailing (0 Jt Ton), Nihil Konflik Laut", f"{sulbar['air']:.1f} / 5", get_likert_label(sulbar['air']), "Penurunan Mutu Air Sungai Akibat Limbah PKS Monokultur Sawit & Erosi"],
        ["Pilar 3: Lahan", "Bencana (143 Kejadian), Deforestasi (133k Ha), Lindung (1.2k Ha), IUP (0.26%)", f"{sulbar['lahan']:.1f} / 5", get_likert_label(sulbar['lahan']), "Kepadatan Konsesi Tambang Nikel Terendah Se-Sulawesi (Hanya 0,26%)"],
        ["Pilar 4: Sosial", "FPIC (0 Kasus), Korban (1 Jiwa), Kriminalisasi (1 Insiden), Gap SPA (0%)", f"{sulbar['sosial']:.1f} / 5", get_likert_label(sulbar['sosial']), "Relatif Minim Konflik Tambang Ekstraktif Nikel Skala Masif"],
        ["Pilar 5: Veto", "Obral Izin (27 IUP Baru), Nihil Korporat Ilegal Teridentifikasi", f"{sulbar['veto']:.1f} / 5", get_likert_label(sulbar['veto']), "Aktivitas Perizinan Tambang Nikel Terbatas di Kawasan Pesisir"],
        ["SKOR KOMPOSIT SULBAR", "Agregasi 5 Pilar EWM Weighted Average (Z-Score Standardization)", f"{sulbar['total_likert']:.1f} / 5", sulbar['likert_label'], "STATUS TERJAGA: BIOREGION NON-HILIRISASI NIKEL"]
    ]

    # 7. Ekstrak Data Empiris Gorontalo (6.6.5)
    gorontalo = all_prov_results['Gorontalo']
    table_eval_gorontalo = extract_prov_eval_rows(gorontalo)
    rekap_gorontalo = [
        ["Pilar 1: Udara", "PLTU (0 MW), NO2 (3.8e-6), ISPA (2.4x), B3 (0 Jt Ton), CO2 (53.7 Jt Ton)", f"{gorontalo['udara']:.1f} / 5", get_likert_label(gorontalo['udara']), "Konsentrasi NO2 Satelit Terbersih Se-Sulawesi & Nihil PLTU Captive"],
        ["Pilar 2: Air", "IKA (58.14), Diare (0.98x), Tailing (0 Jt Ton), Nihil Konflik Laut", f"{gorontalo['air']:.1f} / 5", get_likert_label(gorontalo['air']), "Bebas Tailing Tambang Nikel, Tekanan Sedimen di Danau Limboto"],
        ["Pilar 3: Lahan", "Bencana (0 Kejadian), Deforestasi (98k Ha), Lindung (2.0k Ha), IUP (0.46%)", f"{gorontalo['lahan']:.1f} / 5", get_likert_label(gorontalo['lahan']), "Deforestasi Primer & Emisi Karbon Terendah Se-Sulawesi"],
        ["Pilar 4: Sosial", "FPIC (0 Kasus), Korban (0 Jiwa), Kriminalisasi (0 Insiden), Gap SPA (0%)", f"{gorontalo['sosial']:.1f} / 5", get_likert_label(gorontalo['sosial']), "Bebas Konflik Perampasan Ruang Hidup Skala Masif Tambang Nikel"],
        ["Pilar 5: Veto", "Obral Izin (7 IUP Baru), Korporat Ilegal (1 Perusahaan), Obral Terendah", f"{gorontalo['veto']:.1f} / 5", get_likert_label(gorontalo['veto']), "Penerbitan IUP Tambang Terendah Se-Sulawesi (Hanya 7 IUP Pasca-2014)"],
        ["SKOR KOMPOSIT GORONTALO", "Agregasi 5 Pilar EWM Weighted Average (Z-Score Standardization)", f"{gorontalo['total_likert']:.1f} / 5", gorontalo['likert_label'], "STATUS TERJAGA: EMISI NO2 & TEKANAN LAHAN TERENDAH"]
    ]

    # Flowchart Mermaid LR Regional (Satu Flowchart Terpusat untuk Seluruh Provinsi)
    mermaid_regional = """flowchart LR
    subgraph S1["1. Matriks Empiris Regional"]
        A1["20 Indikator Empiris Multisektor<br/><i>6 Provinsi Se-Pulau Sulawesi</i>"]
        A2["Sensor NASA, KLHK, ESDM,<br/><i>Kemenkes, GFW, BNPB & KPA</i>"]
    end
    subgraph S2["2. Standardisasi & Pembobotan Objektif"]
        B1["Z-Score Regional: Z = (x - mean) / std<br/><i>Inversi Indikator Positif (IKA)</i>"]
        B2["Entropy Weight Method (EWM)<br/><i>Dispersi Informasi Shannon W_j</i>"]
    end
    subgraph S3["3. Transformasi & Agregasi Pilar"]
        C1["Mapping Skala Likert Diskret (0 - 5)<br/><i>Threshold Outlier Ekstrem >= +1.0σ</i>"]
        C2["EWM Weighted Average per Pilar<br/><i>Udara, Air, Lahan, Sosial, Veto</i>"]
    end
    subgraph S4["4. Sintesis Komposit & Klasifikasi"]
        D1["Skor Komposit 0.0 - 5.0 (WSM 0-10)<br/><i>Mean 5 Pilar Terbobot EWM</i>"]
        D2["Vonis Status Daya Dukung<br/><i>Melampaui / Mendekati / Tidak</i>"]
    end
    A1 & A2 --> B1 & B2
    B1 --> C1
    B2 & C1 --> C2 --> D1 --> D2"""

    png_regional = str(tool_dir / "mermaid_flowchart_6_6_regional.png")
    dl_regional = download_mermaid_png(mermaid_regional, png_regional)

    # =========================================================================
    # [2/4] MEMBANGUN DOKUMEN WORD (DOCX) TERPUSAT
    # =========================================================================
    print("[2/4] Membangun DOCX Metodologi_Bab6_Skoring_Provinsi.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after = Pt(2)
    run_hdr = p_hdr.add_run("CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH TINGKAT PROVINSI")
    run_hdr.bold = True
    run_hdr.font.size = Pt(8.0)
    run_hdr.font.color.rgb = C_RED

    add_h1(doc, "BAB VI: AUDIT FORENSIK METODOLOGI D3TLH")
    add_h2(doc, "SUB-BAB 6.6: ALGORITMA SKORING TINGKAT PROVINSI (MODEL HYBRID Z-SCORE & EWM)")

    # -------------------------------------------------------------------------
    # KERANGKA METODOLOGI & FORMULASI MATEMATIS UNIVERSAL (BERLAKU 6 PROVINSI)
    # -------------------------------------------------------------------------
    add_note_box(
        doc,
        "KERANGKA METODOLOGI MULTIKRITERIA REGIONAL (BERLAKU MENGIKAT UNTUK 6 PROVINSI)",
        "Model evaluasi daya dukung dan daya tampung lingkungan hidup tingkat provinsi dirancang menggunakan pendekatan terstandarisasi berbasis Hybrid Z-Score Anomali Deviasi Standar dan Entropy Weight Method (EWM) sesuai Nature Scientific Reports (Sun et al., 2024). "
        "Metodologi, formula normalisasi, matriks entropi, dan bobot objektif indikator dihitung secara simultan dari matriks 6 provinsi se-Pulau Sulawesi, sehingga menjadi dasar penilaian yang adil, objektif, dan bebas bias subjektivitas."
    )

    add_h4(doc, "A. Pengantar & Kerangka Narasi Metodologis")
    add_p(doc, [
        ("Sebagaimana ditampilkan pada antarmuka Streamlit ", False, False),
        ("Dashboard Page 6 (Audit D3TLH - Tab Bedah Matematika Z-Score + EWM per Provinsi)", True, False),
        (", evaluasi tingkat provinsi bertujuan untuk mengatasi kelemahan mendasar dokumen AMDAL dan D3TLH konvensional yang kerap mengaburkan krisis lingkungan lokal melalui teknik perataan agregat wilayah (", False, False),
        ("dilution effect", True, True),
        ("). Dalam metodologi pemerintah, beban pencemaran masif di suatu kawasan industri tambang sering kali tampak 'aman' hanya karena dirata-ratakan dengan luas daratan pulau secara keseluruhan.", False, False)
    ])
    add_p(doc, [
        ("Untuk mendobrak bias tersebut, riset ini menerapkan ", False, False),
        ("Model Hybrid Z-Score Anomali dan Pembobotan Objektif Entropi (EWM)", True, False),
        (". Pendekatan ini secara otomatis memberikan bobot evaluasi tertinggi pada indikator-indikator yang memiliki tingkat ketimpangan spasial paling ekstrem (seperti timbulan limbah B3, tailing tambang, korban krisis agraria, dan PLTU captive batubara). "
         "Dengan demikian, provinsi yang menjadi episentrum industri ekstraktif terdeteksi secara akurat berada pada status anomali krisis tanpa terdistorsi oleh luas wilayah administratif.", False, False)
    ])

    add_h4(doc, "B. Alur Logika Metodologis Regional (Flowchart 6 Provinsi)")
    add_p(doc, [
        ("Seluruh proses komputasi skoring dari 6 provinsi di Pulau Sulawesi dijalankan melalui satu alur algoritma regional terintegrasi sebagaimana disajikan pada ", False, False),
        ("Bagan Alur 6.6", True, False),
        (":", False, False)
    ])
    add_caption(doc, "Bagan Alur 6.6: Alur Logika Algoritma Skoring Tingkat Provinsi (Model Hybrid Z-Score & EWM)")
    if dl_regional:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(png_regional, width=Cm(15.0))
        except Exception as exc:
            print(f"[WARN] Gagal pasang gambar Regional: {exc}")

    add_h4(doc, "C. Formulasi Matematis Universal & Definisi Variabel")
    add_p(doc, [("Operasionalisasi matematis di bawah ini berlaku universal dan mengikat untuk seluruh 6 provinsi di Pulau Sulawesi:", False, False)])

    add_formula_box(
        doc,
        "Tahap 1: Standardisasi Deviasi Z-Score Regional (Anomali Spasial)",
        "Z_ij = (x_ij - mean(x_j)) / std(x_j)    ;    Khusus IKA: Z_ika = - (ika_i - mean(ika)) / std(ika)",
        [
            ("x_ij", "Nilai empiris aktual provinsi i pada indikator j."),
            ("mean(x_j)", "Rata-rata aritmatika indikator j dari seluruh 6 provinsi di Pulau Sulawesi (nilai B)."),
            ("std(x_j)", "Standar deviasi sampel indikator j se-Sulawesi (nilai C; disubstitusi 1.0 jika deviasi 0)."),
            ("Inversi IKA", "Indeks Kualitas Air diinversi tandanya (dikalikan -1) karena nilai IKA tinggi mencerminkan kondisi air baik, sedangkan nilai rendah mencerminkan krisis pencemaran.")
        ]
    )

    add_formula_box(
        doc,
        "Tahap 2: Pembobotan Objektif Entropi Informasi (Entropy Weight Method - EWM)",
        "r_ij = (x_ij - min(x_j)) / (max(x_j) - min(x_j))  -->  P_ij = r_ij / SUM(r_ij)\n"
        "E_j = - (1 / ln(n)) * SUM(P_ij * ln(P_ij + eps))  -->  D_j = 1 - E_j  -->  W_j = D_j / SUM(D_j)",
        [
            ("r_ij", "Matriks ternormalisasi skala Min-Max [0, 1] per kolom indikator."),
            ("P_ij", "Proporsi probabilitas kontribusi provinsi i terhadap total nilai indikator j."),
            ("E_j", "Nilai entropi informasi Shannon indikator j (konstanta k = 1 / ln(6) = 0.5581). Semakin timpang nilainya antar-provinsi, semakin kecil nilai entropi informasinya."),
            ("D_j", "Derajat divergensi atau koefisien dispersi informasi indikator j (D_j = 1 - E_j)."),
            ("W_j", "Bobot objektif final EWM untuk indikator j. Indikator yang paling timpang antar-provinsi (misal PLTU, B3, dan Tailing) otomatis memperoleh bobot evaluasi tertinggi.")
        ]
    )

    add_formula_box(
        doc,
        "Tahap 3: Pemetaan Z-Score ke Skala Likert Diskret (0.0 - 5.0)",
        "L_ij = 5.0 (jika Z >= +1.0) ; 4.0 (0.5 <= Z < 1.0) ; 3.0 (0.0 <= Z < 0.5) ; 2.0 (-0.5 <= Z < 0.0) ; 1.0 (-1.0 <= Z < -0.5) ; 0.0 (Z < -1.0)",
        [
            ("Skor 5.0 (Z >= +1.0σ)", "Outlier Kritis Ekstrem / Red Alert (Beban indikator melampaui batas rata-rata regional lebih dari 1 standar deviasi)."),
            ("Skor 4.0 (+0.5σ s/d +1.0σ)", "Kerentanan Tinggi / Kondisi Buruk."),
            ("Skor 3.0 (0.0 s/d +0.5σ)", "Ambang Batas Waspada / Kondisi Sedang."),
            ("Skor 1.0 - 2.0 (Z < 0.0)", "Rendah / Waspada (Di bawah rata-rata tekanan lingkungan regional)."),
            ("Skor 0.0 (Z < -1.0σ)", "Bebas Risiko Signifikan.")
        ]
    )

    add_formula_box(
        doc,
        "Tahap 4 & 5: Agregasi EWM Weighted Average per Pilar & Skor Komposit Provinsi",
        "Skor_Pilar = SUM(L_ij * W_j) / SUM(W_j)    ;    Skor_Komposit = (Udara + Air + Lahan + Sosial + Veto) / 5.0",
        [
            ("Skor_Pilar", "Rata-rata tertimbang skor Likert dalam satu matriks menggunakan bobot objektif EWM masing-masing indikator."),
            ("Skor_Komposit", "Rata-rata unweighted dari 5 pilar daya dukung (skala 0.0 - 5.0). Ekuivalen Skor WSM 0 - 10 diperoleh dari Skor Komposit dikali 2.0."),
            ("Vonis Status Ekologis", "Melampaui Batas (Skor >= 4.0), Mendekati Batas (Skor = 3.0), Tidak Melampaui Batas (Skor < 3.0).")
        ]
    )

    add_formula_box(
        doc,
        "Contoh Persamaan Substitusi Riil: Indikator Kapasitas PLTU Captive & Komposit Sulteng",
        "1. Substitusi Z-Score: Z = (7.325 MW - 1.637,50 MW) / 2.882,26 MW = +1,97σ\n"
        "2. Substitusi EWM Shannon: Min-Max r_sulteng = (7.325 - 0) / 7.325 = 1,000 ; Proporsi P_sulteng = 1,000 / 1,341 = 0,745\n"
        "   Entropi Ej = -0,5581 * SUM(P * ln P) = 0,3948  -->  Dj = 1 - 0,3948 = 0,6052  -->  W_pltu = 0,6052 / 7,8331 = 0,0773 (7,73%)\n"
        "3. Substitusi Likert Diskret: Z = +1,97σ >= +1,0σ  -->  Skor Likert = 5,0 / 5 (Melampaui Batas / Red Alert)\n"
        "4. Substitusi Pilar Udara: Skor = [(5,0*0,0773) + (4,0*0,0224) + (5,0*0,0461) + (5,0*0,0829) + (5,0*0,0395)] / 0,2682 = 4,92 / 5\n"
        "5. Substitusi Komposit Total: Skor = (4,92 + 3,30 + 4,70 + 2,50 + 4,40) / 5 = 3,96 / 5,0  -->  WSM = 7,92 / 10.0 (Melampaui Batas)",
        [
            ("Substitusi Z-Score", "Nilai aktual 7.325 MW milik Sulteng diselisihkan terhadap rerata 6 provinsi (1.637,50 MW), lalu dibagi standar deviasi regional (2.882,26 MW), menghasilkan deviasi anomali Z = +1,97σ."),
            ("Substitusi EWM", "Tingginya disparitas PLTU se-Sulawesi (Sulteng 7.325 MW, Sultra 1.900 MW, Sulsel 600 MW, 3 provinsi lain 0 MW) menghasilkan divergensi Dj = 0,6052 sehingga indikator ini memperoleh bobot objektif tertinggi W = 0,0773 (7,73%)."),
            ("Substitusi Likert", "Karena nilai Z = +1,97σ melampaui ambang batas atas regional (+1,0σ), maka langsung dipetakan ke skor tertinggi yaitu 5,0 (Krisis Parah / Melampaui Batas)."),
            ("Substitusi Pilar", "Total bobot kelima indikator pilar udara adalah 0,2682. Hasil agregasi tertimbang menghasilkan Skor Pilar Udara sebesar 4,92 / 5 (dibulatkan 4,9 / 5)."),
            ("Substitusi Komposit", "Rata-rata unweighted 5 pilar menghasilkan Skor Komposit 3,96 / 5,0 (dibulatkan 4,0 / 5) atau WSM 7,92 / 10.0 dengan vonis status Melampaui Batas (RED ALERT).")
        ]
    )

    add_caption(doc, "Tabel 6.12: Matriks Parameter Regional Se-Sulawesi (Rata-rata, Deviasi Standar, dan Bobot Objektif EWM 20 Indikator Empiris)")
    add_table_styled(
        doc,
        ["Pilar", "Indikator Empiris", "Rata-rata (B)", "Deviasi (C)", "Entropi (Ej)", "Divergensi (Dj)", "Bobot EWM (Wj)"],
        table_regional_meta,
        [2.2, 4.5, 2.5, 2.5, 1.8, 1.8, 1.7],
        ["C", "L", "R", "R", "C", "C", "C"]
    )

    # -------------------------------------------------------------------------
    # SEKSI 6.6.1: EVALUASI EMPIRIS SULAWESI TENGAH
    # -------------------------------------------------------------------------
    doc.add_page_break()
    add_h3(doc, "6.6.1 Evaluasi Empiris D3TLH: Provinsi Sulawesi Tengah (Sulteng)")
    add_note_box(
        doc,
        "PROFIL EMPIRIS: Provinsi Sulawesi Tengah (Episentrum Hilirisasi & PLTU Captive)",
        "Kabupaten/Kota: 13 Daerah  |  Pusat Industri: Kawasan IMIP Morowali & Smelter Palu  |  Populasi BPS: 2.985.734 Jiwa\n"
        "Karakteristik Krisis: Konsentrasi PLTU captive batubara terbesar nasional, hotspot satelit troposferik NO2 tertinggi, timbulan limbah B3 raksasa, dan laju deforestasi primer masif."
    )

    add_h4(doc, "A. Narasi Temuan Lapangan Sulteng")
    add_p(doc, [
        ("Hasil komputasi algoritma Z-Score EWM membuktikan bahwa ", False, False),
        ("Provinsi Sulawesi Tengah berada pada status RED ALERT (Skor Komposit 4.0 / 5.0 — Melampaui Batas)", True, False),
        (". Dari 20 indikator yang diuji, sebanyak 13 indikator berada pada kategori ", False, False),
        ("Melampaui Batas (Skor Likert 4.0 hingga 5.0)", True, False),
        (", dengan tekanan polusi udara dan perusakan lanskap daratan yang telah melampaui kapasitas asimilasi ekosistem.", False, False)
    ])

    add_h4(doc, "B. Matriks Hasil Uji Empiris (Sulteng)")
    add_caption(doc, "Tabel 6.13: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Tengah (Model Hybrid Z-Score & EWM)")
    add_table_styled(
        doc,
        ["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"],
        table_eval_sulteng,
        [2.2, 4.3, 3.0, 1.8, 1.7, 1.8, 2.2],
        ["C", "L", "R", "C", "C", "C", "C"]
    )

    add_caption(doc, "Tabel 6.14: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Tengah")
    add_table_styled(
        doc,
        ["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulteng"],
        rekap_sulteng,
        [2.5, 4.5, 2.2, 2.8, 5.0],
        ["C", "L", "C", "C", "L"]
    )

    add_h4(doc, "C. Analisis Temuan Empiris (Sulteng)")
    add_p(doc, [
        ("1. ", True, False), (f"Daya Tampung Udara (Skor {sulteng['udara']:.1f} / 5 — {get_likert_label(sulteng['udara'])}): ", True, False),
        (f"Sulawesi Tengah memikul beban polusi udara terparah se-Sulawesi. Kapasitas PLTU captive batubara mencapai 7.325,0 MW (Z = +1.97σ), timbulan limbah B3 menyentuh 25,30 Juta Ton (Z = +1.97σ), emisi karbon 291,34 Juta Ton CO2e (Z = +1.67σ), dan rasio anomali ISPA mencapai 3,50x lipat (Z = +1.66σ).\n", False, False),
        ("2. ", True, False), (f"Daya Tampung Air (Skor {sulteng['air']:.1f} / 5 — {get_likert_label(sulteng['air'])}): ", True, False),
        (f"Meskipun rerata IKA bernilai 62,07, timbulan tailing dan slag mencapai 24,50 Juta Ton/Tahun (Z = +1.97σ) serta memicu lonjakan morbiditas diare sebesar 1,52x lipat dibanding kontrol (Z = +1.34σ, Likert 5.0).\n", False, False),
        ("3. ", True, False), (f"Daya Dukung Lahan (Skor {sulteng['lahan']:.1f} / 5 — {get_likert_label(sulteng['lahan'])}): ", True, False),
        (f"Kehancuran tutupan daratan terberat dengan total deforestasi primer 481.908 Ha (Z = +1.57σ), perambahan 19.804 Ha di kawasan hutan lindung (Z = +1.89σ), serta 383.304 Ha deforestasi komoditas tambang/sawit (Z = +1.69σ).\n", False, False),
        ("4. ", True, False), (f"Daya Dukung Sosial (Skor {sulteng['sosial']:.1f} / 5 — {get_likert_label(sulteng['sosial'])}): ", True, False),
        (f"Tercatat 12.231 jiwa masyarakat adat dan petani terancam kehilangan ruang hidup, serta terjadi 6 insiden kriminalisasi warga dan aktivis lingkungan hidup (Z = +0.71σ, Likert 4.0).\n", False, False),
        ("5. ", True, False), (f"Veto Kebijakan (Skor {sulteng['veto']:.1f} / 5 — {get_likert_label(sulteng['veto'])}): ", True, False),
        (f"Terjadi kegagalan pengendalian perizinan dengan diterbitkannya 260 IUP baru pasca-2014 (Z = +1.64σ, Likert 5.0) dan pembiaran 3 korporasi besar beroperasi ilegal di kawasan hutan.\n", False, False),
        ("6. ", True, False), (f"Vonis Komposit Sulawesi Tengah (Skor {sulteng['total_likert']:.1f} / 5.0 — {sulteng['likert_label']}): ", True, False),
        (f"Secara agregat, Sulawesi Tengah memperoleh Skor Komposit 4.0 / 5.0 (Ekuivalen WSM 7.92 / 10.0) dengan status MELAMPAUI BATAS (RED ALERT).", False, False)
    ])

    # -------------------------------------------------------------------------
    # SEKSI 6.6.2: EVALUASI EMPIRIS SULAWESI TENGGARA
    # -------------------------------------------------------------------------
    doc.add_page_break()
    add_h3(doc, "6.6.2 Evaluasi Empiris D3TLH: Provinsi Sulawesi Tenggara (Sultra)")
    add_note_box(
        doc,
        "PROFIL EMPIRIS: Provinsi Sulawesi Tenggara (Episentrum Konflik Agraria & Kepadatan IUP Ekstrem)",
        "Kabupaten/Kota: 17 Daerah  |  Pusat Industri: Smelter Morosi, Konawe, Kolaka & Pulau Wawonii  |  Populasi BPS: 2.624.875 Jiwa\n"
        "Karakteristik Krisis: Kepadatan konsesi IUP tambang nikel tertinggi se-Sulawesi (11,72% daratan), korban perampasan ruang hidup terbesar (39.821 jiwa), sengketa ruang tangkap nelayan pesisir, dan pelanggaran persetujuan warga (FPIC) masif."
    )

    add_h4(doc, "A. Narasi Temuan Lapangan Sultra")
    add_p(doc, [
        ("Berdasarkan hasil pemetaan empiris Z-Score EWM, Provinsi Sulawesi Tenggara memperlihatkan profil krisis yang sangat kontras dengan Sulawesi Tengah. "
         "Jika Sulawesi Tengah didominasi oleh polusi PLTU dan deforestasi hulu, maka ", False, False),
        ("Sulawesi Tenggara mengalami ledakan krisis daya dukung sosial, perampasan ruang hidup masyarakat pesisir, dan kepadatan konsesi tambang tertinggi se-Pulau Sulawesi.", True, False)
    ])
    add_p(doc, [
        ("Konsesi tambang nikel di Sulawesi Tenggara mencaplok ", False, False),
        ("11,72% dari total daratan provinsi (Z = +1.50σ, Skor Likert 5.0 — Outlier Ekstrem Terpadat Se-Sulawesi)", True, False),
        (". Dampak destruktifnya memicu krisis kemanusiaan langsung terhadap ", False, False),
        ("39.821 jiwa petani dan nelayan (Z = +1.95σ, Skor Likert 5.0 — mencakup 73% total korban se-Sulawesi)", True, False),
        (", serta 5 kasus manipulasi persetujuan warga (FPIC) di wilayah kepulauan.", False, False)
    ])

    add_h4(doc, "B. Matriks Hasil Uji Empiris (Sultra)")
    add_caption(doc, "Tabel 6.15: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Tenggara (Model Hybrid Z-Score & EWM)")
    add_table_styled(
        doc,
        ["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"],
        table_eval_sultra,
        [2.2, 4.3, 3.0, 1.8, 1.7, 1.8, 2.2],
        ["C", "L", "R", "C", "C", "C", "C"]
    )

    add_caption(doc, "Tabel 6.16: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Tenggara")
    add_table_styled(
        doc,
        ["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sultra"],
        rekap_sultra,
        [2.5, 4.5, 2.2, 2.8, 5.0],
        ["C", "L", "C", "C", "L"]
    )

    add_h4(doc, "C. Analisis Temuan Empiris (Sultra)")
    add_p(doc, [
        ("1. ", True, False), (f"Daya Tampung Udara (Skor {sultra['udara']:.1f} / 5 — {get_likert_label(sultra['udara'])}): ", True, False),
        (f"Kapasitas PLTU captive beroperasi mencapai 1.900,0 MW (Z = +0.09σ), konsentrasi gas NO2 troposferik 6,62e-06 mol/m² (Z = +0.82σ, Likert 4.0), dan pelepasan emisi karbon deforestasi 189,02 Juta Ton CO2e (Z = +0.59σ, Likert 4.0). Tekanan polusi terkonsentrasi di kawasan industri Morosi (Konawe).\n", False, False),
        ("2. ", True, False), (f"Daya Tampung Air (Skor {sultra['air']:.1f} / 5 — {get_likert_label(sultra['air'])}): ", True, False),
        (f"Meskipun IKA sungai bernilai 65,32, tercatat timbulan tailing dan slag 6,52 Juta Ton (Z = +0.12σ) serta meletusnya 5 kasus sengketa ruang tangkap laut antara nelayan pesisir dan armada tongkang tambang (Z = +0.87σ, Likert 4.0).\n", False, False),
        ("3. ", True, False), (f"Daya Dukung Lahan (Skor {sultra['lahan']:.1f} / 5 — {get_likert_label(sultra['lahan'])}): ", True, False),
        (f"Ditemukan anomali ekstrem pada Kepadatan Konsesi IUP yang mencaplok 11,72% dari total daratan daratan provinsi (Z = +1.50σ, Likert 5.0 - Outlier Tertinggi se-Sulawesi), menyebabkan 337.434 Ha deforestasi hutan primer (Z = +0.67σ) dan perambahan 8.236 Ha hutan lindung.\n", False, False),
        ("4. ", True, False), (f"Daya Dukung Sosial (Skor {sultra['sosial']:.1f} / 5 — {get_likert_label(sultra['sosial'])}): ", True, False),
        (f"Sulawesi Tenggara mencatat krisis sosial terparah di seluruh bioregion Pulau Sulawesi. Sebanyak 39.821 jiwa warga terancam kehilangan ruang hidup (Z = +1.95σ, Likert 5.0), terjadi 5 kasus manipulasi persetujuan FPIC (Z = +1.86σ, Likert 5.0), 4 insiden kriminalisasi aktivis, dan defisit fasilitas SPA Puskesmas sebesar 17,92% (Z = +0.79σ, Likert 4.0).\n", False, False),
        ("5. ", True, False), (f"Veto Kebijakan (Skor {sultra['veto']:.1f} / 5 — {get_likert_label(sultra['veto'])}): ", True, False),
        (f"Diterbitkan 160 IUP baru pasca-2014 (Z = +0.64σ, Likert 4.0), memperparah pengkapalan bijih nikel ilegal di pulau-pulau kecil.\n", False, False),
        ("6. ", True, False), (f"Vonis Komposit Sulawesi Tenggara (Skor {sultra['total_likert']:.1f} / 5.0 — {sultra['likert_label']}): ", True, False),
        (f"Secara agregat, Sulawesi Tenggara memperoleh Skor Komposit 3.4 / 5.0 (Ekuivalen WSM 6.78 / 10.0) dengan status MENDEKATI BATAS. Namun, pengujian forensik membuktikan bahwa Pilar Daya Dukung Sosial (4.5 / 5) dan Kepadatan Konsesi Tambang (11,72%) telah berada pada status RED ALERT (MELAMPAUI BATAS EKSTREM).", False, False)
    ])

    # -------------------------------------------------------------------------
    # SEKSI 6.6.3: EVALUASI EMPIRIS SULAWESI SELATAN
    # -------------------------------------------------------------------------
    doc.add_page_break()
    add_h3(doc, "6.6.3 Evaluasi Empiris D3TLH: Provinsi Sulawesi Selatan (Sulsel)")
    add_note_box(
        doc,
        "PROFIL EMPIRIS: Provinsi Sulawesi Selatan (Episentrum Bencana Alam, Konflik Pesisir & Kriminalisasi)",
        "Kabupaten/Kota: 24 Daerah  |  Pusat Industri: KIMA Makassar, Smelter Huadi Bantaeng, Vale Sorowako Luwu Timur & PLTU Jeneponto  |  Populasi BPS: 9.073.509 Jiwa\n"
        "Karakteristik Krisis: Frekuensi bencana hidrometeorologi banjir bandang dan longsor tertinggi se-Sulawesi (669 kejadian), sengketa ruang laut nelayan pesisir terbanyak (7 kasus), insiden kriminalisasi warga tertinggi (9 kasus), tambang ilegal marak (10 korporasi), dan cemaran karsinogenik Cr6+."
    )

    add_h4(doc, "A. Narasi Temuan Lapangan Sulsel")
    add_p(doc, [
        ("Sebagai provinsi dengan populasi terbesar (9,07 juta jiwa) dan pusat gravitasi ekonomi regional, ", False, False),
        ("Provinsi Sulawesi Selatan mencatat Skor Komposit 2.6 / 5.0 (Status: Mendekati Batas)", True, False),
        (". Kendati secara agregat tidak berada pada status Melampaui Batas layaknya Sulteng, ", False, False),
        ("audit forensik Z-score membongkar anomali outlier ekstrem pada 5 indikator kritis (Skor Likert 5.0 / Red Alert)", True, False),
        (" yang memperlihatkan kerentanan ekologis struktural di kawasan pesisir, daerah aliran sungai (DAS), dan ruang hidup agraria.", False, False)
    ])
    add_p(doc, [
        ("Sulawesi Selatan mencatat rekor tertinggi se-Sulawesi pada tiga variabel destruktif sekaligus: ", False, False),
        ("kejadian bencana hidrometeorologi sebanyak 669 kali (Z = +1.63σ, Likert 5.0)", True, False),
        (", meletusnya ", False, False),
        ("7 kasus konflik ruang tangkap laut nelayan vs tambang pasir laut dan tongkang (Z = +1.56σ, Likert 5.0)", True, False),
        (", serta represi hukum dengan ", False, False),
        ("9 insiden kriminalisasi petani dan aktivis pembela HAM (Z = +1.57σ, Likert 5.0)", True, False),
        (". Selain itu, maraknya operasi ", False, False),
        ("10 korporasi tambang ilegal di kawasan lindung (Z = +1.97σ, Likert 5.0)", True, False),
        (" dan cemaran Heksavalen Kromium (Cr6+) menegaskan darurat tata kelola lingkungan hidup di provinsi ini.", False, False)
    ])

    add_h4(doc, "B. Matriks Hasil Uji Empiris (Sulsel)")
    add_caption(doc, "Tabel 6.17: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Selatan (Model Hybrid Z-Score & EWM)")
    add_table_styled(
        doc,
        ["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"],
        table_eval_sulsel,
        [2.2, 4.3, 3.0, 1.8, 1.7, 1.8, 2.2],
        ["C", "L", "R", "C", "C", "C", "C"]
    )

    add_caption(doc, "Tabel 6.18: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Selatan")
    add_table_styled(
        doc,
        ["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulsel"],
        rekap_sulsel,
        [2.5, 4.5, 2.2, 2.8, 5.0],
        ["C", "L", "C", "C", "L"]
    )

    add_h4(doc, "C. Analisis Temuan Empiris (Sulsel)")
    add_p(doc, [
        ("1. ", True, False), (f"Daya Tampung Udara (Skor {sulsel['udara']:.1f} / 5 — {get_likert_label(sulsel['udara'])}): ", True, False),
        (f"Kapasitas PLTU captive beroperasi mencapai 600,0 MW (Punagaya Jeneponto & Barru, Z = -0.36σ), emisi karbon 138,73 Jt Ton CO2e (Z = +0.05σ), namun konsentrasi gas NO2 satelit menyentuh 6,40e-06 mol/m² (Z = +0.65σ, Likert 4.0) yang mencerminkan beban emisi perkotaan dan industri KIMA.\n", False, False),
        ("2. ", True, False), (f"Daya Tampung Air (Skor {sulsel['air']:.1f} / 5 — {get_likert_label(sulsel['air'])}): ", True, False),
        (f"Rerata IKA bernilai 58,50 poin (Z = +0.35σ), namun terdeteksi kontaminasi toksik Heksavalen Kromium Cr6+ (Z = +2.03σ, Likert 5.0) di perairan Luwu Timur, serta meletusnya 7 kasus konflik ruang laut nelayan pesisir vs tambang pasir laut dan tongkang (Z = +1.56σ, Likert 5.0 — tertinggi se-Sulawesi).\n", False, False),
        ("3. ", True, False), (f"Daya Dukung Lahan (Skor {sulsel['lahan']:.1f} / 5 — {get_likert_label(sulsel['lahan'])}): ", True, False),
        (f"Sulsel memikul bencana hidrometeorologi terparah se-Sulawesi dengan 669 kejadian banjir bandang dan longsor (Z = +1.63σ, Likert 5.0), dipicu laju deforestasi primer 261.147 Ha (Z = +0.19σ) dan perambahan 5.314 Ha hutan lindung di kawasan hulu DAS.\n", False, False),
        ("4. ", True, False), (f"Daya Dukung Sosial (Skor {sulsel['sosial']:.1f} / 5 — {get_likert_label(sulsel['sosial'])}): ", True, False),
        (f"Meskipun jumlah warga terdampak langsung tercatat 2.257 jiwa (Z = -0.43σ), Sulawesi Selatan mencatat eskalasi represi terberat dengan 9 insiden kriminalisasi warga dan aktivis lingkungan (Z = +1.57σ, Likert 5.0 — rekor tertinggi se-Sulawesi).\n", False, False),
        ("5. ", True, False), (f"Veto Kebijakan (Skor {sulsel['veto']:.1f} / 5 — {get_likert_label(sulsel['veto'])}): ", True, False),
        (f"Ditemukan 10 korporasi tambang beroperasi ilegal di kawasan hutan lindung dan DAS (Z = +1.97σ, Likert 5.0 — tertinggi se-Sulawesi) serta diterbitkannya 105 IUP baru pasca-2014 (Z = +0.09σ).\n", False, False),
        ("6. ", True, False), (f"Vonis Komposit Sulawesi Selatan (Skor {sulsel['total_likert']:.1f} / 5.0 — {sulsel['likert_label']}): ", True, False),
        (f"Secara agregat, Sulawesi Selatan memperoleh Skor Komposit 2.6 / 5.0 (Ekuivalen WSM 5.29 / 10.0) dengan status MENDEKATI BATAS. Namun, audit forensik membuktikan kondisi darurat pada 5 indikator outlier kritis: Bencana Alam (669 kejadian), Konflik Pesisir (7 kasus), Kriminalisasi Pejuang HAM (9 insiden), Tambang Ilegal (10 korporasi), dan Toksisitas Cr6+ yang berada pada status RED ALERT (MELAMPAUI BATAS EKSTREM).", False, False)
    ])

    # -------------------------------------------------------------------------
    # SEKSI 6.6.4: EVALUASI EMPIRIS SULAWESI BARAT
    # -------------------------------------------------------------------------
    doc.add_page_break()
    add_h3(doc, "6.6.4 Evaluasi Empiris D3TLH: Provinsi Sulawesi Barat (Sulbar)")
    add_note_box(
        doc,
        "PROFIL EMPIRIS: Provinsi Sulawesi Barat (Bioregion Non-Hilirisasi & Dominasi Agromaritim)",
        "Kabupaten/Kota: 6 Daerah  |  Basis Perekonomian: Perkebunan Sawit, Kakao & Perikanan Tangkap  |  Populasi BPS: 1.419.229 Jiwa\n"
        "Karakteristik Krisis: Bebas dari ekspansi PLTU captive batubara (0 MW) dan nihil timbulan limbah B3/tailing nikel, namun menghadapi tekanan mutu air sungai (IKA 55,93 poin) akibat limbah PKS perkebunan sawit monokultur dan erosi DAS."
    )

    add_h4(doc, "A. Narasi Temuan Lapangan Sulbar")
    add_p(doc, [
        ("Sebagai wilayah pemekaran dengan 6 kabupaten dan populasi 1,42 juta jiwa, ", False, False),
        ("Provinsi Sulawesi Barat mencatatkan Skor Komposit 1.2 / 5.0 (Status: Tidak Melampaui Batas)", True, False),
        (". Profil ekologis Sulbar menjadi ", False, False),
        ("bukti empiris pembanding (control baseline)", True, False),
        (" yang sangat berharga dalam laporan ini. Karena tidak menjadi lokasi hilirisasi industri nikel berskala raksasa, Sulbar terhindar dari akumulasi limbah B3, tailing laut, dan PLTU captive batubara yang menghancurkan daya lentur lingkungan hidup seperti di Sulteng dan Sultra.", False, False)
    ])
    add_p(doc, [
        ("Kepadatan konsesi tambang nikel di Sulbar tercatat hanya ", False, False),
        ("0,26% dari luas daratan provinsi (Z = -1.09σ, Likert 0.0 — Terendah se-Sulawesi)", True, False),
        (". Namun demikian, evaluasi D3TLH mencatat ", False, False),
        ("dua peringatan ekologis lokal (local stressor)", True, False),
        (": yaitu penurunan Indeks Kualitas Air (IKA 55,93 poin, Z = +1.11σ, Likert 5.0) dan angka morbiditas diare sebesar 1,27x lipat (Z = +0.63σ, Likert 4.0), yang dipicu oleh pelepasan limbah cair pabrik kelapa sawit (PKS) monokultur di Pasangkayu dan Mamuju Tengah serta defisit sanitasi dasar perdesaan.", False, False)
    ])

    add_h4(doc, "B. Matriks Hasil Uji Empiris (Sulbar)")
    add_caption(doc, "Tabel 6.19: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Barat (Model Hybrid Z-Score & EWM)")
    add_table_styled(
        doc,
        ["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"],
        table_eval_sulbar,
        [2.2, 4.3, 3.0, 1.8, 1.7, 1.8, 2.2],
        ["C", "L", "R", "C", "C", "C", "C"]
    )

    add_caption(doc, "Tabel 6.20: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Barat")
    add_table_styled(
        doc,
        ["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulbar"],
        rekap_sulbar,
        [2.5, 4.5, 2.2, 2.8, 5.0],
        ["C", "L", "C", "C", "L"]
    )

    add_h4(doc, "C. Analisis Temuan Empiris (Sulbar)")
    add_p(doc, [
        ("1. ", True, False), (f"Daya Tampung Udara (Skor {sulbar['udara']:.1f} / 5 — {get_likert_label(sulbar['udara'])}): ", True, False),
        (f"Sulbar bebas dari beban PLTU captive batubara (0,0 MW, Z = -0.57σ) dan nihil timbulan limbah B3 industri smelter (0,0 Ton, Z = -0.54σ). Emisi karbon deforestasi tercatat 82,51 Jt Ton CO2e (Z = -0.55σ) dan konsentrasi NO2 troposferik sebesar 6,00e-06 mol/m² (Z = +0.34σ).\n", False, False),
        ("2. ", True, False), (f"Daya Tampung Air (Skor {sulbar['air']:.1f} / 5 — {get_likert_label(sulbar['air'])}): ", True, False),
        (f"Nihil pembuangan tailing/slag tambang (0,0 Ton), namun mencatat IKA terendah se-Sulawesi (55,93 poin, Z = +1.11σ, Likert 5.0) dan insidensi diare 1,27x lipat (Z = +0.63σ, Likert 4.0) yang bersumber dari limbah cair PKS kelapa sawit dan sedimentasi erosi DAS Lariang & Karama.\n", False, False),
        ("3. ", True, False), (f"Daya Dukung Lahan (Skor {sulbar['lahan']:.1f} / 5 — {get_likert_label(sulbar['lahan'])}): ", True, False),
        (f"Kepadatan konsesi tambang nikel paling rendah se-Sulawesi (hanya 0,26% daratan, Z = -1.09σ, Likert 0.0), deforestasi primer 133.263 Ha (Z = -0.61σ), perambahan hutan lindung 1.251 Ha, dan frekuensi bencana alam sebanyak 143 kejadian.\n", False, False),
        ("4. ", True, False), (f"Daya Dukung Sosial (Skor {sulbar['sosial']:.1f} / 5 — {get_likert_label(sulbar['sosial'])}): ", True, False),
        (f"Minim konflik agraria struktural skala besar yang melibatkan industri tambang (1 jiwa terdampak, 1 insiden kriminalisasi, dan nihil sengketa FPIC), mencerminkan relasi sosial-ekologis yang relatif stabil.\n", False, False),
        ("5. ", True, False), (f"Veto Kebijakan (Skor {sulbar['veto']:.1f} / 5 — {get_likert_label(sulbar['veto'])}): ", True, False),
        (f"Penerbitan konsesi tambang nikel baru pasca-2014 sangat terbatas (27 IUP, Z = -0.68σ, Likert 1.0) dan tidak teridentifikasi adanya korporasi tambang ilegal skala besar yang beroperasi di kawasan hutan.\n", False, False),
        ("6. ", True, False), (f"Vonis Komposit Sulawesi Barat (Skor {sulbar['total_likert']:.1f} / 5.0 — {sulbar['likert_label']}): ", True, False),
        (f"Secara agregat, Sulawesi Barat memperoleh Skor Komposit 1.2 / 5.0 (Ekuivalen WSM 2.36 / 10.0) dengan status TIDAK MELAMPAUI BATAS. Status terjaga ini membuktikan bahwa tanpa intervensi industri smelter nikel dan PLTU batubara, daya lentur lingkungan hidup regional dapat dipertahankan dalam ambang aman.", False, False)
    ])

    # -------------------------------------------------------------------------
    # SEKSI 6.6.5: EVALUASI EMPIRIS GORONTALO
    # -------------------------------------------------------------------------
    doc.add_page_break()
    add_h3(doc, "6.6.5 Evaluasi Empiris D3TLH: Provinsi Gorontalo")
    add_note_box(
        doc,
        "PROFIL EMPIRIS: Provinsi Gorontalo (Bioregion Terjaga & Bebas Polusi Smelter Nikel)",
        "Kabupaten/Kota: 6 Daerah  |  Pusat Agraria: Lembah Pertanian Jagung & Danau Limboto  |  Populasi BPS: 1.171.681 Jiwa\n"
        "Karakteristik Krisis: Kualitas atmosfer NO2 satelit paling bersih se-Sulawesi (3,76e-06 mol/m²), deforestasi primer dan emisi karbon terendah, nihil PLTU captive batubara maupun limbah B3 smelter, namun memikul anomali ISPA akibat faktor mikroklimat topografi cekungan dan residu pembakaran biomassa jagung."
    )

    add_h4(doc, "A. Narasi Temuan Lapangan Gorontalo")
    add_p(doc, [
        ("Sebagai provinsi dengan luas daratan dan populasi terkecil di Pulau Sulawesi (1,17 juta jiwa), ", False, False),
        ("Provinsi Gorontalo mencatatkan Skor Komposit 1.2 / 5.0 (Status: Tidak Melampaui Batas)", True, False),
        (". Bersama dengan Sulawesi Barat, Gorontalo berada pada kuadran ", False, False),
        ("ekologis terjaga (low-stress environment)", True, False),
        (" yang mempertegas validitas model Z-Score EWM: ketika suatu wilayah tidak dieksploitasi oleh mega-proyek hilirisasi nikel dan PLTU captive batubara, integritas daya dukung lingkungannya tetap berada di bawah ambang batas bahaya.", False, False)
    ])
    add_p(doc, [
        ("Gorontalo membukukan rekor ", False, False),
        ("konsentrasi gas troposferik NO2 paling bersih se-Pulau Sulawesi (3,76e-06 mol/m², Z = -1.40σ, Likert 0.0)", True, False),
        (", laju deforestasi primer paling rendah (98.063 Ha, Z = -0.83σ), serta pelepasan emisi karbon deforestasi terendah (53,66 Jt Ton CO2e, Z = -0.85σ). Kepadatan izin tambang nikel hanya mencapai ", False, False),
        ("0,46% dari daratan provinsi (Z = -1.04σ, Likert 0.0)", True, False),
        (". Satu-satunya anomali yang mencuat adalah rasio morbiditas ISPA sebesar 2,41x lipat (Z = +0.79σ, Likert 4.0) yang dipicu oleh faktor mikroklimat topografi cekungan Lembah Limboto, debu jalanan pedesaan, serta pembakaran residu tongkol jagung pascapanen, bukan dari cerobong batubara.", False, False)
    ])

    add_h4(doc, "B. Matriks Hasil Uji Empiris (Gorontalo)")
    add_caption(doc, "Tabel 6.21: Bedah Matematika 20 Indikator Empiris Provinsi Gorontalo (Model Hybrid Z-Score & EWM)")
    add_table_styled(
        doc,
        ["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"],
        table_eval_gorontalo,
        [2.2, 4.3, 3.0, 1.8, 1.7, 1.8, 2.2],
        ["C", "L", "R", "C", "C", "C", "C"]
    )

    add_caption(doc, "Tabel 6.22: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Gorontalo")
    add_table_styled(
        doc,
        ["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Gorontalo"],
        rekap_gorontalo,
        [2.5, 4.5, 2.2, 2.8, 5.0],
        ["C", "L", "C", "C", "L"]
    )

    add_h4(doc, "C. Analisis Temuan Empiris (Gorontalo)")
    add_p(doc, [
        ("1. ", True, False), (f"Daya Tampung Udara (Skor {gorontalo['udara']:.1f} / 5 — {get_likert_label(gorontalo['udara'])}): ", True, False),
        (f"Nihil beban PLTU captive batubara (0,0 MW, Z = -0.57σ) dan nihil limbah B3 smelter (0,0 Ton, Z = -0.54σ). Tingkat polusi NO2 satelit paling rendah se-Sulawesi (3,76e-06 mol/m², Z = -1.40σ, Likert 0.0). Catatan anomali ISPA 2,41x lipat (Z = +0.79σ, Likert 4.0) berkorelasi dengan asap pembakaran biomassa jagung dan dinamika inversi suhu Lembah Limboto.\n", False, False),
        ("2. ", True, False), (f"Daya Tampung Air (Skor {gorontalo['air']:.1f} / 5 — {get_likert_label(gorontalo['air'])}): ", True, False),
        (f"Nihil pembuangan tailing nikel (0,0 Ton) dan nihil konflik ruang laut pesisir. IKA berada pada level 58,14 poin (Z = +0.46σ, Likert 3.0) akibat tekanan sedimentasi erosi DAS Bone-Bolango dan eutrofikasi Danau Limboto dari limbah domestik.\n", False, False),
        ("3. ", True, False), (f"Daya Dukung Lahan (Skor {gorontalo['lahan']:.1f} / 5 — {get_likert_label(gorontalo['lahan'])}): ", True, False),
        (f"Gorontalo membukukan kehilangan tutupan hutan alam primer terendah se-Sulawesi (98.063 Ha, Z = -0.83σ, Likert 1.0) dengan kepadatan konsesi tambang nikel hanya 0,46% daratan (Z = -1.04σ, Likert 0.0) serta nihil catatan bencana longsor/banjir skala masif dalam periode audit.\n", False, False),
        ("4. ", True, False), (f"Daya Dukung Sosial (Skor {gorontalo['sosial']:.1f} / 5 — {get_likert_label(gorontalo['sosial'])}): ", True, False),
        (f"Bebas dari letupan konflik agraria industri tambang (nihil warga terdampak kehilangan ruang hidup, nihil sengketa FPIC, dan nihil kriminalisasi pejuang lingkungan).\n", False, False),
        ("5. ", True, False), (f"Veto Kebijakan (Skor {gorontalo['veto']:.1f} / 5 — {get_likert_label(gorontalo['veto'])}): ", True, False),
        (f"Pemerintah daerah mencatat laju obral izin tambang terendah se-Sulawesi (hanya 7 IUP baru pasca-2014, Z = -0.88σ, Likert 1.0) dan hanya teridentifikasi 1 aktivitas tambang rakyat/ilegal skala kecil di kawasan hulu (Z = -0.45σ).\n", False, False),
        ("6. ", True, False), (f"Vonis Komposit Gorontalo (Skor {gorontalo['total_likert']:.1f} / 5.0 — {gorontalo['likert_label']}): ", True, False),
        (f"Secara agregat, Gorontalo memperoleh Skor Komposit 1.2 / 5.0 (Ekuivalen WSM 2.31 / 10.0) dengan status TIDAK MELAMPAUI BATAS. Status aman ini menjadi bukti konklusif bahwa kelestarian bioregion Sulawesi bertumpu pada pembatasan ekspansi industri ekstraktif nikel.", False, False)
    ])

    docx_path = tool_dir / "Metodologi_Bab6_Skoring_Provinsi.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    # =========================================================================
    # [3/4] MEMBANGUN DOKUMEN HTML SATUAN
    # =========================================================================
    print("[3/4] Membangun HTML dan Markdown...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Bab VI: Audit Forensik Metodologi D3TLH — Sub-bab 6.6 Algoritma Skoring Tingkat Provinsi</title>
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; color: #202426; max-width: 1150px; margin: 0 auto; padding: 25px; background-color: #FAFAFA; }}
  h1 {{ color: #143642; border-bottom: 2px solid #143642; padding-bottom: 8px; font-size: 24px; }}
  h2 {{ color: #A8201A; margin-top: 25px; font-size: 20px; }}
  h3 {{ color: #0F4C5C; margin-top: 30px; font-size: 18px; border-bottom: 1px solid #0F4C5C; padding-bottom: 5px; }}
  h4 {{ color: #143642; margin-top: 15px; font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px; }}
  .note-box {{ background-color: #F4F5F6; border-left: 5px solid #143642; padding: 12px 18px; margin: 15px 0; font-size: 13px; }}
  .formula {{ background-color: #F8F9FA; border-left: 4px solid #0F4C5C; padding: 10px 15px; margin: 10px 0; font-family: Consolas, monospace; font-size: 13px; color: #143642; font-weight: bold; }}
  .table-container {{ overflow-x: auto; margin: 15px 0; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 12px; background-color: #FFF; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
  th {{ background-color: #143642; color: #FFF; padding: 9px 10px; text-align: center; border: 1px solid #143642; font-size: 11px; }}
  td {{ padding: 7px 10px; border: 1px solid #E0E0E0; }}
  tr:nth-child(even) {{ background-color: #F9FAFA; }}
  tr.total-row {{ background-color: #EAEBED; font-weight: bold; }}
  .text-center {{ text-align: center; }}
  .text-right {{ text-align: right; }}
  .text-left {{ text-align: left; }}
  .badge-danger {{ background-color: #A8201A; color: white; padding: 2px 6px; border-radius: 3px; font-weight: bold; font-size: 10px; display: inline-block; }}
  .badge-warning {{ background-color: #E36414; color: white; padding: 2px 6px; border-radius: 3px; font-weight: bold; font-size: 10px; display: inline-block; }}
  .badge-success {{ background-color: #1E7E34; color: white; padding: 2px 6px; border-radius: 3px; font-weight: bold; font-size: 10px; display: inline-block; }}
  .mermaid {{ margin: 20px 0; text-align: center; background: #FFF; padding: 15px; border: 1px solid #E0E0E0; }}
  .table-caption {{ font-weight: bold; font-style: italic; color: #143642; margin-top: 15px; margin-bottom: 5px; font-size: 13px; }}
  .divider {{ margin: 40px 0; border-top: 2px dashed #CCCCCC; }}
</style>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true}});</script>
</head>
<body>

<h1>BAB VI: AUDIT FORENSIK METODOLOGI D3TLH</h1>
<h2>SUB-BAB 6.6: ALGORITMA SKORING TINGKAT PROVINSI (MODEL HYBRID Z-SCORE & EWM)</h2>

<div class="note-box">
  <strong>KERANGKA METODOLOGI MULTIKRITERIA REGIONAL (BERLAKU MENGIKAT UNTUK 6 PROVINSI)</strong><br>
  Model evaluasi daya dukung dan daya tampung lingkungan hidup tingkat provinsi dirancang menggunakan pendekatan terstandarisasi berbasis Hybrid Z-Score Anomali Deviasi Standar dan Entropy Weight Method (EWM) sesuai Nature Scientific Reports (Sun et al., 2024). Metodologi, formula normalisasi, matriks entropi, dan bobot objektif indikator dihitung secara simultan dari matriks 6 provinsi se-Pulau Sulawesi.
</div>

<h4>A. Pengantar & Kerangka Narasi Metodologis</h4>
<p>Sebagaimana ditampilkan pada antarmuka Streamlit <strong>Dashboard Page 6 (Audit D3TLH - Tab Bedah Matematika Z-Score + EWM per Provinsi)</strong>, evaluasi tingkat provinsi bertujuan untuk mengatasi kelemahan mendasar dokumen AMDAL dan D3TLH konvensional yang kerap mengaburkan krisis lingkungan lokal melalui teknik perataan agregat wilayah (<em>dilution effect</em>). Dalam metodologi pemerintah, beban pencemaran masif di suatu kawasan industri tambang sering kali tampak 'aman' hanya karena dirata-ratakan dengan luas daratan pulau secara keseluruhan.</p>
<p>Untuk mendobrak bias tersebut, riset ini menerapkan <strong>Model Hybrid Z-Score Anomali dan Pembobotan Objektif Entropi (EWM)</strong>. Pendekatan ini secara otomatis memberikan bobot evaluasi tertinggi pada indikator-indikator yang memiliki tingkat ketimpangan spasial paling ekstrem (seperti timbulan limbah B3, tailing tambang, korban krisis agraria, dan PLTU captive batubara). Dengan demikian, provinsi yang menjadi episentrum industri ekstraktif terdeteksi secara akurat berada pada status anomali krisis tanpa terdistorsi oleh luas wilayah administratif.</p>

<h4>B. Alur Logika Metodologis Regional (Flowchart 6 Provinsi)</h4>
<div class="mermaid">
{mermaid_regional}
</div>

<h4>C. Formulasi Matematis Universal & Definisi Variabel</h4>
<div class="formula">Tahap 1: Z_ij = (x_ij - mean(x_j)) / std(x_j) &nbsp;&nbsp;|&nbsp;&nbsp; Khusus IKA: Z_ika = - (ika_i - mean(ika)) / std(ika)</div>
<div class="formula">Tahap 2: r_ij = (x_ij - min(x_j)) / (max(x_j) - min(x_j)) &nbsp;➔&nbsp; P_ij = r_ij / &Sigma;r_ij &nbsp;➔&nbsp; E_j = -k * &Sigma;(P_ij * ln(P_ij)) &nbsp;➔&nbsp; W_j = (1 - E_j) / &Sigma;(1 - E_j)</div>
<div class="formula">Tahap 3: L_ij = 5.0 (Z &ge; +1.0) ; 4.0 (0.5 &le; Z < 1.0) ; 3.0 (0.0 &le; Z < 0.5) ; 2.0 (-0.5 &le; Z < 0.0) ; 1.0 (-1.0 &le; Z < -0.5) ; 0.0 (Z < -1.0)</div>
<div class="formula">Tahap 4 & 5: Skor_Pilar = &Sigma;(L_ij * W_j) / &Sigma;W_j &nbsp;&nbsp;|&nbsp;&nbsp; Skor_Komposit = (Udara + Air + Lahan + Sosial + Veto) / 5.0</div>
<div class="formula" style="background-color: #EEF4F8; border-left: 4px solid #143642; margin-top: 15px;">
  <strong>Contoh Persamaan Substitusi Riil (Indikator PLTU Captive & Komposit Sulteng):</strong><br>
  1. Substitusi Z-Score: Z = (7.325 MW - 1.637,50 MW) / 2.882,26 MW = +1,97&sigma;<br>
  2. Substitusi EWM Shannon: r_sulteng = 1,000 ; P_sulteng = 0,745 &nbsp;➔&nbsp; E_pltu = 0,3948 &nbsp;➔&nbsp; D_pltu = 0,6052 &nbsp;➔&nbsp; W_pltu = 0,0773 (7,73%)<br>
  3. Substitusi Likert: Z = +1,97&sigma; &ge; +1,0&sigma; &nbsp;➔&nbsp; Skor Likert = 5,0 / 5 (Melampaui Batas / Red Alert)<br>
  4. Substitusi Pilar Udara: Skor = [(5,0*0,0773) + (4,0*0,0224) + (5,0*0,0461) + (5,0*0,0829) + (5,0*0,0395)] / 0,2682 = 4,92 / 5<br>
  5. Substitusi Komposit Total: Skor = (4,92 + 3,30 + 4,70 + 2,50 + 4,40) / 5 = 3,96 / 5,0 &nbsp;➔&nbsp; WSM: 7,92 / 10.0 (Melampaui Batas)
</div>

<div class="table-caption">Tabel 6.12: Matriks Parameter Regional Se-Sulawesi (Rata-rata, Deviasi Standar, dan Bobot Objektif EWM 20 Indikator Empiris)</div>
{html_table(["Pilar", "Indikator Empiris", "Rata-rata (B)", "Deviasi (C)", "Entropi (Ej)", "Divergensi (Dj)", "Bobot EWM (Wj)"], table_regional_meta)}

<div class="divider"></div>

<!-- SEKSI 6.6.1 SULTENG -->
<h3>6.6.1 Evaluasi Empiris D3TLH: Provinsi Sulawesi Tengah (Sulteng)</h3>
<div class="note-box">
  <strong>PROFIL EMPIRIS: Provinsi Sulawesi Tengah (Episentrum Hilirisasi & PLTU Captive)</strong><br>
  Kabupaten/Kota: 13 Daerah  |  Pusat Industri: Kawasan IMIP Morowali & Smelter Palu  |  Populasi BPS: 2.985.734 Jiwa<br>
  Karakteristik Krisis: Konsentrasi PLTU captive batubara terbesar nasional, hotspot satelit troposferik NO2 tertinggi, timbulan limbah B3 raksasa, dan laju deforestasi primer masif.
</div>

<h4>A. Narasi Temuan Lapangan Sulteng</h4>
<p>Hasil komputasi algoritma Z-Score EWM membuktikan bahwa <strong>Provinsi Sulawesi Tengah berada pada status RED ALERT (Skor Komposit 4.0 / 5.0 — Melampaui Batas)</strong>. Dari 20 indikator yang diuji, sebanyak 13 indikator berada pada kategori <strong>Melampaui Batas (Skor Likert 4.0 hingga 5.0)</strong>, dengan tekanan polusi udara dan perusakan lanskap daratan yang telah melampaui kapasitas asimilasi ekosistem.</p>

<h4>B. Matriks Hasil Uji Empiris (Sulteng)</h4>
<div class="table-caption">Tabel 6.13: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Tengah (Model Hybrid Z-Score & EWM)</div>
{html_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sulteng)}

<div class="table-caption">Tabel 6.14: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Tengah</div>
{html_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulteng"], rekap_sulteng)}

<h4>C. Analisis Temuan Empiris (Sulteng)</h4>
<p><strong>1. Daya Tampung Udara (Skor {sulteng['udara']:.1f} / 5 — {get_likert_label(sulteng['udara'])}):</strong> Beban PLTU captive batubara 7.325,0 MW (Z = +1.97&sigma;), limbah B3 25,30 Jt Ton (Z = +1.97&sigma;), emisi CO2 291,34 Jt Ton, dan anomali ISPA 3,50x lipat (Z = +1.66&sigma;).<br>
<strong>2. Daya Tampung Air (Skor {sulteng['air']:.1f} / 5 — {get_likert_label(sulteng['air'])}):</strong> Timbulan tailing/slag 24,50 Jt Ton (Z = +1.97&sigma;) dan morbiditas diare 1,52x lipat (Z = +1.34&sigma;).<br>
<strong>3. Daya Dukung Lahan (Skor {sulteng['lahan']:.1f} / 5 — {get_likert_label(sulteng['lahan'])}):</strong> Deforestasi primer 481.908 Ha (Z = +1.57&sigma;), perambahan hutan lindung 19.804 Ha (Z = +1.89&sigma;), dan 458 kejadian bencana hidrometeorologi.<br>
<strong>4. Daya Dukung Sosial (Skor {sulteng['sosial']:.1f} / 5 — {get_likert_label(sulteng['sosial'])}):</strong> 12.231 jiwa terdampak konflik agraria dan 6 insiden kriminalisasi pembela HAM.<br>
<strong>5. Veto Kebijakan (Skor {sulteng['veto']:.1f} / 5 — {get_likert_label(sulteng['veto'])}):</strong> Obral 260 IUP baru pasca-2014 (Z = +1.64&sigma;) dan impunitas korporat ilegal.<br>
<strong>6. Vonis Komposit Sulteng (Skor {sulteng['total_likert']:.1f} / 5.0 — {sulteng['likert_label']}):</strong> Status <span class="badge-danger">MELAMPAUI BATAS (RED ALERT)</span> membuktikan keruntuhan daya dukung lingkungan akibat ekspansi smelter nikel.</p>

<div class="divider"></div>

<!-- SEKSI 6.6.2 SULTRA -->
<h3>6.6.2 Evaluasi Empiris D3TLH: Provinsi Sulawesi Tenggara (Sultra)</h3>
<div class="note-box">
  <strong>PROFIL EMPIRIS: Provinsi Sulawesi Tenggara (Episentrum Konflik Agraria & Kepadatan IUP Ekstrem)</strong><br>
  Kabupaten/Kota: 17 Daerah  |  Pusat Industri: Smelter Morosi, Konawe, Kolaka & Pulau Wawonii  |  Populasi BPS: 2.624.875 Jiwa<br>
  Karakteristik Krisis: Kepadatan konsesi IUP tambang nikel tertinggi se-Sulawesi (11,72% daratan), korban perampasan ruang hidup terbesar (39.821 jiwa), sengketa ruang tangkap nelayan pesisir, dan pelanggaran persetujuan warga (FPIC) masif.
</div>

<h4>A. Narasi Temuan Lapangan Sultra</h4>
<p>Berdasarkan hasil pemetaan empiris Z-Score EWM, Provinsi Sulawesi Tenggara memperlihatkan profil anomali yang sangat kontras dengan Sulawesi Tengah. Jika Sulawesi Tengah didominasi oleh polusi PLTU dan deforestasi hulu, maka <strong>Sulawesi Tenggara mengalami ledakan krisis daya dukung sosial, perampasan ruang hidup masyarakat pesisir, dan kepadatan konsesi tambang tertinggi se-Pulau Sulawesi</strong>. Konsesi tambang nikel mencaplok <strong>11,72% daratan provinsi (Z = +1.50&sigma;, Likert 5.0)</strong> dan memicu perampasan ruang hidup terhadap <strong>39.821 jiwa (Z = +1.95&sigma;, Likert 5.0 — mencakup 73% korban se-Sulawesi)</strong>.</p>

<h4>B. Matriks Hasil Uji Empiris (Sultra)</h4>
<div class="table-caption">Tabel 6.15: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Tenggara (Model Hybrid Z-Score & EWM)</div>
{html_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sultra)}

<div class="table-caption">Tabel 6.16: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Tenggara</div>
{html_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sultra"], rekap_sultra)}

<h4>C. Analisis Temuan Empiris (Sultra)</h4>
<p><strong>1. Daya Tampung Udara (Skor {sultra['udara']:.1f} / 5 — {get_likert_label(sultra['udara'])}):</strong> Kapasitas 1.900 MW PLTU captive (Morosi/Konawe), emisi karbon 189,02 Jt Ton CO2e (Z = +0.59&sigma;), dan NO2 satelit 6,62e-06 mol/m².<br>
<strong>2. Daya Tampung Air (Skor {sultra['air']:.1f} / 5 — {get_likert_label(sultra['air'])}):</strong> IKA 65,32, beban tailing 6,52 Jt Ton, dan meletusnya 5 kasus konflik ruang tangkap laut nelayan vs tongkang nikel (Z = +0.87&sigma;, Likert 4.0).<br>
<strong>3. Daya Dukung Lahan (Skor {sultra['lahan']:.1f} / 5 — {get_likert_label(sultra['lahan'])}):</strong> Kepadatan Konsesi IUP mencapai 11,72% daratan provinsi (Z = +1.50&sigma;, Likert 5.0 - Outlier Ekstrem Se-Sulawesi) yang menggerus 337.434 Ha hutan alam primer.<br>
<strong>4. Daya Dukung Sosial (Skor {sultra['sosial']:.1f} / 5 — {get_likert_label(sultra['sosial'])}):</strong> Krisis sosial terparah se-Sulawesi dengan 39.821 jiwa warga terancam kehilangan ruang hidup (Z = +1.95&sigma;, Likert 5.0), 5 kasus manipulasi persetujuan FPIC (Z = +1.86&sigma;, Likert 5.0), dan defisit SPA 17,92%.<br>
<strong>5. Veto Kebijakan (Skor {sultra['veto']:.1f} / 5 — {get_likert_label(sultra['veto'])}):</strong> Obral 160 IUP baru pasca-2014 (Z = +0.64&sigma;, Likert 4.0).<br>
<strong>6. Vonis Komposit Sulawesi Tenggara (Skor {sultra['total_likert']:.1f} / 5.0 — {sultra['likert_label']}):</strong> Status <span class="badge-warning">MENDEKATI BATAS</span>, dengan catatan kritis bahwa Pilar Sosial (4.5 / 5) dan Kepadatan Konsesi Tambang (11,72%) telah berada pada status RED ALERT (MELAMPAUI BATAS EKSTREM).</p>

<div class="divider"></div>

<!-- SEKSI 6.6.3 SULSEL -->
<h3>6.6.3 Evaluasi Empiris D3TLH: Provinsi Sulawesi Selatan (Sulsel)</h3>
<div class="note-box">
  <strong>PROFIL EMPIRIS: Provinsi Sulawesi Selatan (Episentrum Bencana Alam, Konflik Pesisir & Kriminalisasi)</strong><br>
  Kabupaten/Kota: 24 Daerah  |  Pusat Industri: KIMA Makassar, Smelter Huadi Bantaeng, Vale Sorowako Luwu Timur & PLTU Jeneponto  |  Populasi BPS: 9.073.509 Jiwa<br>
  Karakteristik Krisis: Frekuensi bencana hidrometeorologi banjir bandang dan longsor tertinggi se-Sulawesi (669 kejadian), sengketa ruang laut nelayan pesisir terbanyak (7 kasus), insiden kriminalisasi pejuang HAM tertinggi (9 kasus), tambang ilegal marak (10 korporasi), dan cemaran karsinogenik Cr6+.
</div>

<h4>A. Narasi Temuan Lapangan Sulsel</h4>
<p>Sebagai provinsi dengan populasi terbesar (9,07 juta jiwa) dan pusat gravitasi ekonomi regional, <strong>Provinsi Sulawesi Selatan mencatat Skor Komposit 2.6 / 5.0 (Status: Mendekati Batas)</strong>. Kendati secara agregat tidak berada pada status Melampaui Batas layaknya Sulteng, <strong>audit forensik Z-score membongkar anomali outlier ekstrem pada 5 indikator kritis (Skor Likert 5.0 / Red Alert)</strong> yang memperlihatkan kerentanan ekologis struktural di kawasan pesisir, daerah aliran sungai (DAS), dan ruang hidup agraria.</p>
<p>Sulawesi Selatan mencatat rekor tertinggi se-Sulawesi pada tiga variabel destruktif sekaligus: <strong>kejadian bencana hidrometeorologi sebanyak 669 kali (Z = +1.63&sigma;, Likert 5.0)</strong>, meletusnya <strong>7 kasus konflik ruang tangkap laut nelayan vs tambang pasir laut dan tongkang (Z = +1.56&sigma;, Likert 5.0)</strong>, serta represi hukum dengan <strong>9 insiden kriminalisasi petani dan aktivis pembela HAM (Z = +1.57&sigma;, Likert 5.0)</strong>. Selain itu, maraknya operasi <strong>10 korporasi tambang ilegal di kawasan lindung (Z = +1.97&sigma;, Likert 5.0)</strong> dan cemaran Heksavalen Kromium (Cr6+) menegaskan darurat tata kelola lingkungan hidup di provinsi ini.</p>

<h4>B. Matriks Hasil Uji Empiris (Sulsel)</h4>
<div class="table-caption">Tabel 6.17: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Selatan (Model Hybrid Z-Score & EWM)</div>
{html_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sulsel)}

<div class="table-caption">Tabel 6.18: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Selatan</div>
{html_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulsel"], rekap_sulsel)}

<h4>C. Analisis Temuan Empiris (Sulsel)</h4>
<p><strong>1. Daya Tampung Udara (Skor {sulsel['udara']:.1f} / 5 — {get_likert_label(sulsel['udara'])}):</strong> Kapasitas PLTU captive beroperasi mencapai 600,0 MW (Punagaya Jeneponto & Barru, Z = -0.36&sigma;), emisi karbon 138,73 Jt Ton CO2e (Z = +0.05&sigma;), namun konsentrasi gas NO2 satelit menyentuh 6,40e-06 mol/m² (Z = +0.65&sigma;, Likert 4.0) yang mencerminkan beban emisi perkotaan dan industri KIMA.<br>
<strong>2. Daya Tampung Air (Skor {sulsel['air']:.1f} / 5 — {get_likert_label(sulsel['air'])}):</strong> Rerata IKA bernilai 58,50 poin (Z = +0.35&sigma;), namun terdeteksi kontaminasi toksik Heksavalen Kromium Cr6+ (Z = +2.03&sigma;, Likert 5.0) di perairan Luwu Timur, serta meletusnya 7 kasus konflik ruang laut nelayan pesisir vs tambang pasir laut dan tongkang (Z = +1.56&sigma;, Likert 5.0 — tertinggi se-Sulawesi).<br>
<strong>3. Daya Dukung Lahan (Skor {sulsel['lahan']:.1f} / 5 — {get_likert_label(sulsel['lahan'])}):</strong> Sulsel memikul bencana hidrometeorologi terparah se-Sulawesi dengan 669 kejadian banjir bandang dan longsor (Z = +1.63&sigma;, Likert 5.0), dipicu laju deforestasi primer 261.147 Ha (Z = +0.19&sigma;) dan perambahan 5.314 Ha hutan lindung di kawasan hulu DAS.<br>
<strong>4. Daya Dukung Sosial (Skor {sulsel['sosial']:.1f} / 5 — {get_likert_label(sulsel['sosial'])}):</strong> Meskipun jumlah warga terdampak langsung tercatat 2.257 jiwa (Z = -0.43&sigma;), Sulawesi Selatan mencatat eskalasi represi terberat dengan 9 insiden kriminalisasi warga dan aktivis lingkungan (Z = +1.57&sigma;, Likert 5.0 — rekor tertinggi se-Sulawesi).<br>
<strong>5. Veto Kebijakan (Skor {sulsel['veto']:.1f} / 5 — {get_likert_label(sulsel['veto'])}):</strong> Ditemukan 10 korporasi tambang beroperasi ilegal di kawasan hutan lindung dan DAS (Z = +1.97&sigma;, Likert 5.0 — tertinggi se-Sulawesi) serta diterbitkannya 105 IUP baru pasca-2014 (Z = +0.09&sigma;).<br>
<strong>6. Vonis Komposit Sulawesi Selatan (Skor {sulsel['total_likert']:.1f} / 5.0 — {sulsel['likert_label']}):</strong> Status <span class="badge-warning">MENDEKATI BATAS</span> (WSM: 5.29 / 10.0), dengan catatan kritis bahwa dimensi penegakan hukum tambang ilegal, frekuensi bencana alam, dan represi kriminalisasi warga telah berada pada status RED ALERT (MELAMPAUI BATAS EKSTREM).</p>

<div class="divider"></div>

<!-- SEKSI 6.6.4 SULBAR -->
<h3>6.6.4 Evaluasi Empiris D3TLH: Provinsi Sulawesi Barat (Sulbar)</h3>
<div class="note-box">
  <strong>PROFIL EMPIRIS: Provinsi Sulawesi Barat (Bioregion Non-Hilirisasi & Dominasi Agromaritim)</strong><br>
  Kabupaten/Kota: 6 Daerah  |  Basis Perekonomian: Perkebunan Sawit, Kakao & Perikanan Tangkap  |  Populasi BPS: 1.419.229 Jiwa<br>
  Karakteristik Krisis: Bebas dari ekspansi PLTU captive batubara (0 MW) dan nihil timbulan limbah B3/tailing nikel, namun menghadapi tekanan mutu air sungai (IKA 55,93 poin) akibat limbah PKS perkebunan sawit monokultur dan erosi DAS.
</div>

<h4>A. Narasi Temuan Lapangan Sulbar</h4>
<p>Sebagai wilayah pemekaran dengan 6 kabupaten dan populasi 1,42 juta jiwa, <strong>Provinsi Sulawesi Barat mencatatkan Skor Komposit 1.2 / 5.0 (Status: Tidak Melampaui Batas)</strong>. Profil ekologis Sulbar menjadi <strong>bukti empiris pembanding (control baseline)</strong> yang sangat berharga dalam laporan ini. Karena tidak menjadi lokasi hilirisasi industri nikel berskala raksasa, Sulbar terhindar dari akumulasi limbah B3, tailing laut, dan PLTU captive batubara yang menghancurkan daya lentur lingkungan hidup seperti di Sulteng dan Sultra.</p>
<p>Kepadatan konsesi tambang nikel di Sulbar tercatat hanya <strong>0,26% dari luas daratan provinsi (Z = -1.09&sigma;, Likert 0.0 — Terendah se-Sulawesi)</strong>. Namun demikian, evaluasi D3TLH mencatat <strong>dua peringatan ekologis lokal (local stressor)</strong>: yaitu penurunan Indeks Kualitas Air (IKA 55,93 poin, Z = +1.11&sigma;, Likert 5.0) dan angka morbiditas diare sebesar 1,27x lipat (Z = +0.63&sigma;, Likert 4.0), yang dipicu oleh pelepasan limbah cair pabrik kelapa sawit (PKS) monokultur di Pasangkayu dan Mamuju Tengah serta defisit sanitasi dasar perdesaan.</p>

<h4>B. Matriks Hasil Uji Empiris (Sulbar)</h4>
<div class="table-caption">Tabel 6.19: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Barat (Model Hybrid Z-Score & EWM)</div>
{html_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sulbar)}

<div class="table-caption">Tabel 6.20: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Barat</div>
{html_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulbar"], rekap_sulbar)}

<h4>C. Analisis Temuan Empiris (Sulbar)</h4>
<p><strong>1. Daya Tampung Udara (Skor {sulbar['udara']:.1f} / 5 — {get_likert_label(sulbar['udara'])}):</strong> Sulbar bebas dari beban PLTU captive batubara (0,0 MW, Z = -0.57&sigma;) dan nihil timbulan limbah B3 industri smelter (0,0 Ton, Z = -0.54&sigma;). Emisi karbon deforestasi tercatat 82,51 Jt Ton CO2e (Z = -0.55&sigma;) dan konsentrasi NO2 troposferik sebesar 6,00e-06 mol/m² (Z = +0.34&sigma;).<br>
<strong>2. Daya Tampung Air (Skor {sulbar['air']:.1f} / 5 — {get_likert_label(sulbar['air'])}):</strong> Nihil pembuangan tailing/slag tambang (0,0 Ton), namun mencatat IKA terendah se-Sulawesi (55,93 poin, Z = +1.11&sigma;, Likert 5.0) dan insidensi diare 1,27x lipat (Z = +0.63&sigma;, Likert 4.0) yang bersumber dari limbah cair PKS kelapa sawit dan sedimentasi erosi DAS Lariang & Karama.<br>
<strong>3. Daya Dukung Lahan (Skor {sulbar['lahan']:.1f} / 5 — {get_likert_label(sulbar['lahan'])}):</strong> Kepadatan konsesi tambang nikel paling rendah se-Sulawesi (hanya 0,26% daratan, Z = -1.09&sigma;, Likert 0.0), deforestasi primer 133.263 Ha (Z = -0.61&sigma;), perambahan hutan lindung 1.251 Ha, dan frekuensi bencana alam sebanyak 143 kejadian.<br>
<strong>4. Daya Dukung Sosial (Skor {sulbar['sosial']:.1f} / 5 — {get_likert_label(sulbar['sosial'])}):</strong> Minim konflik agraria struktural skala besar yang melibatkan industri tambang (1 jiwa terdampak, 1 insiden kriminalisasi, dan nihil sengketa FPIC), mencerminkan relasi sosial-ekologis yang relatif stabil.<br>
<strong>5. Veto Kebijakan (Skor {sulbar['veto']:.1f} / 5 — {get_likert_label(sulbar['veto'])}):</strong> Penerbitan konsesi tambang nikel baru pasca-2014 sangat terbatas (27 IUP, Z = -0.68&sigma;, Likert 1.0) dan tidak teridentifikasi adanya korporasi tambang ilegal skala besar yang beroperasi di kawasan hutan.<br>
<strong>6. Vonis Komposit Sulawesi Barat (Skor {sulbar['total_likert']:.1f} / 5.0 — {sulbar['likert_label']}):</strong> Status <span class="badge-success">TIDAK MELAMPAUI BATAS</span> (WSM: 2.36 / 10.0). Status terjaga ini membuktikan bahwa tanpa intervensi industri smelter nikel dan PLTU batubara, daya lentur lingkungan hidup regional dapat dipertahankan dalam ambang aman.</p>

<div class="divider"></div>

<!-- SEKSI 6.6.5 GORONTALO -->
<h3>6.6.5 Evaluasi Empiris D3TLH: Provinsi Gorontalo</h3>
<div class="note-box">
  <strong>PROFIL EMPIRIS: Provinsi Gorontalo (Bioregion Terjaga & Bebas Polusi Smelter Nikel)</strong><br>
  Kabupaten/Kota: 6 Daerah  |  Pusat Agraria: Lembah Pertanian Jagung & Danau Limboto  |  Populasi BPS: 1.171.681 Jiwa<br>
  Karakteristik Krisis: Kualitas atmosfer NO2 satelit paling bersih se-Sulawesi (3,76e-06 mol/m²), deforestasi primer dan emisi karbon terendah, nihil PLTU captive batubara maupun limbah B3 smelter, namun memikul anomali ISPA akibat faktor mikroklimat topografi cekungan dan residu pembakaran biomassa jagung.
</div>

<h4>A. Narasi Temuan Lapangan Gorontalo</h4>
<p>Sebagai provinsi dengan luas daratan dan populasi terkecil di Pulau Sulawesi (1,17 juta jiwa), <strong>Provinsi Gorontalo mencatatkan Skor Komposit 1.2 / 5.0 (Status: Tidak Melampaui Batas)</strong>. Bersama dengan Sulawesi Barat, Gorontalo berada pada kuadran <strong>ekologis terjaga (low-stress environment)</strong> yang mempertegas validitas model Z-Score EWM: ketika suatu wilayah tidak dieksploitasi oleh mega-proyek hilirisasi nikel dan PLTU captive batubara, integritas daya dukung lingkungannya tetap berada di bawah ambang batas bahaya.</p>
<p>Gorontalo membukukan rekor <strong>konsentrasi gas troposferik NO2 paling bersih se-Pulau Sulawesi (3,76e-06 mol/m², Z = -1.40&sigma;, Likert 0.0)</strong>, laju deforestasi primer paling rendah (98.063 Ha, Z = -0.83&sigma;), serta pelepasan emisi karbon deforestasi terendah (53,66 Jt Ton CO2e, Z = -0.85&sigma;). Kepadatan izin tambang nikel hanya mencapai <strong>0,46% dari daratan provinsi (Z = -1.04&sigma;, Likert 0.0)</strong>. Satu-satunya anomali yang mencuat adalah rasio morbiditas ISPA sebesar 2,41x lipat (Z = +0.79&sigma;, Likert 4.0) yang dipicu oleh faktor mikroklimat topografi cekungan Lembah Limboto, debu jalanan pedesaan, serta pembakaran residu tongkol jagung pascapanen, bukan dari cerobong batubara.</p>

<h4>B. Matriks Hasil Uji Empiris (Gorontalo)</h4>
<div class="table-caption">Tabel 6.21: Bedah Matematika 20 Indikator Empiris Provinsi Gorontalo (Model Hybrid Z-Score & EWM)</div>
{html_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_gorontalo)}

<div class="table-caption">Tabel 6.22: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Gorontalo</div>
{html_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Gorontalo"], rekap_gorontalo)}

<h4>C. Analisis Temuan Empiris (Gorontalo)</h4>
<p><strong>1. Daya Tampung Udara (Skor {gorontalo['udara']:.1f} / 5 — {get_likert_label(gorontalo['udara'])}):</strong> Nihil beban PLTU captive batubara (0,0 MW, Z = -0.57&sigma;) dan nihil limbah B3 smelter (0,0 Ton, Z = -0.54&sigma;). Tingkat polusi NO2 satelit paling rendah se-Sulawesi (3,76e-06 mol/m², Z = -1.40&sigma;, Likert 0.0). Catatan anomali ISPA 2,41x lipat (Z = +0.79&sigma;, Likert 4.0) berkorelasi dengan asap pembakaran biomassa jagung dan dinamika inversi suhu Lembah Limboto.<br>
<strong>2. Daya Tampung Air (Skor {gorontalo['air']:.1f} / 5 — {get_likert_label(gorontalo['air'])}):</strong> Nihil pembuangan tailing nikel (0,0 Ton) dan nihil konflik ruang laut pesisir. IKA berada pada level 58,14 poin (Z = +0.46&sigma;, Likert 3.0) akibat tekanan sedimentasi erosi DAS Bone-Bolango dan eutrofikasi Danau Limboto dari limbah domestik.<br>
<strong>3. Daya Dukung Lahan (Skor {gorontalo['lahan']:.1f} / 5 — {get_likert_label(gorontalo['lahan'])}):</strong> Gorontalo membukukan kehilangan tutupan hutan alam primer terendah se-Sulawesi (98.063 Ha, Z = -0.83&sigma;, Likert 1.0) dengan kepadatan konsesi tambang nikel hanya 0,46% daratan (Z = -1.04&sigma;, Likert 0.0) serta nihil catatan bencana longsor/banjir skala masif dalam periode audit.<br>
<strong>4. Daya Dukung Sosial (Skor {gorontalo['sosial']:.1f} / 5 — {get_likert_label(gorontalo['sosial'])}):</strong> Bebas dari letupan konflik agraria industri tambang (nihil warga terdampak kehilangan ruang hidup, nihil sengketa FPIC, dan nihil kriminalisasi pejuang lingkungan).<br>
<strong>5. Veto Kebijakan (Skor {gorontalo['veto']:.1f} / 5 — {get_likert_label(gorontalo['veto'])}):</strong> Pemerintah daerah mencatat laju obral izin tambang terendah se-Sulawesi (hanya 7 IUP baru pasca-2014, Z = -0.88&sigma;, Likert 1.0) dan hanya teridentifikasi 1 aktivitas tambang rakyat/ilegal skala kecil di kawasan hulu (Z = -0.45&sigma;).<br>
<strong>6. Vonis Komposit Gorontalo (Skor {gorontalo['total_likert']:.1f} / 5.0 — {gorontalo['likert_label']}):</strong> Status <span class="badge-success">TIDAK MELAMPAUI BATAS</span> (WSM: 2.31 / 10.0). Status aman ini menjadi bukti konklusif bahwa kelestarian bioregion Sulawesi bertumpu pada pembatasan ekspansi industri ekstraktif nikel.</p>

</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab6_Skoring_Provinsi.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    # =========================================================================
    # [4/4] MEMBANGUN DOKUMEN MARKDOWN SATUAN
    # =========================================================================
    md_lines = [
        "# BAB VI: AUDIT FORENSIK METODOLOGI D3TLH",
        "## SUB-BAB 6.6: ALGORITMA SKORING TINGKAT PROVINSI (MODEL HYBRID Z-SCORE & EWM)",
        "",
        '> **KERANGKA METODOLOGI MULTIKRITERIA REGIONAL (BERLAKU MENGIKAT UNTUK 6 PROVINSI)**  ',
        '> Model evaluasi daya dukung dan daya tampung lingkungan hidup tingkat provinsi dirancang menggunakan pendekatan terstandarisasi berbasis Hybrid Z-Score Anomali Deviasi Standar dan Entropy Weight Method (EWM) sesuai Nature Scientific Reports (Sun et al., 2024). Metodologi, formula normalisasi, matriks entropi, dan bobot objektif indikator dihitung secara simultan dari matriks 6 provinsi se-Pulau Sulawesi.',
        "",
        "#### A. Pengantar & Kerangka Narasi Metodologis",
        "Sebagaimana ditampilkan pada antarmuka Streamlit **Dashboard Page 6 (Audit D3TLH - Tab Bedah Matematika Z-Score + EWM per Provinsi)**, evaluasi tingkat provinsi bertujuan untuk mengatasi kelemahan mendasar dokumen AMDAL dan D3TLH konvensional yang kerap mengaburkan krisis lingkungan lokal melalui teknik perataan agregat wilayah (*dilution effect*). Dalam metodologi pemerintah, beban pencemaran masif di suatu kawasan industri tambang sering kali tampak 'aman' hanya karena dirata-ratakan dengan luas daratan pulau secara keseluruhan.",
        "",
        "Untuk mendobrak bias tersebut, riset ini menerapkan **Model Hybrid Z-Score Anomali dan Pembobotan Objektif Entropi (EWM)**. Pendekatan ini secara otomatis memberikan bobot evaluasi tertinggi pada indikator-indikator yang memiliki tingkat ketimpangan spasial paling ekstrem (seperti timbulan limbah B3, tailing tambang, korban krisis agraria, dan PLTU captive batubara). Dengan demikian, provinsi yang menjadi episentrum industri ekstraktif terdeteksi secara akurat berada pada status anomali krisis tanpa terdistorsi oleh luas wilayah administratif.",
        "",
        "#### B. Alur Logika Metodologis Regional (Flowchart 6 Provinsi)",
        "```mermaid",
        mermaid_regional,
        "```",
        "",
        "#### C. Formulasi Matematis Universal & Definisi Variabel",
        "```text",
        "1. Z-Score Regional: Z_ij = (x_ij - mean(x_j)) / std(x_j)  |  Khusus IKA: Z_ika = - (ika_i - mean(ika)) / std(ika)",
        "2. Min-Max Normalisasi: r_ij = (x_ij - min(x_j)) / (max(x_j) - min(x_j))",
        "3. Proporsi Probabilitas: P_ij = r_ij / SUM(r_ij)",
        "4. Entropi Informasi Shannon: E_j = - (1 / ln(n)) * SUM(P_ij * ln(P_ij + eps))",
        "5. Koefisien Dispersi Informasi: D_j = 1 - E_j",
        "6. Bobot Objektif EWM Final: W_j = D_j / SUM(D_j)",
        "7. Mapping Likert Diskret: Z >= 1.0 -> 5.0 ; 0.5 <= Z < 1.0 -> 4.0 ; 0.0 <= Z < 0.5 -> 3.0 ; -0.5 <= Z < 0.0 -> 2.0 ; -1.0 <= Z < -0.5 -> 1.0 ; Z < -1.0 -> 0.0",
        "8. EWM Weighted Average Pilar: Skor_Pilar = SUM(L_ij * W_j) / SUM(W_j)",
        "9. Skor Komposit Total: (Udara + Air + Lahan + Sosial + Veto) / 5.0",
        "```",
        "",
        "```text",
        "Contoh Persamaan Substitusi Riil (Indikator PLTU Captive & Komposit Sulteng):",
        "1. Substitusi Z-Score: Z = (7.325 MW - 1.637,50 MW) / 2.882,26 MW = +1,97σ",
        "2. Substitusi EWM Shannon: r_sulteng = 1,000 ; P_sulteng = 0,745 -> Ej = 0,3948 -> Dj = 0,6052 -> W_pltu = 0,6052 / 7,8331 = 0,0773 (7,73%)",
        "3. Substitusi Likert Diskret: Z = +1,97σ >= +1,0σ -> Skor Likert = 5,0 / 5 (Melampaui Batas / Red Alert)",
        "4. Substitusi Pilar Udara: Skor = [(5,0*0,0773) + (4,0*0,0224) + (5,0*0,0461) + (5,0*0,0829) + (5,0*0,0395)] / 0,2682 = 4,92 / 5",
        "5. Substitusi Komposit Total: Skor = (4,92 + 3,30 + 4,70 + 2,50 + 4,40) / 5 = 3,96 / 5,0 -> WSM: 7,92 / 10.0 (Melampaui Batas)",
        "```",
        "",
        "##### Tabel 6.12: Matriks Parameter Regional Se-Sulawesi (Rata-rata, Deviasi Standar, dan Bobot Objektif EWM 20 Indikator Empiris)",
        markdown_table(["Pilar", "Indikator Empiris", "Rata-rata (B)", "Deviasi (C)", "Entropi (Ej)", "Divergensi (Dj)", "Bobot EWM (Wj)"], table_regional_meta),
        "",
        "---",
        "",
        "### 6.6.1 Evaluasi Empiris D3TLH: Provinsi Sulawesi Tengah (Sulteng)",
        '> **PROFIL EMPIRIS: Provinsi Sulawesi Tengah (Episentrum Hilirisasi & PLTU Captive)**  ',
        '> Kabupaten/Kota: 13 Daerah  |  Pusat Industri: Kawasan IMIP Morowali & Smelter Palu  |  Populasi BPS: 2.985.734 Jiwa  ',
        '> Karakteristik Krisis: Konsentrasi PLTU captive batubara terbesar nasional, hotspot satelit troposferik NO2 tertinggi, timbulan limbah B3 raksasa, dan laju deforestasi primer masif.',
        "",
        "#### A. Narasi Temuan Lapangan Sulteng",
        "Hasil komputasi algoritma Z-Score EWM membuktikan bahwa **Provinsi Sulawesi Tengah berada pada status RED ALERT (Skor Komposit 4.0 / 5.0 — Melampaui Batas)**. Dari 20 indikator yang diuji, sebanyak 13 indikator berada pada kategori **Melampaui Batas (Skor Likert 4.0 hingga 5.0)**, dengan tekanan polusi udara dan perusakan lanskap daratan yang telah melampaui kapasitas asimilasi ekosistem.",
        "",
        "#### B. Matriks Hasil Uji Empiris (Sulteng)",
        "##### Tabel 6.13: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Tengah (Model Hybrid Z-Score & EWM)",
        markdown_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sulteng),
        "",
        "##### Tabel 6.14: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Tengah",
        markdown_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulteng"], rekap_sulteng),
        "",
        "#### C. Analisis Temuan Empiris (Sulteng)",
        f"1. **Daya Tampung Udara (Skor {sulteng['udara']:.1f} / 5 — {get_likert_label(sulteng['udara'])}):** Beban PLTU captive batubara 7.325,0 MW (Z = +1.97σ), timbulan limbah B3 25,30 Jt Ton (Z = +1.97σ), emisi CO2 291,34 Jt Ton, dan anomali ISPA 3,50x lipat (Z = +1.66σ).",
        f"2. **Daya Tampung Air (Skor {sulteng['air']:.1f} / 5 — {get_likert_label(sulteng['air'])}):** Timbulan tailing/slag 24,50 Jt Ton (Z = +1.97σ) dan morbiditas diare 1,52x lipat (Z = +1.34σ).",
        f"3. **Daya Dukung Lahan (Skor {sulteng['lahan']:.1f} / 5 — {get_likert_label(sulteng['lahan'])}):** Deforestasi primer 481.908 Ha (Z = +1.57σ), perambahan hutan lindung 19.804 Ha (Z = +1.89σ), dan 458 kejadian bencana hidrometeorologi.",
        f"4. **Daya Dukung Sosial (Skor {sulteng['sosial']:.1f} / 5 — {get_likert_label(sulteng['sosial'])}):** 12.231 jiwa terdampak konflik agraria dan 6 insiden kriminalisasi pembela HAM.",
        f"5. **Veto Kebijakan (Skor {sulteng['veto']:.1f} / 5 — {get_likert_label(sulteng['veto'])}):** Obral 260 IUP baru pasca-2014 (Z = +1.64σ) dan impunitas korporat ilegal.",
        f"6. **Vonis Komposit Sulteng (Skor {sulteng['total_likert']:.1f} / 5.0 — {sulteng['likert_label']}):** Status **Melampaui Batas** *(STATUS: RED ALERT)* membuktikan keruntuhan daya dukung lingkungan akibat ekspansi smelter nikel.",
        "",
        "---",
        "",
        "### 6.6.2 Evaluasi Empiris D3TLH: Provinsi Sulawesi Tenggara (Sultra)",
        '> **PROFIL EMPIRIS: Provinsi Sulawesi Tenggara (Episentrum Konflik Agraria & Kepadatan IUP Ekstrem)**  ',
        '> Kabupaten/Kota: 17 Daerah  |  Pusat Industri: Smelter Morosi, Konawe, Kolaka & Pulau Wawonii  |  Populasi BPS: 2.624.875 Jiwa  ',
        '> Karakteristik Krisis: Kepadatan konsesi IUP tambang nikel tertinggi se-Sulawesi (11,72% daratan), korban perampasan ruang hidup terbesar (39.821 jiwa), sengketa ruang tangkap nelayan pesisir, dan pelanggaran persetujuan warga (FPIC) masif.',
        "",
        "#### A. Narasi Temuan Lapangan Sultra",
        "Berdasarkan hasil pemetaan empiris Z-Score EWM, Provinsi Sulawesi Tenggara memperlihatkan profil anomali yang sangat kontras dengan Sulawesi Tengah. Jika Sulawesi Tengah didominasi oleh polusi PLTU dan deforestasi hulu, maka **Sulawesi Tenggara mengalami ledakan krisis daya dukung sosial, perampasan ruang hidup masyarakat pesisir, dan kepadatan konsesi tambang tertinggi se-Pulau Sulawesi**. Konsesi tambang nikel mencaplok **11,72% daratan provinsi (Z = +1.50σ, Likert 5.0)** dan memicu perampasan ruang hidup terhadap **39.821 jiwa (Z = +1.95σ, Likert 5.0 — mencakup 73% korban se-Sulawesi)**.",
        "",
        "#### B. Matriks Hasil Uji Empiris (Sultra)",
        "##### Tabel 6.15: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Tenggara (Model Hybrid Z-Score & EWM)",
        markdown_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sultra),
        "",
        "##### Tabel 6.16: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Tenggara",
        markdown_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sultra"], rekap_sultra),
        "",
        "#### C. Analisis Temuan Empiris (Sultra)",
        f"1. **Daya Tampung Udara (Skor {sultra['udara']:.1f} / 5 — {get_likert_label(sultra['udara'])}):** Kapasitas 1.900 MW PLTU captive (Morosi/Konawe), emisi karbon 189,02 Jt Ton CO2e (Z = +0.59σ), dan NO2 satelit 6,62e-06 mol/m².",
        f"2. **Daya Tampung Air (Skor {sultra['air']:.1f} / 5 — {get_likert_label(sultra['air'])}):** IKA 65,32, beban tailing 6,52 Jt Ton, dan meletusnya 5 kasus konflik ruang tangkap laut nelayan vs tongkang nikel (Z = +0.87σ, Likert 4.0).",
        f"3. **Daya Dukung Lahan (Skor {sultra['lahan']:.1f} / 5 — {get_likert_label(sultra['lahan'])}):** Kepadatan Konsesi IUP mencapai 11,72% daratan provinsi (Z = +1.50σ, Likert 5.0 - Outlier Ekstrem Se-Sulawesi) yang menggerus 337.434 Ha hutan alam primer.",
        f"4. **Daya Dukung Sosial (Skor {sultra['sosial']:.1f} / 5 — {get_likert_label(sultra['sosial'])}):** Krisis sosial terparah se-Sulawesi dengan 39.821 jiwa warga terancam kehilangan ruang hidup (Z = +1.95σ, Likert 5.0), 5 kasus manipulasi persetujuan FPIC (Z = +1.86σ, Likert 5.0), dan defisit SPA 17,92%.",
        f"5. **Veto Kebijakan (Skor {sultra['veto']:.1f} / 5 — {get_likert_label(sultra['veto'])}):** Obral 160 IUP baru pasca-2014 (Z = +0.64σ, Likert 4.0).",
        f"6. **Vonis Komposit Sulawesi Tenggara (Skor {sultra['total_likert']:.1f} / 5.0 — {sultra['likert_label']}):** Status **Mendekati Batas**, dengan catatan kritis bahwa Pilar Sosial (4.5 / 5) dan Kepadatan Konsesi Tambang (11,72%) telah berada pada status **Melampaui Batas Ekstrem** *(RED ALERT)*.",
        "",
        "---",
        "",
        "### 6.6.3 Evaluasi Empiris D3TLH: Provinsi Sulawesi Selatan (Sulsel)",
        '> **PROFIL EMPIRIS: Provinsi Sulawesi Selatan (Episentrum Bencana Alam, Konflik Pesisir & Kriminalisasi)**  ',
        '> Kabupaten/Kota: 24 Daerah  |  Pusat Industri: KIMA Makassar, Smelter Huadi Bantaeng, Vale Sorowako Luwu Timur & PLTU Jeneponto  |  Populasi BPS: 9.073.509 Jiwa  ',
        '> Karakteristik Krisis: Frekuensi bencana hidrometeorologi banjir bandang dan longsor tertinggi se-Sulawesi (669 kejadian), sengketa ruang laut nelayan pesisir terbanyak (7 kasus), insiden kriminalisasi pejuang HAM tertinggi (9 kasus), tambang ilegal marak (10 korporasi), dan cemaran karsinogenik Cr6+.',
        "",
        "#### A. Narasi Temuan Lapangan Sulsel",
        "Sebagai provinsi dengan populasi terbesar (9,07 juta jiwa) dan pusat gravitasi ekonomi regional, **Provinsi Sulawesi Selatan mencatat Skor Komposit 2.6 / 5.0 (Status: Mendekati Batas)**. Kendati secara agregat tidak berada pada status Melampaui Batas layaknya Sulteng, **audit forensik Z-score membongkar anomali outlier ekstrem pada 5 indikator kritis (Skor Likert 5.0 / Red Alert)** yang memperlihatkan kerentanan ekologis struktural di kawasan pesisir, daerah aliran sungai (DAS), dan ruang hidup agraria.",
        "",
        "Sulawesi Selatan mencatat rekor tertinggi se-Sulawesi pada tiga variabel destruktif sekaligus: **kejadian bencana hidrometeorologi sebanyak 669 kali (Z = +1.63σ, Likert 5.0)**, meletusnya **7 kasus konflik ruang tangkap laut nelayan vs tambang pasir laut dan tongkang (Z = +1.56σ, Likert 5.0)**, serta represi hukum dengan **9 insiden kriminalisasi petani dan aktivis pembela HAM (Z = +1.57σ, Likert 5.0)**. Selain itu, maraknya operasi **10 korporasi tambang ilegal di kawasan lindung (Z = +1.97σ, Likert 5.0)** dan cemaran Heksavalen Kromium (Cr6+) menegaskan darurat tata kelola lingkungan hidup di provinsi ini.",
        "",
        "#### B. Matriks Hasil Uji Empiris (Sulsel)",
        "##### Tabel 6.17: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Selatan (Model Hybrid Z-Score & EWM)",
        markdown_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sulsel),
        "",
        "##### Tabel 6.18: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Selatan",
        markdown_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulsel"], rekap_sulsel),
        "",
        "#### C. Analisis Temuan Empiris (Sulsel)",
        f"1. **Daya Tampung Udara (Skor {sulsel['udara']:.1f} / 5 — {get_likert_label(sulsel['udara'])}):** Kapasitas PLTU captive beroperasi mencapai 600,0 MW (Punagaya Jeneponto & Barru, Z = -0.36σ), emisi karbon 138,73 Jt Ton CO2e (Z = +0.05σ), namun konsentrasi gas NO2 satelit menyentuh 6,40e-06 mol/m² (Z = +0.65σ, Likert 4.0) yang mencerminkan beban emisi perkotaan dan industri KIMA.",
        f"2. **Daya Tampung Air (Skor {sulsel['air']:.1f} / 5 — {get_likert_label(sulsel['air'])}):** Rerata IKA bernilai 58,50 poin (Z = +0.35σ), namun terdeteksi kontaminasi toksik Heksavalen Kromium Cr6+ (Z = +2.03σ, Likert 5.0) di perairan Luwu Timur, serta meletusnya 7 kasus konflik ruang laut nelayan pesisir vs tambang pasir laut dan tongkang (Z = +1.56σ, Likert 5.0 — tertinggi se-Sulawesi).",
        f"3. **Daya Dukung Lahan (Skor {sulsel['lahan']:.1f} / 5 — {get_likert_label(sulsel['lahan'])}):** Sulsel memikul bencana hidrometeorologi terparah se-Sulawesi dengan 669 kejadian banjir bandang dan longsor (Z = +1.63σ, Likert 5.0), dipicu laju deforestasi primer 261.147 Ha (Z = +0.19σ) dan perambahan 5.314 Ha hutan lindung di kawasan hulu DAS.",
        f"4. **Daya Dukung Sosial (Skor {sulsel['sosial']:.1f} / 5 — {get_likert_label(sulsel['sosial'])}):** Meskipun jumlah warga terdampak langsung tercatat 2.257 jiwa (Z = -0.43σ), Sulawesi Selatan mencatat eskalasi represi terberat dengan 9 insiden kriminalisasi warga dan aktivis lingkungan (Z = +1.57σ, Likert 5.0 — rekor tertinggi se-Sulawesi).",
        f"5. **Veto Kebijakan (Skor {sulsel['veto']:.1f} / 5 — {get_likert_label(sulsel['veto'])}):** Ditemukan 10 korporasi tambang beroperasi ilegal di kawasan hutan lindung dan DAS (Z = +1.97σ, Likert 5.0 — tertinggi se-Sulawesi) serta diterbitkannya 105 IUP baru pasca-2014 (Z = +0.09σ).",
        f"6. **Vonis Komposit Sulawesi Selatan (Skor {sulsel['total_likert']:.1f} / 5.0 — {sulsel['likert_label']}):** Status **Mendekati Batas** (WSM: 5.29 / 10.0), dengan catatan kritis bahwa dimensi penegakan hukum tambang ilegal, frekuensi bencana alam, dan represi kriminalisasi warga telah berada pada status **Melampaui Batas Ekstrem (RED ALERT)**.",
        "",
        "---",
        "",
        "### 6.6.4 Evaluasi Empiris D3TLH: Provinsi Sulawesi Barat (Sulbar)",
        '> **PROFIL EMPIRIS: Provinsi Sulawesi Barat (Bioregion Non-Hilirisasi & Dominasi Agromaritim)**  ',
        '> Kabupaten/Kota: 6 Daerah  |  Basis Perekonomian: Perkebunan Sawit, Kakao & Perikanan Tangkap  |  Populasi BPS: 1.419.229 Jiwa  ',
        '> Karakteristik Krisis: Bebas dari ekspansi PLTU captive batubara (0 MW) dan nihil timbulan limbah B3/tailing nikel, namun menghadapi tekanan mutu air sungai (IKA 55,93 poin) akibat limbah PKS perkebunan sawit monokultur dan erosi DAS.',
        "",
        "#### A. Narasi Temuan Lapangan Sulbar",
        "Sebagai wilayah pemekaran dengan 6 kabupaten dan populasi 1,42 juta jiwa, **Provinsi Sulawesi Barat mencatatkan Skor Komposit 1.2 / 5.0 (Status: Tidak Melampaui Batas)**. Profil ekologis Sulbar menjadi **bukti empiris pembanding (control baseline)** yang sangat berharga dalam laporan ini. Karena tidak menjadi lokasi hilirisasi industri nikel berskala raksasa, Sulbar terhindar dari akumulasi limbah B3, tailing laut, dan PLTU captive batubara yang menghancurkan daya lentur lingkungan hidup seperti di Sulteng dan Sultra.",
        "",
        "Kepadatan konsesi tambang nikel di Sulbar tercatat hanya **0,26% dari luas daratan provinsi (Z = -1.09σ, Likert 0.0 — Terendah se-Sulawesi)**. Namun demikian, evaluasi D3TLH mencatat **dua peringatan ekologis lokal (local stressor)**: yaitu penurunan Indeks Kualitas Air (IKA 55,93 poin, Z = +1.11σ, Likert 5.0) dan angka morbiditas diare sebesar 1,27x lipat (Z = +0.63σ, Likert 4.0), yang dipicu oleh pelepasan limbah cair pabrik kelapa sawit (PKS) monokultur di Pasangkayu dan Mamuju Tengah serta defisit sanitasi dasar perdesaan.",
        "",
        "#### B. Matriks Hasil Uji Empiris (Sulbar)",
        "##### Tabel 6.19: Bedah Matematika 20 Indikator Empiris Provinsi Sulawesi Barat (Model Hybrid Z-Score & EWM)",
        markdown_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_sulbar),
        "",
        "##### Tabel 6.20: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Sulawesi Barat",
        markdown_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Sulbar"], rekap_sulbar),
        "",
        "#### C. Analisis Temuan Empiris (Sulbar)",
        f"1. **Daya Tampung Udara (Skor {sulbar['udara']:.1f} / 5 — {get_likert_label(sulbar['udara'])}):** Sulbar bebas dari beban PLTU captive batubara (0,0 MW, Z = -0.57σ) dan nihil timbulan limbah B3 industri smelter (0,0 Ton, Z = -0.54σ). Emisi karbon deforestasi tercatat 82,51 Jt Ton CO2e (Z = -0.55σ) dan konsentrasi NO2 troposferik sebesar 6,00e-06 mol/m² (Z = +0.34σ).",
        f"2. **Daya Tampung Air (Skor {sulbar['air']:.1f} / 5 — {get_likert_label(sulbar['air'])}):** Nihil pembuangan tailing/slag tambang (0,0 Ton), namun mencatat IKA terendah se-Sulawesi (55,93 poin, Z = +1.11σ, Likert 5.0) dan insidensi diare 1,27x lipat (Z = +0.63σ, Likert 4.0) yang bersumber dari limbah cair PKS kelapa sawit dan sedimentasi erosi DAS Lariang & Karama.",
        f"3. **Daya Dukung Lahan (Skor {sulbar['lahan']:.1f} / 5 — {get_likert_label(sulbar['lahan'])}):** Kepadatan konsesi tambang nikel paling rendah se-Sulawesi (hanya 0,26% daratan, Z = -1.09σ, Likert 0.0), deforestasi primer 133.263 Ha (Z = -0.61σ), perambahan hutan lindung 1.251 Ha, dan frekuensi bencana alam sebanyak 143 kejadian.",
        f"4. **Daya Dukung Sosial (Skor {sulbar['sosial']:.1f} / 5 — {get_likert_label(sulbar['sosial'])}):** Minim konflik agraria struktural skala besar yang melibatkan industri tambang (1 jiwa terdampak, 1 insiden kriminalisasi, dan nihil sengketa FPIC), mencerminkan relasi sosial-ekologis yang relatif stabil.",
        f"5. **Veto Kebijakan (Skor {sulbar['veto']:.1f} / 5 — {get_likert_label(sulbar['veto'])}):** Penerbitan konsesi tambang nikel baru pasca-2014 sangat terbatas (27 IUP, Z = -0.68σ, Likert 1.0) dan tidak teridentifikasi adanya korporasi tambang ilegal skala besar yang beroperasi di kawasan hutan.",
        f"6. **Vonis Komposit Sulawesi Barat (Skor {sulbar['total_likert']:.1f} / 5.0 — {sulbar['likert_label']}):** Status **Tidak Melampaui Batas** (WSM: 2.36 / 10.0). Status terjaga ini membuktikan bahwa tanpa intervensi industri smelter nikel dan PLTU batubara, daya lentur lingkungan hidup regional dapat dipertahankan dalam ambang aman.",
        "",
        "---",
        "",
        "### 6.6.5 Evaluasi Empiris D3TLH: Provinsi Gorontalo",
        '> **PROFIL EMPIRIS: Provinsi Gorontalo (Bioregion Terjaga & Bebas Polusi Smelter Nikel)**  ',
        '> Kabupaten/Kota: 6 Daerah  |  Pusat Agraria: Lembah Pertanian Jagung & Danau Limboto  |  Populasi BPS: 1.171.681 Jiwa  ',
        '> Karakteristik Krisis: Kualitas atmosfer NO2 satelit paling bersih se-Sulawesi (3,76e-06 mol/m²), deforestasi primer dan emisi karbon terendah, nihil PLTU captive batubara maupun limbah B3 smelter, namun memikul anomali ISPA akibat faktor mikroklimat topografi cekungan dan residu pembakaran biomassa jagung.',
        "",
        "#### A. Narasi Temuan Lapangan Gorontalo",
        "Sebagai provinsi dengan luas daratan dan populasi terkecil di Pulau Sulawesi (1,17 juta jiwa), **Provinsi Gorontalo mencatatkan Skor Komposit 1.2 / 5.0 (Status: Tidak Melampaui Batas)**. Bersama dengan Sulawesi Barat, Gorontalo berada pada kuadran **ekologis terjaga (low-stress environment)** yang mempertegas validitas model Z-Score EWM: ketika suatu wilayah tidak dieksploitasi oleh mega-proyek hilirisasi nikel dan PLTU captive batubara, integritas daya dukung lingkungannya tetap berada di bawah ambang batas bahaya.",
        "",
        "Gorontalo membukukan rekor **konsentrasi gas troposferik NO2 paling bersih se-Pulau Sulawesi (3,76e-06 mol/m², Z = -1.40σ, Likert 0.0)**, laju deforestasi primer paling rendah (98.063 Ha, Z = -0.83σ), serta pelepasan emisi karbon deforestasi terendah (53,66 Jt Ton CO2e, Z = -0.85σ). Kepadatan izin tambang nikel hanya mencapai **0,46% dari daratan provinsi (Z = -1.04σ, Likert 0.0)**. Satu-satunya anomali yang mencuat adalah rasio morbiditas ISPA sebesar 2,41x lipat (Z = +0.79σ, Likert 4.0) yang dipicu oleh faktor mikroklimat topografi cekungan Lembah Limboto, debu jalanan pedesaan, serta pembakaran residu tongkol jagung pascapanen, bukan dari cerobong batubara.",
        "",
        "#### B. Matriks Hasil Uji Empiris (Gorontalo)",
        "##### Tabel 6.21: Bedah Matematika 20 Indikator Empiris Provinsi Gorontalo (Model Hybrid Z-Score & EWM)",
        markdown_table(["Pilar", "Indikator Empiris", "Fakta Mentah (A)", "Nilai Z-Score", "Bobot EWM", "Skor Likert", "Status Ekologis"], table_eval_gorontalo),
        "",
        "##### Tabel 6.22: Rekapitulasi Skor 5 Pilar & Status Ekologis Komposit Provinsi Gorontalo",
        markdown_table(["Pilar / Dimensi", "Cakupan Indikator Kunci", "Skor Likert Pilar (0-5)", "Status Ekologis", "Interpretasi Temuan Lapangan Gorontalo"], rekap_gorontalo),
        "",
        "#### C. Analisis Temuan Empiris (Gorontalo)",
        f"1. **Daya Tampung Udara (Skor {gorontalo['udara']:.1f} / 5 — {get_likert_label(gorontalo['udara'])}):** Nihil beban PLTU captive batubara (0,0 MW, Z = -0.57σ) dan nihil limbah B3 smelter (0,0 Ton, Z = -0.54σ). Tingkat polusi NO2 satelit paling rendah se-Sulawesi (3,76e-06 mol/m², Z = -1.40σ, Likert 0.0). Catatan anomali ISPA 2,41x lipat (Z = +0.79σ, Likert 4.0) berkorelasi dengan asap pembakaran biomassa jagung dan dinamika inversi suhu Lembah Limboto.",
        f"2. **Daya Tampung Air (Skor {gorontalo['air']:.1f} / 5 — {get_likert_label(gorontalo['air'])}):** Nihil pembuangan tailing nikel (0,0 Ton) dan nihil konflik ruang laut pesisir. IKA berada pada level 58,14 poin (Z = +0.46σ, Likert 3.0) akibat tekanan sedimentasi erosi DAS Bone-Bolango dan eutrofikasi Danau Limboto dari limbah domestik.",
        f"3. **Daya Dukung Lahan (Skor {gorontalo['lahan']:.1f} / 5 — {get_likert_label(gorontalo['lahan'])}):** Gorontalo membukukan kehilangan tutupan hutan alam primer terendah se-Sulawesi (98.063 Ha, Z = -0.83σ, Likert 1.0) dengan kepadatan konsesi tambang nikel hanya 0,46% daratan (Z = -1.04σ, Likert 0.0) serta nihil catatan bencana longsor/banjir skala masif dalam periode audit.",
        f"4. **Daya Dukung Sosial (Skor {gorontalo['sosial']:.1f} / 5 — {get_likert_label(gorontalo['sosial'])}):** Bebas dari letupan konflik agraria industri tambang (nihil warga terdampak kehilangan ruang hidup, nihil sengketa FPIC, dan nihil kriminalisasi pejuang lingkungan).",
        f"5. **Veto Kebijakan (Skor {gorontalo['veto']:.1f} / 5 — {get_likert_label(gorontalo['veto'])}):** Pemerintah daerah mencatat laju obral izin tambang terendah se-Sulawesi (hanya 7 IUP baru pasca-2014, Z = -0.88σ, Likert 1.0) dan hanya teridentifikasi 1 aktivitas tambang rakyat/ilegal skala kecil di kawasan hulu (Z = -0.45σ).",
        f"6. **Vonis Komposit Gorontalo (Skor {gorontalo['total_likert']:.1f} / 5.0 — {gorontalo['likert_label']}):** Status **Tidak Melampaui Batas** (WSM: 2.31 / 10.0). Status aman ini menjadi bukti konklusif bahwa kelestarian bioregion Sulawesi bertumpu pada pembatasan ekspansi industri ekstraktif nikel.",
        ""
    ]

    md_path = tool_dir / "Metodologi_Bab6_Skoring_Provinsi.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")
    print("[4/4] Selesai membangun Sub-bab 6.6 untuk Provinsi Sulawesi Tengah, Sulawesi Tenggara, Sulawesi Selatan, Sulawesi Barat, dan Gorontalo.")


if __name__ == "__main__":
    generate_bab6_skoring_provinsi()
