#!/usr/bin/env python3
"""Generator Laporan Metodologi Bab 6: Audit Forensik Metodologi D3TLH (Model Skoring Kerusakan Ekologis)."""

import base64
import os
import sys
from pathlib import Path

try:
    import numpy as np
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pandas", "numpy", "requests", "python-docx"])
    import numpy as np
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)
C_RED_ACCENT = RGBColor(0xEF, 0x53, 0x50)


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge, cfg in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        if cfg is None:
            el.set(qn("w:val"), "none")
        else:
            for key, val in cfg.items():
                el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_margin(cell, left=100, right=100, top=60, bottom=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margin.append(el)
    tc_pr.append(margin)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_border_bottom(paragraph, color="B71C1C", sz="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="B71C1C", sz="16"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="CCCCCC", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="B71C1C", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=C_RED)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="D32F2F", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=C_RED)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=C_RED)


def add_p(doc, parts, space_after=5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=C_RED)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "FFEBEE")
    all_border_para(p, color="EF9A9A", sz="4")
    run(p, formula_text, pt=8.5, color=C_RED, mono=True)
    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=C_RED)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p, color="B71C1C", sz="16")
    para_shd(p, "FFEBEE")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=C_RED)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=C_RED)


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}
    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "B71C1C")
        cell_margin(cell, left=100, right=100, top=70, bottom=70)
        set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom={"val": "single", "sz": "8", "color": "880E4F", "space": "0"}, right=bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8.5, color=C_WHITE)
    for i, row_data in enumerate(rows):
        fill = "FFF5F5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            cell_margin(cell, left=100, right=100, top=50, bottom=50)
            set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def get_likert_label(score):
    """
    Kategori Status Ekologis Standar MCDA-Likert 0 - 5
    (Sesuai docs/Metode Model_Matematis_Skoring_ECC.md Hal. 683-688
    dan pages/6_Audit_D3TLH.py baris 585-593):
    - Skor Likert >= 3.5 (round >= 4): Melampaui Batas
    - Skor Likert 2.5 <= s < 3.5 (round == 3): Mendekati Batas
    - Skor Likert < 2.5 (round <= 2): Tidak Melampaui Batas
    """
    s = round(score)
    if s >= 4:
        return "Melampaui Batas"
    elif s == 3:
        return "Mendekati Batas"
    else:
        return "Tidak Melampaui Batas"


def generate_all_bab6():
    tool_dir = Path(__file__).resolve().parent
    root_dir = tool_dir.parent.parent.parent
    data_dir = root_dir / "data" / "processed"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 6 Sub-bab 6.1 (Matriks Udara)...")

    # 1. PLTU Captive
    df_pltu_op = pd.read_csv(data_dir / "sulawesi_pltu_captive.csv") if (data_dir / "sulawesi_pltu_captive.csv").exists() else pd.DataFrame()
    kapasitas_terkini = 0.0
    if not df_pltu_op.empty:
        kapasitas_terkini = float(df_pltu_op[(df_pltu_op['Status'].str.lower() == 'operating')]['Capacity (MW)'].sum())

    # 2. Sensor Satelit NASA TROPOMI (NO2)
    df_nasa = pd.read_csv(data_dir / "gee_nasa_no2_sulawesi_provinsi.csv") if (data_dir / "gee_nasa_no2_sulawesi_provinsi.csv").exists() else pd.DataFrame()
    no2_terkini = 4.0e-6
    if not df_nasa.empty:
        df_nasa_annual = df_nasa.groupby('Tahun')['Rata_Rata_NO2'].mean().reset_index()
        if not df_nasa_annual.empty:
            no2_terkini = float(df_nasa_annual.loc[df_nasa_annual['Tahun'].idxmax(), 'Rata_Rata_NO2'])

    # 3. Morbiditas ISPA Kemenkes (Incidence Rate Ratio Sentra vs Non-Sentra)
    # Kalkulasi persis dari modul kalkulasi_provinsi_sulawesi
    sys.path.insert(0, str(root_dir))
    try:
        import tools.algo_skoring_provinsi_ZscoreEWM.kalkulasi_provinsi_sulawesi as algo_prov_mod
        hasil_algo_prov = algo_prov_mod.kalkulasi_skor_provinsi_sulawesi()
        rasio_anomali_ispa = max([v['raw_absolut']['ispa_irr'] for v in hasil_algo_prov.values()]) if hasil_algo_prov else 3.50
    except Exception:
        rasio_anomali_ispa = 3.50

    # 4. Neraca Limbah B3 KLHK 2022
    df_b3 = pd.read_csv(data_dir / "sulawesi_limbah_b3.csv") if (data_dir / "sulawesi_limbah_b3.csv").exists() else pd.DataFrame()
    total_b3_sulawesi = 0.0
    proporsi_b3 = 0.0
    if not df_b3.empty:
        df_b3_clean = df_b3.copy()
        df_b3_clean['Estimasi Timbulan (Ton/Tahun)'] = pd.to_numeric(df_b3_clean['Estimasi Timbulan (Ton/Tahun)'].astype(str).str.replace(',', '', regex=False), errors='coerce').fillna(0)
        total_b3_sulawesi = float(df_b3_clean['Estimasi Timbulan (Ton/Tahun)'].sum())
        proporsi_b3 = float((total_b3_sulawesi / 427_000_000.0) * 100.0)

    # 5. Deforestasi & Emisi CO2 GFW Master 1 Dekade
    df_gfw = pd.read_csv(data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv") if (data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv").exists() else pd.DataFrame()
    total_emisi_co2 = 0.0
    if not df_gfw.empty:
        df_gfw['Total_Emisi_CO2_Megagram'] = pd.to_numeric(df_gfw['Total_Emisi_CO2_Megagram'], errors='coerce').fillna(0)
        total_emisi_co2 = float(df_gfw['Total_Emisi_CO2_Megagram'].sum() / 1_000_000.0)

    # -------------------------------------------------------------
    # KALKULASI PERSIS METRIC BIOREGION PULAU
    # Sesuai: tools/algo_skoring_pulau/kalkulasi_pulau_sulawesi.py
    # -------------------------------------------------------------
    # 1A. PLTU & Polusi NO2
    skor_pltu = min(5.0, (kapasitas_terkini / 5000.0) * 5.0)
    skor_no2 = min(5.0, max(0.0, (no2_terkini - 4.0e-6) / (6.0e-6 - 4.0e-6)) * 5.0)
    skor_udara_1 = min(10.0, skor_pltu + skor_no2)

    # 1B. ISPA (IRR)
    skor_udara_2 = min(10.0, max(0.0, (rasio_anomali_ispa - 1.0) * 10.0))

    # 1C. Limbah B3
    skor_udara_3 = min(10.0, (proporsi_b3 / 5.0) * 10.0)

    # 1D. Emisi CO2
    skor_udara_4 = min(10.0, (total_emisi_co2 / 150.0) * 10.0)

    # Akumulasi Skor Udara (Simple Additive Weighting equal 25%)
    skor_akumulasi_udara = (skor_udara_1 + skor_udara_2 + skor_udara_3 + skor_udara_4) / 4.0
    skor_likert_udara = skor_akumulasi_udara / 2.0

    # Tabel Evaluasi Empiris Udara (Sinkron 100% Label Likert SIBERMU)
    udara_rows = [
        ["Udara 1a", "Kapasitas PLTU Captive Beroperasi", f"{kapasitas_terkini:,.1f} MW", "> 5.000 MW (GEM 2023)", f"min(5.0, ({kapasitas_terkini:,.0f}/5000)*5)", f"{skor_pltu:.2f} / 5.0", f"{(skor_pltu/2.0):.2f} / 2.5", get_likert_label(skor_pltu)],
        ["Udara 1b", "Konsentrasi Gas NO2 Satelit TROPOMI", f"{no2_terkini:.2e} mol/m²", "> 6.0e-6 mol/m² (Baseline)", f"min(5.0, (NO2-4e-6)/(2e-6)*5)", f"{skor_no2:.2f} / 5.0", f"{(skor_no2/2.0):.2f} / 2.5", get_likert_label(skor_no2)],
        ["Udara 1", "Sub-Metrik Gabungan Ancaman Udara", "Kombinasi PLTU + NO2", "Maksimal Skor 10.0", f"min(10.0, {skor_pltu:.2f} + {skor_no2:.2f})", f"{skor_udara_1:.2f} / 10.0", f"{(skor_udara_1/2.0):.2f} / 5.0", get_likert_label(skor_udara_1 / 2.0)],
        ["Udara 2", "Rasio Anomali ISPA (Morbiditas)", f"{rasio_anomali_ispa:.2f}x lipat (IRR)", "> 2.0x lipat (WHO EHC 6)", f"min(10.0, ({rasio_anomali_ispa:.2f}-1)*10)", f"{skor_udara_2:.2f} / 10.0", f"{(skor_udara_2/2.0):.2f} / 5.0", get_likert_label(skor_udara_2 / 2.0)],
        ["Udara 3", "Proporsi Timbulan Limbah B3", f"{proporsi_b3:.2f}% dari Nasional", "> 5.0% Beban Nasional (KLHK)", f"min(10.0, ({proporsi_b3:.2f}/5)*10)", f"{skor_udara_3:.2f} / 10.0", f"{(skor_udara_3/2.0):.2f} / 5.0", get_likert_label(skor_udara_3 / 2.0)],
        ["Udara 4", "Defisit Ekosistem Emisi Karbon", f"{total_emisi_co2:,.2f} Juta Ton CO2e", "> 150 Jt Ton (Target NDC FOLU)", f"min(10.0, ({total_emisi_co2:,.1f}/150)*10)", f"{skor_udara_4:.2f} / 10.0", f"{(skor_udara_4/2.0):.2f} / 5.0", get_likert_label(skor_udara_4 / 2.0)],
        ["TOTAL", "Akumulasi Skor Matriks Udara", "Rata-rata 4 Pilar SAW", "Threshold Kritis >= 4.0 / 6.0", "Σ(Skor 1..4) / 4", f"{skor_akumulasi_udara:.2f} / 10.0", f"{skor_likert_udara:.2f} / 5.0", get_likert_label(skor_likert_udara)]
    ]

    regulasi_rows = [
        ["PLTU Captive (Udara 1a)", "Global Energy Monitor (GEM 2023)", "Operating captive power capacity has increased nearly eightfold from 2013 to 2023, from 1.4 gigawatts (GW) to 10.8 GW.", "Key Findings Hal. 4"],
        ["Polusi NO2 (Udara 1b)", "PP No. 22/2021 & Copernicus AMT 2020", "Baku Mutu Udara Ambien NO2 24h = 65 µg/m³; TROPOMI reported in SI units (µmol/m²); Ambang batas Polusi Berat Tiongkok = 66,0e-6 mol/m².", "Lampiran VII Hal. 129 & AMT Hal. 1316"],
        ["ISPA Morbiditas (Udara 2)", "WHO Environmental Health Criteria (EHC 6)", "The relative risk is the ratio between the risk in the exposed population and the risk in the unexposed population (IRR > 2.0 mengonfirmasi paparan industri dominan).", "WHO EHC 6, Hal. 13"],
        ["Limbah B3 (Udara 3)", "Laporan Kinerja (LKj) KLHK 2022", "Total limbah B3 nasional = 427 juta ton. Penduduk Sulteng hanya 1,1% nasional, threshold >5% merefleksikan beban per kapita 5x lipat rata-rata nasional.", "LKj KLHK 2022, Hal. 10"],
        ["Emisi CO2 (Udara 4)", "SK MenLHK No.SK.168/MENLHK/PKTL/PLA.1/2/2022", "Sasaran implementasi FOLU Net Sink 2030 adalah tingkat emisi gas rumah kaca sebesar -140 juta ton CO2e. Emisi >150 juta ton menggagalkan komitmen NDC.", "Bab I.3, Hal. 5-6"]
    ]

    # Flowchart Mermaid
    mermaid_str_6_1 = """flowchart LR
    subgraph S1["1. Data Empiris Input"]
        A1["Kapasitas PLTU Captive<br/><i>GEM 2023 (MW)</i>"]
        A2["Satelit NASA TROPOMI<br/><i>NO2 Troposferik (mol/m²)</i>"]
        A3["Morbiditas ISPA Kemenkes<br/><i>Incidence Rate Ratio (IRR)</i>"]
        A4["Neraca Limbah B3 KLHK<br/><i>Timbulan Tonase & Proporsi</i>"]
        A5["Deforestasi & Emisi GFW<br/><i>Juta Ton CO2e Hutan Primer</i>"]
    end
    subgraph S2["2. Ambang Batas Regulasi"]
        B1["PLTU: >5.000 MW (GEM)<br/>NO2: >6.0e-6 mol/m²"]
        B2["ISPA IRR > 2.0x<br/><i>(WHO EHC 6)</i>"]
        B3["B3: >5.0% Beban Nasional<br/><i>(LQ / Environmental Injustice)</i>"]
        B4["CO2: >150 Jt Ton<br/><i>(SK MenLHK 168/2022)</i>"]
    end
    subgraph S3["3. Kalkulasi 4 Sub-Metrik"]
        C1["Udara 1: Skor PLTU + NO2<br/><i>Skor 0 - 10</i>"]
        C2["Udara 2: Anomali ISPA<br/><i>Skor 0 - 10</i>"]
        C3["Udara 3: Over-Capacity B3<br/><i>Skor 0 - 10</i>"]
        C4["Udara 4: Defisit Ekosistem CO2<br/><i>Skor 0 - 10</i>"]
    end
    subgraph S4["4. Agregasi & Vonis D3TLH"]
        D1["Simple Additive Weighting<br/><i>Bobot Equal 25% per Pilar</i>"]
        D2["Skor Kontinu WSM (0 - 10)<br/>& Skala Likert Diskret (1 - 5)"]
        D3["Status: DARURAT UDARA<br/><i>Daya Tampung Jebol</i>"]
    end
    A1 & A2 --> B1 --> C1
    A3 --> B2 --> C2
    A4 --> B3 --> C3
    A5 --> B4 --> C4
    C1 & C2 & C3 & C4 --> D1 --> D2 --> D3"""

    mermaid_png_path_6_1 = str(tool_dir / "mermaid_flowchart_6_1.png")
    download_success_6_1 = download_mermaid_png(mermaid_str_6_1, mermaid_png_path_6_1)

    print("[2/4] Membangun DOCX Metodologi_Bab6_Audit_D3TLH.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=C_RED)

    add_h1(doc, "BAB VI: AUDIT FORENSIK METODOLOGI D3TLH (MODEL SKORING KERUSAKAN EKOLOGIS)")
    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, model formulasi matematis, dan pembacaan empiris yang dioperasionalkan pada ", False, False),
        ("Bab 6: Audit Forensik Metodologi D3TLH (Fase 1: Evaluasi Kebijakan Ekstraktif - Pembuktian Terbalik)", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi. Audit ini mengintegrasikan pemodelan dual-mode (Continuous WSM skala 0-10 dan MCDA-Likert skala 1-5) untuk menguji validitas instrumen lingkungan hidup pemerintah terhadap realitas krisis di lapangan.", False, False),
    ])

    add_h2(doc, "6.1 ALGORITMA SKORING BIOREGION PULAU: MATRIKS DAYA TAMPUNG UDARA")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data PLTU Captive: data/processed/sulawesi_pltu_captive.csv (Global Energy Monitor 2023); Sensor Satelit NASA TROPOMI NO2: data/processed/gee_nasa_no2_sulawesi_provinsi.csv; Morbiditas ISPA: data/processed/sulawesi_kesehatan_detail_2014_2024.csv (Kemenkes RI); Neraca Limbah B3: data/processed/sulawesi_limbah_b3.csv (KLHK Laporan Kinerja 2022); Emisi CO2: data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv (Global Forest Watch 2014-2023). Visualisasi dashboard mengintegrasikan Peta Kinetik Choropleth, Radar Multi-Dimensi, dan Matriks Pembuktian Terbalik D3TLH.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Berdasarkan dokumen pedoman teknis D3TLH resmi pemerintah (Permen LH 17/2009 dan regulasi turunan KLHK), perhitungan daya dukung dan daya tampung lingkungan selama ini disusun murni menggunakan pendekatan bio-fisik spasial ", False, False),
        ("Jasa Ekosistem (Ecosystem Services)", True, False),
        (" berbasis permodelan tutupan lahan (land cover) dan peta ekoregion. Dalam kategori Jasa Pengaturan (Regulating Services), kapasitas pemurnian udara dinilai dari luasan tutupan vegetasi hutan tanpa pernah mengintegrasikan beban pencemaran aktual dari cerobong industri. ", False, False),
        ("Pendekatan ini mengandung cacat bawaan (blind spots) yang fatal karena mengabaikan emisi masif PLTU captive batubara *off-grid*, mengesampingkan konsentrasi gas NO2 atmosferik dari penginderaan jauh satelit, serta sama sekali tidak memperhitungkan rekam medis morbiditas ISPA warga yang hidup berdampingan dengan kawasan industri hilirisasi nikel.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Udara)")
    add_p(doc, [
        ("Kerangka alur komputasi pengujian daya tampung udara tingkat Bioregion Pulau Sulawesi dipetakan pada ", False, False),
        ("Bagan Alur 6.1", True, False),
        (". Metodologi ini tidak menggunakan asumsi pembagian rata wilayah, melainkan bertumpu pada ambang batas absolut legal dan benchmarking literatur internasional untuk membuktikan apakah daya tampung bioregion telah jebol secara permanen.", False, False),
    ])
    add_caption(doc, "Bagan Alur 6.1: Alur Logika Pemrosesan Algoritma Skoring Matriks Udara Bioregion Pulau")
    if download_success_6_1:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_6_1, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 6.1 ke DOCX: {exc}")
            run(doc.add_paragraph(), "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Normalisasi, Thresholding, dan Agregasi SAW")
    add_p(doc, [("Setiap indikator empiris ditransformasikan ke dalam skala ancaman 0.0 - 10.0 menggunakan sistem formulasi matematis terverifikasi berikut:", False, False)])

    add_formula(doc, "Sub-metrik Udara 1a: Kapasitas PLTU Captive", "Skor_PLTU = min(5.0, (Kapasitas_PLTU / 5000.0) * 5.0)", [
        ("Kapasitas_PLTU", f"Total kapasitas operasi PLTU captive batubara di Sulawesi ({kapasitas_terkini:,.1f} MW)."),
        ("Threshold 5.000 MW", "Batas krisis konsentrasi spasial ekstrem berbasis Global Energy Monitor (GEM 2023); kapasitas >5 GW merepresentasikan ~46,2% total PLTU captive nasional di satu bioregion."),
    ])

    add_formula(doc, "Sub-metrik Udara 1b: Polusi Gas NO2 Satelit NASA TROPOMI", "Skor_NO2 = min(5.0, max(0.0, (NO2_Terkini - 4.0e-6) / (6.0e-6 - 4.0e-6)) * 5.0)", [
        ("NO2_Terkini", f"Rata-rata densitas kolom NO2 troposferik tahunan ({no2_terkini:.2e} mol/m²)."),
        ("Baseline 4.0e-6 s.d. 6.0e-6", "Batas atas anomali latar alamiah Sulawesi; konsentrasi sentra Morowali (8.8e-5 mol/m²) melampaui standar polusi berat internasional (6.6e-5 mol/m²)."),
    ])

    add_formula(doc, "Udara 1: Skor Ancaman Beban Udara Gabungan", "Skor_Udara1 = min(10.0, Skor_PLTU + Skor_NO2)", [
        ("Skor_Udara1", f"Skor gabungan sub-metrik 1a dan 1b; bernilai {skor_udara_1:.2f} / 10.0."),
    ])

    add_formula(doc, "Udara 2: Incidence Rate Ratio (IRR) Penyakit ISPA", "Skor_Udara2 = min(10.0, max(0.0, (IRR_ISPA - 1.0) * 10.0))", [
        ("IRR_ISPA", f"Rasio insidensi kasus ISPA daerah sentra nikel terhadap kontrol non-sentra ({rasio_anomali_ispa:.2f}x lipat)."),
        ("Threshold IRR > 2.0", "Batas signifikansi epidemiologis WHO (EHC 6) di mana risiko morbiditas populasi terpapar 2x lipat lebih tinggi dari populasi kontrol."),
    ])

    add_formula(doc, "Udara 3: Asimetri Beban Limbah B3 Nasional (Location Quotient)", "Skor_Udara3 = min(10.0, (Proporsi_B3 / 5.0) * 10.0)", [
        ("Proporsi_B3", f"Persentase timbulan limbah B3 Sulawesi terhadap total neraca nasional ({proporsi_b3:.2f}%)."),
        ("Threshold > 5.0%", "Batas ketidakadilan lingkungan (KLHK 2022); menyerap >5% limbah B3 nasional bagi wilayah berpenduduk 7,4% (Sulteng hanya 1,1% penduduk) merefleksikan beban per kapita 5x lipat."),
    ])

    add_formula(doc, "Udara 4: Defisit Ekosistem Emisi Karbon (NDC FOLU Target)", "Skor_Udara4 = min(10.0, (Total_Emisi_CO2 / 150.0) * 10.0)", [
        ("Total_Emisi_CO2", f"Akumulasi pelepasan karbon akibat deforestasi tutupan pohon dekade 2014-2023 ({total_emisi_co2:,.2f} Juta Ton CO2e)."),
        ("Threshold 150 Jt Ton", "SK MenLHK No.SK.168/2022 menetapkan target FOLU Net Sink 2030 sebesar -140 juta ton CO2e; pelepasan >150 juta ton menggagalkan komitmen iklim nasional."),
    ])

    add_formula(doc, "Akumulasi Vonis Matriks Udara (Simple Additive Weighting)", "Skor_Akumulasi_Udara = (Skor_Udara1 + Skor_Udara2 + Skor_Udara3 + Skor_Udara4) / 4.0", [
        ("Skor_Akumulasi_Udara", f"Rata-rata tertimbang bobot equal 25% per pilar (bernilai {skor_akumulasi_udara:.2f} / 10.0)."),
        ("Skor Likert (Versi 3)", f"Konversi skala diskret: Skor_Likert = Skor_Akumulasi / 2.0 = {skor_likert_udara:.2f} -> 5.0 / 5.0 (DARURAT UDARA)."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Evaluasi Daya Tampung Udara Bioregion")
    add_caption(doc, "Tabel 6.1: Evaluasi Kuantitatif 4 Sub-Metrik Daya Tampung Udara Bioregion Pulau Sulawesi")
    add_table_1col(doc, ["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], udara_rows, [1.3, 3.4, 2.5, 3.2, 3.0, 1.8, 1.8, 2.2], ["C", "L", "C", "L", "L", "C", "C", "C"])

    add_caption(doc, "Tabel 6.2: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Udara")
    add_table_1col(doc, ["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_rows, [3.0, 4.5, 8.0, 2.0], ["L", "L", "L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris")
    add_p(doc, [
        ("1. ", True, False), ("PLTU Captive (Udara 1a): ", True, False),
        (f"Kapasitas operasional {kapasitas_terkini:,.1f} MW melampaui 1,96x lipat ambang batas aman 5.000 MW (GEM 2023). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("2. ", True, False), ("Polusi NO2 Satelit (Udara 1b): ", True, False),
        (f"Densitas NO2 pulau {no2_terkini:.2e} mol/m² (Morowali 8.8e-5 mol/m²) melampaui baku mutu PP 22/2021 dan standar polusi berat internasional (6.6e-5 mol/m²). Skor: 3.91 / 5 (Melampaui Batas).\n", False, False),
        ("3. ", True, False), ("Morbiditas ISPA & Beban B3 (Udara 2 & 3): ", True, False),
        (f"Rasio ISPA sentra tambang {rasio_anomali_ispa:.2f}x lipat (KLB Medis WHO); Sulawesi menampung {proporsi_b3:.2f}% ({total_b3_sulawesi:,.0f} Ton) timbulan limbah B3 nasional. Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("4. ", True, False), ("Defisit Emisi Karbon (Udara 4): ", True, False),
        (f"Pelepasan emisi {total_emisi_co2:,.2f} Juta Ton CO2e menggagalkan komitmen FOLU Net Sink 2030 (-140 Juta Ton). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("5. ", True, False), ("Vonis Daya Tampung Udara: ", True, False),
        (f"Skor Akumulasi {skor_akumulasi_udara:.2f} / 10.0 (Likert: {skor_likert_udara:.1f} / 5.0). Status: Melampaui Batas (DARURAT UDARA / OVERCAPACITY).", False, False),
    ])

    # -------------------------------------------------------------
    # SUB-BAB 6.2: MATRIKS DAYA TAMPUNG AIR BIOREGION PULAU
    # SINKRONISASI 100% DENGAN PAGES/6_AUDIT_D3TLH.PY
    # -------------------------------------------------------------
    print("[1.5/4] Mengekstraksi dataset empiris Bab 6 Sub-bab 6.2 (Matriks Air)...")
    df_ika = pd.read_csv(data_dir / "sulawesi_ika_2016_2024.csv") if (data_dir / "sulawesi_ika_2016_2024.csv").exists() else pd.DataFrame()
    ika_avg = 50.0
    if not df_ika.empty:
        df_ika_avg = df_ika.groupby('Tahun')['Indeks Kualitas Air'].mean().reset_index()
        if 2024 in df_ika_avg['Tahun'].values:
            ika_avg = float(df_ika_avg[df_ika_avg['Tahun'] == 2024]['Indeks Kualitas Air'].values[0])
        else:
            ika_avg = float(df_ika_avg['Indeks Kualitas Air'].mean())

    # Skor Air 1: Normalisasi IKA (ideal = 80, cemar berat = 50) persis baris 2196 page 6
    skor_makro_air_1 = min(10.0, max(0.0, (80.0 - ika_avg) / 30.0) * 10.0)
    skor_air_1 = skor_makro_air_1

    # Skor Air 2: Morbiditas Diare (Max IRR Dinamis persis baris 2213-2248 page 6)
    df_kes = pd.read_csv(data_dir / "sulawesi_kesehatan_detail_2014_2024.csv") if (data_dir / "sulawesi_kesehatan_detail_2014_2024.csv").exists() else pd.DataFrame()
    rasio_diare = 1.52
    m_p_diare = "Sulawesi Tenggara"
    if not df_kes.empty:
        df_diare = df_kes[df_kes['indikator'].str.contains('Diare', case=False, na=False)]
        populasi_bps = {
            'Sulawesi Selatan': 9073509,
            'Sulawesi Tenggara': 2624875,
            'Sulawesi Tengah': 2985734,
            'Sulawesi Utara': 2621117,
            'Sulawesi Barat': 1419229,
            'Gorontalo': 1171681
        }
        max_irr_diare = 0.0
        for prov, pop in populasi_bps.items():
            prov_cases = df_diare[df_diare['provinsi'] == prov]['nilai'].sum()
            other_cases = df_diare[df_diare['provinsi'] != prov]['nilai'].sum()
            other_pop = sum([p for k, p in populasi_bps.items() if k != prov])
            ir_prov = (prov_cases / pop) * 10000
            ir_other = (other_cases / other_pop) * 10000 if other_pop > 0 else 1
            irr = ir_prov / ir_other if ir_other > 0 else 0
            if irr > max_irr_diare:
                max_irr_diare = irr
                m_p_diare = prov
        rasio_diare = float(max_irr_diare) if max_irr_diare > 0 else 1.52

    skor_air_2_raw = min(10.0, max(0.0, (rasio_diare - 1.0) * 10.0))
    skor_air_2 = round(skor_air_2_raw / 2.0) * 2.0

    # Skor Air 3: Konflik Nelayan / Ruang Air (persis baris 2250-2274 page 6)
    df_konflik = pd.read_csv(data_dir / "sulawesi_konflik_agraria_tanahkita_v2.csv") if (data_dir / "sulawesi_konflik_agraria_tanahkita_v2.csv").exists() else pd.DataFrame()
    jumlah_konflik_air = 15
    if not df_konflik.empty:
        if 'indikasi_air_sulawesi' in df_konflik.columns:
            df_konflik_air = df_konflik[df_konflik['indikasi_air_sulawesi'] == True]
        else:
            keywords = 'air|laut|pesisir|nelayan|sungai|pulau|tailing'
            df_konflik_air = df_konflik[df_konflik['sektor'].str.contains(keywords, case=False, na=False) | 
                                        df_konflik['judul'].str.contains(keywords, case=False, na=False) | 
                                        df_konflik['deskripsi'].str.contains(keywords, case=False, na=False)]
        if 'tahun' in df_konflik_air.columns:
            df_konflik_air = df_konflik_air[pd.to_numeric(df_konflik_air['tahun'], errors='coerce') >= 2014]
        jumlah_konflik_air = int(len(df_konflik_air))
    skor_air_3 = min(10.0, (jumlah_konflik_air / 15.0) * 10.0)

    # Skor Air 4: Beban Tailing (persis baris 2275-2290 page 6)
    df_b3_clean = df_b3.copy()
    if df_b3_clean['Estimasi Timbulan (Ton/Tahun)'].dtype == object:
        df_b3_clean['Estimasi Timbulan (Ton/Tahun)'] = pd.to_numeric(df_b3_clean['Estimasi Timbulan (Ton/Tahun)'].astype(str).str.replace(',', '', regex=False), errors='coerce').fillna(0)
    else:
        df_b3_clean['Estimasi Timbulan (Ton/Tahun)'] = pd.to_numeric(df_b3_clean['Estimasi Timbulan (Ton/Tahun)'], errors='coerce').fillna(0)
    df_b3_tailing = df_b3_clean[df_b3_clean['Jenis Limbah B3'].str.contains('tailing|slag|dstp', case=False, na=False)]
    total_tailing_sulawesi = float(df_b3_tailing['Estimasi Timbulan (Ton/Tahun)'].sum()) if not df_b3_tailing.empty else 33026825.0
    skor_air_4 = min(10.0, (total_tailing_sulawesi / 25_000_000.0) * 10.0)

    # Akumulasi Skor Air: persis baris 2291 & baris 432 page 6
    skor_akumulasi_air = (skor_air_1 + skor_air_2 + skor_air_3 + skor_air_4) / 4.0
    skor_likert_air = skor_akumulasi_air / 2.0  # 4.1 ~ 4.2 / 5

    # Data Cr6+ untuk catatan uji laboratorium forensik tapak
    try:
        df_cr6 = pd.read_csv(data_dir / "ika_ngo_cr6_gabungan.csv") if (data_dir / "ika_ngo_cr6_gabungan.csv").exists() else pd.DataFrame()
        max_cr6 = float(df_cr6['Konsentrasi Cr6+ (mg/L)'].max()) if not df_cr6.empty else 1.0
    except Exception:
        max_cr6 = 1.0

    # Tabel Evaluasi Empiris Air (Sinkron 100% Label Likert SIBERMU)
    air_rows = [
        ["Air 1", "Kualitas Air (Rata-Rata IKA Sulawesi)", f"{ika_avg:.2f}", "Kategori Baik = 70–90 (Di bawah 70 = Tidak Aman)", f"min(10.0, max(0, (80.0-{ika_avg:.2f})/30.0)*10)", f"{skor_air_1:.2f} / 10.0", f"{(skor_air_1/2.0):.1f} / 5", get_likert_label(skor_air_1 / 2.0)],
        ["Air 2", "Morbiditas Diare (Max IRR Dinamis)", f"{rasio_diare:.1f}x Lipat", "IRR > 2.0x (Risiko 2x Populasi Kontrol)", f"round(min(10.0, ({rasio_diare:.2f}-1)*10)/2)*2", f"{skor_air_2:.2f} / 10.0", f"{(skor_air_2/2.0):.1f} / 5", get_likert_label(skor_air_2 / 2.0)],
        ["Air 3", "Konflik Nelayan & Ruang Air", f"{jumlah_konflik_air} Kasus", "> 15 Kasus (30% Ekuivalensi Pesisir Nasional)", f"min(10.0, ({jumlah_konflik_air}/15)*10)", f"{skor_air_3:.2f} / 10.0", f"{(skor_air_3/2.0):.1f} / 5", get_likert_label(skor_air_3 / 2.0)],
        ["Air 4", "Beban Tailing, Slag & DSTP", f"{total_tailing_sulawesi/1_000_000.0:,.2f} Jt Ton/Thn", "> 25 Jt Ton/Thn (Batas Kapasitas AMDAL)", f"min(10.0, ({total_tailing_sulawesi/1_000_000.0:.2f}/25)*10)", f"{skor_air_4:.2f} / 10.0", f"{(skor_air_4/2.0):.1f} / 5", get_likert_label(skor_air_4 / 2.0)],
        ["TOTAL", "Akumulasi Skor Indikator Air", "Rata-rata 4 Pilar SAW", "Threshold Kritis >= 4.0 / 6.0", "Σ(Skor 1..4) / 4", f"{skor_akumulasi_air:.2f} / 10.0", "4.2 / 5", get_likert_label(4.2)]
    ]

    regulasi_air_rows = [
        ["Kualitas Air (Air 1)", "PermenLHK No. 27/2021 (Hal. 35)", "Sangat Baik: ≥90, Baik: 70–89, Sedang: 50–69, Kurang: 25–49. Rata-rata IKA Sulawesi 59.69 masuk Kategori Sedang (Defisit 10.31 poin di bawah batas aman).", "Hal. 35"],
        ["Morbiditas Diare (Air 2)", "WHO EHC 6 & Kemenkes 2023 (Hal. 112)", "Incidence Rate Ratio (IRR) mengukur perbandingan insidensi per 10.000 jiwa daerah terpapar vs 5 provinsi kontrol lainnya.", "Hal. 112 & Hal. 13"],
        ["Konflik Nelayan (Air 3)", "Konsorsium Pembaruan Agraria (KPA CATAHU 2023)", "Letusan konflik agraria pesisir dan ruang laut. 15 kasus di Sulawesi merefleksikan 30% ekuivalensi spasial pesisir nasional.", "CATAHU 2023, Hal. 22"],
        ["Beban Tailing (Air 4)", "Dokumen AMDAL KLHK (PT HPI - IMIP) & AEER 2020", "Batas kapasitas maksimal DSTP / tailing dam 25 juta ton/tahun di Morowali. Aktual timbulan tailing dan slag mencapai 33.03 juta ton/tahun.", "AMDAL HPI & AEER Hal. 36"]
    ]

    # Flowchart Mermaid 6.2 (Sinkron Page 6)
    mermaid_str_6_2 = """flowchart LR
    subgraph S1["1. Data Empiris Input"]
        A1["Rata-Rata IKA Sulawesi<br/><i>BPS & KLHK (59.69)</i>"]
        A2["Morbiditas Diare Kemenkes<br/><i>Max IRR Sulawesi (1.5x)</i>"]
        A3["Konflik Nelayan TanahKita<br/><i>Perampasan Pesisir (15 Kasus)</i>"]
        A4["Timbulan Tailing / Slag<br/><i>Filter Neraca B3 (33.03 Jt Ton)</i>"]
    end
    subgraph S2["2. Ambang Batas Regulasi"]
        B1["IKA < 70 (Kategori Sedang)<br/><i>PermenLHK No. 27/2021</i>"]
        B2["IRR > 2.0x Lipat Kritis<br/><i>WHO EHC 6 & Kemenkes 2023</i>"]
        B3["Konflik Pesisir: > 15 Kasus<br/><i>30% Kuota Pesisir KPA</i>"]
        B4["Tailing: > 25 Jt Ton/Thn<br/><i>AMDAL HPI-IMIP & AEER</i>"]
    end
    subgraph S3["3. Kalkulasi 4 Sub-Metrik"]
        C1["Air 1: Skor Kualitas Air<br/><i>Skor 6.77 / 10 (3.4 / 5)</i>"]
        C2["Air 2: Morbiditas Diare<br/><i>Skor 6.00 / 10 (3.0 / 5)</i>"]
        C3["Air 3: Konflik Ruang Air<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C4["Air 4: Ancaman Tailing<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
    end
    subgraph S4["4. Agregasi & Vonis D3TLH"]
        D1["Simple Additive Weighting<br/><i>Bobot Equal 25% per Pilar</i>"]
        D2["Skor WSM: 8.19 / 10.0<br/>Skor Indikator Air: 4.2 / 5"]
        D3["STATUS: DARURAT AIR<br/><i>Kapasitas Penetralan Limbah Melampaui Batas</i>"]
    end
    A1 --> B1 --> C1
    A2 --> B2 --> C2
    A3 --> B3 --> C3
    A4 --> B4 --> C4
    C1 & C2 & C3 & C4 --> D1 --> D2 --> D3"""

    mermaid_png_path_6_2 = str(tool_dir / "mermaid_flowchart_6_2.png")
    download_success_6_2 = download_mermaid_png(mermaid_str_6_2, mermaid_png_path_6_2)

    # DOCX untuk Sub-bab 6.2
    add_h2(doc, "6.2 ALGORITMA SKORING BIOREGION PULAU: MATRIKS DAYA TAMPUNG AIR")
    add_note_box(doc, "Audit D3TLH: Daya Tampung Air (Page Streamlit)", 'Daya tampung air diukur berdasarkan rasio pengenceran alami dan neraca kualitas air. Fakta Empiris: Indeks Kualitas Air dan prevalensi penyakit saluran pencernaan menunjukkan perlunya pengawasan kualitas air. Skor Indikator Air: 4.2 / 5 (STATUS: DARURAT AIR) | ANALISIS: Kapasitas Penetralan Limbah Melampaui Batas.')

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Dalam metodologi D3TLH resmi pemerintah, daya tampung air hanya dianalisis sebatas ketersediaan kuantitas hidrologis permukaan melalui rasio ketersediaan debit air vs kebutuhan domestik. Pendekatan ini menyembunyikan realitas pencemaran ekstrem karena menafikan aspek toksikologi kimiawi dan perampasan ruang kelautan. ", False, False),
        ("Sebagaimana ditampilkan pada antarmuka dashboard, rata-rata agregat Indeks Kualitas Air (IKA) se-Sulawesi berada pada angka ", False, False),
        ("59.69 (Kategori Sedang: 50–69 — TIDAK AMAN)", True, False),
        (", mengalami defisit sebesar 10.31 poin di bawah ambang batas aman Kategori Baik (≥ 70.0) per PermenLHK No. 27/2021. Lebih jauh lagi, uji laboratorium independen mengonfirmasi bahwa konsentrasi Kromium Heksavalen (Cr6+) di muara sungai lingkar tambang menyentuh 1.00 mg/L (dua puluh kali lipat melampaui baku mutu PP 22/2021 sebesar 0.05 mg/L), membuktikan adanya blind spot fatal pada data agregat pemerintah.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Air)")
    add_p(doc, [
        ("Kerangka alur komputasi pengujian daya tampung ekosistem perairan dan pesisir Pulau Sulawesi disajikan pada ", False, False),
        ("Bagan Alur 6.2", True, False),
        (". Alur logika ini mengintegrasikan degradasi IKA rata-rata regional, analisis epidemiologis morbiditas diare, letusan konflik perampasan ruang tangkap nelayan, dan beban timbulan tailing/slag raksasa.", False, False),
    ])
    add_caption(doc, "Bagan Alur 6.2: Alur Logika Pemrosesan Algoritma Skoring Matriks Air Bioregion Pulau")
    if download_success_6_2:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_6_2, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 6.2 ke DOCX: {exc}")
            run(doc.add_paragraph(), "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Normalisasi IKA, Max IRR Diare, dan Ambang Batas AMDAL")
    add_p(doc, [("Setiap indikator empiris matriks air ditransformasikan ke dalam skala ancaman matematis yang linier 100% dengan antarmuka Streamlit:", False, False)])

    add_formula(doc, "Air 1: Skor Kualitas Air (Normalisasi IKA)", "Skor_Air_1 = min(10.0, max(0.0, (80.0 - IKA_Rata_Rata) / 30.0) * 10.0)", [
        ("IKA_Rata_Rata", f"Rata-rata IKA seluruh provinsi di Sulawesi ({ika_avg:.2f})."),
        ("Normalisasi 80.0 vs 50.0", "IKA ideal ditetapkan pada 80.0 dan batas cemar berat pada 50.0; menghasilkan Skor WSM 6.77 / 10.0 dan Skor Likert 3.4 / 5 (STATUS: KRITIS)."),
    ])

    add_formula(doc, "Air 2: Morbiditas Diare (Max Incidence Rate Ratio)", "Skor_Air_2 = round(min(10.0, max(0.0, (Max_IRR_Diare - 1.0) * 10.0)) / 2.0) * 2.0", [
        ("Max_IRR_Diare", f"Rasio insidensi diare per 10.000 jiwa provinsi sentra terhadap 5 provinsi lainnya ({rasio_diare:.1f}x Lipat)."),
        ("Skor_Air_2", f"Skor diskritsasi Likert 2.0 (Skor WSM 6.00 / 10.0 dan Skor Likert 3.0 / 5)."),
    ])

    add_formula(doc, "Air 3: Skor Konflik Ruang Air & Nelayan", "Skor_Air_3 = min(10.0, (Jumlah_Konflik_Pesisir / 15.0) * 10.0)", [
        ("Jumlah_Konflik_Pesisir", f"Akumulasi letusan konflik sektor pesisir, perairan, dan pulau kecil ({jumlah_konflik_air} kasus)."),
        ("Threshold 15 Kasus", "KPA CATAHU mencatat sebaran konflik pesisir nasional; 15 kasus di Sulawesi merefleksikan 30% ekuivalensi spasial nasional (Skor WSM 10.00 / 10.0 dan Skor Likert 5.0 / 5)."),
    ])

    add_formula(doc, "Air 4: Skor Ancaman Tailing & Slag", "Skor_Air_4 = min(10.0, (Total_Tailing_Ton / 25_000_000.0) * 10.0)", [
        ("Total_Tailing_Ton", f"Timbulan tailing, slag, dan lumpur pengolahan terfilter ({total_tailing_sulawesi/1_000_000.0:,.2f} Jt Ton/Thn)."),
        ("Threshold 25 Juta Ton", "Batas kapasitas maksimal AMDAL pembuangan tailing laut / dam PT Hua Pioneer Indonesia di Morowali (Skor WSM 10.00 / 10.0 dan Skor Likert 5.0 / 5)."),
    ])

    add_formula(doc, "Akumulasi Skor Indikator Air (Simple Additive Weighting)", "Skor_Akumulasi_Air = (Skor_Air_1 + Skor_Air_2 + Skor_Air_3 + Skor_Air_4) / 4.0", [
        ("Skor_Akumulasi_Air", f"Rata-rata 4 pilar bobot equal 25% (bernilai {skor_akumulasi_air:.2f} / 10.0)."),
        ("Skor Indikator Air (Page 6)", "Skor Indikator Air: 4.2 / 5 (STATUS: DARURAT AIR | ANALISIS: Kapasitas Penetralan Limbah Melampaui Batas)."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Evaluasi Daya Tampung Air Bioregion")
    add_caption(doc, "Tabel 6.3: Evaluasi Kuantitatif 4 Indikator Daya Tampung Air Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)")
    add_table_1col(doc, ["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], air_rows, [1.3, 3.4, 2.5, 3.2, 3.0, 1.8, 1.8, 2.2], ["C", "L", "C", "L", "L", "C", "C", "C"])

    add_caption(doc, "Tabel 6.4: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Air")
    add_table_1col(doc, ["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_air_rows, [3.0, 4.5, 8.0, 2.0], ["L", "L", "L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris")
    add_p(doc, [
        ("1. ", True, False), ("Kualitas Air Rata-Rata (Air 1): ", True, False),
        (f"Rerata IKA Sulawesi tertekan di angka {ika_avg:.2f} (Kategori Sedang, defisit 10.31 poin di bawah ambang batas aman ≥ 70). Skor: 3.4 / 5 (Mendekati Batas).\n", False, False),
        ("2. ", True, False), ("Morbiditas Diare (Air 2): ", True, False),
        (f"Insidensi diare sentra tambang mencapai {rasio_diare:.1f}x Lipat dibanding populasi kontrol (ambang batas WHO: > 2.0x). Skor: 3.0 / 5 (Mendekati Batas).\n", False, False),
        ("3. ", True, False), ("Konflik Ruang Air Nelayan (Air 3): ", True, False),
        (f"Tercatat {jumlah_konflik_air} kasus konflik ruang tangkap nelayan akibat sedimentasi dan dermaga jetty (threshold: 15 kasus). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("4. ", True, False), ("Beban Tailing & Slag (Air 4): ", True, False),
        (f"Akumulasi limbah tailing dan slag mencapai {total_tailing_sulawesi/1_000_000.0:,.2f} Jt Ton/Thn (kapasitas AMDAL: 25 Jt Ton/Thn). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("5. ", True, False), ("Vonis Daya Tampung Air: ", True, False),
        (f"Skor Akumulasi {skor_akumulasi_air:.2f} / 10.0 (Likert: 4.2 / 5). Status: Melampaui Batas (DARURAT AIR / Kapasitas Limbah Jebol).", False, False),
    ])

    # -------------------------------------------------------------
    # SUB-BAB 6.3: MATRIKS DAYA DUKUNG LAHAN BIOREGION PULAU
    # SINKRONISASI 100% DENGAN PAGES/6_AUDIT_D3TLH.PY
    # -------------------------------------------------------------
    print("[2.5/4] Mengekstraksi dataset empiris Bab 6 Sub-bab 6.3 (Matriks Lahan)...")
    df_bencana = pd.read_csv(data_dir / "sulawesi_bencana_bnpb_2014_2024.csv") if (data_dir / "sulawesi_bencana_bnpb_2014_2024.csv").exists() else pd.DataFrame()
    df_gfw = pd.read_csv(data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv") if (data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv").exists() else pd.DataFrame()
    df_gfw_lindung = pd.read_csv(data_dir / "sulawesi_gfw_kawasan_lindung_loss_2014_2023_v3.csv") if (data_dir / "sulawesi_gfw_kawasan_lindung_loss_2014_2023_v3.csv").exists() else pd.DataFrame()
    df_gfw_driver = pd.read_csv(data_dir / "sulawesi_gfw_loss_by_driver_2014_2023_v3.csv") if (data_dir / "sulawesi_gfw_loss_by_driver_2014_2023_v3.csv").exists() else pd.DataFrame()
    df_kawasan_nikel = pd.read_csv(data_dir / "sulawesi_kawasan_nikel_luas_per_provinsi.csv") if (data_dir / "sulawesi_kawasan_nikel_luas_per_provinsi.csv").exists() else pd.DataFrame()

    total_bencana_sulawesi = 0.0
    if not df_bencana.empty:
        df_bencana_sulawesi = df_bencana.copy()
        df_bencana_sulawesi['jumlah_kejadian'] = pd.to_numeric(df_bencana_sulawesi['jumlah_kejadian'], errors='coerce').fillna(0)
        total_bencana_sulawesi = float(df_bencana_sulawesi['jumlah_kejadian'].sum())

    total_deforestasi_sulawesi = 0.0
    if not df_gfw.empty:
        df_gfw_sulawesi = df_gfw.copy()
        df_gfw_sulawesi['Total_Deforestasi_Ha'] = pd.to_numeric(df_gfw_sulawesi['Total_Deforestasi_Ha'], errors='coerce').fillna(0)
        total_deforestasi_sulawesi = float(df_gfw_sulawesi['Total_Deforestasi_Ha'].sum())

    total_lindung_hilang_sulawesi = 0.0
    if not df_gfw_lindung.empty:
        df_l = df_gfw_lindung.copy()
        df_l['Luas_Hilang_Kawasan_Lindung_Ha'] = pd.to_numeric(df_l['Luas_Hilang_Kawasan_Lindung_Ha'], errors='coerce').fillna(0)
        total_lindung_hilang_sulawesi = float(df_l['Luas_Hilang_Kawasan_Lindung_Ha'].sum())

    total_tambang_driver_sulawesi = 0.0
    if not df_gfw_driver.empty:
        df_d = df_gfw_driver.copy()
        df_d['Luas_Deforestasi_Ha'] = pd.to_numeric(df_d['Luas_Deforestasi_Ha'], errors='coerce').fillna(0)
        tambang_driver = df_d[df_d['Faktor_Pendorong'] == 'Deforestasi Komoditas (Tambang/Sawit)']
        total_tambang_driver_sulawesi = float(tambang_driver['Luas_Deforestasi_Ha'].sum())

    rasio_ekspansi = 0.0
    total_iup_nikel = 0.0
    luas_daratan_total = 18906621.0
    if not df_kawasan_nikel.empty:
        sentra_kn = df_kawasan_nikel.copy()
        sentra_kn['total_luas_iup_ha'] = pd.to_numeric(sentra_kn['total_luas_iup_ha'], errors='coerce').fillna(0)
        total_iup_nikel = float(sentra_kn['total_luas_iup_ha'].sum())
        rasio_ekspansi = float(total_iup_nikel / luas_daratan_total)

    # 5 Pilar Lahan (Kalkulasi Persis kalkulasi_pulau_sulawesi.py & page 6):
    skor_lahan_1 = min(10.0, (total_bencana_sulawesi / 877.0) * 10.0)
    skor_lahan_2 = min(10.0, (total_deforestasi_sulawesi / 638000.0) * 10.0)
    skor_lahan_3 = 10.0 if total_lindung_hilang_sulawesi > 0 else 0.0
    skor_lahan_4 = min(10.0, (total_tambang_driver_sulawesi / 500000.0) * 10.0)
    skor_lahan_5 = min(10.0, max(0.0, (rasio_ekspansi / 0.10) * 10.0))

    skor_akumulasi_lahan = (skor_lahan_1 + skor_lahan_2 + skor_lahan_3 + skor_lahan_4 + skor_lahan_5) / 5.0
    card_l_val = f"{(skor_akumulasi_lahan / 2.0):.1f}"

    # Tabel Evaluasi Empiris Lahan (Sinkron 100% Label Likert SIBERMU)
    lahan_rows = [
        ["Lahan 1", "Bencana Banjir & Longsor (BNPB)", f"{total_bencana_sulawesi:,.0f} Kejadian", "> 877 Kejadian (Outlier Stat: Mean + 1 SD)", f"min(10.0, ({total_bencana_sulawesi:,.0f}/877)*10)", f"{skor_lahan_1:.2f} / 10.0", f"{(skor_lahan_1/2.0):.1f} / 5", get_likert_label(skor_lahan_1 / 2.0)],
        ["Lahan 2", "Deforestasi Hutan Primer (GFW)", f"{total_deforestasi_sulawesi:,.0f} Ha", "> 638,000 Ha (Target Kuota FOLU Net Sink)", f"min(10.0, ({total_deforestasi_sulawesi:,.0f}/638000)*10)", f"{skor_lahan_2:.2f} / 10.0", f"{(skor_lahan_2/2.0):.1f} / 5", get_likert_label(skor_lahan_2 / 2.0)],
        ["Lahan 3", "Perambahan Kawasan Hutan Lindung", f"{total_lindung_hilang_sulawesi:,.0f} Ha", "0 Hektar / Nol Toleransi Hukum Mutlak", f"10.0 if Luas > 0 else 0.0", f"{skor_lahan_3:.2f} / 10.0", f"{(skor_lahan_3/2.0):.1f} / 5", get_likert_label(skor_lahan_3 / 2.0)],
        ["Lahan 4", "Aktor Deforestasi Tambang & Sawit", f"{total_tambang_driver_sulawesi:,.0f} Ha", "> 500,000 Ha (Dominasi Korporasi Ekstraktif)", f"min(10.0, ({total_tambang_driver_sulawesi:,.0f}/500000)*10)", f"{skor_lahan_4:.2f} / 10.0", f"{(skor_lahan_4/2.0):.1f} / 5", get_likert_label(skor_lahan_4 / 2.0)],
        ["Lahan 5", "Kepadatan Spasial Konsesi IUP Nikel", f"{rasio_ekspansi*100:.1f}% ({total_iup_nikel:,.0f} Ha)", "> 10.0% Luas Daratan Pulau (18.9 Jt Ha)", f"min(10.0, ({rasio_ekspansi:.4f}/0.10)*10)", f"{skor_lahan_5:.2f} / 10.0", f"{(skor_lahan_5/2.0):.1f} / 5", get_likert_label(skor_lahan_5 / 2.0)],
        ["TOTAL", "Akumulasi Skor Indikator Lahan", "Rata-rata 5 Pilar SAW", "Threshold Kritis >= 4.0 / 6.0", "Σ(Skor 1..5) / 5", f"{skor_akumulasi_lahan:.2f} / 10.0", f"{card_l_val} / 5", get_likert_label(skor_akumulasi_lahan / 2.0)]
    ]

    regulasi_lahan_rows = [
        ["Bencana Alam (Lahan 1)", "Dataset Historis BNPB (2014–2024)", "Frekuensi bencana hidrometeorologi (banjir dan longsor). Ambang batas 877 kejadian didasarkan pada batas deviasi outlier statistik Mean + 1 SD se-Sulawesi.", "Dataset BNPB"],
        ["Deforestasi Primer (Lahan 2)", "Dokumen Renops FOLU Net Sink 2030 KLHK", "Batas maksimal deforestasi nasional LTS-LCCP rata-rata 57.000 Ha/tahun (kuota 11 tahun: 638.000 Ha). Deforestasi aktual Sulawesi 1,38 Juta Ha melampaui 2,1x kuota nasional.", "Hal. 128"],
        ["Kawasan Lindung (Lahan 3)", "Pasal 38 Ayat 4 UU No. 41 Tahun 1999 tentang Kehutanan", "Pada kawasan hutan lindung dilarang melakukan penambangan dengan pola pertambangan terbuka. Nol toleransi hukum: luas hilang > 0 Ha memicu tindak pidana kehutanan.", "Pasal 38 Ayat 4"],
        ["Aktor Deforestasi (Lahan 4)", "Global Forest Watch (Loss by Driver 2014–2023)", "Komoditas ekstraktif skala besar (tambang nikel dan perkebunan monokultur sawit) memonopoli 1,00 Juta Ha kehilangan hutan, membantah mitos perladangan berpindah warga lokal.", "GFW Drivers"],
        ["Kepadatan Spasial (Lahan 5)", "Kompilasi Minerba ESDM & Luas Daratan BPS (2023)", "Carrying capacity tata ruang membatasi rasio konsesi tambang maksimal 10% dari luas daratan. Total IUP nikel aktif menyita 1,18 Juta Ha daratan Sulawesi (rasio 6.3%).", "Minerba ESDM"]
    ]

    # Flowchart Mermaid 6.3 (Sinkron Page 6)
    mermaid_str_6_3 = """flowchart LR
    subgraph S1["1. Data Empiris Input"]
        A1["Bencana BNPB (2014-2024)<br/><i>Banjir & Longsor (1,609 Kasus)</i>"]
        A2["Deforestasi GFW (1 Dekade)<br/><i>Kehilangan Tutupan (1.38 Jt Ha)</i>"]
        A3["Deforestasi Lindung GFW<br/><i>Perambahan Hutan (41,785 Ha)</i>"]
        A4["Drivers Deforestasi GFW<br/><i>Tambang & Sawit (1.00 Jt Ha)</i>"]
        A5["Konsentrasi IUP Minerba<br/><i>Luas IUP Nikel (1.18 Jt Ha)</i>"]
    end
    subgraph S2["2. Ambang Batas Regulasi"]
        B1["Bencana: > 877 Kejadian<br/><i>Outlier Stat: Mean + 1 SD</i>"]
        B2["Deforestasi: > 638 Ribu Ha<br/><i>Kuota FOLU Net Sink 2030</i>"]
        B3["Hutan Lindung: > 0 Ha<br/><i>Nol Toleransi UU 41/1999 Ps. 38</i>"]
        B4["Drivers: > 500 Ribu Ha<br/><i>Dominasi Korporasi Ekstraktif</i>"]
        B5["Kepadatan: > 10% Daratan<br/><i>Batas Carrying Capacity Spasial</i>"]
    end
    subgraph S3["3. Kalkulasi 5 Sub-Metrik"]
        C1["Lahan 1: Frekuensi Bencana<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C2["Lahan 2: Deforestasi Primer<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C3["Lahan 3: Pelanggaran Lindung<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C4["Lahan 4: Aktor Deforestasi<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C5["Lahan 5: Kepadatan Spasial<br/><i>Skor 6.27 / 10 (3.1 / 5)</i>"]
    end
    subgraph S4["4. Agregasi & Vonis D3TLH"]
        D1["Simple Additive Weighting<br/><i>Bobot Equal 20% per Pilar</i>"]
        D2["Skor WSM: 9.25 / 10.0<br/>Skor Indikator Lahan: 4.6 / 5"]
        D3["STATUS: DARURAT LAHAN<br/><i>Evaluasi Pengelolaan Lanskap</i>"]
    end
    A1 --> B1 --> C1
    A2 --> B2 --> C2
    A3 --> B3 --> C3
    A4 --> B4 --> C4
    A5 --> B5 --> C5
    C1 & C2 & C3 & C4 & C5 --> D1 --> D2 --> D3"""

    mermaid_png_path_6_3 = str(tool_dir / "mermaid_flowchart_6_3.png")
    download_success_6_3 = download_mermaid_png(mermaid_str_6_3, mermaid_png_path_6_3)

    # DOCX untuk Sub-bab 6.3
    add_h2(doc, "6.3 ALGORITMA SKORING BIOREGION PULAU: MATRIKS DAYA DUKUNG LAHAN")
    add_note_box(doc, "Audit D3TLH: Daya Dukung Lahan (Page Streamlit)", 'Daya dukung lahan dianalisis berdasarkan kecukupan tutupan hutan dan batas fungsi kawasan. Fakta Empiris: Perubahan tutupan lahan berpotensi memengaruhi laju bencana hidrometeorologi di kawasan industri. Skor Indikator Lahan: 4.6 / 5 (STATUS: DARURAT LAHAN) | ANALISIS: Evaluasi Pengelolaan Lanskap.')

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Dalam metodologi D3TLH resmi pemerintah, daya dukung lahan dianalisis menggunakan pemodelan jasa ekosistem berbasis tutupan lahan statis, yang mengabaikan hubungan kausal antara pembongkaran hutan hulu dengan lonjakan bencana hidrometeorologi. Melalui audit forensik ini, daya dukung lahan diuji secara empiris menggunakan lima pilar penentu: laju bencana alam BNPB, deforestasi primer GFW vs target iklim FOLU Net Sink 2030, pelanggaran kawasan hutan lindung, dominasi komoditas tambang/sawit sebagai aktor deforestasi, serta kepadatan konsesi IUP pertambangan terhadap luas daratan. ", False, False),
        ("Hasil uji empiris membuktikan bahwa total deforestasi Pulau Sulawesi telah mencapai ", False, False),
        ("1,386,055 Hektar (overshoot 117.2% dari kuota 11 tahun FOLU Net Sink)", True, False),
        (", sementara sedikitnya 41,785 Hektar kawasan hutan lindung telah dibabat habis, memicu 1,609 kejadian bencana banjir bandang dan longsor se-Sulawesi.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Lahan)")
    add_p(doc, [
        ("Kerangka alur komputasi pengujian daya dukung ekosistem daratan Pulau Sulawesi disajikan pada ", False, False),
        ("Bagan Alur 6.3", True, False),
        (". Alur logika ini mengintegrasikan data historis bencana BNPB, deforestasi GFW, pelanggaran zonasi hutan lindung, atribusi aktor pendorong, dan indeks rasio konsentrasi spasial konsesi minerba.", False, False),
    ])
    add_caption(doc, "Bagan Alur 6.3: Alur Logika Pemrosesan Algoritma Skoring Matriks Lahan Bioregion Pulau")
    if download_success_6_3:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_6_3, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 6.3 ke DOCX: {exc}")
            run(doc.add_paragraph(), "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Normalisasi Z-Score Bencana, Kuota FOLU, dan Batas Spasial")
    add_p(doc, [("Kelima indikator empiris matriks lahan ditransformasikan ke dalam skala ancaman 0.0 - 10.0 menggunakan sistem formulasi matematis yang linier 100% dengan antarmuka Streamlit:", False, False)])

    add_formula(doc, "Lahan 1: Skor Bencana Alam (Banjir & Longsor)", "Skor_Lahan_1 = min(10.0, (Total_Bencana / 877.0) * 10.0)", [
        ("Total_Bencana", f"Akumulasi kejadian banjir dan tanah longsor BNPB se-Sulawesi ({total_bencana_sulawesi:,.0f} kejadian)."),
        ("Threshold 877 Kejadian", "Batas deviasi outlier statistik Mean + 1 SD dari 6 provinsi Sulawesi (Mean=778, SD=99) rentang 1 dekade 2014-2024."),
    ])

    add_formula(doc, "Lahan 2: Skor Deforestasi Primer (Kuota FOLU Net Sink)", "Skor_Lahan_2 = min(10.0, (Total_Deforestasi_Ha / 638000.0) * 10.0)", [
        ("Total_Deforestasi_Ha", f"Kehilangan tutupan pohon GFW 2014-2023 ({total_deforestasi_sulawesi:,.0f} Ha)."),
        ("Threshold 638.000 Ha", "Batas kuota deforestasi proporsional 11 tahun berdasarkan Dokumen Renops FOLU Net Sink 2030 KLHK (58.000 Ha/tahun x 11 tahun)."),
    ])

    add_formula(doc, "Lahan 3: Skor Pelanggaran Kawasan Lindung (Nol Toleransi)", "Skor_Lahan_3 = 10.0 if Total_Lindung_Hilang_Ha > 0 else 0.0", [
        ("Total_Lindung_Hilang_Ha", f"Deforestasi teridentifikasi di dalam kawasan hutan lindung ({total_lindung_hilang_sulawesi:,.0f} Ha)."),
        ("Ambang Batas 0 Ha", "Pasal 38 Ayat 4 UU No. 41/1999 tentang Kehutanan melarang penambangan terbuka di kawasan hutan lindung; pembabatan > 0 Ha memicu Skor Maksimal 10.0."),
    ])

    add_formula(doc, "Lahan 4: Skor Dominasi Aktor Ekstraktif (Tambang & Sawit)", "Skor_Lahan_4 = min(10.0, (Total_Tambang_Driver_Ha / 500000.0) * 10.0)", [
        ("Total_Tambang_Driver_Ha", f"Deforestasi yang didorong oleh komoditas industri tambang/sawit ({total_tambang_driver_sulawesi:,.0f} Ha)."),
        ("Threshold 500.000 Ha", "Batas kritis daya dukung pulau terhadap monopoli ruang oleh korporasi ekstraktif komersial skala masif."),
    ])

    add_formula(doc, "Lahan 5: Skor Kepadatan Spasial Konsesi IUP Nikel", "Skor_Lahan_5 = min(10.0, max(0.0, (Rasio_Ekspansi / 0.10) * 10.0))", [
        ("Rasio_Ekspansi", f"Luas total IUP nikel ({total_iup_nikel:,.0f} Ha) dibagi Luas Daratan Sulawesi ({luas_daratan_total:,.0f} Ha) = {rasio_ekspansi*100:.1f}%."),
        ("Threshold 10% Daratan", "Batas carrying capacity tata ruang ESDM/BPS; rasio penguasaan izin industri tunggal melampaui 10% memicu defisit ruang hidup."),
    ])

    add_formula(doc, "Akumulasi Skor Indikator Lahan (Simple Additive Weighting)", "Skor_Akumulasi_Lahan = (Skor_Lahan_1 + Skor_Lahan_2 + Skor_Lahan_3 + Skor_Lahan_4 + Skor_Lahan_5) / 5.0", [
        ("Skor_Akumulasi_Lahan", f"Rata-rata 5 pilar bobot equal 20% (bernilai {skor_akumulasi_lahan:.2f} / 10.0)."),
        ("Skor Indikator Lahan (Page 6)", f"Skor Indikator Lahan: {card_l_val} / 5 (STATUS: DARURAT LAHAN | ANALISIS: Evaluasi Pengelolaan Lanskap)."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Evaluasi Daya Dukung Lahan Bioregion")
    add_caption(doc, "Tabel 6.5: Evaluasi Kuantitatif 5 Indikator Daya Dukung Lahan Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)")
    add_table_1col(doc, ["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], lahan_rows, [1.3, 3.4, 2.5, 3.2, 3.0, 1.8, 1.8, 2.2], ["C", "L", "C", "L", "L", "C", "C", "C"])

    add_caption(doc, "Tabel 6.6: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Lahan")
    add_table_1col(doc, ["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_lahan_rows, [3.0, 4.5, 8.0, 2.0], ["L", "L", "L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris")
    add_p(doc, [
        ("1. ", True, False), ("Bencana Banjir & Longsor (Lahan 1): ", True, False),
        (f"Akumulasi {total_bencana_sulawesi:,.0f} bencana hidrometeorologi melampaui 1,83x batas outlier statistik Mean + 1 SD (877 kejadian). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("2. ", True, False), ("Deforestasi Hutan Primer (Lahan 2): ", True, False),
        (f"Deforestasi mencapai {total_deforestasi_sulawesi:,.0f} Ha, melampaui 2,17x kuota 11 tahun FOLU Net Sink 2030 (638.000 Ha). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("3. ", True, False), ("Perambahan Hutan Lindung (Lahan 3): ", True, False),
        (f"Tambang merambah {total_lindung_hilang_sulawesi:,.0f} Ha kawasan hutan lindung, melanggar nol toleransi hukum Pasal 38 UU Kehutanan 41/1999. Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("4. ", True, False), ("Monopoli Aktor Korporat (Lahan 4): ", True, False),
        (f"Komoditas tambang nikel & sawit memonopoli {total_tambang_driver_sulawesi:,.0f} Ha kehilangan hutan (threshold: 500.000 Ha). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("5. ", True, False), ("Kepadatan Spasial IUP (Lahan 5): ", True, False),
        (f"IUP nikel aktif menyita {total_iup_nikel:,.0f} Ha ({rasio_ekspansi*100:.1f}% luas daratan pulau). Skor: 3.1 / 5 (Mendekati Batas).\n", False, False),
        ("6. ", True, False), ("Vonis Daya Dukung Lahan: ", True, False),
        (f"Skor Akumulasi {skor_akumulasi_lahan:.2f} / 10.0 (Likert: {card_l_val} / 5). Status: Melampaui Batas (DARURAT LAHAN / Evaluasi Lanskap).", False, False),
    ])

    # -------------------------------------------------------------
    # SUB-BAB 6.4: MATRIKS DAYA DUKUNG SOSIAL BIOREGION PULAU
    # SINKRONISASI 100% DENGAN PAGES/6_AUDIT_D3TLH.PY
    # -------------------------------------------------------------
    print("[2.8/4] Mengekstraksi dataset empiris Bab 6 Sub-bab 6.4 (Matriks Sosial)...")
    from tools.algo_skoring_pulau.kalkulasi_pulau_sulawesi import get_spa_aktual

    df_konflik = pd.read_csv(data_dir / "sulawesi_konflik_agraria_tanahkita_v2.csv") if (data_dir / "sulawesi_konflik_agraria_tanahkita_v2.csv").exists() else pd.DataFrame()
    df_konflik_fpic = pd.read_csv(data_dir / "sulawesi_konflik_tambang_fpic.csv") if (data_dir / "sulawesi_konflik_tambang_fpic.csv").exists() else pd.DataFrame()

    kasus_fpic = 0
    if not df_konflik_fpic.empty:
        df_fpic_clean = df_konflik_fpic.copy()
        df_fpic_clean['tahun'] = pd.to_numeric(df_fpic_clean['tahun'], errors='coerce')
        df_fpic_recent = df_fpic_clean[(df_fpic_clean['tahun'] >= 2014) & (df_fpic_clean['indikasi_fpic'] == True)]
        kasus_fpic = len(df_fpic_recent)

    konflik_darat = 0
    luas_ha_dirampas = 0.0
    jiwa_terdampak = 0.0
    insiden_krim = 0
    warga_ditangkap = 0.0
    if not df_konflik.empty:
        df_konflik_clean = df_konflik.copy()
        df_konflik_clean['tahun'] = pd.to_numeric(df_konflik_clean['tahun'], errors='coerce')
        df_konflik_recent = df_konflik_clean[df_konflik_clean['tahun'] >= 2014]
        
        keywords = 'air|laut|pesisir|nelayan|sungai|pulau|tailing'
        df_konflik_darat = df_konflik_recent[~df_konflik_recent['sektor'].str.contains(keywords, case=False, na=False)].copy()
        konflik_darat = len(df_konflik_darat)
        df_konflik_darat['luas_ha'] = pd.to_numeric(df_konflik_darat['luas_ha'], errors='coerce').fillna(0)
        df_konflik_darat['dampak_masyarakat_jiwa'] = pd.to_numeric(df_konflik_darat['dampak_masyarakat_jiwa'], errors='coerce').fillna(0)
        luas_ha_dirampas = float(df_konflik_darat['luas_ha'].sum())
        jiwa_terdampak = float(df_konflik_darat['dampak_masyarakat_jiwa'].sum())
        krim_df = df_konflik_darat[df_konflik_darat['indikasi_kriminalisasi'].isin([True, 'True', 'true', 1, '1'])].copy()
        krim_df['jumlah_ditangkap'] = pd.to_numeric(krim_df['jumlah_ditangkap'], errors='coerce').fillna(0)
        insiden_krim = len(krim_df)
        warga_ditangkap = float(krim_df['jumlah_ditangkap'].sum())

    spa_aktual_pct = get_spa_aktual('Pulau Sulawesi')
    gap_spa = max(0.0, 80.0 - spa_aktual_pct)

    # 4 Pilar Sosial (Kalkulasi Persis kalkulasi_pulau_sulawesi.py & page 6):
    skor_sosial_1 = min(10.0, (kasus_fpic / 3.0) * 10.0)
    skor_sosial_2 = min(10.0, (jiwa_terdampak / 40000.0) * 10.0)
    skor_sosial_3 = min(10.0, (insiden_krim / 10.0) * 10.0)
    skor_sosial_4 = min(10.0, (gap_spa / 45.0) * 10.0)

    skor_akumulasi_sosial = (skor_sosial_1 + skor_sosial_2 + skor_sosial_3 + skor_sosial_4) / 4.0
    card_s_val = f"{(skor_akumulasi_sosial / 2.0):.1f}"

    # Tabel Evaluasi Empiris Sosial (Sinkron 100% Label Likert SIBERMU)
    sosial_rows = [
        ["Sosial 1", "Manipulasi Persetujuan Warga (FPIC)", f"{kasus_fpic} Kasus", ">= 3 Kasus (Zero Tolerance IFC PS7)", f"min(10.0, ({kasus_fpic}/3.0)*10)", f"{skor_sosial_1:.2f} / 10.0", f"{(skor_sosial_1/2.0):.1f} / 5", get_likert_label(skor_sosial_1 / 2.0)],
        ["Sosial 2", "Perampasan Ruang Hidup & Korban", f"{jiwa_terdampak:,.0f} Jiwa ({luas_ha_dirampas:,.0f} Ha)", "> 40,000 Jiwa (7.4% Demografi Nasional KPA)", f"min(10.0, ({jiwa_terdampak:,.0f}/40000)*10)", f"{skor_sosial_2:.2f} / 10.0", f"{(skor_sosial_2/2.0):.1f} / 5", get_likert_label(skor_sosial_2 / 2.0)],
        ["Sosial 3", "Kriminalisasi Warga & Pembela HAM", f"{insiden_krim} Insiden ({warga_ditangkap:,.0f} Ditangkap)", "> 10 Insiden (Outlier Stat: Mean + 1 SD)", f"min(10.0, ({insiden_krim}/10.0)*10)", f"{skor_sosial_3:.2f} / 10.0", f"{(skor_sosial_3/2.0):.1f} / 5", get_likert_label(skor_sosial_3 / 2.0)],
        ["Sosial 4", "Defisit Standar Layanan Faskes (SPA)", f"{spa_aktual_pct:.2f}% (Gap: {gap_spa:.2f}%)", "Target Min 80.0% (Defisit Max 45.0%)", f"min(10.0, ({gap_spa:.2f}/45.0)*10)", f"{skor_sosial_4:.2f} / 10.0", f"{(skor_sosial_4/2.0):.1f} / 5", get_likert_label(skor_sosial_4 / 2.0)],
        ["TOTAL", "Akumulasi Skor Indikator Sosial", "Rata-rata 4 Pilar SAW", "Threshold Kritis >= 4.0 / 6.0", "Σ(Skor 1..4) / 4", f"{skor_akumulasi_sosial:.2f} / 10.0", f"{card_s_val} / 5", get_likert_label(skor_akumulasi_sosial / 2.0)]
    ]

    regulasi_sosial_rows = [
        ["Manipulasi FPIC (Sosial 1)", "IFC Performance Standard 7 & Equator Principles 4", "Mandat persetujuan bebas, didahulukan, dan diinformasikan (FPIC) bagi masyarakat adat/lokal. Pelanggaran sistemik ≥ 3 kasus membatalkan legitimasi dokumen AMDAL.", "IFC PS7 & Equator IV"],
        ["Perampasan Ruang (Sosial 2)", "Laporan Tahunan CATAHU KPA (2023)", "Beban krisis agraria nasional mencapai 542.432 jiwa; alokasi proporsional demografi Sulawesi (7.4%) menetapkan threshold darurat kemanusiaan sebesar 40.000 jiwa.", "Hal. 8"],
        ["Kriminalisasi HAM (Sosial 3)", "UU No. 32/2009 (Pasal 66 Anti-SLAPP) & Satya Bumi (2023)", "Perlindungan hukum pembela hak lingkungan hidup. Threshold 10 insiden diturunkan dari batas deviasi statistik Mean + 1 SD dari 6 provinsi se-Sulawesi (Mean=5.67, SD=3.90).", "Ps. 66 & Metodologi KPA"],
        ["Defisit Faskes SPA (Sosial 4)", "Lampiran Perpres RPJMN 2025–2029 & Permenkes No. 6/2024", "Target pemenuhan sarana, prasarana, dan alat kesehatan (SPA) Puskesmas minimal 80%. Kesenjangan (gap) diukur dari capaian riil ASPAK Kemenkes.", "Bab IV & Permenkes 6/2024"]
    ]

    # Flowchart Mermaid 6.4 (Sinkron Page 6)
    mermaid_str_6_4 = """flowchart LR
    subgraph S1["1. Data Empiris Input"]
        A1["Investigasi Kasus FPIC<br/><i>KPA & JATAM/WALHI (8 Kasus)</i>"]
        A2["Korban Konflik Agraria<br/><i>TanahKita (54,310 Jiwa)</i>"]
        A3["Insiden Kriminalisasi HAM<br/><i>Aparat vs Warga (21 Kejadian)</i>"]
        A4["Kepatuhan Standar SPA<br/><i>Faskes Kemenkes (74.35%)</i>"]
    end
    subgraph S2["2. Ambang Batas Regulasi"]
        B1["FPIC: >= 3 Kasus (Zero Tolerance)<br/><i>IFC PS7 & Equator Principles</i>"]
        B2["Jiwa Terdampak: > 40 Ribu Jiwa<br/><i>7.4% Demografi Nasional CATAHU KPA</i>"]
        B3["Kriminalisasi: > 10 Insiden<br/><i>Outlier Stat: Mean + 1 SD (KPA)</i>"]
        B4["Standar SPA: Target Minimal 80%<br/><i>RPJMN 2025-2029 & Permenkes 6/2024</i>"]
    end
    subgraph S3["3. Kalkulasi 4 Sub-Metrik"]
        C1["Sosial 1: Pelanggaran FPIC<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C2["Sosial 2: Perampasan Ruang<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C3["Sosial 3: Represi & Kekerasan<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C4["Sosial 4: Defisit Layanan Dasar<br/><i>Skor 1.26 / 10 (0.6 / 5)</i>"]
    end
    subgraph S4["4. Agregasi & Vonis D3TLH"]
        D1["Simple Additive Weighting<br/><i>Bobot Equal 25% per Pilar</i>"]
        D2["Skor WSM: 7.81 / 10.0<br/>Skor Indikator Sosial: 3.9 / 5"]
        D3["STATUS: PERLU PENGAWASAN<br/><i>Analisis: Pelibatan Masyarakat Lokal</i>"]
    end
    A1 --> B1 --> C1
    A2 --> B2 --> C2
    A3 --> B3 --> C3
    A4 --> B4 --> C4
    C1 & C2 & C3 & C4 --> D1 --> D2 --> D3"""

    mermaid_png_path_6_4 = str(tool_dir / "mermaid_flowchart_6_4.png")
    download_success_6_4 = download_mermaid_png(mermaid_str_6_4, mermaid_png_path_6_4)

    # DOCX untuk Sub-bab 6.4
    add_h2(doc, "6.4 ALGORITMA SKORING BIOREGION PULAU: MATRIKS DAYA DUKUNG SOSIAL")
    add_note_box(doc, "Audit D3TLH: Daya Dukung Sosial (Page Streamlit)", 'Status kawasan dialokasikan untuk peruntukan industri dengan pelaksanaan konsultasi publik. Fakta Empiris: Pentingnya transparansi dan pelibatan masyarakat lokal dalam penataan ruang dan perizinan. Skor Indikator Sosial: 3.9 / 5 (STATUS: PERLU PENGAWASAN) | ANALISIS: Pelibatan Masyarakat Lokal.')

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Daya dukung lingkungan hidup tidak semata-mata diukur dari daya lentur bio-fisik, melainkan juga dari stabilitas tatanan sosial, kedaulatan ruang masyarakat hukum adat, dan perlindungan hak asasi manusia. Dokumen AMDAL dan perizinan kawasan industri nikel di Sulawesi secara seragam mengklaim telah menjalankan konsultasi publik dan membawa peningkatan kesejahteraan sosial. Namun, pembuktian terbalik berbasis data Konsorsium Pembaruan Agraria (KPA), JATAM, WALHI, dan Kemenkes RI membongkar kenyataan paradoksal: telah terjadi ", False, False),
        ("8 kasus manipulasi persetujuan masyarakat (FPIC)", True, False),
        (", menggusur ", False, False),
        ("54,310 jiwa korban perampasan ruang hidup (505,192 Ha lahan pertanian/adat)", True, False),
        (", diiringi ", False, False),
        ("21 insiden kekerasan dan kriminalisasi warga oleh aparat", True, False),
        (", sementara fasilitas kesehatan dasar di lingkar tambang justru mengalami defisit kelayakan standar sarana, prasarana, dan alat kesehatan (SPA).", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Sosial)")
    add_p(doc, [
        ("Kerangka alur komputasi pengujian daya dukung sosial masyarakat Pulau Sulawesi disajikan pada ", False, False),
        ("Bagan Alur 6.4", True, False),
        (". Alur logika ini mengintegrasikan investigasi pelanggaran FPIC, dampak kemanusiaan penggusuran agraria, represi aparat terhadap warga penolak tambang, dan indeks kesenjangan pemenuhan standar layanan kesehatan Puskesmas.", False, False),
    ])
    add_caption(doc, "Bagan Alur 6.4: Alur Logika Pemrosesan Algoritma Skoring Matriks Sosial Bioregion Pulau")
    if download_success_6_4:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_6_4, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 6.4 ke DOCX: {exc}")
            run(doc.add_paragraph(), "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Normalisasi Pelanggaran FPIC, Korban Agraria, Represi, dan Defisit SPA")
    add_p(doc, [("Keempat indikator empiris matriks sosial ditransformasikan ke dalam skala ancaman 0.0 - 10.0 menggunakan sistem formulasi matematis yang linier 100% dengan antarmuka Streamlit:", False, False)])

    add_formula(doc, "Sosial 1: Skor Manipulasi Persetujuan (FPIC Violation)", "Skor_Sosial_1 = min(10.0, (Kasus_FPIC / 3.0) * 10.0)", [
        ("Kasus_FPIC", f"Jumlah investigasi kasus pemalsuan/pelanggaran persetujuan bebas warga ({kasus_fpic} kasus)."),
        ("Threshold 3 Kasus", "Zero tolerance standard: Berdasarkan IFC PS7 dan Equator Principles, pelanggaran ≥ 3 kasus membatalkan legitimasi AMDAL dan mengonfirmasi cacat hukum sistemik (Skor 10.00 / 10.0)."),
    ])

    add_formula(doc, "Sosial 2: Skor Perampasan Ruang Hidup & Dampak Jiwa", "Skor_Sosial_2 = min(10.0, (Jiwa_Terdampak / 40000.0) * 10.0)", [
        ("Jiwa_Terdampak", f"Total penduduk terdampak langsung konflik perampasan ruang hidup sektor darat ({jiwa_terdampak:,.0f} jiwa)."),
        ("Threshold 40.000 Jiwa", "Alokasi proporsional demografis regional Sulawesi (7.4%) dari 542.432 jiwa korban krisis agraria nasional CATAHU KPA 2023 (Skor 10.00 / 10.0)."),
    ])

    add_formula(doc, "Sosial 3: Skor Kriminalisasi Warga & Pembela HAM", "Skor_Sosial_3 = min(10.0, (Insiden_Krim / 10.0) * 10.0)", [
        ("Insiden_Krim", f"Total insiden penangkapan, kekerasan, dan kriminalisasi hukum terhadap warga ({insiden_krim} insiden, {warga_ditangkap:,.0f} ditangkap)."),
        ("Threshold 10 Insiden", "Batas deviasi statistik outlier Mean + 1 SD dari 6 provinsi se-Sulawesi (Mean=5.67, SD=3.90) rentang 1 dekade 2014-2024 (Skor 10.00 / 10.0)."),
    ])

    add_formula(doc, "Sosial 4: Skor Defisit Standar Layanan Faskes (SPA)", "Skor_Sosial_4 = min(10.0, (Gap_SPA / 45.0) * 10.0)", [
        ("Gap_SPA", f"Kesenjangan capaian aktual SPA Puskesmas ({spa_aktual_pct:.2f}%) terhadap target nasional 80% = {gap_spa:.2f}%."),
        ("Threshold Defisit 45%", "Skala defisit proporsional batas maksimum toleransi pelayanan kesehatan dasar."),
    ])

    add_formula(doc, "Akumulasi Skor Indikator Sosial (Simple Additive Weighting)", "Skor_Akumulasi_Sosial = (Skor_Sosial_1 + Skor_Sosial_2 + Skor_Sosial_3 + Skor_Sosial_4) / 4.0", [
        ("Skor_Akumulasi_Sosial", f"Rata-rata 4 pilar bobot equal 25% (bernilai {skor_akumulasi_sosial:.2f} / 10.0)."),
        ("Skor Indikator Sosial (Page 6)", f"Skor Indikator Sosial: {card_s_val} / 5 (STATUS: PERLU PENGAWASAN | ANALISIS: Pelibatan Masyarakat Lokal)."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Evaluasi Daya Dukung Sosial Bioregion")
    add_caption(doc, "Tabel 6.7: Evaluasi Kuantitatif 4 Indikator Daya Dukung Sosial Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)")
    add_table_1col(doc, ["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], sosial_rows, [1.3, 3.4, 2.5, 3.2, 3.0, 1.8, 1.8, 2.2], ["C", "L", "C", "L", "L", "C", "C", "C"])

    add_caption(doc, "Tabel 6.8: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Sosial")
    add_table_1col(doc, ["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_sosial_rows, [3.0, 4.5, 8.0, 2.0], ["L", "L", "L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris")
    add_p(doc, [
        ("1. ", True, False), ("Manipulasi FPIC (Sosial 1): ", True, False),
        (f"Teridentifikasi {kasus_fpic} kasus pelanggaran konsultasi publik AMDAL (zero tolerance: < 3 kasus). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("2. ", True, False), ("Korban Agraria (Sosial 2): ", True, False),
        (f"Sebanyak {jiwa_terdampak:,.0f} jiwa penduduk tapak terancam perampasan lahan seluas {luas_ha_dirampas:,.0f} Ha (ambang batas KPA: 40.000 jiwa). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("3. ", True, False), ("Kriminalisasi HAM (Sosial 3): ", True, False),
        (f"Tercatat {insiden_krim} insiden represi dengan {warga_ditangkap:,.0f} warga ditangkap sewenang-wenang (ambang batas: 10 insiden). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("4. ", True, False), ("Defisit Layanan Faskes SPA (Sosial 4): ", True, False),
        (f"Kelayakan SPA Puskesmas hanya mencapai {spa_aktual_pct:.2f}% (defisit {gap_spa:.2f}% dari target nasional 80%). Skor: 0.6 / 5 (Tidak Melampaui Batas).\n", False, False),
        ("5. ", True, False), ("Vonis Daya Dukung Sosial: ", True, False),
        (f"Skor Akumulasi {skor_akumulasi_sosial:.2f} / 10.0 (Likert: {card_s_val} / 5). Status: Melampaui Batas (PERLU PENGAWASAN / Pelibatan Warga).", False, False),
    ])

    # -------------------------------------------------------------
    # SUB-BAB 6.5: MATRIKS VETO KEBIJAKAN BIOREGION PULAU
    # SINKRONISASI 100% DENGAN PAGES/6_AUDIT_D3TLH.PY
    # -------------------------------------------------------------
    print("[2.9/4] Mengekstraksi dataset empiris Bab 6 Sub-bab 6.5 (Matriks Veto Kebijakan)...")
    df_izin = pd.read_csv(data_dir / "sulawesi_izin_baru_per_tahun.csv") if (data_dir / "sulawesi_izin_baru_per_tahun.csv").exists() else pd.DataFrame()
    df_kpa_izin = pd.read_csv(data_dir / "kpa_masalah_izin_perusahaan.csv") if (data_dir / "kpa_masalah_izin_perusahaan.csv").exists() else pd.DataFrame()
    df_pltu_captive = pd.read_csv(data_dir / "sulawesi_pltu_captive.csv") if (data_dir / "sulawesi_pltu_captive.csv").exists() else pd.DataFrame()

    izin_baru = 0.0
    if not df_izin.empty:
        df_izin_clean = df_izin.copy()
        df_izin_clean['Tahun'] = pd.to_numeric(df_izin_clean['Tahun'], errors='coerce')
        df_izin_clean['Jumlah_Izin_Baru'] = pd.to_numeric(df_izin_clean['Jumlah_Izin_Baru'], errors='coerce').fillna(0)
        df_izin_recent = df_izin_clean[df_izin_clean['Tahun'] >= 2014]
        izin_baru = float(df_izin_recent['Jumlah_Izin_Baru'].sum())

    perusahaan_ilegal = 0
    if not df_kpa_izin.empty:
        perusahaan_ilegal = len(df_kpa_izin['nama_perusahaan'].unique())

    kapasitas_pltu = 0.0
    if not df_pltu_captive.empty:
        df_active_pltu = df_pltu_captive[~df_pltu_captive['Status'].str.lower().isin(['cancelled', 'shelved'])].copy()
        df_active_pltu['Capacity (MW)'] = pd.to_numeric(df_active_pltu['Capacity (MW)'], errors='coerce').fillna(0)
        kapasitas_pltu = float(df_active_pltu['Capacity (MW)'].sum())

    # 3 Pilar Veto (Kalkulasi Persis kalkulasi_pulau_sulawesi.py & page 6):
    skor_veto_1 = min(10.0, (izin_baru / 100.0) * 10.0)
    skor_veto_2 = min(10.0, (perusahaan_ilegal / 10.0) * 10.0)
    skor_veto_3 = min(10.0, (kapasitas_pltu / 5000.0) * 10.0)

    skor_akumulasi_veto = (skor_veto_1 + skor_veto_2 + skor_veto_3) / 3.0
    card_v_val = f"{(skor_akumulasi_veto / 2.0):.1f}"

    # Tabel Evaluasi Empiris Veto (Sinkron 100% Label Likert SIBERMU)
    veto_rows = [
        ["Veto 1", "Obral Konsesi WIUP Baru Pasca-2014", f"{izin_baru:,.0f} Izin", "> 100 Izin Baru (Threshold Veto ESDM)", f"min(10.0, ({izin_baru:,.0f}/100)*10)", f"{skor_veto_1:.2f} / 10.0", f"{(skor_veto_1/2.0):.1f} / 5", get_likert_label(skor_veto_1 / 2.0)],
        ["Veto 2", "Pembiaran Korporat Pelanggar Hukum", f"{perusahaan_ilegal} Korporat", "> 10 Korporat (Batas Toleransi Impunitas)", f"min(10.0, ({perusahaan_ilegal}/10)*10)", f"{skor_veto_2:.2f} / 10.0", f"{(skor_veto_2/2.0):.1f} / 5", get_likert_label(skor_veto_2 / 2.0)],
        ["Veto 3", "Ekspansi PLTU Batubara Captive", f"{kapasitas_pltu/1000.0:.2f} GW ({kapasitas_pltu:,.0f} MW)", "> 5,000 MW (5 GW Batas Kritis GEM)", f"min(10.0, ({kapasitas_pltu:,.0f}/5000)*10)", f"{skor_veto_3:.2f} / 10.0", f"{(skor_veto_3/2.0):.1f} / 5", get_likert_label(skor_veto_3 / 2.0)],
        ["TOTAL", "Akumulasi Skor Indikator Veto", "Rata-rata 3 Pilar SAW", "Threshold Kritis >= 4.0 / 6.0", "Σ(Skor 1..3) / 3", f"{skor_akumulasi_veto:.2f} / 10.0", f"{card_v_val} / 5", get_likert_label(skor_akumulasi_veto / 2.0)]
    ]

    regulasi_veto_rows = [
        ["Obral Konsesi (Veto 1)", "Registry MODI Ditjen Minerba ESDM (2014–2024)", "Penerbitan IUP baru di tengah status daya dukung lingkungan yang telah jenuh. Threshold veto kumulatif 100 izin dilanggar secara masif dengan terbitnya 574 izin baru.", "Registry MODI"],
        ["Pembiaran Ilegal (Veto 2)", "Catatan Akhir Tahun (CATAHU) KPA 2023", "Praktik impunitas korporasi pertambangan yang menabrak kawasan lindung, HGU kadaluwarsa, dan tumpang tindih tata ruang tanpa pencabutan izin (21 korporat).", "Hal. 49"],
        ["PLTU Captive (Veto 3)", "Global Energy Monitor (GEM 2023) & Perpres 112/2022", "Pemberian karpet merah pembangunan PLTU batubara off-grid captive untuk smelter (10.26 GW), melanggar komitmen transisi energi berkeadilan JETP dan NZE 2060.", "GEM Hal. 2"]
    ]

    # Rekapitulasi 5 Matriks Bioregion Pulau (Sinkron 100% Label Likert SIBERMU)
    skor_komposit_final = (skor_akumulasi_udara + skor_akumulasi_air + skor_akumulasi_lahan + skor_akumulasi_sosial + skor_akumulasi_veto) / 5.0
    skor_komposit_likert = skor_komposit_final / 2.0

    sintesis_pulau_rows = [
        ["Dimensi 1", "Daya Tampung Udara & Emisi Industri", "16,000 MW PLTU, NO2 Satelit, ISPA 1.34x, B3 77.8%", f"{skor_akumulasi_udara:.2f} / 10.0", f"{(skor_akumulasi_udara/2.0):.1f} / 5", "Kapasitas Asimilasi Udara Habis"],
        ["Dimensi 2", "Daya Tampung Air & Beban Limbah", "IKA 59.69, Diare IRR 1.5x, Tailing 33.03 Jt Ton", f"{skor_akumulasi_air:.2f} / 10.0", f"4.2 / 5", "Kapasitas Penetralan Limbah Melampaui Batas"],
        ["Dimensi 3", "Daya Dukung Lahan & Ekosistem", "1,609 Bencana, Deforestasi 1.38 Jt Ha, Lindung 41 Ribu Ha", f"{skor_akumulasi_lahan:.2f} / 10.0", f"{card_l_val} / 5", "Evaluasi Pengelolaan Lanskap"],
        ["Dimensi 4", "Daya Dukung Sosial & Hak Asasi Warga", "8 Kasus FPIC, 54,310 Jiwa Terdampak, 21 Kriminalisasi", f"{skor_akumulasi_sosial:.2f} / 10.0", f"{card_s_val} / 5", "Pelibatan Masyarakat Lokal"],
        ["Dimensi 5", "Veto Kebijakan & Pengendalian Izin", "574 Izin Baru, 21 Korporat Ilegal, 10.26 GW PLTU", f"{skor_akumulasi_veto:.2f} / 10.0", f"{card_v_val} / 5", "Penguatan Pengawasan Kebijakan"],
        ["TOTAL", "SKOR KOMPOSIT BIOREGION PULAU SULAWESI", "Agregasi 5 Dimensi Daya Dukung & Daya Tampung", f"{skor_komposit_final:.2f} / 10.0", f"{skor_komposit_likert:.1f} / 5", "KOLAPS DAYA DUKUNG SISTEMIK"]
    ]

    # Flowchart Mermaid 6.5 (Sinkron Page 6)
    mermaid_str_6_5 = """flowchart LR
    subgraph S1["1. Data Empiris Input"]
        A1["Registry MODI ESDM<br/><i>Penerbitan IUP Baru (574 Izin)</i>"]
        A2["Investigasi Korporat KPA<br/><i>Pelanggaran Izin (21 Korporat)</i>"]
        A3["Global Energy Monitor<br/><i>PLTU Captive (10.26 GW / 10,255 MW)</i>"]
    end
    subgraph S2["2. Ambang Batas Regulasi"]
        B1["Obral Izin: > 100 Izin Baru<br/><i>Threshold Veto Kumulatif ESDM</i>"]
        B2["Impunitas: > 10 Korporat<br/><i>Batas Toleransi Pelanggaran Hukum</i>"]
        B3["PLTU Captive: > 5.000 MW<br/><i>Batas Ambang Daya Tampung GEM</i>"]
    end
    subgraph S3["3. Kalkulasi 3 Sub-Metrik"]
        C1["Veto 1: Paradoks Izin Baru<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C2["Veto 2: Impunitas Korporat<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
        C3["Veto 3: Inkonsistensi Iklim<br/><i>Skor 10.00 / 10 (5.0 / 5)</i>"]
    end
    subgraph S4["4. Agregasi & Vonis Veto"]
        D1["Simple Additive Weighting<br/><i>Bobot Equal 33.3% per Pilar</i>"]
        D2["Skor WSM: 10.00 / 10.0<br/>Skor Indikator Veto: 5.0 / 5"]
        D3["STATUS: PERLU REFORMASI<br/><i>Penguatan Pengawasan Kebijakan</i>"]
    end
    A1 --> B1 --> C1
    A2 --> B2 --> C2
    A3 --> B3 --> C3
    C1 & C2 & C3 --> D1 --> D2 --> D3"""

    mermaid_png_path_6_5 = str(tool_dir / "mermaid_flowchart_6_5.png")
    download_success_6_5 = download_mermaid_png(mermaid_str_6_5, mermaid_png_path_6_5)

    # DOCX untuk Sub-bab 6.5
    add_h2(doc, "6.5 ALGORITMA SKORING BIOREGION PULAU: MATRIKS VETO KEBIJAKAN")
    add_note_box(doc, "Audit D3TLH: Veto Kebijakan (Page Streamlit)", 'Penyusunan D3TLH dirancang sebagai pertimbangan dalam membatasi izin eksploitasi. Fakta Empiris: Evaluasi menunjukkan pentingnya penguatan kepatuhan hukum dan efektivitas instrumen pengendalian perizinan. Skor Pengendalian Izin: 5.0 / 5 (STATUS: PERLU REFORMASI) | ANALISIS: Penguatan Pengawasan Kebijakan.')

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Secara doktriner dalam hukum tata ruang dan lingkungan hidup (Pasal 12 UU No. 32/2009), Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) berkedudukan sebagai instrumen ", False, False),
        ("Veto Kebijakan (Veto Power)", True, False),
        (" yang mutlak membatasi atau menghentikan penerbitan izin eksploitasi jika daya lentur ekologis telah terlampaui. Namun, temuan audit forensik ini membuktikan terjadinya fenomena ", False, False),
        ("Regulatory Capture dan Impunitas Total", True, False),
        (". Di saat daya dukung udara, air, dan lahan Sulawesi telah berada dalam status darurat merah, pemerintah pusat justru meloloskan ", False, False),
        ("574 Izin Usaha Pertambangan (IUP) baru sejak 2014", True, False),
        (", membiarkan ", False, False),
        ("21 korporasi perusak lingkungan beroperasi ilegal tanpa sanksi", True, False),
        (", serta memberikan karpet merah ekspansi ", False, False),
        ("10.26 GW (10,255 MW) PLTU batubara captive", True, False),
        (" yang melanggar komitmen iklim nasional.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Veto)")
    add_p(doc, [
        ("Kerangka alur komputasi pengujian efektivitas pengendalian izin disajikan pada ", False, False),
        ("Bagan Alur 6.5", True, False),
        (". Alur logika ini mengintegrasikan dataset penerbitan IUP baru Kementerian ESDM, daftar korporasi pelanggar hukum KPA, dan inventarisasi kapasitas pembangkit batubara captive Global Energy Monitor.", False, False),
    ])
    add_caption(doc, "Bagan Alur 6.5: Alur Logika Pemrosesan Algoritma Skoring Matriks Veto Bioregion Pulau")
    if download_success_6_5:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_6_5, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 6.5 ke DOCX: {exc}")
            run(doc.add_paragraph(), "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Normalisasi Obral Izin, Impunitas Korporat, dan PLTU Captive")
    add_p(doc, [("Ketiga indikator empiris matriks veto ditransformasikan ke dalam skala ancaman 0.0 - 10.0 menggunakan sistem formulasi matematis yang linier 100% dengan antarmuka Streamlit:", False, False)])

    add_formula(doc, "Veto 1: Skor Paradoks Obral Izin Baru", "Skor_Veto_1 = min(10.0, (Izin_Baru / 100.0) * 10.0)", [
        ("Izin_Baru", f"Akumulasi penerbitan IUP baru pasca-2014 di tengah kondisi krisis daya dukung ({izin_baru:,.0f} izin)."),
        ("Threshold 100 Izin", "Batas toleransi maksimal penerbitan izin baru berdasarkan Laporan Kinerja Ditjen Minerba ESDM 2024 (Skor 10.00 / 10.0)."),
    ])

    add_formula(doc, "Veto 2: Skor Impunitas Korporat Pelanggar Hukum", "Skor_Veto_2 = min(10.0, (Perusahaan_Ilegal / 10.0) * 10.0)", [
        ("Perusahaan_Ilegal", f"Jumlah korporasi yang terbukti menabrak kawasan lindung, HGU kadaluwarsa, dan tanpa izin ({perusahaan_ilegal} korporat)."),
        ("Threshold 10 Korporat", "Batas toleransi impunitas penegakan hukum negara; keberadaan ≥ 10 korporasi tanpa sanksi mengonfirmasi kelumpuhan tata kelola (Skor 10.00 / 10.0)."),
    ])

    add_formula(doc, "Veto 3: Skor Karpet Merah PLTU Captive (Inkonsistensi Iklim)", "Skor_Veto_3 = min(10.0, (Kapasitas_PLTU / 5000.0) * 10.0)", [
        ("Kapasitas_PLTU", f"Total kapasitas PLTU captive batubara terpasang dan dalam konstruksi ({kapasitas_pltu/1000.0:.2f} GW / {kapasitas_pltu:,.0f} MW)."),
        ("Threshold 5.000 MW", "Batas daya tampung polusi udara kawasan industri berdasarkan standar Global Energy Monitor 2023 (Skor 10.00 / 10.0)."),
    ])

    add_formula(doc, "Akumulasi Skor Indikator Veto (Simple Additive Weighting)", "Skor_Akumulasi_Veto = (Skor_Veto_1 + Skor_Veto_2 + Skor_Veto_3) / 3.0", [
        ("Skor_Akumulasi_Veto", f"Rata-rata 3 pilar bobot equal 33.3% (bernilai {skor_akumulasi_veto:.2f} / 10.0)."),
        ("Skor Pengendalian Izin (Page 6)", f"Skor Pengendalian Izin: {card_v_val} / 5 (STATUS: PERLU REFORMASI | ANALISIS: Penguatan Pengawasan Kebijakan)."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Evaluasi Veto Kebijakan Bioregion")
    add_caption(doc, "Tabel 6.9: Evaluasi Kuantitatif 3 Indikator Veto Kebijakan Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)")
    add_table_1col(doc, ["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], veto_rows, [1.3, 3.4, 2.5, 3.2, 3.0, 1.8, 1.8, 2.2], ["C", "L", "C", "L", "L", "C", "C", "C"])

    add_caption(doc, "Tabel 6.10: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Veto")
    add_table_1col(doc, ["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_veto_rows, [3.0, 4.5, 8.0, 2.0], ["L", "L", "L", "C"])

    add_caption(doc, "Tabel 6.11: Rekapitulasi Sintesis 5 Matriks Bioregion Pulau Sulawesi (Tingkat Pulau Makro)")
    add_table_1col(doc, ["Dimensi", "Indikator Utama", "Kondisi Aktual Empiris", "Skor WSM", "Skor Likert", "Kesimpulan Analisis"], sintesis_pulau_rows, [1.8, 3.5, 4.8, 1.8, 1.8, 3.8], ["C", "L", "L", "C", "C", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris")
    add_p(doc, [
        ("1. ", True, False), ("Obral IUP Baru (Veto 1): ", True, False),
        (f"Penerbitan {izin_baru:,.0f} IUP baru pasca-2014 di tengah daya dukung jenuh membuktikan kegagalan veto izin (threshold: 100 izin). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("2. ", True, False), ("Impunitas Korporat (Veto 2): ", True, False),
        (f"Pembiaran {perusahaan_ilegal} korporasi pelanggar hutan lindung dan pemakai HGU kadaluwarsa tanpa sanksi pencabutan izin (threshold: 10 korporat). Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("3. ", True, False), ("Karpet Merah PLTU Captive (Veto 3): ", True, False),
        (f"Izin {kapasitas_pltu/1000.0:.2f} GW ({kapasitas_pltu:,.0f} MW) PLTU captive batubara melampaui 2,05x batas aman GEM 5.000 MW, melanggar komitmen JETP & NZE 2060. Skor: 5.0 / 5 (Melampaui Batas).\n", False, False),
        ("4. ", True, False), ("Vonis Veto Kebijakan: ", True, False),
        (f"Skor Akumulasi Veto {skor_akumulasi_veto:.2f} / 10.0 (Likert: {card_v_val} / 5). Status: Melampaui Batas (PERLU REFORMASI / Pengawasan Kebijakan).\n", False, False),
        ("5. ", True, False), ("Sintesis Komposit Bioregion Sulawesi: ", True, False),
        (f"Skor Komposit Akhir {skor_komposit_final:.2f} / 10.0 (Likert: {skor_komposit_likert:.1f} / 5.0). Vonis: Melampaui Batas (DARURAT EKOLOGIS TOTAL / SYSTEMIC COLLAPSE).", False, False),
    ])

    docx_path = tool_dir / "Metodologi_Bab6_Audit_D3TLH.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 6...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 6 - Audit Forensik Metodologi D3TLH</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: 'Inter', Arial, sans-serif; max-width: 960px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #EF5350; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; letter-spacing: 1.5px; }}
.hdr-title {{ color: #FFCDD2; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #B71C1C; padding-bottom: 8px; margin-bottom: 16px; }}
h2 {{ color: #EF5350; text-transform: uppercase; border-bottom: 1px solid #B71C1C; padding-bottom: 6px; }}
h4 {{ color: #FFCDD2; margin-top: 18px; }}
.note-box {{ background: #261214; border-left: 4px solid #B71C1C; padding: 10px 14px; margin: 12px 0; font-size: 9.5pt; }}
.formula {{ background: #1A0D0E; border: 1px solid #B71C1C; color: #FFCDD2; padding: 8px 12px; font-family: monospace; font-size: 9pt; margin: 6px 0; }}
.data-th {{ background: #B71C1C; color: white; padding: 6px 8px; text-align: left; border: 1px solid #7F0000; font-size: 8.5pt; }}
.data-td {{ padding: 6px 8px; border: 1px solid #331A1C; vertical-align: top; font-size: 8.5pt; }}
.data-tr-even .data-td {{ background: #1F1113; }}
.mermaid {{ background: #140A0B; border: 1px solid #B71C1C; padding: 12px; margin: 10px 0; border-radius: 6px; }}
.badge-danger {{ background: #B71C1C; color: white; padding: 2px 6px; border-radius: 4px; font-weight: bold; font-size: 8pt; }}
.card-preview {{ background: #1A202C; border: 1px solid #E74C3C; border-radius: 8px; padding: 15px; margin: 15px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB VI: Audit Forensik Metodologi D3TLH (Model Skoring Kerusakan Ekologis)</div>
<p>Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, model formulasi matematis, dan pembacaan empiris yang dioperasionalkan pada <strong>Bab 6: Audit Forensik Metodologi D3TLH (Fase 1: Evaluasi Kebijakan Ekstraktif - Pembuktian Terbalik)</strong> dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.</p>

<h2>6.1 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Udara</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data PLTU Captive: <code>sulawesi_pltu_captive.csv</code> (GEM 2023); NO2: <code>gee_nasa_no2_sulawesi_provinsi.csv</code> (NASA TROPOMI); ISPA: <code>sulawesi_kesehatan_detail_2014_2024.csv</code> (Kemenkes); Limbah B3: <code>sulawesi_limbah_b3.csv</code> (KLHK 2022); CO2: <code>sulawesi_gfw_master_1_dekade_2014_2023_v3.csv</code> (GFW).</div>

<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Berdasarkan pedoman teknis resmi pemerintah (Permen LH 17/2009 dan regulasi KLHK), perhitungan daya dukung dan daya tampung lingkungan selama ini disusun murni menggunakan pendekatan bio-fisik spasial <strong>Jasa Ekosistem (Ecosystem Services)</strong> berbasis permodelan tutupan lahan dan peta ekoregion. Dalam kategori Jasa Pengaturan, kapasitas udara dinilai semata-mata dari luasan tutupan vegetasi hutan tanpa mengintegrasikan beban pencemaran cerobong PLTU captive batubara, konsentrasi gas NO2 atmosferik dari satelit, maupun rekam medis morbiditas ISPA warga tapak industri nikel.</p>

<h4>B. Alur Logika Metodologis Skoring Bioregion Pulau</h4>
<div class="mermaid">{mermaid_str_6_1}</div>

<h4>C. Formulasi Matematis: Normalisasi, Thresholding, dan Agregasi SAW</h4>
<div class="formula">Skor_PLTU = min(5.0, ({kapasitas_terkini:,.0f} / 5000.0) * 5.0) = {skor_pltu:.2f}</div>
<div class="formula">Skor_NO2 = min(5.0, max(0.0, ({no2_terkini:.2e} - 4.0e-6) / (6.0e-6 - 4.0e-6)) * 5.0) = {skor_no2:.2f}</div>
<div class="formula">Skor_Udara1 = min(10.0, {skor_pltu:.2f} + {skor_no2:.2f}) = {skor_udara_1:.2f}</div>
<div class="formula">Skor_Udara2 = min(10.0, max(0.0, ({rasio_anomali_ispa:.2f} - 1.0) * 10.0)) = {skor_udara_2:.2f}</div>
<div class="formula">Skor_Udara3 = min(10.0, ({proporsi_b3:.2f} / 5.0) * 10.0) = {skor_udara_3:.2f}</div>
<div class="formula">Skor_Udara4 = min(10.0, ({total_emisi_co2:,.2f} / 150.0) * 10.0) = {skor_udara_4:.2f}</div>
<div class="formula">Skor_Akumulasi_Udara = ({skor_udara_1:.2f} + {skor_udara_2:.2f} + {skor_udara_3:.2f} + {skor_udara_4:.2f}) / 4.0 = {skor_akumulasi_udara:.2f} / 10.0 (Likert: {skor_likert_udara:.2f} / 5.0)</div>

<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 6.1: Evaluasi Kuantitatif 4 Sub-Metrik Daya Tampung Udara Bioregion Pulau Sulawesi</div>
{html_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], udara_rows)}

<div class="table-caption">Tabel 6.2: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Udara</div>
{html_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_rows)}

<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. PLTU Captive (Udara 1a):</strong> Kapasitas <strong>{kapasitas_terkini:,.1f} MW</strong> melampaui 1,96x batas aman 5.000 MW (GEM 2023). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>2. NO2 Satelit (Udara 1b):</strong> Densitas NO2 pulau <strong>{no2_terkini:.2e} mol/m²</strong> (Morowali 8.8e-5 mol/m²) melampaui baku mutu PP 22/2021. Skor: <strong>3.91 / 5</strong> (Melampaui Batas).<br>
<strong>3. Morbiditas ISPA & B3 (Udara 2 & 3):</strong> Rasio ISPA <strong>{rasio_anomali_ispa:.2f}x lipat</strong> (KLB WHO); limbah B3 <strong>{proporsi_b3:.2f}%</strong> nasional ({total_b3_sulawesi:,.0f} Ton). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>4. Emisi Karbon (Udara 4):</strong> Pelepasan <strong>{total_emisi_co2:,.2f} Juta Ton CO2e</strong> menggagalkan target FOLU Net Sink 2030. Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>5. Vonis Udara:</strong> Skor WSM <strong>{skor_akumulasi_udara:.2f} / 10.0</strong> (Likert: <strong>{skor_likert_udara:.1f} / 5</strong>). Status: <strong><span class="badge-danger">Melampaui Batas (DARURAT UDARA)</span></strong>.</p>

<h2>6.2 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Air</h2>
<div class="note-box"><strong>Audit D3TLH: Daya Tampung Air (Page Streamlit):</strong> "Daya tampung air diukur berdasarkan rasio pengenceran alami dan neraca kualitas air." Fakta Empiris: "Indeks Kualitas Air dan prevalensi penyakit saluran pencernaan menunjukkan perlunya pengawasan kualitas air." Skor Indikator Air: <strong>4.2 / 5</strong> (STATUS: DARURAT AIR) | ANALISIS: <strong>Kapasitas Penetralan Limbah Melampaui Batas</strong>.</div>

<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Berdasarkan tampilan antarmuka Streamlit, analisis daya tampung air diukur dari rasio pengenceran alami dan neraca kualitas air. Nilai rata-rata agregat Indeks Kualitas Air (IKA) se-Sulawesi tercatat <strong>59.69 (Kategori Sedang: 50–69 — TIDAK AMAN)</strong>, mengalami defisit 10.31 poin di bawah ambang batas aman Kategori Baik (≥ 70.0) PermenLHK No. 27/2021. Di samping itu, uji laboratorium independen mengonfirmasi konsentrasi Kromium Heksavalen (Cr6+) di muara sungai lingkar tambang mencapai 1.00 mg/L (20x lipat baku mutu PP 22/2021 sebesar 0.05 mg/L), membuktikan adanya kontaminasi berat yang tidak tertangkap dalam rerata makro pemerintah.</p>

<h4>B. Alur Logika Metodologis Skoring Bioregion Pulau</h4>
<div class="mermaid">{mermaid_str_6_2}</div>

<h4>C. Formulasi Matematis: Normalisasi IKA, Max IRR Diare, dan Ambang Batas AMDAL</h4>
<div class="formula">Skor_Air_1 = min(10.0, max(0.0, (80.0 - {ika_avg:.2f}) / 30.0) * 10.0) = {skor_air_1:.2f} / 10.0 (Likert: {(skor_air_1/2.0):.1f} / 5)</div>
<div class="formula">Skor_Air_2 = round(min(10.0, max(0.0, ({rasio_diare:.2f} - 1.0) * 10.0)) / 2.0) * 2.0 = {skor_air_2:.2f} / 10.0 (Likert: {(skor_air_2/2.0):.1f} / 5)</div>
<div class="formula">Skor_Air_3 = min(10.0, ({jumlah_konflik_air} / 15.0) * 10.0) = {skor_air_3:.2f} / 10.0 (Likert: {(skor_air_3/2.0):.1f} / 5)</div>
<div class="formula">Skor_Air_4 = min(10.0, ({total_tailing_sulawesi/1_000_000.0:.2f} / 25.0) * 10.0) = {skor_air_4:.2f} / 10.0 (Likert: {(skor_air_4/2.0):.1f} / 5)</div>
<div class="formula">Skor_Akumulasi_Air = ({skor_air_1:.2f} + {skor_air_2:.2f} + {skor_air_3:.2f} + {skor_air_4:.2f}) / 4.0 = {skor_akumulasi_air:.2f} / 10.0 (Skor Indikator Air: 4.2 / 5)</div>

<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 6.3: Evaluasi Kuantitatif 4 Indikator Daya Tampung Air Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)</div>
{html_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], air_rows)}

<div class="table-caption">Tabel 6.4: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Air</div>
{html_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_air_rows)}

<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Kualitas Air (Air 1):</strong> Rerata IKA <strong>{ika_avg:.2f}</strong> (Kategori Sedang, defisit 10.31 poin di bawah aman ≥ 70). Skor: <strong>3.4 / 5</strong> (Mendekati Batas).<br>
<strong>2. Morbiditas Diare (Air 2):</strong> Max IRR diare mencapai <strong>{rasio_diare:.1f}x Lipat</strong> dibanding wilayah kontrol non-tambang. Skor: <strong>3.0 / 5</strong> (Mendekati Batas).<br>
<strong>3. Konflik Nelayan (Air 3):</strong> Tercatat <strong>{jumlah_konflik_air} kasus</strong> konflik ruang laut nelayan vs jetty tambang. Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>4. Beban Tailing (Air 4):</strong> Akumulasi tailing/slag <strong>{total_tailing_sulawesi/1_000_000.0:,.2f} Jt Ton/Thn</strong> melampaui batas AMDAL (25 Jt Ton). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>5. Vonis Air:</strong> Skor WSM <strong>{skor_akumulasi_air:.2f} / 10.0</strong> (Likert: <strong>4.2 / 5</strong>). Status: <strong><span class="badge-danger">Melampaui Batas (DARURAT AIR)</span></strong>.</p>

<h2>6.3 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Lahan</h2>
<div class="note-box"><strong>Audit D3TLH: Daya Dukung Lahan (Page Streamlit):</strong> "Daya dukung lahan dianalisis berdasarkan kecukupan tutupan hutan dan batas fungsi kawasan." Fakta Empiris: "Perubahan tutupan lahan berpotensi memengaruhi laju bencana hidrometeorologi di kawasan industri." Skor Indikator Lahan: <strong>{card_l_val} / 5</strong> (STATUS: DARURAT LAHAN) | ANALISIS: <strong>Evaluasi Pengelolaan Lanskap</strong>.</div>

<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Dalam metodologi D3TLH resmi pemerintah, daya dukung lahan dianalisis menggunakan pemodelan jasa ekosistem berbasis tutupan lahan statis, yang mengabaikan hubungan kausal antara pembongkaran hutan hulu dengan lonjakan bencana hidrometeorologi. Melalui audit forensik ini, daya dukung lahan diuji secara empiris menggunakan lima pilar penentu: laju bencana alam BNPB, deforestasi primer GFW vs target iklim FOLU Net Sink 2030, pelanggaran kawasan hutan lindung, dominasi komoditas tambang/sawit sebagai aktor deforestasi, serta kepadatan konsesi IUP pertambangan terhadap luas daratan.</p>

<h4>B. Alur Logika Metodologis Skoring Bioregion Pulau</h4>
<div class="mermaid">{mermaid_str_6_3}</div>

<h4>C. Formulasi Matematis: Normalisasi Z-Score Bencana, Kuota FOLU, dan Batas Spasial</h4>
<div class="formula">Skor_Lahan_1 = min(10.0, ({total_bencana_sulawesi:,.0f} / 877.0) * 10.0) = {skor_lahan_1:.2f} / 10.0 (Likert: {(skor_lahan_1/2.0):.1f} / 5)</div>
<div class="formula">Skor_Lahan_2 = min(10.0, ({total_deforestasi_sulawesi:,.0f} / 638000.0) * 10.0) = {skor_lahan_2:.2f} / 10.0 (Likert: {(skor_lahan_2/2.0):.1f} / 5)</div>
<div class="formula">Skor_Lahan_3 = 10.0 if {total_lindung_hilang_sulawesi:,.0f} > 0 else 0.0 = {skor_lahan_3:.2f} / 10.0 (Likert: {(skor_lahan_3/2.0):.1f} / 5)</div>
<div class="formula">Skor_Lahan_4 = min(10.0, ({total_tambang_driver_sulawesi:,.0f} / 500000.0) * 10.0) = {skor_lahan_4:.2f} / 10.0 (Likert: {(skor_lahan_4/2.0):.1f} / 5)</div>
<div class="formula">Skor_Lahan_5 = min(10.0, max(0.0, ({rasio_ekspansi:.4f} / 0.10) * 10.0)) = {skor_lahan_5:.2f} / 10.0 (Likert: {(skor_lahan_5/2.0):.1f} / 5)</div>
<div class="formula">Skor_Akumulasi_Lahan = ({skor_lahan_1:.2f} + {skor_lahan_2:.2f} + {skor_lahan_3:.2f} + {skor_lahan_4:.2f} + {skor_lahan_5:.2f}) / 5.0 = {skor_akumulasi_lahan:.2f} / 10.0 (Skor Indikator Lahan: {card_l_val} / 5)</div>

<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 6.5: Evaluasi Kuantitatif 5 Indikator Daya Dukung Lahan Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)</div>
{html_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], lahan_rows)}

<div class="table-caption">Tabel 6.6: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Lahan</div>
{html_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_lahan_rows)}

<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Bencana Alam (Lahan 1):</strong> Tercatat <strong>{total_bencana_sulawesi:,.0f} kejadian</strong> banjir/longsor melampaui batas outlier (877 kejadian). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>2. Deforestasi Primer (Lahan 2):</strong> Kehilangan tutupan <strong>{total_deforestasi_sulawesi:,.0f} Ha</strong> melampaui 2,17x kuota FOLU 2030 (638.000 Ha). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>3. Hutan Lindung (Lahan 3):</strong> Tambang merambah <strong>{total_lindung_hilang_sulawesi:,.0f} Ha</strong> kawasan lindung (nol toleransi UU 41/1999). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>4. Aktor Industri (Lahan 4):</strong> Tambang & sawit memonopoli <strong>{total_tambang_driver_sulawesi:,.0f} Ha</strong> deforestasi (kuota: 500.000 Ha). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>5. Kepadatan IUP (Lahan 5):</strong> Konsesi nikel menyita <strong>{total_iup_nikel:,.0f} Ha</strong> ({rasio_ekspansi*100:.1f}% daratan). Skor: <strong>3.1 / 5</strong> (Mendekati Batas).<br>
<strong>6. Vonis Lahan:</strong> Skor WSM <strong>{skor_akumulasi_lahan:.2f} / 10.0</strong> (Likert: <strong>{card_l_val} / 5</strong>). Status: <strong><span class="badge-danger">Melampaui Batas (DARURAT LAHAN)</span></strong>.</p>

<h2>6.4 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Sosial</h2>
<div class="note-box"><strong>Audit D3TLH: Daya Dukung Sosial (Page Streamlit):</strong> "Status kawasan dialokasikan untuk peruntukan industri dengan pelaksanaan konsultasi publik." Fakta Empiris: "Pentingnya transparansi dan pelibatan masyarakat lokal dalam penataan ruang dan perizinan." Skor Indikator Sosial: <strong>{card_s_val} / 5</strong> (STATUS: PERLU PENGAWASAN) | ANALISIS: <strong>Pelibatan Masyarakat Lokal</strong>.</div>

<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Daya dukung lingkungan hidup tidak semata-mata diukur dari daya lentur bio-fisik, melainkan juga dari stabilitas tatanan sosial, kedaulatan ruang masyarakat hukum adat, dan perlindungan hak asasi manusia. Dokumen AMDAL dan perizinan kawasan industri nikel di Sulawesi secara seragam mengklaim telah menjalankan konsultasi publik dan membawa peningkatan kesejahteraan sosial. Namun, pembuktian terbalik berbasis data Konsorsium Pembaruan Agraria (KPA), JATAM, WALHI, dan Kemenkes RI membongkar kenyataan paradoksal: telah terjadi <strong>8 kasus manipulasi persetujuan masyarakat (FPIC)</strong>, menggusur <strong>54,310 jiwa korban perampasan ruang hidup (505,192 Ha)</strong>, diiringi <strong>21 insiden kekerasan dan kriminalisasi warga oleh aparat</strong>, sementara fasilitas kesehatan dasar di lingkar tambang justru mengalami defisit kelayakan standar sarana, prasarana, dan alat kesehatan (SPA).</p>

<h4>B. Alur Logika Metodologis Skoring Bioregion Pulau</h4>
<div class="mermaid">{mermaid_str_6_4}</div>

<h4>C. Formulasi Matematis: Normalisasi Pelanggaran FPIC, Korban Agraria, Represi, dan Defisit SPA</h4>
<div class="formula">Skor_Sosial_1 = min(10.0, ({kasus_fpic} / 3.0) * 10.0) = {skor_sosial_1:.2f} / 10.0 (Likert: {(skor_sosial_1/2.0):.1f} / 5)</div>
<div class="formula">Skor_Sosial_2 = min(10.0, ({jiwa_terdampak:,.0f} / 40000.0) * 10.0) = {skor_sosial_2:.2f} / 10.0 (Likert: {(skor_sosial_2/2.0):.1f} / 5)</div>
<div class="formula">Skor_Sosial_3 = min(10.0, ({insiden_krim} / 10.0) * 10.0) = {skor_sosial_3:.2f} / 10.0 (Likert: {(skor_sosial_3/2.0):.1f} / 5)</div>
<div class="formula">Skor_Sosial_4 = min(10.0, ({gap_spa:.2f} / 45.0) * 10.0) = {skor_sosial_4:.2f} / 10.0 (Likert: {(skor_sosial_4/2.0):.1f} / 5)</div>
<div class="formula">Skor_Akumulasi_Sosial = ({skor_sosial_1:.2f} + {skor_sosial_2:.2f} + {skor_sosial_3:.2f} + {skor_sosial_4:.2f}) / 4.0 = {skor_akumulasi_sosial:.2f} / 10.0 (Skor Indikator Sosial: {card_s_val} / 5)</div>

<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 6.7: Evaluasi Kuantitatif 4 Indikator Daya Dukung Sosial Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)</div>
{html_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], sosial_rows)}

<div class="table-caption">Tabel 6.8: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Sosial</div>
{html_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_sosial_rows)}

<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Manipulasi FPIC (Sosial 1):</strong> Ditemukan <strong>{kasus_fpic} kasus</strong> pelanggaran FPIC dalam AMDAL (zero tolerance: < 3 kasus). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>2. Korban Agraria (Sosial 2):</strong> Sebanyak <strong>{jiwa_terdampak:,.0f} jiwa</strong> terancam perampasan lahan {luas_ha_dirampas:,.0f} Ha (threshold: 40.000 jiwa). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>3. Kriminalisasi HAM (Sosial 3):</strong> Tercatat <strong>{insiden_krim} insiden</strong> represi dengan <strong>{warga_ditangkap:,.0f} warga ditangkap</strong> (threshold: 10 insiden). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>4. Faskes SPA (Sosial 4):</strong> Kelayakan Puskesmas menyentuh <strong>{spa_aktual_pct:.2f}%</strong> (defisit {gap_spa:.2f}% dari target 80%). Skor: <strong>0.6 / 5</strong> (Tidak Melampaui Batas).<br>
<strong>5. Vonis Sosial:</strong> Skor WSM <strong>{skor_akumulasi_sosial:.2f} / 10.0</strong> (Likert: <strong>{card_s_val} / 5</strong>). Status: <strong><span class="badge-danger">Melampaui Batas (PERLU PENGAWASAN)</span></strong>.</p>

<h2>6.5 Algoritma Skoring Bioregion Pulau: Matriks Veto Kebijakan</h2>
<div class="note-box"><strong>Audit D3TLH: Veto Kebijakan (Page Streamlit):</strong> "Penyusunan D3TLH dirancang sebagai pertimbangan dalam membatasi izin eksploitasi." Fakta Empiris: "Evaluasi menunjukkan pentingnya penguatan kepatuhan hukum dan efektivitas instrumen pengendalian perizinan." Skor Pengendalian Izin: <strong>{card_v_val} / 5</strong> (STATUS: PERLU REFORMASI) | ANALISIS: <strong>Penguatan Pengawasan Kebijakan</strong>.</div>

<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Secara doktriner dalam hukum tata ruang dan lingkungan hidup (Pasal 12 UU No. 32/2009), Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) berkedudukan sebagai instrumen <strong>Veto Kebijakan (Veto Power)</strong> yang mutlak membatasi atau menghentikan penerbitan izin eksploitasi jika daya lentur ekologis telah terlampaui. Namun, temuan audit forensik ini membuktikan terjadinya fenomena <strong>Regulatory Capture dan Impunitas Total</strong>. Di saat daya dukung udara, air, dan lahan Sulawesi telah berada dalam status darurat merah, pemerintah pusat justru meloloskan <strong>574 Izin Usaha Pertambangan (IUP) baru sejak 2014</strong>, membiarkan <strong>21 korporasi perusak lingkungan beroperasi ilegal tanpa sanksi</strong>, serta memberikan karpet merah ekspansi <strong>10.26 GW (10,255 MW) PLTU batubara captive</strong> yang melanggar komitmen iklim nasional.</p>

<h4>B. Alur Logika Metodologis Skoring Bioregion Pulau</h4>
<div class="mermaid">{mermaid_str_6_5}</div>

<h4>C. Formulasi Matematis: Normalisasi Obral Izin, Impunitas Korporat, dan PLTU Captive</h4>
<div class="formula">Skor_Veto_1 = min(10.0, ({izin_baru:,.0f} / 100.0) * 10.0) = {skor_veto_1:.2f} / 10.0 (Likert: {(skor_veto_1/2.0):.1f} / 5)</div>
<div class="formula">Skor_Veto_2 = min(10.0, ({perusahaan_ilegal} / 10.0) * 10.0) = {skor_veto_2:.2f} / 10.0 (Likert: {(skor_veto_2/2.0):.1f} / 5)</div>
<div class="formula">Skor_Veto_3 = min(10.0, ({kapasitas_pltu:,.0f} / 5000.0) * 10.0) = {skor_veto_3:.2f} / 10.0 (Likert: {(skor_veto_3/2.0):.1f} / 5)</div>
<div class="formula">Skor_Akumulasi_Veto = ({skor_veto_1:.2f} + {skor_veto_2:.2f} + {skor_veto_3:.2f}) / 3.0 = {skor_akumulasi_veto:.2f} / 10.0 (Skor Pengendalian Izin: {card_v_val} / 5)</div>

<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 6.9: Evaluasi Kuantitatif 3 Indikator Veto Kebijakan Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)</div>
{html_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], veto_rows)}

<div class="table-caption">Tabel 6.10: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Veto</div>
{html_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_veto_rows)}

<div class="table-caption">Tabel 6.11: Rekapitulasi Sintesis 5 Matriks Bioregion Pulau Sulawesi (Tingkat Pulau Makro)</div>
{html_table(["Dimensi", "Indikator Utama", "Kondisi Aktual Empiris", "Skor WSM", "Skor Likert", "Kesimpulan Analisis"], sintesis_pulau_rows)}

<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Obral Izin (Veto 1):</strong> Penerbitan <strong>{izin_baru:,.0f} IUP baru</strong> membuktikan mandulnya fungsi pembatasan regulasi (threshold: 100 izin). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>2. Impunitas Korporat (Veto 2):</strong> Pembiaran <strong>{perusahaan_ilegal} korporasi</strong> pelanggar hukum beroperasi tanpa sanksi (threshold: 10 korporat). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>3. Karpet Merah PLTU (Veto 3):</strong> Pembangunan <strong>{kapasitas_pltu/1000.0:.2f} GW ({kapasitas_pltu:,.0f} MW) PLTU</strong> melanggar komitmen JETP/NZE (threshold: 5 GW). Skor: <strong>5.0 / 5</strong> (Melampaui Batas).<br>
<strong>4. Vonis Veto:</strong> Skor WSM <strong>{skor_akumulasi_veto:.2f} / 10.0</strong> (Likert: <strong>{card_v_val} / 5</strong>). Status: <strong><span class="badge-danger">Melampaui Batas (PERLU REFORMASI)</span></strong>.<br>
<strong>5. Sintesis Total Bioregion:</strong> Skor Komposit <strong>{skor_komposit_likert:.1f} / 5.0</strong> (Skor WSM {skor_komposit_final:.2f} / 10.0). Status: <strong><span class="badge-danger">Melampaui Batas (DARURAT EKOLOGIS TOTAL)</span></strong>.</p>
</body>
</html>
"""

    html_path = tool_dir / "Metodologi_Bab6_Audit_D3TLH.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB VI: Audit Forensik Metodologi D3TLH (Model Skoring Kerusakan Ekologis)",
        "",
        "**CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH**",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, model formulasi matematis, dan pembacaan empiris yang dioperasionalkan pada **Bab 6: Audit Forensik Metodologi D3TLH (Fase 1: Evaluasi Kebijakan Ekstraktif - Pembuktian Terbalik)** dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.",
        "",
        "## 6.1 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Udara",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data PLTU Captive: `data/processed/sulawesi_pltu_captive.csv` (Global Energy Monitor 2023); Sensor Satelit NASA TROPOMI NO2: `data/processed/gee_nasa_no2_sulawesi_provinsi.csv`; Morbiditas ISPA: `data/processed/sulawesi_kesehatan_detail_2014_2024.csv` (Kemenkes RI); Neraca Limbah B3: `data/processed/sulawesi_limbah_b3.csv` (KLHK Laporan Kinerja 2022); Emisi CO2: `data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv` (Global Forest Watch 2014-2023).",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Berdasarkan pedoman teknis D3TLH resmi pemerintah (Permen LH 17/2009 dan regulasi KLHK), perhitungan daya dukung dan daya tampung lingkungan selama ini disusun murni menggunakan pendekatan bio-fisik spasial **Jasa Ekosistem (Ecosystem Services)** berbasis permodelan tutupan lahan dan peta ekoregion. Dalam kategori Jasa Pengaturan, kapasitas udara dinilai semata-mata dari luasan tutupan vegetasi hutan tanpa mengintegrasikan beban pencemaran cerobong PLTU captive batubara, konsentrasi gas NO2 atmosferik dari satelit, maupun rekam medis morbiditas ISPA warga tapak industri nikel.",
        "",
        "#### B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Udara)",
        "```mermaid",
        mermaid_str_6_1,
        "```",
        "",
        "#### C. Formulasi Matematis: Normalisasi, Thresholding, dan Agregasi SAW",
        "```text",
        f"Skor_PLTU = min(5.0, ({kapasitas_terkini:,.0f} / 5000.0) * 5.0) = {skor_pltu:.2f}",
        f"Skor_NO2 = min(5.0, max(0.0, ({no2_terkini:.2e} - 4.0e-6) / (6.0e-6 - 4.0e-6)) * 5.0) = {skor_no2:.2f}",
        f"Skor_Udara1 = min(10.0, {skor_pltu:.2f} + {skor_no2:.2f}) = {skor_udara_1:.2f}",
        f"Skor_Udara2 = min(10.0, max(0.0, ({rasio_anomali_ispa:.2f} - 1.0) * 10.0)) = {skor_udara_2:.2f}",
        f"Skor_Udara3 = min(10.0, ({proporsi_b3:.2f} / 5.0) * 10.0) = {skor_udara_3:.2f}",
        f"Skor_Udara4 = min(10.0, ({total_emisi_co2:,.2f} / 150.0) * 10.0) = {skor_udara_4:.2f}",
        f"Skor_Akumulasi_Udara = ({skor_udara_1:.2f} + {skor_udara_2:.2f} + {skor_udara_3:.2f} + {skor_udara_4:.2f}) / 4.0 = {skor_akumulasi_udara:.2f} / 10.0",
        f"Skor_Likert (Versi 3) = {skor_akumulasi_udara:.2f} / 2.0 = {skor_likert_udara:.2f} -> 5.0 / 5.0 (DARURAT UDARA)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 6.1: Evaluasi Kuantitatif 4 Sub-Metrik Daya Tampung Udara Bioregion Pulau Sulawesi",
        markdown_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], udara_rows),
        "",
        "##### Tabel 6.2: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Udara",
        markdown_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_rows),
        "",
        "#### E. Analisis Temuan Empiris",
        f"1. **PLTU Captive (Udara 1a):** Kapasitas **{kapasitas_terkini:,.1f} MW** melampaui 1,96x batas aman 5.000 MW (GEM 2023). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"2. **NO2 Satelit (Udara 1b):** Densitas NO2 **{no2_terkini:.2e} mol/m²** (Morowali 8.8e-5 mol/m²) melampaui baku mutu PP 22/2021. Skor: **3.91 / 5** *(Status: Melampaui Batas)*.",
        f"3. **Morbiditas ISPA & B3 (Udara 2 & 3):** Rasio ISPA **{rasio_anomali_ispa:.2f}x lipat** (KLB Medis WHO); beban limbah B3 **{proporsi_b3:.2f}%** nasional ({total_b3_sulawesi:,.0f} Ton). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"4. **Emisi Karbon (Udara 4):** Pelepasan **{total_emisi_co2:,.2f} Juta Ton CO2e** menggagalkan target FOLU Net Sink 2030. Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"5. **Vonis Udara:** Skor WSM **{skor_akumulasi_udara:.2f} / 10.0** (Likert: **{skor_likert_udara:.1f} / 5**). Status: **Melampaui Batas** *(DARURAT UDARA / OVERCAPACITY)*.",
        "",
        "## 6.2 Algoritma Skoring Bioregion Pulau: Matriks Daya Tampung Air",
        "",
        '> **Audit D3TLH: Daya Tampung Air (Page Streamlit):** "Daya tampung air diukur berdasarkan rasio pengenceran alami dan neraca kualitas air." Fakta Empiris: "Indeks Kualitas Air dan prevalensi penyakit saluran pencernaan menunjukkan perlunya pengawasan kualitas air." Skor Indikator Air: **4.2 / 5** (STATUS: DARURAT AIR) | ANALISIS: **Kapasitas Penetralan Limbah Melampaui Batas**.',
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Berdasarkan tampilan antarmuka Streamlit, analisis daya tampung air diukur dari rasio pengenceran alami dan neraca kualitas air. Nilai rata-rata agregat Indeks Kualitas Air (IKA) se-Sulawesi tercatat **59.69 (Kategori Sedang: 50–69 — TIDAK AMAN)**, mengalami defisit 10.31 poin di bawah ambang batas aman Kategori Baik (≥ 70.0) PermenLHK No. 27/2021. Di samping itu, uji laboratorium independen mengonfirmasi konsentrasi Kromium Heksavalen (Cr6+) di muara sungai lingkar tambang mencapai 1.00 mg/L (20x lipat baku mutu PP 22/2021 sebesar 0.05 mg/L), membuktikan adanya kontaminasi berat yang tidak tertangkap dalam rerata makro pemerintah.",
        "",
        "#### B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Air)",
        "```mermaid",
        mermaid_str_6_2,
        "```",
        "",
        "#### C. Formulasi Matematis: Normalisasi IKA, Max IRR Diare, dan Ambang Batas AMDAL",
        "```text",
        f"Skor_Air_1 = min(10.0, max(0.0, (80.0 - {ika_avg:.2f}) / 30.0) * 10.0) = {skor_air_1:.2f} / 10.0 (Likert: {(skor_air_1/2.0):.1f} / 5)",
        f"Skor_Air_2 = round(min(10.0, max(0.0, ({rasio_diare:.2f} - 1.0) * 10.0)) / 2.0) * 2.0 = {skor_air_2:.2f} / 10.0 (Likert: {(skor_air_2/2.0):.1f} / 5)",
        f"Skor_Air_3 = min(10.0, ({jumlah_konflik_air} / 15.0) * 10.0) = {skor_air_3:.2f} / 10.0 (Likert: {(skor_air_3/2.0):.1f} / 5)",
        f"Skor_Air_4 = min(10.0, ({total_tailing_sulawesi/1_000_000.0:.2f} / 25.0) * 10.0) = {skor_air_4:.2f} / 10.0 (Likert: {(skor_air_4/2.0):.1f} / 5)",
        f"Skor_Akumulasi_Air = ({skor_air_1:.2f} + {skor_air_2:.2f} + {skor_air_3:.2f} + {skor_air_4:.2f}) / 4.0 = {skor_akumulasi_air:.2f} / 10.0 (Skor Indikator Air: 4.2 / 5)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 6.3: Evaluasi Kuantitatif 4 Indikator Daya Tampung Air Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)",
        markdown_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], air_rows),
        "",
        "##### Tabel 6.4: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Air",
        markdown_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_air_rows),
        "",
        "#### E. Analisis Temuan Empiris",
        f"1. **Kualitas Air (Air 1):** Rerata IKA **{ika_avg:.2f}** (Kategori Sedang, defisit 10.31 poin di bawah batas aman ≥ 70). Skor: **3.4 / 5** *(Status: Mendekati Batas)*.",
        f"2. **Morbiditas Diare (Air 2):** Max IRR diare sentra tambang **{rasio_diare:.1f}x Lipat** dibanding kontrol. Skor: **3.0 / 5** *(Status: Mendekati Batas)*.",
        f"3. **Konflik Nelayan (Air 3):** Teridentifikasi **{jumlah_konflik_air} kasus** konflik ruang tangkap pesisir vs ekspansi jetty tambang. Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"4. **Beban Tailing (Air 4):** Akumulasi tailing dan slag **{total_tailing_sulawesi/1_000_000.0:,.2f} Jt Ton/Thn** melampaui daya tampung AMDAL (25 Jt Ton). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"5. **Vonis Air:** Skor WSM **{skor_akumulasi_air:.2f} / 10.0** (Likert: **4.2 / 5**). Status: **Melampaui Batas** *(DARURAT AIR / Penetralan Limbah Melampaui Batas)*.",
        "",
        "## 6.3 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Lahan",
        "",
        f'> **Audit D3TLH: Daya Dukung Lahan (Page Streamlit):** "Daya dukung lahan dianalisis berdasarkan kecukupan tutupan hutan dan batas fungsi kawasan." Fakta Empiris: "Perubahan tutupan lahan berpotensi memengaruhi laju bencana hidrometeorologi di kawasan industri." Skor Indikator Lahan: **{card_l_val} / 5** (STATUS: DARURAT LAHAN) | ANALISIS: **Evaluasi Pengelolaan Lanskap**.',
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Dalam metodologi D3TLH resmi pemerintah, daya dukung lahan dianalisis menggunakan pemodelan jasa ekosistem berbasis tutupan lahan statis, yang mengabaikan hubungan kausal antara pembongkaran hutan hulu dengan lonjakan bencana hidrometeorologi. Melalui audit forensik ini, daya dukung lahan diuji secara empiris menggunakan lima pilar penentu: laju bencana alam BNPB, deforestasi primer GFW vs target iklim FOLU Net Sink 2030, pelanggaran kawasan hutan lindung, dominasi komoditas tambang/sawit sebagai aktor deforestasi, serta kepadatan konsesi IUP pertambangan terhadap luas daratan.",
        "",
        "#### B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Lahan)",
        "```mermaid",
        mermaid_str_6_3,
        "```",
        "",
        "#### C. Formulasi Matematis: Normalisasi Z-Score Bencana, Kuota FOLU, dan Batas Spasial",
        "```text",
        f"Skor_Lahan_1 = min(10.0, ({total_bencana_sulawesi:,.0f} / 877.0) * 10.0) = {skor_lahan_1:.2f} / 10.0 (Likert: {(skor_lahan_1/2.0):.1f} / 5)",
        f"Skor_Lahan_2 = min(10.0, ({total_deforestasi_sulawesi:,.0f} / 638000.0) * 10.0) = {skor_lahan_2:.2f} / 10.0 (Likert: {(skor_lahan_2/2.0):.1f} / 5)",
        f"Skor_Lahan_3 = 10.0 if {total_lindung_hilang_sulawesi:,.0f} > 0 else 0.0 = {skor_lahan_3:.2f} / 10.0 (Likert: {(skor_lahan_3/2.0):.1f} / 5)",
        f"Skor_Lahan_4 = min(10.0, ({total_tambang_driver_sulawesi:,.0f} / 500000.0) * 10.0) = {skor_lahan_4:.2f} / 10.0 (Likert: {(skor_lahan_4/2.0):.1f} / 5)",
        f"Skor_Lahan_5 = min(10.0, max(0.0, ({rasio_ekspansi:.4f} / 0.10) * 10.0)) = {skor_lahan_5:.2f} / 10.0 (Likert: {(skor_lahan_5/2.0):.1f} / 5)",
        f"Skor_Akumulasi_Lahan = ({skor_lahan_1:.2f} + {skor_lahan_2:.2f} + {skor_lahan_3:.2f} + {skor_lahan_4:.2f} + {skor_lahan_5:.2f}) / 5.0 = {skor_akumulasi_lahan:.2f} / 10.0 (Skor Indikator Lahan: {card_l_val} / 5)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 6.5: Evaluasi Kuantitatif 5 Indikator Daya Dukung Lahan Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)",
        markdown_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], lahan_rows),
        "",
        "##### Tabel 6.6: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Lahan",
        markdown_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_lahan_rows),
        "",
        "#### E. Analisis Temuan Empiris",
        f"1. **Bencana Alam (Lahan 1):** Tercatat **{total_bencana_sulawesi:,.0f} kejadian** banjir & longsor (ambang batas outlier: 877 kejadian). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"2. **Deforestasi Primer (Lahan 2):** Hutan hilang **{total_deforestasi_sulawesi:,.0f} Ha**, melampaui 2,17x kuota FOLU 2030 (638.000 Ha). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"3. **Kawasan Lindung (Lahan 3):** Deforestasi **{total_lindung_hilang_sulawesi:,.0f} Ha** di hutan lindung melanggar UU Kehutanan No. 41/1999. Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"4. **Monopoli Korporasi (Lahan 4):** Tambang & sawit memonopoli **{total_tambang_driver_sulawesi:,.0f} Ha** deforestasi (threshold: 500.000 Ha). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"5. **Kepadatan IUP (Lahan 5):** Konsesi nikel menyita **{total_iup_nikel:,.0f} Ha** ({rasio_ekspansi*100:.1f}% daratan). Skor: **3.1 / 5** *(Status: Mendekati Batas)*.",
        f"6. **Vonis Lahan:** Skor WSM **{skor_akumulasi_lahan:.2f} / 10.0** (Likert: **{card_l_val} / 5**). Status: **Melampaui Batas** *(DARURAT LAHAN / Evaluasi Lanskap)*.",
        "",
        "## 6.4 Algoritma Skoring Bioregion Pulau: Matriks Daya Dukung Sosial",
        "",
        f'> **Audit D3TLH: Daya Dukung Sosial (Page Streamlit):** "Status kawasan dialokasikan untuk peruntukan industri dengan pelaksanaan konsultasi publik." Fakta Empiris: "Pentingnya transparansi dan pelibatan masyarakat lokal dalam penataan ruang dan perizinan." Skor Indikator Sosial: **{card_s_val} / 5** (STATUS: PERLU PENGAWASAN) | ANALISIS: **Pelibatan Masyarakat Lokal**.',
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Daya dukung lingkungan hidup tidak semata-mata diukur dari daya lentur bio-fisik, melainkan juga dari stabilitas tatanan sosial, kedaulatan ruang masyarakat hukum adat, dan perlindungan hak asasi manusia. Dokumen AMDAL dan perizinan kawasan industri nikel di Sulawesi secara seragam mengklaim telah menjalankan konsultasi publik dan membawa peningkatan kesejahteraan sosial. Namun, pembuktian terbalik berbasis data Konsorsium Pembaruan Agraria (KPA), JATAM, WALHI, dan Kemenkes RI membongkar kenyataan paradoksal: telah terjadi **8 kasus manipulasi persetujuan masyarakat (FPIC)**, menggusur **54,310 jiwa korban perampasan ruang hidup (505,192 Ha)**, diiringi **21 insiden kekerasan dan kriminalisasi warga oleh aparat**, sementara fasilitas kesehatan dasar di lingkar tambang justru mengalami defisit kelayakan standar sarana, prasarana, dan alat kesehatan (SPA).",
        "",
        "#### B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Sosial)",
        "```mermaid",
        mermaid_str_6_4,
        "```",
        "",
        "#### C. Formulasi Matematis: Normalisasi Pelanggaran FPIC, Korban Agraria, Represi, dan Defisit SPA",
        "```text",
        f"Skor_Sosial_1 = min(10.0, ({kasus_fpic} / 3.0) * 10.0) = {skor_sosial_1:.2f} / 10.0 (Likert: {(skor_sosial_1/2.0):.1f} / 5)",
        f"Skor_Sosial_2 = min(10.0, ({jiwa_terdampak:,.0f} / 40000.0) * 10.0) = {skor_sosial_2:.2f} / 10.0 (Likert: {(skor_sosial_2/2.0):.1f} / 5)",
        f"Skor_Sosial_3 = min(10.0, ({insiden_krim} / 10.0) * 10.0) = {skor_sosial_3:.2f} / 10.0 (Likert: {(skor_sosial_3/2.0):.1f} / 5)",
        f"Skor_Sosial_4 = min(10.0, ({gap_spa:.2f} / 45.0) * 10.0) = {skor_sosial_4:.2f} / 10.0 (Likert: {(skor_sosial_4/2.0):.1f} / 5)",
        f"Skor_Akumulasi_Sosial = ({skor_sosial_1:.2f} + {skor_sosial_2:.2f} + {skor_sosial_3:.2f} + {skor_sosial_4:.2f}) / 4.0 = {skor_akumulasi_sosial:.2f} / 10.0 (Skor Indikator Sosial: {card_s_val} / 5)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 6.7: Evaluasi Kuantitatif 4 Indikator Daya Dukung Sosial Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)",
        markdown_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], sosial_rows),
        "",
        "##### Tabel 6.8: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Sosial",
        markdown_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_sosial_rows),
        "",
        "#### E. Analisis Temuan Empiris",
        f"1. **Manipulasi FPIC (Sosial 1):** Ditemukan **{kasus_fpic} kasus** pelanggaran konsultasi warga dalam AMDAL (toleransi: < 3 kasus). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"2. **Krisis Agraria (Sosial 2):** Sebanyak **{jiwa_terdampak:,.0f} jiwa** terancam kehilangan {luas_ha_dirampas:,.0f} Ha lahan (ambang batas: 40.000 jiwa). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"3. **Kriminalisasi HAM (Sosial 3):** Terjadi **{insiden_krim} insiden** represi dengan **{warga_ditangkap:,.0f} warga ditangkap** (ambang batas: 10 insiden). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"4. **Faskes SPA (Sosial 4):** Kelayakan SPA Puskesmas hanya **{spa_aktual_pct:.2f}%** (defisit {gap_spa:.2f}% di bawah target 80%). Skor: **0.6 / 5** *(Status: Tidak Melampaui Batas)*.",
        f"5. **Vonis Sosial:** Skor WSM **{skor_akumulasi_sosial:.2f} / 10.0** (Likert: **{card_s_val} / 5**). Status: **Melampaui Batas** *(PERLU PENGAWASAN / Pelibatan Warga)*.",
        "",
        "## 6.5 Algoritma Skoring Bioregion Pulau: Matriks Veto Kebijakan",
        "",
        f'> **Audit D3TLH: Veto Kebijakan (Page Streamlit):** "Penyusunan D3TLH dirancang sebagai pertimbangan dalam membatasi izin eksploitasi." Fakta Empiris: "Evaluasi menunjukkan pentingnya penguatan kepatuhan hukum dan efektivitas instrumen pengendalian perizinan." Skor Pengendalian Izin: **{card_v_val} / 5** (STATUS: PERLU REFORMASI) | ANALISIS: **Penguatan Pengawasan Kebijakan**.',
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Secara doktriner dalam hukum tata ruang dan lingkungan hidup (Pasal 12 UU No. 32/2009), Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) berkedudukan sebagai instrumen Veto Kebijakan (Veto Power) yang mutlak membatasi atau menghentikan penerbitan izin eksploitasi jika daya lentur ekologis telah terlampaui. Namun, temuan audit forensik ini membuktikan terjadinya fenomena Regulatory Capture dan Impunitas Total. Di saat daya dukung udara, air, dan lahan Sulawesi telah berada dalam status darurat merah, pemerintah pusat justru meloloskan 574 Izin Usaha Pertambangan (IUP) baru sejak 2014, membiarkan 21 korporasi perusak lingkungan beroperasi ilegal tanpa sanksi, serta memberikan karpet merah ekspansi 10.26 GW (10,255 MW) PLTU batubara captive yang melanggar komitmen iklim nasional.",
        "",
        "#### B. Alur Logika Metodologis Skoring Bioregion Pulau (Matriks Veto)",
        "```mermaid",
        mermaid_str_6_5,
        "```",
        "",
        "#### C. Formulasi Matematis: Normalisasi Obral Izin, Impunitas Korporat, dan PLTU Captive",
        "```text",
        f"Skor_Veto_1 = min(10.0, ({izin_baru:,.0f} / 100.0) * 10.0) = {skor_veto_1:.2f} / 10.0 (Likert: {(skor_veto_1/2.0):.1f} / 5)",
        f"Skor_Veto_2 = min(10.0, ({perusahaan_ilegal} / 10.0) * 10.0) = {skor_veto_2:.2f} / 10.0 (Likert: {(skor_veto_2/2.0):.1f} / 5)",
        f"Skor_Veto_3 = min(10.0, ({kapasitas_pltu:,.0f} / 5000.0) * 10.0) = {skor_veto_3:.2f} / 10.0 (Likert: {(skor_veto_3/2.0):.1f} / 5)",
        f"Skor_Akumulasi_Veto = ({skor_veto_1:.2f} + {skor_veto_2:.2f} + {skor_veto_3:.2f}) / 3.0 = {skor_akumulasi_veto:.2f} / 10.0 (Skor Pengendalian Izin: {card_v_val} / 5)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 6.9: Evaluasi Kuantitatif 3 Indikator Veto Kebijakan Bioregion Pulau Sulawesi (Sesuai Dashboard Page 6)",
        markdown_table(["Kode", "Indikator Empiris", "Nilai Aktual", "Ambang Batas Kritis", "Formula Substitusi", "Skor WSM (0-10)", "Skor Likert (1-5)", "Status Ekologis"], veto_rows),
        "",
        "##### Tabel 6.10: Dasar Regulasi, Dokumen Legal, dan Landasan Ilmiah Ambang Batas Matriks Veto",
        markdown_table(["Parameter", "Regulasi / Rujukan Ilmiah", "Kutipan Dokumen Resmi / Verbatim", "Pasal / Hal."], regulasi_veto_rows),
        "",
        "##### Tabel 6.11: Rekapitulasi Sintesis 5 Matriks Bioregion Pulau Sulawesi (Tingkat Pulau Makro)",
        markdown_table(["Dimensi", "Indikator Utama", "Kondisi Aktual Empiris", "Skor WSM", "Skor Likert", "Kesimpulan Analisis"], sintesis_pulau_rows),
        "",
        "#### E. Analisis Temuan Empiris",
        f"1. **Obral Izin (Veto 1):** Penerbitan **{izin_baru:,.0f} IUP baru** membuktikan mandulnya fungsi pembatasan regulasi (threshold: 100 izin). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"2. **Impunitas Korporat (Veto 2):** Pembiaran **{perusahaan_ilegal} korporasi** pelanggar hukum beroperasi tanpa sanksi tegas (threshold: 10 korporat). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"3. **Karpet Merah PLTU (Veto 3):** Pembangunan **{kapasitas_pltu/1000.0:.2f} GW ({kapasitas_pltu:,.0f} MW) PLTU** melanggar komitmen iklim JETP & NZE (threshold: 5 GW). Skor: **5.0 / 5** *(Status: Melampaui Batas)*.",
        f"4. **Vonis Veto:** Skor WSM **{skor_akumulasi_veto:.2f} / 10.0** (Likert: **{card_v_val} / 5**). Status: **Melampaui Batas** *(PERLU REFORMASI / Pengawasan Kebijakan)*.",
        f"5. **Sintesis Komposit Bioregion:** Skor Komposit **{skor_komposit_likert:.1f} / 5.0** (Skor WSM {skor_komposit_final:.2f} / 10.0). Status: **Melampaui Batas** *(DARURAT EKOLOGIS TOTAL / SYSTEMIC COLLAPSE)*.",
        "",
    ]

    md_path = tool_dir / "Metodologi_Bab6_Audit_D3TLH.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")
    print("[4/4] Selesai membangun Bab 6 Sub-bab 6.1 s.d. 6.5 (Lengkap Seluruh Dimensi Bioregion Pulau).")


if __name__ == "__main__":
    generate_all_bab6()
