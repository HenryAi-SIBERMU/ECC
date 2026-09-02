#!/usr/bin/env python3
"""
Generator Laporan Metodologi Bab 9: Demografi Sosial - Ketika Hilirisasi Mengubah Struktur Masyarakat

Pilar 1 ditulis langsung dalam generator Python agar selaras dengan SOP dokumentasi Celios2.
Fokus awal: Sub-bab 9.1 Tekanan Demografi di Kabupaten Industri Ekstraktif.
"""

import base64
import sys
from pathlib import Path

try:
    import numpy as np
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([
        sys.executable,
        "-m",
        "pip",
        "install",
        "numpy",
        "pandas",
        "requests",
        "python-docx",
    ])
    import numpy as np
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, cfg):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        for key, val in cfg.items():
            el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def para_border_bottom(paragraph, color="2E7D32", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="2E7D32", sz="16"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="A5D6A7", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="1B5E20", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="2E7D32", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=G_MID)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)


def add_p(doc, parts, space_after=5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "EDF7EE")
    all_border_para(p)
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)

    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=G_DARK)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p)
    para_shd(p, "F1F8E9")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}

    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "2E7D32")
        set_cell_borders(cell, bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8, color=C_WHITE)

    for i, row_data in enumerate(rows):
        fill = "F5FBF5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            set_cell_borders(cell, bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def get_box_stats_hazen(series):
    """Replikasi persis anatomi boxplot dashboard (algoritma default Plotly = numpy method 'hazen')."""
    series = series.dropna()
    if series.empty:
        return {k: 0 for k in ("med", "q1", "q3", "lower", "upper", "min", "max")}
    q1 = np.quantile(series, 0.25, method="hazen")
    med = float(np.median(series))
    q3 = np.quantile(series, 0.75, method="hazen")
    iqr = q3 - q1
    lower_bound = q1 - 1.5 * iqr
    upper_bound = q3 + 1.5 * iqr
    lower_fence = series[series >= lower_bound].min()
    upper_fence = series[series <= upper_bound].max()
    return {"med": med, "q1": float(q1), "q3": float(q3), "lower": float(lower_fence), "upper": float(upper_fence), "min": float(series.min()), "max": float(series.max())}


def generate_all_bab9():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    data_dir = base_dir / "data" / "processed"
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_9"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 9 sub-bab 9.1...")
    df_demo = pd.read_csv(data_dir / "sulawesi_demografi_master_fase4.csv")
    df_demo["tahun"] = pd.to_numeric(df_demo["tahun"], errors="coerce")

    smelter_kabs_91 = sorted(df_demo[df_demo["is_smelter"] == True]["kabupaten"].unique())
    n_smelter_kab_91 = len(smelter_kabs_91)
    latest_year_91 = int(df_demo[df_demo["tahun"] <= 2024]["tahun"].max())
    latest_demo_91 = df_demo[df_demo["tahun"] == latest_year_91]

    smelter_window_91 = df_demo[(df_demo["is_smelter"] == True) & (df_demo["tahun"] <= 2024)]
    non_smelter_window_91 = df_demo[(df_demo["is_smelter"] == False) & (df_demo["tahun"] <= 2024)]
    s_yoy_91 = smelter_window_91["laju_pertumbuhan_yoy_pct"].dropna()
    ns_yoy_91 = non_smelter_window_91["laju_pertumbuhan_yoy_pct"].dropna()

    smelter_avg_yoy_91 = float(s_yoy_91.mean())
    non_smelter_avg_yoy_91 = float(ns_yoy_91.mean())
    smelter_std_91 = float(s_yoy_91.std())
    non_smelter_std_91 = float(ns_yoy_91.std())
    n_s_91 = int(s_yoy_91.count())
    n_ns_91 = int(ns_yoy_91.count())
    smelter_total_pop_latest_91 = float(latest_demo_91[latest_demo_91["is_smelter"] == True]["jumlah_penduduk_rb"].sum())

    s_91 = get_box_stats_hazen(s_yoy_91)
    ns_91 = get_box_stats_hazen(ns_yoy_91)

    box_rows_91 = [
        ["Kabupaten Industri Ekstraktif", f"{s_91['max']:.2f}", f"{s_91['upper']:.2f}", f"{s_91['q3']:.2f}", f"{s_91['med']:.2f}", f"{s_91['q1']:.2f}", f"{s_91['lower']:.2f}", f"{s_91['min']:.2f}"],
        ["Kabupaten Non-Ekstraktif", f"{ns_91['max']:.2f}", f"{ns_91['upper']:.2f}", f"{ns_91['q3']:.2f}", f"{ns_91['med']:.2f}", f"{ns_91['q1']:.2f}", f"{ns_91['lower']:.2f}", f"{ns_91['min']:.2f}"],
    ]
    mean_rows_91 = [
        ["Kabupaten Industri Ekstraktif", f"{smelter_avg_yoy_91:.2f}", f"{smelter_std_91:.2f}", f"{n_s_91}"],
        ["Kabupaten Non-Ekstraktif", f"{non_smelter_avg_yoy_91:.2f}", f"{non_smelter_std_91:.2f}", f"{n_ns_91}"],
    ]
    kab_rows_91 = [[str(i + 1), kab] for i, kab in enumerate(smelter_kabs_91)]

    mermaid_str_9_1 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Demografi Master Fase 4<br/><i>kabupaten, tahun, penduduk, laju YoY, flag smelter</i>"]
    end
    subgraph Proxy_Processing["2. Proxy Migrasi Time-Series Populasi"]
        A --> B["Segmentasi kabupaten<br/>Industri Ekstraktif vs Non-Ekstraktif"]
        B --> C["Window observasi tahun ≤ 2024<br/>laju pertumbuhan YoY per kabupaten-tahun"]
        C --> D["Anatomi Boxplot (kuantil Hazen)<br/>Median, Q1/Q3, IQR, fences, outliers"]
    end
    subgraph Visual_Output["3. Boxplot Komparatif"]
        D --> E["Boxplot sebaran YoY dua kategori<br/>+ tabel metrik anatomi & mean-varians"]
        E --> F["Pembacaan fenomena Boom and Bust demografis"]
    end"""
    mermaid_png_path_9_1 = str(tool_dir / "mermaid_flowchart_9_1.png")
    download_success_9_1 = download_mermaid_png(mermaid_str_9_1, mermaid_png_path_9_1)

    print("[1.5/4] Mengekstraksi dataset empiris Bab 9 sub-bab 9.2...")
    density_92 = df_demo[df_demo["tahun"] <= 2024].copy()
    density_92["Kategori"] = density_92["is_smelter"].map(
        {True: "Kabupaten Industri Ekstraktif", False: "Kabupaten Non-Ekstraktif"}
    )
    density_agg_92 = density_92.groupby(["tahun", "Kategori"], as_index=False)["kepadatan_per_km2"].mean()

    latest_smelter_density_92 = float(latest_demo_91[latest_demo_91["is_smelter"] == True]["kepadatan_per_km2"].mean())
    latest_non_smelter_density_92 = float(latest_demo_91[latest_demo_91["is_smelter"] == False]["kepadatan_per_km2"].mean())
    density_ratio_92 = latest_smelter_density_92 / latest_non_smelter_density_92 if latest_non_smelter_density_92 else 0

    pivot_density_92 = density_agg_92.pivot(index="tahun", columns="Kategori", values="kepadatan_per_km2")
    density_rows_92 = []
    for tahun, row in pivot_density_92.iterrows():
        s_val_92 = row.get("Kabupaten Industri Ekstraktif", float("nan"))
        ns_val_92 = row.get("Kabupaten Non-Ekstraktif", float("nan"))
        if pd.notna(s_val_92) and pd.notna(ns_val_92) and ns_val_92:
            rasio_str_92 = f"{s_val_92 / ns_val_92:.2f}x"
        else:
            rasio_str_92 = "-"
        density_rows_92.append([
            str(int(tahun)),
            f"{s_val_92:,.1f}" if pd.notna(s_val_92) else "-",
            f"{ns_val_92:,.1f}" if pd.notna(ns_val_92) else "-",
            rasio_str_92,
        ])
    kol_s_92 = pivot_density_92["Kabupaten Industri Ekstraktif"].dropna()
    tahun_awal_92 = int(kol_s_92.index.min())
    s_awal_92 = float(kol_s_92.loc[tahun_awal_92])
    ns_awal_92 = float(pivot_density_92.loc[tahun_awal_92, "Kabupaten Non-Ekstraktif"])
    delta_s_92 = latest_smelter_density_92 - s_awal_92
    delta_ns_92 = latest_non_smelter_density_92 - ns_awal_92
    lipat_s_92 = latest_smelter_density_92 / s_awal_92 if s_awal_92 else 0
    lipat_ns_92 = latest_non_smelter_density_92 / ns_awal_92 if ns_awal_92 else 0

    mermaid_str_9_2 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Demografi Master Fase 4<br/><i>kabupaten, tahun, kepadatan per km2, flag smelter</i>"]
    end
    subgraph Density_Processing["2. Comparative Density Analysis"]
        A --> B["Window observasi tahun ≤ 2024"]
        B --> C["Segmentasi kategori<br/>Industri Ekstraktif vs Non-Ekstraktif"]
        C --> D["Rata-rata kepadatan per kategori per tahun"]
        D --> E["Rasio kepadatan ekstraktif / non-ekstraktif"]
    end
    subgraph Visual_Output["3. Area Chart Komparatif"]
        E --> F["Area Chart tren rata-rata kepadatan dua kategori"]
        F --> G["Pembacaan intensifikasi ruang & kebutuhan layanan publik"]
    end"""
    mermaid_png_path_9_2 = str(tool_dir / "mermaid_flowchart_9_2.png")
    download_success_9_2 = download_mermaid_png(mermaid_str_9_2, mermaid_png_path_9_2)

    print("[1.8/4] Mengekstraksi dataset empiris Bab 9 sub-bab 9.3...")
    df_shift_93 = pd.read_csv(data_dir / "sulawesi_employment_shift_fase4.csv")
    df_shift_93["tahun"] = pd.to_numeric(df_shift_93["tahun"], errors="coerce")
    PROPORSI_PERIKANAN_93 = 0.22

    sulteng_shift_93 = df_shift_93[df_shift_93["provinsi"] == "Sulawesi Tengah"].sort_values("tahun")
    sulteng_first_93 = sulteng_shift_93.iloc[0]
    sulteng_last_93 = sulteng_shift_93.iloc[-1]
    pertanian_awal_93 = float(sulteng_first_93["pct_pdrb_pertanian_A"])
    pertanian_akhir_93 = float(sulteng_last_93["pct_pdrb_pertanian_A"])
    industri_awal_93 = float(sulteng_first_93["pct_industri_tambang_BC"])
    industri_akhir_93 = float(sulteng_last_93["pct_industri_tambang_BC"])
    shift_awal_93 = float(sulteng_first_93["agriculture_to_industry_shift_index"])
    shift_akhir_93 = float(sulteng_last_93["agriculture_to_industry_shift_index"])
    shift_multiplier_93 = shift_akhir_93 / shift_awal_93 if shift_awal_93 else 0
    tahun_awal_93 = int(sulteng_first_93["tahun"])
    tahun_akhir_93 = int(sulteng_last_93["tahun"])

    pivot_index_93 = df_shift_93.pivot_table(index="tahun", columns="provinsi", values="agriculture_to_industry_shift_index", aggfunc="mean")
    prov_cols_93 = sorted(pivot_index_93.columns)
    index_rows_93 = []
    for tahun, row in pivot_index_93.iterrows():
        index_rows_93.append([str(int(tahun))] + [f"{row[c]:.3f}" if pd.notna(row[c]) else "-" for c in prov_cols_93])

    ringkas_rows_93 = []
    prov_lampaui_93 = []
    for prov in prov_cols_93:
        seri_93 = df_shift_93[df_shift_93["provinsi"] == prov].sort_values("tahun")
        idx_awal_93 = float(seri_93.iloc[0]["agriculture_to_industry_shift_index"])
        idx_akhir_93 = float(seri_93.iloc[-1]["agriculture_to_industry_shift_index"])
        mult_93 = idx_akhir_93 / idx_awal_93 if idx_awal_93 else 0
        status_93 = "MELAMPAUI AMBANG (B+C > A)" if idx_akhir_93 > 1 else "Di bawah ambang"
        if idx_akhir_93 > 1:
            prov_lampaui_93.append(f"{prov} ({idx_akhir_93:.3f})")
        ringkas_rows_93.append([
            prov,
            f"{idx_awal_93:.3f}",
            f"{idx_akhir_93:.3f}",
            f"{mult_93:.1f}x",
            status_93,
        ])

    mermaid_str_9_3 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Employment Shift Fase 4<br/><i>provinsi, tahun, pct PDRB sektor A/B/C, shift index</i>"]
    end
    subgraph Shift_Processing["2. PDRB Sector Shift Index"]
        A --> B["Blok agraris: Sektor A<br/>dekomposisi Perikanan Tangkap ±22% (estimasi)"]
        A --> C["Blok ekstraktif-industrial: B+C<br/>pertambangan + industri pengolahan"]
        B --> D["Shift Index = (B+C) / A per provinsi-tahun"]
        C --> D
    end
    subgraph Visual_Output["3. Area Chart & Line Index"]
        D --> E["Area Chart komposisi sektor kunci per provinsi"]
        D --> F["Line Chart shift index 6 provinsi<br/>+ garis ambang 1 (B+C melampaui Pertanian)"]
    end
    E --> G["Pembacaan pergeseran pusat gravitasi ekonomi daerah"]
    F --> G"""
    mermaid_png_path_9_3 = str(tool_dir / "mermaid_flowchart_9_3.png")
    download_success_9_3 = download_mermaid_png(mermaid_str_9_3, mermaid_png_path_9_3)

    print("[2/4] Membangun DOCX Metodologi_Bab9_Demografi_Sosial.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)
    add_h1(doc, "BAB IX: METODOLOGI ANALISIS DEMOGRAFI SOSIAL - KETIKA HILIRISASI MENGUBAH STRUKTUR MASYARAKAT")
    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada ", False, False),
        ("Bab 9: Demografi Sosial - Ketika Hilirisasi Mengubah Struktur Masyarakat", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.", False, False),
    ])

    add_h2(doc, "9.1 Tekanan Demografi di Kabupaten Industri Ekstraktif")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Demografi: data/processed/sulawesi_demografi_master_fase4.csv (BPS SIMDASI, klasifikasi Fase 4). Visualisasi dashboard menampilkan Boxplot komparatif sebaran laju pertumbuhan penduduk YoY (kuantil Hazen, semua titik data ditampilkan) antara kabupaten industri ekstraktif dan non-ekstraktif, beserta Tabel Rincian Metrik Anatomi Boxplot dan Tabel Rincian Perhitungan Mean & Varians.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Analisis ini membaca tekanan demografi melalui perubahan jumlah penduduk kabupaten, bukan melalui data migrasi langsung. Dengan pendekatan ini, populasi diperlakukan sebagai sinyal awal: ketika kawasan smelter tumbuh lebih cepat dibanding pola umum wilayah sekitar, maka terdapat indikasi tarikan penduduk, pekerja, dan aktivitas ekonomi baru yang perlu diuji lebih lanjut. ", False, False),
        (f"Fokus pembacaan ditempatkan pada {n_smelter_kab_91} kabupaten prioritas smelter, yaitu {', '.join(smelter_kabs_91)}. Dalam window data yang tersedia, rata-rata (mean) pertumbuhan YoY kabupaten smelter tercatat {smelter_avg_yoy_91:.2f}%, sedangkan wilayah non-smelter berada di sekitar {non_smelter_avg_yoy_91:.2f}%. Pada tahun {latest_year_91}, total populasi kabupaten smelter mencapai {smelter_total_pop_latest_91 / 1000:,.2f} juta jiwa. ", False, False),
        ("Angka-angka ini tidak cukup untuk menyebut asal migran atau arah mobilitas penduduk, tetapi cukup kuat untuk menunjukkan bahwa hilirisasi nikel menciptakan tekanan demografis yang harus dibaca sebagai bagian dari beban sosial, bukan sekadar konsekuensi administratif pembangunan industri.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Proxy Migrasi dari Time-Series Populasi Kabupaten")
    add_p(doc, [
        ("Kerangka pembacaan tekanan demografi berbasis proxy populasi dan anatomi boxplot diilustrasikan pada ", False, False),
        ("Bagan Alur 9.1", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan statistika deskriptif sebaran (kuantil Hazen) komparatif dua kategori wilayah.", False, False),
    ])
    add_caption(doc, "Bagan Alur 9.1: Alur Logika Analisis Proxy Migrasi & Anatomi Boxplot Demografi")
    if download_success_9_1:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_9_1, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Laju Pertumbuhan, Kuantil Hazen, dan Batas Kewajaran")
    add_p(doc, [("Kuantifikasi sebaran tekanan demografi dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Laju Pertumbuhan Penduduk Year-on-Year", "Laju_YoY_k,t (%) = ( P_k,t - P_k,t-1 ) / P_k,t-1 × 100", [
        ("Laju_YoY_k,t (%)", "Persentase perubahan jumlah penduduk kabupaten k dari tahun t-1 ke tahun t (unit observasi: kabupaten-tahun, window ≤ 2024)."),
        ("P_k,t", "Jumlah penduduk kabupaten k pada tahun t (ribu jiwa, BPS SIMDASI)."),
    ])
    add_formula(doc, "Persamaan Kuartil Metode Hazen dan Rentang Interkuartil", "Q_p = Kuantil_Hazen ( Laju_YoY , p )   ;   IQR = Q3 - Q1", [
        ("Q_p", "Kuartil ke-p (Q1 = 0,25; Median = 0,50; Q3 = 0,75) dihitung dengan metode Hazen — algoritma yang identik dengan boxplot default Plotly pada dashboard."),
        ("IQR", "Interquartile Range: ketebalan kotak utama boxplot yang memuat 50% data inti."),
    ])
    add_formula(doc, "Persamaan Batas Kewajaran Data (Fences) dan Pencilan", "Batas_Bawah = Q1 - 1,5 × IQR   ;   Batas_Atas = Q3 + 1,5 × IQR", [
        ("Lower/Upper Fence", "Titik data terdekat yang masih berada di dalam batas kewajaran; titik di luar batas diklasifikasikan sebagai pencilan (outlier)."),
    ])

    add_p(doc, [("Substitusi angka dari dataset aktual ke dalam rumus anatomi sebaran adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Median & IQR Kabupaten Ekstraktif", f"Median = {s_91['med']:.2f}%   ;   Q1 = {s_91['q1']:.2f}%   ;   Q3 = {s_91['q3']:.2f}%   ;   IQR = {s_91['q3'] - s_91['q1']:.2f}")
    add_formula(doc, "Substitusi Median & IQR Kabupaten Non-Ekstraktif", f"Median = {ns_91['med']:.2f}%   ;   Q1 = {ns_91['q1']:.3f}%   ;   Q3 = {ns_91['q3']:.3f}%   ;   IQR = {ns_91['q3'] - ns_91['q1']:.3f}")
    add_formula(doc, "Substitusi Fences & Titik Ekstrem Kabupaten Ekstraktif", f"Lower Fence = {s_91['lower']:.2f}%   ;   Upper Fence = {s_91['upper']:.2f}%   ;   Min = {s_91['min']:.2f}%   ;   Max = {s_91['max']:.2f}%")
    add_formula(doc, "Substitusi Rata-rata YoY Dua Kategori", f"Mean_Ekstraktif = {smelter_avg_yoy_91:.2f}% (N={n_s_91})   vs   Mean_Non-Ekstraktif = {non_smelter_avg_yoy_91:.2f}% (N={n_ns_91})")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Anatomi Boxplot dan Rincian Mean-Varians")
    add_p(doc, [
        ("Rincian metrik anatomi boxplot kedua kategori wilayah disajikan pada ", False, False),
        ("Tabel 9.1", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 9.1: Rincian Metrik Anatomi Boxplot Laju Pertumbuhan Penduduk (YoY %)")
    add_table_1col(doc, ["Kategori", "Max (%)", "Upper Fence (%)", "Q3 (%)", "Median (%)", "Q1 (%)", "Lower Fence (%)", "Min (%)"], box_rows_91, [4.2, 1.7, 2.0, 1.6, 1.8, 1.6, 2.0, 1.7], ["L", "C", "C", "C", "C", "C", "C", "C"])

    add_p(doc, [
        ("Rincian perhitungan mean, deviasi standar, dan jumlah sampel disajikan pada ", False, False),
        ("Tabel 9.2", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 9.2: Rincian Perhitungan Mean & Varians Laju Pertumbuhan YoY")
    add_table_1col(doc, ["Kategori", "Rata-Rata / Mean (%)", "Standard Deviation", "Jumlah Sampel (Tahun-Kabupaten)"], mean_rows_91, [4.6, 3.6, 3.4, 4.4], ["L", "C", "C", "C"])

    add_p(doc, [
        (f"Daftar {n_smelter_kab_91} kabupaten prioritas industri ekstraktif yang menjadi fokus pembacaan disajikan pada ", False, False),
        ("Tabel 9.3", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 9.3: Daftar Kabupaten Prioritas Industri Ekstraktif (Klasifikasi Fase 4)")
    add_table_1col(doc, ["No", "Kabupaten"], kab_rows_91, [2.0, 10.0], ["C", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris: Bukti Matematis Fenomena Boom and Bust")
    add_p(doc, [
        ("1. ", True, False), ("Median Konsisten Lebih Tinggi: ", True, False),
        (f"Nilai tengah pertumbuhan di kawasan industri ekstraktif mencapai {s_91['med']:.2f}%, konsisten lebih tinggi dibandingkan kawasan non-ekstraktif di angka {ns_91['med']:.2f}%.\n", False, False),
        ("2. ", True, False), ("Variabilitas Ekstrem Kawasan Ekstraktif: ", True, False),
        (f"IQR wilayah ekstraktif merentang dari Q1 {s_91['q1']:.2f}% hingga Q3 {s_91['q3']:.2f}% — jauh lebih lebar dibanding non-ekstraktif (Q1 {ns_91['q1']:.3f}% hingga Q3 {ns_91['q3']:.3f}%) yang tumbuh lebih stabil.\n", False, False),
        ("3. ", True, False), ("Bukti Boom and Bust: ", True, False),
        (f"Sebaran data ekstraktif menembus batas kewajaran (fences {s_91['lower']:.2f}% s.d. {s_91['upper']:.2f}%): titik lonjakan tertinggi menyentuh {s_91['max']:.2f}% sementara titik terendah anjlok hingga {s_91['min']:.2f}%. Lonjakan dan kejatuhan tajam ini menjadi bukti matematis fenomena Boom and Bust — masuknya pekerja migran secara masif di awal fase konstruksi pabrik, disusul eksodus drastis ketika proyek operasional menyusut atau terjadi pemutusan kerja massal.", False, False),
    ])

    add_h2(doc, "9.2 Intensifikasi Ruang: Kepadatan Industri Ekstraktif vs Non-Ekstraktif")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Demografi: data/processed/sulawesi_demografi_master_fase4.csv (BPS SIMDASI). Visualisasi dashboard menampilkan Area Chart tren rata-rata kepadatan penduduk (jiwa/km2) dua kategori wilayah beserta tabel agregasi kepadatan per kategori.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Sub-bab ini tidak mengklaim perubahan resmi desa menjadi kota karena data klasifikasi Podes belum menjadi basis utama di halaman ini. Yang dibaca adalah intensifikasi ruang, yaitu tekanan yang muncul ketika pertumbuhan penduduk dan konsentrasi industri bertemu pada wilayah yang sama. ", False, False),
        (f"Rata-rata kepadatan kabupaten smelter pada {latest_year_91} mencapai {latest_smelter_density_92:.1f} jiwa/km2, sedangkan kabupaten non-smelter berada pada {latest_non_smelter_density_92:.1f} jiwa/km2. Rasio smelter terhadap non-smelter sebesar {density_ratio_92:.2f} kali memberi sinyal bahwa kawasan industri membutuhkan kapasitas layanan publik yang berbeda: perumahan, air bersih, sanitasi, transportasi, hingga fasilitas kesehatan. ", False, False),
        ("Dalam kerangka D3TLH, kepadatan bukan sekadar angka demografi, melainkan indikator apakah ruang hidup lokal sedang dipadatkan oleh proyek ekstraktif tanpa perencanaan sosial yang sepadan — dibaca sebagai peta awal tekanan ruang, bukan sebagai klaim urbanisasi formal.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Comparative Density Analysis")
    add_p(doc, [
        ("Kerangka komparasi kepadatan dua kategori wilayah diilustrasikan pada ", False, False),
        ("Bagan Alur 9.2", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan perbandingan rata-rata kepadatan deskriptif antar kategori dari waktu ke waktu.", False, False),
    ])
    add_caption(doc, "Bagan Alur 9.2: Alur Logika Analisis Comparative Density Analysis")
    if download_success_9_2:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_9_2, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Rata-rata Kepadatan dan Rasio Intensifikasi")
    add_p(doc, [("Kuantifikasi intensifikasi ruang dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Rata-rata Kepadatan per Kategori", "D_c,t = [ Σ ( Densitas_k,t ) ] / n_c   ;   untuk seluruh kabupaten k dalam kategori c pada tahun t", [
        ("D_c,t", "Rata-rata kepadatan penduduk (jiwa/km2) kategori c (Industri Ekstraktif / Non-Ekstraktif) pada tahun t (window ≤ 2024)."),
        ("Densitas_k,t", "Kepadatan penduduk kabupaten k pada tahun t (jiwa/km2, BPS SIMDASI)."),
        ("n_c", "Jumlah kabupaten dalam kategori c."),
    ])
    add_formula(doc, "Persamaan Rasio Intensifikasi Ruang", "R_t = D_Ekstraktif,t / D_Non-Ekstraktif,t", [
        ("R_t", "Rasio rata-rata kepadatan kawasan industri ekstraktif terhadap non-ekstraktif pada tahun t; sinyal perbedaan kebutuhan kapasitas layanan publik."),
    ])

    add_p(doc, [("Substitusi angka dari dataset aktual ke dalam rumus intensifikasi ruang adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Rata-rata Kepadatan Tahun Terbaru", f"D_Ekstraktif,{latest_year_91} = {latest_smelter_density_92:.1f} jiwa/km2   ;   D_Non-Ekstraktif,{latest_year_91} = {latest_non_smelter_density_92:.1f} jiwa/km2")
    add_formula(doc, "Substitusi Rasio Intensifikasi", f"R_{latest_year_91} = {latest_smelter_density_92:.1f} / {latest_non_smelter_density_92:.1f} = {density_ratio_92:.2f}x")
    add_formula(doc, "Substitusi Perubahan Sepanjang Window", f"Δ D_Ekstraktif ({tahun_awal_92}-{latest_year_91}) = {delta_s_92:+.1f} jiwa/km2   ;   Δ D_Non-Ekstraktif = {delta_ns_92:+.1f} jiwa/km2")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Agregasi Kepadatan per Kategori")
    add_p(doc, [
        (f"Agregasi rata-rata kepadatan kedua kategori per tahun ({tahun_awal_92}-{latest_year_91}) beserta rasionya disajikan pada ", False, False),
        ("Tabel 9.4", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 9.4: Agregasi Rata-rata Kepadatan Penduduk per Kategori Wilayah ({tahun_awal_92}-{latest_year_91})")
    add_table_1col(doc, ["Tahun", "Industri Ekstraktif (jiwa/km2)", "Non-Ekstraktif (jiwa/km2)", "Rasio (x)"], density_rows_92, [2.2, 4.4, 4.2, 2.6], ["C", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Peta Awal Tekanan Ruang")
    add_p(doc, [
        ("1. ", True, False), ("Profil Kepadatan Dua Kategori: ", True, False),
        (f"Pada {latest_year_91}, rata-rata kepadatan kabupaten industri ekstraktif {latest_smelter_density_92:.1f} jiwa/km2 — rasio {density_ratio_92:.2f} kali terhadap kabupaten non-ekstraktif ({latest_non_smelter_density_92:.1f} jiwa/km2). Kawasan ekstraktif berbasis kabupaten berwilayah luas dan semula berpenduduk jarang, sehingga rasionya di bawah satu.\n", False, False),
        ("2. ", True, False), ("Intensifikasi Jauh Lebih Cepat di Kawasan Ekstraktif: ", True, False),
        (f"Sepanjang window data {tahun_awal_92}-{latest_year_91}, rata-rata kepadatan kawasan ekstraktif melipat {lipat_s_92:.1f} kali (dari {s_awal_92:.1f} menjadi {latest_smelter_density_92:.1f} jiwa/km2, {delta_s_92:+.1f}), jauh melampaui laju kawasan non-ekstraktif yang hanya {lipat_ns_92:.1f} kali — inilah intensifikasi ruang yang dibaca sub-bab ini.\n", False, False),
        ("3. ", True, False), ("Implikasi Kapasitas Layanan Publik: ", True, False),
        ("Pemadatan cepat pada ruang yang semula lengang memberi sinyal kebutuhan kapasitas layanan publik yang berbeda di kawasan industri — perumahan, air bersih, sanitasi, transportasi, hingga fasilitas kesehatan — agar ruang hidup lokal tidak dipadatkan proyek ekstraktif tanpa perencanaan sosial yang sepadan.", False, False),
    ])

    add_h2(doc, "9.3 Pergeseran Ekonomi Agraris ke Tambang dan Industri")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Shift Index: data/processed/sulawesi_employment_shift_fase4.csv dan data/processed/sulawesi_pdrb_sektoral_2016_2024.csv (BPS SIMDASI). Visualisasi dashboard menampilkan Area Chart komposisi PDRB sektor kunci per provinsi (dengan dekomposisi estimasi Perikanan Tangkap) serta Line Chart Shift Index (B+C / A) 6 provinsi dengan garis ambang 1.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Pergeseran pekerjaan tidak dapat diklaim hanya dari PDRB, tetapi struktur PDRB memberi petunjuk kuat tentang arah ekonomi yang sedang dibentuk. Sektor A dibaca sebagai basis agraris, sementara sektor B dan C dibaca sebagai blok ekstraktif-industrial: pertambangan dan industri pengolahan. Rasio B+C terhadap A menjadi shift index; nilai di atas 1 berarti kontribusi tambang dan industri sudah melampaui pertanian. ", False, False),
        (f"Di Sulawesi Tengah, porsi pertanian turun dari {pertanian_awal_93:.2f}% menjadi {pertanian_akhir_93:.2f}%, sementara tambang+industri naik dari {industri_awal_93:.2f}% menjadi {industri_akhir_93:.2f}%. Indeksnya naik dari {shift_awal_93:.3f} ke {shift_akhir_93:.3f}, atau sekitar {shift_multiplier_93:.1f} kali. ", False, False),
        ("Dengan kata lain, data sektoral menunjukkan bahwa hilirisasi tidak hanya menambah pabrik; ia mengubah pusat gravitasi ekonomi daerah, dari ruang produksi agraris menuju rantai ekstraktif yang lebih terkonsentrasi pada modal besar.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis PDRB Sector Shift Index (B+C / A)")
    add_p(doc, [
        ("Kerangka pembacaan pergeseran struktur ekonomi berbasis shift index diilustrasikan pada ", False, False),
        ("Bagan Alur 9.3", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan indeks rasio sektoral deskriptif dengan ambang interpretatif 1.", False, False),
    ])
    add_caption(doc, "Bagan Alur 9.3: Alur Logika Analisis PDRB Sector Shift Index")
    if download_success_9_3:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_9_3, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Shift Index, Dekomposisi Sektor A, dan Multiplier")
    add_p(doc, [("Kuantifikasi pergeseran struktur ekonomi dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan PDRB Sector Shift Index", "Shift_Index_p,t = ( PDRB_B_p,t + PDRB_C_p,t ) / PDRB_A_p,t", [
        ("Shift_Index_p,t", "Rasio blok ekstraktif-industrial terhadap basis agraris provinsi p tahun t; nilai > 1 berarti kontribusi tambang dan industri melampaui pertanian (garis ambang pada grafik)."),
        ("PDRB_A/B/C", "Persentase kontribusi PDRB Sektor A (Pertanian, Kehutanan, Perikanan), B (Pertambangan), dan C (Industri Pengolahan) menurut KBLI."),
    ])
    add_formula(doc, "Persamaan Dekomposisi Estimasi Sektor A", "Perikanan_Tangkap ≈ 0,22 × Sektor_A   ;   Pertanian_Kehutanan ≈ 0,78 × Sektor_A", [
        ("0,22", "Proporsi estimasi Perikanan Tangkap terhadap Sektor A, mengacu pada rata-rata proporsi sub-sektor perikanan di provinsi-provinsi pesisir Sulawesi (Statistik Perikanan BPS Sulawesi, 2016-2024) — konstanta metodologis untuk dekomposisi visual Area Chart."),
    ])
    add_formula(doc, "Persamaan Multiplier Pergeseran", "Multiplier_p = Shift_Index_p,akhir / Shift_Index_p,awal", [
        ("Multiplier_p", "Kelipatan kenaikan shift index provinsi p dari tahun awal ke tahun akhir window data."),
    ])

    add_p(doc, [("Substitusi angka dari dataset aktual (Sulawesi Tengah sebagai episentrum) adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Shift Index Sulawesi Tengah", f"Shift_Index_Sulteng: {shift_awal_93:.3f} ({tahun_awal_93}) → {shift_akhir_93:.3f} ({tahun_akhir_93})   ;   Multiplier = {shift_akhir_93:.3f} / {shift_awal_93:.3f} = {shift_multiplier_93:.1f}x")
    add_formula(doc, "Substitusi Komposisi Sektoral Sulawesi Tengah", f"Pertanian (A): {pertanian_awal_93:.2f}% → {pertanian_akhir_93:.2f}%   ;   Tambang+Industri (B+C): {industri_awal_93:.2f}% → {industri_akhir_93:.2f}%")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Shift Index per Provinsi dan Ringkasan Pergeseran")
    add_p(doc, [
        (f"Nilai shift index seluruh provinsi per tahun ({tahun_awal_93}-{tahun_akhir_93}) disajikan pada ", False, False),
        ("Tabel 9.5", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 9.5: Shift Index (B+C / A) per Provinsi per Tahun ({tahun_awal_93}-{tahun_akhir_93})")
    add_table_1col(doc, ["Tahun"] + prov_cols_93, index_rows_93, [1.6] + [2.4] * len(prov_cols_93), ["C"] * (len(prov_cols_93) + 1))

    add_p(doc, [
        ("Ringkasan pergeseran tiap provinsi beserta status ambangnya disajikan pada ", False, False),
        ("Tabel 9.6", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 9.6: Ringkasan Pergeseran Struktur Ekonomi per Provinsi")
    add_table_1col(doc, ["Provinsi", "Index Awal", "Index Akhir", "Multiplier", "Status Ambang"], ringkas_rows_93, [3.4, 2.2, 2.2, 2.2, 4.6], ["L", "C", "C", "C", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris: Pergeseran Pusat Gravitasi Ekonomi")
    add_p(doc, [
        ("1. ", True, False), ("Episentrum Pergeseran di Sulawesi Tengah: ", True, False),
        (f"Shift index Sulteng melonjak dari {shift_awal_93:.3f} menjadi {shift_akhir_93:.3f} ({shift_multiplier_93:.1f} kali) — blok tambang+industri kini {industri_akhir_93:.2f}% PDRB, jauh melampaui pertanian yang menyusut ke {pertanian_akhir_93:.2f}%.\n", False, False),
        ("2. ", True, False), ("Provinsi Pelampau Ambang: ", True, False),
        (f"Provinsi dengan shift index akhir melampaui ambang 1 (B+C > A): {', '.join(prov_lampaui_93) if prov_lampaui_93 else 'tidak ada'} — sementara provinsi lain masih berbasis agraris.\n", False, False),
        ("3. ", True, False), ("Catatan Metodologis: ", True, False),
        ("Dekomposisi Perikanan Tangkap adalah estimasi (±22% Sektor A) untuk keperluan visual; klaim pergeseran pekerjaan tidak ditarik dari PDRB semata, melainkan dibaca sebagai petunjuk arah ekonomi yang sedang dibentuk hilirisasi — dari ruang produksi agraris menuju rantai ekstraktif yang terkonsentrasi pada modal besar.", False, False),
    ])

    docx_path = tool_dir / "Metodologi_Bab9_Demografi_Sosial.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 9...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 9 - Demografi Sosial</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: Arial, sans-serif; max-width: 920px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #43A047; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; }}
.hdr-title {{ color: #81C784; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #2E7D32; padding-bottom: 8px; }}
h2 {{ color: #81C784; text-transform: uppercase; border-bottom: 1px solid #2E7D32; }}
h4 {{ color: #A5D6A7; }}
.note-box {{ background: #132213; border-left: 4px solid #2E7D32; padding: 10px 14px; margin: 12px 0; }}
.formula {{ background: #0D1B0E; border: 1px solid #2E7D32; color: #A5D6A7; padding: 8px 12px; font-family: monospace; }}
.table-caption {{ color: #A5D6A7; font-weight: 700; font-style: italic; margin-top: 14px; }}
.data-th {{ background: #1B5E20; color: white; padding: 6px; text-align: left; border: 1px solid #2E7D32; }}
.data-td {{ padding: 6px; border: 1px solid #243524; vertical-align: top; }}
.data-tr-even .data-td {{ background: #131B13; }}
.mermaid {{ background: #0D1610; border: 1px solid #2E7D32; padding: 12px; margin: 10px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB IX: Metodologi Analisis Demografi Sosial - Ketika Hilirisasi Mengubah Struktur Masyarakat</div>
<p>Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada <strong>Bab 9: Demografi Sosial</strong>.</p>

<h2>9.1 Tekanan Demografi di Kabupaten Industri Ekstraktif</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Demografi: <code>data/processed/sulawesi_demografi_master_fase4.csv</code> (BPS SIMDASI, klasifikasi Fase 4). Visualisasi dashboard menampilkan Boxplot komparatif sebaran laju pertumbuhan penduduk YoY (kuantil Hazen) antara kabupaten industri ekstraktif dan non-ekstraktif, beserta Tabel Rincian Metrik Anatomi Boxplot dan Tabel Rincian Perhitungan Mean & Varians.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Analisis ini membaca tekanan demografi melalui perubahan jumlah penduduk kabupaten (proxy migrasi), bukan data migrasi langsung. Fokus pada <strong>{n_smelter_kab_91} kabupaten prioritas smelter</strong>: {", ".join(smelter_kabs_91)}. Rata-rata pertumbuhan YoY kabupaten smelter <strong>{smelter_avg_yoy_91:.2f}%</strong> vs non-smelter <strong>{non_smelter_avg_yoy_91:.2f}%</strong>; total populasi kabupaten smelter {latest_year_91} mencapai <strong>{smelter_total_pop_latest_91 / 1000:,.2f} juta jiwa</strong>. Hilirisasi nikel menciptakan tekanan demografis yang harus dibaca sebagai bagian dari beban sosial.</p>
<h4>B. Alur Logika Metodologis Proxy Migrasi dari Time-Series Populasi Kabupaten</h4>
<p>Kerangka pembacaan diilustrasikan pada <strong>Bagan Alur 9.1</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan statistika deskriptif sebaran (kuantil Hazen) komparatif dua kategori wilayah.</p>
<div class="table-caption">Bagan Alur 9.1: Alur Logika Analisis Proxy Migrasi & Anatomi Boxplot Demografi</div>
<div class="mermaid">{mermaid_str_9_1}</div>
<h4>C. Formulasi Matematis: Laju Pertumbuhan, Kuantil Hazen, dan Batas Kewajaran</h4>
<div class="formula">Laju_YoY_k,t (%) = ( P_k,t - P_k,t-1 ) / P_k,t-1 × 100</div>
<div class="formula">Q_p = Kuantil_Hazen ( Laju_YoY , p )   ;   IQR = Q3 - Q1</div>
<div class="formula">Batas_Bawah = Q1 - 1,5 × IQR   ;   Batas_Atas = Q3 + 1,5 × IQR</div>
<p>Substitusi angka dari dataset aktual:</p>
<div class="formula">Ekstraktif: Median = {s_91['med']:.2f}% ; Q1 = {s_91['q1']:.2f}% ; Q3 = {s_91['q3']:.2f}% ; IQR = {s_91['q3'] - s_91['q1']:.2f}</div>
<div class="formula">Non-Ekstraktif: Median = {ns_91['med']:.2f}% ; Q1 = {ns_91['q1']:.3f}% ; Q3 = {ns_91['q3']:.3f}% ; IQR = {ns_91['q3'] - ns_91['q1']:.3f}</div>
<div class="formula">Ekstraktif Fences: {s_91['lower']:.2f}% s.d. {s_91['upper']:.2f}%   ;   Min = {s_91['min']:.2f}% ; Max = {s_91['max']:.2f}%</div>
<div class="formula">Mean_Ekstraktif = {smelter_avg_yoy_91:.2f}% (N={n_s_91})   vs   Mean_Non-Ekstraktif = {non_smelter_avg_yoy_91:.2f}% (N={n_ns_91})</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 9.1: Rincian Metrik Anatomi Boxplot Laju Pertumbuhan Penduduk (YoY %)</div>
{html_table(["Kategori", "Max (%)", "Upper Fence (%)", "Q3 (%)", "Median (%)", "Q1 (%)", "Lower Fence (%)", "Min (%)"], box_rows_91)}
<div class="table-caption">Tabel 9.2: Rincian Perhitungan Mean & Varians Laju Pertumbuhan YoY</div>
{html_table(["Kategori", "Rata-Rata / Mean (%)", "Standard Deviation", "Jumlah Sampel (Tahun-Kabupaten)"], mean_rows_91)}
<div class="table-caption">Tabel 9.3: Daftar Kabupaten Prioritas Industri Ekstraktif (Klasifikasi Fase 4)</div>
{html_table(["No", "Kabupaten"], kab_rows_91)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Median Konsisten Lebih Tinggi:</strong> kawasan ekstraktif {s_91['med']:.2f}% vs non-ekstraktif {ns_91['med']:.2f}%. <strong>2. Variabilitas Ekstrem:</strong> IQR ekstraktif ({s_91['q1']:.2f}% - {s_91['q3']:.2f}%) jauh lebih lebar dibanding non-ekstraktif ({ns_91['q1']:.3f}% - {ns_91['q3']:.3f}%). <strong>3. Bukti Boom and Bust:</strong> sebaran menembus fences ({s_91['lower']:.2f}% s.d. {s_91['upper']:.2f}%) dengan Max {s_91['max']:.2f}% dan Min {s_91['min']:.2f}% — bukti matematis masuknya pekerja migran masif di fase konstruksi lalu eksodus drastis saat operasional menyusut.</p>

<h2>9.2 Intensifikasi Ruang: Kepadatan Industri Ekstraktif vs Non-Ekstraktif</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Demografi: <code>data/processed/sulawesi_demografi_master_fase4.csv</code> (BPS SIMDASI). Visualisasi dashboard menampilkan Area Chart tren rata-rata kepadatan penduduk (jiwa/km2) dua kategori wilayah beserta tabel agregasi kepadatan per kategori.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Yang dibaca sub-bab ini adalah <strong>intensifikasi ruang</strong> — tekanan yang muncul ketika pertumbuhan penduduk dan konsentrasi industri bertemu pada wilayah yang sama (bukan klaim urbanisasi formal Podes). Rata-rata kepadatan kabupaten smelter pada {latest_year_91} mencapai <strong>{latest_smelter_density_92:.1f} jiwa/km2</strong> vs non-smelter <strong>{latest_non_smelter_density_92:.1f} jiwa/km2</strong> — rasio <strong>{density_ratio_92:.2f} kali</strong>, sinyal kebutuhan kapasitas layanan publik yang berbeda di kawasan industri.</p>
<h4>B. Alur Logika Metodologis Comparative Density Analysis</h4>
<p>Kerangka komparasi diilustrasikan pada <strong>Bagan Alur 9.2</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan perbandingan rata-rata kepadatan deskriptif antar kategori dari waktu ke waktu.</p>
<div class="table-caption">Bagan Alur 9.2: Alur Logika Analisis Comparative Density Analysis</div>
<div class="mermaid">{mermaid_str_9_2}</div>
<h4>C. Formulasi Matematis: Rata-rata Kepadatan dan Rasio Intensifikasi</h4>
<div class="formula">D_c,t = [ Σ ( Densitas_k,t ) ] / n_c   ;   untuk seluruh kabupaten k dalam kategori c pada tahun t</div>
<div class="formula">R_t = D_Ekstraktif,t / D_Non-Ekstraktif,t</div>
<p>Substitusi angka dari dataset aktual:</p>
<div class="formula">D_Ekstraktif,{latest_year_91} = {latest_smelter_density_92:.1f} jiwa/km2   ;   D_Non-Ekstraktif,{latest_year_91} = {latest_non_smelter_density_92:.1f} jiwa/km2</div>
<div class="formula">R_{latest_year_91} = {latest_smelter_density_92:.1f} / {latest_non_smelter_density_92:.1f} = {density_ratio_92:.2f}x</div>
<div class="formula">Δ D_Ekstraktif ({tahun_awal_92}-{latest_year_91}) = {delta_s_92:+.1f} jiwa/km2   ;   Δ D_Non-Ekstraktif = {delta_ns_92:+.1f} jiwa/km2</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 9.4: Agregasi Rata-rata Kepadatan Penduduk per Kategori Wilayah ({tahun_awal_92}-{latest_year_91})</div>
{html_table(["Tahun", "Industri Ekstraktif (jiwa/km2)", "Non-Ekstraktif (jiwa/km2)", "Rasio (x)"], density_rows_92)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Profil Kepadatan Dua Kategori:</strong> pada {latest_year_91}, kepadatan ekstraktif {latest_smelter_density_92:.1f} jiwa/km2 — rasio {density_ratio_92:.2f}x terhadap non-ekstraktif ({latest_non_smelter_density_92:.1f} jiwa/km2); kawasan ekstraktif berbasis kabupaten luas berpenduduk jarang sehingga rasionya di bawah satu. <strong>2. Intensifikasi Jauh Lebih Cepat:</strong> sepanjang {tahun_awal_92}-{latest_year_91} kepadatan ekstraktif melipat <strong>{lipat_s_92:.1f}x</strong> (dari {s_awal_92:.1f} ke {latest_smelter_density_92:.1f} jiwa/km2) vs non-ekstraktif hanya {lipat_ns_92:.1f}x — inilah intensifikasi ruang yang dibaca sub-bab ini. <strong>3. Implikasi:</strong> pemadatan cepat pada ruang yang semula lengang menuntut kapasitas layanan publik berbeda (perumahan, air bersih, sanitasi, transportasi, faskes) agar ruang hidup lokal tidak dipadatkan tanpa perencanaan sosial yang sepadan.</p>

<h2>9.3 Pergeseran Ekonomi Agraris ke Tambang dan Industri</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Shift Index: <code>data/processed/sulawesi_employment_shift_fase4.csv</code> dan <code>data/processed/sulawesi_pdrb_sektoral_2016_2024.csv</code> (BPS SIMDASI). Visualisasi dashboard menampilkan Area Chart komposisi PDRB sektor kunci per provinsi (dengan dekomposisi estimasi Perikanan Tangkap) serta Line Chart Shift Index (B+C / A) 6 provinsi dengan garis ambang 1.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Sektor A dibaca sebagai basis agraris, sektor B dan C sebagai blok ekstraktif-industrial. Rasio B+C terhadap A menjadi <strong>shift index</strong>; nilai di atas 1 berarti kontribusi tambang dan industri sudah melampaui pertanian. Di Sulawesi Tengah, porsi pertanian turun dari <strong>{pertanian_awal_93:.2f}%</strong> menjadi <strong>{pertanian_akhir_93:.2f}%</strong>, sementara tambang+industri naik dari <strong>{industri_awal_93:.2f}%</strong> menjadi <strong>{industri_akhir_93:.2f}%</strong>. Indeksnya naik dari <strong>{shift_awal_93:.3f}</strong> ke <strong>{shift_akhir_93:.3f}</strong> ({shift_multiplier_93:.1f} kali) — hilirisasi mengubah pusat gravitasi ekonomi daerah dari ruang produksi agraris menuju rantai ekstraktif yang terkonsentrasi pada modal besar.</p>
<h4>B. Alur Logika Metodologis PDRB Sector Shift Index (B+C / A)</h4>
<p>Kerangka pembacaan diilustrasikan pada <strong>Bagan Alur 9.3</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan indeks rasio sektoral deskriptif dengan ambang interpretatif 1.</p>
<div class="table-caption">Bagan Alur 9.3: Alur Logika Analisis PDRB Sector Shift Index</div>
<div class="mermaid">{mermaid_str_9_3}</div>
<h4>C. Formulasi Matematis: Shift Index, Dekomposisi Sektor A, dan Multiplier</h4>
<div class="formula">Shift_Index_p,t = ( PDRB_B_p,t + PDRB_C_p,t ) / PDRB_A_p,t</div>
<div class="formula">Perikanan_Tangkap ≈ 0,22 × Sektor_A   ;   Pertanian_Kehutanan ≈ 0,78 × Sektor_A</div>
<div class="formula">Multiplier_p = Shift_Index_p,akhir / Shift_Index_p,awal</div>
<p>Substitusi angka dari dataset aktual (Sulawesi Tengah sebagai episentrum):</p>
<div class="formula">Shift_Index_Sulteng: {shift_awal_93:.3f} ({tahun_awal_93}) → {shift_akhir_93:.3f} ({tahun_akhir_93})   ;   Multiplier = {shift_multiplier_93:.1f}x</div>
<div class="formula">Pertanian (A): {pertanian_awal_93:.2f}% → {pertanian_akhir_93:.2f}%   ;   Tambang+Industri (B+C): {industri_awal_93:.2f}% → {industri_akhir_93:.2f}%</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 9.5: Shift Index (B+C / A) per Provinsi per Tahun ({tahun_awal_93}-{tahun_akhir_93})</div>
{html_table(["Tahun"] + prov_cols_93, index_rows_93)}
<div class="table-caption">Tabel 9.6: Ringkasan Pergeseran Struktur Ekonomi per Provinsi</div>
{html_table(["Provinsi", "Index Awal", "Index Akhir", "Multiplier", "Status Ambang"], ringkas_rows_93)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Episentrum Pergeseran di Sulawesi Tengah:</strong> shift index melonjak {shift_awal_93:.3f} → {shift_akhir_93:.3f} ({shift_multiplier_93:.1f}x); blok tambang+industri kini {industri_akhir_93:.2f}% PDRB vs pertanian {pertanian_akhir_93:.2f}%. <strong>2. Provinsi Pelampau Ambang:</strong> {", ".join(prov_lampaui_93) if prov_lampaui_93 else "tidak ada"}. <strong>3. Catatan Metodologis:</strong> dekomposisi Perikanan Tangkap adalah estimasi (±22% Sektor A) untuk keperluan visual; klaim pergeseran pekerjaan tidak ditarik dari PDRB semata.</p>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab9_Demografi_Sosial.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB IX: METODOLOGI ANALISIS DEMOGRAFI SOSIAL - KETIKA HILIRISASI MENGUBAH STRUKTUR MASYARAKAT",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada **Bab 9: Demografi Sosial - Ketika Hilirisasi Mengubah Struktur Masyarakat**.",
        "",
        "## 9.1 Tekanan Demografi di Kabupaten Industri Ekstraktif",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Demografi: `data/processed/sulawesi_demografi_master_fase4.csv` (BPS SIMDASI, klasifikasi Fase 4). Visualisasi dashboard menampilkan Boxplot komparatif sebaran laju pertumbuhan penduduk YoY (kuantil Hazen, semua titik data ditampilkan) antara kabupaten industri ekstraktif dan non-ekstraktif, beserta Tabel Rincian Metrik Anatomi Boxplot dan Tabel Rincian Perhitungan Mean & Varians.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Analisis ini membaca tekanan demografi melalui perubahan jumlah penduduk kabupaten (*proxy* migrasi), bukan melalui data migrasi langsung — populasi diperlakukan sebagai sinyal awal tarikan penduduk, pekerja, dan aktivitas ekonomi baru. Fokus pembacaan pada **{n_smelter_kab_91} kabupaten prioritas smelter**: {', '.join(smelter_kabs_91)}. Rata-rata pertumbuhan YoY kabupaten smelter tercatat **{smelter_avg_yoy_91:.2f}%** vs non-smelter **{non_smelter_avg_yoy_91:.2f}%**; total populasi kabupaten smelter pada {latest_year_91} mencapai **{smelter_total_pop_latest_91 / 1000:,.2f} juta jiwa**. Hilirisasi nikel menciptakan tekanan demografis yang harus dibaca sebagai bagian dari beban sosial, bukan sekadar konsekuensi administratif pembangunan industri.",
        "",
        "#### B. Alur Logika Metodologis Proxy Migrasi dari Time-Series Populasi Kabupaten",
        "Kerangka pembacaan tekanan demografi berbasis proxy populasi dan anatomi boxplot diilustrasikan pada **Bagan Alur 9.1** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan statistika deskriptif sebaran (kuantil Hazen) komparatif dua kategori wilayah.",
        "",
        "##### Bagan Alur 9.1: Alur Logika Analisis Proxy Migrasi & Anatomi Boxplot Demografi",
        "```mermaid",
        mermaid_str_9_1,
        "```",
        "",
        "#### C. Formulasi Matematis: Laju Pertumbuhan, Kuantil Hazen, dan Batas Kewajaran",
        "Kuantifikasi sebaran tekanan demografi dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Laju_YoY_k,t (%) = ( P_k,t - P_k,t-1 ) / P_k,t-1 × 100",
        "Q_p = Kuantil_Hazen ( Laju_YoY , p )   ;   IQR = Q3 - Q1",
        "Batas_Bawah = Q1 - 1,5 × IQR   ;   Batas_Atas = Q3 + 1,5 × IQR",
        "```",
        "",
        "Metode kuantil Hazen dipakai karena identik dengan algoritma boxplot default Plotly pada dashboard, sehingga tabel statis dan tooltip grafik cocok sempurna.",
        "",
        "Substitusi angka dari dataset aktual:",
        "",
        "```text",
        f"Ekstraktif: Median = {s_91['med']:.2f}% ; Q1 = {s_91['q1']:.2f}% ; Q3 = {s_91['q3']:.2f}% ; IQR = {s_91['q3'] - s_91['q1']:.2f}",
        f"Non-Ekstraktif: Median = {ns_91['med']:.2f}% ; Q1 = {ns_91['q1']:.3f}% ; Q3 = {ns_91['q3']:.3f}% ; IQR = {ns_91['q3'] - ns_91['q1']:.3f}",
        f"Ekstraktif Fences: {s_91['lower']:.2f}% s.d. {s_91['upper']:.2f}%   ;   Min = {s_91['min']:.2f}% ; Max = {s_91['max']:.2f}%",
        f"Mean_Ekstraktif = {smelter_avg_yoy_91:.2f}% (N={n_s_91})   vs   Mean_Non-Ekstraktif = {non_smelter_avg_yoy_91:.2f}% (N={n_ns_91})",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 9.1: Rincian Metrik Anatomi Boxplot Laju Pertumbuhan Penduduk (YoY %)",
        markdown_table(["Kategori", "Max (%)", "Upper Fence (%)", "Q3 (%)", "Median (%)", "Q1 (%)", "Lower Fence (%)", "Min (%)"], box_rows_91),
        "",
        "##### Tabel 9.2: Rincian Perhitungan Mean & Varians Laju Pertumbuhan YoY",
        markdown_table(["Kategori", "Rata-Rata / Mean (%)", "Standard Deviation", "Jumlah Sampel (Tahun-Kabupaten)"], mean_rows_91),
        "",
        "##### Tabel 9.3: Daftar Kabupaten Prioritas Industri Ekstraktif (Klasifikasi Fase 4)",
        markdown_table(["No", "Kabupaten"], kab_rows_91),
        "",
        "#### E. Analisis Temuan Empiris: Bukti Matematis Fenomena Boom and Bust",
        f"1. **Median Konsisten Lebih Tinggi:** nilai tengah pertumbuhan kawasan industri ekstraktif mencapai **{s_91['med']:.2f}%**, konsisten lebih tinggi dibanding kawasan non-ekstraktif ({ns_91['med']:.2f}%).",
        f"2. **Variabilitas Ekstrem Kawasan Ekstraktif:** IQR wilayah ekstraktif merentang Q1 {s_91['q1']:.2f}% hingga Q3 {s_91['q3']:.2f}% — jauh lebih lebar dibanding non-ekstraktif (Q1 {ns_91['q1']:.3f}% hingga Q3 {ns_91['q3']:.3f}%) yang tumbuh lebih stabil.",
        f"3. **Bukti Boom and Bust:** sebaran data ekstraktif menembus batas kewajaran (fences {s_91['lower']:.2f}% s.d. {s_91['upper']:.2f}%): lonjakan tertinggi **{s_91['max']:.2f}%** dan kejatuhan terendah **{s_91['min']:.2f}%** — bukti matematis masuknya pekerja migran secara masif di awal fase konstruksi pabrik, disusul eksodus drastis ketika proyek operasional menyusut atau terjadi pemutusan kerja massal.",
        "",
        "## 9.2 Intensifikasi Ruang: Kepadatan Industri Ekstraktif vs Non-Ekstraktif",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Demografi: `data/processed/sulawesi_demografi_master_fase4.csv` (BPS SIMDASI). Visualisasi dashboard menampilkan Area Chart tren rata-rata kepadatan penduduk (jiwa/km2) dua kategori wilayah beserta tabel agregasi kepadatan per kategori.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Yang dibaca sub-bab ini adalah **intensifikasi ruang** — tekanan yang muncul ketika pertumbuhan penduduk dan konsentrasi industri bertemu pada wilayah yang sama (bukan klaim urbanisasi formal Podes). Rata-rata kepadatan kabupaten smelter pada {latest_year_91} mencapai **{latest_smelter_density_92:.1f} jiwa/km2**, sedangkan kabupaten non-smelter berada pada **{latest_non_smelter_density_92:.1f} jiwa/km2**. Rasio **{density_ratio_92:.2f} kali** memberi sinyal kebutuhan kapasitas layanan publik yang berbeda: perumahan, air bersih, sanitasi, transportasi, hingga fasilitas kesehatan. Dalam kerangka D3TLH, kepadatan adalah indikator apakah ruang hidup lokal sedang dipadatkan oleh proyek ekstraktif tanpa perencanaan sosial yang sepadan.",
        "",
        "#### B. Alur Logika Metodologis Comparative Density Analysis",
        "Kerangka komparasi kepadatan dua kategori wilayah diilustrasikan pada **Bagan Alur 9.2** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan perbandingan rata-rata kepadatan deskriptif antar kategori dari waktu ke waktu.",
        "",
        "##### Bagan Alur 9.2: Alur Logika Analisis Comparative Density Analysis",
        "```mermaid",
        mermaid_str_9_2,
        "```",
        "",
        "#### C. Formulasi Matematis: Rata-rata Kepadatan dan Rasio Intensifikasi",
        "Kuantifikasi intensifikasi ruang dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "D_c,t = [ Σ ( Densitas_k,t ) ] / n_c   ;   untuk seluruh kabupaten k dalam kategori c pada tahun t",
        "R_t = D_Ekstraktif,t / D_Non-Ekstraktif,t",
        "```",
        "",
        "Substitusi angka dari dataset aktual:",
        "",
        "```text",
        f"D_Ekstraktif,{latest_year_91} = {latest_smelter_density_92:.1f} jiwa/km2   ;   D_Non-Ekstraktif,{latest_year_91} = {latest_non_smelter_density_92:.1f} jiwa/km2",
        f"R_{latest_year_91} = {latest_smelter_density_92:.1f} / {latest_non_smelter_density_92:.1f} = {density_ratio_92:.2f}x",
        f"Δ D_Ekstraktif ({tahun_awal_92}-{latest_year_91}) = {delta_s_92:+.1f} jiwa/km2   ;   Δ D_Non-Ekstraktif = {delta_ns_92:+.1f} jiwa/km2",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 9.4: Agregasi Rata-rata Kepadatan Penduduk per Kategori Wilayah ({tahun_awal_92}-{latest_year_91})",
        markdown_table(["Tahun", "Industri Ekstraktif (jiwa/km2)", "Non-Ekstraktif (jiwa/km2)", "Rasio (x)"], density_rows_92),
        "",
        "#### E. Analisis Temuan Empiris: Peta Awal Tekanan Ruang",
        f"1. **Profil Kepadatan Dua Kategori:** pada {latest_year_91}, rata-rata kepadatan kabupaten industri ekstraktif **{latest_smelter_density_92:.1f} jiwa/km2** — rasio {density_ratio_92:.2f} kali terhadap kabupaten non-ekstraktif ({latest_non_smelter_density_92:.1f} jiwa/km2); kawasan ekstraktif berbasis kabupaten berwilayah luas dan semula berpenduduk jarang, sehingga rasionya di bawah satu.",
        f"2. **Intensifikasi Jauh Lebih Cepat di Kawasan Ekstraktif:** sepanjang window data {tahun_awal_92}-{latest_year_91}, rata-rata kepadatan kawasan ekstraktif melipat **{lipat_s_92:.1f} kali** (dari {s_awal_92:.1f} menjadi {latest_smelter_density_92:.1f} jiwa/km2, {delta_s_92:+.1f}), jauh melampaui laju kawasan non-ekstraktif yang hanya {lipat_ns_92:.1f} kali — inilah intensifikasi ruang yang dibaca sub-bab ini.",
        "3. **Implikasi Kapasitas Layanan Publik:** pemadatan cepat pada ruang yang semula lengang memberi sinyal kebutuhan kapasitas layanan publik yang berbeda di kawasan industri — perumahan, air bersih, sanitasi, transportasi, hingga fasilitas kesehatan — agar ruang hidup lokal tidak dipadatkan proyek ekstraktif tanpa perencanaan sosial yang sepadan.",
        "",
        "## 9.3 Pergeseran Ekonomi Agraris ke Tambang dan Industri",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Shift Index: `data/processed/sulawesi_employment_shift_fase4.csv` dan `data/processed/sulawesi_pdrb_sektoral_2016_2024.csv` (BPS SIMDASI). Visualisasi dashboard menampilkan Area Chart komposisi PDRB sektor kunci per provinsi (dengan dekomposisi estimasi Perikanan Tangkap) serta Line Chart Shift Index (B+C / A) 6 provinsi dengan garis ambang 1.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Pergeseran pekerjaan tidak dapat diklaim hanya dari PDRB, tetapi struktur PDRB memberi petunjuk kuat tentang arah ekonomi yang sedang dibentuk. Sektor A dibaca sebagai basis agraris, sektor B dan C sebagai blok ekstraktif-industrial. Rasio B+C terhadap A menjadi *shift index*; nilai di atas 1 berarti kontribusi tambang dan industri sudah melampaui pertanian. Di Sulawesi Tengah, porsi pertanian turun dari **{pertanian_awal_93:.2f}%** menjadi **{pertanian_akhir_93:.2f}%**, sementara tambang+industri naik dari **{industri_awal_93:.2f}%** menjadi **{industri_akhir_93:.2f}%**. Indeksnya naik dari **{shift_awal_93:.3f}** ke **{shift_akhir_93:.3f}**, atau sekitar **{shift_multiplier_93:.1f} kali** — hilirisasi mengubah pusat gravitasi ekonomi daerah dari ruang produksi agraris menuju rantai ekstraktif yang lebih terkonsentrasi pada modal besar.",
        "",
        "#### B. Alur Logika Metodologis PDRB Sector Shift Index (B+C / A)",
        "Kerangka pembacaan pergeseran struktur ekonomi berbasis shift index diilustrasikan pada **Bagan Alur 9.3** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan indeks rasio sektoral deskriptif dengan ambang interpretatif 1.",
        "",
        "##### Bagan Alur 9.3: Alur Logika Analisis PDRB Sector Shift Index",
        "```mermaid",
        mermaid_str_9_3,
        "```",
        "",
        "#### C. Formulasi Matematis: Shift Index, Dekomposisi Sektor A, dan Multiplier",
        "Kuantifikasi pergeseran struktur ekonomi dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Shift_Index_p,t = ( PDRB_B_p,t + PDRB_C_p,t ) / PDRB_A_p,t",
        "Perikanan_Tangkap ≈ 0,22 × Sektor_A   ;   Pertanian_Kehutanan ≈ 0,78 × Sektor_A",
        "Multiplier_p = Shift_Index_p,akhir / Shift_Index_p,awal",
        "```",
        "",
        "Proporsi 0,22 adalah estimasi Perikanan Tangkap terhadap Sektor A, mengacu rata-rata proporsi sub-sektor perikanan di provinsi pesisir Sulawesi (Statistik Perikanan BPS Sulawesi, 2016-2024).",
        "",
        "Substitusi angka dari dataset aktual (Sulawesi Tengah sebagai episentrum):",
        "",
        "```text",
        f"Shift_Index_Sulteng: {shift_awal_93:.3f} ({tahun_awal_93}) → {shift_akhir_93:.3f} ({tahun_akhir_93})   ;   Multiplier = {shift_akhir_93:.3f} / {shift_awal_93:.3f} = {shift_multiplier_93:.1f}x",
        f"Pertanian (A): {pertanian_awal_93:.2f}% → {pertanian_akhir_93:.2f}%   ;   Tambang+Industri (B+C): {industri_awal_93:.2f}% → {industri_akhir_93:.2f}%",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 9.5: Shift Index (B+C / A) per Provinsi per Tahun ({tahun_awal_93}-{tahun_akhir_93})",
        markdown_table(["Tahun"] + prov_cols_93, index_rows_93),
        "",
        "##### Tabel 9.6: Ringkasan Pergeseran Struktur Ekonomi per Provinsi",
        markdown_table(["Provinsi", "Index Awal", "Index Akhir", "Multiplier", "Status Ambang"], ringkas_rows_93),
        "",
        "#### E. Analisis Temuan Empiris: Pergeseran Pusat Gravitasi Ekonomi",
        f"1. **Episentrum Pergeseran di Sulawesi Tengah:** shift index Sulteng melonjak dari {shift_awal_93:.3f} menjadi **{shift_akhir_93:.3f}** ({shift_multiplier_93:.1f} kali) — blok tambang+industri kini {industri_akhir_93:.2f}% PDRB, jauh melampaui pertanian yang menyusut ke {pertanian_akhir_93:.2f}%.",
        f"2. **Provinsi Pelampau Ambang:** provinsi dengan shift index akhir melampaui ambang 1 (B+C > A): **{', '.join(prov_lampaui_93) if prov_lampaui_93 else 'tidak ada'}** — sementara provinsi lain masih berbasis agraris.",
        "3. **Catatan Metodologis:** dekomposisi Perikanan Tangkap adalah estimasi (±22% Sektor A) untuk keperluan visual; klaim pergeseran pekerjaan tidak ditarik dari PDRB semata, melainkan dibaca sebagai petunjuk arah ekonomi yang sedang dibentuk hilirisasi — dari ruang produksi agraris menuju rantai ekstraktif yang terkonsentrasi pada modal besar.",
        "",
    ]
    md_path = tool_dir / "Metodologi_Bab9_Demografi_Sosial.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")
    print("[4/4] Selesai membangun Bab 9.")


if __name__ == "__main__":
    generate_all_bab9()
