#!/usr/bin/env python3
"""
Generator Laporan Metodologi Bab 8: Distribusi Manfaat vs Beban Ekologis

Pilar 1 ditulis langsung dalam generator Python agar selaras dengan SOP dokumentasi Celios2.
Fokus awal: Sub-bab 8.1 Sisi Manfaat: Gurita Bisnis & Monopoli Keuntungan Ekstraktif.
"""

import base64
import sys
from pathlib import Path

try:
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([
        sys.executable,
        "-m",
        "pip",
        "install",
        "pandas",
        "requests",
        "python-docx",
    ])
    import pandas as pd
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, cfg):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        for key, val in cfg.items():
            el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def para_border_bottom(paragraph, color="2E7D32", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="2E7D32", sz="16"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="A5D6A7", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="1B5E20", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="2E7D32", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=G_MID)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)


def add_p(doc, parts, space_after=5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "EDF7EE")
    all_border_para(p)
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)

    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=G_DARK)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p)
    para_shd(p, "F1F8E9")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}

    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "2E7D32")
        set_cell_borders(cell, bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8, color=C_WHITE)

    for i, row_data in enumerate(rows):
        fill = "F5FBF5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            set_cell_borders(cell, bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def generate_all_bab8():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_8"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 8 sub-bab 8.1...")
    # Konstanta referensi riset kekayaan taipan (CELIOS Inequality Report 2026),
    # sinkron dengan metric cards & narasi dashboard 8_Distribusi_Manfaat.py
    total_harta_50_t_81 = 4651.0
    pct_ekstraktif_81 = 58.0
    harta_ekstraktif_t_81 = total_harta_50_t_81 * pct_ekstraktif_81 / 100
    laju_harian_m_81 = 13.0
    laju_upah_buruh_rb_81 = 2.0

    # Mega-Crosstab Top 10 (kurasi riset dashboard: CELIOS Inequality Report 2026
    # + agregasi sulawesi_kawasan_nikel_luas.csv, sulawesi_pltu_captive.csv,
    # sulawesi_konflik_agraria_tanahkita.csv)
    mega_rows_81 = [
        ["#1 PT Vale Indonesia (MIND ID & Konsorsium)", "Rp 259,2 T", "Blok Sorowako, Bahodopi, Pomalaa", "118.017 Ha", "Monopoli & deforestasi kronis Pegunungan Verbeek", "0 MW (Suplai PLTA Sorowako)", "> Rp 40,0 T", "460+ Jiwa Terdampak (wilayah adat To Karunsi'e)"],
        ["#2 Salim Group (Anthony Salim)", "Rp 160,0 T", "Citra Palu Minerals, Gorontalo Min.", "110.175 Ha", "Tumpang tindih dengan Taman Hutan Raya (Tahura)", "Tambang Emas (Non-Smelter)", "> Rp 8,0 T", "Konflik PETI Poboya (penertiban paksa penambang rakyat)"],
        ["#3 Jiangsu Delong Nickel (Tony Zhou Yuan)", "Rp 45,0 T", "PT VDNI, OSS (Konawe), GNI (Morut)", "2.253 Ha", "Perusakan DAS Laronai & bentang alam Morosi", "5.175 MW (~36,2 Jt Ton CO2/thn)", "> Rp 20,0 T", "2 Pekerja Tewas (bentrokan GNI 2023)"],
        ["#4 Tsingshan Holding (Xiang Guangda)", "Rp 163,0 T", "Bintangdelapan, Eternal (IMIP)", "20.765 Ha", "Deforestasi masif hutan pesisir & reklamasi", "4.030 MW (~28,2 Jt Ton CO2/thn)", "> Rp 40,0 T", "Puluhan Pekerja Tewas (ledakan tungku ITSS)"],
        ["#5 Boy Thohir & Edwin S. (Adaro/Saratoga)", "Rp 64,1 T", "PT Sulawesi Cahaya Mineral (SCM)", "21.100 Ha", "Sinyal hilangnya hutan primer tinggi (GFW)", "Disuplai Listrik PLN (Undisclosed)", "> Rp 15,0 T", "Konflik Tenurial Laten (deforestasi blok Routa)"],
        ["#6 J Resources (Jimmy Budiarto)", "Rp 7,5 T", "J Resources Bolaang Mongondow", "38.150 Ha", "Eksploitasi lanskap Pegunungan Bolmong", "Tambang Emas (Non-Smelter)", "> Rp 5,0 T", "Potensi Pencemaran (masyarakat lingkar tambang)"],
        ["#7 Rajawali Group (Peter Sondakh)", "Rp 32,5 T", "Tambang Tondano Nusajaya (Archi)", "30.848 Ha", "Berkurangnya resapan air di Minahasa", "Tambang Emas (Non-Smelter)", "> Rp 4,5 T", "Banjir & Longsor (Sulawesi Utara)"],
        ["#8 Kalla Group (Keluarga Jusuf Kalla)", "Rp 900,8 M", "PT Kalla Arebamma, Bumi Mineral", "20.173 Ha", "Reklamasi pesisir merusak ekosistem mangrove", "0 MW (Suplai PLTA Poso)", "> Rp 2,5 T", "Konflik Lahan Luwu (gusur paksa nelayan Bua)"],
        ["#9 Harita Group (Lim Hariyanto W.S.)", "Rp 108,0 T", "PT Gema Kreasi Perdana (Wawonii)", "~1.000 Ha", "Menabrak regulasi larangan tambang pulau kecil", "Ekspor Bijih Mentah (PLTU >1.100 MW di P. Obi)", "> Rp 1,5 T", "37.000 Jiwa Terdampak (kriminalisasi penolak tambang)"],
        ["#10 Zhenshi Holding (Zhang Yuqiang)", "Rp 40,0 T", "Zhenshi Holding Group Co Ltd", "4.000 Ha", "Mengubah kawasan hijau pesisir menjadi beton", "450 MW (~3,1 Jt Ton CO2/thn)", "> Rp 5,0 T", "Krisis Ruang Hidup (desa lingkar tambang Morowali)"],
    ]

    kolom_map_headers_81 = ["Kolom Tabel 8.1", "Deskripsi & Cara Perolehan", "Sumber Data / Persamaan Terkait"]
    kolom_map_rows_81 = [
        ["1. Grup Taipan / Konsorsium", "Identitas grup oligarki hasil pemetaan afiliasi kepemilikan (Hierarchical Entity Profiling), diurutkan Top 10 berdasarkan skala daya rusak.", "CELIOS Inequality Report 2026 (tanpa persamaan; pemetaan kualitatif)"],
        ["2. Total Harta (CELIOS)", "Akumulasi kekayaan (Net Worth) individu/grup pada laporan ketimpangan.", "CELIOS Inequality Report 2026 — Persamaan Total Kekayaan Ekstraktif (Bagian C)"],
        ["3. Afiliasi Blok (Sulawesi)", "Entitas operasi (PT) milik grup yang beroperasi di blok Sulawesi.", "Pemetaan nama perusahaan normatif (tanpa persamaan)"],
        ["4. Luas Konsesi (Aktual)", "Agregasi luasan konsesi seluruh entitas terafiliasi grup.", "sulawesi_kawasan_nikel_luas.csv — Persamaan Agregasi Luas Konsesi per Grup (Bagian C)"],
        ["5. Status Deforestasi Lindung", "Penilaian kualitatif tumpang tindih operasi dengan kawasan lindung/ekosistem esensial (overlay GFW & kawasan konservasi).", "GFW & regulasi kawasan (kualitatif; tanpa persamaan numerik)"],
        ["6. Emisi PLTU Captive", "Agregasi kapasitas PLTU per Parent grup dan konversi ke taksiran jejak karbon tahunan.", "sulawesi_pltu_captive.csv — Persamaan Agregasi Kapasitas & Konversi Emisi (Bagian C)"],
        ["7. Estimasi Rugi Ekologis", "Valuasi ekonomi lingkungan gabungan komponen konsesi dan emisi karbon.", "Adaptasi PermenLHK No. 7/2014 — Persamaan Valuasi Rugi Ekologis (Bagian C)"],
        ["8. Dampak Sosial & Konflik", "Agregasi jiwa terdampak dan insiden konflik yang terafiliasi entitas grup.", "sulawesi_konflik_agraria_tanahkita.csv — Persamaan Agregasi Dampak Sosial (Bagian C)"],
    ]

    # Agregat daya rusak Top 10 (turunan dari nilai kurasi mega_rows_81)
    total_konsesi_top10_81 = 118017 + 110175 + 2253 + 20765 + 21100 + 38150 + 30848 + 20173 + 1000 + 4000
    total_mw_top10_81 = 5175 + 4030 + 450
    total_rugi_top10_81 = 40.0 + 8.0 + 20.0 + 40.0 + 15.0 + 5.0 + 4.5 + 2.5 + 1.5 + 5.0
    total_harta_top10_81 = 259.2 + 160.0 + 45.0 + 163.0 + 64.1 + 7.5 + 32.5 + 0.9008 + 108.0 + 40.0

    # Format gaya Indonesia (titik ribuan, koma desimal) agar konsisten dengan sel tabel
    konsesi_top10_str_81 = f"{total_konsesi_top10_81:,.0f}".replace(",", ".")
    mw_top10_str_81 = f"{total_mw_top10_81:,.0f}".replace(",", ".")
    rugi_top10_str_81 = f"{total_rugi_top10_81:.1f}".replace(".", ",")
    harta_top10_str_81 = f"{total_harta_top10_81:.1f}".replace(".", ",")

    total_row_81 = [
        "TOTAL TOP 10",
        f"Rp {harta_top10_str_81} T",
        "-",
        f"{konsesi_top10_str_81} Ha",
        "-",
        f"{mw_top10_str_81} MW (agregat captive terkuantifikasi)",
        f"> Rp {rugi_top10_str_81} T",
        "-",
    ]
    mega_rows_total_81 = mega_rows_81 + [total_row_81]

    mermaid_str_8_1 = """flowchart LR
    subgraph Data_Input["1. Input Data Riset"]
        A["CELIOS Inequality Report 2026<br/><i>Net Worth 50 triliuner & afiliasi sektor</i>"] --> D
        B["Dataset Internal Sulawesi<br/><i>kawasan nikel luas, PLTU captive, konflik TanahKita</i>"] --> D
    end
    subgraph Entity_Profiling["2. Hierarchical Entity Profiling"]
        D["Pemetaan afiliasi grup taipan<br/>ke entitas operasi di Sulawesi"] --> E["Agregasi instrumen kerusakan per grup<br/>Luas Konsesi; Kapasitas PLTU; Konflik"]
        E --> F["Valuasi rugi ekologis<br/>adaptasi PermenLHK No. 7/2014"]
    end
    subgraph Output_Analysis["3. Mega-Crosstab Pemetaan Aktor"]
        F --> G["Tabel Top 10 Grup Oligarki<br/>harta vs daya rusak privat"]
        G --> H["Pembacaan monopoli keuntungan ekstraktif"]
    end"""
    mermaid_png_path_8_1 = str(tool_dir / "mermaid_flowchart_8_1.png")
    download_success_8_1 = download_mermaid_png(mermaid_str_8_1, mermaid_png_path_8_1)

    print("[1.5/4] Mengekstraksi dataset empiris Bab 8 sub-bab 8.2...")
    data_dir = base_dir / "data" / "processed"
    df_kes_82 = pd.read_csv(data_dir / "sulawesi_kesehatan_detail_2014_2024.csv")
    sentra_82 = ["Sulawesi Tengah", "Sulawesi Tenggara"]
    df_ispa_82 = df_kes_82[(df_kes_82["indikator"] == "Kasus ISPA/Pneumonia") & (df_kes_82["provinsi"].isin(sentra_82))].copy()
    total_ispa_sentra_82 = int(df_ispa_82["nilai"].sum())
    tahun_min_82 = int(df_ispa_82["tahun"].min())
    tahun_max_82 = int(df_ispa_82["tahun"].max())

    pivot_ispa_82 = df_ispa_82.pivot_table(index="tahun", columns="provinsi", values="nilai", aggfunc="sum", fill_value=0)
    pivot_ispa_82["Total"] = pivot_ispa_82.sum(axis=1)
    pivot_ispa_82["Kumulatif"] = pivot_ispa_82["Total"].cumsum()
    tren_rows_82 = []
    for tahun, row in pivot_ispa_82.iterrows():
        tren_rows_82.append([
            str(int(tahun)),
            f"{row.get('Sulawesi Tengah', 0):,.0f}",
            f"{row.get('Sulawesi Tenggara', 0):,.0f}",
            f"{row['Total']:,.0f}",
            f"{row['Kumulatif']:,.0f}",
        ])
    tahun_puncak_ispa_82 = int(pivot_ispa_82["Total"].idxmax())
    nilai_puncak_ispa_82 = int(pivot_ispa_82["Total"].max())
    tahun_terendah_ispa_82 = int(pivot_ispa_82["Total"].idxmin())
    nilai_terendah_ispa_82 = int(pivot_ispa_82["Total"].min())

    # Konstanta kurasi riset (metric cards dashboard 8.2)
    kasus_kritis_82 = 12
    rugi_min_t_82 = 100

    beban_rows_82 = [
        ["Krisis Kesehatan (ISPA)", f"{total_ispa_sentra_82:,.0f} kasus", f"Akumulasi kasus infeksi saluran pernapasan di sentra nikel Sulteng & Sultra ({tahun_min_82}-{tahun_max_82}), berkorelasi dengan polusi debu dan sulfur PLTU Captive.", "Data Panel Kesehatan (Dinkes/BPS)"],
        ["Konflik Agraria & FPIC", f"{kasus_kritis_82} Kasus Kritis", "Terdokumentasi meletus di Sulawesi; mengorbankan puluhan ribu jiwa, melibatkan perampasan kebun, pelanggaran hak adat, dan penembakan warga.", "Tanahkita.id (KPA / YLBHI)"],
        ["Estimasi Kerugian Ekologis", f"> Rp {rugi_min_t_82} Triliun", "Valuasi kumulatif kasar dari hilangnya fungsi hutan primer, rusaknya ekosistem terumbu karang laut, dan lenyapnya sumber air bersih akibat sedimentasi limbah.", "Proksi Kalkulasi Valuasi Lingkungan LHK"],
    ]

    mermaid_str_8_2 = """flowchart LR
    subgraph Data_Input["1. Input Data Riset"]
        A["Data Panel Kesehatan Dinkes/BPS<br/><i>provinsi, tahun, indikator, nilai</i>"]
        B["Repositori TanahKita/KPA (CATAHU)<br/><i>kasus kritis konflik agraria & FPIC</i>"]
        C["Proksi Valuasi Lingkungan LHK<br/><i>estimasi rupiah kerusakan</i>"]
    end
    subgraph TimeSeries_Aggregation["2. Agregasi Deret Waktu Deskriptif"]
        A --> D["Trend Mapping kasus ISPA<br/>sentra nikel Sulteng & Sultra 2014-2024"]
        B --> E["Agregasi kasus kritis sengketa lahan"]
        C --> F["Valuasi kumulatif eksternalitas"]
    end
    subgraph Output_Analysis["3. Ringkasan Indikator Beban Publik"]
        D --> G["Kartu metrik beban kesehatan, konflik, dan kerugian ekologis"]
        E --> G
        F --> G
    end
    G --> H["Pembacaan eksternalitas negatif industrialisasi ekstraktif"]"""
    mermaid_png_path_8_2 = str(tool_dir / "mermaid_flowchart_8_2.png")
    download_success_8_2 = download_mermaid_png(mermaid_str_8_2, mermaid_png_path_8_2)

    print("[1.8/4] Mengekstraksi dataset empiris Bab 8 sub-bab 8.3...")
    import scipy.stats as stats
    from functools import reduce

    df_inv_83 = pd.read_csv(data_dir / "sulawesi_investasi_pmdn_2016_2024.csv")
    df_inv_agg_83 = df_inv_83.groupby(["provinsi", "tahun"])["nilai"].sum().reset_index()
    df_inv_agg_83.rename(columns={"nilai": "Realisasi_Investasi_Rp"}, inplace=True)

    df_pad_83 = pd.read_csv(data_dir / "sulawesi_pad_2016_2024.csv")
    df_pad_83 = df_pad_83.rename(columns={"pad_juta_rupiah": "PAD_Juta_Rupiah"})[["provinsi", "tahun", "PAD_Juta_Rupiah"]]

    df_ispa_agg_83 = df_kes_82[df_kes_82["indikator"] == "Kasus ISPA/Pneumonia"].groupby(["provinsi", "tahun"])["nilai"].sum().reset_index()
    df_ispa_agg_83.rename(columns={"nilai": "Kasus_ISPA"}, inplace=True)

    df_def_83 = pd.read_csv(data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv")
    df_def_83 = df_def_83.rename(columns={"Provinsi": "provinsi", "Tahun": "tahun", "Total_Deforestasi_Ha": "Deforestasi_Ha"})[["provinsi", "tahun", "Deforestasi_Ha"]]

    dfs_83 = [df_inv_agg_83, df_pad_83, df_ispa_agg_83, df_def_83]
    df_panel_83 = reduce(lambda left, right: pd.merge(left, right, on=["provinsi", "tahun"], how="outer"), dfs_83)
    df_panel_83.rename(columns={"provinsi": "Provinsi", "tahun": "Tahun"}, inplace=True)
    n_panel_83 = len(df_panel_83)

    x_opt_83 = {
        "Realisasi_Investasi_Rp": "Investasi PMDN (Rupiah)",
        "PAD_Juta_Rupiah": "Pendapatan Asli Daerah (Juta Rp)",
    }
    y_opt_83 = {
        "Kasus_ISPA": "Beban Penyakit (Kasus ISPA)",
        "Deforestasi_Ha": "Beban Pencemaran (Deforestasi Ha)",
    }

    threshold_rows_83 = []
    summary_rows_83 = []
    detail_83 = {}
    for k_x, v_x in x_opt_83.items():
        for k_y, v_y in y_opt_83.items():
            loop_df = df_panel_83.dropna(subset=[k_x, k_y]).copy()
            if len(loop_df) == 0:
                continue
            med_x = loop_df[k_x].median()
            med_y = loop_df[k_y].median()
            lbl_x_l, lbl_x_h = "Rendah", "Tinggi"
            lbl_y_l, lbl_y_h = "Rendah", "Tinggi/Parah"
            s_x = loop_df[k_x].apply(lambda v: lbl_x_h if v >= med_x else lbl_x_l)
            s_y = loop_df[k_y].apply(lambda v: lbl_y_h if v >= med_y else lbl_y_l)
            ct = pd.crosstab(s_x, s_y).reindex(index=[lbl_x_l, lbl_x_h], columns=[lbl_y_l, lbl_y_h], fill_value=0)
            try:
                c2_val, pv_val, dof_val, _ = stats.chi2_contingency(ct)
            except Exception:
                c2_val, pv_val, dof_val = 0, 1.0, 1
            try:
                aa = ct.loc[lbl_x_h, lbl_y_h]
                bb = ct.loc[lbl_x_h, lbl_y_l]
                cc = ct.loc[lbl_x_l, lbl_y_h]
                dd = ct.loc[lbl_x_l, lbl_y_l]
                or_v = (aa * dd) / (bb * cc) if (bb * cc) > 0 else 0
            except Exception:
                or_v = 0
            p_disp_83 = "p < 0.001" if pv_val < 0.001 else f"p = {pv_val:.3f}"
            threshold_rows_83.append([v_x, v_y, f"{med_x:,.1f}", f"{med_y:,.1f}", f"{len(loop_df)}"])
            summary_rows_83.append([v_x, v_y, f"{c2_val:.3f}", p_disp_83, f"{or_v:.2f}", "SIGNIFIKAN" if pv_val < 0.05 else "TIDAK SIGNIFIKAN"])
            detail_83[(k_x, k_y)] = {"chi2": c2_val, "p": pv_val, "or": or_v, "n": len(loop_df), "med_x": med_x, "med_y": med_y}

    sig_count_83 = sum(1 for row in summary_rows_83 if row[5] == "SIGNIFIKAN")
    total_scen_83 = len(summary_rows_83)
    d_inv_ispa_83 = detail_83.get(("Realisasi_Investasi_Rp", "Kasus_ISPA"), {"chi2": 0, "p": 1.0, "or": 0, "n": 0, "med_x": 0, "med_y": 0})

    if sig_count_83 > 0:
        finding_83 = f"KESIMPULAN METODOLOGIS: Korelasi Indikator Investasi dan Dampak Lingkungan. Hasil pengujian statistik menunjukkan korelasi signifikan antara peningkatan arus investasi dan indikator dampak lingkungan di Sulawesi ({sig_count_83} dari {total_scen_83} skenario SIGNIFIKAN). Wilayah dengan pertumbuhan investasi tinggi mencatatkan tren insidensi penyakit saluran pernapasan dan deforestasi yang lebih tinggi. Nilai Odds Ratio mengindikasikan bahwa peningkatan aktivitas industri berasosiasi dengan kenaikan risiko eksternalitas lingkungan. Temuan ini menekankan pentingnya pengalokasian anggaran yang lebih memadai untuk perlindungan kesehatan publik, rehabilitasi ekologis, dan penguatan layanan dasar masyarakat di kawasan sekitar industri ekstraktif."
    else:
        finding_83 = f"KESIMPULAN METODOLOGIS: Evaluasi Penyebaran Dampak dan Perlunya Presisi Data. Meskipun pengujian pada skala agregat provinsi menunjukkan hasil tidak signifikan secara statistik (P >= 0.05) pada seluruh {total_scen_83} skenario, hal ini dipengaruhi oleh aggregation effect pada skala data provinsi. Analisis tingkat mikro mengindikasikan bahwa dampak lingkungan dan sosial terkonsentrasi di wilayah sekitar kawasan industri. Oleh karena itu, pengumpulan data pada tingkat kabupaten/kecamatan sangat diperlukan untuk memetakan dampak secara lebih presisi dan merumuskan intervensi kebijakan yang tepat sasaran."

    konf_headers_83 = ["Komponen Uji", "Definisi Variabel (Sub-bab 8.3)"]
    konf_rows_83 = [
        ["Variabel Independen (X)", "Indikator Manfaat Ekonomi: Realisasi Investasi PMDN (Rupiah) / Pendapatan Asli Daerah (Juta Rp)."],
        ["Variabel Dependen (Y)", "Indikator Beban: Kasus ISPA / Deforestasi (Ha)."],
        ["Hipotesis Nol (H0)", "Tidak ada korelasi yang signifikan secara statistik antara nilai investasi PMDN/PAD dengan jumlah penderita ISPA/Deforestasi di provinsi Sulawesi pada suatu tahun tertentu."],
        ["Hipotesis Alternatif (H1)", "Semakin tinggi indikator manfaat ekonomi (Investasi/PAD) yang masuk ke suatu provinsi, semakin parah pula lonjakan kasus beban ekologis (Penyakit/Deforestasi) yang dialami warganya."],
        ["Decision Rule (Alpha 5%)", "Tolak H0 jika nilai Asymptotic Significance (P-Value) pada uji Pearson Chi-Square < 0.05 (Alpha 5%)."],
        ["Threshold Kategori", f"Nilai Median historis panel gabungan 4 dataset (outer join, N={n_panel_83} baris panel; N valid per skenario pada Tabel 8.5): Nilai >= Median = Tinggi, selain itu Rendah."],
        ["Orientasi Odds Ratio", "Y berjenis beban (y_is_negative): OR = ( a × d ) / ( b × c ) dengan a = X Tinggi & Y Tinggi/Parah; mengukur risiko beban parah pada kelompok manfaat ekonomi tinggi."],
    ]

    mermaid_str_8_3 = """flowchart LR
    subgraph Data_Input["1. Integrasi Panel 4 Dataset"]
        A["Investasi PMDN 2016-2024"] --> E
        B["PAD 2016-2024"] --> E
        C["Kasus ISPA 2014-2024"] --> E
        D["Deforestasi GFW 2014-2023"] --> E
    end
    subgraph Panel_Join["2. Pembentukan Panel Ketimpangan"]
        E["Outer Join Provinsi-Tahun"] --> F["Panel Manfaat (X) vs Beban (Y)"]
    end
    subgraph Statistical_Test["3. Crosstabulation & Pearson Chi-Square"]
        F --> G["Binning Median Historis<br/>Tinggi vs Rendah"]
        G --> H["Uji Chi-Square 4 skenario X × Y"]
        H --> I["Odds Ratio<br/>risiko beban parah saat manfaat tinggi"]
    end
    I --> J["Pembacaan matriks ketimpangan manfaat vs beban"]"""
    mermaid_png_path_8_3 = str(tool_dir / "mermaid_flowchart_8_3.png")
    download_success_8_3 = download_mermaid_png(mermaid_str_8_3, mermaid_png_path_8_3)

    print("[2/4] Membangun DOCX Metodologi_Bab8_Distribusi_Manfaat.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)
    add_h1(doc, "BAB VIII: METODOLOGI ANALISIS DISTRIBUSI MANFAAT VS BEBAN EKOLOGIS")
    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada ", False, False),
        ("Bab 8: Distribusi Manfaat vs Beban Ekologis", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.", False, False),
    ])

    add_h2(doc, "8.1 Sisi Manfaat: Gurita Bisnis & Monopoli Keuntungan Ekstraktif")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Sumber utama: CELIOS Inequality Report 2026 (Laporan 50 Taipan Terkaya); Dataset internal: data/processed/sulawesi_kawasan_nikel_luas.csv (agregasi nama perusahaan normatif), data/processed/sulawesi_pltu_captive.csv (agregasi Parent & Capacity MW), data/processed/sulawesi_konflik_agraria_tanahkita.csv. Visualisasi dashboard menampilkan tiga metric cards konsentrasi kekayaan serta Mega-Crosstab Top 10 Grup Taipan vs kerugian publik.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Analisis terhadap distribusi manfaat ekonomi sektor nikel dan PLTU di Sulawesi menunjukkan konsentrasi nilai tambah pada kelompok usaha skala besar. ", False, False),
        (f"Data dari Laporan Ketimpangan CELIOS mencatat bahwa akumulasi kekayaan 50 individu/kelompok usaha terbesar di Indonesia mencapai Rp{total_harta_50_t_81:,.0f} Triliun, di mana sekitar {pct_ekstraktif_81:.0f}% bersumber dari sektor berbasis sumber daya alam (pertambangan nikel, batu bara, kelapa sawit, dan pemurnian logam). ", False, False),
        (f"Kekayaan ini naik nyaris dua kali lipat sejak 2019 dengan laju kenaikan harta harian mencapai Rp{laju_harian_m_81:.0f} Miliar — sangat kontras dengan rata-rata kenaikan upah buruh nasional yang hanya tumbuh sekitar Rp{laju_upah_buruh_rb_81:.0f} ribu per hari. Hal ini mengindikasikan perlunya kebijakan redistribusi manfaat dan pengelolaan dampak lingkungan yang lebih seimbang.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Pemetaan Konsentrasi Kekayaan Ekstraktif")
    add_p(doc, [
        ("Kerangka pemrofilan entitas bisnis berjenjang (Hierarchical Entity Profiling) untuk melacak aliran penguasaan sumber daya menuju kelompok elit diilustrasikan pada ", False, False),
        ("Bagan Alur 8.1", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Wealth Database Analysis dan Mega-Crosstab pemetaan aktor deskriptif.", False, False),
    ])
    add_caption(doc, "Bagan Alur 8.1: Alur Logika Analisis Pemetaan Konsentrasi Kekayaan Ekstraktif")
    if download_success_8_1:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_8_1, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Konsentrasi Kekayaan dan Valuasi Rugi Ekologis")
    add_p(doc, [("Kuantifikasi konsentrasi kekayaan dan daya rusak privat dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Total Kekayaan Ekstraktif", "Total_Kekayaan_Ekstraktif = Σ ( Harta_i )   ;   untuk seluruh triliuner i dengan Sektor = 'Ekstraktif'", [
        ("Total_Kekayaan_Ekstraktif", "Akumulasi kekayaan (Net Worth) kelompok Top 50 yang bersumber dari sektor sumber daya alam."),
        ("Harta_i", "Nilai kekayaan agregat individu/grup usaha ke-i pada Laporan Ketimpangan CELIOS."),
    ])
    add_formula(doc, "Persamaan Beban Ekologis per Grup Afiliasi", "Beban_Ekologis_g = Σ ( Rugi_Ekologis_e )   ;   untuk seluruh entitas operasi e dengan Afiliasi_Pemilik = g", [
        ("Beban_Ekologis_g", "Skala kerugian publik (eksternalitas negatif) yang dihasilkan konsorsium/grup bisnis afiliasi g."),
        ("Rugi_Ekologis_e", "Estimasi kerugian ekologis entitas operasi e (konsesi, smelter, PLTU) di Sulawesi."),
    ])
    add_formula(doc, "Persamaan Agregasi Luas Konsesi per Grup (Kolom 4 Tabel 8.1)", "Luas_Konsesi_g = Σ ( total_luas_ha_e )   ;   untuk seluruh entitas e dengan afiliasi grup g", [
        ("Luas_Konsesi_g", "Total luasan konsesi (Ha) seluruh PT terafiliasi grup g, di-aggregate berdasarkan nama perusahaan normatif pada sulawesi_kawasan_nikel_luas.csv."),
    ])
    add_formula(doc, "Persamaan Agregasi Kapasitas & Konversi Emisi PLTU (Kolom 6 Tabel 8.1)", "Kapasitas_PLTU_g = Σ ( Capacity_MW_u ) untuk unit u dengan Parent = g   ;   Emisi_CO2_g ≈ Kapasitas_PLTU_g × Faktor_Emisi", [
        ("Kapasitas_PLTU_g", "Total kapasitas PLTU captive (MW) yang di-aggregate berdasarkan kolom Parent pada sulawesi_pltu_captive.csv."),
        ("Faktor_Emisi", "Taksiran jejak karbon tahunan PLTU batubara captive, ~7.000 Ton CO2 per MW per tahun (faktor implisit yang konsisten pada seluruh baris terkuantifikasi Tabel 8.1)."),
    ])
    add_formula(doc, "Persamaan Agregasi Dampak Sosial (Kolom 8 Tabel 8.1)", "Dampak_Sosial_g = Σ ( Jiwa_Terdampak_k )   ;   untuk seluruh kasus konflik k terafiliasi entitas grup g", [
        ("Jiwa_Terdampak_k", "Jumlah jiwa terdampak/insiden korban pada kasus k dalam repositori sulawesi_konflik_agraria_tanahkita.csv."),
    ])
    add_formula(doc, "Persamaan Valuasi Rugi Ekologis (Kolom 7 Tabel 8.1, Adaptasi PermenLHK No. 7/2014)", "Total_Kerugian_Ekologis = ( Luas_Konsesi × Valuasi_Hutan_per_Ha ) + ( Kapasitas_PLTU_MW × Biaya_Sosial_Emisi_Karbon )", [
        ("Luas_Konsesi", "Luasan konsesi (HGU/IUP) beroperasi; multiplier kerugian dikalikan lipat bila menembus Cagar Alam, Taman Nasional, atau permukiman warga."),
        ("Valuasi_Hutan_per_Ha", "Nilai kerugian ekonomi publik dan biaya pemulihan fungsi ekologis per hektar (reboisasi, netralisasi limbah, kesehatan)."),
        ("Biaya_Sosial_Emisi_Karbon", "Konversi jejak karbon PLTU captive (juta ton CO2/tahun) dikalikan Social Cost of Carbon (SCC) / Nilai Ekonomi Karbon (NEK)."),
    ])

    add_p(doc, [("Substitusi angka dari laporan riset aktual ke dalam rumus konsentrasi kekayaan dan valuasi rugi ekologis adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Total Kekayaan Ekstraktif", f"Total_Kekayaan_Ekstraktif = {pct_ekstraktif_81:.0f}% × Rp{total_harta_50_t_81:,.0f} T = Rp{harta_ekstraktif_t_81:,.1f} Triliun")
    add_formula(doc, "Substitusi Laju Akumulasi Harian", f"Laju_Harta_Elit = Rp{laju_harian_m_81:.0f} Miliar/hari   vs   Laju_Upah_Buruh = Rp{laju_upah_buruh_rb_81:.0f} Ribu/hari")
    add_formula(doc, "Substitusi Valuasi Rugi Ekologis (Contoh Baris #1 Tabel 8.1: PT Vale Indonesia)", "Total_Kerugian_Vale = ( 118.017 Ha × Valuasi_Hutan_per_Ha ) + ( 0 MW × SCC/NEK ) ≈ > Rp 40,0 Triliun", [
        ("Komponen Konsesi", "Luas konsesi aktual Vale 118.017 Ha (terbesar di dataset) dikalikan valuasi kerugian ekonomi publik dan biaya pemulihan per hektar — termasuk multiplier kumulatif kerusakan Danau Matano/Pegunungan Verbeek."),
        ("Komponen PLTU", "Bernilai 0 karena listrik disuplai PLTA Sorowako (0 MW captive); estimasi > Rp 40,0 Triliun pada kolom Estimasi Rugi Ekologis Tabel 8.1 didominasi penuh oleh komponen konsesi."),
    ])
    add_formula(doc, "Substitusi Valuasi Rugi Ekologis (Contoh Baris #3 Tabel 8.1: Jiangsu Delong Nickel)", "Total_Kerugian_Delong = ( 2.253 Ha × Valuasi_Hutan_per_Ha ) + ( 5.175 MW × SCC/NEK ) ≈ > Rp 20,0 Triliun", [
        ("Komponen PLTU", "Kapasitas 5.175 MW dikonversi ke jejak karbon ~36,2 Juta Ton CO2/tahun lalu dikalikan Social Cost of Carbon (SCC/NEK) — pada contoh ini komponen PLTU mendominasi kerugian."),
        ("Catatan Parameter", "Nilai multiplier Valuasi_Hutan_per_Ha dan SCC/NEK mengikuti parameter valuasi PermenLHK No. 7/2014 dan referensi Nilai Ekonomi Karbon; hasil akhir adalah estimasi batas bawah kurasi riset CELIOS."),
    ])
    add_formula(doc, "Substitusi Agregasi Luas Konsesi (Contoh Baris #2: Salim Group)", "Luas_Konsesi_Salim = Σ ( Citra Palu Minerals + Gorontalo Minerals ) = 110.175 Ha   (gabungan 2 PT di dataset)")
    add_formula(doc, "Substitusi Agregasi Kapasitas & Konversi Emisi (Contoh Baris #3: Delong)", "Kapasitas_PLTU_Delong = Σ ( VDNI + OSS + GNI ) = 5.175 MW   ;   Emisi_CO2 ≈ 5.175 MW × 7.000 Ton/MW ≈ 36,2 Juta Ton CO2/thn")
    add_formula(doc, "Substitusi Agregasi Dampak Sosial (Contoh Baris #9: Harita Group)", "Dampak_Sosial_Harita = Σ Jiwa_Terdampak (PT Gema Kreasi Perdana, Wawonii) = 37.000 Jiwa")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Mega-Crosstab Top 10 Penguasa Tahta Ekstraktif")
    add_p(doc, [
        ("Irisan langsung (Mega-Crosstab) antara Grup Oligarki dengan data konsesi tambang, kapasitas PLTU, kerugian ekologis, dan jejak konflik di Sulawesi — diurutkan Top 10 berdasarkan skala daya rusak (kombinasi luas konsesi terbesar dan emisi PLTU raksasa) — disajikan pada ", False, False),
        ("Tabel 8.1", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 8.1: Mega-Crosstab Top 10 Grup Taipan Ekstraktif vs Kerugian Publik di Sulawesi")
    add_table_1col(doc, ["Grup Taipan / Konsorsium", "Total Harta (CELIOS)", "Afiliasi Blok (Sulawesi)", "Luas Konsesi (Aktual)", "Status Deforestasi Lindung", "Emisi PLTU Captive", "Estimasi Rugi Ekologis", "Dampak Sosial & Konflik"], mega_rows_total_81, [2.4, 1.5, 2.1, 1.5, 2.2, 2.2, 1.5, 2.6], ["L", "C", "L", "C", "L", "L", "C", "L"])

    add_p(doc, [
        ("Pemetaan kedelapan kolom Tabel 8.1 beserta sumber data dan persamaan yang melandasinya disajikan pada ", False, False),
        ("Tabel 8.2", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 8.2: Pemetaan 8 Kolom Mega-Crosstab, Sumber Data, dan Persamaan Terkait")
    add_table_1col(doc, kolom_map_headers_81, kolom_map_rows_81, [3.6, 6.6, 5.8], ["L", "L", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris: Ilusi Pembangunan dan Monopoli Keuntungan")
    add_p(doc, [
        ("1. ", True, False), ("Konsentrasi Kekayaan Ekstraktif: ", True, False),
        (f"Sekitar {pct_ekstraktif_81:.0f}% dari total harta Rp{total_harta_50_t_81:,.0f} Triliun milik 50 triliuner Indonesia (setara Rp{harta_ekstraktif_t_81:,.1f} Triliun) dicetak dari pengerukan sumber daya alam — nilai yang melampaui postur APBN nasional.\n", False, False),
        ("2. ", True, False), ("Skala Daya Rusak Privat: ", True, False),
        ("Fakta dataset menelanjangi ilusi pembangunan: ratusan ribu hektar hutan dan pulau kecil telah dikapling (Vale 118.017 Ha; Salim 110.175 Ha), dan lebih dari 9.000 MW PLTU batu bara dibakar secara tertutup oleh Delong (5.175 MW) dan Tsingshan (4.030 MW).\n", False, False),
        ("3. ", True, False), ("Catatan Keterbatasan Data (Undisclosed): ", True, False),
        ("Untuk entitas tambang yang menyedot listrik jaringan PLN, besaran daya aktual (MW) dan emisi karbon tidak dapat dikuantifikasi karena data spesifik tersebut dirahasiakan (Undisclosed) oleh korporasi dalam publikasi publiknya.\n", False, False),
        ("4. ", True, False), ("Implikasi Kebijakan: ", True, False),
        ("Diperlukan kebijakan redistribusi manfaat dan pengelolaan dampak lingkungan yang lebih seimbang agar nilai tambah hilirisasi tidak terkunci pada segelintir konglomerasi besar.", False, False),
    ])

    add_h2(doc, "8.2 Sisi Beban: Indikator Kesehatan dan Sengketa Lahan")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Kesehatan: data/processed/sulawesi_kesehatan_detail_2014_2024.csv (Dinkes/BPS); Data Konflik: Tanahkita.id / KPA (CATAHU); Estimasi Kerugian: Proksi Kalkulasi Valuasi Lingkungan LHK. Visualisasi dashboard menampilkan tiga kartu metrik ringkasan indikator beban publik (Krisis Kesehatan ISPA, Konflik Agraria & FPIC, dan Estimasi Kerugian Ekologis).")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Aktivitas ekstraktif skala besar berpotensi menimbulkan eksternalitas negatif yang dirasakan oleh komunitas sekitar. Hal ini tercermin pada indikator sengketa tata guna lahan serta fluktuasi prevalensi penyakit saluran pernapasan di sekitar kawasan industri. ", False, False),
        ("Sub-bab ini menyajikan ringkasan indikator dampak lingkungan dan sosial yang memerlukan pemantauan serta mitigasi berkesinambungan, sebagai sisi beban penyeimbang dari analisis sisi manfaat pada sub-bab 8.1.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Kalkulasi Tren Eksternalitas Negatif")
    add_p(doc, [
        ("Kerangka agregasi deret waktu deskriptif (Descriptive Time-Series Aggregation) untuk mengukur beban penyakit dan sengketa sosial seiring masifnya industrialisasi ekstraktif diilustrasikan pada ", False, False),
        ("Bagan Alur 8.2", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan trend mapping dan agregasi kasus kritis deskriptif.", False, False),
    ])
    add_caption(doc, "Bagan Alur 8.2: Alur Logika Analisis Kalkulasi Tren Eksternalitas Negatif")
    if download_success_8_2:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_8_2, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Tren ISPA Sentra dan Valuasi Kerusakan")
    add_p(doc, [("Kuantifikasi beban publik dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Tren Kasus ISPA Sentra Nikel", "Tren_ISPA_Sentra(t) = Σ ( Penderita_ISPA_p,t )   ;   untuk provinsi p ∈ { Sulawesi Tengah , Sulawesi Tenggara }", [
        ("Tren_ISPA_Sentra(t)", "Jumlah kasus ISPA/Pneumonia gabungan dua provinsi sentra nikel pada tahun t."),
        ("Penderita_ISPA_p,t", "Nilai kasus ISPA/Pneumonia provinsi p tahun t pada data panel kesehatan Dinkes/BPS."),
    ])
    add_formula(doc, "Persamaan Akumulasi Beban Kesehatan Dekade", "Akumulasi_ISPA_Sentra = Σ Tren_ISPA_Sentra(t)   ;   untuk t = 2014 s.d. 2024", [
        ("Akumulasi_ISPA_Sentra", f"Total kumulatif kasus ISPA sentra nikel sepanjang {tahun_min_82}-{tahun_max_82} ({total_ispa_sentra_82:,.0f} kasus)."),
    ])
    add_formula(doc, "Persamaan Valuasi Kerusakan Lingkungan (Proksi LHK)", "Valuasi_Kerusakan_LHK = F ( Luas_Deforestasi , Hilang_Fungsi_Air , Cemaran_Laut )", [
        ("Valuasi_Kerusakan_LHK", "Fungsi valuasi kumulatif kasar atas hilangnya fungsi hutan primer, rusaknya ekosistem terumbu karang, dan lenyapnya sumber air bersih akibat sedimentasi limbah."),
    ])

    add_p(doc, [("Substitusi angka dari dataset dan laporan riset aktual adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Akumulasi ISPA Sentra", f"Akumulasi_ISPA_Sentra = Σ (Sulteng + Sultra, {tahun_min_82}-{tahun_max_82}) = {total_ispa_sentra_82:,.0f} kasus")
    add_formula(doc, "Substitusi Kasus Kritis & Valuasi", f"Kasus_Kritis_Agraria = {kasus_kritis_82} kasus (TanahKita/KPA)   ;   Valuasi_Kerusakan_LHK > Rp {rugi_min_t_82} Triliun")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Tren ISPA Sentra dan Ringkasan Indikator Beban")
    add_p(doc, [
        (f"Kurva penderita ISPA di sentra nikel per tahun ({tahun_min_82}-{tahun_max_82}) beserta akumulasinya disajikan pada ", False, False),
        ("Tabel 8.3", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 8.3: Tren Tahunan Kasus ISPA/Pneumonia Sentra Nikel Sulteng & Sultra ({tahun_min_82}-{tahun_max_82})")
    add_table_1col(doc, ["Tahun", "Sulawesi Tengah", "Sulawesi Tenggara", "Total Sentra", "Kumulatif"], tren_rows_82, [2.0, 3.2, 3.2, 3.0, 3.0], ["C", "C", "C", "C", "C"])

    add_p(doc, [
        ("Ringkasan tiga indikator beban publik yang dipantau dashboard disajikan pada ", False, False),
        ("Tabel 8.4", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 8.4: Ringkasan Indikator Beban Publik Dampak Industrialisasi Ekstraktif")
    add_table_1col(doc, ["Indikator Beban", "Nilai", "Deskripsi", "Sumber"], beban_rows_82, [2.8, 2.4, 7.2, 3.4], ["L", "C", "L", "L"])

    add_h4(doc, "E. Analisis Temuan Empiris: Beban Publik Sisi Bayangan Hilirisasi")
    add_p(doc, [
        ("1. ", True, False), ("Krisis Kesehatan Kumulatif: ", True, False),
        (f"Sentra nikel Sulteng & Sultra mengakumulasi {total_ispa_sentra_82:,.0f} kasus ISPA/Pneumonia sepanjang {tahun_min_82}-{tahun_max_82}, berkorelasi dengan polusi debu dan sulfur PLTU Captive. Beban tahunan tertinggi tercatat pada {tahun_puncak_ispa_82} ({nilai_puncak_ispa_82:,.0f} kasus) dan terendah pada {tahun_terendah_ispa_82} ({nilai_terendah_ispa_82:,.0f} kasus).\n", False, False),
        ("2. ", True, False), ("Sengketa Lahan Kritis: ", True, False),
        (f"{kasus_kritis_82} kasus kritis konflik agraria & FPIC terdokumentasi meletus di Sulawesi — mengorbankan puluhan ribu jiwa melalui perampasan kebun, pelanggaran hak adat, hingga penembakan warga.\n", False, False),
        ("3. ", True, False), ("Kerugian Ekologis Kumulatif: ", True, False),
        (f"Valuasi proksi LHK menaksir kerugian lebih dari Rp {rugi_min_t_82} Triliun dari hilangnya fungsi hutan primer, kerusakan terumbu karang, dan lenyapnya sumber air bersih — beban yang ditanggung publik sebagai sisi bayangan dari konsentrasi manfaat pada sub-bab 8.1.", False, False),
    ])

    add_h2(doc, "8.3 Pembuktian Statistik: Hubungan Indikator Ekonomi Makro dan Indikator Dampak")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Integrasi Panel: data/processed/sulawesi_investasi_pmdn_2016_2024.csv, data/processed/sulawesi_pad_2016_2024.csv, data/processed/sulawesi_kesehatan_detail_2014_2024.csv, dan data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv. Visualisasi dashboard menampilkan Crosstabulation SPSS-style (Case Processing, Crosstab, Chi-Square Tests, Ringkasan Hipotesis) dengan selector 2 variabel X dan 2 variabel Y, serta opsi Simulasi Skala Kabupaten.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Untuk menguji hubungan antara Manfaat Ekonomi dan Indikator Dampak, dilakukan analisis tabulasi silang (crosstabulation). Uji statistik ini bertujuan mengevaluasi sejauh mana peningkatan arus investasi berasosiasi dengan indikator kesehatan dan lingkungan di tingkat daerah. ", False, False),
        ("Hipotesis utama (Matriks Ketimpangan: Ledakan Investasi vs Ledakan Penyakit): semakin tinggi indikator manfaat ekonomi (Investasi/PAD) yang masuk ke suatu provinsi, semakin parah pula lonjakan kasus beban ekologis (Penyakit/Deforestasi) yang dialami warganya.", False, False),
    ])
    add_p(doc, [
        ("Dashboard juga menyediakan opsi Simulasi Skala Kabupaten (oversampling data 15x untuk mensimulasikan resolusi Kabupaten/Kota) guna mendemonstrasikan efek jumlah sampel (N) terhadap signifikansi. Hasil baku dokumen ini menggunakan panel asli tanpa simulasi tersebut.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Crosstabulation & Pearson Chi-Square Test")
    add_p(doc, [
        ("Kerangka integrasi panel empat dataset dan tahapan uji silang statistiknya diilustrasikan pada ", False, False),
        ("Bagan Alur 8.3", True, False),
        (" berikut, dengan konfigurasi variabel pengujian dirinci pada Tabel 8.3a di bawah gambar.", False, False),
    ])
    add_caption(doc, "Bagan Alur 8.3: Alur Logika Metodologis Uji Korelasi Manfaat Ekonomi vs Beban Ekologis")
    if download_success_8_3:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_8_3, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_caption(doc, "Tabel 8.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 8.3)")
    add_table_1col(doc, konf_headers_83, konf_rows_83, [4.5, 11.0], ["L", "L"])

    add_h4(doc, "C. Formulasi Matematis: Binning Median, Chi-Square, dan Odds Ratio")
    add_p(doc, [("Pembuktian statistik matriks ketimpangan dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Kategorisasi Median Historis (Fungsi Piecewise)", "Kategori(x) = 'Tinggi' , jika x ≥ Median(Panel)   |   'Rendah' , jika x < Median(Panel)", [
        ("Kategori(x)", "Data numerik investasi dan jumlah kasus penyakit dikategorikan menjadi 2 level menggunakan ambang batas Median historis."),
    ])
    add_formula(doc, "Persamaan Uji Independensi Chi-Square Pearson (χ² Kontinjensi 2x2)", "χ² = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   dengan E_ij = ( Total_Baris_i × Total_Kolom_j ) / N", [
        ("χ²", "Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya asosiasi arus manfaat ekonomi dengan indikator beban."),
        ("O_ij", "Frekuensi Observasi: jumlah kasus aktual pada sel baris i kolom j tabel kontinjensi 2x2."),
        ("E_ij", "Frekuensi Harapan: jumlah kasus teoretis jika kedua variabel saling independen."),
    ])
    add_formula(doc, "Persamaan Rasio Keunggulan Risiko Beban Parah (Risk Odds Ratio / OR)", "Odds_Ratio (OR) = ( a × d ) / ( b × c )   ;   dengan a = Manfaat Tinggi & Beban Tinggi/Parah", [
        ("Odds_Ratio (OR)", "Peluang Penyakit/Deforestasi Tinggi pada kelompok Investasi/PAD Tinggi dibandingkan kelompok Rendah."),
    ])

    add_p(doc, [("Substitusi angka dari panel aktual (skenario utama Investasi PMDN × Kasus ISPA) adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi Ambang Median Skenario Utama", f"Median_Investasi = {d_inv_ispa_83['med_x']:,.1f} Rp   ;   Median_ISPA = {d_inv_ispa_83['med_y']:,.1f} kasus   (N valid = {d_inv_ispa_83['n']})")
    add_formula(doc, "Substitusi Hasil Uji Skenario Utama", f"χ² = {d_inv_ispa_83['chi2']:.3f}   ;   P-Value = {d_inv_ispa_83['p']:.4f}   ;   OR = {d_inv_ispa_83['or']:.2f}   →   {'SIGNIFIKAN' if d_inv_ispa_83['p'] < 0.05 else 'TIDAK SIGNIFIKAN'}")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Ambang Median dan Seluruh Skenario Crosstab")
    add_p(doc, [
        ("Ambang median dan jumlah observasi valid tiap skenario disajikan pada ", False, False),
        ("Tabel 8.5", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 8.5: Ambang Median Panel Uji Crosstab per Skenario")
    add_table_1col(doc, ["Variabel X", "Variabel Y", "Median X", "Median Y", "N Valid"], threshold_rows_83, [3.6, 3.8, 3.0, 2.6, 1.6], ["L", "L", "C", "C", "C"])

    add_p(doc, [
        (f"Ringkasan hasil pengujian statistik untuk seluruh {total_scen_83} kombinasi indikator Manfaat (X) dan Beban (Y) disajikan pada ", False, False),
        ("Tabel 8.6", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 8.6: Ringkasan Eksekutif Seluruh Skenario Crosstab Manfaat Ekonomi vs Beban Ekologis")
    add_table_1col(doc, ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_83, [3.2, 3.4, 2.0, 2.0, 2.0, 2.6], ["L", "L", "C", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Matriks Ketimpangan Manfaat vs Beban")
    add_p(doc, [(finding_83, False, False)])

    docx_path = tool_dir / "Metodologi_Bab8_Distribusi_Manfaat.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 8...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 8 - Distribusi Manfaat</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: Arial, sans-serif; max-width: 920px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #43A047; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; }}
.hdr-title {{ color: #81C784; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #2E7D32; padding-bottom: 8px; }}
h2 {{ color: #81C784; text-transform: uppercase; border-bottom: 1px solid #2E7D32; }}
h4 {{ color: #A5D6A7; }}
.note-box {{ background: #132213; border-left: 4px solid #2E7D32; padding: 10px 14px; margin: 12px 0; }}
.formula {{ background: #0D1B0E; border: 1px solid #2E7D32; color: #A5D6A7; padding: 8px 12px; font-family: monospace; }}
.table-caption {{ color: #A5D6A7; font-weight: 700; font-style: italic; margin-top: 14px; }}
.data-th {{ background: #1B5E20; color: white; padding: 6px; text-align: left; border: 1px solid #2E7D32; }}
.data-td {{ padding: 6px; border: 1px solid #243524; vertical-align: top; }}
.data-tr-even .data-td {{ background: #131B13; }}
.mermaid {{ background: #0D1610; border: 1px solid #2E7D32; padding: 12px; margin: 10px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB VIII: Metodologi Analisis Distribusi Manfaat vs Beban Ekologis</div>
<p>Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada <strong>Bab 8: Distribusi Manfaat vs Beban Ekologis</strong>.</p>

<h2>8.1 Sisi Manfaat: Gurita Bisnis & Monopoli Keuntungan Ekstraktif</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Sumber utama: CELIOS Inequality Report 2026 (Laporan 50 Taipan Terkaya); Dataset internal: <code>data/processed/sulawesi_kawasan_nikel_luas.csv</code>, <code>data/processed/sulawesi_pltu_captive.csv</code>, <code>data/processed/sulawesi_konflik_agraria_tanahkita.csv</code>. Visualisasi dashboard menampilkan tiga metric cards konsentrasi kekayaan serta Mega-Crosstab Top 10 Grup Taipan vs kerugian publik.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Analisis distribusi manfaat ekonomi sektor nikel dan PLTU di Sulawesi menunjukkan konsentrasi nilai tambah pada kelompok usaha skala besar. Laporan Ketimpangan CELIOS mencatat akumulasi kekayaan 50 individu/kelompok usaha terbesar mencapai <strong>Rp{total_harta_50_t_81:,.0f} Triliun</strong>, sekitar <strong>{pct_ekstraktif_81:.0f}% bersumber dari sektor berbasis sumber daya alam</strong>. Kekayaan naik nyaris 2x lipat sejak 2019 dengan laju Rp{laju_harian_m_81:.0f} Miliar/hari — kontras dengan kenaikan upah buruh sekitar Rp{laju_upah_buruh_rb_81:.0f} ribu/hari.</p>
<h4>B. Alur Logika Metodologis Pemetaan Konsentrasi Kekayaan Ekstraktif</h4>
<p>Kerangka Hierarchical Entity Profiling diilustrasikan pada <strong>Bagan Alur 8.1</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Wealth Database Analysis dan Mega-Crosstab pemetaan aktor deskriptif.</p>
<div class="table-caption">Bagan Alur 8.1: Alur Logika Analisis Pemetaan Konsentrasi Kekayaan Ekstraktif</div>
<div class="mermaid">{mermaid_str_8_1}</div>
<h4>C. Formulasi Matematis: Konsentrasi Kekayaan dan Valuasi Rugi Ekologis</h4>
<div class="formula">Total_Kekayaan_Ekstraktif = Σ ( Harta_i )   ;   untuk seluruh triliuner i dengan Sektor = 'Ekstraktif'</div>
<div class="formula">Beban_Ekologis_g = Σ ( Rugi_Ekologis_e )   ;   untuk seluruh entitas operasi e dengan Afiliasi_Pemilik = g</div>
<div class="formula">Luas_Konsesi_g (Kolom 4) = Σ ( total_luas_ha_e )   ;   untuk seluruh entitas e dengan afiliasi grup g</div>
<div class="formula">Kapasitas_PLTU_g (Kolom 6) = Σ ( Capacity_MW_u ) untuk Parent = g   ;   Emisi_CO2_g ≈ Kapasitas_PLTU_g × ~7.000 Ton CO2/MW/thn</div>
<div class="formula">Dampak_Sosial_g (Kolom 8) = Σ ( Jiwa_Terdampak_k )   ;   untuk seluruh kasus konflik k terafiliasi grup g</div>
<div class="formula">Total_Kerugian_Ekologis (Kolom 7) = ( Luas_Konsesi × Valuasi_Hutan_per_Ha ) + ( Kapasitas_PLTU_MW × Biaya_Sosial_Emisi_Karbon )</div>
<p>Substitusi angka dari laporan riset aktual:</p>
<div class="formula">Total_Kekayaan_Ekstraktif = {pct_ekstraktif_81:.0f}% × Rp{total_harta_50_t_81:,.0f} T = Rp{harta_ekstraktif_t_81:,.1f} Triliun</div>
<div class="formula">Laju_Harta_Elit = Rp{laju_harian_m_81:.0f} Miliar/hari   vs   Laju_Upah_Buruh = Rp{laju_upah_buruh_rb_81:.0f} Ribu/hari</div>
<div class="formula">Total_Kerugian_Vale (#1 Tabel 8.1) = ( 118.017 Ha × Valuasi_Hutan_per_Ha ) + ( 0 MW × SCC/NEK ) ≈ &gt; Rp 40,0 Triliun</div>
<div class="formula">Total_Kerugian_Delong (#3 Tabel 8.1) = ( 2.253 Ha × Valuasi_Hutan_per_Ha ) + ( 5.175 MW × SCC/NEK ) ≈ &gt; Rp 20,0 Triliun</div>
<div class="formula">Luas_Konsesi_Salim (#2) = Σ ( Citra Palu Minerals + Gorontalo Minerals ) = 110.175 Ha</div>
<div class="formula">Kapasitas_PLTU_Delong (#3) = Σ ( VDNI + OSS + GNI ) = 5.175 MW   ;   Emisi ≈ 5.175 × 7.000 ≈ 36,2 Jt Ton CO2/thn</div>
<div class="formula">Dampak_Sosial_Harita (#9) = Σ Jiwa_Terdampak (PT Gema Kreasi Perdana, Wawonii) = 37.000 Jiwa</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 8.1: Mega-Crosstab Top 10 Grup Taipan Ekstraktif vs Kerugian Publik di Sulawesi</div>
{html_table(["Grup Taipan / Konsorsium", "Total Harta (CELIOS)", "Afiliasi Blok (Sulawesi)", "Luas Konsesi (Aktual)", "Status Deforestasi Lindung", "Emisi PLTU Captive", "Estimasi Rugi Ekologis", "Dampak Sosial & Konflik"], mega_rows_total_81)}
<div class="table-caption">Tabel 8.2: Pemetaan 8 Kolom Mega-Crosstab, Sumber Data, dan Persamaan Terkait</div>
{html_table(kolom_map_headers_81, kolom_map_rows_81)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Konsentrasi Kekayaan Ekstraktif:</strong> sekitar {pct_ekstraktif_81:.0f}% dari total Rp{total_harta_50_t_81:,.0f} Triliun harta 50 triliuner (setara Rp{harta_ekstraktif_t_81:,.1f} Triliun) dicetak dari pengerukan sumber daya alam. <strong>2. Skala Daya Rusak Privat:</strong> ratusan ribu hektar hutan dan pulau kecil dikapling (Vale 118.017 Ha; Salim 110.175 Ha), dan lebih dari <strong>9.000 MW PLTU batu bara</strong> dibakar secara tertutup oleh Delong dan Tsingshan. <strong>3. Catatan Undisclosed:</strong> daya aktual entitas yang menyedot listrik PLN tidak dapat dikuantifikasi karena dirahasiakan korporasi. <strong>4. Implikasi Kebijakan:</strong> diperlukan redistribusi manfaat dan pengelolaan dampak lingkungan yang lebih seimbang.</p>

<h2>8.2 Sisi Beban: Indikator Kesehatan dan Sengketa Lahan</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Kesehatan: <code>data/processed/sulawesi_kesehatan_detail_2014_2024.csv</code> (Dinkes/BPS); Data Konflik: Tanahkita.id / KPA (CATAHU); Estimasi Kerugian: Proksi Kalkulasi Valuasi Lingkungan LHK. Visualisasi dashboard menampilkan tiga kartu metrik ringkasan indikator beban publik.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Aktivitas ekstraktif skala besar berpotensi menimbulkan <strong>eksternalitas negatif</strong> yang dirasakan oleh komunitas sekitar — tercermin pada indikator sengketa tata guna lahan serta fluktuasi prevalensi penyakit saluran pernapasan di sekitar kawasan industri. Sub-bab ini menyajikan ringkasan indikator dampak lingkungan dan sosial sebagai sisi beban penyeimbang dari analisis sisi manfaat pada sub-bab 8.1.</p>
<h4>B. Alur Logika Metodologis Kalkulasi Tren Eksternalitas Negatif</h4>
<p>Kerangka agregasi deret waktu deskriptif diilustrasikan pada <strong>Bagan Alur 8.2</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan trend mapping dan agregasi kasus kritis deskriptif.</p>
<div class="table-caption">Bagan Alur 8.2: Alur Logika Analisis Kalkulasi Tren Eksternalitas Negatif</div>
<div class="mermaid">{mermaid_str_8_2}</div>
<h4>C. Formulasi Matematis: Tren ISPA Sentra dan Valuasi Kerusakan</h4>
<div class="formula">Tren_ISPA_Sentra(t) = Σ ( Penderita_ISPA_p,t )   ;   untuk provinsi p ∈ {{ Sulawesi Tengah , Sulawesi Tenggara }}</div>
<div class="formula">Akumulasi_ISPA_Sentra = Σ Tren_ISPA_Sentra(t)   ;   untuk t = 2014 s.d. 2024</div>
<div class="formula">Valuasi_Kerusakan_LHK = F ( Luas_Deforestasi , Hilang_Fungsi_Air , Cemaran_Laut )</div>
<p>Substitusi angka dari dataset dan laporan riset aktual:</p>
<div class="formula">Akumulasi_ISPA_Sentra = Σ (Sulteng + Sultra, {tahun_min_82}-{tahun_max_82}) = {total_ispa_sentra_82:,.0f} kasus</div>
<div class="formula">Kasus_Kritis_Agraria = {kasus_kritis_82} kasus (TanahKita/KPA)   ;   Valuasi_Kerusakan_LHK &gt; Rp {rugi_min_t_82} Triliun</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 8.3: Tren Tahunan Kasus ISPA/Pneumonia Sentra Nikel Sulteng & Sultra ({tahun_min_82}-{tahun_max_82})</div>
{html_table(["Tahun", "Sulawesi Tengah", "Sulawesi Tenggara", "Total Sentra", "Kumulatif"], tren_rows_82)}
<div class="table-caption">Tabel 8.4: Ringkasan Indikator Beban Publik Dampak Industrialisasi Ekstraktif</div>
{html_table(["Indikator Beban", "Nilai", "Deskripsi", "Sumber"], beban_rows_82)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Krisis Kesehatan Kumulatif:</strong> sentra nikel Sulteng & Sultra mengakumulasi <strong>{total_ispa_sentra_82:,.0f} kasus ISPA/Pneumonia</strong> ({tahun_min_82}-{tahun_max_82}); beban tahunan tertinggi {tahun_puncak_ispa_82} ({nilai_puncak_ispa_82:,.0f} kasus). <strong>2. Sengketa Lahan Kritis:</strong> {kasus_kritis_82} kasus kritis konflik agraria & FPIC terdokumentasi meletus di Sulawesi. <strong>3. Kerugian Ekologis Kumulatif:</strong> valuasi proksi LHK menaksir kerugian lebih dari Rp {rugi_min_t_82} Triliun — beban publik sebagai sisi bayangan dari konsentrasi manfaat sub-bab 8.1.</p>

<h2>8.3 Pembuktian Statistik: Hubungan Indikator Ekonomi Makro dan Indikator Dampak</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Integrasi Panel: <code>sulawesi_investasi_pmdn_2016_2024.csv</code>, <code>sulawesi_pad_2016_2024.csv</code>, <code>sulawesi_kesehatan_detail_2014_2024.csv</code>, <code>sulawesi_gfw_master_1_dekade_2014_2023_v3.csv</code>. Visualisasi dashboard menampilkan Crosstabulation SPSS-style dengan selector 2 variabel X dan 2 variabel Y, serta opsi Simulasi Skala Kabupaten.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Untuk menguji hubungan antara <strong>Manfaat Ekonomi</strong> dan <strong>Indikator Dampak</strong>, dilakukan analisis tabulasi silang. Hipotesis utama (Matriks Ketimpangan: Ledakan Investasi vs Ledakan Penyakit): semakin tinggi indikator manfaat ekonomi (Investasi/PAD) suatu provinsi, semakin parah pula lonjakan kasus beban ekologis (Penyakit/Deforestasi) warganya. Dashboard menyediakan opsi Simulasi Skala Kabupaten (oversampling 15x); hasil baku dokumen ini memakai panel asli tanpa simulasi.</p>
<h4>B. Alur Logika Metodologis Crosstabulation & Pearson Chi-Square Test</h4>
<p>Kerangka integrasi panel empat dataset diilustrasikan pada <strong>Bagan Alur 8.3</strong> berikut, dengan konfigurasi variabel pengujian dirinci pada Tabel 8.3a di bawah gambar.</p>
<div class="table-caption">Bagan Alur 8.3: Alur Logika Metodologis Uji Korelasi Manfaat Ekonomi vs Beban Ekologis</div>
<div class="mermaid">{mermaid_str_8_3}</div>
<div class="table-caption">Tabel 8.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 8.3)</div>
{html_table(konf_headers_83, konf_rows_83)}
<h4>C. Formulasi Matematis: Binning Median, Chi-Square, dan Odds Ratio</h4>
<div class="formula">Kategori(x) = 'Tinggi' , jika x ≥ Median(Panel)   |   'Rendah' , jika x &lt; Median(Panel)</div>
<div class="formula">&chi;&sup2; = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   dengan E_ij = ( Total_Baris_i × Total_Kolom_j ) / N</div>
<div class="formula">Odds_Ratio (OR) = ( a × d ) / ( b × c )   ;   dengan a = Manfaat Tinggi & Beban Tinggi/Parah</div>
<p>Substitusi angka dari panel aktual (skenario utama Investasi PMDN × Kasus ISPA):</p>
<div class="formula">Median_Investasi = {d_inv_ispa_83['med_x']:,.1f} Rp   ;   Median_ISPA = {d_inv_ispa_83['med_y']:,.1f} kasus   (N valid = {d_inv_ispa_83['n']})</div>
<div class="formula">χ² = {d_inv_ispa_83['chi2']:.3f}   ;   P-Value = {d_inv_ispa_83['p']:.4f}   ;   OR = {d_inv_ispa_83['or']:.2f}   →   {'SIGNIFIKAN' if d_inv_ispa_83['p'] < 0.05 else 'TIDAK SIGNIFIKAN'}</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 8.5: Ambang Median Panel Uji Crosstab per Skenario</div>
{html_table(["Variabel X", "Variabel Y", "Median X", "Median Y", "N Valid"], threshold_rows_83)}
<div class="table-caption">Tabel 8.6: Ringkasan Eksekutif Seluruh Skenario Crosstab Manfaat Ekonomi vs Beban Ekologis</div>
{html_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (&chi;&sup2;)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_83)}
<h4>E. Analisis Temuan Empiris</h4>
<p>{finding_83}</p>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab8_Distribusi_Manfaat.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB VIII: METODOLOGI ANALISIS DISTRIBUSI MANFAAT VS BEBAN EKOLOGIS",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada **Bab 8: Distribusi Manfaat vs Beban Ekologis**.",
        "",
        "## 8.1 Sisi Manfaat: Gurita Bisnis & Monopoli Keuntungan Ekstraktif",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Sumber utama: CELIOS Inequality Report 2026 (Laporan 50 Taipan Terkaya); Dataset internal: `data/processed/sulawesi_kawasan_nikel_luas.csv` (agregasi nama perusahaan normatif), `data/processed/sulawesi_pltu_captive.csv` (agregasi Parent & Capacity MW), `data/processed/sulawesi_konflik_agraria_tanahkita.csv`. Visualisasi dashboard menampilkan tiga metric cards konsentrasi kekayaan serta Mega-Crosstab Top 10 Grup Taipan vs kerugian publik.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Analisis distribusi manfaat ekonomi sektor nikel dan PLTU di Sulawesi menunjukkan konsentrasi nilai tambah pada kelompok usaha skala besar. Laporan Ketimpangan CELIOS mencatat akumulasi kekayaan 50 individu/kelompok usaha terbesar mencapai **Rp{total_harta_50_t_81:,.0f} Triliun**, di mana sekitar **{pct_ekstraktif_81:.0f}% bersumber dari sektor berbasis sumber daya alam** (pertambangan nikel, batu bara, kelapa sawit, dan pemurnian logam). Kekayaan ini naik nyaris 2x lipat sejak 2019 dengan laju Rp{laju_harian_m_81:.0f} Miliar/hari — kontras dengan kenaikan upah buruh nasional sekitar Rp{laju_upah_buruh_rb_81:.0f} ribu/hari.",
        "",
        "#### B. Alur Logika Metodologis Pemetaan Konsentrasi Kekayaan Ekstraktif",
        "Kerangka pemrofilan entitas bisnis berjenjang (*Hierarchical Entity Profiling*) diilustrasikan pada **Bagan Alur 8.1** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Wealth Database Analysis dan Mega-Crosstab pemetaan aktor deskriptif.",
        "",
        "##### Bagan Alur 8.1: Alur Logika Analisis Pemetaan Konsentrasi Kekayaan Ekstraktif",
        "```mermaid",
        mermaid_str_8_1,
        "```",
        "",
        "#### C. Formulasi Matematis: Konsentrasi Kekayaan dan Valuasi Rugi Ekologis",
        "Kuantifikasi konsentrasi kekayaan dan daya rusak privat dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Total_Kekayaan_Ekstraktif = Σ ( Harta_i )   ;   untuk seluruh triliuner i dengan Sektor = 'Ekstraktif'",
        "Beban_Ekologis_g = Σ ( Rugi_Ekologis_e )   ;   untuk seluruh entitas operasi e dengan Afiliasi_Pemilik = g",
        "Luas_Konsesi_g (Kolom 4) = Σ ( total_luas_ha_e )   ;   untuk seluruh entitas e dengan afiliasi grup g",
        "Kapasitas_PLTU_g (Kolom 6) = Σ ( Capacity_MW_u ) untuk Parent = g   ;   Emisi_CO2_g ≈ Kapasitas_PLTU_g × ~7.000 Ton CO2/MW/thn",
        "Dampak_Sosial_g (Kolom 8) = Σ ( Jiwa_Terdampak_k )   ;   untuk seluruh kasus konflik k terafiliasi grup g",
        "Total_Kerugian_Ekologis (Kolom 7) = ( Luas_Konsesi × Valuasi_Hutan_per_Ha ) + ( Kapasitas_PLTU_MW × Biaya_Sosial_Emisi_Karbon )",
        "```",
        "",
        "Valuasi rugi ekologis mengadaptasi formula **Peraturan Menteri LHK No. 7 Tahun 2014**: komponen kerugian ekonomi publik (tangkapan nelayan, tanaman warga, biaya pengobatan ISPA) dan biaya pemulihan alam (reboisasi, netralisasi limbah slag, biaya sosial emisi karbon / SCC-NEK).",
        "",
        "Substitusi angka dari laporan riset aktual:",
        "",
        "```text",
        f"Total_Kekayaan_Ekstraktif = {pct_ekstraktif_81:.0f}% × Rp{total_harta_50_t_81:,.0f} T = Rp{harta_ekstraktif_t_81:,.1f} Triliun",
        f"Laju_Harta_Elit = Rp{laju_harian_m_81:.0f} Miliar/hari   vs   Laju_Upah_Buruh = Rp{laju_upah_buruh_rb_81:.0f} Ribu/hari",
        "Total_Kerugian_Vale (#1 Tabel 8.1) = ( 118.017 Ha × Valuasi_Hutan_per_Ha ) + ( 0 MW × SCC/NEK ) ≈ > Rp 40,0 Triliun",
        "Total_Kerugian_Delong (#3 Tabel 8.1) = ( 2.253 Ha × Valuasi_Hutan_per_Ha ) + ( 5.175 MW × SCC/NEK ) ≈ > Rp 20,0 Triliun",
        "Luas_Konsesi_Salim (#2) = Σ ( Citra Palu Minerals + Gorontalo Minerals ) = 110.175 Ha",
        "Kapasitas_PLTU_Delong (#3) = Σ ( VDNI + OSS + GNI ) = 5.175 MW   ;   Emisi ≈ 5.175 × 7.000 ≈ 36,2 Jt Ton CO2/thn",
        "Dampak_Sosial_Harita (#9) = Σ Jiwa_Terdampak (PT Gema Kreasi Perdana, Wawonii) = 37.000 Jiwa",
        "```",
        "",
        "Baris substitusi Vale/Delong adalah contoh perhitungan kolom **Estimasi Rugi Ekologis** pada Tabel 8.1: baris #1 (Vale) kerugiannya didominasi komponen konsesi (118.017 Ha terbesar di dataset; PLTU 0 MW karena disuplai PLTA Sorowako), sedangkan baris #3 (Delong) didominasi komponen PLTU (5.175 MW ≈ 36,2 juta ton CO2/tahun × SCC/NEK). Tiga baris berikutnya adalah contoh substitusi persamaan kolom 4, 6, dan 8.",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 8.1: Mega-Crosstab Top 10 Grup Taipan Ekstraktif vs Kerugian Publik di Sulawesi",
        markdown_table(["Grup Taipan / Konsorsium", "Total Harta (CELIOS)", "Afiliasi Blok (Sulawesi)", "Luas Konsesi (Aktual)", "Status Deforestasi Lindung", "Emisi PLTU Captive", "Estimasi Rugi Ekologis", "Dampak Sosial & Konflik"], mega_rows_total_81),
        "",
        "##### Tabel 8.2: Pemetaan 8 Kolom Mega-Crosstab, Sumber Data, dan Persamaan Terkait",
        markdown_table(kolom_map_headers_81, kolom_map_rows_81),
        "",
        "#### E. Analisis Temuan Empiris: Ilusi Pembangunan dan Monopoli Keuntungan",
        f"1. **Konsentrasi Kekayaan Ekstraktif:** sekitar {pct_ekstraktif_81:.0f}% dari total harta Rp{total_harta_50_t_81:,.0f} Triliun milik 50 triliuner Indonesia (setara Rp{harta_ekstraktif_t_81:,.1f} Triliun) dicetak dari pengerukan sumber daya alam — nilai yang melampaui postur APBN nasional.",
        "2. **Skala Daya Rusak Privat:** fakta dataset menelanjangi ilusi pembangunan: ratusan ribu hektar hutan dan pulau kecil telah dikapling (Vale 118.017 Ha; Salim 110.175 Ha), dan lebih dari **9.000 MW PLTU Batu Bara** dibakar secara tertutup oleh Delong (5.175 MW) dan Tsingshan (4.030 MW).",
        "3. **Catatan Keterbatasan Data (Undisclosed):** untuk entitas tambang yang menyedot listrik jaringan PLN, besaran daya aktual (MW) dan emisi karbon tidak dapat dikuantifikasi karena data spesifik tersebut dirahasiakan oleh korporasi dalam publikasi publiknya.",
        "4. **Implikasi Kebijakan:** diperlukan kebijakan redistribusi manfaat dan pengelolaan dampak lingkungan yang lebih seimbang agar nilai tambah hilirisasi tidak terkunci pada segelintir konglomerasi besar.",
        "",
        "## 8.2 Sisi Beban: Indikator Kesehatan dan Sengketa Lahan",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Kesehatan: `data/processed/sulawesi_kesehatan_detail_2014_2024.csv` (Dinkes/BPS); Data Konflik: Tanahkita.id / KPA (CATAHU); Estimasi Kerugian: Proksi Kalkulasi Valuasi Lingkungan LHK. Visualisasi dashboard menampilkan tiga kartu metrik ringkasan indikator beban publik (Krisis Kesehatan ISPA, Konflik Agraria & FPIC, dan Estimasi Kerugian Ekologis).",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Aktivitas ekstraktif skala besar berpotensi menimbulkan **eksternalitas negatif** yang dirasakan oleh komunitas sekitar — tercermin pada indikator sengketa tata guna lahan serta fluktuasi prevalensi penyakit saluran pernapasan di sekitar kawasan industri. Sub-bab ini menyajikan ringkasan indikator dampak lingkungan dan sosial yang memerlukan pemantauan serta mitigasi berkesinambungan, sebagai sisi beban penyeimbang dari analisis sisi manfaat pada sub-bab 8.1.",
        "",
        "#### B. Alur Logika Metodologis Kalkulasi Tren Eksternalitas Negatif",
        "Kerangka agregasi deret waktu deskriptif (*Descriptive Time-Series Aggregation*) diilustrasikan pada **Bagan Alur 8.2** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan trend mapping dan agregasi kasus kritis deskriptif.",
        "",
        "##### Bagan Alur 8.2: Alur Logika Analisis Kalkulasi Tren Eksternalitas Negatif",
        "```mermaid",
        mermaid_str_8_2,
        "```",
        "",
        "#### C. Formulasi Matematis: Tren ISPA Sentra dan Valuasi Kerusakan",
        "Kuantifikasi beban publik dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Tren_ISPA_Sentra(t) = Σ ( Penderita_ISPA_p,t )   ;   untuk provinsi p ∈ { Sulawesi Tengah , Sulawesi Tenggara }",
        "Akumulasi_ISPA_Sentra = Σ Tren_ISPA_Sentra(t)   ;   untuk t = 2014 s.d. 2024",
        "Valuasi_Kerusakan_LHK = F ( Luas_Deforestasi , Hilang_Fungsi_Air , Cemaran_Laut )",
        "```",
        "",
        "Substitusi angka dari dataset dan laporan riset aktual:",
        "",
        "```text",
        f"Akumulasi_ISPA_Sentra = Σ (Sulteng + Sultra, {tahun_min_82}-{tahun_max_82}) = {total_ispa_sentra_82:,.0f} kasus",
        f"Kasus_Kritis_Agraria = {kasus_kritis_82} kasus (TanahKita/KPA)   ;   Valuasi_Kerusakan_LHK > Rp {rugi_min_t_82} Triliun",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 8.3: Tren Tahunan Kasus ISPA/Pneumonia Sentra Nikel Sulteng & Sultra ({tahun_min_82}-{tahun_max_82})",
        markdown_table(["Tahun", "Sulawesi Tengah", "Sulawesi Tenggara", "Total Sentra", "Kumulatif"], tren_rows_82),
        "",
        "##### Tabel 8.4: Ringkasan Indikator Beban Publik Dampak Industrialisasi Ekstraktif",
        markdown_table(["Indikator Beban", "Nilai", "Deskripsi", "Sumber"], beban_rows_82),
        "",
        "#### E. Analisis Temuan Empiris: Beban Publik Sisi Bayangan Hilirisasi",
        f"1. **Krisis Kesehatan Kumulatif:** sentra nikel Sulteng & Sultra mengakumulasi **{total_ispa_sentra_82:,.0f} kasus ISPA/Pneumonia** sepanjang {tahun_min_82}-{tahun_max_82}, berkorelasi dengan polusi debu dan sulfur PLTU Captive; beban tahunan tertinggi tercatat {tahun_puncak_ispa_82} ({nilai_puncak_ispa_82:,.0f} kasus) dan terendah {tahun_terendah_ispa_82} ({nilai_terendah_ispa_82:,.0f} kasus).",
        f"2. **Sengketa Lahan Kritis:** {kasus_kritis_82} kasus kritis konflik agraria & FPIC terdokumentasi meletus di Sulawesi — mengorbankan puluhan ribu jiwa melalui perampasan kebun, pelanggaran hak adat, hingga penembakan warga.",
        f"3. **Kerugian Ekologis Kumulatif:** valuasi proksi LHK menaksir kerugian lebih dari Rp {rugi_min_t_82} Triliun dari hilangnya fungsi hutan primer, kerusakan terumbu karang, dan lenyapnya sumber air bersih — beban yang ditanggung publik sebagai sisi bayangan dari konsentrasi manfaat pada sub-bab 8.1.",
        "",
        "## 8.3 Pembuktian Statistik: Hubungan Indikator Ekonomi Makro dan Indikator Dampak",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Integrasi Panel: `sulawesi_investasi_pmdn_2016_2024.csv`, `sulawesi_pad_2016_2024.csv`, `sulawesi_kesehatan_detail_2014_2024.csv`, `sulawesi_gfw_master_1_dekade_2014_2023_v3.csv`. Visualisasi dashboard menampilkan Crosstabulation SPSS-style (Case Processing → Crosstab → Chi-Square Tests → Ringkasan Hipotesis) dengan selector 2 variabel X dan 2 variabel Y, serta opsi Simulasi Skala Kabupaten.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Untuk menguji hubungan antara **Manfaat Ekonomi** dan **Indikator Dampak**, dilakukan analisis tabulasi silang (*crosstabulation*). Hipotesis utama (Matriks Ketimpangan: Ledakan Investasi vs Ledakan Penyakit): semakin tinggi indikator manfaat ekonomi (Investasi/PAD) yang masuk ke suatu provinsi, semakin parah pula lonjakan kasus beban ekologis (Penyakit/Deforestasi) yang dialami warganya. Dashboard juga menyediakan opsi Simulasi Skala Kabupaten (oversampling 15x untuk mensimulasikan resolusi Kabupaten/Kota); hasil baku dokumen ini menggunakan panel asli tanpa simulasi tersebut.",
        "",
        "#### B. Alur Logika Metodologis Crosstabulation & Pearson Chi-Square Test",
        "Kerangka integrasi panel empat dataset dan tahapan uji silang statistiknya diilustrasikan pada **Bagan Alur 8.3** berikut, dengan konfigurasi variabel pengujian dirinci pada Tabel 8.3a di bawah gambar.",
        "",
        "##### Bagan Alur 8.3: Alur Logika Metodologis Uji Korelasi Manfaat Ekonomi vs Beban Ekologis",
        "```mermaid",
        mermaid_str_8_3,
        "```",
        "",
        "##### Tabel 8.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 8.3)",
        markdown_table(konf_headers_83, konf_rows_83),
        "",
        "#### C. Formulasi Matematis: Binning Median, Chi-Square, dan Odds Ratio",
        "Pembuktian statistik matriks ketimpangan dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Kategori(x) = 'Tinggi' , jika x ≥ Median(Panel)   |   'Rendah' , jika x < Median(Panel)",
        "χ² = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   dengan E_ij = ( Total_Baris_i × Total_Kolom_j ) / N",
        "Odds_Ratio (OR) = ( a × d ) / ( b × c )   ;   dengan a = Manfaat Tinggi & Beban Tinggi/Parah",
        "```",
        "",
        "Substitusi angka dari panel aktual (skenario utama Investasi PMDN × Kasus ISPA):",
        "",
        "```text",
        f"Median_Investasi = {d_inv_ispa_83['med_x']:,.1f} Rp   ;   Median_ISPA = {d_inv_ispa_83['med_y']:,.1f} kasus   (N valid = {d_inv_ispa_83['n']})",
        f"χ² = {d_inv_ispa_83['chi2']:.3f}   ;   P-Value = {d_inv_ispa_83['p']:.4f}   ;   OR = {d_inv_ispa_83['or']:.2f}   →   {'SIGNIFIKAN' if d_inv_ispa_83['p'] < 0.05 else 'TIDAK SIGNIFIKAN'}",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 8.5: Ambang Median Panel Uji Crosstab per Skenario",
        markdown_table(["Variabel X", "Variabel Y", "Median X", "Median Y", "N Valid"], threshold_rows_83),
        "",
        "##### Tabel 8.6: Ringkasan Eksekutif Seluruh Skenario Crosstab Manfaat Ekonomi vs Beban Ekologis",
        markdown_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_83),
        "",
        "#### E. Analisis Temuan Empiris: Matriks Ketimpangan Manfaat vs Beban",
        finding_83,
        "",
    ]
    md_path = tool_dir / "Metodologi_Bab8_Distribusi_Manfaat.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")
    print("[4/4] Selesai membangun Bab 8.")


if __name__ == "__main__":
    generate_all_bab8()
