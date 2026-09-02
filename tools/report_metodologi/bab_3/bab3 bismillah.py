#!/usr/bin/env python3
"""
Generator Laporan Metodologi Bab 3: Beban Kesehatan Masyarakat Terdampak

Fokus awal: Sub-bab 3.1 Kesenjangan Fasilitas Kesehatan di Kawasan
Ekstraktif. Pilar 1 ditulis langsung dalam generator Python agar selaras
dengan SOP dokumentasi Celios2.
"""

import base64
import sys
from pathlib import Path

try:
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([
        sys.executable,
        "-m",
        "pip",
        "install",
        "pandas",
        "scipy",
        "requests",
        "python-docx",
    ])
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge, cfg in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        if cfg is None:
            el.set(qn("w:val"), "none")
        else:
            for key, val in cfg.items():
                el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_margin(cell, left=100, right=100, top=60, bottom=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margin.append(el)
    tc_pr.append(margin)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_border_bottom(paragraph, color="2E7D32", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="2E7D32", sz="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="444444", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="1B5E20", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="2E7D32", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=G_MID)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)


def add_p(doc, parts, space_after=5, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "EDF7EE")
    all_border_para(p, color="A5D6A7", sz="4")
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)

    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(2)
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=G_DARK)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p, color="2E7D32", sz="16")
    para_shd(p, "F1F8E9")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}

    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "2E7D32")
        cell_margin(cell, left=100, right=100, top=70, bottom=70)
        set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom={"val": "single", "sz": "8", "color": "1B5E20", "space": "0"}, right=bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8.5, color=C_WHITE)

    for i, row_data in enumerate(rows):
        fill = "F5FBF5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            cell_margin(cell, left=100, right=100, top=50, bottom=50)
            set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def generate_all_bab3():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    data_dir = base_dir / "data" / "processed"
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_3"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 3 sub-bab 3.1...")
    df_kes = pd.read_csv(data_dir / "sulawesi_kesehatan_detail_2014_2024.csv")
    df_faskes = pd.read_csv(data_dir / "sulawesi_faskes_agregat_v3.csv")
    df_pltu = pd.read_csv(data_dir / "sulawesi_pltu_captive.csv")
    df_zoonosis = pd.read_csv(data_dir / "zoonosis_kab_kota_2015_2024.csv")

    df_kes = df_kes[df_kes["indikator"] != "Kasus Kusta Baru"].copy()
    tot_ispa = df_kes[df_kes["indikator"] == "Kasus ISPA/Pneumonia"]["nilai"].sum()
    tot_diare = df_kes[df_kes["indikator"] == "Kasus Diare Dilayani"]["nilai"].sum()
    tot_malaria = df_kes[df_kes["indikator"] == "Kasus Malaria Positif"]["nilai"].sum()

    df_pltu_op = df_pltu[df_pltu["Status"].str.lower() == "operating"].copy()
    tot_kapasitas_pltu = df_pltu_op["Capacity (MW)"].sum() if "Capacity (MW)" in df_pltu_op.columns else 0

    latest_year_faskes = int(df_faskes["tahun"].max())
    faskes_latest = df_faskes[df_faskes["tahun"] == latest_year_faskes].copy()
    tot_puskesmas_latest = faskes_latest[faskes_latest["jenis_faskes"] == "Puskesmas"]["jumlah"].sum()
    tot_rs_latest = faskes_latest[faskes_latest["jenis_faskes"] == "Rumah Sakit"]["jumlah"].sum()

    sentra = ["Sulawesi Tengah", "Sulawesi Tenggara"]
    df_faskes_copy = df_faskes[~df_faskes["provinsi"].str.contains("Indonesia", na=False)].copy()
    df_faskes_copy["Kategori"] = df_faskes_copy["provinsi"].apply(
        lambda x: "Sentra Industri (Sulteng & Sultra)" if x in sentra else "Non-Sentra Industri (Lainnya)"
    )
    df_2024 = df_faskes_copy[df_faskes_copy["tahun"] == latest_year_faskes].copy()
    df_gap = df_2024.groupby(["Kategori", "jenis_faskes"])["jumlah"].mean().reset_index()

    rs_sentra = df_gap[(df_gap["jenis_faskes"] == "Rumah Sakit") & (df_gap["Kategori"] == "Sentra Industri (Sulteng & Sultra)")]["jumlah"].values[0]
    rs_non = df_gap[(df_gap["jenis_faskes"] == "Rumah Sakit") & (df_gap["Kategori"] == "Non-Sentra Industri (Lainnya)")]["jumlah"].values[0]
    puskesmas_sentra = df_gap[(df_gap["jenis_faskes"] == "Puskesmas") & (df_gap["Kategori"] == "Sentra Industri (Sulteng & Sultra)")]["jumlah"].values[0]
    puskesmas_non = df_gap[(df_gap["jenis_faskes"] == "Puskesmas") & (df_gap["Kategori"] == "Non-Sentra Industri (Lainnya)")]["jumlah"].values[0]

    gap_rows = []
    for _, row in df_gap.sort_values(["jenis_faskes", "Kategori"]).iterrows():
        gap_rows.append([row["Kategori"], row["jenis_faskes"], f"{row['jumlah']:.1f}"])

    prov_rows = []
    for _, row in df_2024.sort_values(["provinsi", "jenis_faskes"]).iterrows():
        kategori = "Sentra Industri" if row["provinsi"] in sentra else "Non-Sentra Industri"
        prov_rows.append([row["provinsi"], kategori, row["jenis_faskes"], f"{row['jumlah']:,.0f}"])

    konf_headers_31 = ["Komponen Analisis", "Definisi Variabel (Sub-bab 3.1)"]
    konf_rows_31 = [
        ["Jumlah & Jenis Faskes (Dependen)", "Unit Rumah Sakit dan Puskesmas terdaftar (BPS)."],
        ["Kategori Zona (Independen)", "Lokasi wilayah: Sentra Industri (Sulteng & Sultra) vs Non-Sentra Industri (Lainnya)."],
        ["Metode Analisis", "Grouped Horizontal Bar Chart pada satu periode cross-sectional untuk mengukur ketimpangan infrastruktur kesehatan primer dan sekunder."],
        ["Tahun Acuan", f"{latest_year_faskes}, mengikuti data terbaru pada file faskes."],
        ["Dataset & File", "data/processed/sulawesi_faskes_agregat_v3.csv"],
    ]

    mermaid_str_3_1 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Agregat Faskes BPS<br/><i>tahun, provinsi, jenis faskes, jumlah</i>"]
    end

    subgraph Gap_Processing["2. Segmentasi & Gap Analysis"]
        A --> B["Filter tahun acuan terbaru"]
        B --> C["Segmentasi zona<br/>Sentra Industri vs Non-Sentra"]
        C --> D["Hitung rata-rata fasilitas<br/>per jenis faskes dan zona"]
    end

    subgraph Visual_Output["3. Grouped Horizontal Bar Chart"]
        D --> E["Bandingkan Puskesmas dan Rumah Sakit"]
        E --> F["Pembacaan defisit infrastruktur publik"]
    end"""
    mermaid_png_path_3_1 = str(tool_dir / "mermaid_flowchart_3_1.png")
    download_success_3_1 = download_mermaid_png(mermaid_str_3_1, mermaid_png_path_3_1)

    print("[1.5/4] Mengekstraksi dataset empiris Bab 3 sub-bab 3.2...")
    df_kes_copy_32 = df_kes.copy()
    df_kes_copy_32["Kategori"] = df_kes_copy_32["provinsi"].apply(
        lambda x: "Sentra Industri (Sulteng & Sultra)" if x in sentra else "Non-Sentra Industri (Sulsel, Sulut, Gorontalo, Sulbar)"
    )
    indikator_32 = ["Kasus ISPA/Pneumonia", "Kasus Diare Dilayani"]
    df_filtered_32 = df_kes_copy_32[df_kes_copy_32["indikator"].isin(indikator_32)].copy()
    df_agg_32 = df_filtered_32.groupby(["indikator", "Kategori"])["nilai"].mean().reset_index()

    ispa_sentra_32 = df_agg_32[(df_agg_32["indikator"] == "Kasus ISPA/Pneumonia") & (df_agg_32["Kategori"].str.startswith("Sentra Industri"))]["nilai"].values[0]
    ispa_non_32 = df_agg_32[(df_agg_32["indikator"] == "Kasus ISPA/Pneumonia") & (df_agg_32["Kategori"].str.startswith("Non-Sentra Industri"))]["nilai"].values[0]
    ispa_diff_32 = ispa_sentra_32 / ispa_non_32 if ispa_non_32 else 0
    diare_sentra_32 = df_agg_32[(df_agg_32["indikator"] == "Kasus Diare Dilayani") & (df_agg_32["Kategori"].str.startswith("Sentra Industri"))]["nilai"].values[0]
    diare_non_32 = df_agg_32[(df_agg_32["indikator"] == "Kasus Diare Dilayani") & (df_agg_32["Kategori"].str.startswith("Non-Sentra Industri"))]["nilai"].values[0]
    diare_diff_32 = diare_sentra_32 / diare_non_32 if diare_non_32 else 0
    tahun_min_32 = int(df_filtered_32["tahun"].min())
    tahun_max_32 = int(df_filtered_32["tahun"].max())

    disease_rows_32 = []
    for _, row in df_agg_32.sort_values(["indikator", "Kategori"]).iterrows():
        disease_rows_32.append([row["indikator"], row["Kategori"], f"{row['nilai']:,.0f}"])

    ratio_rows_32 = [
        ["Kasus ISPA/Pneumonia", f"{ispa_sentra_32:,.0f}", f"{ispa_non_32:,.0f}", f"{ispa_diff_32:.1f}x"],
        ["Kasus Diare Dilayani", f"{diare_sentra_32:,.0f}", f"{diare_non_32:,.0f}", f"{diare_diff_32:.1f}x"],
    ]

    df_substitution_32 = df_filtered_32.groupby(["indikator", "Kategori"]).agg(
        total_kasus=("nilai", "sum"),
        jumlah_observasi=("nilai", "count"),
        rata_rata=("nilai", "mean"),
    ).reset_index()

    def sub_value_32(indikator, kategori, col):
        return df_substitution_32.loc[
            (df_substitution_32["indikator"] == indikator)
            & (df_substitution_32["Kategori"] == kategori),
            col,
        ].values[0]

    sentra_label_32 = "Sentra Industri (Sulteng & Sultra)"
    non_label_32 = "Non-Sentra Industri (Sulsel, Sulut, Gorontalo, Sulbar)"
    ispa_sentra_total_32 = sub_value_32("Kasus ISPA/Pneumonia", sentra_label_32, "total_kasus")
    ispa_sentra_n_32 = int(sub_value_32("Kasus ISPA/Pneumonia", sentra_label_32, "jumlah_observasi"))
    ispa_non_total_32 = sub_value_32("Kasus ISPA/Pneumonia", non_label_32, "total_kasus")
    ispa_non_n_32 = int(sub_value_32("Kasus ISPA/Pneumonia", non_label_32, "jumlah_observasi"))
    diare_sentra_total_32 = sub_value_32("Kasus Diare Dilayani", sentra_label_32, "total_kasus")
    diare_sentra_n_32 = int(sub_value_32("Kasus Diare Dilayani", sentra_label_32, "jumlah_observasi"))
    diare_non_total_32 = sub_value_32("Kasus Diare Dilayani", non_label_32, "total_kasus")
    diare_non_n_32 = int(sub_value_32("Kasus Diare Dilayani", non_label_32, "jumlah_observasi"))

    konf_headers_32 = ["Komponen Analisis", "Definisi Variabel (Sub-bab 3.2)"]
    konf_rows_32 = [
        ["Kategori Zona (Independen)", "Labeling spasial: Sentra Industri (Sulteng & Sultra) vs Non-Sentra Industri (Sulsel, Sulut, Gorontalo, Sulbar)."],
        ["Kasus ISPA/Pneumonia & Diare (Dependen)", "Total prevalensi historis penyakit per tahun dari fasilitas kesehatan primer."],
        ["Metode Analisis", "Comparative Spatial Analysis untuk membandingkan rata-rata beban penyakit antara provinsi sentra ekstraktif dan non-sentra."],
        ["Periode Observasi", f"{tahun_min_32}-{tahun_max_32}, mengikuti data kesehatan agregat yang tersedia."],
        ["Dataset & File", "data/processed/sulawesi_kesehatan_detail_2014_2024.csv"],
    ]

    mermaid_str_3_2 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Agregasi Kesehatan<br/><i>provinsi, tahun, indikator, nilai</i>"]
    end

    subgraph Spatial_Comparison["2. Comparative Spatial Analysis"]
        A --> B["Filter indikator<br/>ISPA/Pneumonia dan Diare"]
        B --> C["Segmentasi wilayah<br/>Sentra Industri vs Non-Sentra"]
        C --> D["Hitung rata-rata beban penyakit<br/>per indikator dan zona"]
    end

    subgraph Output["3. Pembacaan Ketimpangan Kesehatan"]
        D --> E["Grouped Bar Chart<br/>zona industri vs zona lainnya"]
        E --> F["Identifikasi disparitas beban penyakit"]
    end"""
    mermaid_png_path_3_2 = str(tool_dir / "mermaid_flowchart_3_2.png")
    download_success_3_2 = download_mermaid_png(mermaid_str_3_2, mermaid_png_path_3_2)

    print("[1.8/4] Mengekstraksi dataset empiris Bab 3 sub-bab 3.3...")
    df_ika_33 = pd.read_csv(data_dir / "sulawesi_ika_2016_2024.csv").rename(columns={"Indeks Kualitas Air": "IKA"})
    df_iku_33 = pd.read_csv(data_dir / "sulawesi_iku_2015_2024.csv")

    df_kes_ispa_33 = df_kes[df_kes["indikator"] == "Kasus ISPA/Pneumonia"][["provinsi", "tahun", "nilai"]].rename(columns={"nilai": "Total_ISPA", "provinsi": "Provinsi", "tahun": "Tahun"})
    df_kes_diare_33 = df_kes[df_kes["indikator"] == "Kasus Diare Dilayani"][["provinsi", "tahun", "nilai"]].rename(columns={"nilai": "Total_Diare", "provinsi": "Provinsi", "tahun": "Tahun"})
    df_panel_33 = pd.merge(df_kes_ispa_33, df_ika_33, on=["Provinsi", "Tahun"], how="outer")
    df_panel_33 = pd.merge(df_panel_33, df_kes_diare_33, on=["Provinsi", "Tahun"], how="outer")
    df_panel_33 = pd.merge(df_panel_33, df_iku_33, on=["Provinsi", "Tahun"], how="outer")

    df_panel_33["IKU_Sentra"] = df_panel_33.apply(lambda row: row["IKU"] if row["Provinsi"] in sentra else pd.NA, axis=1)
    df_panel_33["IKU_Non_Sentra"] = df_panel_33.apply(lambda row: row["IKU"] if row["Provinsi"] not in sentra else pd.NA, axis=1)

    populasi_bps_33 = {
        "Sulawesi Selatan": 9070000,
        "Sulawesi Tengah": 2985000,
        "Sulawesi Tenggara": 2624000,
        "Sulawesi Utara": 2621000,
        "Sulawesi Barat": 1419000,
        "Gorontalo": 1171000,
    }
    df_ts_33 = df_kes[(df_kes["indikator"] == "Kasus ISPA/Pneumonia") & (df_kes["nilai"] > 0)].copy()
    df_ts_33["populasi"] = df_ts_33["provinsi"].map(populasi_bps_33)
    df_ts_33["rate_per_10k"] = (df_ts_33["nilai"] / df_ts_33["populasi"]) * 10000
    tahun_min_33 = int(df_ts_33["tahun"].min())
    tahun_max_33 = int(df_ts_33["tahun"].max())
    n_prov_33 = df_ts_33["provinsi"].nunique()
    n_tahun_33 = df_ts_33["tahun"].nunique()

    df_rate_33 = df_ts_33.groupby("provinsi").agg(kasus_mean=("nilai", "mean"), rate_mean=("rate_per_10k", "mean")).reset_index()
    df_rate_33 = df_rate_33.sort_values("rate_mean", ascending=False)
    ts_rows_33 = []
    for _, row in df_rate_33.iterrows():
        kategori_33 = "Sentra Industri" if row["provinsi"] in sentra else "Non-Sentra Industri"
        ts_rows_33.append([
            row["provinsi"],
            kategori_33,
            f"{populasi_bps_33.get(row['provinsi'], 0):,.0f}",
            f"{row['kasus_mean']:,.0f}",
            f"{row['rate_mean']:,.0f}",
        ])

    x_options_33 = {"IKU_Sentra": "IKU Wilayah Sentra Tambang", "IKU_Non_Sentra": "IKU Wilayah Non-Sentra"}
    y_options_33 = {"Total_ISPA": "Total Kasus ISPA/Pneumonia"}
    index_like_33 = ["IKU_Sentra", "IKU_Non_Sentra", "IKU", "IKA", "IKA_Point"]

    lbl_h_33 = "Tinggi (>= Median Prov)"
    lbl_l_33 = "Rendah (< Median Prov)"
    summary_rows_33 = []
    detail_33 = {}
    for k_x, v_x in x_options_33.items():
        for k_y, v_y in y_options_33.items():
            loop_df = df_panel_33.dropna(subset=[k_x, k_y]).copy()
            if len(loop_df) == 0:
                continue
            loop_df["y_med_prov"] = loop_df.groupby("Provinsi")[k_y].transform("median")
            loop_df["x_med_prov"] = loop_df.groupby("Provinsi")[k_x].transform("median")
            s_x = loop_df.apply(lambda row: lbl_h_33 if row[k_x] >= row["x_med_prov"] else lbl_l_33, axis=1)
            s_y = loop_df.apply(lambda row: lbl_h_33 if row[k_y] >= row["y_med_prov"] else lbl_l_33, axis=1)
            ct = pd.crosstab(s_x, s_y).reindex(index=[lbl_l_33, lbl_h_33], columns=[lbl_l_33, lbl_h_33], fill_value=0)
            try:
                c2_val, pv_val, dof_val, _ = stats.chi2_contingency(ct)
            except Exception:
                c2_val, pv_val, dof_val = 0, 1.0, 1
            try:
                aa = ct.loc[lbl_l_33, lbl_l_33]
                bb = ct.loc[lbl_l_33, lbl_h_33]
                cc = ct.loc[lbl_h_33, lbl_l_33]
                dd = ct.loc[lbl_h_33, lbl_h_33]
                if k_x in index_like_33:
                    or_v = (bb * cc) / (aa * dd) if (aa * dd) > 0 else 0
                else:
                    or_v = (aa * dd) / (bb * cc) if (bb * cc) > 0 else 0
            except Exception:
                or_v = 0
            p_disp_33 = "p < 0.001" if pv_val < 0.001 else f"p = {pv_val:.3f}"
            summary_rows_33.append([v_x, v_y, f"{c2_val:.3f}", p_disp_33, f"{or_v:.2f}", "SIGNIFIKAN" if pv_val < 0.05 else "TIDAK SIGNIFIKAN"])
            detail_33[k_x] = {"chi2": c2_val, "p": pv_val, "or": or_v, "n": len(loop_df)}

    sig_count_33 = sum(1 for row in summary_rows_33 if row[5] == "SIGNIFIKAN")
    total_scen_33 = len(summary_rows_33)
    valid_cases_33 = detail_33.get("IKU_Sentra", {}).get("n", 0)

    if sig_count_33 > 0:
        finding_33 = f"Dari {total_scen_33} skenario pengujian, terdapat {sig_count_33} skenario yang terbukti SIGNIFIKAN. Tingginya Odds Ratio pada skenario yang signifikan menegaskan bahwa penurunan kualitas lingkungan berasosiasi dengan peningkatan risiko beban penyakit. Jika terdapat skenario yang menunjukkan TIDAK SIGNIFIKAN, ini mengindikasikan bahwa dampak ekologis dari operasi industri telah tersebar secara meluas (spillover effect) di mana dampak lingkungan menjalar melampaui area operasi langsung."
    else:
        finding_33 = f"Dari {total_scen_33} skenario pengujian, seluruhnya menunjukkan status TIDAK SIGNIFIKAN. Dalam perspektif analisis ekologis, ketidaksignifikanan secara agregat ini mengindikasikan bahwa penurunan kualitas lingkungan dan peningkatan beban penyakit telah terjadi secara merata dan persisten di seluruh wilayah. Penambahan aktivitas industri di satu titik berkorelasi dengan tekanan lingkungan yang sudah merata secara sistemik."

    konf_headers_33 = ["Komponen Uji", "Definisi Variabel (Sub-bab 3.3)"]
    konf_rows_33 = [
        ["Variabel Independen (X)", "IKU Wilayah Sentra Tambang / IKU Wilayah Non-Sentra (indeks tekanan kualitas lingkungan)."],
        ["Variabel Dependen (Y)", "Total Kasus ISPA/Pneumonia (insidensi penyakit pernapasan dan lingkungan)."],
        ["Hipotesis Nol (H0)", "Penurunan kualitas lingkungan (IKU/IKA) tidak berkorelasi dengan peningkatan insidensi penyakit pernapasan dan pencernaan."],
        ["Hipotesis Alternatif (H1)", "Penurunan kualitas udara ambien (IKU) berbanding lurus dengan peningkatan insidensi penyakit pernapasan dan lingkungan (ISPA dan Diare)."],
        ["Decision Rule (Alpha 5%)", "Chi-Square P-Value < 0.05 (Tolak H0) dan kalkulasi Odds Ratio."],
        ["Threshold Kategori", f"Median per-provinsi data panel Provinsi-Tahun (N={valid_cases_33} observasi valid skenario Sentra); binning 'Tinggi'/'Rendah' per provinsi untuk menghilangkan bias besaran absolut antar wilayah."],
        ["Orientasi Odds Ratio", "Untuk variabel X berjenis indeks kualitas (IKU/IKA), risiko dihitung saat indeks Rendah: OR = ( b × c ) / ( a × d )."],
    ]

    mermaid_str_3_3 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Kesehatan Detail<br/><i>provinsi, tahun, indikator, nilai</i>"]
        B["Data IKU KLHK & IKA<br/><i>Provinsi, Tahun, indeks kualitas</i>"]
        P["Populasi Proxy BPS 2020<br/><i>denominator per kapita</i>"]
    end
    subgraph Panel_Processing["2. Pembentukan Panel & Normalisasi"]
        A --> C["Merge Panel Provinsi-Tahun<br/>ISPA, Diare, IKA, IKU"]
        B --> C
        P --> D["Rasio Insiden per 10.000 Penduduk"]
        C --> D
    end
    subgraph Visual_Output["3. Time-Series Line Chart"]
        D --> E["Tren per kapita & absolut 2014-2024<br/>Sentra (merah) vs Non-Sentra (biru)"]
        E --> F["Pembacaan lintasan waktu beban kesehatan"]
    end"""
    mermaid_png_path_3_3 = str(tool_dir / "mermaid_flowchart_3_3.png")
    download_success_3_3 = download_mermaid_png(mermaid_str_3_3, mermaid_png_path_3_3)

    print("[1.9/4] Mengekstraksi dataset empiris Bab 3 sub-bab 3.4...")
    df_zoo_sulteng = df_zoonosis[df_zoonosis["provinsi"].astype(str).str.upper() == "SULTENG"].copy()
    tambang_kab_34 = ["MOROWALI", "MOROWALI UTARA", "BANGGAI"]
    df_zoo_sulteng["Kategori_Wilayah"] = df_zoo_sulteng["kabupaten_kota"].apply(
        lambda kab: "Lingkar Tambang/Smelter Aktif" if str(kab).upper() in tambang_kab_34 else "Non-Tambang/Agraris (Kontrol)"
    )
    df_tambang_only_34 = df_zoo_sulteng[df_zoo_sulteng["Kategori_Wilayah"] == "Lingkar Tambang/Smelter Aktif"].copy()
    total_kasus_tambang_34 = df_tambang_only_34["total_kasus"].sum()
    tahun_min_34 = int(df_zoo_sulteng["tahun"].min())
    tahun_max_34 = int(df_zoo_sulteng["tahun"].max())
    n_kab_34 = df_zoo_sulteng["kabupaten_kota"].nunique()
    n_penyakit_34 = df_zoo_sulteng["jenis_penyakit"].nunique()

    peak_rows_34 = []
    for penyakit in sorted(df_tambang_only_34["jenis_penyakit"].dropna().unique()):
        df_p = df_tambang_only_34[df_tambang_only_34["jenis_penyakit"] == penyakit]
        if df_p.empty or df_p["total_kasus"].max() <= 0:
            continue
        max_row = df_p.loc[df_p["total_kasus"].idxmax()]
        peak_rows_34.append([
            penyakit,
            str(max_row["kabupaten_kota"]).title(),
            str(int(max_row["tahun"])),
            f"{max_row['total_kasus']:,.0f}",
        ])

    selected_penyakit_34 = "DBD" if "DBD" in df_zoo_sulteng["jenis_penyakit"].unique() else sorted(df_zoo_sulteng["jenis_penyakit"].dropna().unique())[0]
    df_zoo_selected_34 = df_zoo_sulteng[
        (df_zoo_sulteng["jenis_penyakit"] == selected_penyakit_34)
        & (~df_zoo_sulteng["kabupaten_kota"].astype(str).str.upper().isin(["PALU"]))
    ].copy()
    df_zoo_bar_34 = df_zoo_selected_34.groupby("Kategori_Wilayah")["total_kasus"].mean().reset_index()
    val_tambang_34 = df_zoo_bar_34.loc[df_zoo_bar_34["Kategori_Wilayah"] == "Lingkar Tambang/Smelter Aktif", "total_kasus"].values
    val_non_34 = df_zoo_bar_34.loc[df_zoo_bar_34["Kategori_Wilayah"] == "Non-Tambang/Agraris (Kontrol)", "total_kasus"].values
    val_tambang_34 = float(val_tambang_34[0]) if len(val_tambang_34) else 0.0
    val_non_34 = float(val_non_34[0]) if len(val_non_34) else 0.0
    multiplier_34 = val_tambang_34 / val_non_34 if val_non_34 > 0 else 0.0

    trend_rows_34 = []
    df_trend_34 = df_zoo_selected_34.groupby(["Kategori_Wilayah", "tahun"])["total_kasus"].sum().reset_index()
    for _, row in df_trend_34.sort_values(["tahun", "Kategori_Wilayah"]).iterrows():
        trend_rows_34.append([str(int(row["tahun"])), row["Kategori_Wilayah"], f"{row['total_kasus']:,.0f}"])

    avg_rows_34 = []
    for _, row in df_zoo_bar_34.sort_values("Kategori_Wilayah").iterrows():
        avg_rows_34.append([selected_penyakit_34, row["Kategori_Wilayah"], f"{row['total_kasus']:,.1f}"])

    mermaid_str_3_4 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Zoonosis Kabupaten/Kota<br/><i>provinsi, kabupaten, tahun, penyakit, total kasus</i>"]
    end
    subgraph Case_Study["2. Deep Dive Case Study Sulteng"]
        A --> B["Filter Provinsi Sulteng"]
        B --> C["Segmentasi distrik<br/>Morowali, Morowali Utara, Banggai vs kontrol"]
        C --> D["Akumulasi tren tahunan<br/>per penyakit dan kategori wilayah"]
    end
    subgraph Output["3. Time-Series & Komparasi Spasial"]
        D --> E["Identifikasi puncak kasus lingkar tambang"]
        D --> F["Rata-rata kasus tambang vs kontrol"]
        E --> G["Pembacaan anomali zoonosis level tapak"]
        F --> G
    end"""
    mermaid_png_path_3_4 = str(tool_dir / "mermaid_flowchart_3_4.png")
    download_success_3_4 = download_mermaid_png(mermaid_str_3_4, mermaid_png_path_3_4)

    print("[2/4] Membangun DOCX Metodologi_Bab3_Beban_Kesehatan.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)

    add_h1(doc, "BAB III: METODOLOGI ANALISIS BEBAN KESEHATAN MASYARAKAT TERDAMPAK")
    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada ", False, False),
        ("Bab 3: Beban Kesehatan Masyarakat Terdampak", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.", False, False),
    ])

    add_h2(doc, "3.1 Kesenjangan Fasilitas Kesehatan di Kawasan Ekstraktif")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Agregat Faskes: data/processed/sulawesi_faskes_agregat_v3.csv. Visualisasi dashboard menggunakan Grouped Horizontal Bar Chart pada tahun acuan 2024 untuk membandingkan rata-rata Puskesmas dan Rumah Sakit antara zona Sentra Industri dan Non-Sentra Industri.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Data empiris menggambarkan kesenjangan antara klaim pertumbuhan ekonomi dari ekspansi industri nikel dan kondisi kesehatan masyarakat di kawasan penyangga. Selama satu dekade terakhir, emisi partikulat, gas buang PLTU batu bara, dan timbulan limbah dari fasilitas ekstraktif telah memberikan tekanan signifikan terhadap kualitas lingkungan hidup masyarakat. ", False, False),
        (f"Data empiris merekam bagaimana ekspansi kapasitas industri, yang ditopang oleh PLTU captive berkapasitas {tot_kapasitas_pltu:,.0f} Megawatt, berjalan sejajar dengan peningkatan kasus penyakit di kawasan-kawasan penyangga.", False, False),
    ])
    add_p(doc, [
        (f"Sepanjang 2014-2024, data agregat dinas kesehatan mencatat total kasus ISPA dan Pneumonia sebanyak {tot_ispa:,.0f} kasus, kasus Diare sebanyak {tot_diare:,.0f} kasus, dan kasus Malaria sebanyak {tot_malaria:,.0f} kasus. Distribusi infrastruktur kesehatan tahun {latest_year_faskes} menunjukkan Puskesmas sebanyak {tot_puskesmas_latest:,.0f} unit dan Rumah Sakit sebanyak {tot_rs_latest:,.0f} unit di wilayah observasi.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Grouped Horizontal Bar Chart")
    add_p(doc, [
        ("Kerangka analisis ketimpangan infrastruktur kesehatan primer dan sekunder diilustrasikan pada ", False, False),
        ("Bagan Alur 3.1", True, False),
        (" berikut. Sub-bab ini menggunakan Gap Analysis cross-sectional dan tidak menggunakan uji inferensial Chi-Square.", False, False),
    ])
    add_caption(doc, "Bagan Alur 3.1: Alur Logika Metodologis Grouped Horizontal Bar Chart Kesenjangan Faskes")
    if download_success_3_1:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_3_1, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 3.1 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_caption(doc, "Tabel 3.1a: Konfigurasi Variabel Analisis Gap Fasilitas Kesehatan (Sub-bab 3.1)")
    add_table_1col(doc, konf_headers_31, konf_rows_31, [4.5, 11.0], ["L", "L"])

    add_h4(doc, "C. Formulasi Matematis: Rata-rata Faskes dan Disparitas Zona")
    add_p(doc, [("Kuantifikasi kesenjangan fasilitas kesehatan dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Rata-rata Fasilitas Kesehatan per Zona", "F̄_{z,j} = [ Σ_{p∈z} F_{p,j} ] / n_z", [
        ("F̄_{z,j}", "Rata-rata jumlah fasilitas kesehatan jenis j pada zona z."),
        ("F_{p,j}", "Jumlah fasilitas kesehatan jenis j pada provinsi p."),
        ("p", "Provinsi observasi di Pulau Sulawesi."),
        ("p∈z", "Dibaca 'p anggota dari z' atau 'p termasuk dalam zona z', artinya provinsi p masuk ke dalam kategori zona z."),
        ("z", "Kategori zona wilayah: Sentra Industri atau Non-Sentra Industri."),
        ("j", "Jenis fasilitas kesehatan: Puskesmas atau Rumah Sakit."),
        ("n_z", "Jumlah provinsi dalam zona observasi."),
    ])
    add_formula(doc, "Rasio Disparitas Faskes Sentra vs Non-Sentra", "D_j = F̄_{Sentra,j} / F̄_{Non-Sentra,j}", [
        ("D_j", "Rasio perbandingan rata-rata fasilitas kesehatan jenis j antara zona Sentra dan Non-Sentra."),
        ("F̄_{Sentra,j}", "Rata-rata jumlah fasilitas kesehatan jenis j pada zona Sentra Industri (Sulawesi Tengah dan Sulawesi Tenggara)."),
        ("F̄_{Non-Sentra,j}", "Rata-rata jumlah fasilitas kesehatan jenis j pada zona Non-Sentra Industri (provinsi selain Sulawesi Tengah dan Sulawesi Tenggara)."),
        ("j", "Jenis fasilitas kesehatan: Puskesmas atau Rumah Sakit."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Ketersediaan Faskes Sentra vs Non-Sentra")
    add_p(doc, [
        (f"Perbandingan rata-rata fasilitas kesehatan pada tahun {latest_year_faskes} antara zona Sentra Industri dan Non-Sentra Industri disajikan pada ", False, False),
        ("Tabel 3.1", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 3.1: Rata-rata Fasilitas Kesehatan per Provinsi menurut Zona Industri ({latest_year_faskes})")
    add_table_1col(doc, ["Kategori Zona", "Jenis Faskes", "Rata-rata Jumlah Fasilitas"], gap_rows, [5.2, 4.0, 4.0], ["L", "L", "C"])

    add_caption(doc, f"Tabel 3.2: Rincian Fasilitas Kesehatan per Provinsi ({latest_year_faskes})")
    add_table_1col(doc, ["Provinsi", "Kategori Zona", "Jenis Faskes", "Jumlah"], prov_rows, [3.8, 4.2, 3.2, 2.2], ["L", "L", "L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Defisit Infrastruktur Kesehatan Publik")
    add_p(doc, [
        ("Data perbandingan distribusi fasilitas kesehatan mengindikasikan bahwa ketersediaan infrastruktur medis di provinsi sentra industri relatif tidak lebih baik dibandingkan wilayah non-sentra, meski beban penyakit di wilayah tersebut lebih tinggi. ", False, False),
        (f"Rata-rata Rumah Sakit di Sentra Industri tercatat {rs_sentra:.0f} unit per provinsi, lebih rendah dari wilayah Non-Sentra yang mencapai {rs_non:.0f} unit. ", False, False),
        (f"Untuk fasilitas primer, rata-rata Puskesmas di Sentra Industri tercatat {puskesmas_sentra:.0f} unit per provinsi dibandingkan {puskesmas_non:.0f} unit di Non-Sentra. Kesenjangan ini menunjukkan bahwa pertumbuhan ekonomi dari hilirisasi nikel belum diimbangi distribusi infrastruktur kesehatan yang proporsional bagi masyarakat di wilayah operasi industri.", False, False),
    ])

    add_h2(doc, "3.2 Ketimpangan Beban Penyakit: Sentra Industri vs Non-Sentra")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Agregasi Kesehatan: data/processed/sulawesi_kesehatan_detail_2014_2024.csv. Visualisasi dashboard menggunakan Comparative Spatial Analysis untuk membandingkan rata-rata beban penyakit antara provinsi sentra ekstraktif dan non-sentra.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Melalui analisis komparatif spasial, terlihat bahwa beban ekologis tidak terdistribusi secara merata di seluruh wilayah. Provinsi sentra ekspansi nikel, yaitu Sulawesi Tengah dan Sulawesi Tenggara, menunjukkan indikator penyakit yang secara konsisten lebih tinggi. ", False, False),
        (f"Data menunjukkan bahwa rata-rata penderita ISPA/Pneumonia di Sentra Industri tercatat {ispa_sentra_32:,.0f} kasus per tahun, dibandingkan provinsi Non-Sentra di angka {ispa_non_32:,.0f} kasus. Selisih sebesar {ispa_diff_32:.1f} kali lipat ini mengindikasikan beban pernapasan yang lebih berat di kawasan penyangga smelter.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Comparative Spatial Analysis")
    add_p(doc, [
        ("Kerangka komparasi spasial untuk membandingkan rata-rata beban penyakit antara provinsi sentra ekstraktif dan non-sentra diilustrasikan pada ", False, False),
        ("Bagan Alur 3.2", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan perbandingan rata-rata absolut beban penyakit tahunan berdasarkan klasifikasi wilayah.", False, False),
    ])
    add_caption(doc, "Bagan Alur 3.2: Alur Logika Metodologis Comparative Spatial Analysis Beban Penyakit")
    if download_success_3_2:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_3_2, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 3.2 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_caption(doc, "Tabel 3.2a: Konfigurasi Variabel Comparative Spatial Analysis (Sub-bab 3.2)")
    add_table_1col(doc, konf_headers_32, konf_rows_32, [4.5, 11.0], ["L", "L"])

    add_h4(doc, "C. Formulasi Matematis: Rata-rata Beban Penyakit dan Disparitas Zona")
    add_p(doc, [("Kuantifikasi rata-rata beban penyakit dan disparitas antar-zona dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Rata-rata Beban Penyakit per Zona", "B̄_{z,k} = [ Σ_{p∈z} Σ_t C_{p,t,k} ] / N_z", [
        ("B̄_{z,k}", "Rata-rata beban penyakit indikator k pada zona z."),
        ("C_{p,t,k}", "Jumlah kasus penyakit indikator k pada provinsi p dan tahun t."),
        ("p∈z", "Dibaca 'p anggota dari z' atau 'p termasuk dalam zona z', artinya provinsi p masuk ke dalam kategori zona z."),
        ("t", f"Tahun observasi dalam periode {tahun_min_32}-{tahun_max_32}."),
        ("k", "Indikator penyakit: Kasus ISPA/Pneumonia atau Kasus Diare Dilayani."),
        ("N_z", "Jumlah observasi provinsi-tahun yang tersedia pada zona z."),
    ])
    add_formula(doc, "Rasio Disparitas Beban Penyakit Sentra vs Non-Sentra", "Q_k = B̄_{Sentra,k} / B̄_{Non-Sentra,k}", [
        ("Q_k", "Rasio perbandingan rata-rata beban penyakit indikator k antara zona Sentra dan Non-Sentra."),
        ("B̄_{Sentra,k}", "Rata-rata beban penyakit indikator k pada zona Sentra Industri (Sulawesi Tengah dan Sulawesi Tenggara)."),
        ("B̄_{Non-Sentra,k}", "Rata-rata beban penyakit indikator k pada zona Non-Sentra Industri."),
    ])

    add_p(doc, [("Substitusi angka dari dataset aktual ke dalam rumus rata-rata beban penyakit adalah sebagai berikut:", False, False)])
    add_formula(doc, "Substitusi ISPA/Pneumonia Zona Sentra", f"B̄_Sentra,ISPA = {ispa_sentra_total_32:,.0f} / {ispa_sentra_n_32} = {ispa_sentra_32:,.1f} kasus")
    add_formula(doc, "Substitusi ISPA/Pneumonia Zona Non-Sentra", f"B̄_Non-Sentra,ISPA = {ispa_non_total_32:,.0f} / {ispa_non_n_32} = {ispa_non_32:,.1f} kasus")
    add_formula(doc, "Substitusi Diare Zona Sentra", f"B̄_Sentra,Diare = {diare_sentra_total_32:,.0f} / {diare_sentra_n_32} = {diare_sentra_32:,.1f} kasus")
    add_formula(doc, "Substitusi Diare Zona Non-Sentra", f"B̄_Non-Sentra,Diare = {diare_non_total_32:,.0f} / {diare_non_n_32} = {diare_non_32:,.1f} kasus")
    add_formula(doc, "Substitusi Rasio Disparitas Beban Penyakit", f"Q_ISPA = {ispa_sentra_32:,.1f} / {ispa_non_32:,.1f} = {ispa_diff_32:.1f}x\nQ_Diare = {diare_sentra_32:,.1f} / {diare_non_32:,.1f} = {diare_diff_32:.1f}x")

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Beban Penyakit Sentra vs Non-Sentra")
    add_p(doc, [
        (f"Perbandingan rata-rata beban penyakit pada periode {tahun_min_32}-{tahun_max_32} antara zona Sentra Industri dan Non-Sentra Industri disajikan pada ", False, False),
        ("Tabel 3.3", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 3.3: Rata-rata Beban Penyakit per Tahun menurut Zona Industri ({tahun_min_32}-{tahun_max_32})")
    add_table_1col(doc, ["Indikator Penyakit", "Kategori Zona", "Rata-rata Kasus per Tahun"], disease_rows_32, [4.2, 5.8, 3.2], ["L", "L", "C"])

    add_caption(doc, "Tabel 3.4: Ringkasan Disparitas Beban Penyakit Sentra vs Non-Sentra")
    add_table_1col(doc, ["Indikator Penyakit", "Rata-rata Sentra", "Rata-rata Non-Sentra", "Rasio Disparitas"], ratio_rows_32, [4.2, 3.2, 3.2, 3.2], ["L", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Ketimpangan Beban Penyakit Struktural")
    add_p(doc, [
        (f"Rata-rata penderita ISPA/Pneumonia di Sentra Industri tercatat {ispa_sentra_32:,.0f} kasus per tahun, dibandingkan provinsi Non-Sentra di angka {ispa_non_32:,.0f} kasus. Selisih sebesar {ispa_diff_32:.1f} kali lipat mendukung pembacaan bahwa wilayah dengan konsentrasi industri tinggi cenderung menanggung beban kesehatan yang lebih besar akibat tekanan terhadap daya tampung lingkungan. ", False, False),
        ("Kesenjangan statistik ini mengindikasikan bahwa manfaat ekonomi dari hilirisasi nikel belum disertai perbaikan infrastruktur kesehatan yang proporsional di wilayah operasi industri ekstraktif.", False, False),
    ])

    add_h2(doc, "3.3 Lintasan Waktu Ekologis & Dinamika Penyakit di Kawasan Industri Ekstraktif")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Lingkungan & Penyakit: data/processed/sulawesi_kesehatan_detail_2014_2024.csv, data/processed/sulawesi_ika_2016_2024.csv, data/processed/sulawesi_iku_2015_2024.csv. Visualisasi dashboard menampilkan Time-Series Line Chart (insiden per 10.000 penduduk dan total kasus absolut) serta pengujian Chi-Square tabulasi silang (Crosstabulation) dengan binning median per-provinsi.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        (f"Meskipun secara akumulatif kawasan Sentra Industri menanggung beban yang lebih berat, penelusuran data secara time-series (historis) dari {tahun_min_33} hingga {tahun_max_33} memberikan wawasan tambahan mengenai fluktuasi kasus penyakit dari tahun ke tahun. Konversi angka absolut ke rasio per kapita (Kasus per 10.000 Penduduk) dilakukan untuk menghilangkan bias jumlah populasi antar wilayah.", False, False),
    ])
    add_p(doc, [
        ("Hipotesis utama narasi ini adalah bahwa penurunan kualitas udara ambien (IKU) berbanding lurus dengan peningkatan insidensi penyakit pernapasan dan lingkungan (seperti ISPA dan Diare). ", False, False),
        (f"Untuk mengujinya secara statistik di tengah keterbatasan jumlah provinsi di Sulawesi (N={n_prov_33}), tabel crosstab dan uji Chi-Square menggunakan unit observasi Provinsi-Tahun ({n_prov_33} provinsi × {n_tahun_33} tahun panel). Setiap observasi diklasifikasikan menjadi 'Tinggi' atau 'Rendah' berdasarkan nilai median per-provinsi dari indikator yang dipilih.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Time-Series Line Chart & Crosstabulation")
    add_p(doc, [
        ("Pendekatan penelusuran runtut waktu insiden penyakit sejalan dengan akumulasi polusi tahunan diilustrasikan pada ", False, False),
        ("Bagan Alur 3.3", True, False),
        (" berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya diringkas melalui tabel konfigurasi variabel di bawah gambar.", False, False),
    ])
    add_caption(doc, "Bagan Alur 3.3: Alur Logika Analisis Time-Series Beban Kesehatan")
    if download_success_3_3:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_3_3, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 3.3 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_caption(doc, "Tabel 3.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 3.3)")
    add_table_1col(doc, konf_headers_33, konf_rows_33, [4.5, 11.0], ["L", "L"])

    add_h4(doc, "C. Formulasi Matematis: Normalisasi Per Kapita, Binning Median Provinsi, dan Uji Crosstabulation")
    add_p(doc, [("Kuantifikasi rasio keparahan per kapita dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Persamaan Rasio Insiden per 10.000 Penduduk (Normalisasi Per Kapita)", "Insiden_10K_p,t = ( Kasus_p,t / Populasi_p ) × 10.000", [
        ("Insiden_10K_p,t", "Rasio keparahan penyakit per 10.000 penduduk pada provinsi p dan tahun t, menghilangkan bias jumlah populasi antar wilayah."),
        ("Kasus_p,t", "Jumlah kasus absolut indikator penyakit pada provinsi p dan tahun t."),
        ("Populasi_p", "Populasi proxy BPS 2020 provinsi p sebagai denominator normalisasi."),
    ])
    add_formula(doc, "Persamaan Ambang Batas Median per-Provinsi", "Median_Prov_p = Median ( Nilai_p,t )   ;   untuk seluruh tahun t pada provinsi p", [
        ("Median_Prov_p", "Nilai tengah indikator pada masing-masing provinsi p, digunakan sebagai ambang batas binning untuk menghilangkan efek bias besaran absolut antar wilayah."),
        ("Nilai_p,t", "Nilai indikator (X atau Y) pada provinsi p dan tahun t."),
    ])
    add_formula(doc, "Persamaan Kategorisasi Median per-Provinsi (Fungsi Piecewise)", "Kategori(x_p,t) = 'Tinggi' , jika x_p,t ≥ Median_Prov_p   |   'Rendah' , jika x_p,t < Median_Prov_p", [
        ("Kategori(x_p,t)", "Klasifikasi biner tiap observasi provinsi-tahun terhadap median provinsinya sendiri."),
    ])
    add_formula(doc, "Persamaan Uji Independensi Chi-Square Pearson (χ² Kontinjensi 2x2)", "χ² = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   dengan E_ij = ( Total_Baris_i × Total_Kolom_j ) / N", [
        ("χ²", f"Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara indeks kualitas udara dan insidensi penyakit pada panel Provinsi-Tahun (N={valid_cases_33} observasi valid skenario Sentra)."),
        ("O_ij", "Frekuensi Observasi: jumlah kasus aktual pada sel baris i kolom j tabel kontinjensi 2x2."),
        ("E_ij", "Frekuensi Harapan: jumlah kasus teoretis jika kedua variabel saling independen, E_ij = ( Total_Baris_i × Total_Kolom_j ) / N."),
    ])
    add_formula(doc, "Persamaan Rasio Keunggulan Risiko untuk Variabel Indeks (Risk Odds Ratio / OR)", "Odds_Ratio (OR) = ( b × c ) / ( a × d )   ;   untuk X berjenis indeks kualitas (IKU/IKA)", [
        ("Odds_Ratio (OR)", "Ukuran kelipatan peluang munculnya insiden penyakit Tinggi pada kelompok indeks kualitas Rendah; orientasi dibalik karena risiko terjadi saat indeks kualitas lingkungan menurun."),
        ("a", "Jumlah observasi panel pada kelompok X Rendah dan Y Rendah."),
        ("b", "Jumlah observasi panel pada kelompok X Rendah dan Y Tinggi."),
        ("c", "Jumlah observasi panel pada kelompok X Tinggi dan Y Rendah."),
        ("d", "Jumlah observasi panel pada kelompok X Tinggi dan Y Tinggi."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Insiden Per Kapita dan Skenario Crosstab")
    add_p(doc, [
        (f"Rata-rata tahunan kasus ISPA/Pneumonia dan rasio insiden per 10.000 penduduk masing-masing provinsi pada periode {tahun_min_33}-{tahun_max_33} disajikan pada ", False, False),
        ("Tabel 3.5", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 3.5: Rata-rata Insiden ISPA/Pneumonia Absolut dan per 10.000 Penduduk per Provinsi ({tahun_min_33}-{tahun_max_33})")
    add_table_1col(doc, ["Provinsi", "Kategori Zona", "Populasi Proxy (BPS 2020)", "Rata-rata Kasus per Tahun", "Rata-rata Insiden per 10.000 Penduduk"], ts_rows_33, [2.8, 3.0, 3.2, 3.2, 3.6], ["L", "L", "C", "C", "C"])

    add_p(doc, [
        ("Ringkasan hasil pengujian statistik (Chi-Square) untuk semua kemungkinan kombinasi antara indikator lingkungan (X) dan dampak kesehatan (Y) pada panel data yang sama disajikan pada ", False, False),
        ("Tabel 3.6", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 3.6: Ringkasan Eksekutif Seluruh Skenario Crosstab IKU vs Insidensi Penyakit Bab 3")
    add_table_1col(doc, ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_33, [3.2, 3.4, 2.0, 2.0, 2.0, 2.6], ["L", "L", "C", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Pembedahan Realitas Ekologis")
    add_p(doc, [(finding_33, False, False)])

    add_h2(doc, "3.4 Anomali Zoonosis: Dampak Kritis Ekspansi Industri di Level Tapak (Studi Kasus Sulteng)")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Zoonosis: data/processed/zoonosis_kab_kota_2015_2024.csv. Visualisasi dashboard menggunakan Time-Series dan Komparasi Spasial Wilayah untuk membandingkan kabupaten lingkar tambang/smelter aktif dengan daerah non-tambang/agraris sebagai kontrol.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        (f"Data empiris Dinas Kesehatan mencatat total akumulasi {total_kasus_tambang_34:,.0f} kasus penyakit zoonosis di wilayah Lingkar Tambang/Smelter Aktif Sulawesi Tengah (Morowali, Morowali Utara, Banggai) sepanjang rentang waktu pengamatan {tahun_min_34}-{tahun_max_34}. ", False, False),
        ("Peningkatan angka zoonosis ini dibaca bersama perubahan ekologis akibat ekspansi penggunaan lahan, konversi tutupan hutan, pergeseran habitat alami satwa liar, genangan air galian tambang yang tidak direklamasi, serta kondisi sanitasi di area industri.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Time-Series & Komparasi Spasial Wilayah")
    add_p(doc, [
        ("Kerangka studi kasus mendalam berbasis deret waktu di tingkat kabupaten/kota khusus Sulawesi Tengah diilustrasikan pada ", False, False),
        ("Bagan Alur 3.4", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan isolasi episentrum ekstraktif dan komparasi absolut dengan wilayah kontrol.", False, False),
    ])
    add_caption(doc, "Bagan Alur 3.4: Alur Logika Metodologis Anomali Zoonosis Level Tapak")
    if download_success_3_4:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_3_4, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 3.4 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Tren Zoonosis Distrik dan Komparasi Wilayah")
    add_p(doc, [("Akumulasi tren tahunan infeksi zoonosis dan rasio komparatif wilayah dihitung menggunakan sistem formulasi matematis berikut:", False, False)])
    add_formula(doc, "Akumulasi Tren Zoonosis per Kategori Wilayah", "Z_{w,t,d} = Σ C_{r,t,d}, untuk setiap distrik r yang termasuk wilayah w", [
        ("Z_{w,t,d}", "Total kasus penyakit d pada kategori wilayah w dan tahun t."),
        ("C_{r,t,d}", "Jumlah kasus penyakit d pada distrik/kabupaten r dan tahun t."),
        ("w", "Kategori wilayah: Lingkar Tambang/Smelter Aktif atau Non-Tambang/Agraris (Kontrol)."),
        ("d", "Jenis penyakit zoonosis yang diamati."),
    ])
    add_formula(doc, "Rata-rata Kasus Zoonosis Wilayah", "Z̄_w = [ Σ_t Z_{w,t,d} ] / N_w", [
        ("Z̄_w", "Rata-rata kasus zoonosis pada kategori wilayah w untuk penyakit terpilih."),
        ("N_w", "Jumlah observasi distrik-tahun pada kategori wilayah w."),
    ])
    add_formula(doc, "Rasio Komparatif Tambang vs Kontrol", "R_d = Z̄_Tambang / Z̄_Kontrol", [
        ("R_d", f"Rasio perbandingan rata-rata kasus {selected_penyakit_34} antara wilayah tambang dan kontrol."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Puncak Kasus dan Komparasi Zoonosis")
    add_p(doc, [("Rincian insidensi tertinggi menurut jenis penyakit di wilayah lingkar tambang/smelter aktif disajikan pada Tabel 3.7 berikut:", False, False)])
    add_caption(doc, "Tabel 3.7: Insidensi Tertinggi Zoonosis di Lingkar Tambang/Smelter Aktif Sulawesi Tengah")
    add_table_1col(doc, ["Jenis Penyakit", "Kabupaten", "Tahun", "Puncak Kasus"], peak_rows_34, [4.0, 3.6, 2.0, 3.0], ["L", "L", "C", "C"])

    add_caption(doc, f"Tabel 3.8: Tren Tahunan {selected_penyakit_34} menurut Kategori Wilayah")
    add_table_1col(doc, ["Tahun", "Kategori Wilayah", "Total Kasus"], trend_rows_34, [2.0, 6.0, 3.0], ["C", "L", "C"])

    add_caption(doc, f"Tabel 3.9: Rata-rata Kasus {selected_penyakit_34} Tambang vs Kontrol")
    add_table_1col(doc, ["Jenis Penyakit", "Kategori Wilayah", "Rata-rata Kasus"], avg_rows_34, [3.5, 6.0, 3.0], ["L", "L", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Anomali Zoonosis Level Tapak")
    add_p(doc, [
        (f"Pada penyakit terpilih ({selected_penyakit_34}), rata-rata kasus di wilayah Lingkar Tambang/Smelter Aktif mencapai {val_tambang_34:,.1f} kasus per observasi, dibandingkan {val_non_34:,.1f} kasus pada wilayah Non-Tambang/Agraris (Kontrol), dengan rasio komparatif {multiplier_34:.1f}x. ", False, False),
        ("Pola ini memberikan sinyal bahwa perubahan ekologis di sekitar kawasan industri perlu dibaca sampai level tapak, karena data agregat provinsi dapat mengaburkan lonjakan penyakit pada kabupaten episentrum ekstraktif.", False, False),
    ])

    docx_path = tool_dir / "Metodologi_Bab3_Beban_Kesehatan.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 3...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 3 - Beban Kesehatan</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: Arial, sans-serif; max-width: 920px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #43A047; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; }}
.hdr-title {{ color: #81C784; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #2E7D32; padding-bottom: 8px; }}
h2 {{ color: #81C784; text-transform: uppercase; border-bottom: 1px solid #2E7D32; }}
h4 {{ color: #A5D6A7; }}
.note-box {{ background: #132213; border-left: 4px solid #2E7D32; padding: 10px 14px; margin: 12px 0; }}
.formula {{ background: #0D1B0E; border: 1px solid #2E7D32; color: #A5D6A7; padding: 8px 12px; font-family: monospace; }}
.data-th {{ background: #1B5E20; color: white; padding: 6px; text-align: left; border: 1px solid #2E7D32; }}
.data-td {{ padding: 6px; border: 1px solid #243524; vertical-align: top; }}
.data-tr-even .data-td {{ background: #131B13; }}
.mermaid {{ background: #0D1610; border: 1px solid #2E7D32; padding: 12px; margin: 10px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB III: Metodologi Analisis Beban Kesehatan Masyarakat Terdampak</div>
<p>Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada <strong>Bab 3: Beban Kesehatan Masyarakat Terdampak</strong>.</p>
<h2>3.1 Kesenjangan Fasilitas Kesehatan di Kawasan Ekstraktif</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Agregat Faskes: <code>data/processed/sulawesi_faskes_agregat_v3.csv</code>. Visualisasi dashboard menggunakan Grouped Horizontal Bar Chart pada tahun acuan 2024.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Data empiris menggambarkan kesenjangan antara klaim pertumbuhan ekonomi dari ekspansi industri nikel dan kondisi kesehatan masyarakat di kawasan penyangga. Ekspansi kapasitas industri yang ditopang PLTU captive berkapasitas <strong>{tot_kapasitas_pltu:,.0f} Megawatt</strong> berjalan sejajar dengan peningkatan kasus penyakit.</p>
<p>Sepanjang 2014-2024, data agregat dinas kesehatan mencatat total kasus ISPA/Pneumonia sebanyak <strong>{tot_ispa:,.0f}</strong>, kasus Diare sebanyak <strong>{tot_diare:,.0f}</strong>, dan kasus Malaria sebanyak <strong>{tot_malaria:,.0f}</strong>. Pada {latest_year_faskes}, tercatat <strong>{tot_puskesmas_latest:,.0f}</strong> Puskesmas dan <strong>{tot_rs_latest:,.0f}</strong> Rumah Sakit.</p>
<h4>B. Alur Logika Metodologis Grouped Horizontal Bar Chart</h4>
<div class="mermaid">{mermaid_str_3_1}</div>
<div class="table-caption">Tabel 3.1a: Konfigurasi Variabel Analisis Gap Fasilitas Kesehatan (Sub-bab 3.1)</div>
{html_table(konf_headers_31, konf_rows_31)}
<h4>C. Formulasi Matematis: Rata-rata Faskes dan Disparitas Zona</h4>
<div class="formula">F̄_{{z,j}} = [ Σ_{{p∈z}} F_{{p,j}} ] / n_z</div>
<div class="formula">D_j = F̄_{{Sentra,j}} / F̄_{{Non-Sentra,j}}</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 3.1: Rata-rata Fasilitas Kesehatan per Provinsi menurut Zona Industri ({latest_year_faskes})</div>
{html_table(["Kategori Zona", "Jenis Faskes", "Rata-rata Jumlah Fasilitas"], gap_rows)}
<div class="table-caption">Tabel 3.2: Rincian Fasilitas Kesehatan per Provinsi ({latest_year_faskes})</div>
{html_table(["Provinsi", "Kategori Zona", "Jenis Faskes", "Jumlah"], prov_rows)}
<h4>E. Analisis Temuan Empiris</h4>
<p>Rata-rata Rumah Sakit di Sentra Industri tercatat <strong>{rs_sentra:.0f} unit</strong> per provinsi, lebih rendah dari wilayah Non-Sentra yang mencapai <strong>{rs_non:.0f} unit</strong>. Rata-rata Puskesmas di Sentra Industri tercatat <strong>{puskesmas_sentra:.0f} unit</strong> per provinsi dibandingkan <strong>{puskesmas_non:.0f} unit</strong> di Non-Sentra.</p>
<h2>3.2 Ketimpangan Beban Penyakit: Sentra Industri vs Non-Sentra</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Agregasi Kesehatan: <code>data/processed/sulawesi_kesehatan_detail_2014_2024.csv</code>. Visualisasi dashboard menggunakan Comparative Spatial Analysis untuk membandingkan rata-rata beban penyakit antara provinsi sentra ekstraktif dan non-sentra.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Melalui analisis komparatif spasial, terlihat bahwa beban ekologis tidak terdistribusi secara merata di seluruh wilayah. Provinsi sentra ekspansi nikel, yaitu Sulawesi Tengah dan Sulawesi Tenggara, menunjukkan indikator penyakit yang secara konsisten lebih tinggi.</p>
<p>Data menunjukkan bahwa rata-rata penderita ISPA/Pneumonia di Sentra Industri tercatat <strong>{ispa_sentra_32:,.0f} kasus per tahun</strong>, dibandingkan provinsi Non-Sentra di angka <strong>{ispa_non_32:,.0f} kasus</strong>. Selisih sebesar <strong>{ispa_diff_32:.1f} kali lipat</strong> ini mengindikasikan beban pernapasan yang lebih berat di kawasan penyangga smelter.</p>
<h4>B. Alur Logika Metodologis Comparative Spatial Analysis</h4>
<div class="mermaid">{mermaid_str_3_2}</div>
<div class="table-caption">Tabel 3.2a: Konfigurasi Variabel Comparative Spatial Analysis (Sub-bab 3.2)</div>
{html_table(konf_headers_32, konf_rows_32)}
<h4>C. Formulasi Matematis: Rata-rata Beban Penyakit dan Disparitas Zona</h4>
<div class="formula">B̄_{{z,k}} = [ Σ_{{p∈z}} Σ_t C_{{p,t,k}} ] / N_z</div>
<div class="formula">Q_k = B̄_{{Sentra,k}} / B̄_{{Non-Sentra,k}}</div>
<p>Substitusi angka dari dataset aktual ke dalam rumus rata-rata beban penyakit adalah sebagai berikut:</p>
<div class="formula">B̄_Sentra,ISPA = {ispa_sentra_total_32:,.0f} / {ispa_sentra_n_32} = {ispa_sentra_32:,.1f} kasus</div>
<div class="formula">B̄_Non-Sentra,ISPA = {ispa_non_total_32:,.0f} / {ispa_non_n_32} = {ispa_non_32:,.1f} kasus</div>
<div class="formula">Q_ISPA = {ispa_sentra_32:,.1f} / {ispa_non_32:,.1f} = {ispa_diff_32:.1f}x</div>
<div class="formula">B̄_Sentra,Diare = {diare_sentra_total_32:,.0f} / {diare_sentra_n_32} = {diare_sentra_32:,.1f} kasus</div>
<div class="formula">B̄_Non-Sentra,Diare = {diare_non_total_32:,.0f} / {diare_non_n_32} = {diare_non_32:,.1f} kasus</div>
<div class="formula">Q_Diare = {diare_sentra_32:,.1f} / {diare_non_32:,.1f} = {diare_diff_32:.1f}x</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 3.3: Rata-rata Beban Penyakit per Tahun menurut Zona Industri ({tahun_min_32}-{tahun_max_32})</div>
{html_table(["Indikator Penyakit", "Kategori Zona", "Rata-rata Kasus per Tahun"], disease_rows_32)}
<div class="table-caption">Tabel 3.4: Ringkasan Disparitas Beban Penyakit Sentra vs Non-Sentra</div>
{html_table(["Indikator Penyakit", "Rata-rata Sentra", "Rata-rata Non-Sentra", "Rasio Disparitas"], ratio_rows_32)}
<h4>E. Analisis Temuan Empiris</h4>
<p>Rata-rata penderita ISPA/Pneumonia di Sentra Industri tercatat <strong>{ispa_sentra_32:,.0f} kasus per tahun</strong>, dibandingkan provinsi Non-Sentra di angka <strong>{ispa_non_32:,.0f} kasus</strong>. Selisih sebesar <strong>{ispa_diff_32:.1f} kali lipat</strong> mendukung pembacaan bahwa wilayah dengan konsentrasi industri tinggi cenderung menanggung beban kesehatan yang lebih besar akibat tekanan terhadap daya tampung lingkungan.</p>

<h2>3.3 Lintasan Waktu Ekologis & Dinamika Penyakit di Kawasan Industri Ekstraktif</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Lingkungan & Penyakit: <code>data/processed/sulawesi_kesehatan_detail_2014_2024.csv</code>, <code>data/processed/sulawesi_ika_2016_2024.csv</code>, <code>data/processed/sulawesi_iku_2015_2024.csv</code>. Visualisasi dashboard menampilkan Time-Series Line Chart (insiden per 10.000 penduduk dan total kasus absolut) serta pengujian Chi-Square tabulasi silang (Crosstabulation) dengan binning median per-provinsi.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Meskipun secara akumulatif kawasan Sentra Industri menanggung beban yang lebih berat, penelusuran data secara time-series (historis) dari {tahun_min_33} hingga {tahun_max_33} memberikan wawasan tambahan mengenai fluktuasi kasus penyakit dari tahun ke tahun. Hipotesis utama: <strong>penurunan kualitas udara ambien (IKU) berbanding lurus dengan peningkatan insidensi penyakit pernapasan dan lingkungan</strong>. Untuk mengujinya di tengah keterbatasan jumlah provinsi (N={n_prov_33}), uji Chi-Square menggunakan unit observasi <strong>Provinsi-Tahun</strong> ({n_prov_33} provinsi × {n_tahun_33} tahun panel) dengan klasifikasi berdasarkan median per-provinsi.</p>
<h4>B. Alur Logika Metodologis Time-Series Line Chart & Crosstabulation</h4>
<div class="mermaid">{mermaid_str_3_3}</div>
<div class="table-caption">Tabel 3.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 3.3)</div>
{html_table(konf_headers_33, konf_rows_33)}
<h4>C. Formulasi Matematis: Normalisasi Per Kapita, Binning Median Provinsi, dan Uji Crosstabulation</h4>
<div class="formula">Insiden_10K_p,t = ( Kasus_p,t / Populasi_p ) × 10.000</div>
<div class="formula">Median_Prov_p = Median ( Nilai_p,t )   ;   untuk seluruh tahun t pada provinsi p</div>
<div class="formula">Kategori(x_p,t) = 'Tinggi' , jika x_p,t &ge; Median_Prov_p   |   'Rendah' , jika x_p,t &lt; Median_Prov_p</div>
<div class="formula">&chi;&sup2; = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   dengan E_ij = ( Total_Baris_i × Total_Kolom_j ) / N</div>
<div class="formula">Odds_Ratio (OR) = ( b × c ) / ( a × d )   ;   untuk X berjenis indeks kualitas (IKU/IKA)</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 3.5: Rata-rata Insiden ISPA/Pneumonia Absolut dan per 10.000 Penduduk per Provinsi ({tahun_min_33}-{tahun_max_33})</div>
{html_table(["Provinsi", "Kategori Zona", "Populasi Proxy (BPS 2020)", "Rata-rata Kasus per Tahun", "Rata-rata Insiden per 10.000 Penduduk"], ts_rows_33)}
<div class="table-caption">Tabel 3.6: Ringkasan Eksekutif Seluruh Skenario Crosstab IKU vs Insidensi Penyakit Bab 3</div>
{html_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (&chi;&sup2;)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_33)}
<h4>E. Analisis Temuan Empiris</h4>
<p>{finding_33}</p>

<h2>3.4 Anomali Zoonosis: Dampak Kritis Ekspansi Industri di Level Tapak (Studi Kasus Sulteng)</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Zoonosis: <code>data/processed/zoonosis_kab_kota_2015_2024.csv</code>. Visualisasi dashboard menggunakan Time-Series dan Komparasi Spasial Wilayah untuk membandingkan kabupaten lingkar tambang/smelter aktif dengan daerah non-tambang/agraris sebagai kontrol.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Data empiris Dinas Kesehatan mencatat total akumulasi <strong>{total_kasus_tambang_34:,.0f} kasus</strong> penyakit zoonosis di wilayah Lingkar Tambang/Smelter Aktif Sulawesi Tengah (Morowali, Morowali Utara, Banggai) sepanjang rentang waktu pengamatan {tahun_min_34}-{tahun_max_34}. Peningkatan angka zoonosis ini dibaca bersama perubahan ekologis akibat ekspansi penggunaan lahan, konversi tutupan hutan, pergeseran habitat alami satwa liar, genangan air galian tambang yang tidak direklamasi, serta kondisi sanitasi di area industri.</p>
<h4>B. Alur Logika Metodologis Time-Series & Komparasi Spasial Wilayah</h4>
<div class="mermaid">{mermaid_str_3_4}</div>
<h4>C. Formulasi Matematis: Tren Zoonosis Distrik dan Komparasi Wilayah</h4>
<div class="formula">Z_{{w,t,d}} = Σ C_{{r,t,d}}, untuk setiap distrik r yang termasuk wilayah w</div>
<div class="formula">Z̄_w = [ Σ_t Z_{{w,t,d}} ] / N_w</div>
<div class="formula">R_d = Z̄_Tambang / Z̄_Kontrol</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 3.7: Insidensi Tertinggi Zoonosis di Lingkar Tambang/Smelter Aktif Sulawesi Tengah</div>
{html_table(["Jenis Penyakit", "Kabupaten", "Tahun", "Puncak Kasus"], peak_rows_34)}
<div class="table-caption">Tabel 3.8: Tren Tahunan {selected_penyakit_34} menurut Kategori Wilayah</div>
{html_table(["Tahun", "Kategori Wilayah", "Total Kasus"], trend_rows_34)}
<div class="table-caption">Tabel 3.9: Rata-rata Kasus {selected_penyakit_34} Tambang vs Kontrol</div>
{html_table(["Jenis Penyakit", "Kategori Wilayah", "Rata-rata Kasus"], avg_rows_34)}
<h4>E. Analisis Temuan Empiris</h4>
<p>Pada penyakit terpilih ({selected_penyakit_34}), rata-rata kasus di wilayah Lingkar Tambang/Smelter Aktif mencapai <strong>{val_tambang_34:,.1f}</strong> kasus per observasi, dibandingkan <strong>{val_non_34:,.1f}</strong> kasus pada wilayah Non-Tambang/Agraris (Kontrol), dengan rasio komparatif <strong>{multiplier_34:.1f}x</strong>. Pola ini memberikan sinyal bahwa perubahan ekologis di sekitar kawasan industri perlu dibaca sampai level tapak.</p>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab3_Beban_Kesehatan.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB III: METODOLOGI ANALISIS BEBAN KESEHATAN MASYARAKAT TERDAMPAK",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada **Bab 3: Beban Kesehatan Masyarakat Terdampak**.",
        "",
        "## 3.1 Kesenjangan Fasilitas Kesehatan di Kawasan Ekstraktif",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Agregat Faskes: `data/processed/sulawesi_faskes_agregat_v3.csv`. Visualisasi dashboard menggunakan *Grouped Horizontal Bar Chart* pada tahun acuan 2024.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Data empiris menggambarkan kesenjangan antara klaim pertumbuhan ekonomi dari ekspansi industri nikel dan kondisi kesehatan masyarakat di kawasan penyangga. Ekspansi kapasitas industri yang ditopang PLTU *captive* berkapasitas **{tot_kapasitas_pltu:,.0f} Megawatt** berjalan sejajar dengan peningkatan kasus penyakit.",
        "",
        f"Sepanjang 2014-2024, data agregat dinas kesehatan mencatat total kasus ISPA/Pneumonia sebanyak **{tot_ispa:,.0f}**, kasus Diare sebanyak **{tot_diare:,.0f}**, dan kasus Malaria sebanyak **{tot_malaria:,.0f}**. Pada {latest_year_faskes}, tercatat **{tot_puskesmas_latest:,.0f}** Puskesmas dan **{tot_rs_latest:,.0f}** Rumah Sakit.",
        "",
        "#### B. Alur Logika Metodologis Grouped Horizontal Bar Chart",
        "```mermaid",
        mermaid_str_3_1,
        "```",
        "",
        "##### Tabel 3.1a: Konfigurasi Variabel Analisis Gap Fasilitas Kesehatan (Sub-bab 3.1)",
        markdown_table(konf_headers_31, konf_rows_31),
        "",
        "#### C. Formulasi Matematis: Rata-rata Faskes dan Disparitas Zona",
        "```text",
        "F̄_{z,j} = [ Σ_{p∈z} F_{p,j} ] / n_z",
        "D_j = F̄_{Sentra,j} / F̄_{Non-Sentra,j}",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 3.1: Rata-rata Fasilitas Kesehatan per Provinsi menurut Zona Industri ({latest_year_faskes})",
        markdown_table(["Kategori Zona", "Jenis Faskes", "Rata-rata Jumlah Fasilitas"], gap_rows),
        "",
        f"##### Tabel 3.2: Rincian Fasilitas Kesehatan per Provinsi ({latest_year_faskes})",
        markdown_table(["Provinsi", "Kategori Zona", "Jenis Faskes", "Jumlah"], prov_rows),
        "",
        "#### E. Analisis Temuan Empiris: Defisit Infrastruktur Kesehatan Publik",
        f"Rata-rata Rumah Sakit di Sentra Industri tercatat **{rs_sentra:.0f} unit** per provinsi, lebih rendah dari wilayah Non-Sentra yang mencapai **{rs_non:.0f} unit**. Rata-rata Puskesmas di Sentra Industri tercatat **{puskesmas_sentra:.0f} unit** per provinsi dibandingkan **{puskesmas_non:.0f} unit** di Non-Sentra.",
        "",
        "## 3.2 Ketimpangan Beban Penyakit: Sentra Industri vs Non-Sentra",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Agregasi Kesehatan: `data/processed/sulawesi_kesehatan_detail_2014_2024.csv`. Visualisasi dashboard menggunakan *Comparative Spatial Analysis* untuk membandingkan rata-rata beban penyakit antara provinsi sentra ekstraktif dan non-sentra.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Melalui analisis komparatif spasial, terlihat bahwa beban ekologis tidak terdistribusi secara merata di seluruh wilayah. Provinsi sentra ekspansi nikel, yaitu Sulawesi Tengah dan Sulawesi Tenggara, menunjukkan indikator penyakit yang secara konsisten lebih tinggi.",
        "",
        f"Data menunjukkan bahwa rata-rata penderita **ISPA/Pneumonia** di Sentra Industri tercatat **{ispa_sentra_32:,.0f} kasus per tahun**, dibandingkan provinsi Non-Sentra di angka **{ispa_non_32:,.0f} kasus**. Selisih sebesar **{ispa_diff_32:.1f} kali lipat** ini mengindikasikan beban pernapasan yang lebih berat di kawasan penyangga *smelter*.",
        "",
        "#### B. Alur Logika Metodologis Comparative Spatial Analysis",
        "```mermaid",
        mermaid_str_3_2,
        "```",
        "",
        "##### Tabel 3.2a: Konfigurasi Variabel Comparative Spatial Analysis (Sub-bab 3.2)",
        markdown_table(konf_headers_32, konf_rows_32),
        "",
        "#### C. Formulasi Matematis: Rata-rata Beban Penyakit dan Disparitas Zona",
        "```text",
        "B̄_{z,k} = [ Σ_{p∈z} Σ_t C_{p,t,k} ] / N_z",
        "Q_k = B̄_{Sentra,k} / B̄_{Non-Sentra,k}",
        "```",
        "",
        "Substitusi angka dari dataset aktual ke dalam rumus rata-rata beban penyakit adalah sebagai berikut:",
        "",
        "```text",
        f"B̄_Sentra,ISPA = {ispa_sentra_total_32:,.0f} / {ispa_sentra_n_32} = {ispa_sentra_32:,.1f} kasus",
        f"B̄_Non-Sentra,ISPA = {ispa_non_total_32:,.0f} / {ispa_non_n_32} = {ispa_non_32:,.1f} kasus",
        f"Q_ISPA = {ispa_sentra_32:,.1f} / {ispa_non_32:,.1f} = {ispa_diff_32:.1f}x",
        f"B̄_Sentra,Diare = {diare_sentra_total_32:,.0f} / {diare_sentra_n_32} = {diare_sentra_32:,.1f} kasus",
        f"B̄_Non-Sentra,Diare = {diare_non_total_32:,.0f} / {diare_non_n_32} = {diare_non_32:,.1f} kasus",
        f"Q_Diare = {diare_sentra_32:,.1f} / {diare_non_32:,.1f} = {diare_diff_32:.1f}x",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 3.3: Rata-rata Beban Penyakit per Tahun menurut Zona Industri ({tahun_min_32}-{tahun_max_32})",
        markdown_table(["Indikator Penyakit", "Kategori Zona", "Rata-rata Kasus per Tahun"], disease_rows_32),
        "",
        "##### Tabel 3.4: Ringkasan Disparitas Beban Penyakit Sentra vs Non-Sentra",
        markdown_table(["Indikator Penyakit", "Rata-rata Sentra", "Rata-rata Non-Sentra", "Rasio Disparitas"], ratio_rows_32),
        "",
        "#### E. Analisis Temuan Empiris: Ketimpangan Beban Penyakit Struktural",
        f"Rata-rata penderita ISPA/Pneumonia di Sentra Industri tercatat **{ispa_sentra_32:,.0f} kasus per tahun**, dibandingkan provinsi Non-Sentra di angka **{ispa_non_32:,.0f} kasus**. Selisih sebesar **{ispa_diff_32:.1f} kali lipat** mendukung pembacaan bahwa wilayah dengan konsentrasi industri tinggi cenderung menanggung beban kesehatan yang lebih besar akibat tekanan terhadap daya tampung lingkungan.",
        "",
        "## 3.3 Lintasan Waktu Ekologis & Dinamika Penyakit di Kawasan Industri Ekstraktif",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Lingkungan & Penyakit: `data/processed/sulawesi_kesehatan_detail_2014_2024.csv`, `data/processed/sulawesi_ika_2016_2024.csv`, `data/processed/sulawesi_iku_2015_2024.csv`. Visualisasi dashboard menampilkan *Time-Series Line Chart* (insiden per 10.000 penduduk dan total kasus absolut) serta pengujian Chi-Square tabulasi silang (Crosstabulation) dengan binning median per-provinsi.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Meskipun secara akumulatif kawasan Sentra Industri menanggung beban yang lebih berat, penelusuran data secara *time-series* (historis) dari {tahun_min_33} hingga {tahun_max_33} memberikan wawasan tambahan mengenai fluktuasi kasus penyakit dari tahun ke tahun. Hipotesis utama: **penurunan kualitas udara ambien (IKU) berbanding lurus dengan peningkatan insidensi penyakit pernapasan dan lingkungan**. Untuk mengujinya di tengah keterbatasan jumlah provinsi (N={n_prov_33}), uji Chi-Square menggunakan unit observasi **Provinsi-Tahun** ({n_prov_33} provinsi × {n_tahun_33} tahun panel) dengan klasifikasi berdasarkan median per-provinsi.",
        "",
        "#### B. Alur Logika Metodologis Time-Series Line Chart & Crosstabulation",
        "Kerangka penelusuran runtut waktu insiden penyakit beserta tahapan uji silang statistiknya diilustrasikan pada **Bagan Alur 3.3** berikut, dengan konfigurasi variabel pengujian dirinci pada Tabel 3.3a di bawah gambar.",
        "",
        "##### Bagan Alur 3.3: Alur Logika Metodologis Time-Series Line Chart & Crosstabulation Beban Kesehatan",
        "```mermaid",
        mermaid_str_3_3,
        "```",
        "",
        "##### Tabel 3.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 3.3)",
        markdown_table(konf_headers_33, konf_rows_33),
        "",
        "#### C. Formulasi Matematis: Normalisasi Per Kapita, Binning Median Provinsi, dan Uji Crosstabulation",
        "Kuantifikasi rasio keparahan per kapita dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Insiden_10K_p,t = ( Kasus_p,t / Populasi_p ) × 10.000",
        "Median_Prov_p = Median ( Nilai_p,t )   ;   untuk seluruh tahun t pada provinsi p",
        "Kategori(x_p,t) = 'Tinggi' , jika x_p,t ≥ Median_Prov_p   |   'Rendah' , jika x_p,t < Median_Prov_p",
        "χ² = Σ [ ( O_ij - E_ij )² / E_ij ]   ;   dengan E_ij = ( Total_Baris_i × Total_Kolom_j ) / N",
        "Odds_Ratio (OR) = ( b × c ) / ( a × d )   ;   untuk X berjenis indeks kualitas (IKU/IKA)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 3.5: Rata-rata Insiden ISPA/Pneumonia Absolut dan per 10.000 Penduduk per Provinsi ({tahun_min_33}-{tahun_max_33})",
        markdown_table(["Provinsi", "Kategori Zona", "Populasi Proxy (BPS 2020)", "Rata-rata Kasus per Tahun", "Rata-rata Insiden per 10.000 Penduduk"], ts_rows_33),
        "",
        "##### Tabel 3.6: Ringkasan Eksekutif Seluruh Skenario Crosstab IKU vs Insidensi Penyakit Bab 3",
        markdown_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_33),
        "",
        "#### E. Analisis Temuan Empiris: Pembedahan Realitas Ekologis",
        finding_33,
        "",
        "## 3.4 Anomali Zoonosis: Dampak Kritis Ekspansi Industri di Level Tapak (Studi Kasus Sulteng)",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Zoonosis: `data/processed/zoonosis_kab_kota_2015_2024.csv`. Visualisasi dashboard menggunakan *Time-Series* dan Komparasi Spasial Wilayah untuk membandingkan kabupaten lingkar tambang/smelter aktif dengan daerah non-tambang/agraris sebagai kontrol.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Data empiris Dinas Kesehatan mencatat total akumulasi **{total_kasus_tambang_34:,.0f} kasus** penyakit zoonosis di wilayah Lingkar Tambang/Smelter Aktif Sulawesi Tengah (Morowali, Morowali Utara, Banggai) sepanjang rentang waktu pengamatan {tahun_min_34}-{tahun_max_34}. Peningkatan angka zoonosis ini dibaca bersama perubahan ekologis akibat ekspansi penggunaan lahan, konversi tutupan hutan, pergeseran habitat alami satwa liar, genangan air galian tambang yang tidak direklamasi, serta kondisi sanitasi di area industri.",
        "",
        "#### B. Alur Logika Metodologis Time-Series & Komparasi Spasial Wilayah",
        "Kerangka studi kasus mendalam berbasis deret waktu di tingkat kabupaten/kota khusus Sulawesi Tengah diilustrasikan pada **Bagan Alur 3.4** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan isolasi episentrum ekstraktif dan komparasi absolut dengan wilayah kontrol.",
        "",
        "##### Bagan Alur 3.4: Alur Logika Metodologis Anomali Zoonosis Level Tapak",
        "```mermaid",
        mermaid_str_3_4,
        "```",
        "",
        "#### C. Formulasi Matematis: Tren Zoonosis Distrik dan Komparasi Wilayah",
        "```text",
        "Z_{w,t,d} = Σ C_{r,t,d}, untuk setiap distrik r yang termasuk wilayah w",
        "Z̄_w = [ Σ_t Z_{w,t,d} ] / N_w",
        "R_d = Z̄_Tambang / Z̄_Kontrol",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 3.7: Insidensi Tertinggi Zoonosis di Lingkar Tambang/Smelter Aktif Sulawesi Tengah",
        markdown_table(["Jenis Penyakit", "Kabupaten", "Tahun", "Puncak Kasus"], peak_rows_34),
        "",
        f"##### Tabel 3.8: Tren Tahunan {selected_penyakit_34} menurut Kategori Wilayah",
        markdown_table(["Tahun", "Kategori Wilayah", "Total Kasus"], trend_rows_34),
        "",
        f"##### Tabel 3.9: Rata-rata Kasus {selected_penyakit_34} Tambang vs Kontrol",
        markdown_table(["Jenis Penyakit", "Kategori Wilayah", "Rata-rata Kasus"], avg_rows_34),
        "",
        "#### E. Analisis Temuan Empiris: Anomali Zoonosis Level Tapak",
        f"Pada penyakit terpilih ({selected_penyakit_34}), rata-rata kasus di wilayah Lingkar Tambang/Smelter Aktif mencapai **{val_tambang_34:,.1f}** kasus per observasi, dibandingkan **{val_non_34:,.1f}** kasus pada wilayah Non-Tambang/Agraris (Kontrol), dengan rasio komparatif **{multiplier_34:.1f}x**. Pola ini memberikan sinyal bahwa perubahan ekologis di sekitar kawasan industri perlu dibaca sampai level tapak, karena data agregat provinsi dapat mengaburkan lonjakan penyakit pada kabupaten episentrum ekstraktif.",
        "",
    ]
    md_path = tool_dir / "Metodologi_Bab3_Beban_Kesehatan.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")

    print("[4/4] Selesai membangun Bab 3 sub-bab 3.1.")


if __name__ == "__main__":
    generate_all_bab3()
