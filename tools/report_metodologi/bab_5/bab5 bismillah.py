#!/usr/bin/env python3
"""Generator Laporan Metodologi Bab 5: Pola Penerbitan Izin di Zona Kritis Ekologis."""

import base64
import sys
from pathlib import Path

try:
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pandas", "scipy", "requests", "python-docx"])
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)


def fmt_p(val):
    if pd.isna(val):
        return "p = NaN"
    if val < 0.001:
        return "p < 0.001"
    return f"p = {val:.3f}"


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge, cfg in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        if cfg is None:
            el.set(qn("w:val"), "none")
        else:
            for key, val in cfg.items():
                el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_margin(cell, left=100, right=100, top=60, bottom=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margin.append(el)
    tc_pr.append(margin)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_border_bottom(paragraph, color="2E7D32", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="2E7D32", sz="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="444444", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="1B5E20", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="2E7D32", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=G_MID)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)


def add_p(doc, parts, space_after=5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "EDF7EE")
    all_border_para(p, color="A5D6A7", sz="4")
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)
    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=G_DARK)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p, color="2E7D32", sz="16")
    para_shd(p, "F1F8E9")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}
    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "2E7D32")
        cell_margin(cell, left=100, right=100, top=70, bottom=70)
        set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom={"val": "single", "sz": "8", "color": "1B5E20", "space": "0"}, right=bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8.5, color=C_WHITE)
    for i, row_data in enumerate(rows):
        fill = "F5FBF5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            cell_margin(cell, left=100, right=100, top=50, bottom=50)
            set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def build_crosstab_summary(df_panel, x_options, y_options):
    summary_rows = []
    threshold_rows = []
    for k_x, v_x in x_options.items():
        for k_y, v_y in y_options.items():
            med_x = df_panel[k_x].median()
            med_y = df_panel[k_y].median()
            lbl_x_h = f"Tinggi (>={med_x:,.1f})"
            lbl_x_l = f"Rendah (<{med_x:,.1f})"
            lbl_y_h = f"Tinggi (>={med_y:,.1f})"
            lbl_y_l = f"Rendah (<{med_y:,.1f})"
            s_x = df_panel[k_x].apply(lambda val: lbl_x_h if val >= med_x else lbl_x_l)
            s_y = df_panel[k_y].apply(lambda val: lbl_y_h if val >= med_y else lbl_y_l)
            ct = pd.crosstab(s_x, s_y).reindex(index=[lbl_x_l, lbl_x_h], columns=[lbl_y_l, lbl_y_h], fill_value=0)
            try:
                c2_val, pv_val, _, _ = stats.chi2_contingency(ct)
            except Exception:
                c2_val, pv_val = 0, 1.0
            try:
                aa = ct.loc[lbl_x_l, lbl_y_l]
                bb = ct.loc[lbl_x_l, lbl_y_h]
                cc = ct.loc[lbl_x_h, lbl_y_l]
                dd = ct.loc[lbl_x_h, lbl_y_h]
                or_v = (aa * dd) / (bb * cc) if (bb * cc) > 0 else 0
            except Exception:
                or_v = 0
            summary_rows.append([v_x, v_y, f"{c2_val:.3f}", fmt_p(pv_val), f"{or_v:.2f}", "SIGNIFIKAN" if pv_val < 0.05 else "TIDAK SIGNIFIKAN"])
            threshold_rows.append([v_x, v_y, f"{med_x:,.1f}", f"{med_y:,.1f}", f"{len(df_panel):,}"])
    return summary_rows, threshold_rows


def generate_all_bab5():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    data_dir = base_dir / "data" / "processed"
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_5"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 5...")
    df_izin = pd.read_csv(data_dir / "sulawesi_izin_baru_per_tahun.csv")
    df_gfw = pd.read_csv(data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv")

    total_izin = int(df_izin["Jumlah_Izin_Baru"].sum())
    total_luas_konsesi = float(df_izin["Total_Luas_Konsesi_Baru_Ha"].sum())
    total_deforestasi = float(df_gfw["Total_Deforestasi_Ha"].sum())
    df_izin_thn = df_izin.groupby("Tahun")[["Jumlah_Izin_Baru", "Total_Luas_Konsesi_Baru_Ha"]].sum().reset_index()
    tahun_puncak = int(df_izin_thn.loc[df_izin_thn["Jumlah_Izin_Baru"].idxmax(), "Tahun"])
    izin_puncak = int(df_izin_thn["Jumlah_Izin_Baru"].max())

    df_panel = pd.merge(df_gfw, df_izin, on=["Provinsi", "Tahun"], how="left").fillna({"Jumlah_Izin_Baru": 0, "Total_Luas_Konsesi_Baru_Ha": 0})
    med_def = df_panel["Total_Deforestasi_Ha"].median()
    df_panel["is_kritis"] = df_panel["Total_Deforestasi_Ha"] > med_def
    izin_kritis = int(df_panel[df_panel["is_kritis"]]["Jumlah_Izin_Baru"].sum())
    izin_total = int(df_panel["Jumlah_Izin_Baru"].sum())
    pct_kritis = (izin_kritis / izin_total * 100) if izin_total > 0 else 0
    top_prov_kritis = df_panel[df_panel["is_kritis"]].groupby("Provinsi")["Jumlah_Izin_Baru"].sum().reset_index()
    top_prov_kritis = top_prov_kritis.loc[top_prov_kritis["Jumlah_Izin_Baru"].idxmax()]
    nama_prov_kritis = top_prov_kritis["Provinsi"]
    jumlah_prov_kritis = int(top_prov_kritis["Jumlah_Izin_Baru"])
    izin_pra_2020 = int(df_izin[df_izin["Tahun"] < 2020]["Jumlah_Izin_Baru"].sum())
    izin_pasca_2020 = int(df_izin[df_izin["Tahun"] >= 2020]["Jumlah_Izin_Baru"].sum())
    rasio_akselerasi = izin_pasca_2020 / izin_pra_2020 if izin_pra_2020 > 0 else 0

    df_gfw_thn = df_gfw.groupby("Tahun")["Total_Deforestasi_Ha"].sum().reset_index()
    df_timeline = pd.merge(df_gfw_thn, df_izin_thn, on="Tahun", how="outer").fillna(0).sort_values("Tahun")
    df_timeline = df_timeline[df_timeline["Tahun"] <= 2023].copy()
    tahun_puncak_def = int(df_timeline.loc[df_timeline["Total_Deforestasi_Ha"].idxmax(), "Tahun"])
    def_puncak = float(df_timeline["Total_Deforestasi_Ha"].max())
    tahun_puncak_luas = int(df_timeline.loc[df_timeline["Total_Luas_Konsesi_Baru_Ha"].idxmax(), "Tahun"])
    luas_puncak = float(df_timeline["Total_Luas_Konsesi_Baru_Ha"].max())
    timeline_rows = [[str(int(r["Tahun"])), f"{r['Total_Deforestasi_Ha']:,.0f}", f"{int(r['Jumlah_Izin_Baru']):,}", f"{r['Total_Luas_Konsesi_Baru_Ha']:,.0f}"] for _, r in df_timeline.iterrows()]
    period_rows = [["Pra-2020", f"{izin_pra_2020:,}", "Periode sebelum akselerasi pasca-2020"], ["Pasca-2020", f"{izin_pasca_2020:,}", f"{rasio_akselerasi:.1f}x dibanding pra-2020"], ["Tahun Kritis Ekologis", f"{izin_kritis:,}", f"{pct_kritis:.1f}% dari izin panel GFW-IUP"]]

    x_options_54 = {"Jumlah_Izin_Baru": "Jumlah Izin Baru (IUP)", "Total_Luas_Konsesi_Baru_Ha": "Luas Konsesi Baru (Hektar)"}
    y_options_54 = {"Total_Deforestasi_Ha": "Total Deforestasi Alam (Hektar)", "Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha": "Deforestasi Komoditas Tambang/Sawit (Hektar)"}
    summary_rows_54, threshold_rows_54 = build_crosstab_summary(df_panel, x_options_54, y_options_54)
    sig_count_54 = sum(1 for row in summary_rows_54 if row[5] == "SIGNIFIKAN")
    total_scenarios_54 = len(summary_rows_54)
    finding_54 = f"Dari {total_scenarios_54} skenario pengujian, terdapat {sig_count_54} skenario yang terbukti SIGNIFIKAN. Tingginya Odds Ratio pada skenario signifikan menegaskan bahwa peningkatan penerbitan izin berasosiasi dengan risiko laju deforestasi yang lebih tinggi." if sig_count_54 > 0 else f"Dari {total_scenarios_54} skenario pengujian, seluruhnya menunjukkan status TIDAK SIGNIFIKAN. Hal ini mengindikasikan bahwa laju deforestasi terjadi secara meluas dan dipengaruhi faktor lain di luar jumlah izin baru."

    konf_rows_54 = [["Variabel Independen (X)", "Jumlah_Izin_Baru atau Total_Luas_Konsesi_Baru_Ha."], ["Variabel Dependen (Y)", "Total_Deforestasi_Ha atau Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha."], ["Hipotesis Nol (H0)", "Tidak ada hubungan signifikan antara klasifikasi tingginya penerbitan IUP baru dan klasifikasi tingginya deforestasi."], ["Decision Rule", "Tolak H0 jika P-Value Pearson Chi-Square < 0.05."], ["Unit Observasi", f"Panel Provinsi-Tahun hasil join data izin dan GFW (N={len(df_panel)})."]]

    mermaid_str_5_1 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Izin Baru Minerbaone<br/><i>Tahun, Jumlah IUP, Luas Konsesi</i>"]
        B["Data Deforestasi GFW<br/><i>Tahun, Total Deforestasi Ha</i>"]
    end
    subgraph Timeline_Mapping["2. Sinkronisasi Waktu"]
        A --> C["Agregasi izin & luas konsesi per tahun"]
        B --> D["Agregasi deforestasi per tahun"]
        C --> E["Merge timeline 2014-2023"]
        D --> E
    end
    subgraph Visual_Analysis["3. Dual-Axis Combo Chart"]
        E --> F["Bar: Deforestasi tahunan"]
        E --> G["Line: Area Konsesi IUP & jumlah izin"]
        F --> H["Identifikasi sinkronisasi krisis ekologis dan keputusan izin"]
        G --> H
    end"""
    mermaid_str_5_4 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Izin Baru Minerbaone<br/><i>Jumlah IUP & Luas Konsesi</i>"]
        B["Data Deforestasi GFW<br/><i>Total & Driver Komoditas</i>"]
    end
    subgraph Panel_Join["2. Panel Join Provinsi-Tahun"]
        A --> C["Merge berdasarkan Provinsi dan Tahun"]
        B --> C
        C --> D["Binning median<br/>Tinggi vs Rendah"]
    end
    subgraph Statistical_Test["3. Crosstabulation & Pearson Chi-Square"]
        D --> E["Tabel kontinjensi 2x2"]
        E --> F["Pearson Chi-Square"]
        E --> G["Odds Ratio"]
    end
    F --> H["Pembacaan korelasi izin dan deforestasi"]
    G --> H"""
    mermaid_png_path_5_1 = str(tool_dir / "mermaid_flowchart_5_1.png")
    mermaid_png_path_5_4 = str(tool_dir / "mermaid_flowchart_5_4.png")
    download_success_5_1 = download_mermaid_png(mermaid_str_5_1, mermaid_png_path_5_1)
    download_success_5_4 = download_mermaid_png(mermaid_str_5_4, mermaid_png_path_5_4)

    print("[1.5/4] Mengekstraksi dataset empiris Bab 5 sub-bab 5.2...")
    df_kawasan_52 = pd.read_csv(data_dir / "sulawesi_gfw_kawasan_lindung_loss_2014_2023_v3.csv")
    df_kawasan_52 = df_kawasan_52[(df_kawasan_52["wdpa_protected_areas__iucn_cat"].astype(str) != "0") & (df_kawasan_52["Tahun"] <= 2023)]

    df_pivot_52 = pd.pivot_table(
        df_kawasan_52,
        values="Luas_Hilang_Kawasan_Lindung_Ha",
        index="Tahun",
        columns="wdpa_protected_areas__iucn_cat",
        aggfunc="sum",
        fill_value=0,
    ).reset_index()

    cat1_tahunan_52 = df_pivot_52[1] if 1 in df_pivot_52.columns else pd.Series(0, index=df_pivot_52.index)
    cat2_tahunan_52 = df_pivot_52[2] if 2 in df_pivot_52.columns else pd.Series(0, index=df_pivot_52.index)
    total_tahunan_52 = cat1_tahunan_52 + cat2_tahunan_52
    kum1_52 = cat1_tahunan_52.cumsum()
    kum2_52 = cat2_tahunan_52.cumsum()
    total_kumulatif_52 = kum1_52 + kum2_52

    cat1_total_52 = float(cat1_tahunan_52.sum())
    cat2_total_52 = float(cat2_tahunan_52.sum())
    total_kehancuran_52 = float(total_kumulatif_52.max())
    pct_cat1_52 = cat1_total_52 / total_kehancuran_52 * 100 if total_kehancuran_52 else 0
    pct_cat2_52 = cat2_total_52 / total_kehancuran_52 * 100 if total_kehancuran_52 else 0
    tahun_min_52 = int(df_pivot_52["Tahun"].min())
    tahun_max_52 = int(df_pivot_52["Tahun"].max())
    idx_puncak_52 = total_tahunan_52.idxmax()
    tahun_puncak_52 = int(df_pivot_52.loc[idx_puncak_52, "Tahun"])
    nilai_puncak_52 = float(total_tahunan_52.max())

    kawasan_rows_52 = []
    for i, row in df_pivot_52.iterrows():
        kawasan_rows_52.append([
            str(int(row["Tahun"])),
            f"{cat1_tahunan_52.loc[i]:,.1f}",
            f"{cat2_tahunan_52.loc[i]:,.1f}",
            f"{total_tahunan_52.loc[i]:,.1f}",
            f"{total_kumulatif_52.loc[i]:,.1f}",
        ])

    mermaid_str_5_2 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data GFW Kawasan Lindung Loss<br/><i>Tahun, kategori IUCN livelihood, luas hilang (Ha)</i>"]
    end
    subgraph Spatial_Overlay["2. Geospatial Overlay & Isolasi"]
        A --> B["Filter kategori livelihood valid<br/>Cat 1: Pertanian & Peternakan; Cat 2: Perkebunan Warga"]
        B --> C["Pivot agregasi luas hilang<br/>per Tahun × Kategori"]
        C --> D["Kalkulasi kumulatif<br/>kerusakan permanen 2014-2023"]
    end
    subgraph Visual_Output["3. Stacked Bar Chart Kumulatif"]
        D --> E["Stacked Bar per kategori<br/>+ garis Total Kehancuran Kumulatif"]
        E --> F["Pembacaan tabrakan tata ruang kawasan livelihood"]
    end"""
    mermaid_png_path_5_2 = str(tool_dir / "mermaid_flowchart_5_2.png")
    download_success_5_2 = download_mermaid_png(mermaid_str_5_2, mermaid_png_path_5_2)

    print("[2/4] Membangun DOCX Metodologi_Bab5_Pola_Perizinan.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)
    add_h1(doc, "BAB V: METODOLOGI ANALISIS POLA PENERBITAN IZIN DI ZONA KRITIS EKOLOGIS")
    add_p(doc, [("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada ", False, False), ("Bab 5: Pola Penerbitan Izin di Zona Kritis Ekologis", True, False), (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.", False, False)])

    add_h2(doc, "5.1 Fakta Penyebab: Sinkronisasi Waktu (Timeline Mapping)")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Izin: data/processed/sulawesi_izin_baru_per_tahun.csv; Data Deforestasi: data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv. Visualisasi dashboard menggunakan Dual-Axis Combo Chart.")
    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [(f"Penelusuran data spasial dan temporal di Sulawesi menunjukkan total deforestasi sebesar {total_deforestasi:,.1f} hektar, sementara penerbitan {total_izin:,} izin tambang baru mencakup luas konsesi {total_luas_konsesi:,.1f} hektar. ", False, False), (f"Puncak penerbitan izin tercatat pada tahun {tahun_puncak} ({izin_puncak} izin), sedangkan {pct_kritis:.1f}% izin panel terbit pada tahun-tahun ketika laju deforestasi provinsi berada di atas median.", False, False)])
    add_h4(doc, "B. Alur Logika Metodologis Sinkronisasi Waktu (Timeline Mapping)")
    add_p(doc, [("Kerangka sinkronisasi waktu antara penerbitan izin dan deforestasi diilustrasikan pada ", False, False), ("Bagan Alur 5.1", True, False), (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Timeline Mapping.", False, False)])
    add_caption(doc, "Bagan Alur 5.1: Alur Logika Metodologis Timeline Mapping Izin vs Deforestasi")
    if download_success_5_1:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_5_1, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)
    add_h4(doc, "C. Formulasi Matematis: Agregasi Tahunan dan Akselerasi Izin")
    add_formula(doc, "Total Deforestasi Tahunan", "D_t = Σ D_{p,t}, untuk seluruh provinsi p pada tahun t", [("D_t", "Total deforestasi tahunan di wilayah Sulawesi."), ("D_{p,t}", "Deforestasi provinsi p pada tahun t.")])
    add_formula(doc, "Total IUP dan Luas Konsesi Tahunan", "I_t = Σ I_{p,t};  L_t = Σ L_{p,t}", [("I_t", "Total jumlah izin tambang baru pada tahun t."), ("L_t", "Total luas konsesi tambang baru pada tahun t.")])
    add_formula(doc, "Substitusi Rasio Akselerasi", f"R = {izin_pasca_2020:,} / {izin_pra_2020:,} = {rasio_akselerasi:.1f}x")
    add_h4(doc, "D. Matriks Hasil Uji Empiris: Sinkronisasi Timeline Izin dan Deforestasi")
    add_caption(doc, "Tabel 5.1: Agregasi Waktu Historis Izin Tambang dan Deforestasi (2014-2023)")
    add_table_1col(doc, ["Tahun", "Total Deforestasi (Ha)", "Jumlah IUP Baru", "Luas Konsesi Baru (Ha)"], timeline_rows, [2.0, 4.0, 3.0, 4.0], ["C", "C", "C", "C"])
    add_caption(doc, "Tabel 5.2: Ringkasan Periode Kritis Penerbitan Izin")
    add_table_1col(doc, ["Periode/Indikator", "Nilai", "Keterangan"], period_rows, [3.8, 2.2, 6.0], ["L", "C", "L"])
    add_h4(doc, "E. Analisis Temuan Empiris: Sinkronisasi Krisis Ekologis dan Keputusan Perizinan")
    add_p(doc, [(f"Puncak deforestasi tahunan tercatat pada tahun {tahun_puncak_def} sebesar {def_puncak:,.0f} hektar, sedangkan puncak luas konsesi IUP baru tercatat pada tahun {tahun_puncak_luas} sebesar {luas_puncak:,.0f} hektar. ", False, False), (f"Provinsi dengan penerbitan izin tertinggi pada periode deforestasi kritis adalah {nama_prov_kritis} dengan {jumlah_prov_kritis} IUP.", False, False)])

    add_h2(doc, "5.2 Fakta Spasial: Tabrakan Tata Ruang di Kawasan Konservasi")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data GFW Overlay Kawasan Lindung: data/processed/sulawesi_gfw_kawasan_lindung_loss_2014_2023_v3.csv (GFW dengan overlay Livelihood Zone Proxy Kategori 1 & 2). Visualisasi dashboard menggunakan Stacked Bar Chart kumulatif per kategori livelihood dengan garis Total Kehancuran Kumulatif 2014-2023.")
    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Dataset spasial menunjukkan pentingnya kepatuhan terhadap batas-batas tata ruang. Kawasan penyangga kehidupan (Hutan Produksi, Kawasan Lindung, dan Area Resapan Air) memerlukan pengawasan agar fungsi ekologisnya tetap terjaga di tengah ekspansi pertambangan. ", False, False),
        ("Analisis dilakukan dengan mengisolasi data tree cover loss (GFW) yang secara spesifik bertumpukan/beririsan dengan poligon Kawasan Livelihood (Zona Pertanian, Peternakan) dan Perkebunan Warga, lalu mengkalkulasi kehancuran agregat kawasan penyangga ekosistem esensial selama satu dekade terakhir akibat penetrasi aktivitas tambang.", False, False),
    ])
    add_h4(doc, "B. Alur Logika Metodologis Overlay Area Kawasan Lindung (GFW)")
    add_p(doc, [
        ("Kerangka agregasi spasial bertingkat untuk mendokumentasikan skala kehancuran kawasan livelihood diilustrasikan pada ", False, False),
        ("Bagan Alur 5.2", True, False),
        (" berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Geospatial Overlay dan kuantifikasi kerusakan kumulatif deskriptif.", False, False),
    ])
    add_caption(doc, "Bagan Alur 5.2: Alur Logika Analisis Overlay Spasial Kawasan Livelihood")
    if download_success_5_2:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_5_2, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)
    add_h4(doc, "C. Formulasi Matematis: Isolasi Overlay dan Akumulasi Kerusakan")
    add_formula(doc, "Persamaan Isolasi Luas Hancur per Kategori Livelihood", "Luas_Hancur_c(t) = Σ ( Loss_i )   ;   untuk seluruh observasi i dengan Kategori_Livelihood = c pada tahun t", [
        ("Luas_Hancur_c(t)", "Total luas kehilangan tutupan (Ha) kategori livelihood c pada tahun t."),
        ("Loss_i", "Luas tree cover loss (Ha) observasi ke-i yang beririsan dengan poligon kawasan livelihood."),
        ("c", "Kategori livelihood: c=1 Zona Pertanian & Peternakan; c=2 Perkebunan Warga."),
    ])
    add_formula(doc, "Persamaan Akumulasi Kerusakan Permanen", "Kumulatif_Hancur_c(T) = Σ Luas_Hancur_c(t)   ;   untuk t = 2014 s.d. T", [
        ("Kumulatif_Hancur_c(T)", "Akumulasi kerusakan kategori c hingga tahun berjalan T; bersifat permanen karena hutan yang hilang tidak kembali."),
    ])
    add_formula(doc, "Persamaan Total Kehancuran Kumulatif", "Total_Kumulatif(T) = Kumulatif_Hancur_1(T) + Kumulatif_Hancur_2(T)", [
        ("Total_Kumulatif(T)", f"Total kehancuran kawasan livelihood hingga tahun T; nilai akhir dekade mencapai {total_kehancuran_52:,.1f} Ha."),
    ])
    add_h4(doc, "D. Matriks Hasil Uji Empiris: Rincian Kehancuran Kawasan Livelihood per Tahun")
    add_caption(doc, f"Tabel 5.3: Rincian Kehancuran Kawasan Livelihood Warga per Tahun ({tahun_min_52}-{tahun_max_52})")
    add_table_1col(doc, ["Tahun", "Pertanian & Peternakan (Ha)", "Perkebunan Warga (Ha)", "Total Tahunan (Ha)", "Total Kumulatif (Ha)"], kawasan_rows_52, [1.8, 3.4, 3.2, 3.0, 3.2], ["C", "C", "C", "C", "C"])
    add_h4(doc, "E. Analisis Temuan Empiris: Fakta Spasial Tabrakan Tata Ruang")
    add_p(doc, [
        ("1. ", True, False), ("Skala Kehancuran Dekade: ", True, False),
        (f"Dalam dekade terakhir, total lebih dari {total_kehancuran_52 / 1000:,.1f} ribu hektar kawasan livelihood (Pertanian, Peternakan, dan Perkebunan) warga tercatat mengalami perubahan tutupan lahan yang beririsan dengan dinamika industri ekstraktif.\n", False, False),
        ("2. ", True, False), ("Komposisi Kategori: ", True, False),
        (f"Zona Pertanian & Peternakan menyumbang {cat1_total_52:,.1f} Ha ({pct_cat1_52:.1f}%) dan Perkebunan Warga {cat2_total_52:,.1f} Ha ({pct_cat2_52:.1f}%) dari total kehancuran kumulatif.\n", False, False),
        ("3. ", True, False), ("Tahun Lonjakan Tertinggi: ", True, False),
        (f"Kehancuran tahunan terbesar tercatat pada tahun {tahun_puncak_52} sebesar {nilai_puncak_52:,.1f} Ha. Karena kerusakan bersifat permanen, akumulasi ini menegaskan pentingnya kepatuhan batas tata ruang dan pengawasan kawasan penyangga ekosistem esensial.", False, False),
    ])

    print("[1.5/4] Mengekstraksi dataset empiris Bab 5 sub-bab 5.3...")
    df_konflik_53 = pd.read_csv(data_dir / "sulawesi_konflik_tambang_fpic.csv")
    df_masalah_53 = pd.read_csv(data_dir / "kpa_masalah_izin_perusahaan.csv")

    total_konflik_53 = len(df_konflik_53)
    konflik_fpic_53 = int(df_konflik_53['indikasi_fpic'].sum())
    total_masalah_izin_53 = len(df_masalah_53)
    perusahaan_masalah_sulawesi_53 = int(df_masalah_53[df_masalah_53['lokasi'].str.contains('Sulawesi', case=False, na=False)]['nama_perusahaan'].nunique())

    mermaid_str_5_3 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Konflik Tambang<br/><i>Kasus, Luas, Indikasi FPIC</i>"]
        B["Data Masalah Izin KPA<br/><i>Perusahaan, Jenis Masalah</i>"]
    end
    subgraph Cross_Dataset["2. Cross-Dataset Integration"]
        A --> C["Kuantifikasi Pelanggaran FPIC"]
        B --> D["Rekam Jejak Perusahaan"]
        C --> E["Agregasi Anomali Perizinan"]
        D --> E
    end
    subgraph Output["3. Realitas Lapangan"]
        E --> F["Identifikasi Masyarakat Dikorbankan"]
    end"""
    mermaid_png_path_5_3 = str(tool_dir / "mermaid_flowchart_5_3.png")
    download_success_5_3 = download_mermaid_png(mermaid_str_5_3, mermaid_png_path_5_3)

    add_h2(doc, "5.3 Realitas Lapangan: Izin Bermasalah, FPIC Diabaikan, Masyarakat Dikorbankan")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Konflik: data/processed/sulawesi_konflik_tambang_fpic.csv; Data Masalah Izin: data/processed/kpa_masalah_izin_perusahaan.csv. Metode: Cross-Dataset Integration.")
    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Laporan dari CRI, Mighty Earth, dan BHRRC mendokumentasikan isu tata kelola perizinan dan pelaksanaan konsultasi publik (FPIC). Penelusuran terhadap database Konsorsium Pembaruan Agraria (KPA) CATAHU 2016-2025 dan Tanahkita.id mengidentifikasi ", False, False),
        (f"{total_masalah_izin_53} kasus permasalahan izin perusahaan", True, False),
        (f" secara nasional, dengan {perusahaan_masalah_sulawesi_53} entitas yang bermasalah di wilayah Sulawesi. ", False, False),
        (f"Di Sulawesi, tercatat {total_konflik_53} kasus konflik pertambangan (2014-2024) dengan {konflik_fpic_53} kasus yang mencatatkan indikasi isu pelaksanaan FPIC.", True, False),
    ])
    add_h4(doc, "B. Alur Logika Metodologis Pelanggaran FPIC")
    add_caption(doc, "Bagan Alur 5.3: Alur Logika Metodologis Integrasi Data Konflik dan Pelanggaran FPIC")
    if download_success_5_3:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_5_3, width=Cm(15))
        except Exception:
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)

    add_h4(doc, "C. Formulasi Matematis: Kuantifikasi Pelanggaran")
    add_p(doc, [("Pemodelan indikator pelanggaran dan rekam jejak perusahaan:", False, False)])
    add_formula(doc, "Total Pelanggaran FPIC", "Total_Pelanggaran_FPIC = Σ Kasus, untuk indikasi_fpic = True")
    add_formula(doc, "Rekam Jejak Oligarki", "Rekam_Jejak = Σ Jenis_Masalah_Izin, dikelompokkan berdasarkan nama_perusahaan")

    add_h4(doc, "D. Matriks Hasil Uji Empiris")
    add_caption(doc, "Tabel 5.3: Metrik Konflik dan Pelanggaran FPIC")
    add_table_1col(doc, ["Indikator", "Total Kasus"], [
        ["Total Konflik Pertambangan Sulawesi", str(total_konflik_53)],
        ["Kasus Indikasi Pelanggaran FPIC", str(konflik_fpic_53)],
        ["Perusahaan Bermasalah di Sulawesi", str(perusahaan_masalah_sulawesi_53)],
        ["Total Permasalahan Izin KPA (Nasional)", str(total_masalah_izin_53)]
    ], [8.0, 4.0], ["L", "C"])
    
    add_h4(doc, "E. Analisis Temuan Empiris: Pembuktian Realitas Lapangan")
    add_p(doc, [
        (f"Dari {total_konflik_53} konflik pertambangan di Sulawesi, {konflik_fpic_53} di antaranya secara eksplisit terkait dengan pelanggaran persetujuan awal tanpa paksaan (FPIC). Hal ini menggarisbawahi perlunya penguatan sistem evaluasi perizinan dan penghormatan terhadap hak-hak komunitas lokal.", False, False)
    ])

    add_h2(doc, "5.4 Pembuktian Empiris: Uji Statistik Korelasi Penerbitan Izin & Deforestasi")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Panel Join dari sulawesi_izin_baru_per_tahun.csv dan sulawesi_gfw_master_1_dekade_2014_2023_v3.csv. Visualisasi dashboard menggunakan Crosstabulation & Pearson Chi-Square Test.")
    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [("Sub-bab ini menggunakan pengujian statistik inferensial untuk membuktikan secara matematis apakah besaran jumlah perizinan baru menjadi prediktor kuat terhadap tingkat kerusakan deforestasi. ", False, False), ("Data numerik berkelanjutan dikategorikan menjadi Tinggi dan Rendah menggunakan ambang batas median dari distribusi panel.", False, False)])
    add_h4(doc, "B. Alur Logika Metodologis Crosstabulation & Pearson Chi-Square Test")
    add_caption(doc, "Bagan Alur 5.4: Alur Logika Metodologis Uji Korelasi Izin dan Deforestasi")
    if download_success_5_4:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(mermaid_png_path_5_4, width=Cm(15))
    else:
        run(doc.add_paragraph(), "[Gambar Flowchart Gagal Diunduh]", color=C_RED, pt=9)
    add_caption(doc, "Tabel 5.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 5.4)")
    add_table_1col(doc, ["Komponen Uji", "Definisi Variabel (Sub-bab 5.4)"], konf_rows_54, [4.5, 11.0], ["L", "L"])
    add_h4(doc, "C. Formulasi Matematis: Binning Median, Chi-Square, dan Odds Ratio")
    add_formula(doc, "Kategorisasi Median Panel", "Kategori(x) = Tinggi jika x ≥ Median(Panel); Rendah jika x < Median(Panel)")
    add_formula(doc, "Uji Chi-Square Pearson", "χ² = Σ [ ( O_i - E_i )² / E_i ]", [("O_i", "Frekuensi observasi pada sel kontinjensi."), ("E_i", "Frekuensi harapan jika X dan Y independen.")])
    add_formula(doc, "Odds Ratio", "OR = ( a × d ) / ( b × c )", [("a,b,c,d", "Empat sel dalam tabel kontinjensi 2x2.")])
    add_h4(doc, "D. Matriks Hasil Uji Empiris: Skenario Crosstab Izin dan Deforestasi")
    add_caption(doc, "Tabel 5.4: Ambang Median Panel Uji Crosstab")
    add_table_1col(doc, ["Variabel X", "Variabel Y", "Median X", "Median Y", "N"], threshold_rows_54, [3.0, 3.4, 2.2, 2.2, 1.4], ["L", "L", "C", "C", "C"])
    add_caption(doc, "Tabel 5.5: Ringkasan Eksekutif Seluruh Skenario Crosstab Izin vs Deforestasi")
    add_table_1col(doc, ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_54, [3.0, 3.4, 1.8, 1.8, 1.8, 2.3], ["L", "L", "C", "C", "C", "C"])
    add_h4(doc, "E. Analisis Temuan Empiris: Signifikansi Korelasi Perizinan dan Ekstraksi Ekologis")
    add_p(doc, [(finding_54, False, False)])

    docx_path = tool_dir / "Metodologi_Bab5_Pola_Perizinan.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 5...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 5 - Pola Penerbitan Izin</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: Arial, sans-serif; max-width: 920px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #43A047; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; }}
.hdr-title {{ color: #81C784; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #2E7D32; padding-bottom: 8px; }}
h2 {{ color: #81C784; text-transform: uppercase; border-bottom: 1px solid #2E7D32; }}
h4 {{ color: #A5D6A7; }}
.note-box {{ background: #132213; border-left: 4px solid #2E7D32; padding: 10px 14px; margin: 12px 0; }}
.formula {{ background: #0D1B0E; border: 1px solid #2E7D32; color: #A5D6A7; padding: 8px 12px; font-family: monospace; }}
.data-th {{ background: #1B5E20; color: white; padding: 6px; text-align: left; border: 1px solid #2E7D32; }}
.data-td {{ padding: 6px; border: 1px solid #243524; vertical-align: top; }}
.data-tr-even .data-td {{ background: #131B13; }}
.mermaid {{ background: #0D1610; border: 1px solid #2E7D32; padding: 12px; margin: 10px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB V: Metodologi Analisis Pola Penerbitan Izin di Zona Kritis Ekologis</div>
<h2>5.1 Fakta Penyebab: Sinkronisasi Waktu (Timeline Mapping)</h2>
<p>Penelusuran data spasial dan temporal di Sulawesi menunjukkan total deforestasi sebesar <strong>{total_deforestasi:,.1f} hektar</strong>, sementara penerbitan <strong>{total_izin:,} izin tambang baru</strong> mencakup luas konsesi <strong>{total_luas_konsesi:,.1f} hektar</strong>.</p>
<div class="mermaid">{mermaid_str_5_1}</div>
<div class="table-caption">Tabel 5.1: Agregasi Waktu Historis Izin Tambang dan Deforestasi (2014-2023)</div>
{html_table(["Tahun", "Total Deforestasi (Ha)", "Jumlah IUP Baru", "Luas Konsesi Baru (Ha)"], timeline_rows)}
<div class="table-caption">Tabel 5.2: Ringkasan Periode Kritis Penerbitan Izin</div>
{html_table(["Periode/Indikator", "Nilai", "Keterangan"], period_rows)}
<h2>5.2 Fakta Spasial: Tabrakan Tata Ruang di Kawasan Konservasi</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data GFW Overlay Kawasan Lindung: <code>data/processed/sulawesi_gfw_kawasan_lindung_loss_2014_2023_v3.csv</code> (GFW dengan overlay Livelihood Zone Proxy Kategori 1 & 2). Visualisasi dashboard menggunakan Stacked Bar Chart kumulatif per kategori livelihood dengan garis Total Kehancuran Kumulatif 2014-2023.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Dataset spasial menunjukkan pentingnya kepatuhan terhadap batas-batas tata ruang. Analisis mengisolasi data tree cover loss (GFW) yang beririsan dengan poligon Kawasan Livelihood (Zona Pertanian, Peternakan) dan Perkebunan Warga, lalu mengkalkulasi kehancuran agregat kawasan penyangga ekosistem esensial selama satu dekade terakhir akibat penetrasi aktivitas tambang.</p>
<h4>B. Alur Logika Metodologis Overlay Area Kawasan Lindung (GFW)</h4>
<p>Kerangka agregasi spasial bertingkat diilustrasikan pada <strong>Bagan Alur 5.2</strong> berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Geospatial Overlay dan kuantifikasi kerusakan kumulatif deskriptif.</p>
<div class="table-caption">Bagan Alur 5.2: Alur Logika Analisis Overlay Spasial Kawasan Livelihood</div>
<div class="mermaid">{mermaid_str_5_2}</div>
<h4>C. Formulasi Matematis: Isolasi Overlay dan Akumulasi Kerusakan</h4>
<div class="formula">Luas_Hancur_c(t) = Σ ( Loss_i )   ;   untuk seluruh observasi i dengan Kategori_Livelihood = c pada tahun t</div>
<div class="formula">Kumulatif_Hancur_c(T) = Σ Luas_Hancur_c(t)   ;   untuk t = 2014 s.d. T</div>
<div class="formula">Total_Kumulatif(T) = Kumulatif_Hancur_1(T) + Kumulatif_Hancur_2(T)</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 5.3: Rincian Kehancuran Kawasan Livelihood Warga per Tahun ({tahun_min_52}-{tahun_max_52})</div>
{html_table(["Tahun", "Pertanian & Peternakan (Ha)", "Perkebunan Warga (Ha)", "Total Tahunan (Ha)", "Total Kumulatif (Ha)"], kawasan_rows_52)}
<h4>E. Analisis Temuan Empiris</h4>
<p><strong>1. Skala Kehancuran Dekade:</strong> total lebih dari <strong>{total_kehancuran_52 / 1000:,.1f} ribu hektar</strong> kawasan livelihood warga mengalami perubahan tutupan lahan yang beririsan dengan dinamika industri ekstraktif. <strong>2. Komposisi Kategori:</strong> Zona Pertanian & Peternakan {cat1_total_52:,.1f} Ha ({pct_cat1_52:.1f}%), Perkebunan Warga {cat2_total_52:,.1f} Ha ({pct_cat2_52:.1f}%). <strong>3. Tahun Lonjakan Tertinggi:</strong> {tahun_puncak_52} sebesar {nilai_puncak_52:,.1f} Ha — karena kerusakan bersifat permanen, akumulasi ini menegaskan pentingnya kepatuhan batas tata ruang.</p>
<h2>5.3 Realitas Lapangan: Izin Bermasalah, FPIC Diabaikan, Masyarakat Dikorbankan</h2>
<div class="note-box"><strong>Metode:</strong> Cross-Dataset Integration (KPA CATAHU + Tanahkita + CRI/Mighty Earth Reports)</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Laporan mendokumentasikan isu tata kelola perizinan dan pelaksanaan konsultasi publik (FPIC). Penelusuran terhadap database Konsorsium Pembaruan Agraria (KPA) CATAHU mengidentifikasi <strong>{total_masalah_izin_53} kasus permasalahan izin perusahaan</strong>. Di Sulawesi, tercatat <strong>{total_konflik_53} kasus konflik pertambangan</strong> dengan <strong>{konflik_fpic_53} kasus yang mencatatkan indikasi isu pelaksanaan FPIC</strong>.</p>
<h4>B. Alur Logika Metodologis Pelanggaran FPIC</h4>
<div class="mermaid">{mermaid_str_5_3}</div>
<h4>C. Formulasi Matematis</h4>
<div class="formula">Total_Pelanggaran_FPIC = Σ Kasus, untuk indikasi_fpic = True</div>
<div class="formula">Rekam_Jejak = Σ Jenis_Masalah_Izin, dikelompokkan berdasarkan nama_perusahaan</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<div class="table-caption">Tabel 5.3: Metrik Konflik dan Pelanggaran FPIC</div>
{html_table(["Indikator", "Total Kasus"], [
    ["Total Konflik Pertambangan Sulawesi", str(total_konflik_53)],
    ["Kasus Indikasi Pelanggaran FPIC", str(konflik_fpic_53)],
    ["Perusahaan Bermasalah di Sulawesi", str(perusahaan_masalah_sulawesi_53)]
])}
<h4>E. Analisis Temuan Empiris</h4>
<p>Dari {total_konflik_53} konflik pertambangan di Sulawesi, {konflik_fpic_53} di antaranya terkait pelanggaran FPIC.</p>

<h2>5.4 Pembuktian Empiris: Uji Statistik Korelasi Penerbitan Izin & Deforestasi</h2>
<div class="mermaid">{mermaid_str_5_4}</div>
<div class="table-caption">Tabel 5.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 5.4)</div>
{html_table(["Komponen Uji", "Definisi Variabel (Sub-bab 5.4)"], konf_rows_54)}
<h4>C. Formulasi Matematis</h4>
<div class="formula">Kategori(x) = Tinggi jika x ≥ Median(Panel); Rendah jika x &lt; Median(Panel)</div>
<div class="formula">χ² = Σ [ ( O_i - E_i )² / E_i ]</div>
<div class="formula">OR = ( a × d ) / ( b × c )</div>
<div class="table-caption">Tabel 5.4: Ambang Median Panel Uji Crosstab</div>
{html_table(["Variabel X", "Variabel Y", "Median X", "Median Y", "N"], threshold_rows_54)}
<div class="table-caption">Tabel 5.5: Ringkasan Eksekutif Seluruh Skenario Crosstab Izin vs Deforestasi</div>
{html_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_54)}
<p>{finding_54}</p>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab5_Pola_Perizinan.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB V: METODOLOGI ANALISIS POLA PENERBITAN IZIN DI ZONA KRITIS EKOLOGIS",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, prosedur pengolahan data, dan pembacaan empiris yang dioperasionalkan pada **Bab 5: Pola Penerbitan Izin di Zona Kritis Ekologis**.",
        "",
        "## 5.1 Fakta Penyebab: Sinkronisasi Waktu (Timeline Mapping)",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Izin: `data/processed/sulawesi_izin_baru_per_tahun.csv`; Data Deforestasi: `data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv`.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Penelusuran data spasial dan temporal di Sulawesi menunjukkan total deforestasi sebesar **{total_deforestasi:,.1f} hektar**, sementara penerbitan **{total_izin:,} izin tambang baru** mencakup luas konsesi **{total_luas_konsesi:,.1f} hektar**. Puncak penerbitan izin tercatat pada tahun **{tahun_puncak}** ({izin_puncak} izin).",
        "",
        f"Sebanyak **{pct_kritis:.1f}%** izin panel terbit pada tahun-tahun ketika laju deforestasi provinsi berada di atas median. Provinsi dengan penerbitan izin tertinggi pada periode deforestasi kritis adalah **{nama_prov_kritis}** dengan **{jumlah_prov_kritis} IUP**.",
        "",
        "#### B. Alur Logika Metodologis Sinkronisasi Waktu (Timeline Mapping)",
        "```mermaid",
        mermaid_str_5_1,
        "```",
        "",
        "#### C. Formulasi Matematis: Agregasi Tahunan dan Akselerasi Izin",
        "```text",
        "D_t = Σ D_{p,t}, untuk seluruh provinsi p pada tahun t",
        "I_t = Σ I_{p,t};  L_t = Σ L_{p,t}",
        f"R = {izin_pasca_2020:,} / {izin_pra_2020:,} = {rasio_akselerasi:.1f}x",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 5.1: Agregasi Waktu Historis Izin Tambang dan Deforestasi (2014-2023)",
        markdown_table(["Tahun", "Total Deforestasi (Ha)", "Jumlah IUP Baru", "Luas Konsesi Baru (Ha)"], timeline_rows),
        "",
        "##### Tabel 5.2: Ringkasan Periode Kritis Penerbitan Izin",
        markdown_table(["Periode/Indikator", "Nilai", "Keterangan"], period_rows),
        "",
        "#### E. Analisis Temuan Empiris: Sinkronisasi Krisis Ekologis dan Keputusan Perizinan",
        f"Puncak deforestasi tahunan tercatat pada tahun **{tahun_puncak_def}** sebesar **{def_puncak:,.0f} hektar**, sedangkan puncak luas konsesi IUP baru tercatat pada tahun **{tahun_puncak_luas}** sebesar **{luas_puncak:,.0f} hektar**.",
        "",
        "## 5.2 Fakta Spasial: Tabrakan Tata Ruang di Kawasan Konservasi",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data GFW Overlay Kawasan Lindung: `data/processed/sulawesi_gfw_kawasan_lindung_loss_2014_2023_v3.csv` (GFW dengan overlay Livelihood Zone Proxy Kategori 1 & 2). Visualisasi dashboard menggunakan Stacked Bar Chart kumulatif per kategori livelihood dengan garis Total Kehancuran Kumulatif 2014-2023.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Dataset spasial menunjukkan pentingnya kepatuhan terhadap batas-batas tata ruang. Analisis mengisolasi data tree cover loss (GFW) yang beririsan dengan poligon Kawasan Livelihood (Zona Pertanian, Peternakan) dan Perkebunan Warga, lalu mengkalkulasi kehancuran agregat kawasan penyangga ekosistem esensial selama satu dekade terakhir akibat penetrasi aktivitas tambang.",
        "",
        "#### B. Alur Logika Metodologis Overlay Area Kawasan Lindung (GFW)",
        "Kerangka agregasi spasial bertingkat diilustrasikan pada **Bagan Alur 5.2** berikut. Sub-bab ini tidak menggunakan uji inferensial Chi-Square, melainkan Geospatial Overlay dan kuantifikasi kerusakan kumulatif deskriptif.",
        "",
        "##### Bagan Alur 5.2: Alur Logika Analisis Overlay Spasial Kawasan Livelihood",
        "```mermaid",
        mermaid_str_5_2,
        "```",
        "",
        "#### C. Formulasi Matematis: Isolasi Overlay dan Akumulasi Kerusakan",
        "```text",
        "Luas_Hancur_c(t) = Σ ( Loss_i )   ;   untuk seluruh observasi i dengan Kategori_Livelihood = c pada tahun t",
        "Kumulatif_Hancur_c(T) = Σ Luas_Hancur_c(t)   ;   untuk t = 2014 s.d. T",
        "Total_Kumulatif(T) = Kumulatif_Hancur_1(T) + Kumulatif_Hancur_2(T)",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        f"##### Tabel 5.3: Rincian Kehancuran Kawasan Livelihood Warga per Tahun ({tahun_min_52}-{tahun_max_52})",
        markdown_table(["Tahun", "Pertanian & Peternakan (Ha)", "Perkebunan Warga (Ha)", "Total Tahunan (Ha)", "Total Kumulatif (Ha)"], kawasan_rows_52),
        "",
        "#### E. Analisis Temuan Empiris: Fakta Spasial Tabrakan Tata Ruang",
        f"1. **Skala Kehancuran Dekade:** total lebih dari **{total_kehancuran_52 / 1000:,.1f} ribu hektar** kawasan livelihood (Pertanian, Peternakan, dan Perkebunan) warga tercatat mengalami perubahan tutupan lahan yang beririsan dengan dinamika industri ekstraktif.",
        f"2. **Komposisi Kategori:** Zona Pertanian & Peternakan menyumbang {cat1_total_52:,.1f} Ha ({pct_cat1_52:.1f}%) dan Perkebunan Warga {cat2_total_52:,.1f} Ha ({pct_cat2_52:.1f}%) dari total kehancuran kumulatif.",
        f"3. **Tahun Lonjakan Tertinggi:** kehancuran tahunan terbesar tercatat pada tahun {tahun_puncak_52} sebesar {nilai_puncak_52:,.1f} Ha — karena kerusakan bersifat permanen, akumulasi ini menegaskan pentingnya kepatuhan batas tata ruang dan pengawasan kawasan penyangga ekosistem esensial.",
        "",
        "## 5.3 Realitas Lapangan: Izin Bermasalah, FPIC Diabaikan, Masyarakat Dikorbankan",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Konflik: `data/processed/sulawesi_konflik_tambang_fpic.csv`; Data Masalah Izin: `data/processed/kpa_masalah_izin_perusahaan.csv`. Metode: Cross-Dataset Integration.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Laporan mendokumentasikan isu tata kelola perizinan dan pelaksanaan konsultasi publik (FPIC). Penelusuran terhadap database Konsorsium Pembaruan Agraria (KPA) CATAHU mengidentifikasi **{total_masalah_izin_53} kasus permasalahan izin perusahaan**. Di Sulawesi, tercatat **{total_konflik_53} kasus konflik pertambangan** dengan **{konflik_fpic_53} kasus yang mencatatkan indikasi isu pelaksanaan FPIC**.",
        "",
        "#### B. Alur Logika Metodologis Pelanggaran FPIC",
        "```mermaid",
        mermaid_str_5_3,
        "```",
        "",
        "#### C. Formulasi Matematis: Kuantifikasi Pelanggaran",
        "```text",
        "Total_Pelanggaran_FPIC = Σ Kasus, untuk indikasi_fpic = True",
        "Rekam_Jejak = Σ Jenis_Masalah_Izin, dikelompokkan berdasarkan nama_perusahaan",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 5.3: Metrik Konflik dan Pelanggaran FPIC",
        markdown_table(["Indikator", "Total Kasus"], [
            ["Total Konflik Pertambangan Sulawesi", str(total_konflik_53)],
            ["Kasus Indikasi Pelanggaran FPIC", str(konflik_fpic_53)],
            ["Perusahaan Bermasalah di Sulawesi", str(perusahaan_masalah_sulawesi_53)]
        ]),
        "",
        "#### E. Analisis Temuan Empiris: Pembuktian Realitas Lapangan",
        f"Dari {total_konflik_53} konflik pertambangan di Sulawesi, {konflik_fpic_53} di antaranya secara eksplisit terkait dengan pelanggaran persetujuan awal tanpa paksaan (FPIC).",
        "",
        "## 5.4 Pembuktian Empiris: Uji Statistik Korelasi Penerbitan Izin & Deforestasi",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Panel Join dari `sulawesi_izin_baru_per_tahun.csv` dan `sulawesi_gfw_master_1_dekade_2014_2023_v3.csv`. Metode: *Crosstabulation & Pearson Chi-Square Test*.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Sub-bab ini menggunakan pengujian statistik inferensial untuk membuktikan secara matematis apakah besaran jumlah perizinan baru menjadi prediktor kuat terhadap tingkat kerusakan deforestasi. Data numerik berkelanjutan dikategorikan menjadi Tinggi dan Rendah menggunakan ambang batas median dari distribusi panel.",
        "",
        "#### B. Alur Logika Metodologis Crosstabulation & Pearson Chi-Square Test",
        "```mermaid",
        mermaid_str_5_4,
        "```",
        "",
        "##### Tabel 5.4a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 5.4)",
        markdown_table(["Komponen Uji", "Definisi Variabel (Sub-bab 5.4)"], konf_rows_54),
        "",
        "#### C. Formulasi Matematis: Binning Median, Chi-Square, dan Odds Ratio",
        "```text",
        "Kategori(x) = Tinggi jika x ≥ Median(Panel); Rendah jika x < Median(Panel)",
        "χ² = Σ [ ( O_i - E_i )² / E_i ]",
        "OR = ( a × d ) / ( b × c )",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "##### Tabel 5.4: Ambang Median Panel Uji Crosstab",
        markdown_table(["Variabel X", "Variabel Y", "Median X", "Median Y", "N"], threshold_rows_54),
        "",
        "##### Tabel 5.5: Ringkasan Eksekutif Seluruh Skenario Crosstab Izin vs Deforestasi",
        markdown_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_54),
        "",
        "#### E. Analisis Temuan Empiris: Signifikansi Korelasi Perizinan dan Ekstraksi Ekologis",
        finding_54,
        "",
    ]
    md_path = tool_dir / "Metodologi_Bab5_Pola_Perizinan.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")
    print("[4/4] Selesai membangun Bab 5.")


if __name__ == "__main__":
    generate_all_bab5()
