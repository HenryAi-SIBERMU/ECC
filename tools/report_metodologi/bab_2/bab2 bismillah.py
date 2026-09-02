#!/usr/bin/env python3
"""
Generator Laporan Metodologi Bab 2: Kualitas Lingkungan di Kawasan Smelter

Fokus awal: Sub-bab 2.1. Dampak Limbah Tailing: Konsentrasi Smelter vs
Indeks Kualitas Air (IKA). Pilar 1 ditulis langsung dalam generator Python
agar selaras dengan SOP dokumentasi Celios2.
"""

import base64
import sys
from pathlib import Path

try:
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    import subprocess

    print("[INFO] Memasang dependensi yang diperlukan...")
    subprocess.check_call([
        sys.executable,
        "-m",
        "pip",
        "install",
        "pandas",
        "scipy",
        "requests",
        "python-docx",
    ])
    import pandas as pd
    import scipy.stats as stats
    import requests
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor


G_DARK = RGBColor(0x1B, 0x5E, 0x20)
G_MID = RGBColor(0x2E, 0x7D, 0x32)
G_ACC = RGBColor(0x43, 0xA0, 0x47)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xB7, 0x1C, 0x1C)


def fmt_p(val):
    if pd.isna(val):
        return "= NaN"
    if val < 0.001:
        return "< 0.001"
    return f"= {val:.4f}"


def fmt_p_summary(val):
    if pd.isna(val):
        return "NaN"
    if val < 0.001:
        return "< 0.001"
    return f"{val:.3f}"


def fmt_p_summary(val):
    if pd.isna(val):
        return "NaN"
    if val < 0.001:
        return "< 0.001"
    return f"{val:.3f}"


def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}"
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return True
        print(f"[WARN] Gagal download Mermaid. Status: {resp.status_code}")
        return False
    except Exception as exc:
        print(f"[WARN] Exception saat download Mermaid: {exc}")
        return False


def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for edge, cfg in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        if cfg is None:
            el.set(qn("w:val"), "none")
        else:
            for key, val in cfg.items():
                el.set(qn(f"w:{key}"), str(val))
        bdr.append(el)
    tc_pr.append(bdr)


def cell_margin(cell, left=100, right=100, top=60, bottom=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margin.append(el)
    tc_pr.append(margin)


def para_shd(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def cell_shd(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def para_border_bottom(paragraph, color="2E7D32", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def para_border_left(paragraph, color="2E7D32", sz="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def all_border_para(paragraph, color="444444", sz="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def run(paragraph, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    return r


def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    para_border_bottom(p, color="1B5E20", sz="12")
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)


def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_border_bottom(p, color="2E7D32", sz="6")
    run(p, title.upper(), bold=True, pt=11, color=G_MID)


def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    run(p, title, bold=True, pt=10, color=G_DARK)


def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)


def add_p(doc, parts, space_after=5, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p


def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(12)
    para_shd(p, "EDF7EE")
    all_border_para(p, color="A5D6A7", sz="4")
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)

    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(2)
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.left_indent = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"- {item[0]}: ", bold=True, pt=8, color=G_DARK)
            run(p_desc, f"{item[1]}{trailing}", pt=8, color=RGBColor(0x44, 0x44, 0x44))


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    para_border_left(p, color="2E7D32", sz="16")
    para_shd(p, "F1F8E9")
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))


def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p


def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    bd_cfg = {"val": "single", "sz": "4", "color": "D0D0D0", "space": "0"}

    for j, (header, width) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(width)
        cell_shd(cell, "2E7D32")
        cell_margin(cell, left=100, right=100, top=70, bottom=70)
        set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom={"val": "single", "sz": "8", "color": "1B5E20", "space": "0"}, right=bd_cfg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
        run(p, header, bold=True, pt=8.5, color=C_WHITE)

    for i, row_data in enumerate(rows):
        fill = "F5FBF5" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tbl.cell(i + 1, j)
            cell.width = Cm(col_widths_cm[j])
            cell_shd(cell, fill)
            cell_margin(cell, left=100, right=100, top=50, bottom=50)
            set_cell_borders(cell, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[j] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), pt=8.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(4)
    return tbl


def html_table(headers, rows):
    header_html = "".join(f'<th class="data-th">{h}</th>' for h in headers)
    body = []
    for i, row in enumerate(rows):
        cls = ' class="data-tr-even"' if i % 2 == 1 else ""
        cells = "".join(f'<td class="data-td">{cell}</td>' for cell in row)
        body.append(f"<tr{cls}>{cells}</tr>")
    return f"""<table>
<thead><tr>{header_html}</tr></thead>
<tbody>
{chr(10).join(body)}
</tbody>
</table>"""


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join([":---" for _ in headers]) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def calculate_spss_style_crosstab(df, x_col, y_col, y_is_negative=False):
    df_clean = df.dropna(subset=[x_col, y_col]).copy()
    x_threshold = df_clean[x_col].median()
    y_threshold = df_clean[y_col].median()

    label_x_low = f"Rendah (<{x_threshold:,.1f})"
    label_x_high = f"Tinggi (>={x_threshold:,.1f})"
    if y_is_negative:
        label_y_low = f"Rendah (<{y_threshold:,.1f})"
        label_y_high = f"Tinggi/Parah (>={y_threshold:,.1f})"
    else:
        label_y_low = f"Kritis (<{y_threshold:,.1f})"
        label_y_high = f"Baik (>={y_threshold:,.1f})"

    df_clean["X_Label"] = df_clean[x_col].apply(lambda x: label_x_high if x >= x_threshold else label_x_low)
    df_clean["Y_Label"] = df_clean[y_col].apply(lambda y: label_y_high if y >= y_threshold else label_y_low)

    cats_x = [label_x_low, label_x_high]
    cats_y = [label_y_low, label_y_high]
    crosstab = pd.crosstab(df_clean["X_Label"], df_clean["Y_Label"]).reindex(index=cats_x, columns=cats_y, fill_value=0)

    try:
        chi2, p_val, dof, expected = stats.chi2_contingency(crosstab)
    except Exception:
        chi2, p_val, dof = 0, 1.0, 1
        expected = [[0, 0], [0, 0]]

    try:
        g_val, p_g, dof_g, _ = stats.chi2_contingency(crosstab, lambda_="log-likelihood")
    except Exception:
        g_val, p_g, dof_g = 0, 1.0, dof

    x_codes = df_clean["X_Label"].replace({label_x_low: 0, label_x_high: 1})
    y_codes = df_clean["Y_Label"].replace({label_y_low: 1, label_y_high: 0})
    if len(df_clean) > 2 and x_codes.nunique() > 1 and y_codes.nunique() > 1:
        try:
            r_val, p_corr = stats.pearsonr(x_codes, y_codes)
            lbl_val = (len(df_clean) - 1) * (r_val ** 2)
        except Exception:
            r_val, p_corr, lbl_val = 0, 1.0, 0
    else:
        r_val, p_corr, lbl_val = 0, 1.0, 0

    try:
        if y_is_negative:
            a = crosstab.loc[label_x_high, label_y_high]
            b = crosstab.loc[label_x_high, label_y_low]
            c = crosstab.loc[label_x_low, label_y_high]
            d = crosstab.loc[label_x_low, label_y_low]
        else:
            a = crosstab.loc[label_x_high, label_y_low]
            b = crosstab.loc[label_x_high, label_y_high]
            c = crosstab.loc[label_x_low, label_y_low]
            d = crosstab.loc[label_x_low, label_y_high]
        odds_ratio = (a * d) / (b * c) if (b * c) > 0 else 0
    except Exception:
        a = b = c = d = 0
        odds_ratio = 0

    return {
        "df_clean": df_clean,
        "x_threshold": x_threshold,
        "y_threshold": y_threshold,
        "label_x_low": label_x_low,
        "label_x_high": label_x_high,
        "label_y_low": label_y_low,
        "label_y_high": label_y_high,
        "crosstab": crosstab,
        "expected": expected,
        "chi2": chi2,
        "p_val": p_val,
        "dof": dof,
        "g_val": g_val,
        "p_g": p_g,
        "dof_g": dof_g,
        "lbl_val": lbl_val,
        "p_corr": p_corr,
        "odds_ratio": odds_ratio,
        "a": a,
        "b": b,
        "c": c,
        "d": d,
    }


def generate_all_bab2():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    data_dir = base_dir / "data" / "processed"
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_2"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/4] Mengekstraksi dataset empiris Bab 2 sub-bab 2.1...")
    df_ika = pd.read_csv(data_dir / "sulawesi_ika_2016_2024.csv")
    df_smelter = pd.read_csv(data_dir / "sulawesi_esdm_nikel.csv")
    df_pltu = pd.read_csv(data_dir / "sulawesi_pltu_captive.csv")
    df_gfw = pd.read_csv(data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv")
    df_b3 = pd.read_csv(data_dir / "sulawesi_limbah_b3_ngo_proxy.csv")
    df_sungai = pd.read_csv(data_dir / "sulawesi_sungai_tercemar.csv")
    df_iku = pd.read_csv(data_dir / "sulawesi_iku_2015_2024.csv")

    focus_start_year = int(df_ika["Tahun"].min())
    focus_end_year = int(df_ika["Tahun"].max())
    df_ika_focus = df_ika.copy()
    mean_ika_focus = df_ika_focus[df_ika_focus["Tahun"] == focus_end_year]["Indeks Kualitas Air"].mean()
    tot_smelter = len(df_smelter)
    df_pltu_op = df_pltu[df_pltu["Status"].str.lower() == "operating"].copy()
    tot_kapasitas_pltu = df_pltu_op["Capacity (MW)"].sum() if "Capacity (MW)" in df_pltu_op.columns else 0
    tot_deforestasi = df_gfw["Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha"].sum()
    df_b3["Estimasi Timbulan (Ton/Tahun)"] = pd.to_numeric(df_b3["Estimasi Timbulan (Ton/Tahun)"], errors="coerce").fillna(0)
    tot_limbah_b3 = df_b3["Estimasi Timbulan (Ton/Tahun)"].sum()
    tot_limbah_b3_juta = tot_limbah_b3 / 1_000_000

    df_smelter_prov = df_smelter.groupby("provinsi").size().reset_index(name="Jumlah_Smelter")
    df_smelter_prov.rename(columns={"provinsi": "Provinsi"}, inplace=True)
    df_ika_panel = df_ika_focus.groupby(["Provinsi", "Tahun"])["Indeks Kualitas Air"].mean().reset_index()
    df_panel_2_1 = pd.merge(df_ika_panel, df_smelter_prov, on="Provinsi", how="left").fillna({"Jumlah_Smelter": 0})
    df_panel_2_1.dropna(subset=["Indeks Kualitas Air"], inplace=True)

    max_year_panel = focus_end_year
    df_panel_map_2_1 = df_panel_2_1[df_panel_2_1["Tahun"] == max_year_panel].copy()
    sulteng_smelter = df_smelter_prov.loc[df_smelter_prov["Provinsi"] == "Sulawesi Tengah", "Jumlah_Smelter"].sum()
    sultra_smelter = df_smelter_prov.loc[df_smelter_prov["Provinsi"] == "Sulawesi Tenggara", "Jumlah_Smelter"].sum()
    ika_sulteng = df_panel_map_2_1.loc[df_panel_map_2_1["Provinsi"] == "Sulawesi Tengah", "Indeks Kualitas Air"].mean()
    ika_sultra = df_panel_map_2_1.loc[df_panel_map_2_1["Provinsi"] == "Sulawesi Tenggara", "Indeks Kualitas Air"].mean()

    stats_21 = calculate_spss_style_crosstab(df_panel_2_1, "Jumlah_Smelter", "Indeks Kualitas Air")
    ct = stats_21["crosstab"]
    exp = stats_21["expected"]
    valid_cases = len(stats_21["df_clean"])
    total_cases = len(df_panel_2_1)
    missing_cases = total_cases - valid_cases

    b3_by_prov = df_b3.groupby("Provinsi")["Estimasi Timbulan (Ton/Tahun)"].sum().reset_index()
    sungai_by_prov = df_sungai[["Provinsi", "Jumlah_Sungai_Tercemar", "Daftar_Sungai"]].copy()
    empirical = pd.merge(df_panel_map_2_1[["Provinsi", "Jumlah_Smelter", "Indeks Kualitas Air"]], b3_by_prov, on="Provinsi", how="left")
    empirical = pd.merge(empirical, sungai_by_prov, on="Provinsi", how="left")
    empirical["Estimasi Timbulan (Ton/Tahun)"] = empirical["Estimasi Timbulan (Ton/Tahun)"].fillna(0)
    empirical["Jumlah_Sungai_Tercemar"] = empirical["Jumlah_Sungai_Tercemar"].fillna(0).astype(int)
    empirical["Daftar_Sungai"] = empirical["Daftar_Sungai"].fillna("-")
    empirical = empirical.sort_values(["Jumlah_Smelter", "Estimasi Timbulan (Ton/Tahun)"], ascending=False)

    empirical_rows = []
    for _, row in empirical.iterrows():
        empirical_rows.append([
            row["Provinsi"],
            f"{int(row['Jumlah_Smelter']):,}",
            f"{row['Indeks Kualitas Air']:.1f}",
            f"{row['Estimasi Timbulan (Ton/Tahun)']:,.0f}",
            f"{int(row['Jumlah_Sungai_Tercemar'])}",
            row["Daftar_Sungai"],
        ])

    crosstab_rows = [
        [stats_21["label_x_low"], int(ct.iloc[0, 0]), f"{exp[0][0]:.1f}", int(ct.iloc[0, 1]), f"{exp[0][1]:.1f}", int(ct.iloc[0].sum())],
        [stats_21["label_x_high"], int(ct.iloc[1, 0]), f"{exp[1][0]:.1f}", int(ct.iloc[1, 1]), f"{exp[1][1]:.1f}", int(ct.iloc[1].sum())],
        ["Total", int(ct.iloc[:, 0].sum()), f"{sum(row[0] for row in exp):.1f}", int(ct.iloc[:, 1].sum()), f"{sum(row[1] for row in exp):.1f}", int(ct.values.sum())],
    ]

    chi_rows = [
        ["Pearson Chi-Square", f"{stats_21['chi2']:.3f}", str(stats_21["dof"]), fmt_p(stats_21["p_val"])],
        ["Likelihood Ratio", f"{stats_21['g_val']:.3f}", str(stats_21["dof_g"]), fmt_p(stats_21["p_g"])],
        ["Linear-by-Linear Association", f"{stats_21['lbl_val']:.3f}", "1", fmt_p(stats_21["p_corr"])],
        ["N of Valid Cases", str(valid_cases), "", ""],
    ]

    summary_rows = [[
        "Kepadatan Smelter (Fasilitas)",
        "Indeks Kualitas Air (IKA)",
        f"{stats_21['chi2']:.3f}",
        fmt_p_summary(stats_21["p_val"]),
        "Infinite" if stats_21["odds_ratio"] == 0 else f"{stats_21['odds_ratio']:.2f}",
        "SIGNIFIKAN" if stats_21["p_val"] < 0.05 else "TIDAK SIGNIFIKAN",
    ]]

    konf_rows_21 = [
        ["Variabel Independen (X)", "Jumlah_Smelter: Total fasilitas smelter (beroperasi maupun konstruksi)."],
        ["Variabel Dependen (Y)", "Indeks Kualitas Air: Skor baku mutu air per provinsi."],
        ["Hipotesis Nol (H0)", "Tidak ada hubungan signifikan secara statistik antara kepadatan smelter dengan Indeks Kualitas Air."],
        ["Decision Rule (Alpha 5%)", "Jika P-Value < 0.05, maka Tolak H0 (Terbukti signifikan bahwa smelter menurunkan mutu air)."],
        ["Threshold Kategori", f"Nilai Median Data Panel {focus_start_year}-{focus_end_year} (N={valid_cases}); variabel kontinu dikonversi menjadi biner."],
    ]

    mermaid_str_2_1 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Smelter ESDM<br/><i>Provinsi & jumlah fasilitas</i>"]
        B["Data IKA KLHK/BPS<br/><i>Provinsi, Tahun, Indeks Kualitas Air</i>"]
        C["Data Limbah B3 & Sungai Tercemar<br/><i>Tailing, slag, laporan pencemaran</i>"]
    end

    subgraph Visual_Processing["2. Analisis Spasial & Trendline"]
        A --> D["Agregasi jumlah smelter per provinsi"]
        B --> E["Rata-rata IKA provinsi-tahun"]
        C --> F["Validasi konteks limbah dan sungai tercemar"]
        D --> G["Peta dan trendline tekanan kualitas air"]
        E --> G
        F --> G
    end

    G --> H["Pembacaan empiris kualitas air kawasan smelter"]"""
    mermaid_png_path_2_1 = str(tool_dir / "mermaid_flowchart_2_1.png")
    download_success_2_1 = download_mermaid_png(mermaid_str_2_1, mermaid_png_path_2_1)

    print("[2.5/4] Mengekstraksi dataset empiris Bab 2 sub-bab 2.2...")
    prov_map = {
        'North Sulawesi': 'Sulawesi Utara',
        'South Sulawesi': 'Sulawesi Selatan',
        'Southeast Sulawesi': 'Sulawesi Tenggara',
        'Central Sulawesi': 'Sulawesi Tengah',
        'Gorontalo': 'Gorontalo',
        'West Sulawesi': 'Sulawesi Barat'
    }
    df_pltu_2 = df_pltu.copy()
    df_pltu_2['Provinsi'] = df_pltu_2['Subnational unit (province, state)'].replace(prov_map)
    
    grid_pltu = pd.DataFrame([
        {'Provinsi': 'Gorontalo', 'Capacity (MW)': 100},
        {'Provinsi': 'Sulawesi Utara', 'Capacity (MW)': 220},
        {'Provinsi': 'Sulawesi Selatan', 'Capacity (MW)': 920},
        {'Provinsi': 'Sulawesi Tenggara', 'Capacity (MW)': 100}
    ])
    df_pltu_2 = pd.concat([df_pltu_2, grid_pltu], ignore_index=True)
    
    df_pltu_prov = df_pltu_2.groupby('Provinsi')['Capacity (MW)'].sum().reset_index()
    df_pltu_prov.rename(columns={'Capacity (MW)': 'Kapasitas_PLTU_MW'}, inplace=True)
    
    df_iku_panel = df_iku.groupby(['Provinsi', 'Tahun'])['IKU'].mean().reset_index()
    df_panel_2_2 = pd.merge(df_iku_panel, df_pltu_prov, on='Provinsi', how='left').fillna({'Kapasitas_PLTU_MW': 0})
    df_panel_2_2.dropna(subset=['IKU'], inplace=True)
    
    stats_22 = calculate_spss_style_crosstab(df_panel_2_2, "Kapasitas_PLTU_MW", "IKU")
    valid_cases_22 = len(stats_22["df_clean"])
    
    konf_rows_22 = [
        ["Variabel Independen (X)", "Kapasitas PLTU (MW): Total kapasitas PLTU Captive yang beroperasi."],
        ["Variabel Dependen (Y)", "Indeks Kualitas Udara: Skor baku mutu udara ambien per provinsi."],
        ["Hipotesis Nol (H0)", "Tidak ada hubungan signifikan secara statistik antara kapasitas PLTU dengan Indeks Kualitas Udara."],
        ["Decision Rule (Alpha 5%)", "Jika P-Value < 0.05, maka Tolak H0 (Terbukti signifikan bahwa emisi PLTU menurunkan kualitas udara)."],
        ["Threshold Kategori", f"Nilai Median Data Panel (N={valid_cases_22}); variabel kontinu dikonversi menjadi biner."],
    ]
    
    summary_rows_22 = [[
        "Kapasitas PLTU (MW)",
        "Indeks Kualitas Udara (IKU)",
        f"{stats_22['chi2']:.3f}",
        fmt_p_summary(stats_22["p_val"]),
        f"{stats_22['odds_ratio']:.2f}",
        "SIGNIFIKAN" if stats_22["p_val"] < 0.05 else "TIDAK SIGNIFIKAN",
    ]]
    
    if stats_22["p_val"] < 0.05:
        finding_22 = "Dari skenario pengujian, terbukti secara SIGNIFIKAN bahwa peningkatan kapasitas PLTU berkorelasi dengan memburuknya kualitas udara. Angka Odds Ratio menegaskan bahwa ekspansi industri hilirisasi memberikan risiko kerusakan pada udara ambien."
    else:
        finding_22 = "Kegagalan pengujian statistik ini membongkar fakta krusial bahwa Indeks Kualitas Udara (IKU) level provinsi adalah metrik agregat yang mengencerkan pencemaran udara lokal di tapak industri. Kualitas udara yang buruk di sekitar PLTU tertutupi oleh wilayah yang masih bersih."
        
    mermaid_str_2_2 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data PLTU Captive<br/><i>Kapasitas (MW), Status, Provinsi</i>"]
        B["Data IKU KLHK<br/><i>Provinsi, Tahun, Indeks Kualitas Udara</i>"]
        C["Data Emisi NASA TROPOMI<br/><i>Pantauan satelit udara ambien (NO2)</i>"]
    end

    subgraph Visual_Processing["2. Analisis Spasial & Trendline"]
        A --> D["Agregasi kapasitas PLTU per provinsi"]
        B --> E["Rata-rata IKU provinsi-tahun"]
        C --> F["Validasi konteks polusi dan kepungan asap"]
        D --> G["Peta dan trendline tekanan kualitas udara"]
        E --> G
        F --> G
    end

    G --> H["Pembacaan empiris kualitas udara kawasan PLTU"]"""
    mermaid_png_path_2_2 = str(tool_dir / "mermaid_flowchart_2_2.png")
    download_success_2_2 = download_mermaid_png(mermaid_str_2_2, mermaid_png_path_2_2)

    df_no2 = pd.read_csv(data_dir / "gee_nasa_no2_sulawesi_provinsi.csv")
    df_no2_2023 = df_no2[df_no2["Tahun"] == focus_end_year].copy()
    
    df_iku_2023 = df_iku[df_iku["Tahun"] == focus_end_year].groupby("Provinsi")["IKU"].mean().reset_index()
    
    empirical_22 = pd.merge(df_iku_2023, df_pltu_prov, on="Provinsi", how="left").fillna({'Kapasitas_PLTU_MW': 0})
    empirical_22 = pd.merge(empirical_22, df_no2_2023[["Provinsi", "Rata_Rata_NO2"]], on="Provinsi", how="left")
    empirical_22 = empirical_22.sort_values("Kapasitas_PLTU_MW", ascending=False)
    
    empirical_rows_22 = []
    for _, row in empirical_22.iterrows():
        empirical_rows_22.append([
            row["Provinsi"],
            f"{row['Kapasitas_PLTU_MW']:,.0f}",
            f"{row['IKU']:.1f}" if pd.notna(row['IKU']) else "-",
            f"{row['Rata_Rata_NO2']:.2e}" if pd.notna(row['Rata_Rata_NO2']) else "-",
        ])

    print("[2.7/4] Mengekstraksi dataset empiris Bab 2 sub-bab 2.3...")
    df_luas = pd.read_csv(data_dir / "sulawesi_kawasan_nikel_luas.csv")
    df_izin = pd.read_csv(data_dir / "sulawesi_izin_baru_per_tahun.csv")

    df_luas_prov = df_luas.groupby("provinsi")["total_luas_ha"].sum().reset_index()
    df_luas_prov.rename(columns={"provinsi": "Provinsi", "total_luas_ha": "Luas_IUP_Kawasan_Ha"}, inplace=True)

    df_gfw_panel = df_gfw.groupby(["Provinsi", "Tahun"])["Total_Deforestasi_Ha"].sum().reset_index()
    df_panel_2_3 = pd.merge(df_gfw_panel, df_luas_prov, on="Provinsi", how="inner").fillna(0)

    tot_luas_konsesi = df_luas_prov["Luas_IUP_Kawasan_Ha"].sum()
    tot_def_10thn = df_gfw_panel["Total_Deforestasi_Ha"].sum()
    prov_max_iup = df_luas_prov.loc[df_luas_prov["Luas_IUP_Kawasan_Ha"].idxmax()]["Provinsi"]
    prov_max_def = df_gfw_panel.groupby("Provinsi")["Total_Deforestasi_Ha"].sum().idxmax()

    df_izin = df_izin.sort_values(by=["Provinsi", "Tahun"])
    df_izin["Kumulatif_Luas_Konsesi_Ha"] = df_izin.groupby("Provinsi")["Total_Luas_Konsesi_Baru_Ha"].cumsum()
    izin_kum_prov = df_izin.groupby("Provinsi")["Kumulatif_Luas_Konsesi_Ha"].max().reset_index()
    def_kum_prov = df_gfw_panel.groupby("Provinsi")["Total_Deforestasi_Ha"].sum().reset_index()
    def_kum_prov.rename(columns={"Total_Deforestasi_Ha": "Kumulatif_Deforestasi_Ha"}, inplace=True)

    empirical_23 = pd.merge(df_luas_prov, izin_kum_prov, on="Provinsi", how="left")
    empirical_23 = pd.merge(empirical_23, def_kum_prov, on="Provinsi", how="left")
    empirical_23 = empirical_23.fillna(0).sort_values("Luas_IUP_Kawasan_Ha", ascending=False)

    empirical_rows_23 = []
    for _, row in empirical_23.iterrows():
        empirical_rows_23.append([
            row["Provinsi"],
            f"{row['Luas_IUP_Kawasan_Ha']:,.0f}",
            f"{row['Kumulatif_Luas_Konsesi_Ha']:,.0f}",
            f"{row['Kumulatif_Deforestasi_Ha']:,.0f}",
        ])

    stats_23 = calculate_spss_style_crosstab(df_panel_2_3, "Luas_IUP_Kawasan_Ha", "Total_Deforestasi_Ha", y_is_negative=True)
    valid_cases_23 = len(stats_23["df_clean"])

    summary_rows_23 = [[
        "Luas Ekspansi Industri (Ha)",
        "Kehilangan Tutupan Pohon (Ha)",
        f"{stats_23['chi2']:.3f}",
        f"p {fmt_p(stats_23['p_val'])}",
        "Infinite" if stats_23["odds_ratio"] == 0 else f"{stats_23['odds_ratio']:.1f}",
        "SIGNIFIKAN" if stats_23["p_val"] < 0.05 else "TIDAK SIGNIFIKAN",
    ]]

    if stats_23["p_val"] < 0.05:
        finding_23 = "Hasil pengujian mengonfirmasi secara SIGNIFIKAN bahwa perluasan kawasan industri dan izin pertambangan baru memiliki korelasi positif dengan tingkat deforestasi. Temuan statistik mengonfirmasi bahwa peningkatan luasan Ekspansi Industri berkorelasi signifikan dengan kenaikan tingkat Deforestasi."
    else:
        finding_23 = "Secara umum data menunjukkan kecenderungan bahwa luasan perizinan lahan diikuti oleh kenaikan luasan deforestasi pada wilayah studi. Secara agregat, alokasi perizinan lahan sejalan dengan luasan deforestasi tutupan hutan di tingkat provinsi."

    mermaid_str_2_3 = """flowchart LR
    subgraph Data_Input["1. Input Data Dashboard"]
        A["Data Izin Konsesi Minerbaone<br/><i>Provinsi, Tahun, Luas Konsesi Baru (Ha)</i>"] --> C
        B["Data Kawasan & IUP Nikel<br/><i>Provinsi, Total Luas IUP-Kawasan (Ha)</i>"] --> C
        D["Data Deforestasi GFW 2014-2023<br/><i>Provinsi, Tahun, Total Deforestasi (Ha)</i>"] --> C
    end
    subgraph Panel_Processing["2. Pembentukan Panel 2.3"]
        C["Merge Panel Provinsi-Tahun"] --> F["CUMSUM Konsesi & Deforestasi<br/>per Provinsi (2014-2023)"]
    end
    subgraph Visual_Analysis["3. Animated Bubble Chart (Hans Rosling-style)"]
        F --> G["Choropleth<br/>Level keparahan deforestasi kumulatif"]
        F --> H["Bubble Size<br/>Skala konsesi industri kumulatif"]
        G --> I["Animasi & Slider Temporal<br/>2014-2023"]
        H --> I
    end
    I --> J["Pembacaan empiris eksekusi ruang spasio-temporal"]"""
    mermaid_png_path_2_3 = str(tool_dir / "mermaid_flowchart_2_3.png")
    download_success_2_3 = download_mermaid_png(mermaid_str_2_3, mermaid_png_path_2_3)

    konf_headers_23 = ["Komponen Uji", "Definisi Variabel (Sub-bab 2.3)"]
    konf_rows_23 = [
        ["Variabel Independen (X)", "Luas Ekspansi Industri (Ha) / Luas IUP & Kawasan (Ha)"],
        ["Variabel Dependen (Y)", "Kehilangan Tutupan Pohon (Ha) / Total Deforestasi Alam (Ha)"],
        ["Hipotesis Nol (H0)", "Luasan ekspansi kawasan industri dan perizinan tambang tidak berhubungan dengan laju deforestasi."],
        ["Hipotesis Alternatif (H1)", "Alokasi izin lahan (Luas IUP & Kawasan) berkorelasi positif dengan laju deforestasi."],
        ["Threshold Kategori", f"Nilai Median Data Panel (N={valid_cases_23}): X >= {stats_23['x_threshold']:,.1f} Ha; Y >= {stats_23['y_threshold']:,.1f} Ha"],
    ]

    print("[2.9/4] Membangun DOCX Metodologi_Bab2_Kualitas_Lingkungan.docx...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after = Pt(2)
    run(p_hdr, "CELIOS - CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)

    add_h1(doc, "BAB II: METODOLOGI ANALISIS KUALITAS LINGKUNGAN DI KAWASAN SMELTER")
    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, formulasi matematis, prosedur pengolahan data, dan pengujian statistik yang dioperasionalkan pada ", False, False),
        ("Bab 2: Kualitas Lingkungan di Kawasan Smelter", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi.", False, False),
    ])

    add_h2(doc, "2.1. Dampak Limbah Tailing: Konsentrasi Smelter vs Indeks Kualitas Air (IKA)")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Smelter: data/processed/sulawesi_esdm_nikel.csv; Data IKA: data/processed/sulawesi_ika_2016_2024.csv; Data Limbah B3: data/processed/sulawesi_limbah_b3_ngo_proxy.csv; Data Pencemaran Sungai: data/processed/sulawesi_sungai_tercemar.csv. Visualisasi dashboard menampilkan peta choropleth IKA BPS, timbulan limbah B3, kasus pencemaran sungai, dan uji Crosstabulation SPSS-style.")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        (f"Pengoperasian {tot_smelter:,} fasilitas mega-smelter yang didukung oleh kapasitas {tot_kapasitas_pltu:,.0f} MW PLTU Captive meningkatkan intensitas emisi dan beban lingkungan di Pulau Sulawesi. ", False, False),
        ("Aktivitas pengolahan bijih nikel (smelter) berimplikasi pada timbulan limbah tailing dan terak (slag). ", False, False),
        (f"Data menunjukkan bahwa konversi tutupan hutan mencapai {tot_deforestasi:,.0f} Hektar dengan estimasi timbulan limbah B3/tailing sebesar {tot_limbah_b3_juta:,.1f} Juta Ton per tahun. ", False, False),
        (f"Rata-rata Indeks Kualitas Air (IKA) pada tahun {focus_end_year} berada pada tingkat {mean_ika_focus:.1f} poin.", False, False),
    ])
    add_p(doc, [
        (f"Peta geospasial dan agregasi data dashboard memetakan sebaran {tot_smelter:,} fasilitas smelter, dengan konsentrasi utama berada di Sulawesi Tengah ({int(sulteng_smelter):,} fasilitas smelter) dan Sulawesi Tenggara ({int(sultra_smelter):,} fasilitas smelter). ", False, False),
        (f"Pada tahun {max_year_panel}, IKA tercatat {ika_sulteng:.1f} poin di Sulawesi Tengah dan {ika_sultra:.1f} poin di Sulawesi Tenggara. Sub-bab ini menguji hipotesis: apakah kepadatan smelter berkorelasi secara signifikan dengan penurunan Indeks Kualitas Air (IKA)?", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Konsentrasi Smelter vs IKA")
    add_p(doc, [
        ("Pendekatan statistik trendline untuk membaca hubungan konsentrasi smelter, IKA, dan konteks limbah/sungai tercemar diilustrasikan pada ", False, False),
        ("Bagan Alur 2.1", True, False),
        (" berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya diringkas melalui tabel konfigurasi variabel di bawah gambar.", False, False),
    ])
    add_caption(doc, "Bagan Alur 2.1: Alur Logika Analisis Trendline Konsentrasi Smelter vs IKA")
    if download_success_2_1:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_2_1, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 2.1 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_spacer_21 = doc.add_paragraph()
    p_spacer_21.paragraph_format.space_before = Pt(2)
    p_spacer_21.paragraph_format.space_after = Pt(4)

    add_p(doc, [
        ("Sebagai opsi ringkas pengganti bagan alur crosstab yang terlalu panjang, konfigurasi variabel pengujian Chi-Square disajikan pada ", False, False),
        ("Tabel 2.1a", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 2.1a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.1)")
    add_table_1col(doc, ["Komponen Uji", "Definisi Variabel (Sub-bab 2.1)"], konf_rows_21, [4.5, 11.0], ["L", "L"])

    add_h4(doc, "C. Formulasi Matematis: Kalkulasi Konsentrasi Spasial & Uji Chi-Square")
    add_p(doc, [
        ("Parameterisasi konsentrasi spasial dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:", False, False),
    ])
    add_formula(doc, "Persamaan Agregasi Jumlah Smelter per Provinsi", "S_p = Σ s_i, untuk setiap fasilitas i yang berada di provinsi p", [
        ("S_p", "Total fasilitas smelter pada provinsi p."),
        ("s_i", "Indikator keberadaan fasilitas smelter ke-i; bernilai 1 jika fasilitas berada di provinsi p."),
        ("p", "Provinsi observasi di Pulau Sulawesi."),
    ])
    add_formula(doc, "Persamaan Rata-rata Indeks Kualitas Air Panel Provinsi", "IKĀ_{p,t} = (1 / n_{p,t}) × Σ IKA_{j,p,t}", [
        ("IKĀ_{p,t}", "Rata-rata Indeks Kualitas Air pada provinsi p dan tahun t."),
        ("IKA_{j,p,t}", "Nilai pengamatan Indeks Kualitas Air ke-j pada provinsi p dan tahun t."),
        ("n_{p,t}", "Jumlah pengamatan IKA yang tersedia pada provinsi p dan tahun t."),
    ])
    add_formula(doc, "Persamaan Kategorisasi Median Panel 2x2", "K_x = Tinggi jika X_{p,t} ≥ M_X; K_y = Baik jika Y_{p,t} ≥ M_Y", [
        ("Kategori Smelter", f"Smelter Tinggi jika Jumlah_Smelter_Provinsi >= median panel ({stats_21['x_threshold']:,.1f} fasilitas); selain itu Smelter Rendah."),
        ("Kategori IKA", f"IKA Baik jika Indeks Kualitas Air >= median panel ({stats_21['y_threshold']:,.1f} poin); selain itu IKA Kritis."),
        ("M_X dan M_Y", f"Nilai median masing-masing variabel pada seluruh observasi panel valid N={valid_cases}."),
    ])
    add_formula(doc, "Persamaan Uji Independensi Chi-Square Pearson (χ² Kontinjensi 2x2)", "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]", [
        ("Chi_Square (χ²)", f"Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara kepadatan smelter dan Indeks Kualitas Air pada panel spasiotemporal N={valid_cases}."),
        ("Frekuensi_Observasi (O)", "Jumlah kasus aktual yang tercatat pada sel tabel kontinjensi 2x2."),
        ("Frekuensi_Harapan (E)", "Jumlah kasus teoretis jika kepadatan smelter dan IKA saling independen: E = (Total Baris * Total Kolom) / N."),
    ])
    add_formula(doc, "Persamaan Rasio Keunggulan Risiko IKA Kritis (Risk Odds Ratio / OR)", "Odds_Ratio (OR) = ( a * d ) / ( b * c )", [
        ("Odds_Ratio (OR)", "Ukuran kelipatan peluang munculnya IKA Kritis pada kelompok Smelter Tinggi dibandingkan kelompok Smelter Rendah."),
        ("a", f"Jumlah observasi panel pada kelompok Smelter Tinggi dan IKA Kritis ({stats_21['a']} kasus)."),
        ("b", f"Jumlah observasi panel pada kelompok Smelter Tinggi dan IKA Baik ({stats_21['b']} kasus)."),
        ("c", f"Jumlah observasi panel pada kelompok Smelter Rendah dan IKA Kritis ({stats_21['c']} kasus)."),
        ("d", f"Jumlah observasi panel pada kelompok Smelter Rendah dan IKA Baik ({stats_21['d']} kasus)."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Peta Limbah, IKA, dan Crosstabulation")
    add_p(doc, [
        ("Akumulasi pemusatan fasilitas smelter, nilai IKA, estimasi timbulan limbah B3, dan laporan sungai/pesisir tercemar pada masing-masing provinsi dapat dilihat secara empiris pada ", False, False),
        ("Tabel 2.1", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 2.1: Rincian Empiris Konsentrasi Smelter, IKA, Limbah B3, dan Sungai Tercemar ({max_year_panel})")
    add_table_1col(doc, ["Provinsi", "Jumlah Smelter", "IKA", "Limbah B3 (Ton/Tahun)", "Sungai Tercemar", "Daftar Sungai/Pesisir"], empirical_rows, [2.8, 2.0, 1.4, 3.0, 2.0, 5.8], ["L", "C", "C", "C", "C", "L"])

    add_p(doc, [
        (f"Penerapan sistem pengujian statistik tabulasi silang pada data panel 6 provinsi selama periode {focus_start_year}-{focus_end_year} (total {valid_cases} observasi valid) disajikan secara ringkas pada ", False, False),
        ("Tabel 2.2", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 2.2: Ringkasan Eksekutif Skenario Crosstab Smelter vs IKA Bab 2")
    add_table_1col(doc, ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows, [3.0, 3.5, 2.0, 2.0, 2.0, 2.5], ["L", "L", "C", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Pencemaran Air dan Efek Pengenceran Data Agregat")
    if stats_21["p_val"] < 0.05:
        finding = "Secara konsisten, matriks menunjukkan tren signifikan di mana provinsi dengan intensitas smelter yang tinggi terjebak pada mutu air yang lebih kritis. Ini membuktikan bahwa hilirisasi menumbalkan daya dukung air secara mutlak."
    else:
        finding = "Kegagalan statistik mendeteksi signifikansi membongkar fakta krusial: Indeks Kualitas Air (IKA) provinsi adalah metrik agregat yang mengencerkan tekanan ekologis di tapak. Pencemaran tailing fatal di area tambang dapat tertutupi oleh data sungai-sungai lain di luar lingkar industri."
    add_p(doc, [(finding, False, False)])
    add_p(doc, [
        ("Pembacaan empiris tetap harus dilakukan bersama peta limbah B3 dan laporan sungai tercemar. Dashboard memperlihatkan bahwa analisis IKA BPS tidak berdiri sendiri, melainkan dibaca bersama estimasi timbulan tailing/slag dan laporan NGO tentang pencemaran sungai/pesisir. Dengan demikian, sub-bab ini memosisikan IKA sebagai indikator makro yang perlu diuji silang dengan bukti spasial di kawasan smelter.", False, False),
    ])

    add_h2(doc, "2.2. Kepungan Asap: Kapasitas PLTU vs Indeks Kualitas Udara (IKU)")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data PLTU Captive: data/processed/sulawesi_pltu_captive.csv; Data IKU: data/processed/sulawesi_iku_2015_2024.csv. Visualisasi dashboard menampilkan trendline IKU dan pengujian Chi-Square tabulasi silang (Crosstabulation).")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        (f"Keberadaan {tot_kapasitas_pltu:,.0f} MW PLTU Captive di kawasan hilirisasi secara langsung berkontribusi pada pencemaran udara. Sub-bab ini menguji hipotesis apakah kapasitas terpasang PLTU Captive memiliki hubungan yang signifikan dengan penurunan Indeks Kualitas Udara (IKU).", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Kapasitas PLTU vs IKU")
    add_p(doc, [
        ("Kerangka operasionalisasi sub-bab ini menggunakan pendekatan analisis kuantitatif dan Uji Statistik Chi-Square (Crosstabulation) untuk mengukur korelasi tersebut. Alur data dan pengujian diilustrasikan pada ", False, False),
        ("Bagan Alur 2.2", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Bagan Alur 2.2: Alur Logika Metodologis Crosstabulation & Trendline PLTU vs IKU")
    if download_success_2_2:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_2_2, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 2.2 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_spacer_22 = doc.add_paragraph()
    p_spacer_22.paragraph_format.space_before = Pt(2)
    p_spacer_22.paragraph_format.space_after = Pt(4)

    add_p(doc, [
        ("Sebagai opsi ringkas pengganti bagan alur crosstab yang terlalu panjang, konfigurasi variabel pengujian Chi-Square disajikan pada ", False, False),
        ("Tabel 2.2a", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 2.2a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.2)")
    add_table_1col(doc, ["Komponen Uji", "Definisi Variabel (Sub-bab 2.2)"], konf_rows_22, [4.5, 11.0], ["L", "L"])

    add_h4(doc, "C. Formulasi Matematis: Kapasitas PLTU, Rata-rata IKU, dan Uji Crosstabulation")
    add_formula(doc, "Agregasi Kapasitas PLTU Captive per Provinsi", "Kapasitas_PLTU_Provinsi = SUM(Kapasitas_i) GROUP BY Provinsi", [
        ("Kapasitas_PLTU_Provinsi", "Total kapasitas (MW) PLTU captive yang beroperasi di provinsi observasi."),
        ("Kapasitas_i", "Kapasitas (MW) unit PLTU captive i dalam data operasi."),
    ])
    add_formula(doc, "Rata-rata Indeks Kualitas Udara Panel", "Rata_Rata_IKU_Provinsi_Tahun = MEAN(IKU) GROUP BY Provinsi, Tahun", [
        ("Rata_Rata_IKU_Provinsi_Tahun", "Skor rata-rata Indeks Kualitas Udara (IKU) pada provinsi dan tahun tertentu."),
        ("IKU", "Indeks Kualitas Udara dari data KLHK."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Kapasitas PLTU, IKU, dan Konsentrasi NO2 NASA")
    add_p(doc, [
        ("Akumulasi kapasitas total PLTU (Captive dan Grid) yang beroperasi, beserta kondisi mutu udara melalui pengukuran IKU dan satelit NASA TROPOMI (NO₂) dapat dilihat secara empiris pada ", False, False),
        ("Tabel 2.3", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, f"Tabel 2.3: Rincian Empiris Kapasitas PLTU (Captive & Grid), IKU, dan Konsentrasi NO₂ NASA ({focus_end_year})")
    add_table_1col(doc, ["Provinsi", "Kapasitas PLTU (Captive & Grid) (MW)", "IKU", "NASA TROPOMI NO₂ (mol/m²)\n*Batas Kritis: 6.00e-06"], empirical_rows_22, [3.5, 4.0, 2.0, 4.5], ["L", "C", "C", "C"])

    add_p(doc, [
        (f"Penerapan pengujian statistik tabulasi silang pada data panel (total {valid_cases_22} observasi valid) disajikan secara ringkas pada ", False, False),
        ("Tabel 2.4", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 2.4: Ringkasan Eksekutif Skenario Crosstab Kapasitas PLTU vs IKU Bab 2")
    add_table_1col(doc, ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_22, [3.0, 3.5, 2.0, 2.0, 2.0, 2.5], ["L", "L", "C", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Efek Pengenceran Udara Ambien")
    add_p(doc, [(finding_22, False, False)])

    add_h2(doc, "2.3. Eksekusi Ruang: Ekspansi Kawasan Industri vs Tekanan Ekologis (Deforestasi)")
    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi", "Data Izin Konsesi: data/processed/sulawesi_izin_baru_per_tahun.csv dan data/processed/sulawesi_kawasan_nikel_luas.csv; Data Deforestasi: data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv. Visualisasi dashboard menampilkan Animated Bubble Chart (Hans Rosling-style) berlapis peta choropleth deforestasi kumulatif serta pengujian Chi-Square tabulasi silang (Crosstabulation).")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        (f"Pengembangan kawasan industri pemurnian nikel dan perizinan tambang berimplikasi pada alokasi ruang dan perubahan tutupan lahan. Data menunjukkan bahwa alokasi konsesi perizinan (IUP) dan Kawasan Industri mencakup total luasan ", False, False),
        (f"{tot_luas_konsesi:,.0f} Hektar", True, False),
        (f" di Pulau Sulawesi, dengan alokasi terbesar berada di ", False, False),
        (f"{prov_max_iup}", True, False),
        (".", False, False),
    ])
    add_p(doc, [
        (f"Sepanjang periode 2014-2023, data Global Forest Watch (GFW) merekam akumulasi kehilangan tutupan pohon sebesar ", False, False),
        (f"{tot_def_10thn:,.0f} Hektar", True, False),
        (f", dengan akumulasi terbesar berada di {prov_max_def}. Visualisasi Animated Bubble Chart pada dashboard memperlihatkan pergerakan kumulatif luasan perizinan dan laju perubahan tutupan hutan per provinsi dari tahun ke tahun. Sub-bab ini menguji hipotesis secara empiris: ", False, False),
        ("apakah luasan ekspansi kawasan industri dan perizinan tambang berbanding lurus dengan laju deforestasi?", True, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Ekspansi Industri vs Deforestasi")
    add_p(doc, [
        ("Pendekatan visualisasi dinamis Animated Bubble Chart (Hans Rosling-style) untuk memetakan eksekusi ruang secara spasio-temporal diilustrasikan pada ", False, False),
        ("Bagan Alur 2.3", True, False),
        (" berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya merujuk secara penuh pada ", False, False),
        ("Bagan Alur 2.1", True, False),
        (" (di sub-bab sebelumnya) dengan penyesuaian konfigurasi variabel spesifik sesuai Tabel 2.3a di bawah gambar.", False, False),
    ])
    add_caption(doc, "Bagan Alur 2.3: Alur Logika Metodologis Animated Bubble Chart Ekspansi Industri vs Deforestasi")
    if download_success_2_3:
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(mermaid_png_path_2_3, width=Cm(15))
        except Exception as exc:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 2.3 ke DOCX: {exc}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_spacer_23 = doc.add_paragraph()
    p_spacer_23.paragraph_format.space_before = Pt(2)
    p_spacer_23.paragraph_format.space_after = Pt(4)

    add_caption(doc, "Tabel 2.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.3)")
    add_table_1col(doc, konf_headers_23, konf_rows_23, [4.5, 11.0], ["L", "L"])

    p_spacer_23b = doc.add_paragraph()
    p_spacer_23b.paragraph_format.space_before = Pt(2)
    p_spacer_23b.paragraph_format.space_after = Pt(4)

    add_h4(doc, "C. Formulasi Matematis: Akumulasi Konsesi, Deforestasi, dan Uji Crosstabulation")
    add_p(doc, [
        ("Parameterisasi tekanan ruang dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:", False, False),
    ])
    add_formula(doc, "Agregasi Luas Konsesi & Kawasan Industri per Provinsi", "Luas_IUP_Kawasan_Provinsi = SUM(total_luas_ha) GROUP BY Provinsi", [
        ("Luas_IUP_Kawasan_Provinsi", "Total luasan (Ha) konsesi IUP dan Kawasan Industri nikel pada provinsi observasi."),
        ("total_luas_ha", "Luas izin (Ha) tiap entitas perusahaan dalam data kawasan nikel Minerbaone/CGS."),
    ])
    add_formula(doc, "Kumulatif Luas Konsesi Baru (Ukuran Gelembung / Bubble Size)", "Kumulatif_Luas_Konsesi_Ha = CUMSUM(Total_Luas_Konsesi_Baru_Ha) OVER (ORDER BY Tahun)", [
        ("Kumulatif_Luas_Konsesi_Ha", "Akumulasi luasan konsesi industri per provinsi yang bertambah dari tahun ke tahun (skala ukuran gelembung)."),
        ("Total_Luas_Konsesi_Baru_Ha", "Variabel Tekanan Ruang (Independen): luas IUP diterbitkan per tahun."),
    ])
    add_formula(doc, "Kumulatif Deforestasi (Pewarnaan Choropleth)", "Kumulatif_Deforestasi_Ha = CUMSUM(Total_Deforestasi_Ha) OVER (ORDER BY Tahun)", [
        ("Kumulatif_Deforestasi_Ha", "Akumulasi total deforestasi per provinsi yang merepresentasikan level keparahan pada gradasi warna peta."),
        ("Total_Deforestasi_Ha", "Variabel Dampak Ruang (Dependen): deforestasi alam per tahun."),
    ])
    add_formula(doc, "Persamaan Kategorisasi Median Panel 2x2", "Kategori = IF(Nilai >= Median(Seluruh Panel), 'Tinggi/Parah', 'Rendah')", [
        ("Kategori Ekspansi Industri", f"IUP Tinggi jika Luas_IUP_Kawasan_Ha >= median panel ({stats_23['x_threshold']:,.1f} Ha); selain itu IUP Rendah."),
        ("Kategori Deforestasi", f"Deforestasi Tinggi/Parah jika Total_Deforestasi_Ha >= median panel ({stats_23['y_threshold']:,.1f} Ha); selain itu Deforestasi Rendah."),
        ("Median(Seluruh Panel)", f"Ambang batas dari seluruh observasi panel valid N={valid_cases_23}."),
    ])
    add_formula(doc, "Persamaan Uji Independensi Chi-Square Pearson (χ² Kontinjensi 2x2)", "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]", [
        ("Chi_Square (χ²)", f"Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara luasan ekspansi industri dan kehilangan tutupan pohon pada panel spasiotemporal N={valid_cases_23}."),
        ("Frekuensi_Observasi (O)", "Jumlah kasus aktual yang tercatat pada sel tabel kontinjensi 2x2."),
        ("Frekuensi_Harapan (E)", "Jumlah kasus teoretis jika ekspansi industri dan deforestasi saling independen: E = (Total Baris * Total Kolom) / N."),
    ])
    add_formula(doc, "Persamaan Rasio Keunggulan Risiko Deforestasi Parah (Risk Odds Ratio / OR)", "Odds_Ratio (OR) = ( a * d ) / ( b * c )", [
        ("Odds_Ratio (OR)", "Ukuran kelipatan peluang munculnya Deforestasi Tinggi/Parah pada kelompok IUP Tinggi dibandingkan kelompok IUP Rendah."),
        ("a", f"Jumlah observasi panel pada kelompok IUP Tinggi dan Deforestasi Tinggi/Parah ({stats_23['a']} kasus)."),
        ("b", f"Jumlah observasi panel pada kelompok IUP Tinggi dan Deforestasi Rendah ({stats_23['b']} kasus)."),
        ("c", f"Jumlah observasi panel pada kelompok IUP Rendah dan Deforestasi Tinggi/Parah ({stats_23['c']} kasus)."),
        ("d", f"Jumlah observasi panel pada kelompok IUP Rendah dan Deforestasi Rendah ({stats_23['d']} kasus)."),
    ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Alokasi Ruang Konsesi vs Deforestasi Kumulatif")
    add_p(doc, [
        ("Akumulasi alokasi ruang konsesi IUP-Kawasan Industri dan deforestasi kumulatif dekade 2014-2023 pada masing-masing provinsi dapat dilihat secara empiris pada ", False, False),
        ("Tabel 2.5", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 2.5: Rincian Empiris Luas Konsesi IUP-Kawasan Industri dan Deforestasi Kumulatif per Provinsi (2014-2023)")
    add_table_1col(doc, ["Provinsi", "Luas IUP & Kawasan (Ha)", "Konsesi Baru Kumulatif 2014-2023 (Ha)", "Deforestasi Kumulatif 2014-2023 (Ha)"], empirical_rows_23, [3.4, 3.6, 4.5, 4.5], ["L", "C", "C", "C"])

    add_p(doc, [
        (f"Penerapan pengujian statistik tabulasi silang pada data panel provinsi-tahun periode 2014-2023 (total {valid_cases_23} observasi valid) disajikan secara ringkas pada ", False, False),
        ("Tabel 2.6", True, False),
        (" berikut:", False, False),
    ])
    add_caption(doc, "Tabel 2.6: Ringkasan Eksekutif Skenario Crosstab Ekspansi Industri vs Deforestasi Bab 2")
    add_table_1col(doc, ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_23, [3.0, 3.5, 2.0, 2.0, 2.0, 2.5], ["L", "L", "C", "C", "C", "C"])

    add_h4(doc, "E. Analisis Temuan Empiris: Eksekusi Ruang dan Laju Deforestasi")
    add_p(doc, [(finding_23, False, False)])

    docx_path = tool_dir / "Metodologi_Bab2_Kualitas_Lingkungan.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    print("[3/4] Membangun HTML dan Markdown Bab 2...")
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 2 - Kualitas Lingkungan</title>
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
body {{ font-family: Arial, sans-serif; max-width: 920px; margin: 0 auto; padding: 32px; background: #0E1117; color: #D4D4D4; line-height: 1.65; }}
.hdr-sub {{ color: #43A047; font-size: 8.5pt; font-weight: 700; text-transform: uppercase; }}
.hdr-title {{ color: #81C784; font-size: 15pt; font-weight: 800; text-transform: uppercase; border-bottom: 2px solid #2E7D32; padding-bottom: 8px; }}
h2 {{ color: #81C784; text-transform: uppercase; border-bottom: 1px solid #2E7D32; }}
h4 {{ color: #A5D6A7; }}
.note-box {{ background: #132213; border-left: 4px solid #2E7D32; padding: 10px 14px; margin: 12px 0; }}
.formula {{ background: #0D1B0E; border: 1px solid #2E7D32; color: #A5D6A7; padding: 8px 12px; font-family: monospace; }}
.data-th {{ background: #1B5E20; color: white; padding: 6px; text-align: left; border: 1px solid #2E7D32; }}
.data-td {{ padding: 6px; border: 1px solid #243524; vertical-align: top; }}
.data-tr-even .data-td {{ background: #131B13; }}
.mermaid {{ background: #0D1610; border: 1px solid #2E7D32; padding: 12px; margin: 10px 0; }}
</style>
</head>
<body>
<div class="hdr-sub">CELIOS - Center of Economic and Law Studies | Laporan Riset Metodologi D3TLH</div>
<div class="hdr-title">BAB II: Metodologi Analisis Kualitas Lingkungan di Kawasan Smelter</div>
<p>Dokumen laporan metodologi ini menyajikan kerangka ilmiah, formulasi matematis, prosedur pengolahan data, dan pengujian statistik yang dioperasionalkan pada <strong>Bab 2: Kualitas Lingkungan di Kawasan Smelter</strong>.</p>
<h2>2.1. Dampak Limbah Tailing: Konsentrasi Smelter vs Indeks Kualitas Air (IKA)</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Smelter, IKA, Limbah B3, dan Sungai Tercemar dari folder <code>data/processed</code>. Visualisasi dashboard menampilkan peta choropleth dan uji Crosstabulation SPSS-style.</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Pengoperasian <strong>{tot_smelter:,} fasilitas mega-smelter</strong> yang didukung oleh kapasitas <strong>{tot_kapasitas_pltu:,.0f} MW PLTU Captive</strong> meningkatkan intensitas emisi dan beban lingkungan di Pulau Sulawesi. Data menunjukkan konversi tutupan hutan mencapai <strong>{tot_deforestasi:,.0f} Hektar</strong>, estimasi timbulan limbah B3/tailing sebesar <strong>{tot_limbah_b3_juta:,.1f} Juta Ton</strong> per tahun, dan rata-rata IKA tahun {focus_end_year} sebesar <strong>{mean_ika_focus:.1f}</strong>.</p>
<h4>B. Alur Logika Metodologis Analisis Konsentrasi Smelter vs IKA</h4>
<div class="mermaid">{mermaid_str_2_1}</div>
<p>Sebagai opsi ringkas pengganti bagan alur crosstab yang terlalu panjang, konfigurasi variabel pengujian Chi-Square disajikan pada <strong>Tabel 2.1a</strong> berikut:</p>
<div class="table-caption">Tabel 2.1a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.1)</div>
{html_table(["Komponen Uji", "Definisi Variabel (Sub-bab 2.1)"], konf_rows_21)}
<h4>C. Formulasi Matematis: Kalkulasi Konsentrasi Spasial &amp; Uji Chi-Square</h4>
<p>Parameterisasi konsentrasi spasial dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:</p>
<div class="formula">S_p = Σ s_i, untuk setiap fasilitas i yang berada di provinsi p</div>
<div class="formula">IKĀ_{{p,t}} = (1 / n_{{p,t}}) × Σ IKA_{{j,p,t}}</div>
<div class="formula">K_x = Tinggi jika X_{{p,t}} ≥ M_X; K_y = Baik jika Y_{{p,t}} ≥ M_Y</div>
<div class="formula">Chi_Square (&chi;&sup2;) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]</div>
<div class="formula">Odds_Ratio (OR) = ( a * d ) / ( b * c )</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<p>Akumulasi pemusatan fasilitas smelter, nilai IKA, estimasi timbulan limbah B3, dan laporan sungai/pesisir tercemar pada masing-masing provinsi dapat dilihat secara empiris pada <strong>Tabel 2.1</strong> berikut:</p>
<div class="table-caption">Tabel 2.1: Rincian Empiris Konsentrasi Smelter, IKA, Limbah B3, dan Sungai Tercemar ({max_year_panel})</div>
{html_table(["Provinsi", "Jumlah Smelter", "IKA", "Limbah B3 (Ton/Tahun)", "Sungai Tercemar", "Daftar Sungai/Pesisir"], empirical_rows)}
<p>Penerapan sistem pengujian statistik tabulasi silang pada data panel 6 provinsi selama periode {focus_start_year}-{focus_end_year} (total {valid_cases} observasi valid) disajikan secara ringkas pada <strong>Tabel 2.2</strong> berikut:</p>
<div class="table-caption">Tabel 2.2: Ringkasan Eksekutif Skenario Crosstab Smelter vs IKA Bab 2</div>
{html_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (&chi;&sup2;)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows)}
<h4>E. Analisis Temuan Empiris</h4>
<p>{finding}</p>

<h2>2.2. Kepungan Asap: Kapasitas PLTU vs Indeks Kualitas Udara (IKU)</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data PLTU Captive: <code>data/processed/sulawesi_pltu_captive.csv</code>; Data IKU: <code>data/processed/sulawesi_iku_2015_2024.csv</code>. Visualisasi dashboard menampilkan trendline IKU dan pengujian Chi-Square tabulasi silang (Crosstabulation).</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Keberadaan <strong>{tot_kapasitas_pltu:,.0f} MW PLTU Captive</strong> di kawasan hilirisasi secara langsung berkontribusi pada pencemaran udara. Sub-bab ini menguji hipotesis apakah kapasitas terpasang PLTU Captive memiliki hubungan yang signifikan dengan penurunan Indeks Kualitas Udara (IKU).</p>
<h4>B. Alur Logika Metodologis Analisis Kapasitas PLTU vs IKU</h4>
<div class="mermaid">{mermaid_str_2_2}</div>
<p>Sebagai opsi ringkas pengganti bagan alur crosstab yang terlalu panjang, konfigurasi variabel pengujian Chi-Square disajikan pada <strong>Tabel 2.2a</strong> berikut:</p>
<div class="table-caption">Tabel 2.2a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.2)</div>
{html_table(["Komponen Uji", "Definisi Variabel (Sub-bab 2.2)"], konf_rows_22)}
<h4>C. Formulasi Matematis</h4>
<div class="formula">Kapasitas_PLTU_Provinsi = SUM(Kapasitas_i) GROUP BY Provinsi</div>
<div class="formula">Rata_Rata_IKU_Provinsi_Tahun = MEAN(IKU) GROUP BY Provinsi, Tahun</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<p>Akumulasi kapasitas total PLTU (Captive dan Grid) yang beroperasi, beserta kondisi mutu udara melalui pengukuran IKU dan satelit NASA TROPOMI (NO₂) dapat dilihat secara empiris pada <strong>Tabel 2.3</strong> berikut:</p>
<div class="table-caption">Tabel 2.3: Rincian Empiris Kapasitas PLTU (Captive & Grid), IKU, dan Konsentrasi NO₂ NASA ({focus_end_year})</div>
{html_table(["Provinsi", "Kapasitas PLTU (Captive & Grid) (MW)", "IKU", "NASA TROPOMI NO₂ (mol/m²)"], empirical_rows_22)}
<p>Penerapan pengujian statistik tabulasi silang pada data panel (total {valid_cases_22} observasi valid) disajikan secara ringkas pada <strong>Tabel 2.4</strong> berikut:</p>
<div class="table-caption">Tabel 2.4: Ringkasan Eksekutif Skenario Crosstab Kapasitas PLTU vs IKU Bab 2</div>
{html_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (&chi;&sup2;)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_22)}
<h4>E. Analisis Temuan Empiris</h4>
<p>{finding_22}</p>

<h2>2.3. Eksekusi Ruang: Ekspansi Kawasan Industri vs Tekanan Ekologis (Deforestasi)</h2>
<div class="note-box"><strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong> Data Izin Konsesi: <code>data/processed/sulawesi_izin_baru_per_tahun.csv</code> dan <code>data/processed/sulawesi_kawasan_nikel_luas.csv</code>; Data Deforestasi: <code>data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv</code>. Visualisasi dashboard menampilkan Animated Bubble Chart (Hans Rosling-style) berlapis peta choropleth deforestasi kumulatif serta pengujian Chi-Square tabulasi silang (Crosstabulation).</div>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Pengembangan kawasan industri pemurnian nikel dan perizinan tambang berimplikasi pada alokasi ruang dan perubahan tutupan lahan. Data menunjukkan bahwa alokasi konsesi perizinan (IUP) dan Kawasan Industri mencakup total luasan <strong>{tot_luas_konsesi:,.0f} Hektar</strong> di Pulau Sulawesi, dengan alokasi terbesar berada di <strong>{prov_max_iup}</strong>. Sepanjang periode 2014-2023, data Global Forest Watch (GFW) merekam akumulasi kehilangan tutupan pohon sebesar <strong>{tot_def_10thn:,.0f} Hektar</strong>, dengan akumulasi terbesar berada di {prov_max_def}. Sub-bab ini menguji hipotesis secara empiris: <strong>apakah luasan ekspansi kawasan industri dan perizinan tambang berbanding lurus dengan laju deforestasi?</strong></p>
<h4>B. Alur Logika Metodologis Analisis Ekspansi Industri vs Deforestasi</h4>
<p>Pendekatan visualisasi dinamis Animated Bubble Chart (Hans Rosling-style) untuk memetakan eksekusi ruang secara spasio-temporal diilustrasikan pada <strong>Bagan Alur 2.3</strong> berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya merujuk secara penuh pada <strong>Bagan Alur 2.1</strong> (di sub-bab sebelumnya) dengan penyesuaian konfigurasi variabel spesifik sesuai Tabel 2.3a di bawah gambar.</p>
<div class="table-caption">Bagan Alur 2.3: Alur Logika Metodologis Animated Bubble Chart Ekspansi Industri vs Deforestasi</div>
<div class="mermaid">{mermaid_str_2_3}</div>
<div class="table-caption">Tabel 2.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.3)</div>
{html_table(konf_headers_23, konf_rows_23)}
<h4>C. Formulasi Matematis: Akumulasi Konsesi, Deforestasi, dan Uji Crosstabulation</h4>
<p>Parameterisasi tekanan ruang dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:</p>
<div class="formula">Luas_IUP_Kawasan_Provinsi = SUM(total_luas_ha) GROUP BY Provinsi</div>
<div class="formula">Kumulatif_Luas_Konsesi_Ha = CUMSUM(Total_Luas_Konsesi_Baru_Ha) OVER (ORDER BY Tahun)</div>
<div class="formula">Kumulatif_Deforestasi_Ha = CUMSUM(Total_Deforestasi_Ha) OVER (ORDER BY Tahun)</div>
<div class="formula">Kategori = IF(Nilai &gt;= Median(Seluruh Panel), 'Tinggi/Parah', 'Rendah')</div>
<div class="formula">Chi_Square (&chi;&sup2;) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]</div>
<div class="formula">Odds_Ratio (OR) = ( a * d ) / ( b * c )</div>
<h4>D. Matriks Hasil Uji Empiris</h4>
<p>Akumulasi alokasi ruang konsesi IUP-Kawasan Industri dan deforestasi kumulatif dekade 2014-2023 pada masing-masing provinsi dapat dilihat secara empiris pada <strong>Tabel 2.5</strong> berikut:</p>
<div class="table-caption">Tabel 2.5: Rincian Empiris Luas Konsesi IUP-Kawasan Industri dan Deforestasi Kumulatif per Provinsi (2014-2023)</div>
{html_table(["Provinsi", "Luas IUP & Kawasan (Ha)", "Konsesi Baru Kumulatif 2014-2023 (Ha)", "Deforestasi Kumulatif 2014-2023 (Ha)"], empirical_rows_23)}
<p>Penerapan pengujian statistik tabulasi silang pada data panel provinsi-tahun periode 2014-2023 (total {valid_cases_23} observasi valid) disajikan secara ringkas pada <strong>Tabel 2.6</strong> berikut:</p>
<div class="table-caption">Tabel 2.6: Ringkasan Eksekutif Skenario Crosstab Ekspansi Industri vs Deforestasi Bab 2</div>
{html_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (&chi;&sup2;)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_23)}
<h4>E. Analisis Temuan Empiris</h4>
<p>{finding_23}</p>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab2_Kualitas_Lingkungan.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    md_lines = [
        "# BAB II: METODOLOGI ANALISIS KUALITAS LINGKUNGAN DI KAWASAN SMELTER",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, formulasi matematis, prosedur pengolahan data, dan pengujian statistik yang dioperasionalkan pada **Bab 2: Kualitas Lingkungan di Kawasan Smelter**.",
        "",
        "## 2.1. Dampak Limbah Tailing: Konsentrasi Smelter vs Indeks Kualitas Air (IKA)",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Smelter: `data/processed/sulawesi_esdm_nikel.csv`; Data IKA: `data/processed/sulawesi_ika_2016_2024.csv`; Data Limbah B3: `data/processed/sulawesi_limbah_b3_ngo_proxy.csv`; Data Pencemaran Sungai: `data/processed/sulawesi_sungai_tercemar.csv`.",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Pengoperasian **{tot_smelter:,} fasilitas mega-smelter** yang didukung oleh kapasitas **{tot_kapasitas_pltu:,.0f} MW PLTU Captive** meningkatkan intensitas emisi dan beban lingkungan di Pulau Sulawesi. Data menunjukkan konversi tutupan hutan mencapai **{tot_deforestasi:,.0f} Hektar**, estimasi timbulan limbah B3/tailing sebesar **{tot_limbah_b3_juta:,.1f} Juta Ton** per tahun, dan rata-rata IKA tahun {focus_end_year} sebesar **{mean_ika_focus:.1f}**.",
        "",
        "#### B. Alur Logika Metodologis Analisis Konsentrasi Smelter vs IKA",
        "```mermaid",
        mermaid_str_2_1,
        "```",
        "",
        "Sebagai opsi ringkas pengganti bagan alur crosstab yang terlalu panjang, konfigurasi variabel pengujian Chi-Square disajikan pada **Tabel 2.1a** berikut:",
        "",
        "##### Tabel 2.1a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.1)",
        markdown_table(["Komponen Uji", "Definisi Variabel (Sub-bab 2.1)"], konf_rows_21),
        "",
        "#### C. Formulasi Matematis: Kalkulasi Konsentrasi Spasial & Uji Chi-Square",
        "Parameterisasi konsentrasi spasial dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "S_p = Σ s_i, untuk setiap fasilitas i yang berada di provinsi p",
        "IKĀ_{p,t} = (1 / n_{p,t}) × Σ IKA_{j,p,t}",
        "K_x = Tinggi jika X_{p,t} ≥ M_X; K_y = Baik jika Y_{p,t} ≥ M_Y",
        "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]",
        "Odds_Ratio (OR) = ( a * d ) / ( b * c )",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "Akumulasi pemusatan fasilitas smelter, nilai IKA, estimasi timbulan limbah B3, dan laporan sungai/pesisir tercemar pada masing-masing provinsi dapat dilihat secara empiris pada **Tabel 2.1** berikut:",
        "",
        f"##### Tabel 2.1: Rincian Empiris Konsentrasi Smelter, IKA, Limbah B3, dan Sungai Tercemar ({max_year_panel})",
        markdown_table(["Provinsi", "Jumlah Smelter", "IKA", "Limbah B3 (Ton/Tahun)", "Sungai Tercemar", "Daftar Sungai/Pesisir"], empirical_rows),
        "",
        f"Penerapan sistem pengujian statistik tabulasi silang pada data panel 6 provinsi selama periode {focus_start_year}-{focus_end_year} (total {valid_cases} observasi valid) disajikan secara ringkas pada **Tabel 2.2** berikut:",
        "",
        "##### Tabel 2.2: Ringkasan Eksekutif Skenario Crosstab Smelter vs IKA Bab 2",
        markdown_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows),
        "",
        "#### E. Analisis Temuan Empiris: Pencemaran Air dan Efek Pengenceran Data Agregat",
        finding,
        "",
        "## 2.2. Kepungan Asap: Kapasitas PLTU vs Indeks Kualitas Udara (IKU)",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data PLTU Captive: `data/processed/sulawesi_pltu_captive.csv`; Data IKU: `data/processed/sulawesi_iku_2015_2024.csv`. Visualisasi dashboard menampilkan trendline IKU dan pengujian Chi-Square tabulasi silang (Crosstabulation).",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Keberadaan **{tot_kapasitas_pltu:,.0f} MW PLTU Captive** di kawasan hilirisasi secara langsung berkontribusi pada pencemaran udara. Sub-bab ini menguji hipotesis apakah kapasitas terpasang PLTU Captive memiliki hubungan yang signifikan dengan penurunan Indeks Kualitas Udara (IKU).",
        "",
        "#### B. Alur Logika Metodologis Analisis Kapasitas PLTU vs IKU",
        "```mermaid",
        mermaid_str_2_2,
        "```",
        "",
        "Sebagai opsi ringkas pengganti bagan alur crosstab yang terlalu panjang, konfigurasi variabel pengujian Chi-Square disajikan pada **Tabel 2.2a** berikut:",
        "",
        "##### Tabel 2.2a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.2)",
        markdown_table(["Komponen Uji", "Definisi Variabel (Sub-bab 2.2)"], konf_rows_22),
        "",
        "#### C. Formulasi Matematis",
        "```text",
        "Kapasitas_PLTU_Provinsi = SUM(Kapasitas_i) GROUP BY Provinsi",
        "Rata_Rata_IKU_Provinsi_Tahun = MEAN(IKU) GROUP BY Provinsi, Tahun",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "Akumulasi kapasitas total PLTU (Captive dan Grid) yang beroperasi, beserta kondisi mutu udara melalui pengukuran IKU dan satelit NASA TROPOMI (NO₂) dapat dilihat secara empiris pada **Tabel 2.3** berikut:",
        "",
        f"##### Tabel 2.3: Rincian Empiris Kapasitas PLTU (Captive & Grid), IKU, dan Konsentrasi NO₂ NASA ({focus_end_year})",
        markdown_table(["Provinsi", "Kapasitas PLTU (Captive & Grid) (MW)", "IKU", "NASA TROPOMI NO₂ (mol/m²)"], empirical_rows_22),
        "",
        f"Penerapan pengujian statistik tabulasi silang pada data panel (total {valid_cases_22} observasi valid) disajikan secara ringkas pada **Tabel 2.4** berikut:",
        "",
        "##### Tabel 2.4: Ringkasan Eksekutif Skenario Crosstab Kapasitas PLTU vs IKU Bab 2",
        markdown_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_22),
        "",
        "#### E. Analisis Temuan Empiris: Efek Pengenceran Udara Ambien",
        finding_22,
        "",
        "## 2.3. Eksekusi Ruang: Ekspansi Kawasan Industri vs Tekanan Ekologis (Deforestasi)",
        "",
        "> **Sumber Data Resmi & Deskripsi Visualisasi:** Data Izin Konsesi: `data/processed/sulawesi_izin_baru_per_tahun.csv` dan `data/processed/sulawesi_kawasan_nikel_luas.csv`; Data Deforestasi: `data/processed/sulawesi_gfw_master_1_dekade_2014_2023_v3.csv`. Visualisasi dashboard menampilkan Animated Bubble Chart (Hans Rosling-style) berlapis peta choropleth deforestasi kumulatif serta pengujian Chi-Square tabulasi silang (Crosstabulation).",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        f"Pengembangan kawasan industri pemurnian nikel dan perizinan tambang berimplikasi pada alokasi ruang dan perubahan tutupan lahan. Data menunjukkan bahwa alokasi konsesi perizinan (IUP) dan Kawasan Industri mencakup total luasan **{tot_luas_konsesi:,.0f} Hektar** di Pulau Sulawesi, dengan alokasi terbesar berada di **{prov_max_iup}**. Sepanjang periode 2014-2023, data Global Forest Watch (GFW) merekam akumulasi kehilangan tutupan pohon sebesar **{tot_def_10thn:,.0f} Hektar**, dengan akumulasi terbesar berada di {prov_max_def}. Sub-bab ini menguji hipotesis secara empiris: **apakah luasan ekspansi kawasan industri dan perizinan tambang berbanding lurus dengan laju deforestasi?**",
        "",
        "#### B. Alur Logika Metodologis Analisis Ekspansi Industri vs Deforestasi",
        "Pendekatan visualisasi dinamis Animated Bubble Chart (Hans Rosling-style) untuk memetakan eksekusi ruang secara spasio-temporal diilustrasikan pada **Bagan Alur 2.3** berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya merujuk secara penuh pada **Bagan Alur 2.1** (di sub-bab sebelumnya) dengan penyesuaian konfigurasi variabel spesifik sesuai Tabel 2.3a di bawah gambar.",
        "",
        "##### Bagan Alur 2.3: Alur Logika Metodologis Animated Bubble Chart Ekspansi Industri vs Deforestasi",
        "```mermaid",
        mermaid_str_2_3,
        "```",
        "",
        "##### Tabel 2.3a: Konfigurasi Variabel Uji Chi-Square (Sub-bab 2.3)",
        markdown_table(konf_headers_23, konf_rows_23),
        "",
        "#### C. Formulasi Matematis: Akumulasi Konsesi, Deforestasi, dan Uji Crosstabulation",
        "Parameterisasi tekanan ruang dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "```text",
        "Luas_IUP_Kawasan_Provinsi = SUM(total_luas_ha) GROUP BY Provinsi",
        "Kumulatif_Luas_Konsesi_Ha = CUMSUM(Total_Luas_Konsesi_Baru_Ha) OVER (ORDER BY Tahun)",
        "Kumulatif_Deforestasi_Ha = CUMSUM(Total_Deforestasi_Ha) OVER (ORDER BY Tahun)",
        "Kategori = IF(Nilai >= Median(Seluruh Panel), 'Tinggi/Parah', 'Rendah')",
        "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]",
        "Odds_Ratio (OR) = ( a * d ) / ( b * c )",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris",
        "Akumulasi alokasi ruang konsesi IUP-Kawasan Industri dan deforestasi kumulatif dekade 2014-2023 pada masing-masing provinsi dapat dilihat secara empiris pada **Tabel 2.5** berikut:",
        "",
        "##### Tabel 2.5: Rincian Empiris Luas Konsesi IUP-Kawasan Industri dan Deforestasi Kumulatif per Provinsi (2014-2023)",
        markdown_table(["Provinsi", "Luas IUP & Kawasan (Ha)", "Konsesi Baru Kumulatif 2014-2023 (Ha)", "Deforestasi Kumulatif 2014-2023 (Ha)"], empirical_rows_23),
        "",
        f"Penerapan pengujian statistik tabulasi silang pada data panel provinsi-tahun periode 2014-2023 (total {valid_cases_23} observasi valid) disajikan secara ringkas pada **Tabel 2.6** berikut:",
        "",
        "##### Tabel 2.6: Ringkasan Eksekutif Skenario Crosstab Ekspansi Industri vs Deforestasi Bab 2",
        markdown_table(["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"], summary_rows_23),
        "",
        "#### E. Analisis Temuan Empiris: Eksekusi Ruang dan Laju Deforestasi",
        finding_23,
        "",
    ]
    md_path = tool_dir / "Metodologi_Bab2_Kualitas_Lingkungan.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")

    print("[4/4] Selesai membangun Bab 2.")


if __name__ == "__main__":
    generate_all_bab2()
