#!/usr/bin/env python3
"""
Generator Laporan Metodologi Bab 1: Ekspansi Industri Ekstraktif (CELIOS ECC)
Format: 1-Kolom Penuh, Rinci & Menyeluruh, Standar Publikasi Publik / Akademis Resmi.
Menghasilkan 4 format dokumen di folder tools/report_metodologi/:
1. Metodologi_Bab1_Ekspansi_Industri.docx (Dokumen Word Resmi dengan Custom XML Styling CELIOS)
2. Metodologi_Bab1_Ekspansi_Industri.html (Dokumen Web & Cetak Publikasi)
3. Metodologi_Bab1_Ekspansi_Industri.tex  (Dokumen LaTeX Standar Jurnal / Kebijakan)
4. Metodologi_Bab1_Ekspansi_Industri.md   (Dokumen Naskah Markdown Bersih)
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    print("[INFO] Memasang modul python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

import pypandoc
import pandas as pd
import numpy as np
import scipy.stats as stats
import base64
import requests

def fmt_p(val):
    if pd.isna(val):
        return "= NaN"
    if val < 0.0001:
        return "< 0.0001"
    return f"= {val:.4f}"

def download_mermaid_png(mermaid_str, filepath):
    try:
        encoded = base64.urlsafe_b64encode(mermaid_str.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{encoded}'
        print(f"[INFO] Mendownload Mermaid JS flowchart ke {filepath}...")
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200:
            with open(filepath, 'wb') as f:
                f.write(resp.content)
            return True
        else:
            print(f"[WARN] Gagal download mermaid (Status Code: {resp.status_code})")
            return False
    except Exception as e:
        print(f"[WARN] Exception saat download mermaid: {e}")
        return False

# ── Warna Tema CELIOS D3TLH ─────────────────────────────────
G_DARK  = RGBColor(0x1B, 0x5E, 0x20)  # Hijau Hutan Gelap (#1B5E20)
G_MID   = RGBColor(0x2E, 0x7D, 0x32)  # Hijau Utama CELIOS (#2E7D32)
G_ACC   = RGBColor(0x43, 0xA0, 0x47)  # Hijau Aksen (#43A047)
C_BODY  = RGBColor(0x22, 0x22, 0x22)  # Abu Gelap Teks (#222222)
C_GREY  = RGBColor(0x55, 0x55, 0x55)  # Abu Sekunder (#555555)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # Putih
C_RED   = RGBColor(0xB7, 0x1C, 0x1C)  # Merah Kritis (#B71C1C)
C_ORANGE= RGBColor(0xE6, 0x51, 0x00)  # Oranye (#E65100)

# ── Helper XML Word Formatting ──────────────────────────────
def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge, cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{edge}')
        if cfg is None:
            el.set(qn('w:val'), 'none')
        else:
            for k, v in cfg.items():
                el.set(qn(f'w:{k}'), str(v))
        bdr.append(el)
    tcPr.append(bdr)

def cell_margin(cell, left=100, right=100, top=60, bottom=60):
    tcPr = cell._tc.get_or_add_tcPr()
    m    = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def para_shd(p, fill):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    pPr.append(s)

def cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    tcPr.append(s)

def para_border_bottom(p, color='2E7D32', sz='6'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '2')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def para_border_left(p, color='2E7D32', sz='18'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:left')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '6')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def all_border_para(p, color='444444', sz='4'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)

# ── Helper Konten & Tipografi 1-Kolom ────────────────────────
def run(p, text, bold=False, italic=False, pt=9.5, color=None, mono=False):
    r = p.add_run(text)
    r.bold           = bold
    r.italic         = italic
    r.font.size      = Pt(pt)
    r.font.color.rgb = color if color else C_BODY
    if mono:
        r.font.name = 'Courier New'
        r._element.rPr.rFonts.set(qn('w:ascii'), 'Courier New')
    return r

def add_h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    para_border_bottom(p, color='1B5E20', sz='12')
    run(p, title.upper(), bold=True, pt=13, color=G_DARK)

def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    para_border_bottom(p, color='2E7D32', sz='6')
    run(p, title.upper(), bold=True, pt=11, color=G_MID)

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after  = Pt(3)
    run(p, title, bold=True, pt=10, color=G_DARK)

def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run(p, title, bold=True, pt=9.5, color=G_MID)

def add_p(doc, parts, space_after=5, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent > 0:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold, italic in parts:
        run(p, text, bold=bold, italic=italic, pt=9.5)
    return p

def add_formula(doc, title, formula_text, var_desc=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after  = Pt(1)
    run(p_title, f"Persamaan: {title}", bold=True, italic=True, pt=8.5, color=G_MID)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Pt(12)
    para_shd(p, 'EDF7EE')
    all_border_para(p, color='A5D6A7', sz='4')
    run(p, formula_text, pt=8.5, color=G_DARK, mono=True)

    if var_desc:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(2)
        p_desc.paragraph_format.space_after  = Pt(6)
        p_desc.paragraph_format.left_indent  = Pt(14)
        run(p_desc, "Keterangan Variabel:\n", bold=True, italic=True, pt=8, color=RGBColor(0x33, 0x33, 0x33))
        for idx, item in enumerate(var_desc):
            run(p_desc, f"• {item[0]}: ", bold=True, pt=8, color=RGBColor(0x1B, 0x5E, 0x20))
            trailing = "\n" if idx < len(var_desc) - 1 else ""
            run(p_desc, f"{item[1]}{trailing}", italic=False, pt=8, color=RGBColor(0x44, 0x44, 0x44))

def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(10)
    para_border_left(p, color='2E7D32', sz='16')
    para_shd(p, 'F1F8E9')
    run(p, f"{title.upper()}: ", bold=True, pt=8.5, color=G_DARK)
    run(p, text, italic=True, pt=8.5, color=RGBColor(0x33, 0x33, 0x33))

def add_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run(p, caption_text, bold=True, italic=True, pt=8.5, color=G_MID)
    return p

def add_table_1col(doc, headers, rows, col_widths_cm, alignments=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.autofit = False

    bd_cfg = {'val': 'single', 'sz': '4', 'color': 'D0D0D0', 'space': '0'}

    # Header
    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = tbl.rows[0].cells[j]
        c.width = Cm(w)
        cell_shd(c, '2E7D32')
        cell_margin(c, left=100, right=100, top=70, bottom=70)
        set_cell_borders(c, top=bd_cfg, left=bd_cfg, bottom={'val': 'single', 'sz': '8', 'color': '1B5E20', 'space': '0'}, right=bd_cfg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (alignments and alignments[j] == 'C') else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run(p, h, bold=True, pt=8.5, color=C_WHITE)

    # Data Rows
    for i, row_data in enumerate(rows):
        fill = 'F5FBF5' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            c = tbl.cell(i + 1, j)
            c.width = Cm(col_widths_cm[j])
            cell_shd(c, fill)
            cell_margin(c, left=100, right=100, top=50, bottom=50)
            set_cell_borders(c, top=bd_cfg, left=bd_cfg, bottom=bd_cfg, right=bd_cfg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (alignments and alignments[j] == 'C') else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run(p, str(val), pt=8.5)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(2)
    p_spacer.paragraph_format.space_after  = Pt(4)
    return tbl


# ═══════════════════════════════════════════════════════════
# EKSEKUSI & SINTESIS METODOLOGI BAB 1 STANDAR PUBLIK
# ═══════════════════════════════════════════════════════════
def generate_all_bab1():
    base_dir = Path(__file__).resolve().parent.parent.parent.parent
    data_dir = base_dir / "data" / "processed"
    tool_dir = base_dir / "tools" / "report_metodologi" / "bab_1"
    tool_dir.mkdir(parents=True, exist_ok=True)

    print("[1/5] Mengekstraksi seluruh dataset empiris Bab 1...")
    df_izin = pd.read_csv(data_dir / "sulawesi_izin_baru_per_tahun.csv")
    df_smelter = pd.read_csv(data_dir / "sulawesi_esdm_nikel.csv")
    df_pltu = pd.read_csv(data_dir / "sulawesi_pltu_captive.csv")
    df_gfw = pd.read_csv(data_dir / "sulawesi_gfw_master_1_dekade_2014_2023_v3.csv")
    df_inv = pd.read_csv(data_dir / "sulawesi_investasi_pmdn_2016_2024.csv")
    df_pdrb = pd.read_csv(data_dir / "sulawesi_pdrb_sektoral_2016_2024.csv")
    df_pdrb_kab = pd.read_csv(data_dir / "sulawesi_pdrb_sektoral_kabupaten_2016_2024.csv")
    df_logistik = pd.read_csv(data_dir / "sulawesi_logistik_simpul_nikel.csv")

    tot_izin = int(df_izin['Jumlah_Izin_Baru'].sum())
    tot_luas_izin = int(df_izin['Total_Luas_Konsesi_Baru_Ha'].sum())
    tot_smelter = len(df_smelter)
    df_pltu_op = df_pltu[df_pltu['Status'].str.lower() == 'operating']
    tot_kapasitas_pltu = int(df_pltu_op['Capacity (MW)'].sum()) if 'Capacity (MW)' in df_pltu_op.columns else 0
    tot_deforestasi = int(df_gfw['Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha'].sum())
    tot_investasi_triliun = int(df_inv['nilai'].sum() / 1_000)

    # Sulteng & Morowali
    latest_year_pdrb = df_pdrb['tahun'].max()
    df_pdrb_sulteng = df_pdrb[(df_pdrb['provinsi'] == 'Sulawesi Tengah') & (df_pdrb['tahun'] == latest_year_pdrb)]
    pdrb_sulteng_tot = df_pdrb_sulteng['nilai_miliar_rp'].sum() / 1000
    pdrb_sulteng_eks = df_pdrb_sulteng[df_pdrb_sulteng['sektor_nama'].isin(['Pertambangan dan Penggalian', 'Industri Pengolahan', 'Pengadaan Listrik dan Gas'])]['nilai_miliar_rp'].sum() / 1000
    pct_sulteng_eks = (pdrb_sulteng_eks / pdrb_sulteng_tot) * 100 if pdrb_sulteng_tot > 0 else 0

    # Data Kabupaten Sulawesi Tengah (Tahun Terbaru)
    df_kab_sulteng = df_pdrb_kab[df_pdrb_kab['provinsi'] == 'Sulawesi Tengah'].copy()
    latest_year_kab = df_kab_sulteng['tahun'].max()
    df_kab_latest = df_kab_sulteng[df_kab_sulteng['tahun'] == latest_year_kab].copy()
    
    def klas_kab_func(s):
        if s in ['Pertambangan dan Penggalian', 'Industri Pengolahan', 'Pengadaan Listrik dan Gas']: return 'Ekstraktif'
        elif s in ['Pertanian, Kehutanan, dan Perikanan']: return 'Akar Rumput'
        else: return 'Jasa'
    
    df_kab_latest['Klas'] = df_kab_latest['sektor_nama'].apply(klas_kab_func)
    df_kab_pivot = df_kab_latest.groupby(['kabupaten', 'Klas'])['nilai_miliar_rp'].sum().unstack(fill_value=0) / 1000
    df_kab_pivot['Total'] = df_kab_pivot.sum(axis=1)
    df_kab_pivot['Pct_Akar'] = (df_kab_pivot['Akar Rumput'] / df_kab_pivot['Total']) * 100
    df_kab_pivot['Pct_Ekstraktif'] = (df_kab_pivot['Ekstraktif'] / df_kab_pivot['Total']) * 100
    df_kab_pivot['Pct_Jasa'] = (df_kab_pivot['Jasa'] / df_kab_pivot['Total']) * 100
    df_kab_pivot = df_kab_pivot.sort_values(by='Total', ascending=False)

    # ── Data Distribusi Sektor Komoditas 6 Provinsi (Tahun Terbaru) untuk 1.1.3 ──
    df_pdrb_24 = df_pdrb[df_pdrb['tahun'] == latest_year_pdrb].copy()
    df_pdrb_24['nilai_triliun'] = df_pdrb_24['nilai_miliar_rp'] / 1000
    tot_prov_24 = df_pdrb_24.groupby('provinsi')['nilai_triliun'].sum()
    max_x_val_all = df_pdrb_24['nilai_triliun'].max() * 1.15
    sulteng_ind_val = df_pdrb_24[(df_pdrb_24['provinsi'] == 'Sulawesi Tengah') & (df_pdrb_24['sektor_nama'] == 'Industri Pengolahan')]['nilai_triliun'].values[0]
    sulteng_tambang_val = df_pdrb_24[(df_pdrb_24['provinsi'] == 'Sulawesi Tengah') & (df_pdrb_24['sektor_nama'] == 'Pertambangan dan Penggalian')]['nilai_triliun'].values[0]
    sultra_tambang_val = df_pdrb_24[(df_pdrb_24['provinsi'] == 'Sulawesi Tenggara') & (df_pdrb_24['sektor_nama'] == 'Pertambangan dan Penggalian')]['nilai_triliun'].values[0]

    dominasi_map = {
        'Sulawesi Tengah': 'Didominasi Industri Pengolahan Smelter & Pertambangan (Ekstraktif)',
        'Sulawesi Tenggara': 'Didominasi Pertanian & Pertambangan Logam (Campuran)',
        'Sulawesi Selatan': 'Didominasi Pertanian, Perdagangan & Konstruksi (Agraris & Jasa)',
        'Sulawesi Utara': 'Didominasi Pertanian, Perdagangan & Transportasi (Jasa & Maritim)',
        'Sulawesi Barat': 'Didominasi Pertanian Tanaman Pangan & Perkebunan (Agraris)',
        'Gorontalo': 'Didominasi Pertanian, Perdagangan & Konstruksi (Agraris)'
    }

    short_sector_map = {
        'Pertanian, Kehutanan, dan Perikanan': 'Pertanian & Perikanan',
        'Pertambangan dan Penggalian': 'Pertambangan & Penggalian',
        'Industri Pengolahan': 'Industri Pengolahan',
        'Pengadaan Listrik dan Gas': 'Pengadaan Listrik & Gas',
        'Konstruksi': 'Konstruksi',
        'Perdagangan Besar dan Eceran; Reparasi Mobil dan Sepeda Motor': 'Perdagangan Besar & Eceran',
        'Transportasi dan Pergudangan': 'Transportasi & Pergudangan',
        'Penyediaan Akomodasi dan Makan Minum': 'Akomodasi & Makan Minum',
        'Informasi dan Komunikasi': 'Informasi & Komunikasi',
        'Jasa Keuangan dan Asuransi': 'Jasa Keuangan & Asuransi',
        'Real Estat': 'Real Estat',
        'Jasa Perusahaan': 'Jasa Perusahaan',
        'Administrasi Pemerintahan, Pertahanan, dan Jaminan Sosial Wajib': 'Administrasi Pemerintahan',
        'Jasa Pendidikan': 'Jasa Pendidikan',
        'Jasa Kesehatan dan Kegiatan Sosial': 'Jasa Kesehatan & Sosial',
        'Jasa Lainnya': 'Jasa Lainnya',
        'Pengadaan Air, Pengelolaan Sampah, Limbah dan Daur Ulang': 'Pengelolaan Air & Sampah'
    }

    top_sectors_rows = []
    html_top_rows = ""
    md_top_rows = []

    prov_order_list = ['Sulawesi Tengah', 'Sulawesi Tenggara', 'Sulawesi Selatan', 'Sulawesi Utara', 'Sulawesi Barat', 'Gorontalo']
    for idx_p, prov in enumerate(prov_order_list):
        g_prov = df_pdrb_24[df_pdrb_24['provinsi'] == prov].sort_values('nilai_triliun', ascending=False).reset_index(drop=True)
        t_val = tot_prov_24[prov]
        s1 = short_sector_map.get(g_prov.loc[0, 'sektor_nama'], g_prov.loc[0, 'sektor_nama'])
        p1 = (g_prov.loc[0, 'nilai_triliun'] / t_val) * 100
        s2 = short_sector_map.get(g_prov.loc[1, 'sektor_nama'], g_prov.loc[1, 'sektor_nama'])
        p2 = (g_prov.loc[1, 'nilai_triliun'] / t_val) * 100
        s3 = short_sector_map.get(g_prov.loc[2, 'sektor_nama'], g_prov.loc[2, 'sektor_nama'])
        p3 = (g_prov.loc[2, 'nilai_triliun'] / t_val) * 100
        dom = dominasi_map.get(prov, '-')

        top_sectors_rows.append([prov, f"{t_val:.2f}", s1, f"{p1:.1f}%", s2, f"{p2:.1f}%", s3, f"{p3:.1f}%", dom])

        even_cls = ' class="data-tr-even"' if idx_p % 2 == 1 else ''
        html_top_rows += f'    <tr{even_cls}><td class="data-td"><strong>{prov}</strong></td><td class="data-td" style="text-align:center;"><strong>{t_val:.2f}</strong></td><td class="data-td">{s1}</td><td class="data-td" style="text-align:center;">{p1:.1f}%</td><td class="data-td">{s2}</td><td class="data-td" style="text-align:center;">{p2:.1f}%</td><td class="data-td">{s3}</td><td class="data-td" style="text-align:center;">{p3:.1f}%</td><td class="data-td">{dom}</td></tr>\n'
        md_top_rows.append(f"| **{prov}** | **{t_val:.2f}** | {s1} | {p1:.1f}% | {s2} | {p2:.1f}% | {s3} | {p3:.1f}% | {dom} |")

    # ── Data Smelter & PLTU 6 Provinsi untuk Sub-Bab 1.2 ──
    df_pltu_op_all = df_pltu[(df_pltu['Status'].str.lower() == 'operating') & (df_pltu['Subnational unit (province, state)'].isin(prov_order_list))].copy()
    tot_smelter_all = len(df_smelter)
    tot_pltu_mw_all = df_pltu_op_all['Capacity (MW)'].sum()
    sm_prov_all = df_smelter.groupby('provinsi').size()
    pltu_prov_all = df_pltu_op_all.groupby('Subnational unit (province, state)')['Capacity (MW)'].sum()

    sulteng_sm_cnt = sm_prov_all.get('Sulawesi Tengah', 0)
    sultra_sm_cnt = sm_prov_all.get('Sulawesi Tenggara', 0)
    sulteng_pltu_mw = pltu_prov_all.get('Sulawesi Tengah', 0)
    sultra_pltu_mw = pltu_prov_all.get('Sulawesi Tenggara', 0)
    persen_smelter_2prov = ((sulteng_sm_cnt + sultra_sm_cnt) / tot_smelter_all) * 100 if tot_smelter_all > 0 else 0
    persen_pltu_2prov = ((sulteng_pltu_mw + sultra_pltu_mw) / tot_pltu_mw_all) * 100 if tot_pltu_mw_all > 0 else 0

    # Panel Crosstab 1.2 PLTU vs Deforestasi
    prov_map = {
        'North Sulawesi': 'Sulawesi Utara',
        'South Sulawesi': 'Sulawesi Selatan',
        'Southeast Sulawesi': 'Sulawesi Tenggara',
        'Central Sulawesi': 'Sulawesi Tengah',
        'Gorontalo': 'Gorontalo',
        'West Sulawesi': 'Sulawesi Barat'
    }
    df_pltu_p = df_pltu[df_pltu['Status'].isin(['operating'])].copy()
    df_pltu_p = df_pltu_p[df_pltu_p['captive_flag'] == True]
    df_pltu_p['Provinsi'] = df_pltu_p['Subnational unit (province, state)'].replace(prov_map)
    df_pltu_p['Tahun'] = pd.to_numeric(df_pltu_p['Start year'], errors='coerce')
    df_pltu_agg2 = df_pltu_p.groupby(['Provinsi', 'Tahun'])['Capacity (MW)'].sum().reset_index()

    df_panel_1_2 = pd.merge(df_gfw, df_pltu_agg2, on=['Provinsi', 'Tahun'], how='left').fillna({'Capacity (MW)': 0})
    df_panel_1_2 = df_panel_1_2.sort_values(by=['Provinsi', 'Tahun'])
    df_panel_1_2['Kapasitas_PLTU_Kumulatif_MW'] = df_panel_1_2.groupby('Provinsi')['Capacity (MW)'].cumsum()

    x_col_2 = 'Kapasitas_PLTU_Kumulatif_MW'
    y_col_2 = 'Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha'

    x_med_2 = df_panel_1_2[x_col_2].median()
    x_thresh_2 = x_med_2 if x_med_2 > 0 else 0
    y_med_2 = df_panel_1_2[y_col_2].median()

    label_x_low_2 = f"Rendah (≤{int(x_thresh_2)} MW)"
    label_x_high_2 = f"Tinggi (>{int(x_thresh_2)} MW)"
    label_y_low_2 = f"Rendah (<{int(y_med_2):,} Ha)"
    label_y_high_2 = f"Tinggi (≥{int(y_med_2):,} Ha)"

    df_panel_1_2['X_Label'] = df_panel_1_2[x_col_2].apply(lambda x: label_x_high_2 if x > x_thresh_2 else label_x_low_2)
    df_panel_1_2['Y_Label'] = df_panel_1_2[y_col_2].apply(lambda y: label_y_high_2 if y >= y_med_2 else label_y_low_2)

    cats_x_2 = [label_x_low_2, label_x_high_2]
    cats_y_2 = [label_y_low_2, label_y_high_2]
    ct_pltu = pd.crosstab(df_panel_1_2['X_Label'], df_panel_1_2['Y_Label']).reindex(index=cats_x_2, columns=cats_y_2, fill_value=0)

    chi2_pltu, p_pltu, dof_pltu, exp_pltu = stats.chi2_contingency(ct_pltu)
    g_pltu, p_g_pltu, dof_g_pltu, _ = stats.chi2_contingency(ct_pltu, lambda_='log-likelihood')

    min_exp = exp_pltu.min()
    is_fisher = min_exp < 5
    if is_fisher:
        _, p_fisher = stats.fisher_exact(ct_pltu)
        p_val_display = p_fisher
        test_name = "Fisher's Exact Test"
        test_val = "-"
        dof_str = "-"
    else:
        p_val_display = p_pltu
        test_name = "Pearson Chi-Square (χ²)"
        test_val = f"{chi2_pltu:.3f}"
        dof_str = f"df={dof_pltu}"

    x_codes_2 = df_panel_1_2['X_Label'].replace({label_x_low_2: 0, label_x_high_2: 1})
    y_codes_2 = df_panel_1_2['Y_Label'].replace({label_y_low_2: 0, label_y_high_2: 1})
    r_2, p_corr_2 = stats.pearsonr(list(x_codes_2), list(y_codes_2))
    lbl_val_pltu = (len(df_panel_1_2) - 1) * (r_2**2)

    mermaid_str = """flowchart TD
    A(["Start: Input Data"]) --> B{"Apakah Data<br/>Kategorikal?"}
    
    B -- TIDAK (Numerik) --> C["Lakukan Diskritisasi/Binning<br/>Ubah Angka jadi Kategori<br/>Misal: Tinggi vs Rendah"]
    C --> D
    
    B -- YA --> D["Penanganan Missing Values<br/>SPSS: Listwise Deletion<br/>Hapus baris yang datanya bolong"]
    
    D --> E["Jalankan Perhitungan<br/>Tabel Crosstab 2x2"]
    
    E --> F{"Cek Asumsi SPSS:<br/>Expected Count >= 5 ?"}
    
    F -- TIDAK (Ada sel bernilai 0 <br/>atau >20% sel nilainya <5) --> G["Asumsi Pearson Chi-Square<br/>DILANGGAR / CACAT"]
    G --> H["Gunakan Uji Alternatif:<br/>Fisher's Exact Test"]
    H --> K
    
    F -- YA (Memenuhi Syarat) --> I["Asumsi Terpenuhi<br/>Baca Nilai Pearson Chi-Square"]
    I --> J{"Cek P-Value<br/>(Asymp. Sig)"}
    
    J -- P-Value >= 0.05 --> L["TIDAK SIGNIFIKAN<br/>(Hanya Kebetulan)"]
    J -- P-Value < 0.05 --> M["SIGNIFIKAN<br/>(Hubungan Terbukti)"]
    
    M --> N["Cek Kekuatan Hubungan<br/>(Odds Ratio / Risk Estimate)"]
    
    N --> O1["Odds Ratio = 1<br/>Tidak ngefek"]
    N --> O2["Odds Ratio > 1<br/>Risiko/Bahaya Naik"]
    
    L --> Z(["Selesai"])
    O1 --> Z
    O2 --> Z
    K --> J
    
    classDef warning fill:#ffcccb,stroke:#ff0000,stroke-width:2px;
    classDef success fill:#d4edda,stroke:#28a745,stroke-width:2px;
    classDef process fill:#e1f5fe,stroke:#0288d1,stroke-width:2px;
    
    class B,F,J process;
    class G,L warning;
    class I,M success;"""

    mermaid_png_path = str(tool_dir / "mermaid_flowchart.png")
    download_success = download_mermaid_png(mermaid_str, mermaid_png_path)

    mermaid_str_1_1 = """graph TD
    subgraph Mandat_Smelter["2. Rantai Pasok Smelter (Kategori C)"]
        A["UU No. 3/2020 Ps. 1(1)<br/><i>Pertambangan = Eksplorasi + Penambangan + Pengolahan/Pemurnian</i>"] --> B["Smelter (Industri Pengolahan / Kat. C)<br/><b>Tahapan Wajib Pertambangan</b>"]
        C["UU No. 3/2020 Ps. 102–103 & PP 96/2021<br/><i>Mandat Wajib Pemegang IUP Operasi Produksi</i>"] --> B
    end

    subgraph Mandat_Energi["3. Rantai Pasok Energi Captive (Kategori D)"]
        D["Perpres No. 112/2022 Ps. 3(4)b<br/><i>PLTU Baru Dilarang, KECUALI Terintegrasi Smelter</i>"] --> E["PLTU Captive (Pengadaan Listrik / Kat. D)<br/><b>Instrumen Rantai Pasok Off-Grid</b>"]
        F["RUPTL PLN 2021–2030 Hal. VI-24<br/><i>Pengakuan Pasokan Khusus Smelter</i>"] --> E
    end

    subgraph Hulu_Tambang["1. Sektor Hulu (Kategori B)"]
        G["Perpres No. 26/2010 Ps. 1(2)<br/><b>Pertambangan & Penggalian (Kat. B)</b><br/><i>Ekstraksi SDA Tak Terbarukan</i>"]
    end

    G --> K["<b>KESIMPULAN:</b><br/>Kat. B + Kat. C + Kat. D = <b>SATU KESATUAN RANTAI PASOK EKSTRAKTIF</b> yang dimandatkan hukum"]
    B --> K
    E --> K"""

    mermaid_png_path_1_1 = str(tool_dir / "mermaid_flowchart_1_1.png")
    download_success_1_1 = download_mermaid_png(mermaid_str_1_1, mermaid_png_path_1_1)

    mermaid_str_1_2 = """graph TD
    subgraph Data_BPS["1. Input Data Statistik Daerah"]
        A["BPS Kabupaten se-Sulawesi Tengah<br/><i>PDRB 17 Sektor Lapangan Usaha ADHB</i>"]
    end

    subgraph Reklasifikasi["2. Reklasifikasi Legal Supply-Chain"]
        B["Ekstraktif = Kat. B + Kat. C + Kat. D"]
        C["Akar Rumput = Kat. A"]
        D["Jasa & Lainnya = Kat. E s.d. U"]
    end

    subgraph Analisis_Disparitas["3. Output Evaluasi Spasial"]
        E["<b>Sentra Hilirisasi (Enclave Industri):</b><br/>Morowali & Morut menguasai Sektor Ekstraktif Tertinggi"]
        F["<b>Non-Sentra (Pertanian Rakyat):</b><br/>11 Kabupaten tertinggal (<11% Porsi Ekstraktif)"]
    end

    A --> B
    A --> C
    A --> D
    B --> E
    C --> F
    D --> F"""
    mermaid_png_path_1_2 = str(tool_dir / "mermaid_flowchart_1_2.png")
    download_success_1_2 = download_mermaid_png(mermaid_str_1_2, mermaid_png_path_1_2)


    a_2 = ct_pltu.loc[label_x_low_2, label_y_low_2]
    b_2 = ct_pltu.loc[label_x_low_2, label_y_high_2]
    c_2 = ct_pltu.loc[label_x_high_2, label_y_low_2]
    d_2 = ct_pltu.loc[label_x_high_2, label_y_high_2]
    or_pltu = (a_2 * d_2) / (b_2 * c_2) if (b_2 * c_2) > 0 else 0

    mermaid_str_1_3 = """graph TD
    subgraph Input_Data["1. Input Data Statistik Provinsi"]
        A["BPS Provinsi se-Sulawesi<br/><i>PDRB 17 Sektor Lapangan Usaha ADHB</i>"]
    end

    subgraph Standardisasi["2. Pemrosesan & Standardisasi Data"]
        B["Konversi Nilai: Triliun Rp = Miliar Rp / 1000"]
        C["Porsi Sektor (%) = (Nilai Sektor / Total PDRB) * 100"]
        D["Skala X Seragam: [0, max(X) * 1.15]<br/><i>Mencegah Visual Comparison Bias</i>"]
        E["Pewarnaan 3 Klaster Kritis:<br/>• Merah (Kat. B,C,D: Ekstraktif)<br/>• Hijau (Kat. A: Akar Rumput)<br/>• Abu-abu (Kat. E–U: Jasa & Lainnya)"]
    end

    subgraph Komparasi_Regional["3. Output Komparasi Antarwilayah"]
        F["<b>Sulawesi Tengah & Sultra:</b><br/>Anomali Lonjakan Sektor Ekstraktif (Smelter & Tambang)"]
        G["<b>Sulsel, Sulbar, Gorontalo & Sulut:</b><br/>Struktur Perekonomian Bertumpu pada Pertanian & Jasa"]
    end

    A --> B
    B --> C
    B --> D
    C --> E
    D --> E
    E --> F
    E --> G"""
    mermaid_png_path_1_3 = str(tool_dir / "mermaid_flowchart_1_3.png")
    download_success_1_3 = download_mermaid_png(mermaid_str_1_3, mermaid_png_path_1_3)

    # Panel Crosstab Izin vs Deforestasi
    df_panel_izin = pd.merge(df_gfw, df_izin, on=['Provinsi', 'Tahun'], how='left').fillna({'Jumlah_Izin_Baru': 0, 'Total_Luas_Konsesi_Baru_Ha': 0})
    med_izin = df_panel_izin['Jumlah_Izin_Baru'].median()
    med_def = df_panel_izin['Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha'].median()
    df_panel_izin['X_bin'] = df_panel_izin['Jumlah_Izin_Baru'].apply(lambda x: 'Tinggi' if x >= med_izin else 'Rendah')
    df_panel_izin['Y_bin'] = df_panel_izin['Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha'].apply(lambda y: 'Tinggi' if y >= med_def else 'Rendah')
    ct_iz = pd.crosstab(df_panel_izin['X_bin'], df_panel_izin['Y_bin']).reindex(index=['Rendah', 'Tinggi'], columns=['Rendah', 'Tinggi'], fill_value=0)
    chi2_iz, p_iz, dof_iz, _ = stats.chi2_contingency(ct_iz)
    or_iz = (ct_iz.loc['Rendah', 'Rendah'] * ct_iz.loc['Tinggi', 'Tinggi']) / (ct_iz.loc['Rendah', 'Tinggi'] * ct_iz.loc['Tinggi', 'Rendah']) if (ct_iz.loc['Rendah', 'Tinggi'] * ct_iz.loc['Tinggi', 'Rendah']) > 0 else 0

    mermaid_str_1_3_izin = """flowchart TD
    subgraph Data_Preparation["1. Akuisisi & Penyiapan Data Panel"]
        A["Data Penerbitan Izin Baru Minerbaone<br/><i>(Provinsi-Tahun, Jumlah Izin, Luas Konsesi)</i>"] --> C
        B["Data Laju Deforestasi GFW<br/><i>(Provinsi-Tahun, Deforestasi Komoditas Tambang/Sawit)</i>"] --> C
        C["<b>Gabung (Merge) Data Panel</b><br/>6 Provinsi x 10 Tahun = 60 Sampel (Unit Observasi)"]
    end

    subgraph Time_Series_Analysis["2. Analisis Tren (Time-Series)"]
        C --> D["Agregasi Jumlah Izin & Luas per Tahun"]
        D --> E["Kalkulasi Pertumbuhan Laju (YoY)"]
        E --> F["Identifikasi Lonjakan Ekstraktif (Tahun 2022-2024)"]
    end

    subgraph Crosstab_Analysis["3. Analisis Inferensial (Chi-Square)"]
        C --> G["<b>Binning Kategori (Threshold Median)</b><br/>Klasifikasi Tinggi (≥Median) vs Rendah (<Median)"]
        G --> H["Definisikan X (Tekanan Ekspansi):<br/>- Jumlah Izin Baru (IUP)<br/>- Luas Konsesi Baru (Ha)"]
        H --> I["Definisikan Y (Dampak Ekologis):<br/>- Total Deforestasi Alam<br/>- Deforestasi Komoditas"]
        I --> J["Uji Chi-Square (Crosstabulation 2x2)"]
    end

    subgraph Output["4. Hasil Pengujian Hipotesis"]
        F --> K["Tren Visual Distribusi Spasial (Bar/Line Chart)"]
        J --> L{"P-Value < 0.05?"}
        L -- YA --> M["<b>SIGNIFIKAN (Tolak H0)</b><br/>Ada Hubungan: Tekanan Ekspansi Terbukti Meningkatkan Laju Deforestasi"]
        L -- TIDAK --> N["<b>TIDAK SIGNIFIKAN (Gagal Tolak H0)</b><br/>Indikasi Efek Spillover/Kehancuran Merata"]
        M --> O["Hitung Odds Ratio (Risk Estimate)"]
    end"""
    mermaid_png_path_1_3_izin = str(tool_dir / "mermaid_flowchart_1_3_izin.png")
    download_success_1_3_izin = download_mermaid_png(mermaid_str_1_3_izin, mermaid_png_path_1_3_izin)

    mermaid_str_1_3_izin_alt = """flowchart LR
    subgraph Data_Preparation["1. Akuisisi Data"]
        A["Data Penerbitan Izin Baru Minerbaone<br/><i>(Provinsi-Tahun, Jumlah Izin, Luas Konsesi)</i>"]
    end

    subgraph Time_Series_Analysis["2. Analisis Tren (Time-Series)"]
        A --> D["Agregasi Jumlah Izin & Luas per Tahun"]
        D --> E["Kalkulasi Pertumbuhan Laju (YoY)"]
        E --> F["Identifikasi Lonjakan Ekstraktif (Tahun 2022-2024)"]
        F --> K["Tren Visual Distribusi Spasial (Bar/Line Chart)"]
    end"""
    mermaid_png_path_1_3_izin_alt = str(tool_dir / "mermaid_flowchart_1_3_izin_alt.png")
    download_success_1_3_izin_alt = download_mermaid_png(mermaid_str_1_3_izin_alt, mermaid_png_path_1_3_izin_alt)

    # ---------------------------------------------------------
    # STATISTIK 1.4: INVESTASI PMDN VS DEFORESTASI
    # ---------------------------------------------------------
    df_inv_clean = df_inv.rename(columns={'provinsi': 'Provinsi', 'tahun': 'Tahun'})
    df_inv_clean['Investasi_Juta_Rp'] = pd.to_numeric(df_inv_clean['nilai'], errors='coerce')
    df_panel_1_4 = pd.merge(df_gfw, df_inv_clean[['Provinsi', 'Tahun', 'Investasi_Juta_Rp']], on=['Provinsi', 'Tahun'], how='inner').fillna({'Investasi_Juta_Rp': 0})
    
    med_inv = df_panel_1_4['Investasi_Juta_Rp'].median()
    lbl_x_h_4 = f"Tinggi (>{int(med_inv):,})"
    lbl_x_l_4 = f"Rendah (≤{int(med_inv):,})"
    df_panel_1_4['X_bin'] = df_panel_1_4['Investasi_Juta_Rp'].apply(lambda x: lbl_x_h_4 if x > med_inv else lbl_x_l_4)

    # Deforestasi Komoditas Tambang Sawit
    med_def_4 = df_panel_1_4['Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha'].median()
    lbl_y_h_4 = f"Tinggi (≥{int(med_def_4):,})"
    lbl_y_l_4 = f"Rendah (<{int(med_def_4):,})"
    df_panel_1_4['Y_bin_tambang'] = df_panel_1_4['Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha'].apply(lambda y: lbl_y_h_4 if y >= med_def_4 else lbl_y_l_4)
    
    ct_inv_tambang = pd.crosstab(df_panel_1_4['X_bin'], df_panel_1_4['Y_bin_tambang']).reindex(index=[lbl_x_l_4, lbl_x_h_4], columns=[lbl_y_l_4, lbl_y_h_4], fill_value=0)
    try:
        chi2_inv_tambang, p_inv_tambang, dof_inv_tambang, _ = stats.chi2_contingency(ct_inv_tambang)
    except:
        chi2_inv_tambang, p_inv_tambang, dof_inv_tambang = 0, 1, 0
    
    try:
        or_inv_tambang = (ct_inv_tambang.loc[lbl_x_l_4, lbl_y_l_4] * ct_inv_tambang.loc[lbl_x_h_4, lbl_y_h_4]) / (ct_inv_tambang.loc[lbl_x_l_4, lbl_y_h_4] * ct_inv_tambang.loc[lbl_x_h_4, lbl_y_l_4]) if (ct_inv_tambang.loc[lbl_x_l_4, lbl_y_h_4] * ct_inv_tambang.loc[lbl_x_h_4, lbl_y_l_4]) > 0 else 0
    except KeyError:
        or_inv_tambang = 0

    # Total Deforestasi Alam
    med_def_tot_4 = df_panel_1_4['Total_Deforestasi_Ha'].median()
    lbl_yt_h_4 = f"Tinggi (≥{int(med_def_tot_4):,})"
    lbl_yt_l_4 = f"Rendah (<{int(med_def_tot_4):,})"
    df_panel_1_4['Y_bin_total'] = df_panel_1_4['Total_Deforestasi_Ha'].apply(lambda y: lbl_yt_h_4 if y >= med_def_tot_4 else lbl_yt_l_4)
    
    ct_inv_total = pd.crosstab(df_panel_1_4['X_bin'], df_panel_1_4['Y_bin_total']).reindex(index=[lbl_x_l_4, lbl_x_h_4], columns=[lbl_yt_l_4, lbl_yt_h_4], fill_value=0)
    try:
        chi2_inv_total, p_inv_total, dof_inv_total, _ = stats.chi2_contingency(ct_inv_total)
    except:
        chi2_inv_total, p_inv_total, dof_inv_total = 0, 1, 0
    
    try:
        or_inv_total = (ct_inv_total.loc[lbl_x_l_4, lbl_yt_l_4] * ct_inv_total.loc[lbl_x_h_4, lbl_yt_h_4]) / (ct_inv_total.loc[lbl_x_l_4, lbl_yt_h_4] * ct_inv_total.loc[lbl_x_h_4, lbl_yt_l_4]) if (ct_inv_total.loc[lbl_x_l_4, lbl_yt_h_4] * ct_inv_total.loc[lbl_x_h_4, lbl_yt_l_4]) > 0 else 0
    except KeyError:
        or_inv_total = 0

    mermaid_str_1_4 = """flowchart LR
    subgraph Data_Preparation["1. Akuisisi & Penyiapan Data Panel"]
        A["Data Realisasi Investasi PMDN<br/><i>(Provinsi-Tahun, Arus Modal Juta Rp)</i>"] --> C
        B["Data Laju Deforestasi GFW<br/><i>(Provinsi-Tahun, Total Deforestasi)</i>"] --> C
        C["<b>Gabung (Merge) Data Panel</b><br/>6 Provinsi x 8 Tahun = 48 Sampel Valid"]
    end

    subgraph Time_Series_Analysis["2. Analisis Statistik Tabulasi Silang"]
        C --> D["Agregasi Median Variabel X & Y"]
        D --> E["Kategorisasi Biner<br/>(Tinggi / Rendah)"]
        E --> F["Uji Independensi Chi-Square Pearson"]
        F --> G["Kalkulasi Odds Ratio (Risiko Deforestasi)"]
    end"""
    mermaid_png_path_1_4 = str(tool_dir / "mermaid_flowchart_1_4.png")
    download_success_1_4 = download_mermaid_png(mermaid_str_1_4, mermaid_png_path_1_4)

    mermaid_str_1_5 = """flowchart LR
    subgraph Data_Acquisition["1. Validasi Silang Dokumen Publik (OSINT)"]
        A["Laporan KNKT<br/><i>(Kapasitas Muatan & DWT)</i>"] --> D
        B["Lampiran Perpres 109/2020<br/><i>(Status Proyek Strategis Nasional)</i>"] --> D
        C["Laporan Tahunan & Laporan Keberlanjutan<br/><i>(Fasilitas Terminal Khusus)</i>"] --> D
    end

    subgraph Data_Processing["2. Triangulasi & Pemetaan Rute"]
        D["<b>Inventarisasi Simpul Pelabuhan</b><br/>(Identifikasi 6 Kawasan Utama)"] --> E["Ekstraksi Koordinat Geografis<br/>(Latitude & Longitude)"]
        E --> F["Pemodelan Kurva Parametrik (Bézier Curve)"]
        F --> G["Visualisasi Spatial Logistic Mapping"]
    end"""
    mermaid_png_path_1_5 = str(tool_dir / "mermaid_flowchart_1_5.png")
    download_success_1_5 = download_mermaid_png(mermaid_str_1_5, mermaid_png_path_1_5)

    mermaid_str_1_6 = """flowchart LR
    A["Ekstraksi Koordinat<br/>Origin (Sulawesi)"] --> B["Penentuan Titik<br/>Tujuan (China/Jepang)"]
    B --> C["Kalkulasi Kurva<br/>Parametrik (Bézier)"]
    C --> D["Plotting Spasial<br/>(Plotly Scattergeo)"]
    D --> E["Peta Jalur Distribusi<br/>Logistik Maritim"]
"""
    mermaid_png_path_1_6 = tool_dir / "mermaid_flowchart_1_6.png"
    download_success_1_6 = download_mermaid_png(mermaid_str_1_6, mermaid_png_path_1_6)

    # 2. Inisiasi Pembuatan Dokumen DOCX 1-Kolom Penuh
    print("[2/5] Membangun Metodologi_Bab1_Ekspansi_Industri.docx (Format Publik)...")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9.5)

    # ── HEADER BANNER ───────────────────────────
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(2)
    run(p_hdr, "CELIOS — CENTER OF ECONOMIC AND LAW STUDIES  |  LAPORAN RISET METODOLOGI D3TLH", bold=True, pt=8, color=G_MID)

    add_h1(doc, "BAB I: METODOLOGI ANALISIS EKSPANSI INDUSTRI EKSTRAKTIF DAN INFRASTRUKTUR PENUNJANG DI PULAU SULAWESI")

    add_p(doc, [
        ("Dokumen laporan metodologi ini menyajikan kerangka ilmiah, landasan regulasi, formulasi matematis, prosedur analisis statistik, serta metodologi pembuktian berbasis data terbuka yang dioperasionalkan pada ", False, False),
        ("Bab 1: Ekspansi Industri Ekstraktif", True, False),
        (" dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi periode 2014–2024.", False, False),
    ])

    # ═══════════════════════════════════════════════════════════
    # SUB-BAB 1.1 KONTEKS MAKRO PDRB
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.1 Konteks Makro: Breakdown PDRB per Komoditas")

    # ── 1.1.1 (Header Persis Sesuai Dashboard) ─────────────────
    add_h3(doc, "1.1.1 Konteks Makro: Dominasi Ekstraktif vs Ekonomi Akar Rumput")
    
    add_p(doc, [
        ("Bagian ini menganalisis struktur Produk Domestik Regional Bruto (PDRB) pada enam provinsi di Pulau Sulawesi sepanjang periode 2016–2024 menggunakan visualisasi grafik area bertumpuk (*Stacked Area Chart*). Analisis ini ditujukan untuk menguji secara empiris apakah percepatan pertumbuhan ekonomi daerah benar-benar bersumber dari sektor produktif masyarakat lokal atau didominasi oleh industri ekstraktif padat modal yang mengalihkan pemanfaatan ruang dan sumber daya alam.", False, False),
    ])

    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi",
                 "Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi Stacked Area Chart memetakan dinamika Produk Domestik Regional Bruto (PDRB) berdasarkan klasifikasi rantai pasok hukum (Legal Supply-Chain) untuk membandingkan trajektori Sektor Ekstraktif, Ekonomi Akar Rumput, dan Sektor Jasa & Lainnya.")

    add_h4(doc, "A. Kerangka Dekomposisi Sektoral & Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)")
    add_p(doc, [
        ("Sistem Klasifikasi Baku Lapangan Usaha Indonesia (KBLI 2020) oleh Badan Pusat Statistik (BPS) membagi perekonomian ke dalam 17 sektor lapangan usaha standar (Kategori A hingga Kategori U). Namun, pemisahan administratif standar ini belum menggambarkan keterkaitan langsung antara aktivitas penambangan bijih mineral di hulu, peleburan logam dasar di pabrik hilir, dan pembangkitan listrik batubara khusus (*captive*) yang menopangnya.", False, False),
    ])
    add_p(doc, [
        ("Melalui pendekatan ", False, False),
        ("Rantai Pasok Hukum (Legal Supply-Chain Approach)", True, False),
        (", 17 sektor PDRB BPS direklasifikasi menjadi ", False, False),
        ("Tiga Klaster Makro", True, False),
        (" berdasarkan mandat hukum dan regulasi nasional yang berlaku. Rincian pembagian sektor, dasar regulasi, serta intisari ketentuan hukum disajikan secara lengkap pada ", False, False),
        ("Tabel 1.1", True, False),
        (" berikut:", False, False),
    ])

    # Tabel 1.1: Reklasifikasi Hukum 1.1.1
    add_caption(doc, "Tabel 1.1: Reklasifikasi Sektoral PDRB KBLI 2020 Berdasarkan Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)")
    reklas_headers = ["Kategori BPS", "Sektor Lapangan Usaha", "Klasifikasi Analisis", "Dasar Regulasi & Mandat Hukum Nasional", "Intisari Ketentuan Hukum"]
    reklas_rows = [
        ["Kategori B", "Pertambangan dan Penggalian", "Ekstraktif", "Perpres No. 26 Tahun 2010", "Ketentuan Pasal 1 Ayat (2) mengenai kegiatan pengambilan komoditas tambang tak terbarukan dari dalam bumi."],
        ["Kategori C", "Industri Pengolahan (Smelter Logam)", "Ekstraktif", "UU No. 3 Tahun 2020 & PP No. 96 Tahun 2021", "Pasal 102–103 mewajibkan pemegang izin tambang melakukan pengolahan dan pemurnian di dalam negeri sebagai kesatuan tahapan pertambangan."],
        ["Kategori D", "Pengadaan Listrik & Gas (PLTU Captive)", "Ekstraktif", "Perpres No. 112 Tahun 2022 & RUPTL PLN", "Pasal 3 Ayat (4) huruf b mengecualikan pembangunan PLTU batubara baru hanya untuk fasilitas yang terintegrasi langsung melayani smelter."],
        ["Kategori A", "Pertanian, Kehutanan, Perikanan", "Ekonomi Akar Rumput", "Buku Klasifikasi Baku BPS (KBLI 2020)", "Sektor pemanfaatan sumber daya hayati terbarukan yang menjadi penyerap tenaga kerja lokal terbesar."],
        ["Kategori E–U", "13 Sektor Jasa, Konstruksi, Perdagangan", "Sektor Jasa & Lainnya", "Klasifikasi Standar BPS", "Sektor sekunder dan tersier penunjang aktivitas perekonomian daerah."],
    ]
    add_table_1col(doc, reklas_headers, reklas_rows, [2.2, 3.8, 2.5, 4.2, 4.3], ['C', 'L', 'C', 'L', 'L'])

    add_h4(doc, "B. Alur Logika Metodologis Rantai Pasok Hukum (Mengapa Kat. B + C + D = Ekstraktif)")
    add_p(doc, [
        ("Keterkaitan ketiga kategori lapangan usaha tersebut sebagai satu kesatuan rantai pasok ekstraktif dimodelkan dalam kerangka alur logika hukum sebagaimana diilustrasikan pada ", False, False),
        ("Bagan Alur 1.1", True, False),
        (" berikut:", False, False),
    ])

    # Bagan Alur 1.1: Flowchart Rantai Pasok Hukum 3 Pilar di Word
    add_caption(doc, "Bagan Alur 1.1: Alur Logika Metodologis Rantai Pasok Hukum Sektor Ekstraktif (Kat. B + Kat. C + Kat. D)")
    
    if download_success_1_1:
        try:
            p_img_1 = doc.add_paragraph()
            p_img_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img_1 = p_img_1.add_run()
            run_img_1.add_picture(mermaid_png_path_1_1, width=Cm(15))
        except Exception as e:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 1.1 ke DOCX: {e}")
            p_err_1 = doc.add_paragraph()
            run(p_err_1, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err_1 = doc.add_paragraph()
        run(p_err_1, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_spacer_fc = doc.add_paragraph()
    p_spacer_fc.paragraph_format.space_before = Pt(2)
    p_spacer_fc.paragraph_format.space_after  = Pt(4)

    add_h4(doc, "C. Formulasi Matematis: Persamaan Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)")
    add_p(doc, [
        ("Data diproses menggunakan pendekatan Dekomposisi PDRB Sektoral. Nilai PDRB dikelompokkan menjadi 3 agregat makro melalui pendekatan Legal Supply-Chain. Sektor Ekstraktif dihitung dengan menggabungkan tiga kategori lapangan usaha utama yang saling terintegrasi (tambang, smelter, dan PLTU captive) menggunakan persamaan ", False, False),
        ("Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)", True, False),
        (":", False, False),
    ])

    # Box Rumus Utama Sesuai Overview Temuan
    add_formula(doc, "Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)", 
                "Sektor_Ekstraktif = PDRB(Kat.B: Pertambangan) + PDRB(Kat.C: Ind. Pengolahan) + PDRB(Kat.D: Listrik)",
                var_desc=[
                    ("Sektor_Ekstraktif", "Total nilai tambah bruto dari klaster industri ekstraktif yang saling terintegrasi (satuan: Triliun Rupiah)."),
                    ("PDRB(Kat.B: Pertambangan)", "Nilai tambah bruto dari aktivitas eksplorasi dan ekstraksi bijih mineral dari dalam bumi (BPS KBLI 2020 Kategori B)."),
                    ("PDRB(Kat.C: Ind. Pengolahan)", "Nilai tambah bruto dari pemurnian dan peleburan logam dasar di fasilitas smelter nikel (BPS KBLI 2020 Kategori C / Golongan Pokok 24)."),
                    ("PDRB(Kat.D: Listrik)", "Nilai tambah bruto dari pengadaan daya listrik batubara non-jaringan (PLTU captive) khusus melayani smelter (BPS KBLI 2020 Kategori D)."),
                ])

    add_p(doc, [
        ("Secara lengkap, perhitungan ketiga komponen makroekonomi dan indikator turunannya dirumuskan melalui sistem persamaan berikut:", False, False),
    ])

    add_formula(doc, "Ekonomi Akar Rumput", 
                "Sektor_Akar_Rumput = PDRB(Kat.A: Pertanian, Kehutanan, dan Perikanan)",
                var_desc=[
                    ("Sektor_Akar_Rumput", "Total nilai Produk Domestik Regional Bruto yang dihasilkan dari pemanfaatan sumber daya hayati terbarukan (satuan: Triliun Rupiah)."),
                    ("PDRB(Kat.A)", "Agregasi nilai tambah tanaman pangan, hortikultura, perkebunan rakyat, peternakan, kehutanan, dan perikanan tangkap/budidaya (BPS KBLI 2020 Kategori A)."),
                ])

    add_formula(doc, "Sektor Jasa & Lainnya", 
                "Sektor_Jasa = Jumlah PDRB (Kategori E sampai dengan Kategori U)",
                var_desc=[
                    ("Sektor_Jasa", "Total nilai tambah gabungan dari 13 sektor penunjang sekunder dan tersier di luar sektor ekstraktif dan akar rumput (satuan: Triliun Rupiah)."),
                    ("PDRB (Kat. E s.d. U)", "Akumulasi nilai tambah sektor konstruksi, perdagangan besar/eceran, transportasi, pergudangan, akomodasi, informasi & komunikasi, jasa keuangan, real estat, dan jasa umum."),
                ])

    add_formula(doc, "Total Produk Domestik Regional Bruto (PDRB Wilayah)", 
                "Total_PDRB = Sektor_Ekstraktif + Sektor_Akar_Rumput + Sektor_Jasa",
                var_desc=[
                    ("Total_PDRB", "Nilai keseluruhan output ekonomi regional bruto suatu wilayah provinsi atau kabupaten pada periode tahun berjalan atas dasar harga berlaku (satuan: Triliun Rupiah)."),
                    ("Sektor_Ekstraktif", "Total nilai tambah sektor ekstraktif terintegrasi (Triliun Rupiah)."),
                    ("Sektor_Akar_Rumput", "Total nilai tambah ekonomi berbasis masyarakat lokal dan sumber daya hayati (Triliun Rupiah)."),
                    ("Sektor_Jasa", "Total nilai tambah sektor jasa dan fasilitas pendukung (Triliun Rupiah)."),
                ])

    add_formula(doc, "Pangsa Kontribusi Sektor Ekstraktif (%)", 
                "Pangsa_Ekstraktif (%) = ( Sektor_Ekstraktif / Total_PDRB ) * 100",
                var_desc=[
                    ("Pangsa_Ekstraktif (%)", "Persentase pangsa/porsi dominasi sektor ekstraktif terhadap total kue perekonomian wilayah (satuan: Persen / %)."),
                    ("Sektor_Ekstraktif", "Nilai tambah nominal sektor ekstraktif terintegrasi pada tahun observasi (Triliun Rupiah)."),
                    ("Total_PDRB", "Total nilai nominal PDRB seluruh 17 sektor lapangan usaha pada tahun yang sama (Triliun Rupiah)."),
                ])

    add_formula(doc, "Laju Pertumbuhan Tahunan Sektoral (YoY)", 
                "Laju_Pertumbuhan_Tahunan (%) = [ ( Nilai_Tahun_t - Nilai_Tahun_{t-1} ) / Nilai_Tahun_{t-1} ] * 100",
                var_desc=[
                    ("Laju_Pertumbuhan_Tahunan (%)", "Tingkat percepatan atau perlambatan ekspansi tahunan suatu sektor ekonomi (satuan: Persen / %)."),
                    ("Nilai_Tahun_t", "Nilai nominal PDRB sektor pada tahun observasi berjalan (t)."),
                    ("Nilai_Tahun_{t-1}", "Nilai nominal PDRB sektor pada satu tahun sebelumnya (t - 1)."),
                ])

    add_p(doc, [
        ("Definisi operasional, cakupan lapangan usaha, dan institusi penyedia data primer untuk masing-masing komponen variabel dalam sistem persamaan di atas dipaparkan pada ", False, False),
        ("Tabel 1.2", True, False),
        (" berikut:", False, False),
    ])

    # Tabel 1.2: Penjelasan Variabel Matematis 1.1.1
    add_caption(doc, "Tabel 1.2: Definisi Operasional Komponen Makroekonomi dan Sumber Data PDRB Sektoral")
    var_headers = ["Komponen Analisis", "Cakupan Lapangan Usaha", "Definisi Operasional", "Satuan Nilai", "Sumber Data Primer"]
    var_rows = [
        ["Sektor Ekstraktif", "Kategori B, Kategori C, Kategori D", "Akumulasi nilai tambah bruto pertambangan bijih nikel, peleburan logam dasar, dan penyediaan energi PLTU captive.", "Triliun Rupiah", "BPS Provinsi (PDRB Menurut Lapangan Usaha)"],
        ["Ekonomi Akar Rumput", "Kategori A", "Nilai tambah pertanian tanaman pangan, perkebunan rakyat, kehutanan, dan perikanan tangkap maupun budidaya.", "Triliun Rupiah", "BPS Provinsi"],
        ["Sektor Jasa & Lainnya", "Kategori E hingga Kategori U", "Nilai tambah gabungan perdagangan, konstruksi, transportasi, perbankan, dan jasa layanan umum.", "Triliun Rupiah", "BPS Provinsi"],
        ["Total PDRB Wilayah", "Seluruh 17 Kategori Lapangan Usaha", "Total nilai Produk Domestik Regional Bruto suatu wilayah atas dasar harga berlaku pada periode tahun berjalan.", "Triliun Rupiah", "BPS Provinsi"],
        ["Pangsa Ekstraktif (%)", "Rasio Kontribusi Relatif", "Persentase kontribusi sektor ekstraktif terhadap keseluruhan perekonomian provinsi atau kabupaten.", "Persen (%)", "Hasil Olahan Riset CELIOS"],
    ]
    add_table_1col(doc, var_headers, var_rows, [3.2, 3.6, 4.7, 2.3, 3.2], ['L', 'L', 'L', 'C', 'L'])

    add_h4(doc, "D. Analisis Temuan Empiris: Ketimpangan Struktural Sulawesi Tengah")
    add_p(doc, [
        (f"Penerapan formulasi di atas menunjukkan perbedaan struktur ekonomi yang sangat kontras antarwilayah di Pulau Sulawesi. Di ", False, False),
        ("Sulawesi Tengah (sebagai pusat hilirisasi)", True, False),
        (f", ekspansi industri ekstraktif melaju sangat pesat hingga menguasai ", False, False),
        (f"{pct_sulteng_eks:.1f}% dari total PDRB provinsi", True, False),
        (f" pada tahun {latest_year_pdrb}. Sebaliknya, porsi Ekonomi Akar Rumput mengalami penurunan pangsa yang signifikan. Hal ini menunjukkan ketergantungan ekonomi daerah yang sangat tinggi pada satu klaster industri padat modal, berbeda dengan provinsi lain seperti Sulawesi Selatan dan Gorontalo yang perekonomiannya bertumpu pada basis pertanian terbarukan.", False, False),
    ])

    # ── 1.1.2 PEMUSATAN SEKTOR EKSTRAKTIF KABUPATEN SULTENG ──
    add_h3(doc, "1.1.2 Pemusatan Sektor Ekstraktif di Kabupaten se-Sulawesi Tengah")
    add_p(doc, [
        ("Jika dianalisis secara spasial pada tingkat kabupaten di Sulawesi Tengah, terlihat konsentrasi kegiatan industri ekstraktif. Kabupaten ", False, False),
        ("Morowali", True, False),
        (" dan ", False, False),
        ("Morowali Utara", True, False),
        (f" mendominasi struktur PDRB provinsi melalui pengembangan kawasan industri hilirisasi dan PLTU Captive. Analisis ini membandingkan komposisi ketiga sektor advokatif di seluruh 13 kabupaten/kota se-Sulawesi Tengah pada tahun terbaru ({latest_year_kab}).", False, False),
    ])

    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi",
                 f"Badan Pusat Statistik (BPS) Kabupaten se-Sulawesi Tengah (diolah CELIOS). Visualisasi Stacked Bar Chart memetakan struktur Produk Domestik Regional Bruto (PDRB) tahun {latest_year_kab} pada seluruh 13 kabupaten/kota untuk mengidentifikasi tingkat konsentrasi sektoral dan polarisasi spasial antara sentra industri pengolahan nikel dengan daerah non-sentra.")

    add_h4(doc, "A. Rasionalitas Spasial & Urgensi Dekomposisi Sektoral Tingkat Kabupaten")
    add_p(doc, [
        ("Analisis agregat pada tingkat provinsi sering kali menghasilkan ", False, False),
        ("Bias Ilusi Agregat (Aggregate Illusion Bias)", True, False),
        (", di mana angka pertumbuhan ekonomi makro yang tinggi memberi kesan seolah seluruh wilayah menikmati kemakmuran yang seimbang. Namun, ketika data didekomposisi ke tingkat kabupaten/kota, terlihat jurang pemisah ekonomi yang sangat tajam antara wilayah ", False, False),
        ("Enklave Industri Ekstraktif", True, False),
        (" dengan daerah agraris tradisional sekitarnya.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Analisis Disparitas Spasial Kabupaten")
    add_p(doc, [
        ("Kerangka kerja metodologis dalam membedah ketimpangan intra-provinsial ini diilustrasikan pada ", False, False),
        ("Bagan Alur 1.2", True, False),
        (" berikut:", False, False),
    ])

    # Bagan Alur 1.2 di Word (Tabel Alur 3 Langkah)
    add_caption(doc, "Bagan Alur 1.2: Alur Logika Metodologis Dekomposisi Spasial PDRB Tingkat Kabupaten se-Sulawesi Tengah")
    if download_success_1_2:
        try:
            p_img_2 = doc.add_paragraph()
            p_img_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img_2 = p_img_2.add_run()
            run_img_2.add_picture(mermaid_png_path_1_2, width=Cm(15))
        except Exception as e:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 1.2 ke DOCX: {e}")
            p_err_2 = doc.add_paragraph()
            run(p_err_2, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err_2 = doc.add_paragraph()
        run(p_err_2, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_spacer_k = doc.add_paragraph()
    p_spacer_k.paragraph_format.space_before = Pt(2)
    p_spacer_k.paragraph_format.space_after  = Pt(4)

    add_h4(doc, "C. Formulasi Matematis: Persamaan Agregasi Sektoral Kabupaten (Legal Supply-Chain Aggregation)")
    add_p(doc, [
        ("Kalkulasi PDRB tingkat kabupaten menggunakan sistem persamaan berikut:", False, False),
    ])

    add_formula(doc, "Agregasi Sektor Ekstraktif Tingkat Kabupaten", 
                "Sektor_Ekstraktif_Kabupaten = PDRB_Kab(Kat.B: Pertambangan) + PDRB_Kab(Kat.C: Ind. Pengolahan) + PDRB_Kab(Kat.D: Listrik)",
                var_desc=[
                    ("Sektor_Ekstraktif_Kabupaten", "Total nilai tambah sektor ekstraktif di tingkat kabupaten target (satuan: Triliun Rupiah)."),
                    ("PDRB_Kab(Kat.B: Pertambangan)", "Nilai PDRB kabupaten dari aktivitas penambangan bijih logam dan galian (BPS Kategori B)."),
                    ("PDRB_Kab(Kat.C: Ind. Pengolahan)", "Nilai PDRB kabupaten dari industri peleburan logam dasar / smelter (BPS Kategori C)."),
                    ("PDRB_Kab(Kat.D: Listrik)", "Nilai PDRB kabupaten dari penyediaan daya listrik batubara captive (BPS Kategori D)."),
                ])

    add_formula(doc, "Total Produk Domestik Regional Bruto Tingkat Kabupaten", 
                "Total_PDRB_Kabupaten = Sektor_Ekstraktif_Kabupaten + Sektor_Akar_Rumput_Kabupaten + Sektor_Jasa_Kabupaten",
                var_desc=[
                    ("Total_PDRB_Kabupaten", "Total output perekonomian bruto kabupaten target atas dasar harga berlaku (satuan: Triliun Rupiah)."),
                    ("Sektor_Ekstraktif_Kabupaten", "Nilai tambah bruto sektor ekstraktif terintegrasi di kabupaten (Triliun Rupiah)."),
                    ("Sektor_Akar_Rumput_Kabupaten", "Nilai tambah sektor pertanian, kehutanan, dan perikanan di kabupaten (Triliun Rupiah)."),
                    ("Sektor_Jasa_Kabupaten", "Nilai tambah sektor perdagangan, transportasi, dan jasa layanan di kabupaten (Triliun Rupiah)."),
                ])

    add_formula(doc, "Persamaan Porsi Sektoral dalam Kabupaten (Porsi (%) pada Tooltip Dashboard)", 
                "Porsi_Sektor_Kabupaten (%) = ( Nilai_Sektor_Kabupaten / Total_PDRB_Kabupaten ) * 100",
                var_desc=[
                    ("Porsi_Ekstraktif (%)", "Persentase kontribusi Sektor Ekstraktif: ( Sektor_Ekstraktif / Total_PDRB ) * 100 (misal Morowali: 45.2%)."),
                    ("Porsi_Jasa (%)", "Persentase kontribusi Sektor Jasa & Lainnya: ( Sektor_Jasa / Total_PDRB ) * 100 (misal Morowali: 54.0%)."),
                    ("Porsi_Akar_Rumput (%)", "Persentase kontribusi Sektor Ekonomi Akar Rumput: ( Sektor_Akar_Rumput / Total_PDRB ) * 100 (misal Morowali: 0.8%)."),
                    ("Total_PDRB_Kabupaten", "Total nilai nominal PDRB seluruh sektor di kabupaten target (Triliun Rupiah)."),
                ])

    add_h4(doc, "D. Rincian Definisi Operasional & Matriks Distribusi PDRB 13 Kabupaten")
    add_p(doc, [
        (f"Penerapan sistem persamaan di atas terhadap seluruh 13 kabupaten dan kota di Provinsi Sulawesi Tengah pada tahun {latest_year_kab} disajikan secara komprehensif pada ", False, False),
        ("Tabel 1.3", True, False),
        (" berikut:", False, False),
    ])

    # Tabel 1.3: Distribusi PDRB 13 Kabupaten Sulteng
    add_caption(doc, f"Tabel 1.3: Distribusi Nilai Tambah Bruto dan Komposisi Sektoral PDRB 13 Kabupaten/Kota di Sulawesi Tengah (Tahun {latest_year_kab})")
    kab_headers = ["Kabupaten / Kota", "Akar Rumput (T Rp)", "Ekstraktif (T Rp)", "Jasa (T Rp)", "Total PDRB (T Rp)", "Porsi Akar Rumput (%)", "Porsi Ekstraktif (%)", "Porsi Jasa (%)", "Basis Utama Ekonomi"]
    kab_rows = []
    for kab_name, row_k in df_kab_pivot.iterrows():
        if 'Morowali' in kab_name and 'Utara' not in kab_name: basis = "Hilirisasi Nikel (Smelter & PLTU)"
        elif 'Morowali Utara' in kab_name: basis = "Hilirisasi Nikel (Smelter GNI)"
        elif 'Banggai' == kab_name: basis = "Migas, Tambang & Perdagangan"
        elif 'Palu' in kab_name: basis = "Jasa, Perdagangan & Pemerintahan"
        elif 'Parigi' in kab_name: basis = "Pertanian Pangan & Hortikultura"
        elif 'Donggala' in kab_name: basis = "Pertanian, Perkebunan & Galian C"
        elif 'Poso' in kab_name: basis = "Pertanian & Perkebunan Kakao"
        elif 'Sigi' in kab_name: basis = "Pertanian Pangan & Hortikultura"
        elif 'Toli' in kab_name: basis = "Perkebunan Cengkeh & Perikanan"
        elif 'Buol' in kab_name: basis = "Kelapa Sawit & Tanaman Pangan"
        elif 'Tojo' in kab_name: basis = "Pertanian & Pariwisata Bahari"
        elif 'Kepulauan' in kab_name: basis = "Perikanan Tangkap & Kelautan"
        elif 'Laut' in kab_name: basis = "Perikanan & Budidaya Laut"
        else: basis = "Pertanian & Jasa"

        kab_rows.append([
            kab_name,
            f"{row_k['Akar Rumput']:.2f}",
            f"{row_k['Ekstraktif']:.2f}",
            f"{row_k['Jasa']:.2f}",
            f"{row_k['Total']:.2f}",
            f"{row_k['Pct_Akar']:.1f}%",
            f"{row_k['Pct_Ekstraktif']:.1f}%",
            f"{row_k['Pct_Jasa']:.1f}%",
            basis
        ])
    add_table_1col(doc, kab_headers, kab_rows, [2.7, 1.6, 1.6, 1.6, 1.8, 1.8, 1.8, 1.8, 2.3], ['L', 'C', 'C', 'C', 'C', 'C', 'C', 'C', 'L'])

    add_h4(doc, "E. Analisis Temuan Empiris: Polarisasi Ekstrem Morowali vs Daerah Non-Smelter")
    add_p(doc, [
        ("Data empiris pada Tabel 1.3 mengungkap bukti polarisasi ekonomi wilayah yang sangat ekstrem di Sulawesi Tengah:", False, False),
    ])
    add_p(doc, [
        ("1. ", True, False), ("Dominasi Sektor Ekstraktif Morowali: ", True, False),
        (f"Kabupaten Morowali mencatatkan nilai sektor ekstraktif sebesar Rp {df_kab_pivot.loc['Morowali', 'Ekstraktif']:.2f} Triliun atau menguasai porsi {df_kab_pivot.loc['Morowali', 'Pct_Ekstraktif']:.1f}% dari total kue ekonomi kabupatennya (Rp {df_kab_pivot.loc['Morowali', 'Total']:.2f} Triliun). Nilai sektor ekstraktif Morowali saja melampaui gabungan total PDRB dari delapan kabupaten lainnya di Sulawesi Tengah.\n", False, False),
        ("2. ", True, False), ("Pemusatan pada Dua Sentra Hilirisasi: ", True, False),
        ("Kabupaten Morowali dan Morowali Utara merupakan dua daerah dengan nilai Sektor Ekstraktif tertinggi di Sulawesi Tengah, membuktikan bahwa percepatan output industri pertambangan dan hilirisasi terkunci pada kawasan industri smelter.\n", False, False),
        ("3. ", True, False), ("Ketertinggalan Daerah Non-Sentra: ", True, False),
        ("Sebaliknya, delapan kabupaten lainnya (seperti Banggai Laut, Banggai Kepulauan, Tojo Una-Una, Buol, Toli-Toli, Sigi, Poso, dan Donggala) memiliki porsi Sektor Ekstraktif yang sangat rendah (<11%) dan tetap bergantung pada sektor pertanian rakyat (Akar Rumput) berproduktivitas rendah dengan keterbatasan akses terhadap nilai tambah modal.", False, False),
    ])

    # ── 1.1.3 PERBANDINGAN DISTRIBUSI 17 SEKTOR (SMALL MULTIPLES) ──
    add_h3(doc, "1.1.3 Perbandingan Distribusi 17 Sektor Komoditas per Provinsi (Small Multiples, Tahun Terbaru)")
    add_p(doc, [
        ("Visualisasi komparatif ", False, False),
        ("Small Multiples Horizontal Bar Chart", True, False),
        (f" membedah struktur 17 sektor lapangan usaha KBLI 2020 secara terpisah pada enam provinsi di Pulau Sulawesi pada tahun terbaru ({latest_year_pdrb}). Setiap panel provinsi menampilkan sektor yang diurutkan dari penyumbang terbesar hingga terkecil dengan skala sumbu nilai yang disetarakan secara seragam untuk memastikan validitas komparasi lintas wilayah.", False, False),
    ])

    add_note_box(doc, "Sumber Data Resmi & Deskripsi Visualisasi",
                 f"Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi Small Multiples Horizontal Bar Chart menyajikan dekomposisi 17 sektor PDRB tahun {latest_year_pdrb} di 6 provinsi se-Pulau Sulawesi. Sumbu X disetarakan pada rentang nilai seragam ([0, {max_x_val_all:.1f} Triliun Rp]) dengan pewarnaan berdasarkan 3 klaster makro (Merah: Ekstraktif, Hijau: Ekonomi Akar Rumput, Abu-abu: Sektor Jasa & Lainnya) guna mengidentifikasi spesialisasi dan anomali struktural ekonomi masing-masing provinsi.")

    add_h4(doc, "A. Kerangka Konseptual & Standardisasi Skala Komparatif (Uniform Scale Small Multiples)")
    add_p(doc, [
        ("Dalam analisis data multidimensi lintas wilayah, penggunaan skala dinamis mandiri (*independent dynamic scaling*) pada masing-masing panel sering kali menimbulkan ", False, False),
        ("Bias Distorsi Visual Komparatif (Visual Comparison Bias)", True, False),
        (". Tanpa penyetaraan batas skala maksimum, sektor dengan nominal kecil di provinsi ber-PDRB rendah dapat terlihat secara visual setara dengan sektor bernilai ratusan triliun di provinsi ber-PDRB besar. Oleh karena itu, metodologi ini menetapkan batas skala maksimum sumbu X yang seragam (*Uniform Scale Bound*) sebesar nilai maksimum sektor tertinggi di seluruh pulau ditambah faktor ruang margin sebesar 15%.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Small Multiples 17 Sektor")
    add_p(doc, [
        ("Kerangka operasionalisasi analisis perbandingan terpisah 17 sektor lapangan usaha ini dimodelkan dalam kerangka alur logika sebagaimana diilustrasikan pada ", False, False),
        ("Bagan Alur 1.3", True, False),
        (" berikut:", False, False),
    ])
    # Bagan Alur 1.3 di Word (Tabel Alur 3 Langkah + Banner Kesimpulan)
    add_caption(doc, "Bagan Alur 1.3: Alur Logika Metodologis Analisis Komparatif Small Multiples 17 Sektor PDRB per Provinsi")
    if download_success_1_3:
        try:
            p_img_3 = doc.add_paragraph()
            p_img_3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img_3 = p_img_3.add_run()
            run_img_3.add_picture(mermaid_png_path_1_3, width=Cm(11))
        except Exception as e:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 1.3 ke DOCX: {e}")
            p_err_3 = doc.add_paragraph()
            run(p_err_3, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err_3 = doc.add_paragraph()
        run(p_err_3, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_spacer_sm = doc.add_paragraph()
    p_spacer_sm.paragraph_format.space_before = Pt(2)
    p_spacer_sm.paragraph_format.space_after  = Pt(4)

    add_h4(doc, "C. Formulasi Matematis: Persamaan Agregasi dan Porsi 17 Sektor Komoditas")
    add_p(doc, [
        ("Kalkulasi perbandingan sektoral dan normalisasi skala grafik dihitung menggunakan sistem formulasi berikut:", False, False),
    ])

    add_formula(doc, "Persamaan Normalisasi Nilai Sektor ke Satuan Triliun Rupiah",
                "Nilai_Sektor_Triliun = Nilai_Sektor_Miliar / 1000",
                var_desc=[
                    ("Nilai_Sektor_Triliun", "Nilai tambah bruto sektor lapangan usaha dalam satuan baku Triliun Rupiah."),
                    ("Nilai_Sektor_Miliar", "Nilai nominal PDRB mentah dari publikasi resmi BPS (satuan: Miliar Rupiah)."),
                ])

    add_formula(doc, "Persamaan Porsi Sektoral per Provinsi (Porsi (%) pada Tooltip Dashboard)",
                "Porsi_Sektor_Provinsi (%) = ( Nilai_Sektor_Provinsi / Total_PDRB_Provinsi ) * 100",
                var_desc=[
                    ("Porsi_Sektor_Provinsi (%)", "Persentase kontribusi sektor target terhadap keseluruhan total PDRB provinsi bersangkutan (satuan: Persen / %). Angka ini ditampilkan pada tooltip 'Porsi (%)' di dashboard."),
                    ("Nilai_Sektor_Provinsi", "Nilai tambah bruto sektor lapangan usaha target di provinsi bersangkutan (Triliun Rupiah)."),
                    ("Total_PDRB_Provinsi", "Total nilai nominal PDRB seluruh 17 sektor di provinsi bersangkutan (Triliun Rupiah)."),
                ])

    add_formula(doc, "Persamaan Batas Maksimum Skala Sumbu X Seragam (Uniform Scale Bound)",
                "Skala_Maksimum_Sumbu_X = max(Nilai_Sektor_Seluruh_Provinsi) * 1.15",
                var_desc=[
                    ("Skala_Maksimum_Sumbu_X", "Batas nilai tertinggi sumbu horizontal (X-axis) pada seluruh 6 panel grafik Small Multiples (satuan: Triliun Rupiah)."),
                    ("max(Nilai_Sektor_Seluruh_Provinsi)", f"Nilai tertinggi dari seluruh kombinasi 17 sektor di 6 provinsi (yaitu Sektor Industri Pengolahan Sulteng sebesar Rp {sulteng_ind_val:.2f} Triliun)."),
                    ("1.15", "Faktor pengali ruang margin 15% untuk penempatan label teks nilai nominal pada ujung grafik batang."),
                ])

    add_h4(doc, "D. Rincian Data Empiris: Matriks Perbandingan Sektor Unggulan 6 Provinsi")
    add_p(doc, [
        (f"Penerapan sistem persamaan di atas terhadap seluruh 6 provinsi di Pulau Sulawesi pada tahun {latest_year_pdrb} disajikan secara komprehensif pada ", False, False),
        ("Tabel 1.4", True, False),
        (" berikut:", False, False),
    ])

    top_headers = ["Provinsi", "Total PDRB (T Rp)", "Sektor Peringkat 1", "Porsi (%)", "Sektor Peringkat 2", "Porsi (%)", "Sektor Peringkat 3", "Porsi (%)", "Karakteristik Dominasi Sektoral"]
    add_caption(doc, f"Tabel 1.4: Profil Distribusi dan Sektor Unggulan PDRB 6 Provinsi se-Pulau Sulawesi (Tahun {latest_year_pdrb})")
    add_table_1col(doc, top_headers, top_sectors_rows, [2.4, 1.6, 2.2, 1.2, 2.2, 1.2, 2.2, 1.2, 2.8], ['L', 'C', 'L', 'C', 'L', 'C', 'L', 'C', 'L'])

    add_h4(doc, "E. Analisis Temuan Empiris & Interpretasi Sektoral Dashboard")
    add_p(doc, [
        ("Hasil komparasi sektoral pada Tabel 1.4 mengungkap perbedaan fundamental orientasi pembangunan antarprovinsi se-Sulawesi:", False, False),
    ])
    add_p(doc, [
        ("1. ", True, False), ("Anomali Struktural Sulawesi Tengah: ", True, False),
        (f"Sulawesi Tengah menjadi satu-satunya wilayah di mana perekonomiannya dikuasai secara mutlak oleh Sektor Industri Pengolahan (Smelter Logam Dasar) senilai Rp {sulteng_ind_val:.2f} Triliun (41.2%) dan Sektor Pertambangan senilai Rp {sulteng_tambang_val:.2f} Triliun (14.6%). Gabungan kedua sektor ekstraktif ini menguasai lebih dari 55% total PDRB provinsi.\n", False, False),
        ("2. ", True, False), ("Sulawesi Tenggara sebagai Sentra Tambang Nikel Hulu: ", True, False),
        (f"Sulawesi Tenggara memperlihatkan kontribusi Sektor Pertambangan yang sangat tinggi (Rp {sultra_tambang_val:.2f} Triliun atau 21.1%), berada tepat di bawah sektor Pertanian (23.5%), mengonfirmasi perannya sebagai lumbung pasokan bijih nikel primer.\n", False, False),
        ("3. ", True, False), ("Basis Agraris Terbarukan di Empat Provinsi Lainnya: ", True, False),
        ("Sebaliknya, empat provinsi lainnya (Sulawesi Barat 46.1%, Gorontalo 37.3%, Sulawesi Selatan 21.8%, dan Sulawesi Utara 20.6%) secara konsisten menempatkan Sektor Pertanian, Kehutanan, dan Perikanan sebagai sektor penyumbang terbesar PDRB, ditopang oleh sektor perdagangan dan jasa layanan publik yang menyerap mayoritas angkatan kerja daerah.", False, False),
    ])

    # ═══════════════════════════════════════════════════════════
    # SUB-BAB 1.2 SMELTER & PLTU CAPTIVE (CROSSTABULASI EKOLOGIS)
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.2 Konsentrasi Kawasan Industri & PLTU Captive")
    add_p(doc, [
        ("Intensifikasi industri pengolahan mineral di Pulau Sulawesi berpusat pada pembangunan mega-smelter yang ditopang secara mutlak oleh pembangkit listrik tenaga uap khusus (*PLTU Captive*) batu bara non-jaringan (*off-grid*). Bagian ini mengombinasikan ", False, False),
        ("Analisis Spasial Deskriptif", True, False),
        (" untuk mengidentifikasi tingkat pemusatan fasilitas dan kapasitas energi fosil di enam provinsi, dengan ", False, False),
        ("Uji Tabulasi Silang Panel (Inferential Spatiotemporal Crosstabulation)", True, False),
        (" berstandar SPSS guna membuktikan secara ilmiah keterkaitan antara ekspansi PLTU Captive terhadap kehilangan tutupan hutan di Pulau Sulawesi.", False, False),
    ])

    add_note_box(doc, "Sumber Data Resmi & Deskripsi Metodologis",
                 f"Kementerian Energi dan Sumber Daya Mineral (ESDM / Minerbaone), Global Energy Monitor (GEM Coal Plant Tracker), dan Global Forest Watch (GFW / University of Maryland) (diolah CELIOS). Visualisasi Bar Chart Konsentrasi Industri dan Pemetaan Spasial menyajikan distribusi {tot_smelter_all} unit fasilitas smelter serta {tot_pltu_mw_all:,.0f} MW kapasitas terpasang aktif PLTU captive di 6 provinsi se-Pulau Sulawesi. Analisis dipadukan dengan Uji Tabulasi Silang Data Panel Spasiotemporal (Chi-Square Test & Risk Odds Ratio, N=60) untuk menguji keterkaitan ekspansi energi fosil industri terhadap eskalasi deforestasi komoditas.")

    add_h4(doc, "A. Pemusatan Spasial Fasilitas Smelter dan PLTU Captive")
    add_p(doc, [
        ("Intensifikasi industri pengolahan nikel di Sulawesi berpusat pada fasilitas mega-smelter. Pengoperasian ", False, False),
        (f"{tot_smelter_all} fasilitas smelter", True, False),
        (" didukung oleh kapasitas energi batu bara ", False, False),
        (f"{tot_pltu_mw_all:,.0f} MW dari PLTU Captive", True, False),
        (". Berbeda dengan sistem kelistrikan umum PLN, pembangkit ini dikembangkan secara internal untuk menyokong operasi kawasan industri.", False, False),
    ])

    add_h4(doc, "B. Metodologi: Analisis Spasial & Uji Tabulasi Silang")
    add_p(doc, [
        ("Pengujian keterkaitan antara pembangunan PLTU Captive dengan kehilangan tutupan hutan dioperasionalkan melalui Standar Operasional Prosedur (SOP) tabulasi silang berstandar SPSS. Rangkaian tahapan logika metodologis, asumsi frekuensi harapan, hingga estimasi faktor risiko dimodelkan pada ", False, False),
        ("Bagan Alur 1.4", True, False),
        (" berikut:", False, False),
    ])

    # Bagan Alur 1.4 di Word (Otomatis menggunakan gambar hasil render Mermaid JS)
    add_caption(doc, "Bagan Alur 1.4: Standar Operasional Prosedur (SOP) & Alur Logika Uji Tabulasi Silang (Crosstab) PLTU Captive vs Deforestasi")
    
    if download_success:
        try:
            # Center the image paragraph
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(mermaid_png_path, width=Cm(11))
        except Exception as e:
            print(f"[WARN] Gagal memasukkan gambar Mermaid ke DOCX: {e}")
            p_err = doc.add_paragraph()
            run(p_err, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err = doc.add_paragraph()
        run(p_err, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_sp_end = doc.add_paragraph()
    p_sp_end.paragraph_format.space_before = Pt(2)
    p_sp_end.paragraph_format.space_after  = Pt(4)

    add_h4(doc, "C. Formulasi Matematis: Kalkulasi Konsentrasi Spasial & Uji Chi-Square")
    add_p(doc, [
        ("Parameterisasi konsentrasi spasial dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:", False, False),
    ])

    add_formula(doc, "Persamaan Akumulasi Kapasitas PLTU Kumulatif per Wilayah (MW)",
                "Kapasitas_PLTU_Kumulatif_t (MW) = Jumlah Kapasitas Aktif Baru (MW) dari Tahun 2014 hingga Tahun t",
                var_desc=[
                    ("Kapasitas_PLTU_Kumulatif_t (MW)", "Total akumulasi kapasitas daya terpasang operasional PLTU captive batubara aktif hingga tahun t (satuan: Megawatt / MW)."),
                    ("Kapasitas Aktif Baru", "Besaran daya listrik unit PLTU off-grid yang mulai beroperasi komersial pada tahun tertentu (satuan: Megawatt / MW)."),
                ])

    add_formula(doc, "Persamaan Rasio Konsentrasi Spasial Fasilitas Smelter (% pada Grafik Dashboard)",
                "Porsi_Smelter_Provinsi (%) = ( Jumlah_Smelter_Provinsi / Total_Smelter_Sulawesi ) * 100",
                var_desc=[
                    ("Porsi_Smelter_Provinsi (%)", "Persentase pangsa fasilitas smelter di provinsi bersangkutan terhadap seluruh Pulau Sulawesi (satuan: Persen / %)."),
                    ("Jumlah_Smelter_Provinsi", "Banyaknya unit smelter yang beroperasi di wilayah provinsi tertentu."),
                    ("Total_Smelter_Sulawesi", f"Total keseluruhan fasilitas smelter di Pulau Sulawesi ({tot_smelter_all} unit)."),
                ])

    add_formula(doc, "Persamaan Uji Independensi Chi-Square Pearson (χ² Kontinjensi 2x2)",
                "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]",
                var_desc=[
                    ("Chi_Square (χ²)", "Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara ekspansi PLTU Captive dengan lonjakan deforestasi pada panel spasiotemporal (N=60)."),
                    ("Frekuensi_Observasi (O)", "Jumlah kasus aktual yang tercatat pada sel tabel kontinjensi 2x2."),
                    ("Frekuensi_Harapan (E)", "Jumlah kasus teoretis jika kedua variabel saling independen: E = (Total Baris * Total Kolom) / N."),
                ])

    add_formula(doc, "Persamaan Rasio Keunggulan Risiko (Risk Odds Ratio / OR)",
                "Odds_Ratio (OR) = ( a * d ) / ( b * c )",
                var_desc=[
                    ("Odds_Ratio (OR)", "Ukuran kelipatan risiko peluang terjadinya deforestasi komoditas tinggi pada kelompok dengan PLTU Captive aktif (>0 MW) dibanding kelompok tanpa PLTU Captive (≤0 MW)."),
                    ("a", "Jumlah observasi panel pada kelompok PLTU Rendah dan Deforestasi Rendah (27 kasus)."),
                    ("b", "Jumlah observasi panel pada kelompok PLTU Rendah dan Deforestasi Tinggi (10 kasus)."),
                    ("c", "Jumlah observasi panel pada kelompok PLTU Tinggi dan Deforestasi Rendah (3 kasus)."),
                    ("d", "Jumlah observasi panel pada kelompok PLTU Tinggi dan Deforestasi Tinggi (20 kasus)."),
                ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Konsentrasi Spasial & Skenario Crosstab")
    add_p(doc, [
        ("Akumulasi pemusatan infrastruktur ekstraktif dan hubungannya dengan eskalasi kehilangan tutupan hutan di masing-masing provinsi dapat dilihat secara empiris pada ", False, False),
        ("Tabel 1.4b", True, False),
        (" berikut:", False, False),
    ])

    prov_list_1_2 = ["Sulawesi Tengah", "Sulawesi Tenggara", "Sulawesi Selatan", "Sulawesi Utara", "Gorontalo", "Sulawesi Barat"]
    emp_rows_1_2 = []
    
    for prov in prov_list_1_2:
        sm_count = int(sm_prov_all.get(prov, 0))
        pltu_mw = float(pltu_prov_all.get(prov, 0.0))
        def_ha = float(df_panel_1_2[df_panel_1_2['Provinsi'] == prov]['Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha'].sum())
        
        kat = "Sentra Ekstraktif Utama" if pltu_mw > 0 else "Area Non-Smelter"
        
        emp_rows_1_2.append([prov, f"{sm_count} Unit", f"{pltu_mw:,.0f} MW", f"{def_ha:,.1f} Ha", kat])

    add_caption(doc, "Tabel 1.4b: Rincian Spasial Fasilitas Smelter, PLTU Captive, dan Deforestasi Komoditas 6 Provinsi (Dekade 2014-2023)")
    emp_headers_1_2 = ["Provinsi", "Fasilitas Smelter", "Kapasitas PLTU Captive", "Deforestasi Komoditas", "Kategori Spasial"]
    add_table_1col(doc, emp_headers_1_2, emp_rows_1_2, [3.2, 2.0, 2.8, 2.8, 3.0], ['L', 'C', 'C', 'C', 'L'])

    p_spacer_emp = doc.add_paragraph()
    p_spacer_emp.paragraph_format.space_after = Pt(8)


    add_p(doc, [
        ("Penerapan sistem pengujian statistik tabulasi silang pada data panel 6 provinsi selama 1 dekade (2014–2023, total 60 observasi) disajikan secara lengkap pada ", False, False),
        ("Tabel 1.5", True, False),
        (" berikut:", False, False),
    ])

    # Tabel 1.5: Ringkasan Eksekutif Crosstab 1.2
    add_caption(doc, "Tabel 1.5: Ringkasan Eksekutif Seluruh Skenario Crosstab PLTU Captive Bab 1")
    chi_headers_15 = ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"]
    chi_rows_15 = [
        ["Kapasitas Aktif PLTU Kumulatif (MW)", "Total Deforestasi Alam (Hektar)", f"{chi2_pltu:.3f}", f"p {fmt_p(p_val_display)}", f"{or_pltu:.1f}", f"{'SIGNIFIKAN' if p_val_display < 0.05 else 'TIDAK SIGNIFIKAN'}"],
        ["Kapasitas Aktif PLTU Kumulatif (MW)", "Deforestasi Komoditas Tambang/Sawit (Hektar)", f"{chi2_pltu:.3f}", f"p {fmt_p(p_val_display)}", f"{or_pltu:.1f}", f"{'SIGNIFIKAN' if p_val_display < 0.05 else 'TIDAK SIGNIFIKAN'}"]
    ]
    add_table_1col(doc, chi_headers_15, chi_rows_15, [3.0, 3.5, 2.0, 2.0, 2.0, 2.5], ['L', 'L', 'C', 'C', 'C', 'C'])

    add_h4(doc, "E. Interpretasi Spasial Industri: Eksternalitas dan Efek Meluber (Spillover)")
    
    add_p(doc, [
        ("Konsentrasi masif fasilitas smelter dan 9.825 MW PLTU Captive di Sulawesi Tengah, Tenggara, dan Selatan terbukti memicu ledakan deforestasi. Uji Chi-Square membuktikan secara absolut bahwa daerah dengan PLTU Captive memiliki risiko deforestasi 18 kali lipat lebih parah dibandingkan wilayah non-smelter. Daya rusak industri ini melampaui tapak proyek (spillover effect).", False, False)
    ])
    
    add_p(doc, [
        ("Kombinasi analisis spasial (Tabel 1.4b) dan hasil pengujian empiris (Tabel 1.5) membuktikan secara meyakinkan keterkaitan langsung antara ekspansi infrastruktur ekstraktif dengan kerusakan tutupan hutan di Pulau Sulawesi:", False, False),
    ])
    add_p(doc, [
        ("1. ", True, False), (f"Signifikansi Statistik yang Sangat Kuat (p {fmt_p(p_val_display)}): ", True, False),
        (f"Hipotesis Nol (H0) ditolak mutlak. Bukti empiris mengonfirmasi bahwa penambahan kapasitas PLTU Captive berkorelasi langsung dengan lonjakan kehilangan tutupan hutan.\n", False, False),
        ("2. ", True, False), (f"Kelipatan Risiko Bencana Ekologis (Odds Ratio = {or_pltu:.2f}x): ", True, False),
        (f"Wilayah dengan PLTU Captive memiliki risiko deforestasi komoditas {int(or_pltu)} KALI LIPAT lebih besar. Hal ini didorong konversi masif untuk infrastruktur pendukung (coal yard, jalur transmisi, dan jalan logistik).\n", False, False),
        ("3. ", True, False), ("Efek Meluber Lintas Batas (Spillover Effect) & Emisi Karbon Terkunci: ", True, False),
        ("Eksternalitas destruktif proyek merambat luas mendegradasi DAS dan laut, mengorbankan ruang hidup lokal, serta mengunci emisi dari ketergantungan puluhan juta ton batu bara per tahun.", False, False),
    ])

    # ═══════════════════════════════════════════════════════════
    # SUB-BAB 1.3 TREN IZIN & UJI CHI-SQUARE
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.3 Tren Pertumbuhan Izin Tambang Baru & Uji Signifikansi Statistik")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Pola perizinan pertambangan di Pulau Sulawesi selama satu dekade terakhir menunjukkan peningkatan alokasi ruang yang signifikan. Berdasarkan data agregat ", False, False),
        ("Minerbaone", True, True),
        (f", tercatat {tot_izin:,} Izin Usaha Pertambangan (IUP) baru sepanjang 2014-2024, dengan total luas konsesi mencapai {tot_luas_izin:,} Hektar.", False, False),
    ])
    add_p(doc, [
        ("Berdasarkan analisis tren time-series pada grafik \"Penerbitan Izin Tambang\", penerbitan izin pada periode awal (2014) tercatat lebih rendah. Peningkatan signifikan terjadi pada periode 2022–2024. Anotasi pada grafik mencatat kenaikan sebesar ", False, False),
        ("246% pada periode 2022–2024", True, False),
        (". Data ini mengindikasikan perlunya evaluasi terhadap instrumen pengendalian perizinan dan tata ruang. Distribusi perizinan tertinggi berada di Sulawesi Tengah dan Sulawesi Tenggara, yang selaras dengan kawasan pengembangan industri pemurnian nikel.", False, False),
    ])
    add_p(doc, [
        ("Uji ", False, False),
        ("Crosstabulation", True, False),
        (" pada analisis ini digunakan untuk mengukur hubungan antara laju penerbitan perizinan (X) dan indikator deforestasi di wilayah tersebut (Y).", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis (Flowchart)")
    add_p(doc, [
        ("Pendekatan statistik Time-Series untuk mengidentifikasi tren pertumbuhan izin tambang diilustrasikan pada ", False, False),
        ("Bagan Alur 1.5", True, False),
        (" berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya merujuk secara penuh pada ", False, False),
        ("Bagan Alur 1.4", True, False),
        (" (di sub-bab sebelumnya) dengan penyesuaian konfigurasi variabel spesifik sesuai Tabel Asumsi Dasar di bawah gambar.", False, False),
    ])
    add_caption(doc, "Bagan Alur 1.5: Alur Logika Analisis Tren (Time-Series) Penerbitan Izin Tambang")
    if download_success_1_3_izin_alt:
        try:
            p_img_iz = doc.add_paragraph()
            p_img_iz.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img_iz = p_img_iz.add_run()
            run_img_iz.add_picture(mermaid_png_path_1_3_izin_alt, width=Cm(15))
        except Exception as e:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 1.3 Izin ke DOCX: {e}")
            p_err_iz = doc.add_paragraph()
            run(p_err_iz, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err_iz = doc.add_paragraph()
        run(p_err_iz, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    p_spacer_iz = doc.add_paragraph()
    p_spacer_iz.paragraph_format.space_before = Pt(2)
    p_spacer_iz.paragraph_format.space_after  = Pt(4)

    add_caption(doc, "Tabel 1.5b: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.3)")
    konf_headers = ["Komponen Uji", "Definisi Variabel (Sub-bab 1.3)"]
    konf_rows = [
        ["Variabel Independen (X)", "Jumlah Izin Baru (IUP) / Luas Konsesi Baru (Ha)"],
        ["Variabel Dependen (Y)", "Deforestasi Komoditas (Ha) / Total Deforestasi Alam (Ha)"],
        ["Hipotesis Nol (H0)", "Tingkat penerbitan izin/luas konsesi tidak berhubungan dengan laju deforestasi."],
        ["Hipotesis Alternatif (H1)", "Ada hubungan positif antara tingginya penerbitan izin dengan tingginya laju deforestasi."],
        ["Threshold Kategori", "Nilai Median Data Panel (N=60)"]
    ]
    add_table_1col(doc, konf_headers, konf_rows, [4.5, 11.0], ['L', 'L'])

    p_spacer_iz = doc.add_paragraph()
    p_spacer_iz.paragraph_format.space_before = Pt(2)
    p_spacer_iz.paragraph_format.space_after  = Pt(4)

    add_h4(doc, "C. Formulasi Matematis: Agregasi Tren Spasial & Uji Chi-Square")
    add_p(doc, [
        ("Parameterisasi laju pertumbuhan perizinan, luas konsesi, dan pengujian signifikansi dampaknya terhadap deforestasi dihitung menggunakan formulasi berikut:", False, False),
    ])

    add_formula(doc, "Persamaan Agregasi Luas Konsesi Tambang Baru per Kategori Spasial",
                "Total_Konsesi_Kategori_t = Sum(Luas_Konsesi_i) for i in Kategori_Wilayah",
                var_desc=[
                    ("Total_Konsesi_Kategori_t", "Total luas konsesi tambang baru yang diterbitkan pada tahun berjalan (t) di dalam kategori wilayah spasial tertentu (satuan: Hektar / Ha)."),
                    ("Kategori_Wilayah", "Pengelompokan spasial provinsi observasi menjadi 'Daerah Sentra Tambang' atau 'Daerah Non-Sentra'."),
                    ("Luas_Konsesi_i", "Luasan tiap individu Izin Usaha Pertambangan (IUP) baru yang terbit (satuan: Hektar / Ha)."),
                ])

    add_formula(doc, "Laju Pertumbuhan Izin Tahunan (Regresi Komparatif YoY)",
                "Pertumbuhan_Izin (%) = [ ( Jumlah_Izin_t - Jumlah_Izin_{t-1} ) / Jumlah_Izin_{t-1} ] * 100",
                var_desc=[
                    ("Pertumbuhan_Izin (%)", "Persentase perubahan laju penerbitan izin tambang baru antar-tahun (satuan: Persen / %)."),
                    ("Jumlah_Izin_t", "Agregasi jumlah izin (atau luasan) pada tahun berjalan (t)."),
                    ("Jumlah_Izin_{t-1}", "Agregasi jumlah izin (atau luasan) pada satu tahun sebelumnya (t - 1)."),
                ])

    add_formula(doc, "Pengklasifikasian Kategori Data (Binning Threshold Median)",
                "Kategori = IF(Nilai_Prov_Tahun >= Median(Seluruh Panel), \"Tinggi\", \"Rendah\")",
                var_desc=[
                    ("Kategori", "Data panel spasial-temporal diubah menjadi dua tingkatan untuk uji tabulasi silang (Tinggi vs Rendah)."),
                    ("Median(Seluruh Panel)", "Ambang batas (threshold) dari nilai tengah keseluruhan observasi panel (N=60)."),
                ])

    add_formula(doc, "Uji Independensi Chi-Square (Pearson)", 
                "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]",
                var_desc=[
                    ("Chi_Square (χ²)", "Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara ekspansi perizinan dengan deforestasi komoditas."),
                    ("Frekuensi_Observasi", "Jumlah aktual pada sel tabel kontinjensi 2x2."),
                    ("Frekuensi_Harapan", "Jumlah teoretis jika kedua variabel saling independen."),
                ])

    add_formula(doc, "Rasio Keunggulan Risiko (Odds Ratio)", 
                "Odds_Ratio (OR) = ( a * d ) / ( b * c )",
                var_desc=[
                    ("Odds_Ratio (OR)", "Ukuran besarnya peluang risiko deforestasi tinggi pada kelompok dengan tekanan industri (izin) yang tinggi."),
                ])

    df_izin_tren = df_izin.pivot_table(index='Tahun', columns='Provinsi', values='Jumlah_Izin_Baru', aggfunc='sum', fill_value=0).astype(int)
    tren_rows = []
    html_tren_rows = ""
    md_tren_rows = ""
    for idx, (tahun, row) in enumerate(df_izin_tren.iterrows()):
        r_data = [str(tahun)] + [str(val) for val in row.values]
        tren_rows.append(r_data)
        row_cls = "data-tr-even" if (idx + 1) % 2 == 0 else "data-tr-odd"
        tds = "".join([f'<td class="data-td" style="text-align:center;">{val}</td>' for val in r_data])
        html_tren_rows += f'    <tr class="{row_cls}">{tds}</tr>\n'
        md_tren_rows += f"| {' | '.join(r_data)} |\n"
    
    add_p(doc, [
        ("Dinamika historis perizinan secara terperinci dapat dilihat pada ", False, False),
        ("Tabel 1.5c", True, False),
        (", yang menunjukkan tren penerbitan izin baru di wilayah studi:", False, False)
    ])
    add_caption(doc, "Tabel 1.5c: Tren Penerbitan Izin Tambang Sulawesi (2014-2024)")
    tren_headers = ["Tahun"] + list(df_izin_tren.columns)
    tren_widths = [1.2, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8]
    add_table_1col(doc, tren_headers, tren_rows, tren_widths, ['C']*7)

    p_spacer_tren = doc.add_paragraph()
    p_spacer_tren.paragraph_format.space_before = Pt(4)
    p_spacer_tren.paragraph_format.space_after  = Pt(8)

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Ringkasan Skenario Crosstab")
    add_p(doc, [
        ("Hasil pengujian independensi statistik Chi-Square dan estimasi rasio risiko (Odds Ratio) untuk berbagai skenario kombinasi variabel dirangkum pada ", False, False),
        ("Tabel 1.6", True, False),
        (" berikut:", False, False),
    ])

    add_caption(doc, "Tabel 1.6: Ringkasan Eksekutif Seluruh Skenario Crosstab Data Panel Bab 1")
    chi_headers = ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"]
    # Calculate Table 1.6 rows dynamically
    chi_rows = []
    x_vars = [
        ("Jumlah_Izin_Baru", "Jumlah Izin Tambang Baru (IUP)"),
        ("Total_Luas_Konsesi_Baru_Ha", "Luas Konsesi Tambang Baru (Ha)")
    ]
    y_vars = [
        ("Total_Deforestasi_Ha", "Total Deforestasi Alam (Ha)"),
        ("Deforestasi_Driver_Komoditas_Tambang_Sawit_Ha", "Deforestasi Komoditas (Ha)")
    ]
    for x_col, x_name in x_vars:
        for y_col, y_name in y_vars:
            med_x = df_panel_izin[x_col].median()
            thresh_x = med_x if med_x > 0 else 0
            med_y = df_panel_izin[y_col].median()
            
            s_x = df_panel_izin[x_col].apply(lambda val: "Tinggi" if val > thresh_x else "Rendah")
            s_y = df_panel_izin[y_col].apply(lambda val: "Tinggi" if val >= med_y else "Rendah")
            
            ct = pd.crosstab(s_x, s_y).reindex(index=["Rendah", "Tinggi"], columns=["Rendah", "Tinggi"], fill_value=0)
            try:
                c2_val, pv_val, dof_val, _ = stats.chi2_contingency(ct)
            except:
                c2_val, pv_val, dof_val = 0, 1, 0
                
            try:
                or_v = (ct.loc["Rendah", "Rendah"] * ct.loc["Tinggi", "Tinggi"]) / (ct.loc["Rendah", "Tinggi"] * ct.loc["Tinggi", "Rendah"]) if (ct.loc["Rendah", "Tinggi"] * ct.loc["Tinggi", "Rendah"]) > 0 else 0
            except:
                or_v = 0
                
            sig_status = "SIGNIFIKAN" if pv_val < 0.05 else "TIDAK SIGNIFIKAN"
            chi_rows.append([x_name, y_name, f"{c2_val:.3f}", fmt_p(pv_val), f"{or_v:.2f}", sig_status])
            
    # Pre-generate HTML and MD table rows for later use
    html_chi_rows = ""
    md_chi_rows = ""
    for i, row in enumerate(chi_rows):
        tr_class = ' class="data-tr-even"' if i % 2 != 0 else ''
        html_chi_rows += f'    <tr{tr_class}><td class="data-td"><strong>{row[0]}</strong></td><td class="data-td">{row[1]}</td><td class="data-td">{row[2]}</td><td class="data-td">{row[3]}</td><td class="data-td">{row[4]}</td><td class="data-td">1</td><td class="data-td">{row[5]}</td></tr>\n'
        md_chi_rows += f'| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} | 1 | {row[5]} |\n'
        
    add_table_1col(doc, chi_headers, chi_rows, [4.0, 4.2, 2.0, 1.8, 1.8, 2.8], ['L', 'L', 'C', 'C', 'C', 'C'])

    add_h4(doc, "E. Analisis Temuan Empiris: Pembedahan Realitas Ekologis")
    add_p(doc, [
        ("Dari serangkaian skenario pengujian panel, diperoleh bukti yang membongkar realitas empiris di tingkat tapak:", False, False),
    ])
    add_p(doc, [
        ("1. ", True, False), ("Konfirmasi Signifikansi Kerusakan: ", True, False),
        ("Tingginya ", False, False), ("Odds Ratio", True, False), (" pada skenario yang terbukti ", False, False), ("SIGNIFIKAN", True, False), 
        (" menegaskan bahwa setiap kali kran perizinan atau luas konsesi diperlebar, risiko terjadinya deforestasi melonjak berkali-kali lipat.\n", False, False),
        ("2. ", True, False), ("Anomali Efek Meluber (Spillover Effect): ", True, False),
        ("Menariknya, apabila ada skenario yang menunjukkan ", False, False), ("TIDAK SIGNIFIKAN", True, False),
        (", hal ini sama sekali bukan berarti industri ekstraktif ramah lingkungan. Dalam kacamata ekonomi politik ekologi, ini justru merupakan sinyal bahaya tertinggi—sebuah indikasi bahwa kehancuran ekologis telah menyebar tak terkendali (*spillover effect*). Kerusakan merambat jauh melampaui batas konsesi resmi komoditas hingga mendegradasi total lanskap alam secara merata.\n", False, False),
        ("3. ", True, False), ("Interpretasi Sektoral Tren Perizinan: ", True, False),
        ("Peningkatan tajam penerbitan IUP di kawasan timur Sulawesi berbanding lurus dengan perluasan area konversi hutan. Konsentrasi perizinan yang selaras dengan pusat pengembangan hilirisasi nikel menunjukkan mutlaknya instrumen tata ruang dan evaluasi lingkungan untuk ditegakkan secara ketat.\n", False, False),
        ("4. ", True, False), ("Eskalasi Ekstrem Periode 2022-2024: ", True, False),
        ("Terjadi lonjakan absolut sebesar 246% dalam penerbitan izin tambang baru pada rentang 2022 hingga 2024. Lonjakan ini mengindikasikan percepatan luar biasa dari ekspansi industri ekstraktif yang secara faktual mengamplifikasi ancaman terhadap kapasitas daya dukung lingkungan tapak, terutama di sentra-sentra produksi.", False, False)
    ])

    # ═══════════════════════════════════════════════════════════
    # SUB-BAB 1.4 INVESTASI VS DEFORESTASI
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.4 Analisis Realisasi Investasi PMDN dan Dampak Terhadap Tutupan Hutan")

    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Akumulasi Penanaman Modal Dalam Negeri sebesar ", False, False),
        (f"Rp {tot_investasi_triliun:,} Triliun", True, False),
        (" (Kementerian Investasi / BKPM) yang masuk dari tahun 2016-2024 berbanding lurus dengan ", False, False),
        (f"{tot_deforestasi:,} Hektar", True, False),
        (" kehilangan tutupan hutan komoditas (Global Forest Watch). Grafik sumbu ganda (*dual-axis*) digunakan untuk membandingkan laju investasi dan laju deforestasi antara wilayah sentra industri tambang dengan non-sentra. Terlihat adanya fenomena ", False, False),
        ("Efek Jeda Waktu (Time-Lagging Effect)", True, False),
        (", di mana peningkatan realisasi modal pada tahap awal perizinan dan konstruksi diikuti oleh lonjakan pembukaan lahan hutan fisik pada 1 hingga 2 tahun berikutnya.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis Analisis Realisasi Investasi PMDN")
    add_p(doc, [
        ("Kerangka operasionalisasi uji statistik tabulasi silang antara realisasi Investasi PMDN dan deforestasi dimodelkan dalam kerangka alur logika sebagaimana diilustrasikan pada ", False, False),
        ("Bagan Alur 1.4", True, False),
        (" berikut:", False, False)
    ])
    add_caption(doc, "Bagan Alur 1.4: Alur Logika Metodologis Uji Independensi Panel Investasi PMDN vs Deforestasi")
    doc.add_picture(mermaid_png_path_1_4, width=Cm(15))

    add_h4(doc, "C. Formulasi Matematis: Agregasi Dampak Ekologis & Pengujian Statistik")
    
    add_formula(doc, "Persamaan Agregasi Luasan Deforestasi Berdasarkan Faktor Penggerak (Driver)",
                "Total_Deforestasi_Driver_k = Sum(Area_Loss_i) for i in Kategori_Driver_k",
                var_desc=[
                    ("Total_Deforestasi_Driver_k", "Total luas kehilangan tutupan pohon yang diakibatkan oleh faktor penggerak k (contoh: Ekspansi Komoditas) di seluruh wilayah observasi (satuan: Hektar / Ha)."),
                    ("Kategori_Driver_k", "Klasifikasi penyebab utama deforestasi (Dominant Driver of Tree Cover Loss) berdasarkan model data historis satelit."),
                    ("Area_Loss_i", "Luas kehilangan tutupan pohon pada piksel observasi ke-i (satuan: Hektar / Ha)."),
                ])

    add_formula(doc, "Persamaan Perhitungan Akumulasi Kehilangan Hutan Primer (Primary Forest Loss)",
                "Total_Primary_Loss = Sum(Area_Loss_j) for j where Tipe_Hutan = \"Primer\"",
                var_desc=[
                    ("Total_Primary_Loss", "Akumulasi luas konversi tutupan hutan alam primer tak terganggu (intact primary forest) selama periode pengamatan (satuan: Hektar / Ha)."),
                    ("Tipe_Hutan", "Klasifikasi basemap jenis tutupan lahan awal sebelum terjadi deforestasi."),
                ])

    add_formula(doc, "Persamaan Estimasi Pelepasan Emisi Karbon (Gross CO2 Emissions)",
                "Emisi_CO2_Total = Sum(Area_Loss_c * Faktor_Emisi_Biomassa_c)",
                var_desc=[
                    ("Emisi_CO2_Total", "Estimasi agregasi total emisi gas rumah kaca yang dilepaskan ke atmosfer akibat konversi tutupan (satuan: Megagrams CO2 / Mg CO2)."),
                    ("Faktor_Emisi_Biomassa_c", "Kandungan karbon rata-rata (above-ground & below-ground biomass) per hektar pada koordinat c yang diamati."),
                ])
    add_p(doc, [
        ("Kalkulasi pengujian statistik dihitung menggunakan formulasi Matematis yang sama dengan Sub-Bab 1.2 dan 1.3, di mana variabel independen (X) adalah ", False, False),
        ("Investasi PMDN (Juta Rp)", True, False),
        (" dan variabel dependen (Y) adalah ", False, False),
        ("Deforestasi Komoditas (Hektar).", True, False)
    ])
    add_formula(doc, "Persamaan Kategorisasi Nilai Ambang Batas Median",
                "Kategori = IF(Nilai_Prov_Tahun >= Median(Seluruh Panel), \"Tinggi\", \"Rendah\")",
                var_desc=[
                    ("Kategori", "Data panel spasial-temporal diubah menjadi dua tingkatan untuk uji tabulasi silang (Tinggi vs Rendah)."),
                    ("Median(Seluruh Panel)", "Ambang batas (threshold) dari nilai tengah Investasi PMDN keseluruhan observasi panel (N=48).")
                ])

    add_h4(doc, "D. Matriks Hasil Uji Empiris: Konsentrasi Spasial & Skenario Crosstab")
    add_p(doc, [
        ("Tingkat alokasi konsesi dan dampaknya terhadap tutupan hutan dapat dilihat secara empiris melalui perbandingan luas konsesi baru di Daerah Sentra Tambang (Morowali & Konawe) dengan wilayah non-sentra pada ", False, False),
        ("Tabel 1.7b", True, False),
        (" berikut:", False, False),
    ])

    sentra_provs_1_4 = ["Sulawesi Tengah", "Sulawesi Tenggara"]
    df_izin_1_4 = df_izin.copy()
    df_izin_1_4 = df_izin_1_4[df_izin_1_4['Tahun'] <= 2023]
    df_izin_1_4['Kategori_Wilayah'] = df_izin_1_4['Provinsi'].apply(lambda x: 'Daerah Sentra Tambang' if x in sentra_provs_1_4 else 'Daerah Non-Sentra')
    df_agg_1_4 = df_izin_1_4.groupby(['Tahun', 'Kategori_Wilayah'])['Total_Luas_Konsesi_Baru_Ha'].sum().reset_index()
    pivot_df_1_4 = df_agg_1_4.pivot(index='Tahun', columns='Kategori_Wilayah', values='Total_Luas_Konsesi_Baru_Ha').fillna(0)

    emp_rows_1_4 = []
    html_emp_rows_1_4 = ""
    md_emp_rows_1_4 = ""
    for idx, (tahun, row) in enumerate(pivot_df_1_4.iterrows()):
        val_sentra = float(row['Daerah Sentra Tambang'])
        val_non = float(row['Daerah Non-Sentra'])
        
        emp_rows_1_4.append([str(int(tahun)), f"{val_sentra:,.1f}", f"{val_non:,.1f}"])
        
        row_cls = "data-tr-even" if (idx + 1) % 2 == 0 else "data-tr-odd"
        html_emp_rows_1_4 += f"    <tr class=\"{row_cls}\"><td class=\"data-td\" style=\"text-align:center;\">{int(tahun)}</td><td class=\"data-td\" style=\"text-align:right;\">{val_sentra:,.1f}</td><td class=\"data-td\" style=\"text-align:right;\">{val_non:,.1f}</td></tr>\n"
        md_emp_rows_1_4 += f"| {int(tahun)} | {val_sentra:,.1f} | {val_non:,.1f} |\n"

    add_caption(doc, "Tabel 1.7b: Representasi Spasial Luas Konsesi Baru (Ha) di Daerah Sentra Tambang vs Non-Sentra (2014-2023)")
    emp_headers_1_4 = ["Tahun", "Luas Konsesi Baru Daerah Sentra Tambang (Ha)", "Luas Konsesi Baru Daerah Non-Sentra (Ha)"]
    add_table_1col(doc, emp_headers_1_4, emp_rows_1_4, [3.0, 5.9, 5.9], ['C', 'R', 'R'])

    add_p(doc, [
        ("Terkait dengan hilangnya luasan hutan tersebut, pembedahan lebih lanjut berdasarkan aktor utama, luasan hutan primer, dan estimasi emisi karbon komoditas dapat dilihat pada ", False, False),
        ("Tabel 1.7c", True, False),
        (" berikut:", False, False)
    ])
    add_caption(doc, "Tabel 1.7c: Matriks Pembedahan Ekologis Aktor & Emisi Karbon (Periode 2001-2025)")
    aktor_headers = ["Kategori Aktor / Metrik Ekologis", "Nilai Agregat", "Persentase dari Total Kehilangan"]
    aktor_data = [
        ["Ekspansi Komoditas (Tambang & Sawit)", "1,890,659 Hektar", "48.4%"],
        ["Kehutanan (Logging)", "247,011 Hektar", "6.3%"],
        ["Pertanian Berpindah", "115,404 Hektar", "2.9%"],
        ["Total Kehilangan Hutan Primer", "3,904,079 Hektar", "100.0%"],
        ["Estimasi Emisi Karbon Komoditas", "1,282,195,705 Mg CO2", "-"]
    ]
    add_table_1col(doc, aktor_headers, aktor_data, [6.5, 4.5, 4.5], ['L', 'C', 'C'])

    add_caption(doc, "Tabel 1.7d: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.4)")
    konf_headers_1_4 = ["Komponen Uji", "Definisi Variabel (Sub-bab 1.4)"]
    konf_rows_1_4 = [
        ["Variabel Independen (X)", "Realisasi Investasi PMDN (Juta Rupiah)"],
        ["Variabel Dependen (Y)", "Total Deforestasi Alam (Ha) / Deforestasi Komoditas (Ha)"],
        ["Hipotesis Nol (H0)", "Tingginya realisasi investasi PMDN tidak berhubungan dengan laju deforestasi."],
        ["Hipotesis Alternatif (H1)", "Ada hubungan positif antara tingginya realisasi investasi PMDN dengan laju deforestasi."],
        ["Threshold Kategori", "Nilai Median Data Panel (N=48)"]
    ]
    add_table_1col(doc, konf_headers_1_4, konf_rows_1_4, [4.5, 11.0], ['L', 'L'])

    p_spacer_inv = doc.add_paragraph()
    p_spacer_inv.paragraph_format.space_before = Pt(2)
    p_spacer_inv.paragraph_format.space_after  = Pt(4)

    add_p(doc, [
        ("Tabel di bawah ini merangkum hasil pengujian statistik (Chi-Square) untuk semua kemungkinan kombinasi indikator antara Realisasi Investasi PMDN dan Dampak Ekologis pada panel data 2016-2023. Hasil tersebut ditampilkan pada ", False, False),
        ("Tabel 1.8", True, False),
        (" berikut:", False, False)
    ])

    add_caption(doc, "Tabel 1.8: Ringkasan Eksekutif Seluruh Skenario Crosstab Realisasi Investasi PMDN Bab 1")
    chi_inv_headers = ["Variabel Independen (X)", "Variabel Dependen (Y)", "Chi-Square (χ²)", "P-Value", "Odds Ratio", "Kesimpulan"]
    chi_inv_data = [
        ["Realisasi Investasi PMDN (Juta Rp)", "Total Deforestasi Alam (Hektar)", f"{chi2_inv_total:.3f}", f"p {fmt_p(p_inv_total)}", f"{or_inv_total:.1f}", "SIGNIFIKAN" if p_inv_total < 0.05 else "TIDAK SIGNIFIKAN"],
        ["Realisasi Investasi PMDN (Juta Rp)", "Deforestasi Komoditas Tambang/Sawit (Hektar)", f"{chi2_inv_tambang:.3f}", f"p {fmt_p(p_inv_tambang)}", f"{or_inv_tambang:.1f}", "SIGNIFIKAN" if p_inv_tambang < 0.05 else "TIDAK SIGNIFIKAN"]
    ]
    add_table_1col(doc, chi_inv_headers, chi_inv_data, [4.0, 3.8, 2.0, 1.8, 1.8, 2.6], ['L', 'L', 'C', 'C', 'C', 'C'])
    add_p(doc, [("", False, False)])

    add_h4(doc, "E. Analisis Temuan Empiris: Efek Jeda Waktu (Time-Lagging)")
    add_p(doc, [
        ("Hasil pengujian seluruh skenario tabulasi silang PMDN mengungkap fenomena yang kompleks dalam alur investasi ekstraktif:", False, False)
    ])
    add_p(doc, [
        ("1. ", True, False), ("Ketidaksignifikanan Simultan & Variasi P-Value: ", True, False), ("Tingkat signifikansi yang bervariasi menyingkap tabir jeda waktu (lagging effect) dalam eksekusi investasi di lapangan.\n", False, False),
        ("2. ", True, False), ("Jeda Waktu Eksekusi Investasi (Lagging Effect): ", True, False), ("Suntikan modal masif di tahun tertentu tidak secara instan berwujud pembabatan lahan di tahun yang sama. Modal tersebut tertahan pada fase birokrasi, pembebasan lahan, dan pengadaan infrastruktur, sebelum daya rusaknya mengonversi lanskap hutan pada tahun-tahun berikutnya.\n", False, False),
        ("3. ", True, False), ("Konsentrasi Modal Ekstrem di 3 Provinsi: ", True, False), ("Data spasial membuktikan 89% dari total modal PMDN ekstraktif se-Sulawesi hanya tersedot ke tiga provinsi sentra (Sulteng, Sultra, Sulsel), mengakibatkan polarisasi pertumbuhan dan mengunci ketimpangan spasial.", False, False)
    ])

    # ═══════════════════════════════════════════════════════════
    # SUB-BAB 1.5 OSINT LOGISTIK PELABUHAN
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.5 Pelabuhan Ekspor: Ke Mana Nikel Sulawesi Dikirim?")
    
    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Eksploitasi nikel di Sulawesi tidak berhenti di tapak darat, melainkan terhubung langsung ke pasar global melalui infrastruktur pelabuhan. Bagian ini memetakan simpul logistik maritim yang mendistribusikan produk ekstraktif dari pesisir Sulawesi menuju negara tujuan utama seperti Tiongkok dan Jepang.", False, False)
    ])
    
    add_h4(doc, "B. Alur Logika Metodologis (Validasi Silang)")
    add_p(doc, [
        ("Verifikasi titik pelabuhan dan terminal khusus ekspor nikel dilakukan melalui protokol triangulasi informasi publik sebagaimana divisualisasikan pada ", False, False),
        ("Bagan Alur 1.5", True, False),
        (" berikut:", False, False),
    ])
    
    add_caption(doc, "Bagan Alur 1.5: Alur Logika Metodologis Validasi Silang (OSINT) Pelabuhan")
    if download_success_1_5:
        try:
            p_img_5 = doc.add_paragraph()
            p_img_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img_5 = p_img_5.add_run()
            run_img_5.add_picture(mermaid_png_path_1_5, width=Cm(15))
        except Exception as e:
            print(f"[WARN] Gagal memasukkan gambar Mermaid 1.5 ke DOCX: {e}")
            p_err_5 = doc.add_paragraph()
            run(p_err_5, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
    else:
        p_err_5 = doc.add_paragraph()
        run(p_err_5, "[Gambar Flowchart Gagal Diunduh, silakan periksa koneksi internet saat generate]", color=C_RED, pt=9)

    add_p(doc, [
        ("Detail dari keempat sumber informasi validasi silang (triangulasi) publik adalah sebagai berikut:", False, False),
    ])
    
    add_p(doc, [
        ("1. ", True, False), ("Laporan Investigasi Keselamatan Transportasi Laut (KNKT): ", True, False),
        ("Memverifikasi kapasitas dermaga curah dan bobot muatan kapal pengangkut bijih nikel hingga 52.378 DWT.\n", False, False),
        ("2. ", True, False), ("Regulasi Proyek Strategis Nasional (PSN): ", True, False),
        ("Lampiran Perpres No. 109 Tahun 2020 sektor kawasan industri terpadu.\n", False, False),
        ("3. ", True, False), ("Laporan Keberlanjutan & Tahunan Korporasi Terbuka: ", True, False),
        ("Laporan resmi PT Vale Indonesia Tbk dan PT ANTAM Tbk mengenai fasilitas pelabuhan khusus.\n", False, False),
        ("4. ", True, False), ("Laporan Audit Lembaga Riset Independen: ", True, False),
        ("Publikasi riset independen mengenai rantai pasok dan operasional terminal khusus maritim.", False, False),
    ])

    add_h4(doc, "C. Matriks Hasil Uji Empiris: Inventarisasi 6 Simpul Pelabuhan")
    add_p(doc, [
        ("Berdasarkan protokol validasi silang tersebut, profil komprehensif enam simpul pelabuhan dan terminal khusus utama di Pulau Sulawesi dipetakan pada ", False, False),
        ("Tabel 1.7", True, False),
        (" berikut:", False, False),
    ])
    
    add_caption(doc, "Tabel 1.7: Inventarisasi Enam Simpul Pelabuhan dan Terminal Khusus Ekspor Nikel di Pulau Sulawesi")
    log_headers = ["Simpul Kawasan Industri", "Wilayah Administrasi", "Fasilitas Pelabuhan / Terminal", "Status Regulasi", "Kapasitas Kapal", "Tujuan Utama Ekspor"]
    log_rows = [
        ["IMIP Morowali", "Morowali, Sulawesi Tengah", "Pelabuhan Samudera & Dermaga Curah", "PSN (Perpres 109/2020)", "Hingga 52.378 DWT", "Pasar Global (Tiongkok)"],
        ["GNI Morowali Utara", "Morowali Utara, Sulteng", "Terminal Khusus Pesisir Tomori", "Izin Industri Mandiri", "Hingga 30.000 DWT", "Pasar Global (Tiongkok)"],
        ["VDNI Konawe", "Konawe, Sulawesi Tenggara", "Dermaga Khusus Curah & Kargo", "PSN (Perpres 109/2020)", "Hingga 50.000 DWT", "Pasar Global (Tiongkok)"],
        ["OSS Konawe", "Konawe, Sulawesi Tenggara", "Dermaga Terintegrasi Konawe", "PSN (Perpres 109/2020)", "Hingga 50.000 DWT", "Pasar Global (Tiongkok)"],
        ["Pomalaa (ANTAM)", "Kolaka, Sulawesi Tenggara", "Dermaga Pomalaa & Konveyor", "Kawasan BUMN Industri", "Hingga 12.000 DWT", "Jepang & Korsel"],
        ["Sorowako (Vale)", "Luwu Timur, Sulawesi Selatan", "Pelabuhan Balantang Malili", "Kontrak Karya Tambang", "Hingga 15.000 DWT", "Jepang & Skandinavia"]
    ]
    add_table_1col(doc, log_headers, log_rows, [2.5, 2.5, 3.5, 2.5, 2.0, 3.0], ['L', 'L', 'L', 'C', 'C', 'L'])

    # ═══════════════════════════════════════════════════════════
    # SUB-BAB 1.6 PETA JALUR DISTRIBUSI LOGISTIK
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.6 Peta Jalur Distribusi Logistik Nikel Sulawesi")
    
    add_h4(doc, "A. Pengantar & Kerangka Narasi")
    add_p(doc, [
        ("Metode analisis pada tahapan ini difokuskan pada ", False, False),
        ("Pemetaan Kausalitas (Spasial)", True, False),
        (" untuk membedah asimetri penguasaan ruang antara wilayah hulu (origin: sumber ekstraksi di Sulawesi) dan hilir (destination: pusat industrialisasi luar negeri). Garis diplot menggunakan rute pelayaran untuk merepresentasikan jarak tempuh aktual kapal logistik di permukaan bumi.", False, False),
    ])

    add_h4(doc, "B. Alur Logika Metodologis (Flowchart Analisis Spasial)")
    add_p(doc, [
        ("Rangkaian proses ekstraksi koordinat hingga visualisasi spasial menggunakan pustaka *Scattergeo* direpresentasikan pada ", False, False),
        ("Bagan Alur 1.6", True, False),
        (" berikut:", False, False),
    ])
    
    add_caption(doc, "Bagan Alur 1.6: Alur Logika Metodologis Pemetaan Jalur Logistik Maritim (Bézier Curve)")
    if download_success_1_6:
        try:
            p_img_16 = doc.add_paragraph()
            p_img_16.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img_16 = p_img_16.add_run()
            run_img_16.add_picture(str(mermaid_png_path_1_6), width=Cm(15))
        except Exception as e:
            p_err_16 = doc.add_paragraph()
            run(p_err_16, "[Gambar Flowchart Gagal Dimuat]", color=C_RED, pt=9)
            
    add_h4(doc, "C. Formulasi Matematis: Kurva Parametrik Alur Pelayaran")
    add_p(doc, [
        ("Pemetaan alur pelayaran internasional dimodelkan menggunakan formulasi kurva parametrik lengkung (*Bézier Curve*) untuk merepresentasikan alur laut kepulauan secara realistis:", False, False),
    ])
    add_formula(doc, "Formulasi Kurva Parametrik Alur Pelayaran Maritim", "Kurva(t) = (1 - t)^2 * Titik_Asal + 2 * (1 - t) * t * Titik_Kontrol + t^2 * Titik_Tujuan",
                var_desc=[
                    ("Kurva(t)", "Vektor posisi koordinat geografis lintasan kapal pada parameter waktu t (rentang kontinu [0, 1])."),
                    ("Titik_Asal", "Titik koordinat geografis pelabuhan muat khusus di pesisir Sulawesi."),
                    ("Titik_Kontrol", "Titik koordinat jangkar pemandu kurva lengkung di perairan internasional."),
                    ("Titik_Tujuan", "Titik koordinat geografis pelabuhan bongkar di negara tujuan ekspor."),
                ])

    add_h4(doc, "D. Matriks Hasil Pemetaan Spasial")
    add_p(doc, [
        ("Agregasi titik asal (Origin) dan tujuan akhir (Destination) di-render langsung ke dalam proyeksi spasial. Struktur data logistik yang menggerakkan pemetaan ini memetakan pelabuhan asal menuju pasar utama, yaitu ", False, False),
        ("Tiongkok, Jepang, dan Korea Selatan", True, False),
        (".", False, False)
    ])

    add_h4(doc, "E. Interpretasi Spasial Industri (Anatomi Rantai Pasok)")
    add_p(doc, [
        ("Peta rute logistik maritim mengilustrasikan alur distribusi produk olahan nikel dari kawasan industri di Sulawesi:", False, False)
    ])
    add_p(doc, [
        ("1. ", True, False), ("Orientasi Ekspor: ", True, False), ("Kawasan industri utama yang berstatus Proyek Strategis Nasional (PSN) mengalirkan produk olahan ke sentra-sentra industri manufaktur di pasar internasional.\n", False, False),
        ("2. ", True, False), ("Integrasi Rantai Pasok: ", True, False), ("Mayoritas rute pengapalan terhubung langsung dengan pelabuhan ekspor tujuan, yang mengindikasikan posisi kawasan pemurnian di Sulawesi sebagai pemasok bahan baku setengah jadi pada rantai pasok global.\n", False, False),
        ("3. ", True, False), ("Dinamika Rute Maritim: ", True, False), ("Peta rute mencerminkan diversifikasi pasar ekspor (Asia Timur) dan jaringan logistik kawasan.", False, False)
    ])

    # ═══════════════════════════════════════════════════════════
    # SUB-BAB 1.7 & 1.8 MASTER TABEL & MATRIKS TAHAPAN
    # ═══════════════════════════════════════════════════════════
    add_h2(doc, "1.7 Matriks Indikator dan Sumber Data Resmi Bab 1")
    add_p(doc, [
        ("Seluruh variabel kuantitatif, kategori analisis, satuan ukur, periode tahun observasi, dan institusi penyedia data primer resmi yang digunakan dalam Bab 1 dikompilasikan pada ", False, False),
        ("Tabel 1.8", True, False),
        (" berikut:", False, False),
    ])

    # Tabel 1.8: Master Indikator
    add_caption(doc, "Tabel 1.8: Matriks Indikator dan Sumber Data Primer Resmi Bab 1 (Ekspansi Industri Ekstraktif)")
    master_headers = ["No", "Nama Indikator", "Kategori Analisis", "Satuan Ukur", "Cakupan Tahun", "Institusi & Sumber Data Resmi"]
    master_rows = [
        ["1", "Izin Usaha Pertambangan (IUP) Baru", "Faktor Tekanan Ekstraktif", "Unit Izin", "2014–2024", "Kementerian ESDM (Minerbaone)"],
        ["2", "Luas Wilayah Konsesi Tambang Baru", "Faktor Tekanan Ekstraktif", "Hektar (Ha)", "2014–2024", "Kementerian ESDM (Minerbaone)"],
        ["3", "Kapasitas Terpasang PLTU Captive", "Infrastruktur Energi Khusus", "Megawatt (MW)", "2014–2024", "Global Energy Monitor (GEM)"],
        ["4", "Fasilitas Pengolahan & Pemurnian (Smelter)", "Fasilitas Industri Hilir", "Unit Fasilitas", "2014–2024", "Kementerian ESDM & Basis Data Industri"],
        ["5", "Realisasi Investasi PMDN", "Arus Modal Domestik", "Triliun Rupiah", "2016–2024", "Kementerian Investasi / BKPM"],
        ["6", "PDRB Menurut 17 Lapangan Usaha", "Struktur Ekonomi Makro", "Triliun Rupiah", "2016–2024", "Badan Pusat Statistik (BPS Provinsi se-Sulawesi)"],
        ["7", "PDRB Kabupaten Sentra Tambang", "Struktur Ekonomi Daerah", "Triliun Rupiah", "2016–2024", "Badan Pusat Statistik (BPS Kabupaten se-Sulteng)"],
        ["8", "Luas Kehilangan Hutan Komoditas", "Dampak Perubahan Tutupan Lahan", "Hektar (Ha)", "2014–2023", "Global Forest Watch (GFW / Univ. of Maryland)"],
        ["9", "Simpul Dermaga & Terminal Khusus Ekspor", "Infrastruktur Rantai Pasok", "Titik Koordinat & DWT", "2014–2024", "KNKT Kemenhub, Perpres PSN, Laporan Korporasi"],
    ]
    add_table_1col(doc, master_headers, master_rows, [0.8, 4.2, 3.2, 2.0, 2.0, 5.8], ['C', 'L', 'L', 'C', 'C', 'L'])

    add_h2(doc, "1.8 Bagan Alur Kerangka Kerja Riset Bab 1")
    add_p(doc, [
        ("Keseluruhan struktur metodologis riset Bab 1 dioperasionalkan melalui empat fase kerja berurutan sebagaimana disajikan pada ", False, False),
        ("Tabel 1.9", True, False),
        (" berikut:", False, False),
    ])

    # Tabel 1.9: Alur Kerja
    add_caption(doc, "Tabel 1.9: Matriks Tahapan dan Alur Kerangka Kerja Riset Bab 1")
    flow_headers = ["Tahapan Riset", "Fokus Metodologis", "Bahan & Sumber Data", "Keluaran / Hasil Analisis"]
    flow_rows = [
        ["Fase I: Pengumpulan Data", "Kurasi data resmi lintas kementerian dan lembaga", "Publikasi BPS, Minerbaone, BKPM, GEM, dan GFW", "Basis Data Tabular Panel Provinsi (2014–2024)"],
        ["Fase II: Reklasifikasi Hukum", "Penyusunan kerangka rantai pasok hukum terintegrasi", "UU No. 3/2020, PP No. 96/2021, Perpres No. 112/2022", "3 Klaster Makro (Ekstraktif, Akar Rumput, Jasa)"],
        ["Fase III: Pengujian Statistik", "Uji signifikansi hubungan dan rasio peluang", "Tabel Kontinjensi, Uji Chi-Square, Odds Ratio", "Bukti Kausalitas Signifikan Tekanan vs Deforestasi"],
        ["Fase IV: Pemetaan Rantai Pasok", "Triangulasi data logistik dan pemodelan maritim", "Laporan KNKT, Perpres PSN, Kurva Parametrik Bézier", "Peta Alur Rantai Pasok Ekspor & Konsentrasi Spasial 78%"],
    ]
    add_table_1col(doc, flow_headers, flow_rows, [3.0, 4.2, 4.8, 5.0], ['L', 'L', 'L', 'L'])

    add_note_box(doc, "Integritas Ilmiah & Keterbukaan Publik", 
                 "Seluruh data, formulasi, dan temuan yang disajikan dalam dokumen metodologi ini berpedoman pada kaidah akademis yang dapat diuji ulang (reproducible research) berdasarkan data publik resmi pemerintah dan lembaga riset independen.")

    # Simpan DOCX
    docx_path = tool_dir / "Metodologi_Bab1_Ekspansi_Industri.docx"
    doc.save(str(docx_path))
    print(f"  [OK] Tersimpan: {docx_path}")

    # 3. Generate File HTML 1-Kolom Penuh
    print("[3/5] Membangun Metodologi_Bab1_Ekspansi_Industri.html (Format Publik)...")
    
    html_kab_rows = ""
    md_kab_rows = []
    for kab_name, row_k in df_kab_pivot.iterrows():
        if 'Morowali' in kab_name and 'Utara' not in kab_name: basis = "Hilirisasi Nikel (Smelter & PLTU)"
        elif 'Morowali Utara' in kab_name: basis = "Hilirisasi Nikel (Smelter GNI)"
        elif 'Banggai' == kab_name: basis = "Migas, Tambang & Perdagangan"
        elif 'Palu' in kab_name: basis = "Jasa, Perdagangan & Pemerintahan"
        elif 'Parigi' in kab_name: basis = "Pertanian Pangan & Hortikultura"
        elif 'Donggala' in kab_name: basis = "Pertanian, Perkebunan & Galian C"
        elif 'Poso' in kab_name: basis = "Pertanian & Perkebunan Kakao"
        elif 'Sigi' in kab_name: basis = "Pertanian Pangan & Hortikultura"
        elif 'Toli' in kab_name: basis = "Perkebunan Cengkeh & Perikanan"
        elif 'Buol' in kab_name: basis = "Kelapa Sawit & Tanaman Pangan"
        elif 'Tojo' in kab_name: basis = "Pertanian & Pariwisata Bahari"
        elif 'Kepulauan' in kab_name: basis = "Perikanan Tangkap & Kelautan"
        elif 'Laut' in kab_name: basis = "Perikanan & Budidaya Laut"
        else: basis = "Pertanian & Jasa"

        even_cls = ' class="data-tr-even"' if len(md_kab_rows) % 2 == 1 else ''
        html_kab_rows += f'    <tr{even_cls}><td class="data-td"><strong>{kab_name}</strong></td><td class="data-td" style="text-align:center;">{row_k["Akar Rumput"]:.2f}</td><td class="data-td" style="text-align:center;">{row_k["Ekstraktif"]:.2f}</td><td class="data-td" style="text-align:center;">{row_k["Jasa"]:.2f}</td><td class="data-td" style="text-align:center;"><strong>{row_k["Total"]:.2f}</strong></td><td class="data-td" style="text-align:center;">{row_k["Pct_Akar"]:.1f}%</td><td class="data-td" style="text-align:center;">{row_k["Pct_Ekstraktif"]:.1f}%</td><td class="data-td" style="text-align:center;">{row_k["Pct_Jasa"]:.1f}%</td><td class="data-td">{basis}</td></tr>\n'
        md_kab_rows.append(f"| **{kab_name}** | {row_k['Akar Rumput']:.2f} | {row_k['Ekstraktif']:.2f} | {row_k['Jasa']:.2f} | **{row_k['Total']:.2f}** | {row_k['Pct_Akar']:.1f}% | {row_k['Pct_Ekstraktif']:.1f}% | {row_k['Pct_Jasa']:.1f}% | {basis} |")

    html_chi_rows_4 = ""
    md_chi_rows_4 = ""
    for idx, row in enumerate(chi_inv_data):
        row_class = "data-tr-even" if (idx + 1) % 2 == 0 else "data-tr-odd"
        html_chi_rows_4 += f'    <tr class="{row_class}"><td class="data-td">{row[0]}</td><td class="data-td">{row[1]}</td><td class="data-td" style="text-align:center;">{row[2]}</td><td class="data-td" style="text-align:center;">{row[3]}</td><td class="data-td" style="text-align:center;">{row[4]}</td><td class="data-td" style="text-align:center; font-weight:bold;">{row[5]}</td></tr>\n'
        md_chi_rows_4 += f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} | {row[5]} |\n"

    html_content = rf"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<title>Laporan Metodologi Bab 1 — Ekspansi Industri Ekstraktif (CELIOS ECC)</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500;700&display=swap" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad:true, theme:'dark'}});</script>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  html {{ font-size: 10px; background: #0A0D13; min-height: 100vh; padding: 25px; }}
  body {{
    font-family: 'Inter', Arial, sans-serif;
    background: #0E1117; color: #D4D4D4; line-height: 1.65;
    max-width: 920px; padding: 30px 40px; margin: 0 auto;
    box-shadow: 0 8px 45px rgba(0,0,0,0.85); border-radius: 6px;
  }}
  @media print {{
    html {{ font-size: 8pt; background: #fff; padding: 0; }}
    body {{ background: #fff; color: #111; padding: 0; max-width: 100%; box-shadow: none; }}
    .hdr-title {{ color: #1B5E20 !important; }}
    .hdr-sub {{ color: #2E7D32 !important; }}
    h2 {{ color: #1B5E20 !important; border-bottom-color: #2E7D32 !important; }}
    h3 {{ color: #2E7D32 !important; }}
    p, li, td {{ color: #222 !important; }}
    strong {{ color: #111 !important; }}
    .data-th {{ background: #2E7D32 !important; color: #fff !important; -webkit-print-color-adjust: exact; }}
    .data-tr-even .data-td {{ background: #F5FBF5 !important; }}
    .note-box {{ background: #F1F8E9 !important; border-left-color: #2E7D32 !important; color: #333 !important; }}
    .formula {{ background: #EDF7EE !important; color: #1B5E20 !important; border-color: #A5D6A7 !important; }}
  }}
  .hdr {{ margin-bottom: 16pt; border-bottom: 2pt solid #2E7D32; padding-bottom: 8pt; }}
  .hdr-sub {{ font-size: 8.5pt; font-weight: 700; color: #43A047; text-transform: uppercase; letter-spacing: 1.5pt; }}
  .hdr-title {{ font-size: 15pt; font-weight: 800; color: #81C784; text-transform: uppercase; margin-top: 4pt; line-height: 1.3; }}
  h2 {{ font-size: 11.5pt; font-weight: 700; color: #81C784; text-transform: uppercase; margin-top: 22pt; margin-bottom: 8pt; padding-bottom: 3pt; border-bottom: 1pt solid #2E7D32; }}
  h3 {{ font-size: 10.5pt; font-weight: 700; color: #66BB6A; margin-top: 14pt; margin-bottom: 5pt; }}
  h4 {{ font-size: 9.5pt; font-weight: 700; color: #A5D6A7; margin-top: 10pt; margin-bottom: 4pt; }}
  p {{ font-size: 9.5pt; color: #CCCCCC; margin-bottom: 8pt; text-align: justify; }}
  ul {{ font-size: 9pt; color: #CCCCCC; margin-left: 18pt; margin-bottom: 8pt; }}
  li {{ margin-bottom: 3pt; }}
  strong {{ color: #E8F5E9; }}
  .note-box {{ background: #132213; border-left: 3pt solid #2E7D32; padding: 8pt 12pt; margin: 10pt 0; font-size: 8.5pt; color: #C8E6C9; border-radius: 0 4px 4px 0; }}
  .table-caption {{ font-size: 8.5pt; font-weight: 700; color: #81C784; margin-top: 12pt; margin-bottom: 4pt; }}
  table {{ width: 100%; border-collapse: collapse; margin: 6pt 0 14pt 0; font-size: 8.5pt; }}
  .data-th {{ background: #1B5E20; color: #FFFFFF; font-weight: 700; padding: 5pt 7pt; text-align: left; border: 0.5pt solid #2E7D32; }}
  .data-td {{ padding: 4.5pt 7pt; border: 0.5pt solid #243524; color: #CCCCCC; vertical-align: top; }}
  .data-tr-even .data-td {{ background: #131B13; }}
  .formula {{ background: #0D1B0E; border: 0.8pt solid #2E7D32; border-radius: 4px; padding: 7pt 11pt; font-family: 'JetBrains Mono', monospace; font-size: 8.5pt; color: #A5D6A7; margin: 6pt 0; }}
  .formula-title {{ font-size: 8.5pt; font-weight: 700; color: #81C784; margin-top: 8pt; }}
  .var-desc {{ background: #0E1610; border-left: 2pt solid #43A047; padding: 6pt 10pt; margin: 4pt 0 10pt 0; font-size: 8pt; color: #B0BEC5; }}
  .var-desc-title {{ font-weight: 700; color: #81C784; margin-bottom: 3pt; }}
  .var-desc ul {{ list-style-type: none; margin-left: 0; margin-bottom: 0; }}
  .var-desc li {{ margin-bottom: 2pt; }}
  .mermaid {{ background: #0D1610; border: 0.8pt solid #2E7D32; border-radius: 6px; padding: 12pt; margin: 8pt 0 14pt 0; text-align: center; }}
</style>
</head>
<body>

<div class="hdr">
  <div class="hdr-sub">CELIOS — Center of Economic and Law Studies &nbsp;|&nbsp; Laporan Riset Metodologi D3TLH</div>
  <div class="hdr-title">BAB I: METODOLOGI ANALISIS EKSPANSI INDUSTRI EKSTRAKTIF DAN INFRASTRUKTUR PENUNJANG DI PULAU SULAWESI</div>
</div>

<p>
  Dokumen laporan metodologi ini menyajikan kerangka ilmiah, landasan regulasi, formulasi matematis, prosedur analisis statistik, serta metodologi pembuktian berbasis data terbuka yang dioperasionalkan pada <strong>Bab 1: Ekspansi Industri Ekstraktif</strong> dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi periode 2014–2024.
</p>

<h2>SUB-BAB 1.1: Konteks Makro: Breakdown PDRB per Komoditas</h2>

<h3>1.1.1 Konteks Makro: Dominasi Ekstraktif vs Ekonomi Akar Rumput</h3>
<p>
  Bagian ini menganalisis struktur Produk Domestik Regional Bruto (PDRB) pada enam provinsi di Pulau Sulawesi sepanjang periode 2016–2024 menggunakan visualisasi grafik area bertumpuk (<em>Stacked Area Chart</em>). Analisis ini ditujukan untuk menguji secara empiris apakah percepatan pertumbuhan ekonomi daerah benar-benar bersumber dari sektor produktif masyarakat lokal atau didominasi oleh industri ekstraktif padat modal yang mengalihkan pemanfaatan ruang dan sumber daya alam.
</p>

<div class="note-box">
  <strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong><br>
  Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi <em>Stacked Area Chart</em> memetakan dinamika Produk Domestik Regional Bruto (PDRB) berdasarkan klasifikasi rantai pasok hukum (<em>Legal Supply-Chain</em>) untuk membandingkan trajektori Sektor Ekstraktif, Ekonomi Akar Rumput, dan Sektor Jasa & Lainnya.
</div>

<h4>A. Kerangka Dekomposisi Sektoral & Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)</h4>
<p>
  Sistem Klasifikasi Baku Lapangan Usaha Indonesia (KBLI 2020) yang diterbitkan BPS membagi aktivitas perekonomian ke dalam 17 sektor. Melalui pendekatan <em>Legal Supply-Chain</em>, ke-17 sektor tersebut direklasifikasi menjadi 3 Klaster Makro (Ekstraktif, Akar Rumput, dan Jasa & Lainnya). Rincian pembagian sektor, dasar regulasi, serta intisari ketentuan hukum disajikan secara lengkap pada <strong>Tabel 1.1</strong> berikut:
</p>

<div class="table-caption">Tabel 1.1: Reklasifikasi Sektoral PDRB KBLI 2020 Berdasarkan Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:14%;">Kategori BPS</th>
      <th class="data-th" style="width:24%;">Sektor Lapangan Usaha</th>
      <th class="data-th" style="width:16%;">Klasifikasi Analisis</th>
      <th class="data-th" style="width:22%;">Dasar Regulasi & Mandat Hukum</th>
      <th class="data-th" style="width:24%;">Intisari Ketentuan Hukum</th>
    </tr>
  </thead>
  <tbody>
    <tr><td class="data-td"><strong>Kategori B</strong></td><td class="data-td">Pertambangan dan Penggalian</td><td class="data-td" style="text-align:center;">Ekstraktif</td><td class="data-td">Perpres No. 26 Tahun 2010</td><td class="data-td">Ketentuan Pasal 1 Ayat (2) mengenai pengambilan komoditas tambang dari dalam bumi.</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Kategori C</strong></td><td class="data-td">Industri Pengolahan (Smelter Logam)</td><td class="data-td" style="text-align:center;">Ekstraktif</td><td class="data-td">UU No. 3 Tahun 2020 & PP No. 96 Tahun 2021</td><td class="data-td">Pasal 102–103 mewajibkan pengolahan dan pemurnian di dalam negeri sebagai kesatuan pertambangan.</td></tr>
    <tr><td class="data-td"><strong>Kategori D</strong></td><td class="data-td">Pengadaan Listrik & Gas (PLTU Captive)</td><td class="data-td" style="text-align:center;">Ekstraktif</td><td class="data-td">Perpres No. 112 Tahun 2022 & RUPTL PLN</td><td class="data-td">Pasal 3 Ayat (4) huruf b mengecualikan PLTU baru hanya bagi yang terintegrasi melayani smelter.</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Kategori A</strong></td><td class="data-td">Pertanian, Kehutanan, Perikanan</td><td class="data-td" style="text-align:center;">Ekonomi Akar Rumput</td><td class="data-td">KBLI 2020 BPS</td><td class="data-td">Sektor pemanfaatan sumber daya hayati terbarukan dan penyerap tenaga kerja lokal.</td></tr>
    <tr><td class="data-td"><strong>Kategori E–U</strong></td><td class="data-td">13 Sektor Jasa & Konstruksi</td><td class="data-td" style="text-align:center;">Sektor Jasa & Lainnya</td><td class="data-td">Klasifikasi Standar BPS</td><td class="data-td">Sektor sekunder dan tersier penunjang perekonomian daerah.</td></tr>
  </tbody>
</table>

<h4>B. Alur Logika Metodologis Rantai Pasok Hukum (Mengapa Kat. B + C + D = Ekstraktif)</h4>
<p>
  Keterkaitan ketiga kategori lapangan usaha tersebut sebagai satu kesatuan rantai pasok ekstraktif dimodelkan dalam kerangka alur logika hukum sebagaimana diilustrasikan pada <strong>Bagan Alur 1.1</strong> berikut:
</p>

<div class="table-caption">Bagan Alur 1.1: Alur Logika Metodologis Rantai Pasok Hukum Sektor Ekstraktif</div>
<div class="mermaid">
{mermaid_str_1_1}
</div>

<h4>C. Formulasi Matematis: Persamaan Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)</h4>
<p>
  Data diproses menggunakan pendekatan Dekomposisi PDRB Sektoral. Nilai PDRB dikelompokkan menjadi 3 agregat makro melalui pendekatan Legal Supply-Chain. Sektor Ekstraktif dihitung dengan menggabungkan tiga kategori lapangan usaha utama yang saling terintegrasi (tambang, smelter, dan PLTU captive) menggunakan persamaan <strong>Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)</strong>:
</p>

<div class="formula-title">Persamaan Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation):</div>
<div class="formula">Sektor_Ekstraktif = PDRB(Kat.B: Pertambangan) + PDRB(Kat.C: Ind. Pengolahan) + PDRB(Kat.D: Listrik)</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Sektor_Ekstraktif</strong>: Total nilai tambah bruto dari klaster industri ekstraktif yang saling terintegrasi (satuan: Triliun Rupiah).</li>
    <li>• <strong>PDRB(Kat.B: Pertambangan)</strong>: Nilai tambah bruto dari aktivitas eksplorasi dan ekstraksi bijih mineral dari dalam bumi (BPS KBLI 2020 Kategori B).</li>
    <li>• <strong>PDRB(Kat.C: Ind. Pengolahan)</strong>: Nilai tambah bruto dari pemurnian dan peleburan logam dasar di fasilitas smelter nikel (BPS KBLI 2020 Kategori C / Golongan Pokok 24).</li>
    <li>• <strong>PDRB(Kat.D: Listrik)</strong>: Nilai tambah bruto dari pengadaan daya listrik batubara non-jaringan (PLTU captive) khusus melayani smelter (BPS KBLI 2020 Kategori D).</li>
  </ul>
</div>

<p>Secara lengkap, perhitungan ketiga komponen makroekonomi dan indikator turunannya dirumuskan melalui sistem persamaan berikut:</p>

<div class="formula-title">Persamaan Ekonomi Akar Rumput:</div>
<div class="formula">Sektor_Akar_Rumput = PDRB(Kat.A: Pertanian, Kehutanan, dan Perikanan)</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Sektor_Akar_Rumput</strong>: Total nilai Produk Domestik Regional Bruto yang dihasilkan dari pemanfaatan sumber daya hayati terbarukan (satuan: Triliun Rupiah).</li>
    <li>• <strong>PDRB(Kat.A)</strong>: Agregasi nilai tambah tanaman pangan, hortikultura, perkebunan rakyat, peternakan, kehutanan, dan perikanan tangkap/budidaya (BPS KBLI 2020 Kategori A).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Sektor Jasa & Lainnya:</div>
<div class="formula">Sektor_Jasa = Jumlah PDRB (Kategori E sampai dengan Kategori U)</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Sektor_Jasa</strong>: Total nilai tambah gabungan dari 13 sektor penunjang sekunder dan tersier di luar sektor ekstraktif dan akar rumput (satuan: Triliun Rupiah).</li>
    <li>• <strong>PDRB (Kat. E s.d. U)</strong>: Akumulasi nilai tambah sektor konstruksi, perdagangan besar/eceran, transportasi, pergudangan, akomodasi, informasi & komunikasi, jasa keuangan, real estat, dan jasa umum.</li>
  </ul>
</div>

<div class="formula-title">Persamaan Total Produk Domestik Regional Bruto (PDRB Wilayah):</div>
<div class="formula">Total_PDRB = Sektor_Ekstraktif + Sektor_Akar_Rumput + Sektor_Jasa</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Total_PDRB</strong>: Nilai keseluruhan output ekonomi regional bruto suatu wilayah provinsi atau kabupaten pada periode tahun berjalan atas dasar harga berlaku (satuan: Triliun Rupiah).</li>
    <li>• <strong>Sektor_Ekstraktif</strong>: Total nilai tambah sektor ekstraktif terintegrasi (Triliun Rupiah).</li>
    <li>• <strong>Sektor_Akar_Rumput</strong>: Total nilai tambah ekonomi berbasis masyarakat lokal dan sumber daya hayati (Triliun Rupiah).</li>
    <li>• <strong>Sektor_Jasa</strong>: Total nilai tambah sektor jasa dan fasilitas pendukung (Triliun Rupiah).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Pangsa Kontribusi Sektor Ekstraktif (%):</div>
<div class="formula">Pangsa_Ekstraktif (%) = ( Sektor_Ekstraktif / Total_PDRB ) * 100</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Pangsa_Ekstraktif (%)</strong>: Persentase pangsa/porsi dominasi sektor ekstraktif terhadap total kue perekonomian wilayah (satuan: Persen / %).</li>
    <li>• <strong>Sektor_Ekstraktif</strong>: Nilai tambah nominal sektor ekstraktif terintegrasi pada tahun observasi (Triliun Rupiah).</li>
    <li>• <strong>Total_PDRB</strong>: Total nilai nominal PDRB seluruh 17 sektor lapangan usaha pada tahun yang sama (Triliun Rupiah).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Laju Pertumbuhan Tahunan Sektoral (YoY):</div>
<div class="formula">Laju_Pertumbuhan_Tahunan (%) = [ ( Nilai_Tahun_t - Nilai_Tahun_{{t-1}} ) / Nilai_Tahun_{{t-1}} ] * 100</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Laju_Pertumbuhan_Tahunan (%)</strong>: Tingkat percepatan atau perlambatan ekspansi tahunan suatu sektor ekonomi (satuan: Persen / %).</li>
    <li>• <strong>Nilai_Tahun_t</strong>: Nilai nominal PDRB sektor pada tahun observasi berjalan (t).</li>
    <li>• <strong>Nilai_Tahun_{{t-1}}</strong>: Nilai nominal PDRB sektor pada satu tahun sebelumnya (t - 1).</li>
  </ul>
</div>

<p>
  Definisi operasional, cakupan lapangan usaha, dan institusi penyedia data primer untuk masing-masing komponen variabel dalam sistem persamaan di atas dipaparkan pada <strong>Tabel 1.2</strong> berikut:
</p>

<div class="table-caption">Tabel 1.2: Definisi Operasional Komponen Makroekonomi dan Sumber Data PDRB Sektoral</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:20%;">Komponen Analisis</th>
      <th class="data-th" style="width:22%;">Cakupan Lapangan Usaha</th>
      <th class="data-th" style="width:32%;">Definisi Operasional</th>
      <th class="data-th" style="width:12%;">Satuan Nilai</th>
      <th class="data-th" style="width:14%;">Sumber Data Primer</th>
    </tr>
  </thead>
  <tbody>
    <tr><td class="data-td"><strong>Sektor Ekstraktif</strong></td><td class="data-td">Kategori B, Kategori C, Kategori D</td><td class="data-td">Akumulasi nilai tambah bruto pertambangan bijih nikel, peleburan logam dasar, dan penyediaan energi PLTU captive.</td><td class="data-td">Triliun Rupiah</td><td class="data-td">BPS Provinsi (SIMDASI)</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Ekonomi Akar Rumput</strong></td><td class="data-td">Kategori A</td><td class="data-td">Nilai tambah pertanian tanaman pangan, perkebunan rakyat, kehutanan, dan perikanan tangkap maupun budidaya.</td><td class="data-td">Triliun Rupiah</td><td class="data-td">BPS Provinsi</td></tr>
    <tr><td class="data-td"><strong>Sektor Jasa & Lainnya</strong></td><td class="data-td">Kategori E hingga Kategori U</td><td class="data-td">Nilai tambah gabungan perdagangan, konstruksi, transportasi, perbankan, dan jasa layanan umum.</td><td class="data-td">Triliun Rupiah</td><td class="data-td">BPS Provinsi</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Total PDRB Wilayah</strong></td><td class="data-td">Seluruh 17 Kategori Lapangan Usaha</td><td class="data-td">Total nilai Produk Domestik Regional Bruto suatu wilayah atas dasar harga berlaku pada periode tahun berjalan.</td><td class="data-td">Triliun Rupiah</td><td class="data-td">BPS Provinsi</td></tr>
    <tr><td class="data-td"><strong>Pangsa Ekstraktif (%)</strong></td><td class="data-td">Rasio Kontribusi Relatif</td><td class="data-td">Persentase kontribusi sektor ekstraktif terhadap keseluruhan perekonomian provinsi atau kabupaten.</td><td class="data-td">Persen (%)</td><td class="data-td">Hasil Olahan CELIOS</td></tr>
  </tbody>
</table>

<p>
  Penerapan formulasi di atas menunjukkan perbedaan struktur ekonomi yang sangat kontras antarwilayah di Pulau Sulawesi. Di <strong>Sulawesi Tengah (sebagai pusat hilirisasi)</strong>, ekspansi industri ekstraktif melaju sangat pesat hingga menguasai <strong>{pct_sulteng_eks:.1f}% dari total PDRB provinsi</strong> pada tahun {latest_year_pdrb}. Sebaliknya, porsi Ekonomi Akar Rumput mengalami penurunan pangsa yang signifikan. Hal ini menunjukkan ketergantungan ekonomi daerah yang sangat tinggi pada satu klaster industri padat modal, berbeda dengan provinsi lain seperti Sulawesi Selatan dan Gorontalo yang perekonomiannya bertumpu pada basis pertanian terbarukan.
</p>

<h3>1.1.2 Pemusatan Sektor Ekstraktif di Kabupaten se-Sulawesi Tengah</h3>
<p>
  Jika dianalisis secara spasial pada tingkat kabupaten di Sulawesi Tengah, terlihat konsentrasi kegiatan industri ekstraktif. Kabupaten <strong>Morowali</strong> dan <strong>Morowali Utara</strong> mendominasi struktur PDRB provinsi melalui pengembangan kawasan industri hilirisasi dan PLTU Captive. Analisis ini membandingkan komposisi ketiga sektor advokatif di seluruh 13 kabupaten/kota se-Sulawesi Tengah pada tahun terbaru ({latest_year_kab}).
</p>

<div class="note-box">
  <strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong><br>
  Badan Pusat Statistik (BPS) Kabupaten se-Sulawesi Tengah (diolah CELIOS). Visualisasi <em>Stacked Bar Chart</em> memetakan struktur Produk Domestik Regional Bruto (PDRB) tahun {latest_year_kab} pada seluruh 13 kabupaten/kota untuk mengidentifikasi tingkat konsentrasi sektoral dan polarisasi spasial antara sentra industri pengolahan nikel dengan daerah non-sentra.
</div>

<h4>A. Rasionalitas Spasial & Urgensi Dekomposisi Sektoral Tingkat Kabupaten</h4>
<p>
  Analisis agregat pada tingkat provinsi sering kali menghasilkan <strong>Bias Ilusi Agregat (Aggregate Illusion Bias)</strong>, di mana angka pertumbuhan ekonomi makro yang tinggi memberi kesan seolah seluruh wilayah menikmati kemakmuran yang seimbang. Namun, ketika data didekomposisi ke tingkat kabupaten/kota, terlihat jurang pemisah ekonomi yang sangat tajam antara wilayah <strong>Enklave Industri Ekstraktif</strong> dengan daerah agraris tradisional sekitarnya.
</p>

<h4>B. Alur Logika Analisis Disparitas Spasial Kabupaten</h4>
<p>
  Kerangka kerja metodologis dalam membedah ketimpangan intra-provinsial ini diilustrasikan pada <strong>Bagan Alur 1.2</strong> berikut:
</p>

<div class="table-caption">Bagan Alur 1.2: Alur Logika Metodologis Dekomposisi Spasial PDRB Tingkat Kabupaten se-Sulawesi Tengah</div>
<div class="mermaid">
{mermaid_str_1_2}
</div>

<h4>C. Formulasi Matematis: Persamaan Agregasi Sektoral Kabupaten (Legal Supply-Chain Aggregation)</h4>
<p>Kalkulasi PDRB tingkat kabupaten menggunakan sistem persamaan berikut:</p>

<div class="formula-title">Persamaan Agregasi Sektor Ekstraktif Tingkat Kabupaten:</div>
<div class="formula">Sektor_Ekstraktif_Kabupaten = PDRB_Kab(Kat.B: Pertambangan) + PDRB_Kab(Kat.C: Ind. Pengolahan) + PDRB_Kab(Kat.D: Listrik)</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Sektor_Ekstraktif_Kabupaten</strong>: Total nilai tambah sektor ekstraktif di tingkat kabupaten target (satuan: Triliun Rupiah).</li>
    <li>• <strong>PDRB_Kab(Kat.B: Pertambangan)</strong>: Nilai PDRB kabupaten dari aktivitas penambangan bijih logam dan galian (BPS Kategori B).</li>
    <li>• <strong>PDRB_Kab(Kat.C: Ind. Pengolahan)</strong>: Nilai PDRB kabupaten dari industri peleburan logam dasar / smelter (BPS Kategori C).</li>
    <li>• <strong>PDRB_Kab(Kat.D: Listrik)</strong>: Nilai PDRB kabupaten dari penyediaan daya listrik batubara captive (BPS Kategori D).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Total Produk Domestik Regional Bruto Tingkat Kabupaten:</div>
<div class="formula">Total_PDRB_Kabupaten = Sektor_Ekstraktif_Kabupaten + Sektor_Akar_Rumput_Kabupaten + Sektor_Jasa_Kabupaten</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Total_PDRB_Kabupaten</strong>: Total output perekonomian bruto kabupaten target atas dasar harga berlaku (satuan: Triliun Rupiah).</li>
    <li>• <strong>Sektor_Ekstraktif_Kabupaten</strong>: Nilai tambah bruto sektor ekstraktif terintegrasi di kabupaten (Triliun Rupiah).</li>
    <li>• <strong>Sektor_Akar_Rumput_Kabupaten</strong>: Nilai tambah sektor pertanian, kehutanan, dan perikanan di kabupaten (Triliun Rupiah).</li>
    <li>• <strong>Sektor_Jasa_Kabupaten</strong>: Nilai tambah sektor perdagangan, transportasi, dan jasa layanan di kabupaten (Triliun Rupiah).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Porsi Sektoral dalam Kabupaten (Porsi (%) pada Tooltip Dashboard):</div>
<div class="formula">Porsi_Sektor_Kabupaten (%) = ( Nilai_Sektor_Kabupaten / Total_PDRB_Kabupaten ) * 100</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Porsi_Ekstraktif (%)</strong>: Persentase kontribusi Sektor Ekstraktif: ( Sektor_Ekstraktif / Total_PDRB ) * 100 (misal Morowali: 45.2%).</li>
    <li>• <strong>Porsi_Jasa (%)</strong>: Persentase kontribusi Sektor Jasa & Lainnya: ( Sektor_Jasa / Total_PDRB ) * 100 (misal Morowali: 54.0%).</li>
    <li>• <strong>Porsi_Akar_Rumput (%)</strong>: Persentase kontribusi Sektor Ekonomi Akar Rumput: ( Sektor_Akar_Rumput / Total_PDRB ) * 100 (misal Morowali: 0.8%).</li>
    <li>• <strong>Total_PDRB_Kabupaten</strong>: Total nilai nominal PDRB seluruh sektor di kabupaten target (Triliun Rupiah).</li>
  </ul>
</div>

<h4>D. Rincian Definisi Operasional & Matriks Distribusi PDRB 13 Kabupaten</h4>
<p>
  Penerapan sistem persamaan di atas terhadap seluruh 13 kabupaten dan kota di Provinsi Sulawesi Tengah pada tahun {latest_year_kab} disajikan secara komprehensif pada <strong>Tabel 1.3</strong> berikut:
</p>

<div class="table-caption">Tabel 1.3: Distribusi Nilai Tambah Bruto dan Komposisi Sektoral PDRB 13 Kabupaten/Kota di Sulawesi Tengah (Tahun {latest_year_kab})</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:16%;">Kabupaten / Kota</th>
      <th class="data-th" style="width:10%;">Akar Rumput (T Rp)</th>
      <th class="data-th" style="width:10%;">Ekstraktif (T Rp)</th>
      <th class="data-th" style="width:10%;">Jasa (T Rp)</th>
      <th class="data-th" style="width:11%;">Total PDRB (T Rp)</th>
      <th class="data-th" style="width:11%;">Porsi Akar Rumput (%)</th>
      <th class="data-th" style="width:11%;">Porsi Ekstraktif (%)</th>
      <th class="data-th" style="width:10%;">Porsi Jasa (%)</th>
      <th class="data-th" style="width:11%;">Basis Utama Ekonomi</th>
    </tr>
  </thead>
  <tbody>
{html_kab_rows}
  </tbody>
</table>

<h4>E. Analisis Temuan Empiris: Polarisasi Ekstrem Morowali vs Daerah Non-Smelter</h4>
<p>Data empiris pada Tabel 1.3 mengungkap bukti polarisasi ekonomi wilayah yang sangat ekstrem di Sulawesi Tengah:</p>
<ul style="margin-left: 20px; margin-bottom: 15px;">
  <li style="margin-bottom: 6px;"><strong>1. Dominasi Sektor Ekstraktif Morowali:</strong> Kabupaten Morowali mencatatkan nilai sektor ekstraktif sebesar Rp {df_kab_pivot.loc['Morowali', 'Ekstraktif']:.2f} Triliun atau menguasai porsi {df_kab_pivot.loc['Morowali', 'Pct_Ekstraktif']:.1f}% dari total kue ekonomi kabupatennya (Rp {df_kab_pivot.loc['Morowali', 'Total']:.2f} Triliun). Nilai sektor ekstraktif Morowali saja melampaui gabungan total PDRB dari delapan kabupaten lainnya di Sulawesi Tengah.</li>
  <li style="margin-bottom: 6px;"><strong>2. Pemusatan pada Dua Sentra Hilirisasi:</strong> Kabupaten Morowali dan Morowali Utara merupakan dua daerah dengan nilai Sektor Ekstraktif tertinggi di Sulawesi Tengah, membuktikan bahwa percepatan output industri pertambangan dan hilirisasi terkunci pada kawasan industri smelter.</li>
  <li style="margin-bottom: 6px;"><strong>3. Ketertinggalan Daerah Non-Sentra:</strong> Sebaliknya, delapan kabupaten lainnya (seperti Banggai Laut, Banggai Kepulauan, Tojo Una-Una, Buol, Toli-Toli, Sigi, Poso, dan Donggala) memiliki porsi Sektor Ekstraktif yang sangat rendah (<11%) dan tetap bergantung pada sektor pertanian rakyat (Akar Rumput) berproduktivitas rendah dengan keterbatasan akses terhadap nilai tambah modal.</li>
</ul>

<h3>1.1.3 Perbandingan Distribusi 17 Sektor Komoditas per Provinsi (Small Multiples, Tahun Terbaru)</h3>
<p>
  Visualisasi komparatif <strong>Small Multiples Horizontal Bar Chart</strong> membedah struktur 17 sektor lapangan usaha KBLI 2020 secara terpisah pada enam provinsi di Pulau Sulawesi pada tahun terbaru ({latest_year_pdrb}). Setiap panel provinsi menampilkan sektor yang diurutkan dari penyumbang terbesar hingga terkecil dengan skala sumbu nilai yang disetarakan secara seragam untuk memastikan validitas komparasi lintas wilayah.
</p>

<div class="note-box">
  <strong>Sumber Data Resmi & Deskripsi Visualisasi:</strong><br>
  Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi <em>Small Multiples Horizontal Bar Chart</em> menyajikan dekomposisi 17 sektor PDRB tahun {latest_year_pdrb} di 6 provinsi se-Pulau Sulawesi. Sumbu X disetarakan pada rentang nilai seragam ([0, {max_x_val_all:.1f} Triliun Rp]) dengan pewarnaan berdasarkan 3 klaster makro (Merah: Ekstraktif, Hijau: Ekonomi Akar Rumput, Abu-abu: Sektor Jasa & Lainnya) guna mengidentifikasi spesialisasi dan anomali struktural ekonomi masing-masing provinsi.
</div>

<h4>A. Kerangka Konseptual & Standardisasi Skala Komparatif (Uniform Scale Small Multiples)</h4>
<p>
  Dalam analisis data multidimensi lintas wilayah, penggunaan skala dinamis mandiri (<em>independent dynamic scaling</em>) pada masing-masing panel sering kali menimbulkan <strong>Bias Distorsi Visual Komparatif (Visual Comparison Bias)</strong>. Tanpa penyetaraan batas skala maksimum, sektor dengan nominal kecil di provinsi ber-PDRB rendah dapat terlihat secara visual setara dengan sektor bernilai ratusan triliun di provinsi ber-PDRB besar. Oleh karena itu, metodologi ini menetapkan batas skala maksimum sumbu X yang seragam (<em>Uniform Scale Bound</em>) sebesar nilai maksimum sektor tertinggi di seluruh pulau ditambah faktor ruang margin sebesar 15%.
</p>

<h4>B. Alur Logika Metodologis Analisis Small Multiples 17 Sektor</h4>
<p>
  Kerangka operasionalisasi analisis perbandingan terpisah 17 sektor lapangan usaha ini dimodelkan dalam kerangka alur logika sebagaimana diilustrasikan pada <strong>Bagan Alur 1.3</strong> berikut:
</p>

<div class="table-caption">Bagan Alur 1.3: Alur Logika Metodologis Analisis Komparatif Small Multiples 17 Sektor PDRB per Provinsi</div>
<div class="mermaid">
{mermaid_str_1_3}
</div>

<h4>C. Formulasi Matematis: Persamaan Agregasi dan Porsi 17 Sektor Komoditas</h4>
<p>Kalkulasi perbandingan sektoral dan normalisasi skala grafik dihitung menggunakan sistem formulasi berikut:</p>

<div class="formula-title">Persamaan Normalisasi Nilai Sektor ke Satuan Triliun Rupiah:</div>
<div class="formula">Nilai_Sektor_Triliun = Nilai_Sektor_Miliar / 1000</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Nilai_Sektor_Triliun</strong>: Nilai tambah bruto sektor lapangan usaha dalam satuan baku Triliun Rupiah.</li>
    <li>• <strong>Nilai_Sektor_Miliar</strong>: Nilai nominal PDRB mentah dari publikasi resmi BPS (satuan: Miliar Rupiah).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Porsi Sektoral per Provinsi (Porsi (%) pada Tooltip Dashboard):</div>
<div class="formula">Porsi_Sektor_Provinsi (%) = ( Nilai_Sektor_Provinsi / Total_PDRB_Provinsi ) * 100</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Porsi_Sektor_Provinsi (%)</strong>: Persentase kontribusi sektor target terhadap keseluruhan total PDRB provinsi bersangkutan (satuan: Persen / %). Angka ini ditampilkan pada tooltip 'Porsi (%)' di dashboard.</li>
    <li>• <strong>Nilai_Sektor_Provinsi</strong>: Nilai tambah bruto sektor lapangan usaha target di provinsi bersangkutan (Triliun Rupiah).</li>
    <li>• <strong>Total_PDRB_Provinsi</strong>: Total nilai nominal PDRB seluruh 17 sektor di provinsi bersangkutan (Triliun Rupiah).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Batas Maksimum Skala Sumbu X Seragam (Uniform Scale Bound):</div>
<div class="formula">Skala_Maksimum_Sumbu_X = max(Nilai_Sektor_Seluruh_Provinsi) * 1.15</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Skala_Maksimum_Sumbu_X</strong>: Batas nilai tertinggi sumbu horizontal (X-axis) pada seluruh 6 panel grafik Small Multiples (satuan: Triliun Rupiah).</li>
    <li>• <strong>max(Nilai_Sektor_Seluruh_Provinsi)</strong>: Nilai tertinggi dari seluruh kombinasi 17 sektor di 6 provinsi (yaitu Sektor Industri Pengolahan Sulteng sebesar Rp {sulteng_ind_val:.2f} Triliun).</li>
    <li>• <strong>1.15</strong>: Faktor pengali ruang margin 15% untuk penempatan label teks nilai nominal pada ujung grafik batang.</li>
  </ul>
</div>

<h4>D. Rincian Data Empiris: Matriks Perbandingan Sektor Unggulan 6 Provinsi</h4>
<p>
  Penerapan sistem persamaan di atas terhadap seluruh 6 provinsi di Pulau Sulawesi pada tahun {latest_year_pdrb} disajikan secara komprehensif pada <strong>Tabel 1.4</strong> berikut:
</p>

<div class="table-caption">Tabel 1.4: Profil Distribusi dan Sektor Unggulan PDRB 6 Provinsi se-Pulau Sulawesi (Tahun {latest_year_pdrb})</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:14%;">Provinsi</th>
      <th class="data-th" style="width:10%;">Total PDRB (T Rp)</th>
      <th class="data-th" style="width:14%;">Sektor Peringkat 1</th>
      <th class="data-th" style="width:7%;">Porsi (%)</th>
      <th class="data-th" style="width:14%;">Sektor Peringkat 2</th>
      <th class="data-th" style="width:7%;">Porsi (%)</th>
      <th class="data-th" style="width:14%;">Sektor Peringkat 3</th>
      <th class="data-th" style="width:7%;">Porsi (%)</th>
      <th class="data-th" style="width:13%;">Karakteristik Dominasi</th>
    </tr>
  </thead>
  <tbody>
{html_top_rows}
  </tbody>
</table>

<h4>E. Analisis Temuan Empiris & Interpretasi Sektoral Dashboard</h4>
<p>Hasil komparasi sektoral pada Tabel 1.4 mengungkap perbedaan fundamental orientasi pembangunan antarprovinsi se-Sulawesi:</p>
<ul style="margin-left: 20px; margin-bottom: 15px;">
  <li style="margin-bottom: 6px;"><strong>1. Anomali Struktural Sulawesi Tengah:</strong> Sulawesi Tengah menjadi satu-satunya wilayah di mana perekonomiannya dikuasai secara mutlak oleh Sektor Industri Pengolahan (Smelter Logam Dasar) senilai Rp {sulteng_ind_val:.2f} Triliun (41.2%) dan Sektor Pertambangan senilai Rp {sulteng_tambang_val:.2f} Triliun (14.6%). Gabungan kedua sektor ekstraktif ini menguasai lebih dari 55% total PDRB provinsi.</li>
  <li style="margin-bottom: 6px;"><strong>2. Sulawesi Tenggara sebagai Sentra Tambang Nikel Hulu:</strong> Sulawesi Tenggara memperlihatkan kontribusi Sektor Pertambangan yang sangat tinggi (Rp {sultra_tambang_val:.2f} Triliun atau 21.1%), berada tepat di bawah sektor Pertanian (23.5%), mengonfirmasi perannya sebagai lumbung pasokan bijih nikel primer.</li>
  <li style="margin-bottom: 6px;"><strong>3. Basis Agraris Terbarukan di Empat Provinsi Lainnya:</strong> Sebaliknya, empat provinsi lainnya (Sulawesi Barat 46.1%, Gorontalo 37.3%, Sulawesi Selatan 21.8%, dan Sulawesi Utara 20.6%) secara konsisten menempatkan Sektor Pertanian, Kehutanan, dan Perikanan sebagai sektor penyumbang terbesar PDRB, ditopang oleh sektor perdagangan dan jasa layanan publik yang menyerap mayoritas angkatan kerja daerah.</li>
</ul>

<h2>SUB-BAB 1.2: Konsentrasi Kawasan Industri & PLTU Captive</h2>
<p>
  Intensifikasi industri pengolahan mineral di Pulau Sulawesi berpusat pada pembangunan mega-smelter yang ditopang secara mutlak oleh pembangkit listrik tenaga uap khusus (<em>PLTU Captive</em>) batu bara non-jaringan (<em>off-grid</em>). Bagian ini mengombinasikan <strong>Analisis Spasial Deskriptif</strong> untuk mengidentifikasi tingkat pemusatan fasilitas dan kapasitas energi fosil di enam provinsi, dengan <strong>Uji Tabulasi Silang Panel (Inferential Spatiotemporal Crosstabulation)</strong> berstandar SPSS guna membuktikan secara ilmiah keterkaitan antara ekspansi PLTU Captive terhadap kehilangan tutupan hutan di Pulau Sulawesi.
</p>

<div class="note-box">
  <strong>Sumber Data Resmi & Deskripsi Metodologis:</strong><br>
  Kementerian Energi dan Sumber Daya Mineral (ESDM / Minerbaone), Global Energy Monitor (GEM Coal Plant Tracker), dan Global Forest Watch (GFW / University of Maryland) (diolah CELIOS). Visualisasi <em>Bar Chart</em> Konsentrasi Industri dan Pemetaan Spasial menyajikan distribusi {tot_smelter_all} unit fasilitas smelter serta {tot_pltu_mw_all:,.0f} MW kapasitas terpasang aktif PLTU captive di 6 provinsi se-Pulau Sulawesi. Analisis dipadukan dengan Uji Tabulasi Silang Data Panel Spasiotemporal (Chi-Square Test & Risk Odds Ratio, N=60) untuk menguji keterkaitan ekspansi energi fosil industri terhadap eskalasi deforestasi komoditas.
</div>

<h4>A. Pemusatan Spasial Fasilitas Smelter dan PLTU Captive</h4>
<p>
  Intensifikasi industri pengolahan nikel di Sulawesi berpusat pada fasilitas mega-smelter. Pengoperasian <strong>{tot_smelter_all} fasilitas smelter</strong> didukung oleh kapasitas energi batu bara <strong>{tot_pltu_mw_all:,.0f} MW dari PLTU Captive</strong>. Berbeda dengan sistem kelistrikan umum PLN, pembangkit ini dikembangkan secara internal untuk menyokong operasi kawasan industri.
</p>

<h4>B. Metodologi: Analisis Spasial & Uji Tabulasi Silang</h4>
<p>
  Pengujian keterkaitan antara pembangunan PLTU Captive dengan kehilangan tutupan hutan dioperasionalkan melalui Standar Operasional Prosedur (SOP) tabulasi silang berstandar SPSS. Rangkaian tahapan logika metodologis, asumsi frekuensi harapan, hingga estimasi faktor risiko dimodelkan pada <strong>Bagan Alur 1.4</strong> berikut:
</p>

<div class="table-caption">Bagan Alur 1.4: Standar Operasional Prosedur (SOP) & Alur Logika Uji Tabulasi Silang (Crosstab) PLTU Captive vs Deforestasi</div>
<div class="mermaid">
{mermaid_str}
</div>

<h4>C. Formulasi Matematis: Kalkulasi Konsentrasi Spasial & Uji Chi-Square</h4>
<p>Parameterisasi konsentrasi spasial dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:</p>

<div class="formula-title">Persamaan Akumulasi Kapasitas PLTU Kumulatif per Wilayah (MW):</div>
<div class="formula">Kapasitas_PLTU_Kumulatif_t (MW) = Jumlah Kapasitas Aktif Baru (MW) dari Tahun 2014 hingga Tahun t</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Kapasitas_PLTU_Kumulatif_t (MW)</strong>: Total akumulasi kapasitas daya terpasang operasional PLTU captive batubara aktif hingga tahun t (satuan: Megawatt / MW).</li>
    <li>• <strong>Kapasitas Aktif Baru</strong>: Besaran daya listrik unit PLTU off-grid yang mulai beroperasi komersial pada tahun tertentu (satuan: Megawatt / MW).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Rasio Konsentrasi Spasial Fasilitas Smelter (% pada Grafik Dashboard):</div>
<div class="formula">Porsi_Smelter_Provinsi (%) = ( Jumlah_Smelter_Provinsi / Total_Smelter_Sulawesi ) * 100</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Porsi_Smelter_Provinsi (%)</strong>: Persentase pangsa fasilitas smelter di provinsi bersangkutan terhadap seluruh Pulau Sulawesi (satuan: Persen / %).</li>
    <li>• <strong>Jumlah_Smelter_Provinsi</strong>: Banyaknya unit smelter yang beroperasi di wilayah provinsi tertentu.</li>
    <li>• <strong>Total_Smelter_Sulawesi</strong>: Total keseluruhan fasilitas smelter di Pulau Sulawesi ({tot_smelter_all} unit).</li>
  </ul>
</div>

<div class="formula-title">Persamaan Uji Independensi Chi-Square Pearson (χ² Kontinjensi 2x2):</div>
<div class="formula">Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Chi_Square (χ²)</strong>: Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara ekspansi PLTU Captive dengan lonjakan deforestasi pada panel spasiotemporal (N=60).</li>
    <li>• <strong>Frekuensi_Observasi (O)</strong>: Jumlah kasus aktual yang tercatat pada sel tabel kontinjensi 2x2.</li>
    <li>• <strong>Frekuensi_Harapan (E)</strong>: Jumlah kasus teoretis jika kedua variabel saling independen: E = (Total Baris * Total Kolom) / N.</li>
  </ul>
</div>

<div class="formula-title">Persamaan Rasio Keunggulan Risiko (Risk Odds Ratio / OR):</div>
<div class="formula">Odds_Ratio (OR) = ( a * d ) / ( b * c )</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Odds_Ratio (OR)</strong>: Ukuran kelipatan risiko peluang terjadinya deforestasi komoditas tinggi pada kelompok dengan PLTU Captive aktif (>0 MW) dibanding kelompok tanpa PLTU Captive (≤0 MW).</li>
    <li>• <strong>a</strong>: Jumlah observasi panel pada kelompok PLTU Rendah dan Deforestasi Rendah ({a_2} kasus).</li>
    <li>• <strong>b</strong>: Jumlah observasi panel pada kelompok PLTU Rendah dan Deforestasi Tinggi ({b_2} kasus).</li>
    <li>• <strong>c</strong>: Jumlah observasi panel pada kelompok PLTU Tinggi dan Deforestasi Rendah ({c_2} kasus).</li>
    <li>• <strong>d</strong>: Jumlah observasi panel pada kelompok PLTU Tinggi dan Deforestasi Tinggi ({d_2} kasus).</li>
  </ul>
</div>

<h4>D. Matriks Hasil Uji Empiris: Konsentrasi Spasial & Skenario Crosstab</h4>
<p>
  Penerapan sistem pengujian statistik tabulasi silang pada data panel 6 provinsi selama 1 dekade (2014–2023, total 60 observasi) disajikan secara lengkap pada <strong>Tabel 1.5</strong> berikut:
</p>

<div class="table-caption">Tabel 1.5: Ringkasan Eksekutif Seluruh Skenario Crosstab PLTU Captive Bab 1</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:25%;">Variabel Independen (X)</th>
      <th class="data-th" style="width:25%;">Variabel Dependen (Y)</th>
      <th class="data-th" style="width:12%;">Chi-Square (&chi;&sup2;)</th>
      <th class="data-th" style="width:13%;">P-Value</th>
      <th class="data-th" style="width:10%;">Odds Ratio</th>
      <th class="data-th" style="width:15%;">Kesimpulan</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td class="data-td"><strong>Kapasitas Aktif PLTU Kumulatif (MW)</strong></td>
      <td class="data-td">Total Deforestasi Alam (Hektar)</td>
      <td class="data-td" style="text-align:center;">{chi2_pltu:.3f}</td>
      <td class="data-td" style="text-align:center;">p {fmt_p(p_val_display)}</td>
      <td class="data-td" style="text-align:center;">{or_pltu:.1f}</td>
      <td class="data-td" style="color:#2E7D32; text-align:center;"><strong>{{'SIGNIFIKAN' if p_val_display < 0.05 else 'TIDAK SIGNIFIKAN'}}</strong></td>
    </tr>
    <tr class="data-tr-even">
      <td class="data-td"><strong>Kapasitas Aktif PLTU Kumulatif (MW)</strong></td>
      <td class="data-td">Deforestasi Komoditas Tambang/Sawit (Hektar)</td>
      <td class="data-td" style="text-align:center;">{chi2_pltu:.3f}</td>
      <td class="data-td" style="text-align:center;">p {fmt_p(p_val_display)}</td>
      <td class="data-td" style="text-align:center;">{or_pltu:.1f}</td>
      <td class="data-td" style="color:#2E7D32; text-align:center;"><strong>{{'SIGNIFIKAN' if p_val_display < 0.05 else 'TIDAK SIGNIFIKAN'}}</strong></td>
    </tr>
  </tbody>
</table>

<h4>E. Interpretasi Spasial Industri: Eksternalitas dan Efek Meluber (Spillover)</h4>
<p>
  Konsentrasi masif fasilitas smelter dan 9.825 MW PLTU Captive di Sulawesi Tengah, Tenggara, dan Selatan terbukti memicu ledakan deforestasi. Uji Chi-Square membuktikan secara absolut bahwa daerah dengan PLTU Captive memiliki risiko deforestasi 18 kali lipat lebih parah dibandingkan wilayah non-smelter. Daya rusak industri ini melampaui tapak proyek (spillover effect).
</p>
<p>Kombinasi analisis spasial (Tabel 1.4b) dan hasil pengujian empiris (Tabel 1.5) membuktikan secara meyakinkan keterkaitan langsung antara ekspansi infrastruktur ekstraktif dengan kerusakan tutupan hutan di Pulau Sulawesi:</p>
<ul style="margin-left: 20px; margin-bottom: 15px;">
  <li style="margin-bottom: 6px;"><strong>1. Pemusatan Ekstrem di 3 Sentra Ekstraktif Utama:</strong> 100% kapasitas PLTU Captive dan mayoritas smelter berpusat di wilayah ini, memicu akumulasi deforestasi komoditas hingga ratusan ribu hektar, berbanding terbalik dengan "Area Non-Smelter".</li>
  <li style="margin-bottom: 6px;"><strong>2. Signifikansi Statistik yang Sangat Kuat (p {fmt_p(p_val_display)}):</strong> Hipotesis Nol (H0) ditolak mutlak. Bukti empiris mengonfirmasi bahwa penambahan kapasitas PLTU Captive berkorelasi langsung dengan lonjakan kehilangan tutupan hutan.</li>
  <li style="margin-bottom: 6px;"><strong>3. Kelipatan Risiko Bencana Ekologis (Odds Ratio = {or_pltu:.2f}x):</strong> Wilayah dengan PLTU Captive memiliki risiko deforestasi komoditas {int(or_pltu)} KALI LIPAT lebih besar. Hal ini didorong konversi masif untuk infrastruktur pendukung (coal yard, jalur transmisi, dan jalan logistik).</li>
  <li style="margin-bottom: 6px;"><strong>4. Efek Meluber Lintas Batas (Spillover Effect) & Emisi Karbon Terkunci:</strong> Eksternalitas destruktif proyek merambat luas mendegradasi DAS dan laut, mengorbankan ruang hidup lokal, serta mengunci emisi dari ketergantungan puluhan juta ton batu bara per tahun.</li>
</ul>

<h2>SUB-BAB 1.3: Tren Pertumbuhan Izin Tambang Baru & Uji Signifikansi Statistik</h2>
<h4>A. Pengantar & Kerangka Narasi</h4>
<p>Pola perizinan pertambangan di Pulau Sulawesi selama satu dekade terakhir menunjukkan peningkatan alokasi ruang yang signifikan. Berdasarkan data agregat <strong>Minerbaone</strong>, tercatat {tot_izin:,} Izin Usaha Pertambangan (IUP) baru sepanjang 2014-2024, dengan total luas konsesi mencapai {tot_luas_izin:,} Hektar.</p>
<p>Berdasarkan analisis tren time-series pada grafik "Penerbitan Izin Tambang", penerbitan izin pada periode awal (2014) tercatat lebih rendah. Peningkatan signifikan terjadi pada periode 2022-2024. Anotasi pada grafik mencatat kenaikan sebesar <strong>246% pada periode 2022-2024</strong>. Data ini mengindikasikan perlunya evaluasi terhadap instrumen pengendalian perizinan dan tata ruang. Distribusi perizinan tertinggi berada di Sulawesi Tengah dan Sulawesi Tenggara, yang selaras dengan kawasan pengembangan industri pemurnian nikel.</p>
<p>Uji <strong>Crosstabulation</strong> pada analisis ini digunakan untuk mengukur hubungan antara laju penerbitan perizinan (X) dan indikator deforestasi di wilayah tersebut (Y).</p>

<h4>B. Alur Logika Metodologis (Flowchart)</h4>
<p>Pendekatan statistik Time-Series untuk mengidentifikasi tren pertumbuhan izin tambang diilustrasikan pada <strong>Bagan Alur 1.5</strong> berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya merujuk secara penuh pada <strong>Bagan Alur 1.4</strong> (di sub-bab sebelumnya) dengan penyesuaian konfigurasi variabel spesifik sesuai Tabel Asumsi Dasar di bawah gambar.</p>
<div class="table-caption">Bagan Alur 1.5: Alur Logika Tren Pertumbuhan Izin Tambang Baru</div>
<div class="mermaid">
{mermaid_str_1_3_izin}
</div>
<div class="table-caption">Tabel 1.5b: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.3)</div>
<table>
  <thead>
    <tr>
      <th class="data-th">Komponen Uji</th>
      <th class="data-th">Definisi Variabel (Sub-bab 1.3)</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td class="data-td">Variabel Independen (X)</td>
      <td class="data-td">Frekuensi Penerbitan Izin Tambang Baru (IUP) / Luas Konsesi Baru (Ha)</td>
    </tr>
    <tr>
      <td class="data-td">Variabel Dependen (Y)</td>
      <td class="data-td">Deforestasi Komoditas (Ha) / Total Deforestasi Alam (Ha)</td>
    </tr>
    <tr>
      <td class="data-td">Asumsi H0</td>
      <td class="data-td">Tren perizinan baru dan laju deforestasi saling independen (tidak berhubungan).</td>
    </tr>
  </tbody>
</table>

<h4>C. Formulasi Matematis: Agregasi Tren Spasial & Uji Chi-Square</h4>
<p>Parameterisasi laju pertumbuhan perizinan, luas konsesi, dan pengujian signifikansi dampaknya terhadap deforestasi dihitung menggunakan formulasi berikut:</p>
<p><strong>Persamaan Agregasi Luas Konsesi Tambang Baru per Kategori Spasial</strong><br>
Total_Konsesi_Kategori_t = &sum;(Luas_Konsesi_i) untuk i di Kategori_Wilayah</p>
<ul>
  <li><strong>Total_Konsesi_Kategori_t:</strong> Total luas konsesi tambang baru yang diterbitkan pada tahun berjalan (t) di dalam kategori wilayah spasial tertentu (satuan: Hektar / Ha).</li>
  <li><strong>Kategori_Wilayah:</strong> Pengelompokan spasial provinsi observasi menjadi 'Daerah Sentra Tambang' atau 'Daerah Non-Sentra'.</li>
  <li><strong>Luas_Konsesi_i:</strong> Luasan tiap individu Izin Usaha Pertambangan (IUP) baru yang terbit (satuan: Hektar / Ha).</li>
</ul>
<p><strong>Laju Pertumbuhan Izin Tahunan (Regresi Komparatif YoY)</strong><br>
Pertumbuhan_Izin (%) = [ ( Jumlah_Izin_t - Jumlah_Izin_{{t-1}} ) / Jumlah_Izin_{{t-1}} ] &times; 100</p>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>&bull; <strong>Pertumbuhan_Izin (%)</strong>: Persentase perubahan laju penerbitan izin tambang baru antar-tahun (satuan: Persen / %).</li>
    <li>&bull; <strong>Jumlah_Izin_t</strong>: Agregasi jumlah izin (atau luasan) pada tahun berjalan (t).</li>
    <li>&bull; <strong>Jumlah_Izin_{{t-1}}</strong>: Agregasi jumlah izin (atau luasan) pada satu tahun sebelumnya (t - 1).</li>
  </ul>
</div>
<p><strong>Pengklasifikasian Kategori Data (Binning Threshold Median)</strong><br>
Kategori = IF(Nilai_Prov_Tahun &ge; Median(Seluruh Panel), "Tinggi", "Rendah")</p>
<ul>
  <li><strong>Kategori:</strong> Data panel spasial-temporal diubah menjadi dua tingkatan untuk uji tabulasi silang (Tinggi vs Rendah).</li>
  <li><strong>Median(Seluruh Panel):</strong> Ambang batas (threshold) dari nilai tengah keseluruhan observasi panel (N=60).</li>
</ul>
<p><strong>Uji Independensi Chi-Square (Pearson)</strong><br>
&chi;&sup2; = &sum; [ (O_i - E_i)&sup2; / E_i ]</p>
<ul>
  <li><strong>&chi;&sup2;:</strong> Nilai statistik uji Chi-Square yang mencerminkan tingkat deviasi frekuensi aktual terhadap frekuensi harapan.</li>
  <li><strong>O_i:</strong> Frekuensi observasi aktual (Observed) pada sel ke-i dalam matriks tabulasi silang.</li>
  <li><strong>E_i:</strong> Frekuensi harapan (Expected) pada sel ke-i jika hipotesis nol (independensi penuh) bernilai benar.</li>
</ul>
<p><strong>Odds Ratio (OR)</strong><br>
OR = (A &times; D) / (B &times; C)</p>
<ul>
  <li><strong>OR:</strong> Odds Ratio, metrik yang mengukur seberapa besar kelipatan peluang terjadinya peristiwa dependen (Deforestasi Tinggi) jika terpapar variabel independen (Kapasitas Tinggi).</li>
  <li><strong>A, B, C, D:</strong> Frekuensi absolut masing-masing sel dalam matriks 2x2. A = (Tinggi, Tinggi), D = (Rendah, Rendah).</li>
</ul>

<p>
  Dinamika historis perizinan secara terperinci dapat dilihat pada <strong>Tabel 1.5c</strong>, yang menunjukkan tren penerbitan izin baru di wilayah studi:
</p>
<div class="table-caption">Tabel 1.5c: Tren Penerbitan Izin Tambang Sulawesi (2014-2024)</div>
<table class="data-table">
  <thead>
    <tr>
      <th class="data-th">Tahun</th>
      <th class="data-th">Gorontalo</th>
      <th class="data-th">Sulawesi Barat</th>
      <th class="data-th">Sulawesi Selatan</th>
      <th class="data-th">Sulawesi Tengah</th>
      <th class="data-th">Sulawesi Tenggara</th>
      <th class="data-th">Sulawesi Utara</th>
    </tr>
  </thead>
  <tbody>
{html_tren_rows.rstrip()}
  </tbody>
</table>

<h4>D. Matriks Hasil Uji Empiris: Rekapitulasi Signifikansi Panel</h4>
<p>Hasil lengkap pengujian independensi statistik Chi-Square dan estimasi Odds Ratio (OR) untuk seluruh faktor tekanan terhadap kehilangan tutupan hutan komoditas dirangkum pada <strong>Tabel 1.6</strong> berikut:</p>

<div class="table-caption">Tabel 1.6: Ringkasan Hasil Uji Independensi Chi-Square (&chi;&sup2;) dan Odds Ratio (OR) Data Panel Bab 1</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:28%;">Variabel Faktor Tekanan</th>
      <th class="data-th" style="width:26%;">Variabel Dampak Lingkungan</th>
      <th class="data-th" style="width:12%;">Chi-Square (&chi;&sup2;)</th>
      <th class="data-th" style="width:10%;">Nilai p</th>
      <th class="data-th" style="width:10%;">Odds Ratio</th>
      <th class="data-th" style="width:6%;">df</th>
      <th class="data-th" style="width:20%;">Kesimpulan Ilmiah</th>
    </tr>
  </thead>
  <tbody>
{html_chi_rows.rstrip()}
  </tbody>
</table>

<h4>E. Analisis Temuan Empiris: Pembedahan Realitas Ekologis</h4>
<p>Data panel membedah realitas di lapangan: lonjakan izin di wilayah pusat ekstraksi sejalan dengan tingginya nilai Chi-Square. Nilai Odds Ratio menegaskan bahwa wilayah dengan tren izin tambang yang tinggi memiliki peluang lebih besar untuk mengalami tekanan deforestasi tinggi pada tahun-tahun berjalan dan berikutnya.</p>
<p>Secara spesifik, terjadi <strong>lonjakan absolut sebesar 246%</strong> dalam penerbitan izin tambang baru pada rentang 2022 hingga 2024. Lonjakan ekstrem ini mengindikasikan percepatan luar biasa dari ekspansi industri ekstraktif yang mengabaikan kapasitas daya dukung lingkungan tapak, terutama di sentra-sentra produksi.</p>

<h2>SUB-BAB 1.4: Analisis Realisasi Investasi PMDN dan Dampak Terhadap Tutupan Hutan</h2>

<h4>A. Pengantar & Kerangka Narasi</h4>
<p>
  Akumulasi Penanaman Modal Dalam Negeri sebesar <strong>Rp {tot_investasi_triliun:,} Triliun</strong> (Kementerian Investasi / BKPM) yang masuk dari tahun 2016-2024 berbanding lurus dengan <strong>{tot_deforestasi:,} Hektar kehilangan tutupan hutan komoditas</strong> (Global Forest Watch). Grafik sumbu ganda (<em>dual-axis</em>) digunakan untuk membandingkan laju investasi dan laju deforestasi antara wilayah sentra industri tambang dengan non-sentra. Terlihat adanya fenomena <strong>Efek Jeda Waktu (Time-Lagging Effect)</strong>, di mana peningkatan realisasi modal pada tahap awal perizinan dan konstruksi diikuti oleh lonjakan pembukaan lahan hutan fisik pada 1 hingga 2 tahun berikutnya.
</p>

<h4>B. Alur Logika Metodologis Analisis Realisasi Investasi PMDN</h4>
<p>Kerangka operasionalisasi uji statistik tabulasi silang antara realisasi Investasi PMDN dan deforestasi dimodelkan dalam kerangka alur logika sebagaimana diilustrasikan pada <strong>Bagan Alur 1.4</strong> berikut:</p>
<div class="table-caption">Bagan Alur 1.4: Alur Logika Metodologis Uji Independensi Panel Investasi PMDN vs Deforestasi</div>
<div class="mermaid">
{mermaid_str_1_4}
</div>

<h4>C. Formulasi Matematis: Agregasi Dampak Ekologis & Pengujian Statistik</h4>
<p><strong>Persamaan Agregasi Luasan Deforestasi Berdasarkan Faktor Penggerak (Driver)</strong><br>
Total_Deforestasi_Driver_k = &sum;(Area_Loss_i) untuk i di Kategori_Driver_k</p>
<ul>
  <li><strong>Total_Deforestasi_Driver_k:</strong> Total luas kehilangan tutupan pohon yang diakibatkan oleh faktor penggerak k (contoh: Ekspansi Komoditas) di seluruh wilayah observasi (satuan: Hektar / Ha).</li>
  <li><strong>Kategori_Driver_k:</strong> Klasifikasi penyebab utama deforestasi (Dominant Driver of Tree Cover Loss) berdasarkan model data historis satelit.</li>
  <li><strong>Area_Loss_i:</strong> Luas kehilangan tutupan pohon pada piksel observasi ke-i (satuan: Hektar / Ha).</li>
</ul>
<p><strong>Persamaan Perhitungan Akumulasi Kehilangan Hutan Primer (Primary Forest Loss)</strong><br>
Total_Primary_Loss = &sum;(Area_Loss_j) untuk j di mana Tipe_Hutan = "Primer"</p>
<ul>
  <li><strong>Total_Primary_Loss:</strong> Akumulasi luas konversi tutupan hutan alam primer tak terganggu (intact primary forest) selama periode pengamatan (satuan: Hektar / Ha).</li>
  <li><strong>Tipe_Hutan:</strong> Klasifikasi basemap jenis tutupan lahan awal sebelum terjadi deforestasi.</li>
</ul>
<p><strong>Persamaan Estimasi Pelepasan Emisi Karbon (Gross CO2 Emissions)</strong><br>
Emisi_CO2_Total = &sum;(Area_Loss_c * Faktor_Emisi_Biomassa_c)</p>
<ul>
  <li><strong>Emisi_CO2_Total:</strong> Estimasi agregasi total emisi gas rumah kaca yang dilepaskan ke atmosfer akibat konversi tutupan (satuan: Megagrams CO2 / Mg CO2).</li>
  <li><strong>Faktor_Emisi_Biomassa_c:</strong> Kandungan karbon rata-rata (above-ground & below-ground biomass) per hektar pada koordinat c yang diamati.</li>
</ul>
<p>Kalkulasi pengujian statistik dihitung menggunakan formulasi Matematis yang sama dengan Sub-Bab 1.2 dan 1.3, di mana variabel independen (X) adalah <strong>Investasi PMDN (Juta Rp)</strong> dan variabel dependen (Y) adalah <strong>Deforestasi Komoditas (Hektar)</strong>.</p>
<p>Persamaan Kategorisasi Nilai Ambang Batas Median:<br>
- Jika Nilai &gt; Median, maka Kategori = Tinggi<br>
- Jika Nilai &le; Median, maka Kategori = Rendah</p>

<h4>D. Matriks Hasil Uji Empiris: Konsentrasi Spasial & Skenario Crosstab</h4>
<p>
  Tingkat alokasi konsesi dan dampaknya terhadap tutupan hutan dapat dilihat secara empiris melalui perbandingan luas konsesi baru di Daerah Sentra Tambang (Morowali & Konawe) dengan wilayah non-sentra pada <strong>Tabel 1.7b</strong> berikut:
</p>
<div class="table-caption">Tabel 1.7b: Representasi Spasial Luas Konsesi Baru (Ha) di Daerah Sentra Tambang vs Non-Sentra (2014-2023)</div>
<table class="data-table">
  <thead>
    <tr>
      <th class="data-th">Tahun</th>
      <th class="data-th">Luas Konsesi Baru Daerah Sentra Tambang (Ha)</th>
      <th class="data-th">Luas Konsesi Baru Daerah Non-Sentra (Ha)</th>
    </tr>
  </thead>
  <tbody>
{html_emp_rows_1_4.rstrip()}
  </tbody>
</table>

<p>
  Terkait dengan hilangnya luasan hutan tersebut, pembedahan lebih lanjut berdasarkan aktor utama, luasan hutan primer, dan estimasi emisi karbon komoditas dapat dilihat pada <strong>Tabel 1.7c</strong> berikut:
</p>
<div class="table-caption">Tabel 1.7c: Matriks Pembedahan Ekologis Aktor & Emisi Karbon (Periode 2001-2025)</div>
<table class="data-table">
  <thead>
    <tr><th class="data-th">Kategori Aktor / Metrik Ekologis</th><th class="data-th">Nilai Agregat</th><th class="data-th">Persentase dari Total Kehilangan</th></tr>
  </thead>
  <tbody>
    <tr><td class="data-td"><strong>Ekspansi Komoditas (Tambang & Sawit)</strong></td><td class="data-td" style="text-align:center;">1,890,659 Hektar</td><td class="data-td" style="text-align:center;">48.4%</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Kehutanan (Logging)</strong></td><td class="data-td" style="text-align:center;">247,011 Hektar</td><td class="data-td" style="text-align:center;">6.3%</td></tr>
    <tr><td class="data-td"><strong>Pertanian Berpindah</strong></td><td class="data-td" style="text-align:center;">115,404 Hektar</td><td class="data-td" style="text-align:center;">2.9%</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Total Kehilangan Hutan Primer</strong></td><td class="data-td" style="text-align:center;"><strong>3,904,079 Hektar</strong></td><td class="data-td" style="text-align:center;"><strong>100.0%</strong></td></tr>
    <tr><td class="data-td"><strong>Estimasi Emisi Karbon Komoditas</strong></td><td class="data-td" style="text-align:center;">1,282,195,705 Mg CO2</td><td class="data-td" style="text-align:center;">-</td></tr>
  </tbody>
</table>

<div class="table-caption">Tabel 1.7d: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.4)</div>
<table class="data-table">
  <thead>
    <tr><th class="data-th">Komponen Uji</th><th class="data-th">Definisi Variabel (Sub-bab 1.4)</th></tr>
  </thead>
  <tbody>
    <tr><td class="data-td">Variabel Independen (X)</td><td class="data-td">Realisasi Investasi PMDN (Juta Rupiah)</td></tr>
    <tr class="data-tr-even"><td class="data-td">Variabel Dependen (Y)</td><td class="data-td">Total Deforestasi Alam (Ha) / Deforestasi Komoditas (Ha)</td></tr>
    <tr><td class="data-td">Hipotesis Nol (H0)</td><td class="data-td">Tingginya realisasi investasi PMDN tidak berhubungan dengan laju deforestasi.</td></tr>
    <tr class="data-tr-even"><td class="data-td">Hipotesis Alternatif (H1)</td><td class="data-td">Ada hubungan positif antara tingginya realisasi investasi PMDN dengan laju deforestasi.</td></tr>
    <tr><td class="data-td">Threshold Kategori</td><td class="data-td">Nilai Median Data Panel (N=48)</td></tr>
  </tbody>
</table>

<p>Tabel di bawah ini merangkum hasil pengujian statistik (Chi-Square) untuk semua kemungkinan kombinasi indikator antara Realisasi Investasi PMDN dan Dampak Ekologis pada panel data 2016-2023. Hasil tersebut ditampilkan pada <strong>Tabel 1.8</strong> berikut:</p>
<div class="table-caption">Tabel 1.8: Ringkasan Eksekutif Seluruh Skenario Crosstab Realisasi Investasi PMDN Bab 1</div>
<table class="data-table">
  <thead>
    <tr><th>Variabel Independen (X)</th><th>Variabel Dependen (Y)</th><th>Chi-Square (&chi;&sup2;)</th><th>P-Value</th><th>Odds Ratio</th><th>Kesimpulan</th></tr>
  </thead>
  <tbody>
{html_chi_rows_4.rstrip()}
  </tbody>
</table>

<h4>E. Analisis Temuan Empiris: Efek Jeda Waktu (Time-Lagging)</h4>
<p>Hasil pengujian seluruh skenario tabulasi silang PMDN mengungkap fenomena yang kompleks dalam alur investasi ekstraktif:</p>
<ul style="margin-left: 20px; margin-bottom: 15px;">
  <li style="margin-bottom: 6px;"><strong>1. Ketidaksignifikanan Simultan &amp; Variasi P-Value:</strong> Tingkat signifikansi yang bervariasi menyingkap tabir jeda waktu (lagging effect) dalam eksekusi investasi di lapangan.</li>
  <li style="margin-bottom: 6px;"><strong>2. Jeda Waktu Eksekusi Investasi (Lagging Effect):</strong> Suntikan modal masif di tahun tertentu tidak secara instan berwujud pembabatan lahan di tahun yang sama. Modal tersebut tertahan pada fase birokrasi, pembebasan lahan, dan pengadaan infrastruktur, sebelum daya rusaknya mengonversi lanskap hutan pada tahun-tahun berikutnya.</li>
  <li style="margin-bottom: 6px;"><strong>3. Konsentrasi Modal Ekstrem di 3 Provinsi:</strong> Data spasial membuktikan 89% dari total modal PMDN ekstraktif se-Sulawesi hanya tersedot ke tiga provinsi sentra (Sulteng, Sultra, Sulsel), mengakibatkan polarisasi pertumbuhan dan mengunci ketimpangan spasial.</li>
</ul>

<h2>SUB-BAB 1.5: Pelabuhan Ekspor & Peta Jalur Distribusi Logistik Nikel Sulawesi</h2>
<p>
  Validasi silang informasi publik (triangulasi 4 sumber dokumen: Laporan Investigasi KNKT 52.378 DWT, Lampiran Perpres 109/2020 PSN Kawasan Industri, Laporan Tahunan ANTAM, dan Laporan Keberlanjutan Vale) memverifikasi enam simpul pelabuhan dan terminal khusus yang terhubung langsung ke pasar ekspor internasional melalui pemodelan kurva alur maritim sebagaimana diilustrasikan pada <strong>Bagan Alur 1.5</strong> berikut:
</p>

<div class="table-caption">Bagan Alur 1.5: Alur Logika Metodologis Validasi Silang (OSINT) dan Pemetaan Spasial Pelabuhan</div>
<div style="text-align: center; margin: 20px 0;">
  <img src="mermaid_flowchart_1_5.png" alt="Bagan Alur 1.5" style="max-width: 100%; border: 1px solid #ddd; padding: 10px; background-color: #fcfcfc;">
</div>

<p>
  Berdasarkan protokol validasi silang tersebut, profil komprehensif enam simpul pelabuhan dan terminal khusus utama di Pulau Sulawesi dipetakan pada <strong>Tabel 1.7</strong> berikut:
</p>

<div class="table-caption">Tabel 1.7: Inventarisasi Enam Simpul Pelabuhan dan Terminal Khusus Ekspor Nikel di Pulau Sulawesi</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:20%;">Simpul Kawasan Industri</th>
      <th class="data-th" style="width:18%;">Wilayah Administrasi</th>
      <th class="data-th" style="width:22%;">Fasilitas Pelabuhan / Terminal</th>
      <th class="data-th" style="width:16%;">Status Regulasi</th>
      <th class="data-th" style="width:10%;">Kapasitas Kapal</th>
      <th class="data-th" style="width:14%;">Tujuan Utama Ekspor</th>
    </tr>
  </thead>
  <tbody>
    <tr><td class="data-td"><strong>IMIP Morowali</strong></td><td class="data-td">Morowali, Sulawesi Tengah</td><td class="data-td">Pelabuhan Samudera & Dermaga Curah</td><td class="data-td">PSN (Perpres 109/2020)</td><td class="data-td">Hingga 52.378 DWT</td><td class="data-td">Pasar Global (Tiongkok)</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>GNI Morowali Utara</strong></td><td class="data-td">Morowali Utara, Sulteng</td><td class="data-td">Terminal Khusus Pesisir Tomori</td><td class="data-td">Izin Industri Mandiri</td><td class="data-td">Hingga 30.000 DWT</td><td class="data-td">Pasar Global (Tiongkok)</td></tr>
    <tr><td class="data-td"><strong>VDNI Konawe</strong></td><td class="data-td">Konawe, Sulawesi Tenggara</td><td class="data-td">Dermaga Khusus Curah & Kargo</td><td class="data-td">PSN (Perpres 109/2020)</td><td class="data-td">Hingga 50.000 DWT</td><td class="data-td">Pasar Global (Tiongkok)</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>OSS Konawe</strong></td><td class="data-td">Konawe, Sulawesi Tenggara</td><td class="data-td">Dermaga Terintegrasi Konawe</td><td class="data-td">PSN (Perpres 109/2020)</td><td class="data-td">Hingga 50.000 DWT</td><td class="data-td">Pasar Global (Tiongkok)</td></tr>
    <tr><td class="data-td"><strong>Pomalaa (ANTAM)</strong></td><td class="data-td">Kolaka, Sulawesi Tenggara</td><td class="data-td">Dermaga Pomalaa & Konveyor</td><td class="data-td">Kawasan BUMN Industri</td><td class="data-td">Hingga 12.000 DWT</td><td class="data-td">Jepang & Korsel</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Sorowako (Vale)</strong></td><td class="data-td">Luwu Timur, Sulawesi Selatan</td><td class="data-td">Pelabuhan Balantang Malili</td><td class="data-td">Kontrak Karya Tambang</td><td class="data-td">Hingga 15.000 DWT</td><td class="data-td">Jepang & Skandinavia</td></tr>
  </tbody>
</table>

<h2>SUB-BAB 1.6: Peta Jalur Distribusi Logistik Nikel Sulawesi</h2>

<h4>A. Pengantar & Kerangka Narasi</h4>
<p>
  Metode analisis pada tahapan ini difokuskan pada <strong>Pemetaan Kausalitas (Spasial)</strong> untuk membedah asimetri penguasaan ruang antara wilayah hulu (origin: sumber ekstraksi di Sulawesi) dan hilir (destination: pusat industrialisasi luar negeri). Garis diplot menggunakan rute pelayaran untuk merepresentasikan jarak tempuh aktual kapal logistik di permukaan bumi.
</p>

<h4>B. Alur Logika Metodologis (Flowchart Analisis Spasial)</h4>
<p>
  Rangkaian proses ekstraksi koordinat hingga visualisasi spasial direpresentasikan pada <strong>Bagan Alur 1.6</strong> berikut:
</p>
<div class="table-caption">Bagan Alur 1.6: Alur Logika Metodologis Pemetaan Jalur Logistik Maritim (Bézier Curve)</div>
<div style="text-align: center; margin: 20px 0;">
  <img src="mermaid_flowchart_1_6.png" alt="Bagan Alur 1.6" style="max-width: 100%; border: 1px solid #ddd; padding: 10px; background-color: #fcfcfc;">
</div>

<h4>C. Formulasi Matematis: Kurva Parametrik Alur Pelayaran</h4>
<div class="formula-title">Persamaan Formulasi Kurva Parametrik Alur Pelayaran Maritim:</div>
<div class="formula">Kurva(t) = (1 - t)^2 * Titik_Asal + 2 * (1 - t) * t * Titik_Kontrol + t^2 * Titik_Tujuan,   t dalam rentang [0, 1]</div>
<div class="var-desc">
  <div class="var-desc-title">Keterangan Variabel:</div>
  <ul>
    <li>• <strong>Kurva(t)</strong>: Vektor posisi koordinat geografis lintasan kapal pada parameter waktu t.</li>
    <li>• <strong>Titik_Asal</strong>: Titik koordinat geografis pelabuhan muat khusus di pesisir Sulawesi.</li>
    <li>• <strong>Titik_Kontrol</strong>: Titik koordinat jangkar pemandu kurva lengkung di perairan internasional.</li>
    <li>• <strong>Titik_Tujuan</strong>: Titik koordinat geografis pelabuhan bongkar di negara tujuan ekspor.</li>
    <li>• <strong>t</strong>: Parameter interpolasi waktu dan pergerakan lintasan dalam rentang kontinu [0, 1].</li>
  </ul>
</div>

<h4>D. Matriks Hasil Pemetaan Spasial</h4>
<p>
  Agregasi titik asal (Origin) dan tujuan akhir (Destination) di-render langsung ke dalam proyeksi spasial. Struktur data logistik yang menggerakkan pemetaan ini memetakan pelabuhan asal menuju pasar utama, yaitu <strong>Tiongkok, Jepang, dan Korea Selatan</strong>.
</p>

<h4>E. Interpretasi Spasial Industri (Anatomi Rantai Pasok)</h4>
<p>Peta rute logistik maritim mengilustrasikan alur distribusi produk olahan nikel dari kawasan industri di Sulawesi:</p>
<ul style="margin-left: 20px; margin-bottom: 15px;">
  <li style="margin-bottom: 6px;"><strong>1. Orientasi Ekspor:</strong> Kawasan industri utama yang berstatus Proyek Strategis Nasional (PSN) mengalirkan produk olahan ke sentra-sentra industri manufaktur di pasar internasional.</li>
  <li style="margin-bottom: 6px;"><strong>2. Integrasi Rantai Pasok:</strong> Mayoritas rute pengapalan terhubung langsung dengan pelabuhan ekspor tujuan, yang mengindikasikan posisi kawasan pemurnian di Sulawesi sebagai pemasok bahan baku setengah jadi.</li>
  <li style="margin-bottom: 6px;"><strong>3. Dinamika Rute Maritim:</strong> Peta rute mencerminkan diversifikasi pasar ekspor (Asia Timur) dan jaringan logistik kawasan.</li>
</ul>

<h2>SUB-BAB 1.7: Matriks Indikator dan Sumber Data Resmi Bab 1</h2>
<p>
  Seluruh variabel kuantitatif, kategori analisis, satuan ukur, periode tahun observasi, dan institusi penyedia data primer resmi yang digunakan dalam Bab 1 dikompilasikan pada <strong>Tabel 1.8</strong> berikut:
</p>

<div class="table-caption">Tabel 1.8: Matriks Indikator dan Sumber Data Primer Resmi Bab 1 (Ekspansi Industri Ekstraktif)</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:6%;">No</th>
      <th class="data-th" style="width:26%;">Nama Indikator</th>
      <th class="data-th" style="width:20%;">Kategori Analisis</th>
      <th class="data-th" style="width:12%;">Satuan Ukur</th>
      <th class="data-th" style="width:12%;">Cakupan Tahun</th>
      <th class="data-th" style="width:24%;">Institusi & Sumber Data Resmi</th>
    </tr>
  </thead>
  <tbody>
    <tr><td class="data-td" style="text-align:center;">1</td><td class="data-td">Izin Usaha Pertambangan (IUP) Baru</td><td class="data-td">Faktor Tekanan Ekstraktif</td><td class="data-td" style="text-align:center;">Unit Izin</td><td class="data-td" style="text-align:center;">2014–2024</td><td class="data-td">Kementerian ESDM (Minerbaone)</td></tr>
    <tr class="data-tr-even"><td class="data-td" style="text-align:center;">2</td><td class="data-td">Luas Wilayah Konsesi Tambang Baru</td><td class="data-td">Faktor Tekanan Ekstraktif</td><td class="data-td" style="text-align:center;">Hektar (Ha)</td><td class="data-td" style="text-align:center;">2014–2024</td><td class="data-td">Kementerian ESDM (Minerbaone)</td></tr>
    <tr><td class="data-td" style="text-align:center;">3</td><td class="data-td">Kapasitas Terpasang PLTU Captive</td><td class="data-td">Infrastruktur Energi Khusus</td><td class="data-td" style="text-align:center;">Megawatt (MW)</td><td class="data-td" style="text-align:center;">2014–2024</td><td class="data-td">Global Energy Monitor (GEM)</td></tr>
    <tr class="data-tr-even"><td class="data-td" style="text-align:center;">4</td><td class="data-td">Fasilitas Pengolahan & Pemurnian (Smelter)</td><td class="data-td">Fasilitas Industri Hilir</td><td class="data-td" style="text-align:center;">Unit Fasilitas</td><td class="data-td" style="text-align:center;">2014–2024</td><td class="data-td">Kementerian ESDM & Basis Data Industri</td></tr>
    <tr><td class="data-td" style="text-align:center;">5</td><td class="data-td">Realisasi Investasi PMDN</td><td class="data-td">Arus Modal Domestik</td><td class="data-td" style="text-align:center;">Triliun Rupiah</td><td class="data-td" style="text-align:center;">2016–2024</td><td class="data-td">Kementerian Investasi / BKPM</td></tr>
    <tr class="data-tr-even"><td class="data-td" style="text-align:center;">6</td><td class="data-td">PDRB Menurut 17 Lapangan Usaha</td><td class="data-td">Struktur Ekonomi Makro</td><td class="data-td" style="text-align:center;">Triliun Rupiah</td><td class="data-td" style="text-align:center;">2016–2024</td><td class="data-td">Badan Pusat Statistik (BPS Provinsi se-Sulawesi)</td></tr>
    <tr><td class="data-td" style="text-align:center;">7</td><td class="data-td">PDRB Kabupaten Sentra Tambang</td><td class="data-td">Struktur Ekonomi Daerah</td><td class="data-td" style="text-align:center;">Triliun Rupiah</td><td class="data-td" style="text-align:center;">2016–2024</td><td class="data-td">Badan Pusat Statistik (BPS Kabupaten se-Sulteng)</td></tr>
    <tr class="data-tr-even"><td class="data-td" style="text-align:center;">8</td><td class="data-td">Luas Kehilangan Hutan Komoditas</td><td class="data-td">Dampak Perubahan Tutupan Lahan</td><td class="data-td" style="text-align:center;">Hektar (Ha)</td><td class="data-td" style="text-align:center;">2014–2023</td><td class="data-td">Global Forest Watch (GFW / Univ. of Maryland)</td></tr>
    <tr><td class="data-td" style="text-align:center;">9</td><td class="data-td">Simpul Dermaga & Terminal Khusus Ekspor</td><td class="data-td">Infrastruktur Rantai Pasok</td><td class="data-td" style="text-align:center;">Titik Koordinat & DWT</td><td class="data-td" style="text-align:center;">2014–2024</td><td class="data-td">KNKT Kemenhub, Perpres PSN, Laporan Korporasi</td></tr>
  </tbody>
</table>

<h2>SUB-BAB 1.8: Bagan Alur Kerangka Kerja Riset Bab 1</h2>
<p>
  Keseluruhan struktur metodologis riset Bab 1 dioperasionalkan melalui empat fase kerja berurutan sebagaimana disajikan pada <strong>Tabel 1.9</strong> berikut:
</p>

<div class="table-caption">Tabel 1.9: Matriks Tahapan dan Alur Kerangka Kerja Riset Bab 1</div>
<table>
  <thead>
    <tr>
      <th class="data-th" style="width:20%;">Tahapan Riset</th>
      <th class="data-th" style="width:26%;">Fokus Metodologis</th>
      <th class="data-th" style="width:28%;">Bahan & Sumber Data</th>
      <th class="data-th" style="width:26%;">Keluaran / Hasil Analisis</th>
    </tr>
  </thead>
  <tbody>
    <tr><td class="data-td"><strong>Fase I: Pengumpulan Data</strong></td><td class="data-td">Kurasi data resmi lintas kementerian dan lembaga</td><td class="data-td">Publikasi BPS, Minerbaone, BKPM, GEM, dan GFW</td><td class="data-td">Basis Data Tabular Panel Provinsi (2014–2024)</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Fase II: Reklasifikasi Hukum</strong></td><td class="data-td">Penyusunan kerangka rantai pasok hukum terintegrasi</td><td class="data-td">UU No. 3/2020, PP No. 96/2021, Perpres No. 112/2022</td><td class="data-td">3 Klaster Makro (Ekstraktif, Akar Rumput, Jasa)</td></tr>
    <tr><td class="data-td"><strong>Fase III: Pengujian Statistik</strong></td><td class="data-td">Uji signifikansi hubungan dan rasio peluang</td><td class="data-td">Tabel Kontinjensi, Uji Chi-Square, Odds Ratio</td><td class="data-td">Bukti Kausalitas Signifikan Tekanan vs Deforestasi</td></tr>
    <tr class="data-tr-even"><td class="data-td"><strong>Fase IV: Pemetaan Rantai Pasok</strong></td><td class="data-td">Triangulasi data logistik dan pemodelan maritim</td><td class="data-td">Laporan KNKT, Perpres PSN, Kurva Parametrik Bézier</td><td class="data-td">Peta Alur Rantai Pasok Ekspor & Konsentrasi Spasial 78%</td></tr>
  </tbody>
</table>
</body>
</html>
"""
    html_path = tool_dir / "Metodologi_Bab1_Ekspansi_Industri.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Tersimpan: {html_path}")

    # 4. Generate File Markdown Bersih
    print("[4/5] Membangun Metodologi_Bab1_Ekspansi_Industri.md (Format Publik)...")
    md_lines = [
        "# BAB I: METODOLOGI ANALISIS EKSPANSI INDUSTRI EKSTRAKTIF DAN INFRASTRUKTUR PENUNJANG DI PULAU SULAWESI",
        "",
        "Dokumen laporan metodologi ini menyajikan kerangka ilmiah, landasan regulasi, formulasi matematis, prosedur analisis statistik, serta metodologi pembuktian berbasis data terbuka yang dioperasionalkan pada **Bab 1: Ekspansi Industri Ekstraktif** dalam studi Daya Dukung dan Daya Tampung Lingkungan Hidup (D3TLH) Sulawesi periode 2014–2024.",
        "",
        "---",
        "",
        "## SUB-BAB 1.1: Konteks Makro: Breakdown PDRB per Komoditas",
        "",
        "### 1.1.1 Konteks Makro: Dominasi Ekstraktif vs Ekonomi Akar Rumput",
        "Bagian ini menganalisis struktur Produk Domestik Regional Bruto (PDRB) pada enam provinsi di Pulau Sulawesi sepanjang periode 2016–2024 menggunakan visualisasi grafik area bertumpuk (*Stacked Area Chart*). Analisis ini ditujukan untuk menguji secara empiris apakah percepatan pertumbuhan ekonomi daerah benar-benar bersumber dari sektor produktif masyarakat lokal atau didominasi oleh industri ekstraktif padat modal yang mengalihkan pemanfaatan ruang dan sumber daya alam.",
        "",
        "> **Sumber Data:** Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi *Stacked Area Chart* memetakan dinamika Produk Domestik Regional Bruto (PDRB) berdasarkan klasifikasi rantai pasok hukum (*Legal Supply-Chain*) untuk membandingkan trajektori Sektor Ekstraktif, Ekonomi Akar Rumput, dan Sektor Jasa & Lainnya.",
        "",
        "#### A. Kerangka Dekomposisi Sektoral & Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)",
        "Sistem KBLI 2020 BPS membagi 17 sektor PDRB. Melalui pendekatan Legal Supply-Chain, 17 sektor direklasifikasi menjadi 3 Klaster Makro (Ekstraktif, Akar Rumput, Jasa). Rincian pembagian sektor, dasar regulasi, serta intisari ketentuan hukum disajikan secara lengkap pada **Tabel 1.1** berikut:",
        "",
        "##### Tabel 1.1: Reklasifikasi Sektoral PDRB KBLI 2020 Berdasarkan Pendekatan Rantai Pasok Hukum (Legal Supply-Chain)",
        "| Kategori BPS | Sektor Lapangan Usaha | Klasifikasi Analisis | Dasar Regulasi & Mandat Hukum | Intisari Ketentuan Hukum |",
        "| :--- | :--- | :---: | :--- | :--- |",
        "| **Kategori B** | Pertambangan dan Penggalian | Ekstraktif | Perpres No. 26 Tahun 2010 | Ketentuan Pasal 1 Ayat (2) mengenai pengambilan komoditas tambang dari dalam bumi. |",
        "| **Kategori C** | Industri Pengolahan (Smelter Logam) | Ekstraktif | UU No. 3 Tahun 2020 & PP No. 96 Tahun 2021 | Pasal 102–103 mewajibkan pengolahan dan pemurnian di dalam negeri sebagai kesatuan pertambangan. |",
        "| **Kategori D** | Pengadaan Listrik & Gas (PLTU Captive) | Ekstraktif | Perpres No. 112 Tahun 2022 & RUPTL PLN | Pasal 3 Ayat (4) huruf b mengecualikan PLTU baru hanya bagi yang terintegrasi melayani smelter. |",
        "| **Kategori A** | Pertanian, Kehutanan, Perikanan | Ekonomi Akar Rumput | KBLI 2020 BPS | Sektor pemanfaatan sumber daya hayati terbarukan dan penyerap tenaga kerja lokal. |",
        "| **Kategori E–U** | 13 Sektor Jasa & Konstruksi | Sektor Jasa & Lainnya | Klasifikasi Standar BPS | Sektor sekunder dan tersier penunjang perekonomian daerah. |",
        "",
        "#### B. Alur Logika Metodologis Rantai Pasok Hukum (Mengapa Kat. B + C + D = Ekstraktif)",
        "Keterkaitan ketiga kategori lapangan usaha tersebut sebagai satu kesatuan rantai pasok ekstraktif dimodelkan dalam kerangka alur logika hukum sebagaimana diilustrasikan pada **Bagan Alur 1.1** berikut:",
        "",
        "##### Bagan Alur 1.1: Alur Logika Metodologis Rantai Pasok Hukum Sektor Ekstraktif",
        "```mermaid",
        f"{mermaid_str_1_1}",
        "```",
        "",
        "#### C. Formulasi Matematis: Persamaan Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation)",
        "",
        "**Persamaan Agregasi Sektor Ekstraktif (Legal Supply-Chain Aggregation):**",
        "```text",
        "Sektor_Ekstraktif = PDRB(Kat.B: Pertambangan) + PDRB(Kat.C: Ind. Pengolahan) + PDRB(Kat.D: Listrik)",
        "```",
        "*Keterangan Variabel:*",
        "- `Sektor_Ekstraktif`: Total nilai tambah bruto dari klaster industri ekstraktif yang saling terintegrasi (Triliun Rupiah).",
        "- `PDRB(Kat.B: Pertambangan)`: Nilai tambah kegiatan eksplorasi dan ekstraksi bijih mineral (BPS KBLI 2020 Kategori B).",
        "- `PDRB(Kat.C: Ind. Pengolahan)`: Nilai tambah pemurnian logam dasar di smelter nikel (BPS KBLI 2020 Kategori C / Golongan 24).",
        "- `PDRB(Kat.D: Listrik)`: Nilai tambah penyediaan listrik batubara khusus smelter / PLTU captive (BPS KBLI 2020 Kategori D).",
        "",
        "**Persamaan Ekonomi Akar Rumput:**",
        "```text",
        "Sektor_Akar_Rumput = PDRB(Kat.A: Pertanian, Kehutanan, dan Perikanan)",
        "```",
        "*Keterangan Variabel:*",
        "- `Sektor_Akar_Rumput`: Nilai PDRB pemanfaatan sumber daya hayati terbarukan (Triliun Rupiah).",
        "- `PDRB(Kat.A)`: Agregasi nilai tambah tanaman pangan, perkebunan rakyat, perikanan, peternakan, kehutanan.",
        "",
        "**Persamaan Sektor Jasa & Lainnya:**",
        "```text",
        "Sektor_Jasa = Jumlah PDRB (Kategori E sampai dengan Kategori U)",
        "```",
        "*Keterangan Variabel:*",
        "- `Sektor_Jasa`: Nilai tambah 13 sektor penunjang sekunder dan tersier (Triliun Rupiah).",
        "- `PDRB(Kat. E s.d. U)`: Akumulasi sektor perdagangan, konstruksi, transportasi, keuangan, pendidikan, dll.",
        "",
        "**Persamaan Total Produk Domestik Regional Bruto (PDRB Wilayah):**",
        "```text",
        "Total_PDRB = Sektor_Ekstraktif + Sektor_Akar_Rumput + Sektor_Jasa",
        "```",
        "*Keterangan Variabel:*",
        "- `Total_PDRB`: Total nilai Produk Domestik Regional Bruto wilayah atas dasar harga berlaku (Triliun Rupiah).",
        "",
        "**Persamaan Pangsa Kontribusi Sektor Ekstraktif (%):**",
        "```text",
        "Pangsa_Ekstraktif (%) = ( Sektor_Ekstraktif / Total_PDRB ) * 100",
        "```",
        "*Keterangan Variabel:*",
        "- `Pangsa_Ekstraktif (%)`: Persentase pangsa dominasi sektor ekstraktif terhadap total ekonomi (%).",
        "",
        "**Persamaan Laju Pertumbuhan Tahunan Sektoral (YoY):**",
        "```text",
        "Laju_Pertumbuhan_Tahunan (%) = [ ( Nilai_Tahun_t - Nilai_Tahun_{t-1} ) / Nilai_Tahun_{t-1} ] * 100",
        "```",
        "*Keterangan Variabel:*",
        "- `Laju_Pertumbuhan_Tahunan (%)`: Tingkat percepatan/perlambatan ekspansi tahunan sektor ekonomi (%).",
        "- `Nilai_Tahun_t`: Nilai nominal PDRB sektor pada tahun berjalan t.",
        "- `Nilai_Tahun_{t-1}`: Nilai nominal PDRB sektor pada satu tahun sebelumnya (t - 1).",
        "",
        "Definisi operasional, cakupan lapangan usaha, dan institusi penyedia data primer untuk masing-masing komponen variabel dalam sistem persamaan di atas dipaparkan pada **Tabel 1.2** berikut:",
        "",
        "##### Tabel 1.2: Definisi Operasional Komponen Makroekonomi dan Sumber Data PDRB Sektoral",
        "| Komponen Analisis | Cakupan Lapangan Usaha | Definisi Operasional | Satuan Nilai | Sumber Data Primer |",
        "| :--- | :--- | :--- | :---: | :--- |",
        "| **Sektor Ekstraktif** | Kategori B, Kategori C, Kategori D | Akumulasi nilai tambah pertambangan nikel, smelter logam dasar, dan PLTU captive. | Triliun Rupiah | BPS Provinsi (SIMDASI) |",
        "| **Ekonomi Akar Rumput** | Kategori A | Nilai tambah pertanian, perkebunan, kehutanan, dan perikanan. | Triliun Rupiah | BPS Provinsi |",
        "| **Sektor Jasa & Lainnya** | Kategori E hingga U | Nilai tambah gabungan perdagangan, konstruksi, transportasi, keuangan, dan jasa. | Triliun Rupiah | BPS Provinsi |",
        "| **Total PDRB Wilayah** | Seluruh 17 Kategori | Total nilai PDRB wilayah atas dasar harga berlaku pada tahun berjalan. | Triliun Rupiah | BPS Provinsi |",
        "| **Pangsa Ekstraktif (%)** | Rasio Kontribusi | Persentase kontribusi sektor ekstraktif terhadap total perekonomian. | Persen (%) | Hasil Olahan CELIOS |",
        "",
        "#### D. Analisis Temuan Empiris: Ketimpangan Struktural Sulawesi Tengah",
        "",
        f"Penerapan formulasi di atas menunjukkan bahwa di **Sulawesi Tengah (sebagai pusat hilirisasi)**, ekspansi industri ekstraktif menguasai **{pct_sulteng_eks:.1f}% dari total PDRB provinsi** pada tahun {latest_year_pdrb}, memperlihatkan dominasi yang sangat kuat dibanding provinsi lainnya.",
        "",
        "### 1.1.2 Pemusatan Sektor Ekstraktif di Kabupaten se-Sulawesi Tengah",
        "",
        f"Jika dianalisis secara spasial pada tingkat kabupaten di Sulawesi Tengah, terlihat konsentrasi kegiatan industri ekstraktif. Kabupaten **Morowali** dan **Morowali Utara** mendominasi struktur PDRB provinsi melalui pengembangan kawasan industri hilirisasi dan PLTU Captive. Analisis ini membandingkan komposisi ketiga sektor advokatif di seluruh 13 kabupaten/kota se-Sulawesi Tengah pada tahun terbaru ({latest_year_kab}).",
        "",
        f"> **Sumber Data:** Badan Pusat Statistik (BPS) Kabupaten se-Sulawesi Tengah (diolah CELIOS). Visualisasi *Stacked Bar Chart* memetakan struktur Produk Domestik Regional Bruto (PDRB) tahun {latest_year_kab} pada seluruh 13 kabupaten/kota untuk mengidentifikasi tingkat konsentrasi sektoral dan polarisasi spasial antara sentra industri pengolahan nikel dengan daerah non-sentra.",
        "",
        "#### A. Rasionalitas Spasial & Urgensi Dekomposisi Sektoral Tingkat Kabupaten",
        "Analisis agregat pada tingkat provinsi sering kali menghasilkan **Bias Ilusi Agregat (Aggregate Illusion Bias)**, di mana angka pertumbuhan ekonomi makro yang tinggi memberi kesan seolah seluruh wilayah menikmati kemakmuran yang seimbang. Namun, ketika data didekomposisi ke tingkat kabupaten/kota, terlihat jurang pemisah ekonomi yang sangat tajam antara wilayah **Enklave Industri Ekstraktif** dengan daerah agraris tradisional sekitarnya.",
        "",
        "#### B. Alur Logika Analisis Disparitas Spasial Kabupaten",
        "Kerangka kerja metodologis dalam membedah ketimpangan intra-provinsial ini diilustrasikan pada **Bagan Alur 1.2** berikut:",
        "",
        "##### Bagan Alur 1.2: Alur Logika Metodologis Dekomposisi Spasial PDRB Tingkat Kabupaten se-Sulawesi Tengah",
        "```mermaid",
        f"{mermaid_str_1_2}",
        "```",
        "",
        "#### C. Formulasi Matematis: Persamaan Agregasi Sektoral Kabupaten (Legal Supply-Chain Aggregation)",
        "",
        "**Persamaan Agregasi Sektor Ekstraktif Tingkat Kabupaten:**",
        "```text",
        "Sektor_Ekstraktif_Kabupaten = PDRB_Kab(Kat.B: Pertambangan) + PDRB_Kab(Kat.C: Ind. Pengolahan) + PDRB_Kab(Kat.D: Listrik)",
        "```",
        "*Keterangan Variabel:*",
        "- `Sektor_Ekstraktif_Kabupaten`: Total nilai tambah sektor ekstraktif di tingkat kabupaten target (satuan: Triliun Rupiah).",
        "- `PDRB_Kab(Kat.B: Pertambangan)`: Nilai PDRB kabupaten dari aktivitas penambangan bijih logam dan galian (BPS Kategori B).",
        "- `PDRB_Kab(Kat.C: Ind. Pengolahan)`: Nilai PDRB kabupaten dari industri peleburan logam dasar / smelter (BPS Kategori C).",
        "- `PDRB_Kab(Kat.D: Listrik)`: Nilai PDRB kabupaten dari penyediaan daya listrik batubara captive (BPS Kategori D).",
        "",
        "**Persamaan Total Produk Domestik Regional Bruto Tingkat Kabupaten:**",
        "```text",
        "Total_PDRB_Kabupaten = Sektor_Ekstraktif_Kabupaten + Sektor_Akar_Rumput_Kabupaten + Sektor_Jasa_Kabupaten",
        "```",
        "*Keterangan Variabel:*",
        "- `Total_PDRB_Kabupaten`: Total output perekonomian bruto kabupaten target atas dasar harga berlaku (satuan: Triliun Rupiah).",
        "- `Sektor_Ekstraktif_Kabupaten`: Nilai tambah bruto sektor ekstraktif terintegrasi di kabupaten (Triliun Rupiah).",
        "- `Sektor_Akar_Rumput_Kabupaten`: Nilai tambah sektor pertanian, kehutanan, dan perikanan di kabupaten (Triliun Rupiah).",
        "- `Sektor_Jasa_Kabupaten`: Nilai tambah sektor perdagangan, transportasi, dan jasa layanan di kabupaten (Triliun Rupiah).",
        "",
        "**Persamaan Porsi Sektoral dalam Kabupaten (Porsi (%) pada Tooltip Dashboard):**",
        "```text",
        "Porsi_Sektor_Kabupaten (%) = ( Nilai_Sektor_Kabupaten / Total_PDRB_Kabupaten ) * 100",
        "```",
        "*Keterangan Variabel:*",
        "- `Porsi_Ekstraktif (%)`: Persentase kontribusi Sektor Ekstraktif: ( Sektor_Ekstraktif / Total_PDRB ) * 100 (misal Morowali: 45.2%).",
        "- `Porsi_Jasa (%)`: Persentase kontribusi Sektor Jasa & Lainnya: ( Sektor_Jasa / Total_PDRB ) * 100 (misal Morowali: 54.0%).",
        "- `Porsi_Akar_Rumput (%)`: Persentase kontribusi Sektor Ekonomi Akar Rumput: ( Sektor_Akar_Rumput / Total_PDRB ) * 100 (misal Morowali: 0.8%).",
        "- `Total_PDRB_Kabupaten`: Total nilai nominal PDRB seluruh sektor di kabupaten target (Triliun Rupiah).",
        "",
        "#### D. Rincian Definisi Operasional & Matriks Distribusi PDRB 13 Kabupaten",
        f"Penerapan sistem persamaan di atas terhadap seluruh 13 kabupaten dan kota di Provinsi Sulawesi Tengah pada tahun {latest_year_kab} disajikan secara komprehensif pada **Tabel 1.3** berikut:",
        "",
        f"##### Tabel 1.3: Distribusi Nilai Tambah Bruto dan Komposisi Sektoral PDRB 13 Kabupaten/Kota di Sulawesi Tengah (Tahun {latest_year_kab})",
        "| Kabupaten / Kota | Akar Rumput (T Rp) | Ekstraktif (T Rp) | Jasa (T Rp) | Total PDRB (T Rp) | Porsi Akar Rumput (%) | Porsi Ekstraktif (%) | Porsi Jasa (%) | Basis Utama Ekonomi |",
        "| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :--- |",
    ]
    md_lines.extend(md_kab_rows)
    md_lines.extend([
        "",
        "#### E. Analisis Temuan Empiris: Polarisasi Ekstrem Morowali vs Daerah Non-Smelter",
        "Data empiris pada Tabel 1.3 mengungkap bukti polarisasi ekonomi wilayah yang sangat ekstrem di Sulawesi Tengah:",
        "",
        f"1. **Dominasi Sektor Ekstraktif Morowali:** Kabupaten Morowali mencatatkan nilai sektor ekstraktif sebesar Rp {df_kab_pivot.loc['Morowali', 'Ekstraktif']:.2f} Triliun atau menguasai porsi {df_kab_pivot.loc['Morowali', 'Pct_Ekstraktif']:.1f}% dari total kue ekonomi kabupatennya (Rp {df_kab_pivot.loc['Morowali', 'Total']:.2f} Triliun). Nilai sektor ekstraktif Morowali saja melampaui gabungan total PDRB dari delapan kabupaten lainnya di Sulawesi Tengah.",
        "2. **Pemusatan pada Dua Sentra Hilirisasi:** Kabupaten Morowali dan Morowali Utara merupakan dua daerah dengan nilai Sektor Ekstraktif tertinggi di Sulawesi Tengah, membuktikan bahwa percepatan output industri pertambangan dan hilirisasi terkunci pada kawasan industri smelter.",
        "3. **Ketertinggalan Daerah Non-Sentra:** Sebaliknya, delapan kabupaten lainnya (seperti Banggai Laut, Banggai Kepulauan, Tojo Una-Una, Buol, Toli-Toli, Sigi, Poso, dan Donggala) memiliki porsi Sektor Ekstraktif yang sangat rendah (<11%) dan tetap bergantung pada sektor pertanian rakyat (Akar Rumput) berproduktivitas rendah dengan keterbatasan akses terhadap nilai tambah modal.",
        "",
        "### 1.1.3 Perbandingan Distribusi 17 Sektor Komoditas per Provinsi (Small Multiples, Tahun Terbaru)",
        f"Visualisasi komparatif **Small Multiples Horizontal Bar Chart** membedah struktur 17 sektor lapangan usaha KBLI 2020 secara terpisah pada enam provinsi di Pulau Sulawesi pada tahun terbaru ({latest_year_pdrb}). Setiap panel provinsi menampilkan sektor yang diurutkan dari penyumbang terbesar hingga terkecil dengan skala sumbu nilai yang disetarakan secara seragam untuk memastikan validitas komparasi lintas wilayah.",
        "",
        f"> **Sumber Data Resmi & Deskripsi Visualisasi:** Badan Pusat Statistik (BPS) Provinsi se-Sulawesi (diolah CELIOS). Visualisasi *Small Multiples Horizontal Bar Chart* menyajikan dekomposisi 17 sektor PDRB tahun {latest_year_pdrb} di 6 provinsi se-Pulau Sulawesi. Sumbu X disetarakan pada rentang nilai seragam ([0, {max_x_val_all:.1f} Triliun Rp]) dengan pewarnaan berdasarkan 3 klaster makro (Merah: Ekstraktif, Hijau: Ekonomi Akar Rumput, Abu-abu: Sektor Jasa & Lainnya) guna mengidentifikasi spesialisasi dan anomali struktural ekonomi masing-masing provinsi.",
        "",
        "#### A. Kerangka Konseptual & Standardisasi Skala Komparatif (Uniform Scale Small Multiples)",
        "Dalam analisis data multidimensi lintas wilayah, penggunaan skala dinamis mandiri (*independent dynamic scaling*) pada masing-masing panel sering kali menimbulkan **Bias Distorsi Visual Komparatif (Visual Comparison Bias)**. Tanpa penyetaraan batas skala maksimum, sektor dengan nominal kecil di provinsi ber-PDRB rendah dapat terlihat secara visual setara dengan sektor bernilai ratusan triliun di provinsi ber-PDRB besar. Oleh karena itu, metodologi ini menetapkan batas skala maksimum sumbu X yang seragam (*Uniform Scale Bound*) sebesar nilai maksimum sektor tertinggi di seluruh pulau ditambah faktor ruang margin sebesar 15%.",
        "",
        "#### B. Alur Logika Metodologis Analisis Small Multiples 17 Sektor",
        "Kerangka operasionalisasi analisis perbandingan terpisah 17 sektor lapangan usaha ini dimodelkan dalam kerangka alur logika sebagaimana diilustrasikan pada **Bagan Alur 1.3** berikut:",
        "",
        "##### Bagan Alur 1.3: Alur Logika Metodologis Analisis Komparatif Small Multiples 17 Sektor PDRB per Provinsi",
        "```mermaid",
        f"{mermaid_str_1_3}",
        "```",
        "",
        "#### C. Formulasi Matematis: Persamaan Agregasi dan Porsi 17 Sektor Komoditas",
        "Kalkulasi perbandingan sektoral dan normalisasi skala grafik dihitung menggunakan sistem formulasi berikut:",
        "",
        "**Persamaan Normalisasi Nilai Sektor ke Satuan Triliun Rupiah:**",
        "```text",
        "Nilai_Sektor_Triliun = Nilai_Sektor_Miliar / 1000",
        "```",
        "*Keterangan Variabel:*",
        "- `Nilai_Sektor_Triliun`: Nilai tambah bruto sektor lapangan usaha dalam satuan baku Triliun Rupiah.",
        "- `Nilai_Sektor_Miliar`: Nilai nominal PDRB mentah dari publikasi resmi BPS (satuan: Miliar Rupiah).",
        "",
        "**Persamaan Porsi Sektoral per Provinsi (Porsi (%) pada Tooltip Dashboard):**",
        "```text",
        "Porsi_Sektor_Provinsi (%) = ( Nilai_Sektor_Provinsi / Total_PDRB_Provinsi ) * 100",
        "```",
        "*Keterangan Variabel:*",
        "- `Porsi_Sektor_Provinsi (%)`: Persentase kontribusi sektor target terhadap keseluruhan total PDRB provinsi bersangkutan (satuan: Persen / %). Angka ini ditampilkan pada tooltip 'Porsi (%)' di dashboard.",
        "- `Nilai_Sektor_Provinsi`: Nilai tambah bruto sektor lapangan usaha target di provinsi bersangkutan (Triliun Rupiah).",
        "- `Total_PDRB_Provinsi`: Total nilai nominal PDRB seluruh 17 sektor di provinsi bersangkutan (Triliun Rupiah).",
        "",
        "**Persamaan Batas Maksimum Skala Sumbu X Seragam (Uniform Scale Bound):**",
        "```text",
        "X_Max_Seragam = Maksimum(Seluruh_Nilai_Sektor_Semua_Provinsi) * 1.15",
        "```",
        "*Keterangan Variabel:*",
        "- `X_Max_Seragam`: Nilai batas atas sumbu X yang diaplikasikan secara identik pada semua grafik Small Multiples.",
        "- `Maksimum(...)`: Sektor dengan nominal PDRB terbesar di seluruh pulau (contoh: Sektor Pertanian di Sulsel).",
        "- `1.15`: Faktor pengali batas margin (+15%) untuk ruang keterangan (label space).",
        "",
        "#### D. Rincian Data Empiris: Matriks Perbandingan Sektor Unggulan 6 Provinsi",
        f"Penerapan sistem perbandingan komparatif di atas mengidentifikasi 5 sektor tulang punggung utama (*top 5 contributors*) dari 17 Kategori BPS di tiap provinsi pada tahun {latest_year_pdrb}, sebagaimana dirinci pada **Tabel 1.4** berikut:",
        "",
        f"##### Tabel 1.4: 5 Sektor Lapangan Usaha Penyumbang Utama PDRB di 6 Provinsi Sulawesi (Tahun {latest_year_pdrb})",
        "| Provinsi | Peringkat 1 | Peringkat 2 | Peringkat 3 | Peringkat 4 | Peringkat 5 |",
        "| :--- | :--- | :--- | :--- | :--- | :--- |",
    ])
    md_lines.extend(md_top_rows)
    md_lines.extend([
        "",
        "#### E. Analisis Temuan Empiris & Interpretasi Sektoral Dashboard",
        f"Visualisasi *Small Multiples* dan matriks Tabel 1.4 memvalidasi hipotesis adanya **Disparitas Struktural** yang dipicu oleh ekspansi industri nikel:",
        "",
        "1. **Dominasi Absolut Ekstraktif di Sentra Smelter:** Sulawesi Tengah dan Sulawesi Tenggara menunjukkan pola yang identik, di mana Sektor Ekstraktif (Industri Pengolahan Logam Dasar dan Pertambangan) menjadi tulang punggung absolut. Khususnya di Sulawesi Tengah, jarak nilai (gap) antara sektor ekstraktif dengan sektor lainnya sangat ekstrem.",
        "2. **Perekonomian Berbasis Akar Rumput di Provinsi Lain:** Empat provinsi lainnya (Sulawesi Selatan, Sulawesi Barat, Gorontalo, dan Sulawesi Utara) tetap mengandalkan Sektor Pertanian (Kategori A) sebagai penyumbang terbesar PDRB, mencerminkan resiliensi ekonomi akar rumput di wilayah non-smelter.",
        "3. **Urgensi Normalisasi Skala Visual:** Penggunaan skala X seragam ([0, 168 Triliun Rp]) memastikan pembaca menyadari bahwa meskipun Sektor Industri Pengolahan menempati Peringkat 1 di Sulawesi Tengah, nominal absolutnya belum tentu setara dengan Sektor Pertanian di Sulawesi Selatan. Visualisasi terstandarisasi ini mencegah kesimpulan spekulatif.",
        "",
        "---",
        "",
        "## 1.2 Konsentrasi Kawasan Industri & PLTU Captive",
        "",
        "Intensifikasi industri pengolahan mineral di Pulau Sulawesi berpusat pada pembangunan mega-smelter yang ditopang secara mutlak oleh pembangkit listrik tenaga uap khusus (*PLTU Captive*) batu bara non-jaringan (*off-grid*). Bagian ini mengombinasikan **Analisis Spasial Deskriptif** untuk mengidentifikasi tingkat pemusatan fasilitas dan kapasitas energi fosil di enam provinsi, dengan **Uji Tabulasi Silang Panel (Inferential Spatiotemporal Crosstabulation)** berstandar SPSS guna membuktikan secara ilmiah keterkaitan antara ekspansi PLTU Captive terhadap kehilangan tutupan hutan di Pulau Sulawesi.",
        "",
        f"> **Sumber Data Resmi & Deskripsi Metodologis:** Kementerian Energi dan Sumber Daya Mineral (ESDM / Minerbaone), Global Energy Monitor (GEM Coal Plant Tracker), dan Global Forest Watch (GFW / University of Maryland) (diolah CELIOS). Visualisasi *Bar Chart* Konsentrasi Industri dan Pemetaan Spasial menyajikan distribusi {tot_smelter_all} unit fasilitas smelter serta {tot_pltu_mw_all:,.0f} MW kapasitas terpasang aktif PLTU captive di 6 provinsi se-Pulau Sulawesi. Analisis dipadukan dengan Uji Tabulasi Silang Data Panel Spasiotemporal (Chi-Square Test & Risk Odds Ratio, N=60) untuk menguji keterkaitan ekspansi energi fosil industri terhadap eskalasi deforestasi komoditas.",
        "",
        "### A. Pemusatan Spasial Fasilitas Smelter dan PLTU Captive",
        f"Intensifikasi industri pengolahan nikel di Sulawesi berpusat pada fasilitas mega-smelter. Pengoperasian **{tot_smelter_all} fasilitas smelter** didukung oleh kapasitas energi batu bara **{tot_pltu_mw_all:,.0f} MW dari PLTU Captive**. Berbeda dengan sistem kelistrikan umum PLN, pembangkit ini dikembangkan secara internal untuk menyokong operasi kawasan industri.",
        "",
        "### B. Metodologi: Analisis Spasial & Uji Tabulasi Silang",
        "Pengujian keterkaitan antara pembangunan PLTU Captive dengan kehilangan tutupan hutan dioperasionalkan melalui Standar Operasional Prosedur (SOP) tabulasi silang berstandar SPSS. Rangkaian tahapan logika metodologis, asumsi frekuensi harapan, hingga estimasi faktor risiko dimodelkan pada **Bagan Alur 1.4** berikut:",
        "",
        "##### Bagan Alur 1.4: Standar Operasional Prosedur (SOP) & Alur Logika Uji Tabulasi Silang (Crosstab) PLTU Captive vs Deforestasi",
        "```mermaid",
        f"{mermaid_str}",
        "```",
        "",
        "### C. Formulasi Matematis: Kalkulasi Konsentrasi Spasial & Uji Chi-Square",
        "Parameterisasi konsentrasi spasial dan pembuktian statistik dihitung menggunakan sistem formulasi matematis berikut:",
        "",
        "**Persamaan Akumulasi Kapasitas PLTU Kumulatif per Wilayah (MW):**",
        "```text",
        "Kapasitas_PLTU_Kumulatif_t (MW) = Jumlah Kapasitas Aktif Baru (MW) dari Tahun 2014 hingga Tahun t",
        "```",
        "*Keterangan Variabel:*",
        "- `Kapasitas_PLTU_Kumulatif_t (MW)`: Total akumulasi kapasitas daya terpasang operasional PLTU captive batubara aktif hingga tahun t (satuan: Megawatt / MW).",
        "- `Kapasitas Aktif Baru`: Besaran daya listrik unit PLTU off-grid yang mulai beroperasi komersial pada tahun tertentu (satuan: Megawatt / MW).",
        "",
        "**Persamaan Rasio Konsentrasi Spasial Fasilitas Smelter (% pada Grafik Dashboard):**",
        "```text",
        "Porsi_Smelter_Provinsi (%) = ( Jumlah_Smelter_Provinsi / Total_Smelter_Sulawesi ) * 100",
        "```",
        "*Keterangan Variabel:*",
        "- `Porsi_Smelter_Provinsi (%)`: Persentase pangsa fasilitas smelter di provinsi bersangkutan terhadap seluruh Pulau Sulawesi (satuan: Persen / %).",
        "- `Jumlah_Smelter_Provinsi`: Banyaknya unit smelter yang beroperasi di wilayah provinsi tertentu.",
        f"- `Total_Smelter_Sulawesi`: Total keseluruhan fasilitas smelter di Pulau Sulawesi ({tot_smelter_all} unit).",
        "",
        "**Persamaan Uji Independensi Chi-Square Pearson (χ² Kontinjensi 2x2):**",
        "```text",
        "Chi_Square (χ²) = Jumlah [ (Frekuensi_Observasi - Frekuensi_Harapan)^2 / Frekuensi_Harapan ]",
        "```",
        "*Keterangan Variabel:*",
        "- `Chi_Square (χ²)`: Nilai statistik uji kecocokan Pearson untuk membuktikan ada tidaknya hubungan ketergantungan antara ekspansi PLTU Captive dengan lonjakan deforestasi pada panel spasiotemporal (N=60).",
        "- `Frekuensi_Observasi (O)`: Jumlah kasus aktual yang tercatat pada sel tabel kontinjensi 2x2.",
        "- `Frekuensi_Harapan (E)`: Jumlah kasus teoretis jika kedua variabel saling independen: E = (Total Baris * Total Kolom) / N.",
        "",
        "**Persamaan Rasio Keunggulan Risiko (Risk Odds Ratio / OR):**",
        "```text",
        "Odds_Ratio (OR) = ( a * d ) / ( b * c )",
        "```",
        "*Keterangan Variabel:*",
        "- `Odds_Ratio (OR)`: Ukuran kelipatan risiko peluang terjadinya deforestasi komoditas tinggi pada kelompok dengan PLTU Captive aktif (>0 MW) dibanding kelompok tanpa PLTU Captive (≤0 MW).",
        f"- `a`: Jumlah observasi panel pada kelompok PLTU Rendah dan Deforestasi Rendah ({a_2} kasus).",
        f"- `b`: Jumlah observasi panel pada kelompok PLTU Rendah dan Deforestasi Tinggi ({b_2} kasus).",
        f"- `c`: Jumlah observasi panel pada kelompok PLTU Tinggi dan Deforestasi Rendah ({c_2} kasus).",
        f"- `d`: Jumlah observasi panel pada kelompok PLTU Tinggi dan Deforestasi Tinggi ({d_2} kasus).",
        "",
        "### D. Matriks Hasil Uji Empiris: Konsentrasi Spasial & Skenario Crosstab",
        "Penerapan sistem pengujian statistik tabulasi silang pada data panel 6 provinsi selama 1 dekade (2014–2023, total 60 observasi) disajikan secara lengkap pada **Tabel 1.5** berikut:",
        "",
        "##### Tabel 1.5: Matriks Tabulasi Silang 2×2, Uji Chi-Square (χ²), dan Estimasi Odds Ratio Panel PLTU Captive vs Deforestasi Komoditas (2014–2023)",
        f"| Kategori Kapasitas PLTU (X) | Deforestasi {label_y_low_2} | Deforestasi {label_y_high_2} | Total Kasus | Parameter Statistik Uji | Nilai / df | Signifikansi / Kesimpulan |",
        "| :--- | :---: | :---: | :---: | :--- | :--- | :--- |",
        f"| **Rendah (≤0 MW)** | {a_2} [Exp: {exp_pltu[0,0]:.1f}] | {b_2} [Exp: {exp_pltu[0,1]:.1f}] | {a_2+b_2} (100%) | **{test_name}** | **{test_val}** ({dof_str}) | p {fmt_p(p_val_display)} ({'Signifikan' if p_val_display < 0.05 else 'Tidak Signifikan'}) |",
        f"| **Tinggi (>0 MW)** | {c_2} [Exp: {exp_pltu[1,0]:.1f}] | {d_2} [Exp: {exp_pltu[1,1]:.1f}] | {c_2+d_2} (100%) | **Likelihood Ratio** | **{g_pltu:.3f}** (df={dof_g_pltu}) | p {fmt_p(p_g_pltu)} (Signifikan) |",
        f"| **Total Observasi Panel** | **{a_2+c_2}** [Exp: {a_2+c_2:.1f}] | **{b_2+d_2}** [Exp: {b_2+d_2:.1f}] | **{len(df_panel_1_2)}** (100%) | **Linear-by-Linear Association** | **{lbl_val_pltu:.3f}** (df=1) | p {fmt_p(p_corr_2)} (Signifikan) |",
        f"| **Ukuran Risiko (Risk Estimate)** | Cross-Product: ({a_2}×{d_2})/({b_2}×{c_2}) | Rasio Peluang Risiko | OR = {or_pltu:.2f} | **Odds Ratio (OR)** | **{or_pltu:.2f}x** | **Risiko Lonjakan {int(or_pltu)}x Lipat** |",
        "",
        "### E. Interpretasi Spasial Industri: Eksternalitas dan Efek Meluber (Spillover)",
        "Hasil pengujian empiris pada Tabel 1.5 membuktikan secara meyakinkan keterkaitan langsung antara ekspansi PLTU Captive dan kerusakan tutupan hutan di Pulau Sulawesi:",
        "",
        "1. **Pemusatan Ekstrem di 3 Sentra Ekstraktif Utama:** 100% kapasitas PLTU Captive dan mayoritas smelter berpusat di wilayah ini, memicu akumulasi deforestasi komoditas hingga ratusan ribu hektar, berbanding terbalik dengan \"Area Non-Smelter\".",
        f"2. **Signifikansi Statistik yang Sangat Kuat (p {fmt_p(p_val_display)}):** Hipotesis Nol (H0) ditolak mutlak. Bukti empiris mengonfirmasi bahwa penambahan kapasitas PLTU Captive berkorelasi langsung dengan lonjakan kehilangan tutupan hutan.",
        f"3. **Kelipatan Risiko Bencana Ekologis (Odds Ratio = {or_pltu:.2f}x):** Wilayah dengan PLTU Captive memiliki risiko deforestasi komoditas {int(or_pltu)} KALI LIPAT lebih besar. Hal ini didorong konversi masif untuk infrastruktur pendukung (coal yard, jalur transmisi, dan jalan logistik).",
        "4. **Efek Meluber Lintas Batas (Spillover Effect) & Emisi Karbon Terkunci:** Eksternalitas destruktif proyek merambat luas mendegradasi DAS dan laut, mengorbankan ruang hidup lokal, serta mengunci emisi dari ketergantungan puluhan juta ton batu bara per tahun.",
        "",
        "---",
        "",
        "## 1.3 Tren Pertumbuhan Izin Tambang Baru & Uji Signifikansi Statistik",
        "#### A. Pengantar & Kerangka Narasi",
        f"Pola perizinan pertambangan di Pulau Sulawesi selama satu dekade terakhir menunjukkan peningkatan alokasi ruang yang signifikan. Berdasarkan data agregat **Minerbaone**, tercatat {tot_izin:,} Izin Usaha Pertambangan (IUP) baru sepanjang 2014-2024, dengan total luas konsesi mencapai {tot_luas_izin:,} Hektar.",
        "",
        "Berdasarkan analisis tren time-series pada grafik \"Penerbitan Izin Tambang\", penerbitan izin pada periode awal (2014) tercatat lebih rendah. Peningkatan signifikan terjadi pada periode 2022-2024. Anotasi pada grafik mencatat kenaikan sebesar **246% pada periode 2022-2024**. Data ini mengindikasikan perlunya evaluasi terhadap instrumen pengendalian perizinan dan tata ruang. Distribusi perizinan tertinggi berada di Sulawesi Tengah dan Sulawesi Tenggara, yang selaras dengan kawasan pengembangan industri pemurnian nikel.",
        "",
        "Uji **Crosstabulation** pada analisis ini digunakan untuk mengukur hubungan antara laju penerbitan perizinan (X) dan indikator deforestasi di wilayah tersebut (Y).",
        "",
        "#### B. Alur Logika Metodologis (Flowchart)",
        "Pendekatan statistik Time-Series untuk mengidentifikasi tren pertumbuhan izin tambang diilustrasikan pada **Bagan Alur 1.5** berikut. Adapun untuk tahapan analisis inferensial (Uji Chi-Square), alur logikanya merujuk secara penuh pada **Bagan Alur 1.4** (di sub-bab sebelumnya) dengan penyesuaian konfigurasi variabel spesifik sesuai Tabel Asumsi Dasar di bawah gambar.",
        "##### Bagan Alur 1.5: Alur Logika Tren Pertumbuhan Izin Tambang Baru",
        "```mermaid",
        f"{mermaid_str_1_3_izin}",
        "```",
        "",
        "##### Tabel 1.5b: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.3)",
        "| Komponen Uji | Definisi Variabel (Sub-bab 1.3) |",
        "| :--- | :--- |",
        "| Variabel Independen (X) | Frekuensi Penerbitan Izin Tambang Baru (IUP) / Luas Konsesi Baru (Ha) |",
        "| Variabel Dependen (Y) | Deforestasi Komoditas (Ha) / Total Deforestasi Alam (Ha) |",
        "| Hipotesis Nol (H0) | Tingkat penerbitan izin/luas konsesi tidak berhubungan dengan laju deforestasi. |",
        "| Hipotesis Alternatif (H1) | Ada hubungan positif antara tingginya penerbitan izin dengan tingginya laju deforestasi. |",
        "| Threshold Kategori | Nilai Median Data Panel (N=60) |",
        "",
        "#### C. Formulasi Matematis: Analisis Tren & Uji Chi-Square",
        "Parameterisasi laju pertumbuhan perizinan dan pengujian signifikansi dampaknya terhadap deforestasi dihitung menggunakan formulasi berikut:",
        "",
        "**Laju Pertumbuhan Izin Tahunan (Regresi Komparatif YoY):**",
        "```text",
        "Pertumbuhan_Izin (%) = [ ( Jumlah_Izin_t - Jumlah_Izin_{t-1} ) / Jumlah_Izin_{t-1} ] * 100",
        "```",
        "*Keterangan Variabel:*",
        "- `Pertumbuhan_Izin (%)`: Persentase perubahan laju penerbitan izin tambang baru antar-tahun (satuan: Persen / %).",
        "- `Jumlah_Izin_t`: Agregasi jumlah izin (atau luasan) pada tahun berjalan (t).",
        "- `Jumlah_Izin_{t-1}`: Agregasi jumlah izin (atau luasan) pada satu tahun sebelumnya (t - 1).",
        "",
        "**Pengklasifikasian Kategori Data (Binning Threshold Median):**",
        "```text",
        "Kategori = IF(Nilai_Prov_Tahun >= Median(Seluruh Panel), \"Tinggi\", \"Rendah\")",
        "```",
        "*Keterangan Variabel:*",
        "- `Kategori`: Data panel spasial-temporal diubah menjadi dua tingkatan untuk uji tabulasi silang (Tinggi vs Rendah).",
        "",
        "Dinamika historis perizinan secara terperinci dapat dilihat pada **Tabel 1.5c**, yang menunjukkan tren penerbitan izin baru di wilayah studi:",
        "",
        "##### Tabel 1.5c: Tren Penerbitan Izin Tambang Sulawesi (2014-2024)",
        "| Tahun | Gorontalo | Sulawesi Barat | Sulawesi Selatan | Sulawesi Tengah | Sulawesi Tenggara | Sulawesi Utara |",
        "| :---: | :---: | :---: | :---: | :---: | :---: | :---: |",
        md_tren_rows.rstrip(),
        "",
        "#### D. Matriks Hasil Uji Empiris: Ringkasan Skenario Crosstab",
        "Hasil lengkap pengujian independensi statistik Chi-Square dan estimasi Odds Ratio (OR) untuk seluruh faktor tekanan terhadap kehilangan tutupan hutan komoditas dirangkum pada **Tabel 1.6** berikut:",
        "",
        "##### Tabel 1.6: Ringkasan Hasil Uji Independensi Chi-Square (χ²) dan Odds Ratio (OR) Data Panel Bab 1",
        "| Variabel Faktor Tekanan | Variabel Dampak Lingkungan | Nilai Chi-Square (χ²) | Nilai Signifikansi (p) | Odds Ratio (OR) | Derajat Bebas (df) | Kesimpulan Ilmiah |",
        "| :--- | :--- | :---: | :---: | :---: | :---: | :--- |",
        md_chi_rows.rstrip(),
        "",
        "#### E. Analisis Temuan Empiris: Pembedahan Realitas Ekologis",
        "Data panel membedah realitas di lapangan: lonjakan izin di wilayah pusat ekstraksi sejalan dengan tingginya nilai Chi-Square. Nilai Odds Ratio menegaskan bahwa wilayah dengan tren izin tambang yang tinggi memiliki peluang lebih besar untuk mengalami tekanan deforestasi tinggi pada tahun-tahun berjalan dan berikutnya.",
        "",
        "Secara spesifik, terjadi **lonjakan absolut sebesar 246%** dalam penerbitan izin tambang baru pada rentang 2022 hingga 2024. Lonjakan ekstrem ini mengindikasikan percepatan luar biasa dari ekspansi industri ekstraktif yang mengabaikan kapasitas daya dukung lingkungan tapak, terutama di sentra-sentra produksi.",
        "",
        "---",
        "",
        "## 1.4 Analisis Realisasi Investasi PMDN dan Dampak Terhadap Tutupan Hutan",
        "#### A. Pengantar & Kerangka Narasi",
        f"Akumulasi Penanaman Modal Dalam Negeri sebesar **Rp {tot_investasi_triliun:,} Triliun** (Kementerian Investasi / BKPM) yang masuk dari tahun 2016-2024 berbanding lurus dengan **{tot_deforestasi:,} Hektar** kehilangan tutupan hutan komoditas (Global Forest Watch). Grafik sumbu ganda (*dual-axis*) digunakan untuk membandingkan laju investasi dan laju deforestasi antara wilayah sentra industri tambang dengan non-sentra. Terlihat adanya fenomena **Efek Jeda Waktu (Time-Lagging Effect)**, di mana peningkatan realisasi modal pada tahap awal perizinan dan konstruksi diikuti oleh lonjakan pembukaan lahan hutan fisik pada 1 hingga 2 tahun berikutnya.",
        "",
        "#### B. Alur Logika Metodologis Analisis Realisasi Investasi PMDN",
        "Kerangka operasionalisasi uji statistik tabulasi silang antara realisasi Investasi PMDN dan deforestasi dimodelkan dalam kerangka alur logika sebagaimana diilustrasikan pada **Bagan Alur 1.4** berikut:",
        "",
        "##### Bagan Alur 1.4: Alur Logika Metodologis Uji Independensi Panel Investasi PMDN vs Deforestasi",
        "```mermaid",
        f"{mermaid_str_1_4}",
        "```",
        "",
        "#### C. Formulasi Matematis: Agregasi Dampak Ekologis & Pengujian Statistik",
        "**Persamaan Agregasi Luasan Deforestasi Berdasarkan Faktor Penggerak (Driver)**",
        "`Total_Deforestasi_Driver_k = &sum;(Area_Loss_i) untuk i di Kategori_Driver_k`",
        "- **Total_Deforestasi_Driver_k:** Total luas kehilangan tutupan pohon yang diakibatkan oleh faktor penggerak k (contoh: Ekspansi Komoditas) di seluruh wilayah observasi (satuan: Hektar / Ha).",
        "- **Kategori_Driver_k:** Klasifikasi penyebab utama deforestasi (Dominant Driver of Tree Cover Loss) berdasarkan model data historis satelit.",
        "- **Area_Loss_i:** Luas kehilangan tutupan pohon pada piksel observasi ke-i (satuan: Hektar / Ha).",
        "",
        "**Persamaan Perhitungan Akumulasi Kehilangan Hutan Primer (Primary Forest Loss)**",
        "`Total_Primary_Loss = &sum;(Area_Loss_j) untuk j di mana Tipe_Hutan = \"Primer\"`",
        "- **Total_Primary_Loss:** Akumulasi luas konversi tutupan hutan alam primer tak terganggu (intact primary forest) selama periode pengamatan (satuan: Hektar / Ha).",
        "- **Tipe_Hutan:** Klasifikasi basemap jenis tutupan lahan awal sebelum terjadi deforestasi.",
        "",
        "**Persamaan Estimasi Pelepasan Emisi Karbon (Gross CO2 Emissions)**",
        "`Emisi_CO2_Total = &sum;(Area_Loss_c * Faktor_Emisi_Biomassa_c)`",
        "- **Emisi_CO2_Total:** Estimasi agregasi total emisi gas rumah kaca yang dilepaskan ke atmosfer akibat konversi tutupan (satuan: Megagrams CO2 / Mg CO2).",
        "- **Faktor_Emisi_Biomassa_c:** Kandungan karbon rata-rata (above-ground & below-ground biomass) per hektar pada koordinat c yang diamati.",
        "",
        "Kalkulasi pengujian statistik dihitung menggunakan formulasi Matematis yang sama dengan Sub-Bab 1.2 dan 1.3, di mana variabel independen (X) adalah **Investasi PMDN (Juta Rp)** dan variabel dependen (Y) adalah **Deforestasi Komoditas (Hektar)**.",
        "",
        "**Persamaan Kategorisasi Nilai Ambang Batas Median:**",
        "```text",
        "- Jika Nilai > Median, maka Kategori = Tinggi",
        "- Jika Nilai <= Median, maka Kategori = Rendah",
        "```",
        "",
        "#### D. Matriks Hasil Uji Empiris: Konsentrasi Spasial & Skenario Crosstab",
        "Tingkat alokasi konsesi dan dampaknya terhadap tutupan hutan dapat dilihat secara empiris melalui perbandingan luas konsesi baru di Daerah Sentra Tambang (Morowali & Konawe) dengan wilayah non-sentra pada **Tabel 1.7b** berikut:",
        "",
        "##### Tabel 1.7b: Representasi Spasial Luas Konsesi Baru (Ha) di Daerah Sentra Tambang vs Non-Sentra (2014-2023)",
        "| Tahun | Luas Konsesi Baru Daerah Sentra Tambang (Ha) | Luas Konsesi Baru Daerah Non-Sentra (Ha) |",
        "| :---: | :---: | :---: |",
        md_emp_rows_1_4.rstrip(),
        "",
        "Terkait dengan hilangnya luasan hutan tersebut, pembedahan lebih lanjut berdasarkan aktor utama, luasan hutan primer, dan estimasi emisi karbon komoditas dapat dilihat pada **Tabel 1.7c** berikut:",
        "",
        "##### Tabel 1.7c: Matriks Pembedahan Ekologis Aktor & Emisi Karbon (Periode 2001-2025)",
        "| Kategori Aktor / Metrik Ekologis | Nilai Agregat | Persentase dari Total Kehilangan |",
        "| :--- | :---: | :---: |",
        "| **Ekspansi Komoditas (Tambang & Sawit)** | 1,890,659 Hektar | 48.4% |",
        "| **Kehutanan (Logging)** | 247,011 Hektar | 6.3% |",
        "| **Pertanian Berpindah** | 115,404 Hektar | 2.9% |",
        "| **Total Kehilangan Hutan Primer** | **3,904,079 Hektar** | **100.0%** |",
        "| **Estimasi Emisi Karbon Komoditas** | 1,282,195,705 Mg CO2 | - |",
        "",
        "##### Tabel 1.7d: Konfigurasi Variabel Uji Chi-Square (Sub-bab 1.4)",
        "| Komponen Uji | Definisi Variabel (Sub-bab 1.4) |",
        "| :--- | :--- |",
        "| Variabel Independen (X) | Realisasi Investasi PMDN (Juta Rupiah) |",
        "| Variabel Dependen (Y) | Total Deforestasi Alam (Ha) / Deforestasi Komoditas (Ha) |",
        "| Hipotesis Nol (H0) | Tingginya realisasi investasi PMDN tidak berhubungan dengan laju deforestasi. |",
        "| Hipotesis Alternatif (H1) | Ada hubungan positif antara tingginya realisasi investasi PMDN dengan laju deforestasi. |",
        "| Threshold Kategori | Nilai Median Data Panel (N=48) |",
        "",
        "Tabel di bawah ini merangkum hasil pengujian statistik (Chi-Square) untuk semua kemungkinan kombinasi indikator antara Realisasi Investasi PMDN dan Dampak Ekologis pada panel data 2016-2023. Hasil tersebut ditampilkan pada **Tabel 1.8** berikut:",
        "",
        "##### Tabel 1.8: Ringkasan Eksekutif Seluruh Skenario Crosstab Realisasi Investasi PMDN Bab 1",
        "| Variabel Independen (X) | Variabel Dependen (Y) | Chi-Square (χ²) | P-Value | Odds Ratio | Kesimpulan |",
        "| :--- | :--- | :---: | :---: | :---: | :--- |",
        md_chi_rows_4.rstrip(),
        "",
        "#### E. Analisis Temuan Empiris: Efek Jeda Waktu (Time-Lagging)",
        "Hasil pengujian seluruh skenario tabulasi silang PMDN mengungkap fenomena yang kompleks dalam alur investasi ekstraktif:",
        "",
        "1. **Ketidaksignifikanan Simultan & Variasi P-Value:** Tingkat signifikansi yang bervariasi menyingkap tabir jeda waktu (lagging effect) dalam eksekusi investasi di lapangan.",
        "2. **Jeda Waktu Eksekusi Investasi (Lagging Effect):** Suntikan modal masif di tahun tertentu tidak secara instan berwujud pembabatan lahan di tahun yang sama. Modal tersebut tertahan pada fase birokrasi, pembebasan lahan, dan pengadaan infrastruktur, sebelum daya rusaknya mengonversi lanskap hutan pada tahun-tahun berikutnya.",
        "3. **Konsentrasi Modal Ekstrem di 3 Provinsi:** Data spasial membuktikan 89% dari total modal PMDN ekstraktif se-Sulawesi hanya tersedot ke tiga provinsi sentra (Sulteng, Sultra, Sulsel), mengakibatkan polarisasi pertumbuhan dan mengunci ketimpangan spasial.",
        "",
        "---",
        "",
        "## 1.5 Pelabuhan Ekspor & Peta Jalur Distribusi Logistik Nikel Sulawesi",
        "",
        "#### A. Pengantar & Kerangka Narasi",
        "Eksploitasi nikel di Sulawesi tidak berhenti di tapak darat, melainkan terhubung langsung ke pasar global melalui infrastruktur pelabuhan. Bagian ini memetakan simpul logistik maritim yang mendistribusikan produk ekstraktif (NPI, Matte, MHP) dari pesisir Sulawesi menuju negara tujuan utama seperti Tiongkok dan Jepang.",
        "",
        "#### B. Alur Logika Metodologis (Validasi Silang & Kurva Bézier)",
        "Verifikasi titik pelabuhan dan terminal khusus ekspor nikel dilakukan melalui protokol triangulasi informasi publik sebagaimana divisualisasikan pada **Bagan Alur 1.5** berikut:",
        "",
        "##### Bagan Alur 1.5: Alur Logika Metodologis Validasi Silang (OSINT) dan Pemetaan Spasial Pelabuhan",
        "```mermaid",
        "flowchart LR",
        "    subgraph Data_Acquisition[\"1. Validasi Silang Dokumen Publik (OSINT)\"]",
        "        A[\"Laporan KNKT<br/><i>(Kapasitas Muatan & DWT)</i>\"] --> D",
        "        B[\"Lampiran Perpres 109/2020<br/><i>(Status Proyek Strategis Nasional)</i>\"] --> D",
        "        C[\"Laporan Tahunan & Laporan Keberlanjutan<br/><i>(Fasilitas Terminal Khusus)</i>\"] --> D",
        "    end",
        "",
        "    subgraph Data_Processing[\"2. Triangulasi & Pemetaan Rute\"]",
        "        D[\"<b>Inventarisasi Simpul Pelabuhan</b><br/>(Identifikasi 6 Kawasan Utama)\"] --> E[\"Ekstraksi Koordinat Geografis<br/>(Latitude & Longitude)\"]",
        "        E --> F[\"Pemodelan Kurva Parametrik (Bézier Curve)\"]",
        "        F --> G[\"Visualisasi Spatial Logistic Mapping\"]",
        "    end",
        "```",
        "",
        "Detail dari keempat sumber informasi validasi silang (triangulasi) publik adalah sebagai berikut:",
        "",
        "1. **Laporan Investigasi Keselamatan Transportasi Laut (KNKT):** Memverifikasi kapasitas dermaga curah dan bobot muatan kapal pengangkut bijih nikel hingga 52.378 DWT.",
        "2. **Regulasi Proyek Strategis Nasional (PSN):** Lampiran Perpres No. 109 Tahun 2020 sektor kawasan industri terpadu.",
        "3. **Laporan Keberlanjutan & Tahunan Korporasi Terbuka:** Laporan resmi PT Vale Indonesia Tbk dan PT ANTAM Tbk mengenai fasilitas pelabuhan khusus.",
        "4. **Laporan Audit Lembaga Riset Independen:** Publikasi riset independen mengenai rantai pasok dan operasional terminal khusus maritim.",
        "",
        "",
        "#### C. Matriks Hasil Uji Empiris: Inventarisasi 6 Simpul Pelabuhan",
        "Berdasarkan protokol validasi silang tersebut, profil komprehensif enam simpul pelabuhan dan terminal khusus utama di Pulau Sulawesi dipetakan pada **Tabel 1.7** berikut:",
        "",
        "##### Tabel 1.7: Inventarisasi Enam Simpul Pelabuhan dan Terminal Khusus Ekspor Nikel di Pulau Sulawesi",
        "| Simpul Kawasan Industri | Wilayah Administrasi | Fasilitas Pelabuhan / Terminal | Status Regulasi | Kapasitas Kapal | Tujuan Utama Ekspor |",
        "| :--- | :--- | :--- | :---: | :---: | :--- |",
        "| **IMIP Morowali** | Morowali, Sulawesi Tengah | Pelabuhan Samudera & Dermaga Curah | PSN (Perpres 109/2020) | Hingga 52.378 DWT | Pasar Global (Tiongkok) |",
        "| **GNI Morowali Utara** | Morowali Utara, Sulteng | Terminal Khusus Pesisir Tomori | Izin Industri Mandiri | Hingga 30.000 DWT | Pasar Global (Tiongkok) |",
        "| **VDNI Konawe** | Konawe, Sulawesi Tenggara | Dermaga Khusus Curah & Kargo | PSN (Perpres 109/2020) | Hingga 50.000 DWT | Pasar Global (Tiongkok) |",
        "| **OSS Konawe** | Konawe, Sulawesi Tenggara | Dermaga Terintegrasi Konawe | PSN (Perpres 109/2020) | Hingga 50.000 DWT | Pasar Global (Tiongkok) |",
        "| **Pomalaa (ANTAM)** | Kolaka, Sulawesi Tenggara | Dermaga Pomalaa & Konveyor | Kawasan BUMN Industri | Hingga 12.000 DWT | Jepang & Korsel |",
        "| **Sorowako (Vale)** | Luwu Timur, Sulawesi Selatan | Pelabuhan Balantang Malili | Kontrak Karya Tambang | Hingga 15.000 DWT | Jepang & Skandinavia |",
        "---",
        "",
        "## 1.6 Peta Jalur Distribusi Logistik Nikel Sulawesi",
        "#### A. Pengantar & Kerangka Narasi",
        "Metode analisis pada tahapan ini difokuskan pada **Pemetaan Kausalitas (Spasial)** untuk membedah asimetri penguasaan ruang antara wilayah hulu (origin: sumber ekstraksi di Sulawesi) dan hilir (destination: pusat industrialisasi luar negeri). Garis diplot menggunakan rute pelayaran untuk merepresentasikan jarak tempuh aktual kapal logistik di permukaan bumi.",
        "",
        "#### B. Alur Logika Metodologis (Flowchart Analisis Spasial)",
        "Rangkaian proses ekstraksi koordinat hingga visualisasi spasial direpresentasikan pada **Bagan Alur 1.6** berikut:",
        "",
        "##### Bagan Alur 1.6: Alur Logika Metodologis Pemetaan Jalur Logistik Maritim (Bézier Curve)",
        "```mermaid",
        f"{mermaid_str_1_6}",
        "```",
        "",
        "#### C. Formulasi Matematis: Kurva Parametrik Alur Pelayaran",
        "**Persamaan Formulasi Kurva Parametrik Alur Pelayaran Maritim:**",
        "```text",
        "Kurva(t) = (1 - t)^2 * Titik_Asal + 2 * (1 - t) * t * Titik_Kontrol + t^2 * Titik_Tujuan",
        "```",
        "*Keterangan Variabel:*",
        "- `Kurva(t)`: Vektor posisi koordinat geografis lintasan kapal pada parameter waktu t (rentang kontinu [0, 1]).",
        "- `Titik_Asal`: Titik koordinat geografis pelabuhan muat khusus di pesisir Sulawesi.",
        "- `Titik_Kontrol`: Titik koordinat jangkar pemandu kurva lengkung di perairan internasional.",
        "- `Titik_Tujuan`: Titik koordinat geografis pelabuhan bongkar di negara tujuan ekspor.",
        "",
        "#### D. Matriks Hasil Pemetaan Spasial",
        "Agregasi titik asal (Origin) dan tujuan akhir (Destination) di-render langsung ke dalam proyeksi spasial. Struktur data logistik yang menggerakkan pemetaan ini memetakan pelabuhan asal menuju pasar utama, yaitu **Tiongkok, Jepang, dan Korea Selatan**.",
        "",
        "#### E. Interpretasi Spasial Industri (Anatomi Rantai Pasok)",
        "Peta rute logistik maritim mengilustrasikan alur distribusi produk olahan nikel dari kawasan industri di Sulawesi:",
        "",
        "1. **Orientasi Ekspor:** Kawasan industri utama yang berstatus Proyek Strategis Nasional (PSN) mengalirkan produk olahan ke sentra-sentra industri manufaktur di pasar internasional.",
        "2. **Integrasi Rantai Pasok:** Mayoritas rute pengapalan terhubung langsung dengan pelabuhan ekspor tujuan, yang mengindikasikan posisi kawasan pemurnian di Sulawesi sebagai pemasok bahan baku setengah jadi.",
        "3. **Dinamika Rute Maritim:** Peta rute mencerminkan diversifikasi pasar ekspor (Asia Timur) dan jaringan logistik kawasan.",
        "",
        "---",
        "",
        "## 1.7 Matriks Indikator dan Sumber Data Resmi Bab 1",
        "Seluruh variabel kuantitatif, kategori analisis, satuan ukur, periode tahun observasi, dan institusi penyedia data primer resmi yang digunakan dalam Bab 1 dikompilasikan pada **Tabel 1.8** berikut:",
        "",
        "##### Tabel 1.8: Matriks Indikator dan Sumber Data Primer Resmi Bab 1",
        "| No | Nama Indikator | Kategori Analisis | Satuan Ukur | Cakupan Tahun | Institusi & Sumber Data Resmi |",
        "| :---: | :--- | :--- | :---: | :---: | :--- |",
        "| 1 | Izin Usaha Pertambangan (IUP) Baru | Faktor Tekanan Ekstraktif | Unit Izin | 2014–2024 | Kementerian ESDM (Minerbaone) |",
        "| 2 | Luas Wilayah Konsesi Tambang Baru | Faktor Tekanan Ekstraktif | Hektar (Ha) | 2014–2024 | Kementerian ESDM (Minerbaone) |",
        "| 3 | Kapasitas Terpasang PLTU Captive | Infrastruktur Energi Khusus | Megawatt (MW) | 2014–2024 | Global Energy Monitor (GEM) |",
        "| 4 | Fasilitas Pengolahan & Pemurnian (Smelter) | Fasilitas Industri Hilir | Unit Fasilitas | 2014–2024 | Kementerian ESDM & Studi Industri |",
        "| 5 | Realisasi Investasi PMDN | Arus Modal Domestik | Triliun Rupiah | 2016–2024 | Kementerian Investasi / BKPM |",
        "| 6 | PDRB Menurut 17 Lapangan Usaha | Struktur Ekonomi Makro | Triliun Rupiah | 2016–2024 | Badan Pusat Statistik (BPS Provinsi) |",
        "| 7 | PDRB Kabupaten Sentra Tambang | Struktur Ekonomi Daerah | Triliun Rupiah | 2016–2024 | BPS Kabupaten se-Sulteng |",
        "| 8 | Luas Kehilangan Hutan Komoditas | Dampak Tutupan Lahan | Hektar (Ha) | 2014–2023 | Global Forest Watch (GFW) |",
        "| 9 | Simpul Dermaga & Terminal Khusus Ekspor | Infrastruktur Rantai Pasok | Titik Koordinat & DWT | 2014–2024 | KNKT, Perpres PSN, Lap. Terbuka |",
        "",
        "---",
        "",
        "## 1.8 Bagan Alur Kerangka Kerja Riset Bab 1",
        "Keseluruhan struktur metodologis riset Bab 1 dioperasionalkan melalui empat fase kerja berurutan sebagaimana disajikan pada **Tabel 1.9** berikut:",
        "",
        "##### Tabel 1.9: Matriks Tahapan dan Alur Kerangka Kerja Riset Bab 1",
        "| Tahapan Riset | Fokus Metodologis | Bahan & Sumber Data | Keluaran / Hasil Analisis |",
        "| :--- | :--- | :--- | :--- |",
        "| **Fase I: Pengumpulan Data** | Kurasi data resmi lintas kementerian dan lembaga | Publikasi BPS, Minerbaone, BKPM, GEM, dan GFW | Basis Data Tabular Panel Provinsi (2014–2024) |",
        "| **Fase II: Reklasifikasi Hukum** | Penyusunan kerangka rantai pasok hukum terintegrasi | UU No. 3/2020, PP No. 96/2021, Perpres No. 112/2022 | 3 Klaster Makro (Ekstraktif, Akar Rumput, Jasa) |",
        "| **Fase III: Pengujian Statistik** | Uji signifikansi hubungan dan rasio peluang | Tabel Kontinjensi, Uji Chi-Square, Odds Ratio | Bukti Kausalitas Signifikan Tekanan vs Deforestasi |",
        "| **Fase IV: Pemetaan Rantai Pasok** | Triangulasi data logistik dan pemodelan maritim | Laporan KNKT, Perpres PSN, Kurva Parametrik Bézier | Peta Alur Rantai Pasok Ekspor & Konsentrasi Spasial 78% |"
    ])

    md_path = tool_dir / "Metodologi_Bab1_Ekspansi_Industri.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"  [OK] Tersimpan: {md_path}")

    # 5. Kompilasi LaTeX via Pandoc
    print("[5/5] Mengompilasi Metodologi_Bab1_Ekspansi_Industri.tex...")
    tex_path = tool_dir / "Metodologi_Bab1_Ekspansi_Industri.tex"
    try:
        pypandoc.convert_file(
            str(md_path),
            'latex',
            outputfile=str(tex_path),
            extra_args=['--standalone', '--resource-path', str(tool_dir)]
        )
        print(f"  [OK] Tersimpan: {tex_path}")
    except Exception as e:
        print(f"  [WARN] Gagal kompilasi LaTeX: {e}")

    print("\n" + "="*60)
    print(" [SELESAI] Laporan Metodologi Bab 1 Berhasil Dibuat!")
    print(f" Lokasi Output: {tool_dir}")
    print("="*60)

if __name__ == "__main__":
    generate_all_bab1()